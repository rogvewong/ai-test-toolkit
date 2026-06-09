"""FastAPI entry point for the AI Test Toolkit.

Provides multi-tenant REST endpoints for each step orchestrator. The API layer
is intentionally thin — it wires incoming requests to the same StepOrchestrator
classes that the CLI drives.

Run locally:
    uvicorn apps.api.main:app --reload --port 8080
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Iterable
from uuid import uuid4

try:
    from fastapi import Cookie, Depends, FastAPI, Header, HTTPException, Request
    from fastapi.responses import (
        HTMLResponse,
        PlainTextResponse,
        StreamingResponse,
        FileResponse,
        Response,
        RedirectResponse,
        JSONResponse,
    )
    from pydantic import BaseModel
except ImportError as _exc:  # pragma: no cover
    raise RuntimeError(
        "FastAPI not installed. Install the 'api' extra: pip install -e '.[api]'"
    ) from _exc

# stdlib html.escape for OAuth callback page text safety
from html import escape  # noqa: E402

from packages.core.config import settings
from packages.core.llm import LlmClient
from packages.core.memory import LayeredMemory, SqliteMemoryStore
from packages.reporting import aggregate, render_html, render_markdown
from packages.tdr import TdrWorkstation
from packages.workflow.base import StepContext, resolve_evidence_dir
from packages.workflow.step1_requirement import Step1Orchestrator
from packages.workflow.step2_testcase import Step2Orchestrator
from packages.workflow.step4_api import Step4Orchestrator
from packages.workflow.step5_ui import Step5Orchestrator
from packages.workflow.step6_agent import Step6Orchestrator
from packages.workflow.step6_agent.executor import ExecutionEnvironment
from packages.core.auth import user_store, UserRecord
from packages.core.auth.user_store import SESSION_COOKIE_NAME, SESSION_TTL_SECONDS

app = FastAPI(title="AI Test Toolkit", version="0.1.0")


@app.on_event("startup")
async def _seed_admin_from_env() -> None:
    """启动时检查 $AITK_ADMIN_USER + $AITK_ADMIN_PASSWORD 是否已设。
    若有 + 数据库里还没这位用户 → 自动建管理员账号(idempotent)。
    Docker 部署时这是给"零交互"建第一个 admin 的钩子。
    """
    import os as _os
    u = (_os.environ.get("AITK_ADMIN_USER") or "").strip()
    p = _os.environ.get("AITK_ADMIN_PASSWORD") or ""
    if not u or not p:
        return
    try:
        # 已 import 在文件头部:`from packages.core.auth import user_store`
        # user_store 这里就是 UserStore 单例
        if user_store.get_user_by_username(u):
            return  # 已经存在,不要重置密码
        user_store.create_user(
            username=u, password=p, role="admin",
            display_name=_os.environ.get("AITK_ADMIN_DISPLAY_NAME") or u,
        )
        print(f"[bootstrap] seeded admin user: {u}")
    except Exception as exc:
        print(f"[bootstrap] failed to seed admin: {exc}")


@app.on_event("startup")
async def _ensure_superadmin() -> None:
    """启动时确保存在超级管理员(idempotent)。
    默认把 zhangyafeng 升为 superadmin;可用 $AITK_SUPERADMIN_USER 覆盖。
    仅当该用户存在且尚非 superadmin 时升级;不创建账号、不改密码。
    """
    import os as _os
    name = (_os.environ.get("AITK_SUPERADMIN_USER") or "zhangyafeng").strip().lower()
    try:
        u = user_store.get_user_by_username(name)
        if u and not u.is_superadmin():
            user_store.set_role(u.id, "superadmin")
            print(f"[bootstrap] promoted superadmin: {name}")
    except Exception as exc:
        print(f"[bootstrap] ensure superadmin failed: {exc}")


@app.on_event("startup")
async def _start_oauth_refresh_loop() -> None:
    """后台周期性刷新 OAuth token — 每 20 分钟查一次,过期前 5 分钟自动续。
    这样即使没人跑工具、或一个 run 跑很久,token 也不会过期。
    """
    import asyncio as _a

    async def _loop() -> None:
        # 启动后先等几秒让 app 完全起来,然后立即刷一次
        await _a.sleep(5)
        while True:
            try:
                r = await ensure_fresh_oauth_token()
                if r.get("status") in ("refreshed", "error", "no_refresh_token"):
                    print(f"[oauth-loop] {r.get('status')}: {r.get('detail','')}")
            except Exception as exc:
                print(f"[oauth-loop] 异常: {exc}")
            await _a.sleep(10 * 60)  # 10 分钟查一次

    _asyncio.create_task(_loop())


# ---------------------------------------------------------------------
# Web 用户账号体系 (与 Claude 凭据完全独立)
# ---------------------------------------------------------------------
# 设计:
# - SESSION_COOKIE_NAME = aitk_session,HttpOnly cookie,30 天 TTL
# - 凡是 /tools /reports /settings 及大多数 /api/* 必须 已登录
# - 白名单(免登录):/login /register /api/auth/* /healthz /static /api/healthz
# - 注册的第一个用户自动变 admin,后续注册都是普通 user
# - admin 可以改 Claude 凭据 + 看所有人的报告;user 只看自己

_AUTH_EXEMPT_PATHS: tuple[str, ...] = (
    "/healthz",
    "/login",
    "/register",
    "/favicon.ico",
)
_AUTH_EXEMPT_PREFIXES: tuple[str, ...] = (
    "/api/auth/",
    "/api/healthz",
    "/static/",
    "/api/static/",
)


def _is_auth_exempt(path: str) -> bool:
    if path in _AUTH_EXEMPT_PATHS:
        return True
    return any(path.startswith(p) for p in _AUTH_EXEMPT_PREFIXES)


@app.middleware("http")
async def session_guard(request: Request, call_next):
    """全局 session 闸:未登录 + 访问非白名单路径 → 网页 302 到 /login,
    API 返回 401。已登录:把 user 注到 request.state.user。
    """
    path = request.url.path
    token = request.cookies.get(SESSION_COOKIE_NAME, "")
    user = user_store.resolve_session(token) if token else None
    request.state.current_user = user
    if user is None and not _is_auth_exempt(path):
        # 第一次部署 — 还没有任何账号,直接放行 /register 引导建第一个 admin
        if user_store.count_users() == 0:
            if path == "/":
                return RedirectResponse("/register", status_code=302)
            if path.startswith("/api/"):
                # API 也允许穿透 — 让前端 hint
                response = await call_next(request)
                return response
        # 已经有账号但当前未登录 → 网页登录,API 401
        if path.startswith("/api/"):
            return JSONResponse({"error": "unauthorized", "detail": "请先登录"}, status_code=401)
        return RedirectResponse(f"/login?next={path}", status_code=302)
    response = await call_next(request)
    return response


def require_user(request: Request) -> UserRecord:
    user = getattr(request.state, "current_user", None)
    if not user:
        raise HTTPException(401, "unauthorized")
    return user


def require_admin(request: Request) -> UserRecord:
    user = require_user(request)
    if not user.is_admin():
        raise HTTPException(403, "需要管理员权限")
    return user


def require_superadmin(request: Request) -> UserRecord:
    user = require_user(request)
    if not user.is_superadmin():
        raise HTTPException(403, "需要超级管理员权限")
    return user


def _set_session_cookie(resp: Response, token: str) -> None:
    # secure=False 是因为先 HTTP 跑;上线 HTTPS 时改 True
    resp.set_cookie(
        SESSION_COOKIE_NAME,
        token,
        max_age=SESSION_TTL_SECONDS,
        httponly=True,
        samesite="lax",
        path="/",
    )


def _clear_session_cookie(resp: Response) -> None:
    resp.delete_cookie(SESSION_COOKIE_NAME, path="/")


# ---------- auth endpoints ----------

class RegisterReq(BaseModel):
    username: str
    password: str
    display_name: str | None = None


class LoginReq(BaseModel):
    username: str
    password: str


@app.post("/api/auth/register")
async def api_auth_register(req: RegisterReq, request: Request) -> JSONResponse:
    """注册端点 — 仅在系统还没用户时开放(首次安装建 admin)。
    一旦有 admin,公开注册自动失效;后续新用户必须由 admin 在管理后台创建。
    """
    if user_store.count_users() > 0:
        raise HTTPException(
            403,
            "公开注册已关闭。新账号需由管理员在「用户管理」创建。",
        )
    try:
        user = user_store.create_user(
            username=req.username,
            password=req.password,
            display_name=(req.display_name or "").strip(),
        )
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    ua = request.headers.get("user-agent", "")[:240]
    sess = user_store.create_session(user.id, user_agent=ua)
    resp = JSONResponse({"ok": True, "user": user.to_dict(), "is_first_admin": user.is_admin()})
    _set_session_cookie(resp, sess.token)
    return resp


@app.post("/api/auth/login")
async def api_auth_login(req: LoginReq, request: Request) -> JSONResponse:
    user = user_store.verify_password(req.username, req.password)
    if not user:
        raise HTTPException(401, "用户名或密码错误")
    ua = request.headers.get("user-agent", "")[:240]
    sess = user_store.create_session(user.id, user_agent=ua)
    resp = JSONResponse({"ok": True, "user": user.to_dict()})
    _set_session_cookie(resp, sess.token)
    return resp


@app.post("/api/auth/logout")
async def api_auth_logout(request: Request) -> JSONResponse:
    token = request.cookies.get(SESSION_COOKIE_NAME, "")
    if token:
        user_store.delete_session(token)
    resp = JSONResponse({"ok": True})
    _clear_session_cookie(resp)
    return resp


@app.get("/api/auth/me")
async def api_auth_me(request: Request) -> dict[str, Any]:
    user = getattr(request.state, "current_user", None)
    if not user:
        return {"authenticated": False, "first_setup_needed": user_store.count_users() == 0}
    return {"authenticated": True, "user": user.to_dict()}


class ChangePasswordReq(BaseModel):
    old_password: str
    new_password: str


@app.post("/api/auth/change_password")
async def api_auth_change_password(req: ChangePasswordReq, request: Request) -> dict[str, Any]:
    """用户主动改密 — 验旧密码 + 自主选择是否改 (即非强制)。"""
    user = require_user(request)
    verified = user_store.verify_password(user.username, req.old_password)
    if not verified:
        raise HTTPException(400, "原密码不正确")
    try:
        user_store.change_password(user.id, req.new_password)
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    return {"ok": True, "msg": "密码已更新,所有现有 session 已撤销,请重新登录"}


# ---------- admin-only user management ----------

class AdminCreateUserReq(BaseModel):
    username: str
    password: str
    display_name: str | None = None
    role: str | None = "user"  # 'user' 或 'admin'


class AdminResetPasswordReq(BaseModel):
    new_password: str


@app.get("/api/auth/users")
async def api_auth_list_users(request: Request) -> dict[str, Any]:
    """管理员看用户列表。"""
    require_admin(request)
    return {"users": [u.to_dict() for u in user_store.list_users()]}


@app.post("/api/auth/users")
async def api_auth_admin_create_user(req: AdminCreateUserReq, request: Request) -> dict[str, Any]:
    """管理员创建新用户。
    Body: {username, password, display_name?, role?}
    """
    actor = require_admin(request)
    role = req.role if req.role in ("user", "admin", "superadmin") else "user"
    # 非超管只能创建普通用户
    if role != "user" and not actor.is_superadmin():
        raise HTTPException(403, "只有超级管理员能创建管理员 / 超级管理员")
    try:
        user = user_store.create_user(
            username=req.username,
            password=req.password,
            display_name=(req.display_name or "").strip(),
            role=role,
        )
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    return {"ok": True, "user": user.to_dict()}


# ---------- 批量创建 ----------

def _gen_random_password(length: int = 10) -> str:
    """生成强但可读的随机密码:字母+数字 (避开易混淆的 0/O/l/1)。"""
    import secrets as _s
    alphabet = "abcdefghjkmnpqrstuvwxyzABCDEFGHJKLMNPQRSTUVWXYZ23456789"
    return "".join(_s.choice(alphabet) for _ in range(length))


class BulkUserItem(BaseModel):
    username: str
    password: str | None = None  # 留空则自动生成
    display_name: str | None = None
    role: str | None = None


class BulkCreateUsersReq(BaseModel):
    users: list[BulkUserItem]
    default_role: str = "user"
    default_password_length: int = 10  # auto-gen 时长度


@app.post("/api/auth/users/bulk")
async def api_auth_admin_bulk_create_users(
    req: BulkCreateUsersReq, request: Request,
) -> dict[str, Any]:
    """批量创建用户。
    - 每个 item 的 password 留空 → 用 _gen_random_password 生成
    - 失败的不会阻塞剩下的;返回 {created:[...], failed:[...]} 让前端展示
    - created 列表里会把刚才用的密码原文带回去(后端不会留明文,只是这一次响应有),
      admin 复制下来发给员工即可
    """
    actor = require_admin(request)
    default_role = req.default_role if req.default_role in ("user", "admin", "superadmin") else "user"
    if not actor.is_superadmin():
        default_role = "user"  # 非超管批量只能建普通用户
    created: list[dict[str, Any]] = []
    failed: list[dict[str, Any]] = []
    seen_in_batch: set[str] = set()
    for raw in req.users:
        uname = (raw.username or "").strip().lower()
        if not uname:
            failed.append({"username": raw.username, "error": "用户名为空"})
            continue
        if uname in seen_in_batch:
            failed.append({"username": uname, "error": "本批中重复"})
            continue
        seen_in_batch.add(uname)
        pwd = raw.password
        was_auto = False
        if not pwd:
            pwd = _gen_random_password(max(8, req.default_password_length))
            was_auto = True
        role = raw.role if raw.role in ("user", "admin", "superadmin") else default_role
        if not actor.is_superadmin():
            role = "user"
        try:
            u = user_store.create_user(
                username=uname,
                password=pwd,
                display_name=(raw.display_name or "").strip(),
                role=role,
            )
        except ValueError as exc:
            failed.append({"username": uname, "error": str(exc)})
            continue
        created.append({
            **u.to_dict(),
            "password": pwd,             # 仅这一次响应里返回原文密码
            "password_auto_generated": was_auto,
        })
    return {
        "ok": True,
        "summary": {"total": len(req.users), "created": len(created), "failed": len(failed)},
        "created": created,
        "failed": failed,
    }


@app.post("/api/auth/users/{user_id}/reset_password")
async def api_auth_admin_reset_password(
    user_id: int, req: AdminResetPasswordReq, request: Request,
) -> dict[str, Any]:
    """管理员给某用户重置密码(不验旧密码)。"""
    actor = require_admin(request)
    target = user_store.get_user(user_id)
    if not target:
        raise HTTPException(404, "用户不存在")
    # 非超管只能重置普通用户的密码
    if not actor.is_superadmin() and target.role != "user":
        raise HTTPException(403, "只有超级管理员能重置管理员 / 超级管理员的密码")
    try:
        user_store.admin_reset_password(user_id, req.new_password)
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    return {"ok": True, "msg": f"已重置 {target.username} 的密码;该用户所有现有 session 已撤销"}


@app.delete("/api/auth/users/{user_id}")
async def api_auth_admin_delete_user(user_id: int, request: Request) -> dict[str, Any]:
    """管理员删除用户。禁止自删,禁止删最后一个 admin。"""
    actor = require_admin(request)
    if user_id == actor.id:
        raise HTTPException(400, "不能删除自己")
    target = user_store.get_user(user_id)
    if not target:
        raise HTTPException(404, "用户不存在")
    # 非超管只能删普通用户
    if not actor.is_superadmin() and target.role != "user":
        raise HTTPException(403, "只有超级管理员能删除管理员 / 超级管理员")
    # 不能删最后一个超级管理员
    if target.is_superadmin() and user_store.count_by_role("superadmin") <= 1:
        raise HTTPException(400, "不能删除最后一个超级管理员")
    user_store.delete_user(user_id)
    return {"ok": True, "msg": f"已删除用户 {target.username}"}


class SetRoleReq(BaseModel):
    role: str


@app.post("/api/auth/users/{user_id}/role")
async def api_auth_set_role(user_id: int, req: SetRoleReq, request: Request) -> dict[str, Any]:
    """超级管理员修改用户角色(user / admin / superadmin)。改后该用户需重新登录。"""
    require_superadmin(request)
    if req.role not in ("user", "admin", "superadmin"):
        raise HTTPException(400, "非法角色")
    target = user_store.get_user(user_id)
    if not target:
        raise HTTPException(404, "用户不存在")
    # 不能把最后一个超级管理员降级
    if target.is_superadmin() and req.role != "superadmin" and user_store.count_by_role("superadmin") <= 1:
        raise HTTPException(400, "不能降级最后一个超级管理员")
    user_store.set_role(user_id, req.role)
    return {"ok": True, "msg": f"已将 {target.username} 设为 {req.role};该用户需重新登录生效"}


# ---------------------------------------------------------------------
# Tenancy
# ---------------------------------------------------------------------

class Tenant(BaseModel):
    tenant_id: str
    project_id: str


async def get_tenant(
    x_tenant_id: str | None = Header(default=None),
    x_project_id: str | None = Header(default=None),
) -> Tenant:
    if not x_tenant_id or not x_project_id:
        raise HTTPException(status_code=400, detail="X-Tenant-Id / X-Project-Id required")
    return Tenant(tenant_id=x_tenant_id, project_id=x_project_id)


async def make_context(tenant: Tenant, inputs: dict[str, Any]) -> StepContext:
    run_id = str(uuid4())
    db_path = str((settings.report_output_dir / "memory.db").resolve())
    settings.report_output_dir.mkdir(parents=True, exist_ok=True)
    store = SqliteMemoryStore(db_path)
    memory = LayeredMemory(
        store=store,
        run_id=run_id,
        project_id=tenant.project_id,
        tenant_id=tenant.tenant_id,
    )
    return StepContext(
        run_id=run_id,
        project_id=tenant.project_id,
        tenant_id=tenant.tenant_id,
        inputs=inputs,
        memory=memory,
        llm=LlmClient(),
        evidence_dir=resolve_evidence_dir(),
    )


# ---------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------

class Step1Request(BaseModel):
    prd: str | dict[str, Any] | None = None
    prototype: str | None = None
    ui_design: str | None = None
    flow_chart: str | None = None
    api_doc: str | None = None
    business_rules: str | None = None


class Step2Request(BaseModel):
    requirement_report: dict[str, Any]


class Step4Request(BaseModel):
    requirement_report: dict[str, Any]
    test_case_report: dict[str, Any]
    api_doc: str | dict[str, Any] | None = None
    execution_results: dict[str, Any] | None = None


class Step5Request(BaseModel):
    requirement_report: dict[str, Any] | None = None
    test_case_report: dict[str, Any] | None = None
    design_assets: list[dict[str, Any]] = []
    actual_snapshots: list[dict[str, Any]] = []
    state_captures: list[dict[str, Any]] = []


class Step6Request(BaseModel):
    test_case_report: dict[str, Any]
    automation_risk: dict[str, Any] = {}
    env_info: dict[str, Any] = {}
    dry_run: bool = True


class TdrRequest(BaseModel):
    run_id: str
    artifacts: list[str]
    reviewer_scores: list[dict[str, dict[str, float]]]
    comments: list[dict[str, Any]] = []


# ---------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------

@app.get("/healthz")
async def healthz() -> dict[str, str]:
    return {"status": "ok", "version": "0.1.0"}


@app.post("/runs/step1")
async def run_step1(req: Step1Request, tenant: Tenant = Depends(get_tenant)) -> dict[str, Any]:
    ctx = await make_context(tenant, req.model_dump())
    report = await Step1Orchestrator(ctx).execute()
    return report.model_dump(mode="json")


@app.post("/runs/step2")
async def run_step2(req: Step2Request, tenant: Tenant = Depends(get_tenant)) -> dict[str, Any]:
    ctx = await make_context(tenant, req.model_dump())
    report = await Step2Orchestrator(ctx).execute()
    return report.model_dump(mode="json")


@app.post("/runs/step4")
async def run_step4(req: Step4Request, tenant: Tenant = Depends(get_tenant)) -> dict[str, Any]:
    ctx = await make_context(tenant, req.model_dump())
    report = await Step4Orchestrator(ctx).execute()
    return report.model_dump(mode="json")


@app.post("/runs/step5")
async def run_step5(req: Step5Request, tenant: Tenant = Depends(get_tenant)) -> dict[str, Any]:
    ctx = await make_context(tenant, req.model_dump())
    report = await Step5Orchestrator(ctx).execute()
    return report.model_dump(mode="json")


@app.post("/runs/step6")
async def run_step6(req: Step6Request, tenant: Tenant = Depends(get_tenant)) -> dict[str, Any]:
    inputs = req.model_dump()
    inputs["execution_env"] = ExecutionEnvironment(
        evidence_dir=Path(settings.evidence_output_dir) / "step6",
        dry_run=req.dry_run,
    )
    ctx = await make_context(tenant, inputs)
    report = await Step6Orchestrator(ctx).execute()
    return report.model_dump(mode="json")


def _normalize_tdr_comment(c: dict[str, Any]) -> dict[str, Any] | None:
    """把 UI / API 提交的 TDR 评论字段，规范化为 add_comment 接受的完整 7 字段。

    宽容映射：
      - text / comment → observation
      - 缺 id          → 自动生成
      - 缺 dimension   → "general"
      - 缺 severity    → "info"
      - 缺其他字段     → 空字符串
    完全空的评论（无 text/observation）返回 None，调用方应跳过。

    被 /tdr 和 /api/tdr/submit 共用，保证两个入口行为一致。
    """
    if not isinstance(c, dict):
        return None
    body_text = c.get("observation") or c.get("text") or c.get("comment") or ""
    if not body_text:
        return None
    return {
        "id":                c.get("id") or f"cmt-{uuid4().hex[:8]}",
        "dimension":         c.get("dimension") or "general",
        "severity":          c.get("severity") or "info",
        "location":          c.get("location") or "",
        "observation":       body_text,
        "suggestion":        c.get("suggestion") or "",
        "requires_response": bool(c.get("requires_response", False)),
    }


@app.post("/tdr")
async def submit_tdr(req: TdrRequest, tenant: Tenant = Depends(get_tenant)) -> dict[str, Any]:
    ws = TdrWorkstation(
        run_id=req.run_id, project_id=tenant.project_id, tenant_id=tenant.tenant_id
    )
    ws.submit_artifacts(req.artifacts)
    ws.score(req.reviewer_scores)
    for raw in req.comments:
        # 兼容两种格式：完整字段 / UI 简短字段（{severity, text}）
        normalized = _normalize_tdr_comment(raw)
        if normalized is None:
            continue
        try:
            ws.add_comment(**normalized)
        except TypeError as e:
            raise HTTPException(422, f"评论字段不被工作站接受：{e}")
    return ws.finalize()


# ---------------------------------------------------------------------
# Standards + TDR toolkit — dashboard-facing JSON endpoints
# ---------------------------------------------------------------------


@app.get("/api/standards")
async def standards_all() -> dict[str, Any]:
    """Return all five governance standards (process / quality / data / tdr / team)."""
    from packages.core.config import standards as S
    return {
        "process": S.process,
        "quality": S.quality,
        "data": S.data,
        "tdr": S.tdr,
        "team": S.team,
    }


@app.get("/api/standards/{name}")
async def standards_one(name: str) -> dict[str, Any]:
    from packages.core.config import standards as S
    mapping = {
        "process": S.process,
        "quality": S.quality,
        "data": S.data,
        "tdr": S.tdr,
        "team": S.team,
    }
    if name not in mapping:
        raise HTTPException(status_code=404, detail=f"unknown standard: {name}")
    return mapping[name]


@app.get("/api/tdr/reviews")
async def tdr_list() -> dict[str, Any]:
    """Enumerate all TDR reviews on disk + summary for dashboard listing.

    扫描两种位置（兼容旧/新落盘）：
      - 顶层文件: <report_dir>/tdr_<run_id>.json （/api/tdr/submit mirror 出来的）
      - 嵌套目录: <report_dir>/tdr/<run_id>/review.json （TdrWorkstation.finalize 写入）
    """
    import json as _json
    base = settings.report_output_dir
    if not base.exists():
        return {"reviews": []}
    reviews: list[dict[str, Any]] = []
    seen_run_ids: set[str] = set()

    # 路径 A: 顶层 tdr_*.json
    for f in sorted(base.glob("tdr_*.json")):
        try:
            data = _json.loads(f.read_text(encoding="utf-8"))
        except Exception:
            continue
        rid = data.get("meta", {}).get("run_id") or f.stem.removeprefix("tdr_")
        if rid in seen_run_ids:
            continue
        seen_run_ids.add(rid)
        reviews.append(_summarize_tdr_review(data, rid, f.name))

    # 路径 B: tdr/<run_id>/review.json（TdrWorkstation 默认写入位置）
    tdr_dir = base / "tdr"
    if tdr_dir.exists():
        for review_file in sorted(tdr_dir.glob("*/review.json")):
            try:
                data = _json.loads(review_file.read_text(encoding="utf-8"))
            except Exception:
                continue
            rid = data.get("meta", {}).get("run_id") or review_file.parent.name
            if rid in seen_run_ids:
                continue  # 已被 A 覆盖
            seen_run_ids.add(rid)
            reviews.append(_summarize_tdr_review(data, rid, f"tdr/{review_file.parent.name}/review.json"))

    return {"reviews": reviews}


def _iter_saved_report_files() -> Iterable[Path]:
    """枚举所有已保存的报告文件（顶层 *.json + 嵌套 tdr/<run_id>/review.json）。

    被 /api/reports、/api/reports/export 等用来扫盘。统一这里让两种位置兼容。
    """
    out_dir = Path(settings.report_output_dir)
    if not out_dir.exists():
        return
    yield from out_dir.glob("*.json")
    nested_tdr = out_dir / "tdr"
    if nested_tdr.exists():
        yield from nested_tdr.glob("*/review.json")


def _find_report_files_by_run_id(run_id: str) -> list[Path]:
    """按 run_id 找所有匹配的报告文件（含嵌套 TDR）。

    返回多个 Path 是因为：
      - 工具报告：<report_dir>/<tool_id>_<run_id>.json
      - TDR mirror: <report_dir>/tdr_<run_id>.json
      - TDR 嵌套:   <report_dir>/tdr/<run_id>/review.json
    /api/reports/{run_id} 取第一个，DELETE 全部删。
    """
    out_dir = Path(settings.report_output_dir)
    if not out_dir.exists():
        return []
    found: list[Path] = list(out_dir.glob(f"*_{run_id}.json"))
    nested = out_dir / "tdr" / run_id / "review.json"
    if nested.exists():
        found.append(nested)
    return found


def _parse_tool_id_from_stem(stem: str, run_id: str | None = None) -> str:
    """从报告 JSON 的 stem 解出 tool_id。

    文件名形如 `<tool_id>_<run_id>.json`，但 tool_id 自身可能含下划线
    （network_resilience / h5_adapt / seo_audit）。简单的 `split("_", 1)[0]`
    会把 network_resilience 截成 network — 导致导出 HTML 标题、`/api/reports/{run_id}`
    返回的 tool_id 都错。

    解析顺序：
      1. 若知道 run_id 且 stem 以 `_<run_id>` 结尾，直接剪掉。
      2. 否则按 TOOL_CATALOG 的 id（含 'tdr'）做最长前缀匹配。
      3. 最后兜底 split("_", 1)[0]，至少不抛异常。
    """
    if run_id and stem.endswith("_" + run_id):
        return stem[: -(len(run_id) + 1)]
    valid_ids = sorted(
        ({t["id"] for t in TOOL_CATALOG} | {"tdr"}),
        key=lambda s: -len(s),
    )
    for tid in valid_ids:
        if stem == tid or stem.startswith(tid + "_"):
            return tid
    return stem.split("_", 1)[0] if "_" in stem else stem


def _summarize_tdr_review(data: dict[str, Any], run_id: str, file_repr: str) -> dict[str, Any]:
    """共用：把 TDR 报告 JSON 收成 dashboard 卡片。"""
    meta = data.get("meta") or {}
    return {
        "run_id": run_id,
        "file": file_repr,
        "decision": data.get("decision"),
        "overall_score": data.get("overall_score"),
        "dimension_scores": data.get("dimension_scores") or {},
        "comments_n": len(data.get("comments") or []),
        "comments_blocker": sum(
            1 for c in (data.get("comments") or [])
            if c.get("severity") in {"blocker", "major"} and not c.get("resolved")
        ),
        "signatures_n": len(data.get("signatures") or []),
        "artifacts_n": len(data.get("artifacts_reviewed") or []),
        "follow_up_n": len(data.get("follow_up_items") or []),
        "created_at": meta.get("created_at") or meta.get("timestamp"),
        "project_id": meta.get("project_id"),
        "tenant_id": meta.get("tenant_id"),
    }


@app.get("/api/tdr/reviews/{run_id}")
async def tdr_get(run_id: str) -> dict[str, Any]:
    """读 TDR 详情 — 同时支持顶层 mirror 和嵌套 review.json 两种存储。"""
    import json as _json
    base = settings.report_output_dir
    # 顶层 mirror（/api/tdr/submit 写出来的）
    flat = base / f"tdr_{run_id}.json"
    if flat.exists():
        return _json.loads(flat.read_text(encoding="utf-8"))
    # 嵌套（TdrWorkstation.finalize 默认位置）
    nested = base / "tdr" / run_id / "review.json"
    if nested.exists():
        return _json.loads(nested.read_text(encoding="utf-8"))
    raise HTTPException(status_code=404, detail=f"no TDR review for {run_id}")


class TdrWorkbenchRequest(BaseModel):
    """Header-less TDR submission for the browser toolkit.

    Tenant / project are optional; fall back to 'default' so the UI has zero
    friction. The TDR spec governance is enforced by the workstation itself
    (dimensions, decision rules, signing).
    """
    run_id: str | None = None
    project_id: str = "default"
    tenant_id: str = "default"
    artifacts: list[str] = []
    reviewer_scores: list[dict[str, dict[str, float]]] = []
    comments: list[dict[str, Any]] = []


@app.post("/api/tdr/submit")
async def tdr_submit_ui(req: TdrWorkbenchRequest) -> dict[str, Any]:
    """Run a TDR review end-to-end from UI inputs, persist to disk, return the full report."""
    import json as _json
    ws = TdrWorkstation(
        run_id=req.run_id or f"ui-{uuid4().hex[:8]}",
        project_id=req.project_id,
        tenant_id=req.tenant_id,
    )
    ws.submit_artifacts(req.artifacts)
    ws.score(req.reviewer_scores)
    for raw in req.comments:
        normalized = _normalize_tdr_comment(raw)
        if normalized is None:
            continue
        try:
            ws.add_comment(**normalized)
        except TypeError as e:
            raise HTTPException(422, f"评论字段不被工作站接受：{e}")
    result = ws.finalize()
    # Mirror the CLI: persist next to the other reports so the listing endpoint picks it up
    path = settings.report_output_dir / f"tdr_{ws.run_id}.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(_json.dumps(result, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    return {"path": str(path), "report": result}


def _load_step_reports(run_id: str) -> dict[str, dict[str, Any]]:
    """Best-effort load of stepN_<run_id>.json files from the reports dir.

    Returns a dict keyed by "step1"…"step6" so aggregate() can render a real
    bundle instead of an empty shell.
    """
    import json as _json
    out: dict[str, dict[str, Any]] = {}
    base = settings.report_output_dir
    if not base.exists():
        return out
    for n in ("1", "2", "4", "5", "6"):
        p = base / f"step{n}_{run_id}.json"
        if p.exists():
            try:
                out[f"step{n}"] = _json.loads(p.read_text(encoding="utf-8"))
            except Exception:
                pass
    return out


def _load_tdr(run_id: str) -> dict[str, Any] | None:
    import json as _json
    p = settings.report_output_dir / f"tdr_{run_id}.json"
    if p.exists():
        try:
            return _json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return None
    return None


def _infer_tenant_from_reports(run_id: str) -> tuple[str, str]:
    """Read tenant_id / project_id out of the first stepN_<run_id>.json on disk.

    The report-viewing endpoints are read-only against files the CLI/API already
    produced — requiring the caller to also know the headers is friction with
    no security benefit. We pull them straight from report.meta.
    """
    reports = _load_step_reports(run_id)
    for _, rpt in reports.items():
        meta = rpt.get("meta") or {}
        t = meta.get("tenant_id")
        p = meta.get("project_id")
        if t and p:
            return str(t), str(p)
    return "default", "unknown"


@app.get("/api/runs")
async def runs_json() -> dict[str, Any]:
    """Dashboard data source — enumerate runs, pull gate/confidence/key metrics per step."""
    import re
    base = settings.report_output_dir
    if not base.exists():
        return {"runs": []}
    grouped: dict[str, dict[str, dict[str, Any]]] = {}
    pattern = re.compile(r"^step(\d)_([0-9a-f-]+)\.json$")
    for f in sorted(base.glob("step*_*.json")):
        m = pattern.match(f.name)
        if not m:
            continue
        step, rid = m.group(1), m.group(2)
        try:
            import json as _json
            data = _json.loads(f.read_text(encoding="utf-8"))
        except Exception:
            continue
        meta = data.get("meta") or {}
        gd = data.get("gate_decision") or {}
        conf = data.get("confidence") or {}
        metrics = data.get("metrics") or {}
        entry: dict[str, Any] = {
            "gate_action": gd.get("action"),
            "gate_reasons": gd.get("reasons") or [],
            "gate_blockers": gd.get("blockers") or [],
            "next_step": gd.get("next_step"),
            "confidence": conf.get("score"),
            "confidence_grade": conf.get("grade"),
            "project_id": meta.get("project_id"),
            "tenant_id": meta.get("tenant_id"),
            "model": meta.get("model_id"),
        }
        # Per-step key counts
        if step == "1":
            entry["materials_n"] = len(data.get("materials") or [])
            entry["modules_n"] = len(data.get("modules") or [])
            entry["flows_n"] = len(data.get("flows") or [])
            entry["ambiguities_n"] = len(data.get("ambiguities") or [])
        elif step == "2":
            entry["p0_n"] = len(data.get("p0_cases") or [])
            entry["p1_n"] = len(data.get("p1_cases") or [])
            entry["p2_n"] = len(data.get("p2_cases") or [])
            entry["automation"] = data.get("automation_summary") or {}
        elif step == "4":
            entry["apis_n"] = len(data.get("api_list") or [])
            entry["functional_n"] = len(data.get("functional_results") or [])
            entry["boundary_n"] = len(data.get("boundary_results") or [])
            entry["security_n"] = len(data.get("security_results") or [])
            entry["defects_n"] = len(data.get("defects") or [])
            entry["pass_rate"] = metrics.get("pass_rate")
        elif step == "5":
            entry["references_n"] = len(data.get("references") or [])
            entry["diffs_n"] = len(data.get("diffs") or [])
            entry["severity_map"] = data.get("severity_map") or {}
            entry["deviation"] = data.get("overall_deviation_ratio")
        elif step == "6":
            entry["steps_n"] = len(data.get("steps") or [])
            entry["failures_n"] = len(data.get("failures") or [])
            entry["success_rate"] = metrics.get("success_rate")
        grouped.setdefault(rid, {})[f"step{step}"] = entry
    runs = []
    for rid, steps in grouped.items():
        any_meta = next(iter(steps.values()), {})
        runs.append({
            "run_id": rid,
            "project_id": any_meta.get("project_id"),
            "tenant_id": any_meta.get("tenant_id"),
            "steps": steps,
        })
    runs.sort(key=lambda r: r["run_id"])
    return {"runs": runs}


@app.get("/api/runs/{run_id}/raw/{step}")
async def run_step_raw(run_id: str, step: str) -> dict[str, Any]:
    """Return the raw step report JSON for dashboard drill-down."""
    if step not in {"1", "2", "4", "5", "6"}:
        raise HTTPException(status_code=400, detail="step must be 1/2/4/5/6")
    p = settings.report_output_dir / f"step{step}_{run_id}.json"
    if not p.exists():
        raise HTTPException(status_code=404, detail=f"no step{step} for run_id={run_id}")
    import json as _json
    return _json.loads(p.read_text(encoding="utf-8"))








@app.get("/")
async def root(request: Request):
    """根路径默认带去 /tools (已登录) 或 /login (未登录)。
    middleware 已经把未登录的拦走;这里走到说明已登录。
    """
    return RedirectResponse("/tools", status_code=302)


# =====================================================================
# Prompts — 提示词集成进工具详情页
# =====================================================================

_STEP_LABELS = [
    ("step1_requirement", "Step 1 需求拆解"),
    ("step2_testcase", "Step 2 测试用例设计"),
    ("step4_api", "Step 4 接口测试"),
    ("step5_ui", "Step 5 UI 一致性比对"),
    ("step6_agent", "Step 6 Agent 自动化执行"),
]


def _all_prompt_steps() -> list[dict[str, Any]]:
    """Load every step's prompts and return JSON-serializable summaries."""
    from packages.core.prompts import load_step

    out: list[dict[str, Any]] = []
    for dir_name, label in _STEP_LABELS:
        try:
            step = load_step(dir_name)
        except Exception as exc:
            out.append({"step_id": dir_name, "name": label, "error": str(exc)})
            continue
        substeps = []
        for fname in step.order:
            tpl = next((t for t in step.templates.values() if t.path.name == fname), None)
            if tpl is None:
                continue
            substeps.append({
                "id": tpl.id,
                "name": tpl.name,
                "version": tpl.version,
                "model_tier": tpl.model_tier.value,
                "temperature": tpl.temperature,
                "max_tokens": tpl.max_tokens,
                "placeholders": tpl.placeholders,
                "output_format": tpl.output_format,
                "output_schema": tpl.output_schema,
                "body": tpl.body,
                "source_path": str(tpl.path),
            })
        out.append({
            "step_id": step.step_id,
            "dir_name": dir_name,
            "display_name": label,
            "name": step.name,
            "version": step.version,
            "common_system_suffix": step.common_system_suffix,
            "substeps": substeps,
        })
    return out


@app.get("/api/prompts")
async def api_prompts_list() -> dict[str, Any]:
    """列出全部 25 个 SOP 提示词（按 5 个 AI 步骤分组）。"""
    steps = _all_prompt_steps()
    total = sum(len(s.get("substeps", [])) for s in steps)
    return {"total": total, "steps": steps}


@app.get("/api/prompts/{step_dir}/{sub_id}")
async def api_prompt_show(step_dir: str, sub_id: str) -> dict[str, Any]:
    """获取单个提示词的完整正文与 frontmatter。"""
    from packages.core.prompts import load_step
    from packages.core.prompts.loader import OVERRIDES_DIR, PROMPTS_DIR

    try:
        step = load_step(step_dir)
    except Exception as exc:
        raise HTTPException(404, f"step not found: {exc}")
    try:
        tpl = step.get(sub_id)
    except KeyError as exc:
        raise HTTPException(404, str(exc))
    is_override = str(tpl.path).startswith(str(OVERRIDES_DIR))
    return {
        "id": tpl.id,
        "name": tpl.name,
        "version": tpl.version,
        "model_tier": tpl.model_tier.value,
        "temperature": tpl.temperature,
        "max_tokens": tpl.max_tokens,
        "placeholders": tpl.placeholders,
        "output_format": tpl.output_format,
        "output_schema": tpl.output_schema,
        "body": tpl.body,
        "common_system_suffix": step.common_system_suffix,
        "source_path": str(tpl.path),
        "is_override": is_override,
    }


class PromptUpdate(BaseModel):
    body: str  # body-only or full markdown with frontmatter


@app.put("/api/prompts/{step_dir}/{sub_id}")
async def api_prompt_update(step_dir: str, sub_id: str, payload: PromptUpdate) -> dict[str, Any]:
    """编辑提示词正文 — 写入 configs/prompts_overrides/，下次运行立即生效。"""
    from packages.core.prompts import load_step, write_prompt_override

    try:
        step = load_step(step_dir)
        tpl = step.get(sub_id)
    except Exception as exc:
        raise HTTPException(404, str(exc))
    fname = Path(tpl.path).name
    target = write_prompt_override(step_dir, fname, payload.body)
    return {"saved": str(target), "is_override": True}


@app.delete("/api/prompts/{step_dir}/{sub_id}")
async def api_prompt_reset(step_dir: str, sub_id: str) -> dict[str, Any]:
    """重置覆盖：删除 overrides 文件，下次运行回退到 configs/prompts/ 原版。"""
    from packages.core.prompts import load_step, reset_prompt_override
    from packages.core.prompts.loader import PROMPTS_DIR

    # We need the original filename from configs/prompts/ (override may have been deleted already)
    try:
        step = load_step(step_dir)
        tpl = step.get(sub_id)
    except Exception as exc:
        raise HTTPException(404, str(exc))
    fname = Path(tpl.path).name
    removed = reset_prompt_override(step_dir, fname)
    return {"reset": removed, "filename": fname}




# =====================================================================
# Tools — 工具集（每个 SOP AI 步骤一个工具）
# =====================================================================

import asyncio as _asyncio
import json as _json
import re as _re
import time as _time
import traceback as _tb

# Tool catalog: AI 工具集
# 每个工具：id / step / name / icon / description / fields / output / 提示词跳转
#
# 命名说明:8 步 SOP 里的 step3(原"开发自测")在本工具集合并到 step4 接口测试,
# 因此 catalog 里只有 step1/2/4/5/6 + network_resilience / h5_adapt / seo_audit。
# 如需补回 step3,在此列表新增 id="step3" 的项并新建 configs/prompts/step3_*/。
TOOL_CATALOG: list[dict[str, Any]] = [
    {
        "id": "step1",
        "step": "step1",
        "prompt_dir": "step1_requirement",
        "name": "需求评审",
        "icon": "📋",
        "tagline": "把 PRD/UI/原型拆成可测结构，识别遗漏与门禁",
        "description": (
            "AI 自动阅读需求材料，输出《需求拆解报告》：模块拆解、主流程/异常流程、"
            "前后端交互链路、状态流转、需求遗漏点。命中关键缺失时直接 reject_with_report 阻止下一步。"
        ),
        "responsible": "AI 主执行",
        "output": "《需求拆解报告》",
        "gate": "材料完整度 < 80% 或 blocker 遗漏 ≥ 3 时阻止进入测试用例设计",
        "prompts": ["step1.1", "step1.2", "step1.3", "step1.4", "step1.5"],
        "endpoint": "/api/tools/step1/run",
        "substeps_optional": True,
        "input": {
            "label": "需求与文档",
            "hint": "粘贴需求文本 / 上传 PRD/原型/UI/流程/接口文档 — 任意拼合",
            "primary_key": "documents",
            "format": "text",
        },
    },
    {
        "id": "step2",
        "step": "step2",
        "prompt_dir": "step2_testcase",
        "name": "测试用例设计",
        "icon": "🧪",
        "tagline": "基于需求拆解报告生成 P0/P1/P2 用例集",
        "description": (
            "覆盖主流程、异常、边界、权限、状态流转、双端差异。每条用例含步骤、预期、"
            "依赖、自动化适配标记。最终输出统一字段的《测试用例集》。"
        ),
        "responsible": "AI 主执行",
        "output": "《测试用例集（P0/P1/P2）》",
        "gate": "主流程未覆盖、优先级错误、用例不可执行时不流转",
        "prompts": ["step2.1", "step2.2", "step2.3", "step2.4", "step2.5"],
        "endpoint": "/api/tools/step2/run",
        "substeps_optional": True,
        "input": {
            "label": "需求与场景",
            "hint": "粘贴需求文档 / 业务流程 / 用例骨架 — 直接生成 P0/P1/边界/权限/状态机用例",
            "primary_key": "documents",
            "format": "text",
        },
    },
    {
        "id": "step4",
        "step": "step4",
        "prompt_dir": "step4_api",
        "name": "接口测试",
        "icon": "🔌",
        "tagline": "覆盖功能、安全、边界值，给版本可测性结论",
        "description": (
            "梳理核心接口清单与业务链路，覆盖功能正确性 / 安全与权限 / 边界值与异常处理。"
            "命中核心接口不通、鉴权失效、关键字段缺失等条件 → 退回提测。"
        ),
        "responsible": "AI + 测试",
        "output": "《接口测试报告》",
        "gate": "核心接口不通 / 鉴权失效 / 主流程接口报错 → 退回提测",
        "prompts": ["step4.1", "step4.2", "step4.3", "step4.4", "step4.5"],
        "endpoint": "/api/tools/step4/run",
        "substeps_optional": True,
        "input": {
            "label": "接口与场景",
            "hint": "粘贴 API 文档 / OpenAPI / Postman / 接口清单 / 业务说明 — 直接出功能/性能/安全/边界/契约用例",
            "primary_key": "documents",
            "format": "text",
        },
    },
    {
        "id": "step6",
        "step": "step6",
        "prompt_dir": "step6_agent",
        "name": "Agent 自动化执行",
        "icon": "🤖",
        "tagline": "Agent 跑 P0 + 关键 P1/P2，归因 + 阻塞清单",
        "description": (
            "执行前检查（环境/账号/数据/依赖）→ P0 全量执行 → 关键 P1/P2 执行 → "
            "失败归因（前端/后端/接口/数据/环境/需求）→ 输出《Agent 自动化执行报告》。"
            "默认 dry-run，需要真实环境时手动开 --live。"
        ),
        "responsible": "AI Agent",
        "output": "《Agent 自动化执行报告》",
        "gate": "页面不可达 / 元素缺失 / 数据或权限阻塞 → 标记阻塞",
        "prompts": ["step6.1", "step6.2", "step6.3", "step6.4", "step6.5"],
        "endpoint": "/api/tools/step6/run",
        "substeps_optional": True,
        "input": {
            "label": "执行材料",
            "hint": "粘贴用例集 / 业务场景 / 环境信息;或【拖入用例 Excel(.xlsx,自动智能识别列)】;也可直接用「用例设计」工具产出的用例 — 逐条真执行 + 判 pass/fail + 失败归因",
            "primary_key": "documents",
            "format": "text",
        },
        "run_options": [
            {"key": "dry_run", "label": "Dry-run（仅推演不真实操作）", "default": True},
        ],
    },
    {
        "id": "network_resilience",
        "step": "network",
        "prompt_dir": "network_resilience",
        "name": "弱网/断网测试",
        "icon": "📶",
        "substeps_optional": True,
        "tagline": "弱网档位 + 离线/恢复 + 容错审计 + 退回判定",
        "description": (
            "识别网络敏感操作 → 设计弱网（慢/丢包/延迟）+ 断网（offline/reconnect/切换/挂起）"
            "用例 → 审计客户端容错 7 维度（超时/重试/幂等/取消/队列/错误分类/可观测）"
            "→ 输出《弱网容灾报告》。命中重复提交/数据丢失/缺幂等等即退回提测。"
        ),
        "responsible": "AI + 测试",
        "output": "《弱网容灾报告》",
        "gate": "断网→恢复出现重复提交/数据丢失/无幂等 → 退回提测",
        "prompts": ["net.1", "net.2", "net.3", "net.4", "net.5"],
        "endpoint": "/api/tools/network_resilience/run",
        "input": {
            "label": "业务材料",
            "hint": "PRD / 接口文档 / 主流程描述 / 客户端实现资料 — 任意拼合；测视频弱网请把宿主机 net_video_collect.py 采的 evidence.md 一并粘进来",
            "primary_key": "documents",
            "format": "text",
        },
        "run_options": [
            {"key": "mode_page_api", "label": "页面/接口弱网（容器实测：各档位加载/操作/断网恢复/资损）", "default": True},
            {"key": "mode_video", "label": "视频播放弱网（先用宿主机 net_video_collect.py 采视频证据并粘入材料）", "default": False},
        ],
    },
    {
        "id": "h5_adapt",
        "step": "h5",
        "prompt_dir": "h5_adapt",
        "name": "H5 适配初审",
        "icon": "📱",
        "substeps_optional": True,
        "tagline": "页面盘点 + 视口/安全区/浏览器矩阵/交互/性能 全维度",
        "description": (
            "针对 H5 / Web Mobile 在 iOS Safari、Android Chrome、微信 X5/MQQ、QQ/UC、"
            "钉钉/飞书/企业微信、抖音/小红书等 15+ 浏览器和 WebView 下的适配深度审计。"
            "覆盖 viewport / safe-area / 像素密度 / 长屏 / 表单键盘 / 触摸手势 / 性能 / 分享 / 唤起 App / 暗色模式 / 无障碍。"
        ),
        "responsible": "AI + 前端",
        "output": "《H5 适配初审报告》",
        "gate": "主流程在微信/iOS Safari/Android Chrome 任一 critical 失败 → 退回",
        "prompts": ["h5.1", "h5.2", "h5.3", "h5.4", "h5.5"],
        "endpoint": "/api/tools/h5_adapt/run",
        "input": {
            "label": "H5 资料（URL + 业务说明）",
            "hint": (
                "建议按下面两段格式贴：\n\n"
                "## 待测 H5 URL\n"
                "<https://...> （系统自动多 viewport 截图）\n\n"
                "## 业务/客诉/已知差异\n"
                "<页面用途 / 已知兼容问题 / UA 样本 / HTML head / CSS 片段>\n\n"
                "工具会自动截取 6 个常见手机分辨率（iPhone SE / 13 / 14 ProMax / Galaxy S20 / iPad / Desktop）。"
            ),
            "primary_key": "documents",
            "format": "text",
        },
    },
    {
        "id": "seo_audit",
        "step": "seo",
        "prompt_dir": "seo_audit",
        "name": "SEO 深度审计",
        "icon": "🔍",
        "substeps_optional": True,
        "tagline": "抓取 + 元数据 + 内容 + Core Web Vitals 四维深度",
        "description": (
            "审计协议/重定向/robots/sitemap/canonical/hreflang → 元数据与 schema.org → "
            "内容结构与可访问性 → Core Web Vitals → 输出按 ROI 排序的整改清单 + 快速胜利。"
        ),
        "responsible": "AI 主执行",
        "output": "《SEO 深度审计报告》",
        "gate": "全站 noindex 误配 / robots 屏蔽全站 / sitemap 大量 4xx → 退回",
        "prompts": ["seo.1", "seo.2", "seo.3", "seo.4", "seo.5"],
        "endpoint": "/api/tools/seo_audit/run",
        "input": {
            "label": "站点资料",
            "hint": "robots.txt / sitemap.xml / 抽样 HTML / Lighthouse 数据 / 业务定位 — 任意拼合",
            "primary_key": "documents",
            "format": "text",
        },
    },
]


@app.get("/api/tools")
async def api_tools_catalog() -> dict[str, Any]:
    """工具集目录（5 个 AI 工具）。"""
    return {"total": len(TOOL_CATALOG), "tools": TOOL_CATALOG}


@app.get("/api/tools/{tool_id}")
async def api_tool_detail(tool_id: str, request: Request) -> dict[str, Any]:
    # 路由顺序兼容：FastAPI 会用 wildcard 抢先匹配 /api/tools/runs，
    # 这里转发到正确的 list 端点。/api/tools/runs/{run_id} 是不同 path
    # 模板，不会被这个 wildcard 拦下。
    if tool_id == "runs":
        return await api_tool_run_list(request)
    tool = next((t for t in TOOL_CATALOG if t["id"] == tool_id), None)
    if not tool:
        raise HTTPException(404, f"unknown tool: {tool_id}")
    return tool


# ----- 异步任务管理 -----
# 简单内存版任务表：{run_id: {status, started_at, finished_at, tool_id, progress, report, error}}
_RUNS: dict[str, dict[str, Any]] = {}
# run_id → asyncio.Task,用于「停止执行」时取消任务(与 _RUNS 分开,避免 Task 进 JSON)
_RUN_TASKS: dict[str, Any] = {}

_RUN_TOOL_NAMES = {"step1": "需求评审", "step2": "测试用例", "step4": "接口测试",
                   "step5": "UI 一致性比对", "step6": "Agent 自动化", "seo_audit": "SEO 审计",
                   "network_resilience": "弱网/断网", "h5_adapt": "H5 适配"}


def _hydrate_runs_from_disk(only_run_id: str | None = None) -> int:
    """把磁盘历史报告(stepN_<run_id>.json 等)补进 _RUNS,使重启后仍能在 app 内查看
    历史运行(否则内存只读 → 重启就 404)。只补不覆盖;only_run_id 给定时只补那一个。"""
    import re as _re
    import json as _json   # 模块级无全局 json,必须局部导入(否则 json.loads → NameError 被 except 吞掉)
    base = settings.report_output_dir
    if not base.exists():
        return 0
    pat = _re.compile(r"^(step[12456]|seo_audit|network_resilience|h5_adapt)_(.+)\.json$")
    n = 0
    for p in sorted(base.glob("*.json"), key=lambda x: x.stat().st_mtime):
        m = pat.match(p.name)
        if not m:
            continue
        tool_id, run_id = m.group(1), m.group(2)
        if only_run_id and run_id != only_run_id:
            continue
        if run_id in _RUNS:
            continue
        try:
            rep = _json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue
        meta = rep.get("meta") or {}
        mt = p.stat().st_mtime
        _RUNS[run_id] = {
            "run_id": run_id, "tool_id": tool_id, "tool_name": _RUN_TOOL_NAMES.get(tool_id, tool_id),
            "status": "succeeded", "progress": "完成(历史报告)",
            "started_at": mt, "finished_at": mt,
            "tenant_id": meta.get("tenant_id", "default"), "project_id": meta.get("project_id", "default"),
            "project_code": meta.get("project_code", ""), "project_name": meta.get("project_name", ""),
            "owner_user_id": meta.get("owner_user_id"), "owner_username": meta.get("owner_username", ""),
            "report": rep, "_from_disk": True,
        }
        n += 1
    return n


@app.on_event("startup")
async def _startup_hydrate_runs():
    try:
        cnt = _hydrate_runs_from_disk()
        if cnt:
            print(f"[startup] 从磁盘恢复 {cnt} 条历史运行报告")
    except Exception as exc:
        print(f"[startup] 历史报告恢复失败: {exc}")


_KNOWN_INPUT_KEYS = {
    "prd", "prototype", "ui_design", "flow_chart", "api_doc", "business_rules",
    "requirement_report", "test_case_report",
    "automation_risk", "env_info",
    "design_assets", "actual_snapshots", "state_captures",
    # network_resilience / seo_audit single-blob input
    "documents", "raw_documents",
    "main_flow", "client_impl",
    "site_info", "robots_txt", "sitemap_xml", "crawl_samples",
    "page_samples", "page_dom_samples", "business_positioning",
    "perf_data", "resource_inventory", "waterfall",
    # h5_adapt advanced keys
    "business_doc", "page_list", "screenshots", "user_scenarios",
    "html_head_samples", "css_snippets", "js_entry_deps", "key_apis",
    "ua_samples", "complaints", "interactive_components", "forms",
    "perf_samples",
}


def _shape_orchestrator_inputs(
    tool: dict[str, Any], raw: dict[str, Any]
) -> dict[str, Any]:
    """Map simplified single-field UI body to the orchestrator's expected keys.

    Logic:
      - If the tool's primary_key is a real input key → assign directly.
      - If primary_key == "_documents" (auto mode for step4/step5):
          - If the value is a dict, pass through any known keys
            (requirement_report, test_case_report, api_doc, …).
          - Otherwise dump as 'documents' raw text — orchestrators tolerate
            missing keys via `(未提供)` placeholders.
      - Run options (e.g. dry_run) are merged last.
    """
    inp_meta = tool.get("input", {}) or {}
    primary = inp_meta.get("primary_key")
    out: dict[str, Any] = {}

    payload = raw.pop(primary, None) if primary and primary != "_documents" else None
    docs = raw.pop("_documents", None) if primary == "_documents" else None

    if primary and primary != "_documents" and payload is not None:
        out[primary] = payload
    elif docs is not None:
        if isinstance(docs, dict):
            for k, v in docs.items():
                if k in _KNOWN_INPUT_KEYS:
                    out[k] = v
            if not out and "raw" in docs:
                out["raw_documents"] = docs["raw"]
        else:
            # Free text — we don't know how to split; orchestrators expect
            # specific keys, so park it under 'raw_documents'. Most prompts
            # treat absent keys as `(未提供)` and proceed.
            out["raw_documents"] = docs

    # Run options (checkboxes etc.) come from the form too.
    for ro in tool.get("run_options", []) or []:
        out[ro["key"]] = bool(raw.pop(ro["key"], ro.get("default", False)))

    # Anything left in raw that is a real orchestrator key (advanced override)
    for k, v in list(raw.items()):
        if k.startswith("__"):
            continue
        if k in _KNOWN_INPUT_KEYS:
            out[k] = v
    # 项目编号 + 名称 透传(build_meta 会读)
    if raw.get("project_code"):
        out["project_code"] = str(raw["project_code"]).strip()
    if raw.get("project_name"):
        out["project_name"] = str(raw["project_name"]).strip()
    return out


_URL_RE = _re.compile(r"https?://[\w\.\-/:?#\[\]@!$&'()*+,;=%]+", _re.IGNORECASE)

# Tool ID → (viewport tuples). Each tuple = (label, width, height).
_TOOL_VIEWPORTS = {
    # h5_adapt 已改为「分析宿主三端真机证据」模式，不在容器内用桌面 Chromium 模拟视口截图。
    "step5": [
        ("Mobile", 375, 812),
        ("Desktop", 1440, 900),
    ],
    # SEO 审计:抓首页 + Mobile 视图,LLM 看 LCP/CLS 时有图可参
    "seo_audit": [
        ("Desktop", 1440, 900),
        ("Mobile", 375, 812),
    ],
    # 弱网/断网:有图能看「断网态错误页 vs 在线态」对比
    "network_resilience": [
        ("Desktop", 1440, 900),
    ],
}


async def _execute_apis_agentic(ctx: Any, state: dict[str, Any]) -> dict[str, Any] | None:
    """step4:AI 真实调用接口的 agentic 循环 — LLM 决定发什么请求,httpx 真发,
    AI 看真实响应找 bug。真实 req/resp 记录注入 documents,供后续 substep 分析。
    """
    docs = (ctx.inputs or {}).get("documents") or ""
    # 有上传文件(接口文档 PDF/文本)时,URL 可能在文件里 → 不因文本无 URL 而跳过;
    # agent 会从附带的文件内容块里读出接口地址再真发请求。
    has_files = bool(getattr(ctx, "files", None))
    if (not isinstance(docs, str) or not _re.search(r"https?://", docs)) and not has_files:
        state.setdefault("logs", []).append({
            "ts": _time.time(), "event": "api.execute.skip", "reason": "材料里没有可调用的 URL,也无上传文件"})
        return None
    import httpx
    import asyncio as _aio
    from packages.core.agent import agent_loop

    # ── 通用动作:send_request(httpx 真发)—— 模式A 与 模式B 的重放/变形都用它 ──
    async def http_request(args: dict[str, Any]) -> str:
        method = (args.get("method") or "GET").upper()
        url = args.get("url")
        if not url:
            return "缺少 url"
        headers = args.get("headers") or {}
        body = args.get("body")
        try:
            async with httpx.AsyncClient(timeout=25.0, follow_redirects=True) as cli:
                kw: dict[str, Any] = {}
                if isinstance(body, (dict, list)):
                    kw["json"] = body
                elif isinstance(body, str) and body:
                    kw["content"] = body
                r = await cli.request(method, url, headers=headers, **kw)
            hdrs = dict(list(r.headers.items())[:10])
            return f"HTTP {r.status_code} {r.reason_phrase}\nresp-headers: {hdrs}\nbody(前1500字符):\n{r.text[:1500]}"
        except Exception as exc:
            return f"请求失败: {type(exc).__name__}: {str(exc)[:200]}"

    tools: dict[str, Any] = {"send_request": http_request}

    # ── 模式B 能力:真驱动前端 + 抓包。材料给了前端就走前端;Playwright 不可用则降级为纯接口(模式A)──
    captured: list[dict[str, Any]] = []
    pw = browser = page = None
    try:
        from playwright.async_api import async_playwright
        pw = await async_playwright().start()
        browser = await pw.chromium.launch(headless=True)
        page = await browser.new_page(viewport={"width": 1440, "height": 900})

        def _on_resp(resp: Any) -> None:
            try:
                req = resp.request
                if req.resource_type in ("xhr", "fetch"):
                    h = req.headers
                    captured.append({
                        "method": req.method, "url": req.url, "status": resp.status,
                        "auth": "有" if (h.get("authorization") or h.get("cookie") or h.get("x-token") or h.get("token")) else "无",
                        "ctype": h.get("content-type", ""), "post": (req.post_data or "")[:400]})
            except Exception:
                pass
        page.on("response", _on_resp)

        async def navigate(a: dict[str, Any]) -> str:
            url = a.get("url")
            if not url:
                return "缺少 url"
            try:
                r = await page.goto(url, timeout=30000, wait_until="domcontentloaded")
                await _aio.sleep(1.8)
                return f"已打开前端 {page.url} | HTTP {r.status if r else '?'} | 标题: {await page.title()}"
            except Exception as e:
                return f"打开失败: {type(e).__name__}: {str(e)[:150]}"

        async def click(a: dict[str, Any]) -> str:
            try:
                if a.get("text"):
                    await page.get_by_text(a["text"], exact=False).first.click(timeout=8000)
                elif a.get("selector"):
                    await page.click(a["selector"], timeout=8000)
                else:
                    return "需提供 text 或 selector"
                await _aio.sleep(1.6)
                return f"已点击 | 当前 {page.url}"
            except Exception as e:
                return f"点击失败: {str(e)[:150]}"

        async def form_input(a: dict[str, Any]) -> str:
            try:
                val = str(a.get("value", ""))
                if a.get("placeholder"):
                    await page.get_by_placeholder(a["placeholder"]).first.fill(val, timeout=8000)
                    tgt = a["placeholder"]
                elif a.get("selector"):
                    await page.fill(a["selector"], val, timeout=8000)
                    tgt = a["selector"]
                else:
                    return "需提供 selector 或 placeholder"
                return f"已填写 {tgt} = {val[:40]}"
            except Exception as e:
                return f"填写失败: {str(e)[:150]}"

        async def inspect(a: dict[str, Any]) -> str:
            try:
                sig = await page.evaluate(
                    "() => ({title:document.title,"
                    "bodyText:(document.body.innerText||'').slice(0,600),"
                    "inputs:[...document.querySelectorAll('input,select,textarea')].map(e=>({name:e.name||e.id,ph:e.placeholder,type:e.type})).slice(0,20),"
                    "buttons:[...document.querySelectorAll('button,a')].map(e=>(e.innerText||'').trim()).filter(Boolean).slice(0,24)})")
                return _json.dumps(sig, ensure_ascii=False)[:1900]
            except Exception as e:
                return f"抽取失败: {str(e)[:150]}"

        async def read_network(a: dict[str, Any]) -> str:
            # 读出前端真实发出的接口请求(抓包)——模式B 的核心:看前端实际调了哪些接口、怎么调
            if not captured:
                return "暂无捕获到接口请求。先 navigate 前端页面,再 click/form_input 触发业务操作,前端才会发请求。"
            seen: set = set()
            out: list[dict[str, Any]] = []
            for c in captured[-60:]:
                k = (c["method"], c["url"].split("?")[0])
                if k in seen:
                    continue
                seen.add(k)
                out.append(c)
            return f"前端真实发出的接口(去重后 {len(out)} 个):\n" + "\n".join(
                f"- {c['method']} {c['url'][:130]} → {c['status']} | 鉴权:{c['auth']} | {c['ctype'][:30]} | body:{c['post'][:120]}" for c in out)

        tools.update({"navigate": navigate, "click": click, "form_input": form_input,
                      "inspect": inspect, "read_network": read_network})
    except Exception as exc:
        state.setdefault("logs", []).append({
            "ts": _time.time(), "event": "api.browser.unavailable",
            "reason": "Playwright 不可用,降级为纯接口模式A", "error": str(exc)[:150]})

    ex_md = Path(__file__).resolve().parent.parent.parent / "configs" / "prompts" / "step4_api" / "_execute.md"
    try:
        sysp = ex_md.read_text(encoding="utf-8")
    except Exception:
        sysp = ("你真实测试 HTTP 接口(双模):材料含前端页面就先 navigate 前端 + click/form_input 触发、"
                "read_network 抓前端真实请求,再 send_request 重放/变形深测;只有接口文档就直接 send_request。"
                "每轮输出 JSON {thought, tool, args, finding, done}。覆盖正常/必填/边界/鉴权/越权/契约。")
    has_browser = "navigate" in tools
    task = (f"测试者提供的物料如下。{'材料里若有前端页面 URL → 走模式B(从真实前端发起请求);' if has_browser else ''}"
            f"只有接口文档 → 走模式A(直接对接口真发请求)。\n\n{docs[:6000]}\n\n"
            f"现在开始,输出第一步动作的 JSON。")
    state["progress"] = "AI 接口测试中（双模）…"
    res = await agent_loop(
        ctx.llm, sysp, task, tools, max_steps=28,
        on_step=lambda r: state.update({"progress": f"接口测试 第{r.get('step')}步: {r.get('tool','')} {((r.get('args') or {}).get('method','') or '')} {((r.get('args') or {}).get('url','') or (r.get('args') or {}).get('text','') or '')[:46]}"}),
        files=getattr(ctx, "files", None),
    )
    try:
        if browser:
            await browser.close()
        if pw:
            await pw.stop()
    except Exception:
        pass

    # 把真实记录(前端抓包 + agent 调用)注入 documents,供 5 个 substep 基于真实结果分析
    lines = ["", "", "## 真实接口测试记录(AI 实测,以下结论须基于这些真实 req/resp)"]
    if captured:
        lines.append(f"### 模式B:从真实前端抓到的接口请求(共 {len({(c['method'], c['url'].split('?')[0]) for c in captured})} 个不同)")
        seen2: set = set()
        for c in captured:
            k = (c["method"], c["url"].split("?")[0])
            if k in seen2:
                continue
            seen2.add(k)
            lines.append(f"- {c['method']} {c['url'][:130]} → {c['status']} | 鉴权:{c['auth']}")
    lines.append("### AI 测试动作记录")
    for t in res.get("transcript", []):
        a = t.get("args") or {}
        ident = a.get("url", "") or a.get("text", "") or a.get("selector", "") or a.get("placeholder", "")
        lines.append(f"- [{t.get('step')}] {t.get('tool','')} {(a.get('method') or '').upper()} {ident} "
                     f"→ {str(t.get('result',''))[:260]}")
    if res.get("findings"):
        lines.append("\nAI 实测中已标记的问题:")
        for f in res["findings"]:
            lines.append(f"- {_json.dumps(f, ensure_ascii=False)[:300]}")
    ctx.inputs["documents"] = docs + "\n".join(lines)
    state.setdefault("logs", []).append({
        "ts": _time.time(), "event": "api.execute", "mode": "B" if captured else ("AB" if has_browser else "A"),
        "steps": res.get("steps"), "captured": len(captured), "findings": len(res.get("findings") or [])})
    return res


async def _run_browser_agent(
    ctx: Any, state: dict[str, Any], prompt_rel: str, name_prefix: str,
    shots_out: list[dict[str, Any]], with_http: bool = False, with_network: bool = False,
    max_steps: int = 16,
) -> dict[str, Any] | None:
    """通用浏览器 agentic 执行 — AI 真驱动 Playwright(导航/抽信号/点击/截图/视口/断网),
    服务 SEO / H5 / step6 / 弱网。真实执行记录注入 documents。
    """
    docs = (ctx.inputs or {}).get("documents") or ""
    has_files = bool(getattr(ctx, "files", None))
    if (not isinstance(docs, str) or not _re.search(r"https?://", docs)) and not has_files:
        state.setdefault("logs", []).append({"ts": _time.time(), "event": "browser.agent.skip", "reason": "材料无 URL,也无上传文件"})
        return None
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        return None
    from packages.core.agent import agent_loop
    import httpx as _httpx
    import asyncio
    sc_dir = Path(settings.evidence_output_dir) / "screenshots"
    sc_dir.mkdir(parents=True, exist_ok=True)
    cap_idx = [0]
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        page = await browser.new_page(viewport={"width": 1440, "height": 900})

        async def navigate(a):
            url = a.get("url")
            if not url:
                return "缺少 url"
            try:
                r = await page.goto(url, timeout=30000, wait_until="domcontentloaded")
                await asyncio.sleep(1.5)
                return f"已打开 {page.url} | HTTP {r.status if r else '?'} | 标题: {await page.title()}"
            except Exception as e:
                return f"打开失败: {type(e).__name__}: {str(e)[:150]}"

        async def inspect(a):
            try:
                sig = await page.evaluate(
                    "() => ({title:document.title,"
                    "metaDesc:(document.querySelector('meta[name=description]')||{}).content||'',"
                    "h1:[...document.querySelectorAll('h1')].map(e=>e.innerText.trim()).slice(0,6),"
                    "h1Count:document.querySelectorAll('h1').length,"
                    "viewportMeta:(document.querySelector('meta[name=viewport]')||{}).content||'',"
                    "lang:document.documentElement.lang||'',"
                    "canonical:(document.querySelector('link[rel=canonical]')||{}).href||'',"
                    "imgTotal:document.querySelectorAll('img').length,"
                    "imgNoAlt:[...document.querySelectorAll('img')].filter(i=>!i.getAttribute('alt')).length,"
                    "links:document.querySelectorAll('a[href]').length,"
                    "docWidth:document.documentElement.scrollWidth,winWidth:window.innerWidth,"
                    "bodyText:(document.body.innerText||'').slice(0,700)})")
                return _json.dumps(sig, ensure_ascii=False)[:2200]
            except Exception as e:
                return f"抽取失败: {str(e)[:150]}"

        async def click(a):
            try:
                if a.get("text"):
                    await page.get_by_text(a["text"], exact=False).first.click(timeout=8000)
                elif a.get("selector"):
                    await page.click(a["selector"], timeout=8000)
                else:
                    return "需提供 text 或 selector"
                await asyncio.sleep(1.5)
                return f"已点击 | 当前 {page.url} | 标题: {await page.title()}"
            except Exception as e:
                return f"点击失败: {str(e)[:150]}"

        async def screenshot(a):
            cap_idx[0] += 1
            fn = f"{name_prefix}_{cap_idx[0]}.png"
            try:
                await page.screenshot(path=str(sc_dir / fn), full_page=bool(a.get("full_page")))
                shots_out.append({"url": page.url, "viewport": a.get("label") or f"页面{cap_idx[0]}",
                                  "width": "", "height": "", "filename": fn})
                return f"已截图 {fn}({a.get('label','')})"
            except Exception as e:
                return f"截图失败: {str(e)[:120]}"

        async def set_viewport(a):
            try:
                w, h = int(a.get("width", 375)), int(a.get("height", 812))
                await page.set_viewport_size({"width": w, "height": h})
                await asyncio.sleep(1.0)
                return f"已切视口 {w}x{h}"
            except Exception as e:
                return f"切视口失败: {str(e)[:120]}"

        handlers: dict[str, Any] = {"navigate": navigate, "inspect": inspect,
                                    "click": click, "screenshot": screenshot, "set_viewport": set_viewport}

        if with_network:
            async def set_network(a):
                mode = (a.get("mode") or "").lower()
                try:
                    if mode in ("offline", "断网"):
                        await page.context.set_offline(True); return "已切断网络(offline)"
                    if mode in ("online", "恢复"):
                        await page.context.set_offline(False); return "已恢复网络(online)"
                    if mode in ("slow", "弱网", "3g"):
                        cdp = await page.context.new_cdp_session(page)
                        await cdp.send("Network.enable")
                        await cdp.send("Network.emulateNetworkConditions", {
                            "offline": False, "latency": 400,
                            "downloadThroughput": 50 * 1024, "uploadThroughput": 20 * 1024})
                        return "已切弱网(慢 3G:50KB/s,400ms 延迟)"
                    return f"未知 mode={mode}(用 offline/online/slow)"
                except Exception as e:
                    return f"切网络失败: {str(e)[:120]}"
            handlers["set_network"] = set_network

        if with_http:
            async def send_request(a):
                method = (a.get("method") or "GET").upper()
                url = a.get("url")
                if not url:
                    return "缺少 url"
                try:
                    async with _httpx.AsyncClient(timeout=25.0, follow_redirects=True) as cli:
                        kw = {}
                        b = a.get("body")
                        if isinstance(b, (dict, list)):
                            kw["json"] = b
                        elif isinstance(b, str) and b:
                            kw["content"] = b
                        r = await cli.request(method, url, headers=a.get("headers") or {}, **kw)
                    return f"HTTP {r.status_code} | body(前1200): {r.text[:1200]}"
                except Exception as e:
                    return f"请求失败: {str(e)[:150]}"
            handlers["send_request"] = send_request

        try:
            ex_md = Path(__file__).resolve().parent.parent.parent / "configs" / "prompts" / prompt_rel
            sysp = ex_md.read_text(encoding="utf-8")
        except Exception:
            sysp = "你真实驱动浏览器测试。每轮输出 JSON{thought, <动作字段:navigate/inspect/click/screenshot/set_viewport...>, finding, done}。系统真实执行并回灌结果。"
        task = f"目标与材料:\n{docs[:5000]}\n\n现在开始,输出第一步动作 JSON。"
        state["progress"] = "AI 真实驱动浏览器中…"
        res = await agent_loop(
            ctx.llm, sysp, task, handlers, max_steps=max_steps,
            on_step=lambda r: state.update({"progress": f"浏览器执行 第{r.get('step')}步: {r.get('action')}"}),
            files=getattr(ctx, "files", None))  # 上传的目标说明文件直传给规划 AI
        try:
            await browser.close()
        except Exception:
            pass
    lines = ["", "", "## 真实浏览器执行记录(AI 实测,结论须基于此)"]
    for t in res.get("transcript", []):
        lines.append(f"- [{t.get('step')}] {t.get('action')} {_json.dumps(t.get('args') or {}, ensure_ascii=False)[:160]} → {str(t.get('result',''))[:260]}")
    if res.get("findings"):
        lines.append("\nAI 实测标记的问题:")
        for f in res["findings"]:
            lines.append(f"- {_json.dumps(f, ensure_ascii=False)[:280]}")
    ctx.inputs["documents"] = docs + "\n".join(lines)
    state.setdefault("logs", []).append({"ts": _time.time(), "event": "browser.agent",
                                          "steps": res.get("steps"), "findings": len(res.get("findings") or [])})
    return res


def _load_uploaded_files(ctx: Any) -> list[dict[str, Any]]:
    """从 documents 文本里解析 file_ref=<id>,加载真实文件元数据(uploads/files/<id>.json)。

    返回 [{path, mime, kind, filename}] 供 base.run_substep 作为 document/image
    内容块经 stdin 直传 LLM。文件缺失/元数据损坏则跳过(不影响其余)。
    """
    docs = (ctx.inputs or {}).get("documents") or ""
    if not isinstance(docs, str):
        try:
            docs = _json.dumps(docs, ensure_ascii=False)
        except Exception:
            docs = str(docs)
    refs = list(dict.fromkeys(_re.findall(r"file_ref=([0-9a-fA-F]{6,32})", docs)))
    if not refs:
        return []
    files_dir = settings.report_output_dir.parent.parent / "uploads" / "files"
    out: list[dict[str, Any]] = []
    for ref in refs[:20]:
        meta_p = files_dir / f"{ref}.json"
        if not meta_p.exists():
            continue
        try:
            meta = _json.loads(meta_p.read_text(encoding="utf-8"))
        except Exception:
            continue
        p = meta.get("path")
        if not p or not Path(p).exists():
            continue
        out.append({
            "path": p,
            "mime": meta.get("mime") or "",
            "kind": meta.get("kind") or "document_text",
            "filename": meta.get("filename") or Path(p).name,
        })
    return out


# ───────────────────────────────────────────────────────────────────────────
# 用例 Excel 智能解析(step6 输入):上传的 .xlsx 用例表 → 智能识别列 →
# 与 step2 用例设计同一套 cases 结构(两路归一,下游执行/判定/归因统一)。
# ───────────────────────────────────────────────────────────────────────────
_CASE_COL_ALIASES = {
    "id": ["用例编号", "用例id", "编号", "id", "caseid", "序号", "用例号", "case号"],
    "title": ["用例标题", "用例名称", "标题", "名称", "title", "用例描述", "用例", "测试点", "case", "摘要"],
    "preconditions": ["前置条件", "前提条件", "预置条件", "前提", "precondition", "preconditions", "前置"],
    "steps": ["测试步骤", "操作步骤", "执行步骤", "步骤", "steps", "step", "操作", "用例步骤"],
    "expected": ["预期结果", "预期", "期望结果", "预期输出", "expected", "结果", "期望", "验证点"],
    "priority": ["优先级", "级别", "priority", "重要程度", "等级"],
    "type": ["用例类型", "类型", "type", "分类", "category", "用例分类"],
    "module": ["模块", "功能模块", "module", "所属模块", "功能点", "所属功能"],
}

def _norm_col(s: str) -> str:
    return str(s or "").strip().lower().replace(" ", "").replace("_", "").replace("-", "")

def _norm_priority(p: str) -> str:
    t = str(p or "").strip().upper()
    for k in ("P0", "P1", "P2", "P3"):
        if k in t:
            return k
    return {"高": "P0", "critical": "P0", "紧急": "P0", "1": "P0",
            "中": "P1", "重要": "P1", "2": "P1",
            "低": "P3", "次要": "P3", "3": "P2", "4": "P3"}.get(str(p or "").strip().lower(), "P2")

def _parse_cases_from_excel(path: str) -> tuple[list[dict[str, Any]], dict[str, int]]:
    """智能识别列名,把测试用例 Excel 解析成与 step2 同构的 cases。容忍非模板表格。"""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    cases: list[dict[str, Any]] = []
    colmap: dict[str, int] = {}
    for ws in wb.worksheets:
        rows = [r for r in ws.iter_rows(values_only=True)]
        if not rows:
            continue
        hi = next((i for i, r in enumerate(rows) if any(c for c in r)), 0)
        header = [str(c or "").strip() for c in rows[hi]]
        cm: dict[str, int] = {}
        for f, al in _CASE_COL_ALIASES.items():
            for ci, h in enumerate(header):
                hn = _norm_col(h)
                if hn and any(_norm_col(a) == hn or (len(_norm_col(a)) >= 2 and _norm_col(a) in hn) for a in al):
                    cm[f] = ci
                    break
        if "title" not in cm and "id" not in cm:
            continue  # 该 sheet 不像用例表
        def g(r, f):
            ci = cm.get(f)
            return str(r[ci]).strip() if ci is not None and ci < len(r) and r[ci] is not None else ""
        sheet_cases: list[dict[str, Any]] = []
        for r in rows[hi + 1:]:
            if not any(c for c in r):
                continue
            title = g(r, "title") or g(r, "id")
            if not title:
                continue
            steps_raw = g(r, "steps")
            steps = [s.strip() for s in _re.split(r"[\n;；]|[①②③④⑤⑥⑦⑧⑨⑩]|\d+\s*[.、)]\s*", steps_raw) if s.strip()]
            sheet_cases.append({
                "id": g(r, "id") or f"TC-{len(cases)+len(sheet_cases)+1:04d}",
                "title": title, "module": g(r, "module"),
                "preconditions": g(r, "preconditions"),
                "steps": steps or ([steps_raw] if steps_raw else []),
                "expected": g(r, "expected"),
                "priority": _norm_priority(g(r, "priority")),
                "type": g(r, "type") or "main",
            })
        if sheet_cases:
            cases.extend(sheet_cases)
            if not colmap:
                colmap = cm
    return cases, colmap

def _inject_excel_cases(ctx: Any) -> int:
    """扫描上传文件里的 .xlsx 用例表 → 解析成统一 cases → 注入 documents + test_case_report。返回用例数。"""
    try:
        files = _load_uploaded_files(ctx)
    except Exception:
        return 0
    all_cases: list[dict[str, Any]] = []
    for f in files:
        name = (f.get("filename") or "").lower()
        mime = (f.get("mime") or "").lower()
        if name.endswith((".xlsx", ".xlsm")) or "spreadsheet" in mime:
            try:
                cases, _cm = _parse_cases_from_excel(f["path"])
                all_cases.extend(cases)
            except Exception:
                pass
    if not all_cases:
        return 0
    lines = ["", "", "## 待执行用例(从上传 Excel 智能解析,已转成与用例设计一致的统一结构)"]
    for c in all_cases:
        steps = " → ".join(c["steps"]) if c.get("steps") else "(未填)"
        lines.append(f"- [{c['id']} | {c.get('priority')} | {c.get('type')}] {c['title']}"
                     f"\n  模块:{c.get('module') or '-'} | 前置条件:{c.get('preconditions') or '-'}"
                     f"\n  步骤:{steps}\n  预期结果:{c.get('expected') or '(未填)'}")
    ctx.inputs["documents"] = (ctx.inputs.get("documents") or "") + "\n".join(lines)
    p0 = [c for c in all_cases if c.get("priority") == "P0"]
    p1 = [c for c in all_cases if c.get("priority") == "P1"]
    p2 = [c for c in all_cases if c.get("priority") not in ("P0", "P1")]
    tcr = ctx.inputs.get("test_case_report")
    if not isinstance(tcr, dict):
        tcr = {}
    tcr["cases"] = (tcr.get("cases") or []) + all_cases
    tcr["p0_cases"] = (tcr.get("p0_cases") or []) + p0
    tcr["p1_cases"] = (tcr.get("p1_cases") or []) + p1
    tcr["p2_cases"] = (tcr.get("p2_cases") or []) + p2
    ctx.inputs["test_case_report"] = tcr
    return len(all_cases)


async def _capture_screenshots_for_tool(
    tool_id: str,
    ctx: Any,
    state: dict[str, Any],
) -> list[dict[str, Any]] | None:
    """Detect URLs in the tool's `documents` input and capture viewport screenshots.

    Only runs for tools listed in `_TOOL_VIEWPORTS` (currently step5 / h5_adapt).
    Saves PNGs to <evidence_output_dir>/screenshots/.
    """
    viewports = _TOOL_VIEWPORTS.get(tool_id)
    if not viewports:
        return None

    docs = (ctx.inputs or {}).get("documents") or ""
    if not isinstance(docs, str):
        return None
    urls = list(dict.fromkeys(m.group(0).rstrip(".,;:!?)」")
                              for m in _URL_RE.finditer(docs)))
    # figma 链接只走宿主助手读设计图,不让 Playwright 当普通页面截(容器访问 figma 是登录墙/403,产生垃圾图)
    urls = [u for u in urls if "figma.com" not in u.lower()]
    urls = urls[:5]  # cap so we don't run for hours

    # APP 上传:解析 app_ref → AI 驱动探索(LLM 看屏决定跳弹窗/翻页/截图),Python 只执行
    app_refs = list(dict.fromkeys(_re.findall(r"app_ref=([0-9a-f]{6,32})", docs)))
    app_shots: list[dict[str, Any]] = []
    if app_refs:
        apps_dir = settings.report_output_dir.parent.parent / "uploads" / "apps"
        try:
            from packages.core.device import run_app_agentic, run_app_and_capture
            sc_dir = Path(settings.evidence_output_dir) / "screenshots"
            # 探索决策提示词(AI 行为写在 prompt 里,可编辑)
            _explore_md = Path(__file__).resolve().parent.parent.parent / "configs" / "prompts" / "step5_ui" / "_explore.md"
            try:
                explore_sys = _explore_md.read_text(encoding="utf-8")
            except Exception:
                explore_sys = "你在探索一个 APP,每轮给你截图+可点元素,输出 JSON{action:tap|back|swipe_up|done,target_index,capture,page_label,done}。先跳过弹窗,再走遍主要界面,每个正式界面 capture=true。"

            def _make_decider():
                async def decide(shot_path: str, clickables: list[dict[str, Any]], context: dict[str, Any] | None = None) -> dict[str, Any]:
                    context = context or {}
                    captured = context.get("captured") or []
                    recent = context.get("recent") or []
                    lines = [f'[{i}] text="{c.get("text","")}" pos=({c.get("cx")},{c.get("cy")})'
                             for i, c in enumerate(clickables[:40])]
                    user = (
                        f"【已截图页面（不要重复截、不要重复进入同形态）】：{captured or '(还没有)'}\n"
                        f"【最近动作】：\n" + ("\n".join("  " + str(s) for s in recent) if recent else "  (无)") + "\n\n"
                        "【当前屏可点元素】：\n" + ("\n".join(lines) if lines else "(无可点元素;可 back 或 done)")
                        + "\n\n请按结构**系统遍历所有不同形态的界面**(底部 tab 全走、列表→详情、登录态界面、设置/我的;"
                        "遇登录/注册自己完成 tap+input)。只有当**所有不同形态界面都覆盖**才 done,不要因张数草草结束。"
                        "输出下一步动作 JSON。")
                    resp = await ctx.llm.complete(
                        system=explore_sys,
                        messages=[{"role": "user", "content": user}],
                        images=[{"path": Path(shot_path), "mime": "image/png", "caption": "当前 APP 屏幕(实拍)"}],
                        max_tokens=900, allow_degrade=False,
                    )
                    try:
                        return resp.json()
                    except Exception:
                        return {"action": "done", "done": True, "thought": "决策解析失败,结束探索"}
                return decide

            for ref in app_refs[:1]:  # agentic 较慢,一次跑 1 个 APP
                apk = apps_dir / f"{ref}.apk"
                if not apk.exists():
                    continue
                state["progress"] = f"AI 探索 APP {ref}…(跳弹窗 + 翻页截图)"
                res = await run_app_agentic(
                    apk_path=str(apk), out_dir=str(sc_dir),
                    name_prefix=f"{tool_id}_{ctx.run_id[:8]}_{ref}",
                    decide=_make_decider(), max_steps=30,  # 遍历所有界面需更多步
                )
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "app.explore",
                    "ok": res.get("ok"), "steps": res.get("steps"), "error": res.get("error")})
                shots_got = res.get("screenshots") or []
                # 兜底:AI 一张都没截到(决策异常等)→ 退回确定式跑一遍
                if res.get("ok") and not shots_got:
                    fb = await run_app_and_capture(
                        apk_path=str(apk), out_dir=str(sc_dir),
                        name_prefix=f"{tool_id}_{ctx.run_id[:8]}_{ref}", explore=True, max_pages=3)
                    shots_got = fb.get("screenshots") or []
                    state.setdefault("logs", []).append({"ts": _time.time(), "event": "app.explore.fallback", "n": len(shots_got)})
                app_shots.extend(shots_got)
        except Exception as exc:
            state.setdefault("logs", []).append({
                "ts": _time.time(), "event": "app.run.failed", "error": str(exc)[:200]})

    # Figma 设计图(P3):扫物料里的 figma 链接 → 取设计图作"设计基线"
    # 优先走前端浏览器读图(持久登录,用户要求);无账号密码则退回 API token。
    figma_shots: list[dict[str, Any]] = []
    try:
        from packages.core.device.figma import (
            parse_figma_links, fetch_figma_image, fetch_figma_frames_via_api,
            fetch_figma_frames_via_export, fetch_figma_via_browser, fetch_figma_via_host_runner)
        from packages.core.auth_config import get_figma_token, get_figma_login
        links = parse_figma_links(docs)
        if links:
            token = get_figma_token()
            login = get_figma_login()
            profile_dir = str(settings.report_output_dir.parent.parent / "figma_profile")
            sc_dir = Path(settings.evidence_output_dir) / "screenshots"
            sc_dir.mkdir(parents=True, exist_ok=True)
            for i, lk in enumerate(links[:3]):
                state["progress"] = f"读取 Figma 设计图 {lk['file_key'][:8]}…"
                fname = f"{tool_id}_{ctx.run_id[:8]}_figma_{i+1}.png"
                fpath = str(sc_dir / fname)
                # 优先级:① 登录态浏览器导出 PDF 拆帧(**零 REST 额度**,只读权限也行 —— 首选)
                #         ② 只读 API token 逐帧渲染(REST,有账号级额度限制) ③ 宿主 /shot 截图
                _run_user = state.get("owner_user_id") or "default"
                res = await fetch_figma_frames_via_export(lk["url"], fpath, max_frames=50, user="default")
                mode = "export_pdf"
                if not res.get("ok"):
                    if token:
                        res = await fetch_figma_frames_via_api(
                            lk["file_key"], token, fpath, prefer_node=lk["node_id"])
                        mode = "api_frames"
                    if not res.get("ok"):
                        res = await fetch_figma_via_host_runner(lk["url"], fpath, user=_run_user)
                        mode = "host_runner"
                _frames = res.get("frames") or ([{"path": fpath, "name": "frame1"}] if res.get("ok") else [])
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "figma.fetch", "mode": mode,
                    "ok": res.get("ok"), "error": res.get("error"), "node": lk["node_id"],
                    "frames": len(_frames)})
                if res.get("ok"):
                    # 逐帧:每个顶层 Frame 一张设计基线图,供逐帧 vs 实拍对照。
                    for fj, fr in enumerate(_frames):
                        ffn = Path(fr["path"]).name
                        figma_shots.append({
                            "url": lk["url"],
                            "viewport": (f"设计基线(Figma)· 帧{fj + 1}" if len(_frames) > 1 else "设计基线(Figma)"),
                            "width": "", "height": "", "filename": ffn, "is_design": True,
                            "node_name": fr.get("name")})
    except Exception as exc:
        state.setdefault("logs", []).append({
            "ts": _time.time(), "event": "figma.fetch.failed", "error": str(exc)[:200]})

    extra_shots = figma_shots + app_shots
    if not urls and not extra_shots:
        return None

    # 没有 URL,只有 APP / Figma 图 → 直接返回,不启动浏览器
    if not urls:
        return extra_shots or None

    try:
        from playwright.async_api import async_playwright  # type: ignore
    except ImportError:
        state.setdefault("logs", []).append({
            "ts": _time.time(), "event": "screenshot.skip",
            "reason": "playwright not installed",
        })
        return extra_shots or None

    out_dir = Path(settings.evidence_output_dir) / "screenshots"
    out_dir.mkdir(parents=True, exist_ok=True)

    state["progress"] = f"截图准备：{len(urls)} URL × {len(viewports)} 视口…"
    captured: list[dict[str, Any]] = list(extra_shots)  # 设计基线 + APP 实拍图打头
    try:
        async with async_playwright() as pw:
            browser = await pw.chromium.launch(headless=True)
            try:
                for url in urls:
                    for vp_label, w, h in viewports:
                        state["progress"] = f"截图 {url[:60]} @ {vp_label}"
                        safe = _re.sub(r"[^\w]", "_", url)[:40]
                        fname = f"{tool_id}_{ctx.run_id[:8]}_{safe}_{w}x{h}.png"
                        fpath = out_dir / fname
                        try:
                            page = await browser.new_page(viewport={"width": w, "height": h})
                            await page.goto(url, timeout=30000, wait_until="domcontentloaded")
                            try:
                                await page.wait_for_load_state("networkidle", timeout=8000)
                            except Exception:
                                pass  # tolerate slow networks; capture what we have
                            await page.screenshot(path=str(fpath), full_page=True)
                            await page.close()
                            captured.append({
                                "url": url, "viewport": vp_label,
                                "width": w, "height": h,
                                "filename": fname,
                                "size": fpath.stat().st_size,
                            })
                        except Exception as exc:
                            captured.append({
                                "url": url, "viewport": vp_label,
                                "width": w, "height": h,
                                "error": str(exc)[:200],
                            })
            finally:
                await browser.close()
    except Exception as exc:
        state.setdefault("logs", []).append({
            "ts": _time.time(), "event": "screenshot.failed",
            "error": str(exc)[:200],
        })
        return captured or None

    return captured or None


_VALID_SEVERITIES = {"critical", "high", "medium", "low", "info"}
_SEVERITY_ALIASES = {
    "blocker":    "high",
    "major":      "high",
    "minor":      "low",
    "suggestion": "low",
    "trivial":    "info",
    "warning":    "medium",
    "warn":       "medium",
    "error":      "high",
    "fatal":      "critical",
}


def _normalize_severity(value: Any) -> str:
    """把任意 severity 值规范化为白名单内的 5 个值之一。

    防御 XSS：模型输出 / 报告导入可能包含 `high" onclick="alert(1)` 这种
    破坏 HTML 属性的串。这里把所有值都映射到固定枚举，确保 class 拼接安全。
    """
    s = str(value or "").lower().strip()
    if s in _VALID_SEVERITIES:
        return s
    if s in _SEVERITY_ALIASES:
        return _SEVERITY_ALIASES[s]
    return "medium"


def _build_executive_summary(report: dict[str, Any], tool: dict[str, Any]) -> dict[str, Any]:
    """Aggregate any existing report into the 5-section structure (统一报告契约):
       测试结论 (含 verdict_summary + KPI) / 风险结论 / 阻碍 / Bug 表 (issues 按 sev×pri 排序) / 执行用例记录.

    优先读 report 顶层契约字段(verdict/verdict_summary/risks/blockers/issues/cases);
    缺失则走 walk-substeps 兜底。"""

    sev_rank = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    pri_rank = {"P0": 0, "P1": 1, "P2": 2, "P3": 3}

    walked_issues: list[dict[str, Any]] = []
    walked_cases: list[dict[str, Any]] = []

    def _walk(x: Any) -> None:
        if isinstance(x, dict):
            if any(k in x for k in ("expected", "scenario")) and "id" in x:
                walked_cases.append(x)
            elif "severity" in x and any(k in x for k in ("issue", "title", "description", "name")):
                walked_issues.append(x)
            for v in x.values():
                _walk(v)
        elif isinstance(x, list):
            for v in x:
                _walk(v)

    _walk(report.get("substeps") or {})

    top_issues = report.get("issues") if isinstance(report.get("issues"), list) else walked_issues
    top_cases = report.get("cases") if isinstance(report.get("cases"), list) else walked_cases
    top_risks = report.get("risks") if isinstance(report.get("risks"), list) else []
    top_blockers = report.get("blockers") if isinstance(report.get("blockers"), list) else []

    sev_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    pri_counts = {"P0": 0, "P1": 0, "P2": 0, "P3": 0, "其他": 0}

    natural_issues: list[dict[str, Any]] = []
    for it in top_issues:
        sev = _normalize_severity(it.get("severity"))
        pri = str(it.get("priority") or "P2").upper()
        sev_counts[sev] = sev_counts.get(sev, 0) + 1
        title = (it.get("title") or it.get("name") or it.get("issue")
                 or it.get("description") or "未命名问题")
        module = (it.get("module") or it.get("endpoint") or it.get("file_path")
                  or it.get("location") or it.get("viewport") or it.get("page")
                  or it.get("viewport_filename") or "")
        repro = it.get("reproduce_steps") or []
        if isinstance(repro, str):
            repro = [repro]
        rcases = it.get("related_test_cases") or []
        if isinstance(rcases, str):
            rcases = [rcases]
        natural_issues.append({
            "issue_id": str(it.get("issue_id") or it.get("id") or "")[:50],
            "title": str(title)[:200],
            "severity": sev,
            "priority": pri,
            "module": str(module)[:200],
            "current": str(it.get("current_behavior") or it.get("current") or it.get("observed")
                           or it.get("description") or it.get("issue") or "")[:800],
            "expected": str(it.get("expected_behavior") or it.get("expected")
                            or it.get("requirement") or "")[:800],
            "fix": str(it.get("fix_suggestion") or it.get("fix") or it.get("recommendation")
                       or it.get("suggestion") or it.get("remediation") or "")[:1000],
            "repro": [str(s)[:200] for s in repro if s][:10],
            "accept": str(it.get("acceptance_criteria") or it.get("acceptance")
                          or it.get("verify") or "")[:600],
            "cases": [str(c)[:50] for c in rcases if c][:10],
            "owner": str(it.get("owner_role") or it.get("owner") or "").lower()[:30],
            "hours": it.get("estimated_hours") or it.get("effort"),
            "impact": str(it.get("impact_scope") or it.get("impact") or "")[:400],
            "evidence": str(it.get("evidence") or it.get("source") or "")[:400],
        })
    natural_issues.sort(key=lambda x: (sev_rank.get(x["severity"], 9), pri_rank.get(x["priority"], 9)))
    natural_issues = natural_issues[:60]

    natural_cases: list[dict[str, Any]] = []
    for c in top_cases:
        pri = str(c.get("priority") or "P2").upper()
        if pri in pri_counts:
            pri_counts[pri] += 1
        elif pri:
            pri_counts["其他"] += 1
        natural_cases.append({
            "id": str(c.get("id") or c.get("case_id") or c.get("tc_id") or "")[:50],
            "title": str(c.get("title") or c.get("name") or c.get("scenario") or "")[:200],
            "priority": pri,
            "type": str(c.get("type") or c.get("kind") or "").lower()[:30],
            "status": str(c.get("status") or "designed").lower()[:30],
            "automation": str(c.get("automation_tag") or c.get("automation") or "").lower()[:30],
            "expected": str(c.get("expected") or c.get("expected_result") or "")[:400],
        })
    natural_cases.sort(key=lambda x: pri_rank.get(x["priority"], 9))
    natural_cases = natural_cases[:200]

    natural_risks: list[dict[str, Any]] = []
    for r in top_risks:
        if isinstance(r, str):
            natural_risks.append({"title": r, "impact": "", "why": "", "severity": "medium"})
        elif isinstance(r, dict):
            natural_risks.append({
                "id": str(r.get("id") or "")[:30],
                "title": str(r.get("title") or r.get("name") or r.get("risk") or "未命名风险")[:200],
                "impact": str(r.get("impact") or r.get("affects") or "")[:400],
                "why": str(r.get("why") or r.get("reason") or r.get("detail") or "")[:400],
                "severity": _normalize_severity(r.get("severity") or "medium"),
            })
    if not natural_risks:
        gate = report.get("gate_decision") or {}
        for r in (gate.get("reasons") or []):
            if r:
                natural_risks.append({"title": str(r)[:200], "impact": "", "why": "", "severity": "medium"})

    natural_blockers: list[dict[str, Any]] = []
    for b in top_blockers:
        if isinstance(b, dict):
            natural_blockers.append({
                "id": str(b.get("id") or "")[:30],
                "title": str(b.get("title") or b.get("name") or "未命名阻碍")[:200],
                "why_blocking": str(b.get("why_blocking") or b.get("reason") or b.get("why") or "")[:500],
                "what_to_unblock": str(b.get("what_to_unblock") or b.get("action") or b.get("fix") or "")[:500],
                "owner_role": str(b.get("owner_role") or b.get("owner") or "").lower()[:30],
                "hours": b.get("estimated_hours") or b.get("effort"),
            })

    # verdict / verdict_summary 优先用 report 顶层
    verdict_text = report.get("verdict")
    verdict_summary = str(report.get("verdict_summary") or "")[:300]
    if verdict_text:
        v = str(verdict_text)
        if "不通过" in v:
            verdict, verdict_class = v, "fail"
        elif "有条件" in v or "警告" in v or "部分" in v:
            verdict, verdict_class = v, "warn"
        else:
            verdict, verdict_class = v, "pass"
    else:
        gate = report.get("gate_decision") or {}
        action = (gate.get("action") or "").lower()
        if "reject" in action or sev_counts["critical"] > 0:
            verdict, verdict_class = "不通过", "fail"
        elif "warn" in action or sev_counts["high"] > 2 or natural_blockers:
            verdict, verdict_class = "有条件通过", "warn"
        elif action == "" and not natural_issues and not natural_cases:
            verdict, verdict_class = "未产出", "skip"
        else:
            verdict, verdict_class = "通过", "pass"

    return {
        "测试结论": {"判定": verdict, "level": verdict_class, "agent": tool.get("name", ""), "summary": verdict_summary},
        "风险结论": natural_risks,
        "阻碍": natural_blockers,
        "问题描述": natural_issues,
        "用例执行": {
            "总数": len(natural_cases),
            "列表": natural_cases,
            "按优先级": pri_counts,
            "说明": "Agent 负责用例设计与方案输出；实际执行需由测试团队 / CI 落地。",
        },
        "严重度分布": sev_counts,
    }


# 各工具的 Excel 差异化配置(标题 / 问题表名 / 编号前缀 / 业务语言 / 用例主表 / H5矩阵)
_EXEC_CFG: dict[str, dict[str, Any]] = {
    "step1":    {"title": "需求评审报告",          "issue_sheet": "需求问题清单", "issue_prefix": "REQ", "biz": True},
    "step2":    {"title": "测试用例设计报告",       "issue_sheet": "设计问题",     "issue_prefix": "TC",  "cases_primary": True},
    "step4":    {"title": "接口测试报告",          "issue_sheet": "接口问题清单", "issue_prefix": "API"},
    "step6":    {"title": "Agent 自动化执行报告",   "issue_sheet": "问题清单",     "issue_prefix": "AG"},
    "h5_adapt": {"title": "H5 适配初审报告",        "issue_sheet": "适配问题清单", "issue_prefix": "H5",  "h5_matrix": True},
}


def _build_seo_template_xlsx(report: dict[str, Any], meta: dict[str, Any]) -> bytes:
    """SEO 报告按用户 Lark 模板:单 sheet,全站总览 + 按页面类型分节(三段头 + 检查清单,状态色标)。"""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook(); ws = wb.active; ws.title = "SEO相关测试"
    for c, w in {"A": 16, "B": 22, "C": 52, "D": 10, "E": 52}.items():
        ws.column_dimensions[c].width = w
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    HDR = PatternFill("solid", fgColor="D6E4FF"); SEC = PatternFill("solid", fgColor="EEF3FF")
    STATUS = {"通过": ("107C10", "DFF6DD"), "警告": ("9A6700", "FFF4CE"), "不通过": ("C42B1C", "FDE7E9")}
    bold = Font(bold=True)
    row = 1
    title = ws.cell(row=row, column=1, value=f"SEO 审计报告（Google 标准）· {meta.get('project_name') or ''} {meta.get('project_code') or ''}".strip())
    title.font = Font(bold=True, size=13); row += 1
    vmap = {"通过": "✅ 通过", "有条件通过": "⚠️ 有条件通过", "不通过": "❌ 不通过"}
    vc = ws.cell(row=row, column=1, value=f"判定：{vmap.get(report.get('verdict'), report.get('verdict') or '')}　|　{report.get('verdict_summary') or ''}")
    vc.alignment = wrap; ws.merge_cells(f"A{row}:E{row}"); ws.row_dimensions[row].height = 36; row += 2
    # 全站总览
    ov = report.get("overview") or {}
    site = [f"综合评定：{ov.get('评定','')}" if ov.get("评定") else "",
            f"核心问题：{ov.get('核心问题','')}" if ov.get("核心问题") else "",
            f"后续动作：{ov.get('后续动作','')}" if ov.get("后续动作") else ""]
    site = [x for x in site if x]
    if site:
        ws.merge_cells(f"A{row}:E{row}")
        cell = ws.cell(row=row, column=1, value="【全站总览】\n" + "\n".join(site))
        cell.alignment = wrap; cell.fill = HDR; cell.font = bold
        ws.row_dimensions[row].height = 28 * (len(site) + 1); row += 2
    # 逐页面类型
    for pt in report.get("page_type_audits") or []:
        ws.merge_cells(f"A{row}:E{row}")
        sc = ws.cell(row=row, column=1, value=f"【{pt.get('page_type','页面')}】"
                     + (f"　(样本 {pt.get('pages_sampled')} 页)" if pt.get("pages_sampled") else ""))
        sc.font = Font(bold=True, size=12, color="1A5FB4"); sc.fill = SEC; row += 1
        po = pt.get("overview") or {}
        ovl = [f"{k}：{po[k]}" for k in ("综合评定", "核心优势", "改进空间") if po.get(k)]
        if ovl:
            ws.merge_cells(f"A{row}:E{row}")
            oc = ws.cell(row=row, column=1, value="\n".join(ovl)); oc.alignment = wrap; oc.fill = SEC
            ws.row_dimensions[row].height = 24 * (len(ovl) + 1); row += 1
        for ci, h in enumerate(["检查维度", "检查项", "测试详情", "状态", "优化建议"], 1):
            hc = ws.cell(row=row, column=ci, value=h); hc.font = bold; hc.fill = HDR
            hc.alignment = center if ci == 4 else wrap; hc.border = border
        row += 1
        last_dim = None
        for it in pt.get("checklist") or []:
            dim = it.get("维度", "")
            ws.cell(row=row, column=1, value=(dim if dim != last_dim else "")).alignment = wrap
            if dim != last_dim:
                ws.cell(row=row, column=1).font = bold
            last_dim = dim
            ws.cell(row=row, column=2, value=it.get("检查项", "")).alignment = wrap
            ws.cell(row=row, column=2).font = bold
            ws.cell(row=row, column=3, value=it.get("测试详情", "")).alignment = wrap
            st = (it.get("状态") or "").strip()
            stc = ws.cell(row=row, column=4, value=st); stc.alignment = center
            if st in STATUS:
                fg, bg = STATUS[st]; stc.font = Font(bold=True, color=fg); stc.fill = PatternFill("solid", fgColor=bg)
            ws.cell(row=row, column=5, value=it.get("优化建议", "")).alignment = wrap
            for ci in range(1, 6):
                ws.cell(row=row, column=ci).border = border
            row += 1
        row += 1
    # 站外/外链等需外部数据(risks)单列附后
    rks = [rk for rk in (report.get("risks") or []) if rk.get("title")]
    if rks:
        ws.merge_cells(f"A{row}:E{row}")
        ws.cell(row=row, column=1, value="【需外部数据补验（GSC / Ahrefs 等，本工具爬站测不到）】").font = bold
        ws.cell(row=row, column=1).fill = SEC; row += 1
        for rk in rks[:12]:
            ws.cell(row=row, column=1, value=f"· {rk.get('title')}").alignment = wrap
            ws.merge_cells(f"A{row}:E{row}"); row += 1
    bio = io.BytesIO(); wb.save(bio); return bio.getvalue()


def _build_network_template_xlsx(report: dict[str, Any], meta: dict[str, Any]) -> bytes:
    """弱网/断网报告(按工具性质:档位×指标矩阵 + 容错checklist + 视频弱网checklist + 资损问题)。"""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook(); ws = wb.active; ws.title = "弱网断网测试"
    thin = Side(style="thin", color="D9D9D9"); bd = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top"); ctr = Alignment(horizontal="center", vertical="center", wrap_text=True)
    HDR = PatternFill("solid", fgColor="D6E4FF"); SEC = PatternFill("solid", fgColor="EEF3FF")
    STA = {"通过": ("107C10", "DFF6DD"), "警告": ("9A6700", "FFF4CE"), "不通过": ("C42B1C", "FDE7E9")}
    bold = Font(bold=True)
    for c, w in {"A": 16, "B": 14, "C": 40, "D": 11, "E": 46, "F": 12, "G": 11, "H": 14}.items():
        ws.column_dimensions[c].width = w
    row = [1]
    def sec(t):
        ws.merge_cells(f"A{row[0]}:E{row[0]}"); s = ws.cell(row=row[0], column=1, value=t)
        s.font = Font(bold=True, size=12, color="1A5FB4"); s.fill = SEC; row[0] += 1
    def checklist(items):
        for ci, h in enumerate(["维度", "检查项", "测试详情", "状态", "优化建议"], 1):
            hc = ws.cell(row=row[0], column=ci, value=h); hc.font = bold; hc.fill = HDR
            hc.alignment = ctr if ci == 4 else wrap; hc.border = bd
        row[0] += 1
        last = None
        for it in items:
            d = it.get("维度", "")
            c1 = ws.cell(row=row[0], column=1, value=(d if d != last else "")); c1.alignment = wrap
            if d != last: c1.font = bold
            last = d
            ws.cell(row=row[0], column=2, value=it.get("检查项", "")).alignment = wrap
            ws.cell(row=row[0], column=2).font = bold
            ws.cell(row=row[0], column=3, value=it.get("测试详情", "")).alignment = wrap
            st = (it.get("状态") or "").strip(); sc = ws.cell(row=row[0], column=4, value=st); sc.alignment = ctr
            if st in STA: fg, bg = STA[st]; sc.font = Font(bold=True, color=fg); sc.fill = PatternFill("solid", fgColor=bg)
            ws.cell(row=row[0], column=5, value=it.get("优化建议", "")).alignment = wrap
            for ci in range(1, 6): ws.cell(row=row[0], column=ci).border = bd
            row[0] += 1
        row[0] += 1
    ws.cell(row=row[0], column=1, value=f"弱网 / 断网容错报告 · {meta.get('project_name') or ''}".strip()).font = Font(bold=True, size=13); row[0] += 1
    vmap = {"通过": "✅ 通过", "有条件通过": "⚠️ 有条件通过", "不通过": "❌ 不通过"}
    vc = ws.cell(row=row[0], column=1, value=f"判定：{vmap.get(report.get('verdict'), report.get('verdict') or '')}　|　{report.get('verdict_summary') or ''}")
    vc.alignment = wrap; ws.merge_cells(f"A{row[0]}:E{row[0]}"); ws.row_dimensions[row[0]].height = 36; row[0] += 2
    ov = report.get("overview") or {}
    site = [f"{k}：{ov[k]}" for k in ("评定", "核心问题", "后续动作") if ov.get(k)]
    if site:
        ws.merge_cells(f"A{row[0]}:E{row[0]}"); c = ws.cell(row=row[0], column=1, value="【总览】\n" + "\n".join(site))
        c.alignment = wrap; c.fill = HDR; c.font = bold; ws.row_dimensions[row[0]].height = 26 * (len(site) + 1); row[0] += 2
    pm = report.get("profile_matrix") or []
    if pm:
        sec("① 弱网档位加载矩阵(各档位真实实测)")
        cols = ["档位", "可达", "加载ms", "首屏FCP", "加载态", "错误UI", "超时", "控制台错误"]
        for ci, h in enumerate(cols, 1):
            hc = ws.cell(row=row[0], column=ci, value=h); hc.font = bold; hc.fill = HDR; hc.alignment = ctr; hc.border = bd
        row[0] += 1
        for m in pm:
            for ci, k in enumerate(["档位", "可达", "加载ms", "FCP", "加载态", "错误UI", "超时", "控制台错误"], 1):
                cc = ws.cell(row=row[0], column=ci, value=m.get(k)); cc.alignment = (wrap if ci == 1 else ctr); cc.border = bd
            row[0] += 1
        row[0] += 1
    if report.get("fault_checklist"):
        sec("② 容错与用户提示 checklist(★含避免静默失败)"); checklist(report["fault_checklist"])
    if report.get("video_checklist"):
        sec("③ 视频播放弱网 checklist(起播/卡顿/ABR/断网/续播/CDN分片)"); checklist(report["video_checklist"])
    iss = report.get("issues") or []
    if iss:
        sec("④ 操作 / 写 / 资损 问题清单")
        for ci, h in enumerate(["编号", "严重度", "问题", "现状 / 影响", "修复建议"], 1):
            hc = ws.cell(row=row[0], column=ci, value=h); hc.font = bold; hc.fill = HDR; hc.alignment = ctr if ci in (1, 2) else wrap; hc.border = bd
        row[0] += 1
        for it in iss[:40]:
            ws.cell(row=row[0], column=1, value=it.get("issue_id", "")).alignment = ctr
            ws.cell(row=row[0], column=2, value=it.get("severity", "")).alignment = ctr
            ws.cell(row=row[0], column=3, value=it.get("title", "")).alignment = wrap
            ws.cell(row=row[0], column=4, value=((it.get("current_behavior") or "") + " / " + (it.get("impact_scope") or "")).strip(" /")).alignment = wrap
            ws.cell(row=row[0], column=5, value=it.get("fix_suggestion", "")).alignment = wrap
            for ci in range(1, 6): ws.cell(row=row[0], column=ci).border = bd
            row[0] += 1
    bio = io.BytesIO(); wb.save(bio); return bio.getvalue()


def _build_testreport_xlsx(report: dict[str, Any], meta: dict[str, Any], tool: dict[str, Any]) -> bytes:
    """通用测试报告(符合软件测试报告/用例规范):测试结论 + 缺陷明细清单 + 用例执行清单 + 风险待验。
    供编排器型工具(h5/step4/step6 等)用;按工具贴标题与缺陷『位置』列名。"""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    name = tool.get("name") or tool.get("id") or "测试"
    loc_label = {"h5_adapt": "端 / 视口 / 页面", "step4": "接口 / 端点", "step6": "用例 / 模块"}.get(tool.get("id"), "模块 / 位置")
    wb = Workbook(); ws = wb.active; ws.title = (name[:28] or "测试报告")
    thin = Side(style="thin", color="D9D9D9"); bd = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top"); ctr = Alignment(horizontal="center", vertical="center", wrap_text=True)
    HDR = PatternFill("solid", fgColor="D6E4FF"); SEC = PatternFill("solid", fgColor="EEF3FF")
    SEV = {"critical": ("C42B1C", "FDE7E9"), "high": ("9A6700", "FFF4CE"), "medium": ("1A5FB4", "EAF1FF"),
           "low": ("107C10", "DFF6DD"), "info": ("616161", "F0F0F0")}
    CST = {"executed_pass": ("107C10", "DFF6DD"), "executed_fail": ("C42B1C", "FDE7E9"),
           "blocked": ("9A6700", "FFF4CE"), "designed": ("616161", "F0F0F0"), "skipped": ("616161", "F0F0F0")}
    bold = Font(bold=True)
    for c, w in {"A": 16, "B": 12, "C": 30, "D": 44, "E": 44, "F": 30}.items():
        ws.column_dimensions[c].width = w
    row = [1]
    def sec(t):
        ws.merge_cells(f"A{row[0]}:F{row[0]}"); s = ws.cell(row=row[0], column=1, value=t)
        s.font = Font(bold=True, size=12, color="1A5FB4"); s.fill = SEC; row[0] += 1
    def hdr(cols):
        for ci, h in enumerate(cols, 1):
            hc = ws.cell(row=row[0], column=ci, value=h); hc.font = bold; hc.fill = HDR
            hc.alignment = ctr if ci <= 2 else wrap; hc.border = bd
        row[0] += 1
    ws.cell(row=row[0], column=1, value=f"{name} · 测试报告　{meta.get('project_name') or ''} {meta.get('project_code') or ''}".strip()).font = Font(bold=True, size=13); row[0] += 1
    vmap = {"通过": "✅ 通过", "有条件通过": "⚠️ 有条件通过", "不通过": "❌ 不通过"}
    vc = ws.cell(row=row[0], column=1, value=f"测试结论：{vmap.get(report.get('verdict'), report.get('verdict') or '')}　|　{report.get('verdict_summary') or ''}")
    vc.alignment = wrap; ws.merge_cells(f"A{row[0]}:F{row[0]}"); ws.row_dimensions[row[0]].height = 40; row[0] += 1
    iss = report.get("issues") or []
    cnt = {}
    for it in iss: cnt[it.get("severity", "info")] = cnt.get(it.get("severity", "info"), 0) + 1
    cases = report.get("cases") or []
    cpass = sum(1 for c in cases if c.get("status") == "executed_pass")
    cfail = sum(1 for c in cases if c.get("status") == "executed_fail")
    ws.merge_cells(f"A{row[0]}:F{row[0]}")
    stat = ws.cell(row=row[0], column=1, value=f"缺陷统计：致命 {cnt.get('critical',0)} · 高 {cnt.get('high',0)} · 中 {cnt.get('medium',0)} · 低 {cnt.get('low',0)}　|　用例：共 {len(cases)},通过 {cpass},失败 {cfail}")
    stat.alignment = wrap; stat.fill = SEC; row[0] += 2
    if iss:
        sec("一、缺陷明细清单")
        hdr(["缺陷编号", "严重度", loc_label, "问题描述 / 现状", "修复建议", "证据"])
        for it in iss[:120]:
            ws.cell(row=row[0], column=1, value=it.get("issue_id", "")).alignment = ctr
            sv = (it.get("severity") or "info"); sc = ws.cell(row=row[0], column=2, value=sv); sc.alignment = ctr
            if sv in SEV: fg, bg = SEV[sv]; sc.font = Font(bold=True, color=fg); sc.fill = PatternFill("solid", fgColor=bg)
            ws.cell(row=row[0], column=3, value=it.get("module", "")).alignment = wrap
            ws.cell(row=row[0], column=4, value=((it.get("title") or "") + "\n" + (it.get("current_behavior") or "")).strip()).alignment = wrap
            ws.cell(row=row[0], column=5, value=it.get("fix_suggestion", "")).alignment = wrap
            ws.cell(row=row[0], column=6, value=it.get("evidence", "")).alignment = wrap
            for ci in range(1, 7): ws.cell(row=row[0], column=ci).border = bd
            row[0] += 1
        row[0] += 1
    if cases:
        sec("二、测试用例执行清单")
        hdr(["用例编号", "优先级", "标题", "预期 / 步骤", "实际 / 证据", "状态"])
        for c in cases[:200]:
            ws.cell(row=row[0], column=1, value=c.get("id", "")).alignment = ctr
            ws.cell(row=row[0], column=2, value=c.get("priority", "")).alignment = ctr
            ws.cell(row=row[0], column=3, value=c.get("title", "")).alignment = wrap
            steps = c.get("steps"); steps = " → ".join(steps) if isinstance(steps, list) else (steps or "")
            ws.cell(row=row[0], column=4, value=((c.get("expected") or "") + ("\n步骤:" + steps if steps else "")).strip()).alignment = wrap
            ws.cell(row=row[0], column=5, value=c.get("evidence", "")).alignment = wrap
            stt = c.get("status", ""); sc = ws.cell(row=row[0], column=6, value=stt); sc.alignment = ctr
            if stt in CST: fg, bg = CST[stt]; sc.font = Font(bold=True, color=fg); sc.fill = PatternFill("solid", fgColor=bg)
            for ci in range(1, 7): ws.cell(row=row[0], column=ci).border = bd
            row[0] += 1
        row[0] += 1
    rks = report.get("risks") or []
    if rks:
        sec("三、风险与待验项(含真机/外部数据需补验)")
        hdr(["编号", "严重度", "风险", "影响", "说明 / 为何需补验", ""])
        for rk in rks[:40]:
            ws.cell(row=row[0], column=1, value=rk.get("id", "")).alignment = ctr
            ws.cell(row=row[0], column=2, value=rk.get("severity", "")).alignment = ctr
            ws.cell(row=row[0], column=3, value=rk.get("title", "")).alignment = wrap
            ws.cell(row=row[0], column=4, value=rk.get("impact", "")).alignment = wrap
            ws.cell(row=row[0], column=5, value=rk.get("why", "")).alignment = wrap
            for ci in range(1, 6): ws.cell(row=row[0], column=ci).border = bd
            row[0] += 1
    bio = io.BytesIO(); wb.save(bio); return bio.getvalue()


def _build_exec_xlsx(r: dict[str, Any], tool: dict[str, Any], report: dict[str, Any]) -> bytes:
    """统一执行报告 Excel:概览 + 问题清单 + 风险 + 阻碍 + 用例 + 截图(+ 按工具特化表)。
    seo_audit 且含 page_type_audits 时,走用户 Lark 模板的单 sheet 按页面类型检查清单格式。"""
    m = report.get("meta") or {}
    meta = {
        "run_id": r.get("run_id") or "",
        "produced_at": m.get("produced_at_utc") or m.get("produced_at") or "",
        "model": m.get("model") or m.get("model_label") or "",
        "project_name": m.get("project_name") or "",
        "project_code": m.get("project_code") or "",
    }
    if tool.get("id") == "seo_audit" and report.get("page_type_audits"):
        try:
            return _build_seo_template_xlsx(report, meta)
        except Exception:
            pass  # 渲染失败则回退通用格式
    if tool.get("id") == "network_resilience" and (
            report.get("profile_matrix") or report.get("fault_checklist") or report.get("video_checklist")):
        try:
            return _build_network_template_xlsx(report, meta)
        except Exception:
            pass
    # 编排器型测试工具:统一走符合测试报告/用例规范的单 sheet(测试结论+缺陷明细+用例执行+风险待验)
    if tool.get("id") in {"h5_adapt", "step4", "step6", "step1", "step2"}:
        try:
            return _build_testreport_xlsx(report, meta, tool)
        except Exception:
            pass
    from packages.reporting.exec_excel import build_exec_xlsx
    summary = _build_executive_summary(report, tool)
    return build_exec_xlsx(summary, report, tool, meta, _EXEC_CFG.get(tool.get("id"), {}))


@app.get("/api/reports/{run_id}/export.{fmt}")
async def api_report_export(run_id: str, fmt: str, request: Request) -> Any:
    """Server-side export — guarantees download dialog in webview.

    Replaces the JS Blob+<a download> path which can fail in pywebview.
    """
    if fmt not in ("json", "md", "html", "xlsx"):
        raise HTTPException(400, f"unsupported format: {fmt}")
    user = require_user(request)
    # Fetch report (memory or disk)
    if run_id in _RUNS:
        r = _RUNS[run_id]
        if not _user_can_see(user, r.get("owner_user_id")):
            raise HTTPException(403, "无权访问此报告")
        report = r.get("report") or {}
    else:
        out_dir = Path(settings.report_output_dir)
        candidate = next(out_dir.glob(f"*_{run_id}.json"), None) if out_dir.exists() else None
        if not candidate:
            raise HTTPException(404, f"report {run_id} not found")
        try:
            file_data = _json.loads(candidate.read_text(encoding="utf-8"))
        except Exception as exc:
            raise HTTPException(500, f"failed to read report: {exc}")
        owner_uid = (file_data.get("meta") or {}).get("owner_user_id")
        if not _user_can_see(user, owner_uid):
            raise HTTPException(403, "无权访问此报告")
        r = {
            "report": file_data,
            "run_id": run_id,
            "tool_id": _parse_tool_id_from_stem(candidate.stem, run_id),
        }
        report = file_data
    tool_id = r.get("tool_id", "?")
    tool = next((t for t in TOOL_CATALOG if t["id"] == tool_id), None) or {"id": tool_id, "name": tool_id, "icon": "?"}
    fname_base = f"{tool_id}_{run_id[:8]}"

    # HTML 报告已下线:任何工具的 export.html 一律重定向到 Excel(MD/JSON 仍可在预览页下载)。
    if fmt == "html":
        return RedirectResponse(f"/api/reports/{run_id}/export.xlsx", status_code=302)

    if fmt == "json":
        body = _json.dumps(report, ensure_ascii=False, indent=2)
        return Response(
            content=body, media_type="application/json; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{fname_base}.json"'},
        )
    if fmt == "md":
        body = _build_markdown_report(r, tool, report)
        return Response(
            content=body, media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{fname_base}.md"'},
        )
    if fmt == "xlsx":
        # 按工具分发到对应 Excel 生成器。seo/network 已改纯 agentic,新报告无结构化采集数据
        # → 走通用执行报告 Excel(issues/cases);仅历史报告(meta 仍带 seo_data/network_data)
        # 才渲染原 10-sheet 富 Excel,保证旧报告打开不退化。
        _rmeta = (report.get("meta") or {}) if isinstance(report, dict) else {}
        if tool_id == "seo_audit" and _rmeta.get("seo_data"):
            xlsx_bytes = _build_seo_xlsx(r, tool, report)
            suffix = "_SEO深度审计报告"
        elif tool_id == "network_resilience" and _rmeta.get("network_data"):
            xlsx_bytes = _build_network_xlsx(r, tool, report)
            suffix = "_弱网断网测试报告"
        else:
            xlsx_bytes = _build_exec_xlsx(r, tool, report)
            suffix = "_" + (tool.get("name") or "报告").replace(" ", "")
        # 文件名含中文 → 必须 RFC 5987 编码(HTTP 头只能 latin-1),
        # 同时给纯 ASCII 回退名,兼容老浏览器。
        from urllib.parse import quote as _q
        full_name = f"{fname_base}{suffix}.xlsx"
        disp = f"attachment; filename=\"{fname_base}.xlsx\"; filename*=UTF-8''{_q(full_name)}"
        return Response(
            content=xlsx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": disp},
        )
    # 兜底:未知格式 → Excel
    return RedirectResponse(f"/api/reports/{run_id}/export.xlsx", status_code=302)


# ══════════════════ SEO 深度审计:真实采集 → Excel ══════════════════
def _seo_baseline_path(host: str):
    d = settings.report_output_dir.parent.parent / "seo_baselines"
    d.mkdir(parents=True, exist_ok=True)
    safe = _re.sub(r"[^0-9a-zA-Z._-]", "_", host or "site")
    return d / f"{safe}.json"


def _load_seo_baseline(host: str) -> dict[str, Any] | None:
    try:
        p = _seo_baseline_path(host)
        if p.exists():
            return _json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        pass
    return None


def _save_seo_baseline(host: str, seo_data: dict[str, Any]) -> None:
    try:
        # 只存摘要(避免基线文件过大)
        _seo_baseline_path(host).write_text(
            _json.dumps({"summary": seo_data.get("summary", {}), "crawled_at": seo_data.get("crawled_at", "")},
                        ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _seo_data_to_prompt(sd: dict[str, Any]) -> str:
    s = sd.get("summary", {}); t = sd.get("tech", {}); sm = sd.get("sitemap", {})
    L = ["## 真实全站 SEO 采集结果(以下分析与问题清单必须基于这些真实数据,不得臆测)"]
    L.append(f"- 爬取 {s.get('pages_crawled')} 页(200 OK {s.get('pages_ok')}),最大深度 {s.get('max_depth')},死链 {s.get('dead_links')},孤儿页 {s.get('orphan_pages')}")
    L.append(f"- 内链 {s.get('internal_links')} 条,通用/空锚文本 {s.get('generic_anchor_pct')}%")
    L.append(f"- 标题重复 {s.get('title_dup')} 页,描述<50字 {s.get('desc_short')} 页,描述缺失 {s.get('desc_missing')} 页")
    L.append(f"- 图片Alt<80% {s.get('alt_below80')} 页,H1 i18n占位符泄漏 {s.get('h1_i18n_leak')} 页,EN页title含中文 {s.get('en_title_cn')} 页,JSON-LD缺失 {s.get('jsonld_missing')} 页")
    L.append(f"- 【Google·T】title缺失 {s.get('title_missing')} 页 / 过长>60字 {s.get('title_long')} 页 / 过短<10 {s.get('title_short')} 页(全站OK {s.get('pages_total_ok')} 页)")
    L.append(f"- 【Google·D】description重复 {s.get('desc_dup')} 页 / 过长>160字 {s.get('desc_long')} 页(K=meta keywords 在 Google 不计权,仅记录)")
    L.append(f"- 【Google·可索引】canonical异常 {s.get('canonical_issue_pages')} 页;noindex {s.get('pages_noindex')} 页" + (f"({', '.join((s.get('noindex_urls') or [])[:5])})" if s.get('pages_noindex') else ""))
    L.append(f"- 【Google·移动优先】无viewport {s.get('pages_no_viewport')} 页;【薄内容】正文<100词 {s.get('pages_thin')} 页;【视频富结果】含VideoObject {s.get('pages_with_video_jsonld')} 页")
    L.append(f"- 技术SEO:HTTP {t.get('http_version')},HTTPS={t.get('https')},HSTS={'有' if t.get('hsts') else '缺'},CSP={'有' if t.get('csp') else '缺'},nosniff={'有' if t.get('x_content_type_options') else '缺'},压缩={t.get('content_encoding') or '无'}")
    L.append(f"- sitemap:HTTP {sm.get('index_status')},{sm.get('url_count')} URL,含lastmod {sm.get('with_lastmod')},爬取覆盖 {sm.get('crawl_coverage_pct')}%,robots中Sitemap协议 {sm.get('protocol_in_robots')}")
    cwv = sd.get("cwv", {}) or {}
    cl = [f"{k}(LCP{v.get('lcp')}ms/CLS{v.get('cls')}/{v.get('verdict')})" for k, v in cwv.items() if isinstance(v, dict) and v.get("lcp")]
    if cl:
        L.append("- Core Web Vitals 实测:" + "; ".join(cl))
    if sd.get("jsonld_dist"):
        L.append("- JSON-LD 类型分布:" + ", ".join(f"{k}×{v}" for k, v in list(sd["jsonld_dist"].items())[:8]))
    for name, info in (sd.get("templates") or {}).items():
        top = ", ".join(f"{k}×{v}" for k, v in sorted(info.get("issues", {}).items(), key=lambda x: -x[1])[:3])
        L.append(f"  · 模板[{name}] {info['pages']}页 通过{info['pass']}/警告{info['warn']}/不通过{info['fail']} 主要问题:{top or '无'}")
    return "\n".join(L)


async def _seo_run_collect(ctx: Any, state: dict[str, Any]) -> None:
    """SEO 真实采集:全站 BFS 爬取 + 逐页解析 + Core Web Vitals 实测 → ctx/state。
    采集摘要注入 documents 供 LLM 综合分析;原始数据存 state 供 Excel 渲染。"""
    docs = (ctx.inputs or {}).get("documents") or ""
    m = _re.search(r"https?://[^\s\"'<>)]+", docs if isinstance(docs, str) else "")
    if not m:
        state.setdefault("logs", []).append({"ts": _time.time(), "event": "seo.skip", "reason": "材料无 URL"})
        return
    entry = m.group(0).rstrip(".,;)")
    from packages.core.seo import crawl_and_audit, measure_cwv
    state["progress"] = "SEO 全站爬取中…"
    data = await crawl_and_audit(
        entry, max_pages=300, max_depth=4,
        on_progress=lambda n, tot, u: state.update({"progress": f"SEO 爬取 {n}/{tot} 页…"}))
    # 每模板取一个代表页实测 CWV
    reps: dict[str, str] = {}
    for p in data.pages:
        if p.status == 200 and p.template not in reps:
            reps[p.template] = p.url
    state["progress"] = "实测 Core Web Vitals…"
    try:
        data.cwv = await measure_cwv(dict(list(reps.items())[:10]),
                                     on_progress=lambda tm, mm: state.update({"progress": f"CWV 实测:{tm}…"}))
    except Exception as exc:
        data.cwv = {"error": str(exc)[:150]}
    try:
        import datetime as _dt
        data.crawled_at = _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    except Exception:
        pass
    sd = data.to_dict()
    state["seo_data"] = sd
    state["seo_baseline"] = _load_seo_baseline(data.host)
    _save_seo_baseline(data.host, sd)
    ctx.inputs["documents"] = (docs or "") + "\n\n" + _seo_data_to_prompt(sd)
    state.setdefault("logs", []).append({
        "ts": _time.time(), "event": "seo.collected", "pages": len(data.pages),
        "cwv": len([k for k, v in data.cwv.items() if isinstance(v, dict)]) if isinstance(data.cwv, dict) else 0})


def _verdict_to_gate(verdict: str) -> str:
    """verdict → gate_decision.action 一致映射(统一报告契约)。"""
    return {
        "通过": "proceed",
        "有条件通过": "proceed_with_warning",
        "不通过": "reject_with_report",
    }.get((verdict or "").strip(), "proceed_with_warning")


async def _seo_synthesize(ctx: Any, state: dict[str, Any]) -> dict[str, Any]:
    """SEO:基于真实采集数据,1 次 LLM 调用产出审计结论(替代 5 子步骤,大幅提速)。"""
    sd = state.get("seo_data") or {}
    data_text = _seo_data_to_prompt(sd)
    system = (
        "你是资深 Google SEO 审计专家。下面是对某站点的【真实全站采集数据】"
        "(确定性爬取 + 逐页解析真实测得 + CWV 实测,不是猜测)。**严格按 Google 现行 SEO 标准"
        "(Google Search Essentials / Search Central)逐层穷尽分析**,产出可执行审计结论(合法 JSON)。"
        "本产品是**视频站**,VideoObject 结构化数据 + 视频可索引 与 T/D 元信息**并重**。必须挖到底。\n\n"
        "【Google 标准 · 8 层逐层核,禁止用『等/类似/若干』含糊带过,适用必列尽,每条引采集字段真值】\n"
        "① 可抓取与可索引(前提,进不来=0 SEO):robots.txt 是否误封 Googlebot;sitemap 存在性/URL 数/lastmod/与实际爬取覆盖率差;"
        "canonical(canonical_issue_pages:缺/跨站/跨语言/跨频道,参数页未归一);**noindex 页 pages_noindex(核心模板被 noindex=critical,逐条看 noindex_urls)**;"
        "死链 4xx/5xx 逐条 URL;3xx 重定向链是否过长;孤儿页;HTTP 状态。\n"
        "② 页面元信息 T / D(K 被 Google 忽略):"
        "T(title):缺失 title_missing / 全站重复 title_dup / 过长 title_long(SERP≈600px,中文全角约 28-30 字截断)/ 过短 title_short;关键词前置、与页面内容一致(否则 Google 改写)。"
        "D(description):重复 desc_dup / 过长 desc_long(SERP 摘要≈中文 70-80 字截断)/ 过短 desc_short / 缺失 desc_missing;非排名因素但决定 SERP 摘要→CTR。"
        "★K(meta keywords):**Google 官方明确不用 meta keywords 作排名信号**——meta_keywords 仅记录现状,**绝不开成 issue、不作排名问题**;堆砌大量无关词最多 info 级。\n"
        "③ 内容质量与 E-E-A-T:**薄内容 pages_thin(正文<100词)**——视频站尤其要有文字上下文(标题/简介/转写),否则 Google 读不到视频里讲了什么;"
        "H1(缺失/多个/i18n 占位泄漏 h1_i18n_leak);标题层级跳级 heading_skip;原创性/满足搜索意图(数据有限处标需人工/外部,不臆断)。\n"
        "④ 结构化数据(富结果·视频站命脉):**VideoObject 覆盖 pages_with_video_jsonld 对比视频详情模板页数——视频页若系统性无 VideoObject = 上不了 Google 视频富结果/视频标签页,视频站核心缺陷(high+)**;"
        "JSON-LD 缺失 jsonld_missing;类型分布;关键字段(thumbnailUrl/uploadDate/contentUrl/duration/SeekToAction)采集只到类型层,字段级标『需 Rich Results Test 深验』。\n"
        "⑤ 页面体验 Core Web Vitals(Google 排名信号):**仅当 cwv 有真实实测才下结论,无则 unknown 绝不编毫秒数**;阈值 LCP 好≤2500ms/需改进≤4000/差>4000,CLS 好≤0.1/≤0.25/差>0.25;"
        "**INP(2024 起取代 FID):实验室爬取无法测真实 INP,需 CrUX/GSC 现场数据 → 标 needs_external_data,不臆造**;"
        "**移动优先索引:无 viewport pages_no_viewport = 移动端不友好,Google 用移动版排名(high)**;HTTPS。\n"
        "⑥ 视频专项(你是视频站):视频可被抓(不被 robots 封/不藏登录后)、VideoObject、视频 sitemap、缩略图唯一——综合 ④⑤,视频不可索引或无 VideoObject 是头号缺陷。\n"
        "⑦ 安全与协议响应头:HTTPS/HSTS/CSP/X-Content-Type-Options/content-encoding(采集数据有的核)。\n"
        "⑧ 站外(外链 backlinks / 真实收录量 / 关键词排名):**爬自己站测不到,一律标 needs_external_data 进 risks(需 Google Search Console / Ahrefs),绝不臆测排名或收录数字**。\n\n"
        "【接地气 · 零幻觉】只列采集数据真实体现的问题;每条 issue 的 current_behavior / evidence "
        "必须引用采集数据里的**具体字段名与值**(如 `summary.dead_links=3`、`pages[].title_len`);"
        "数据没覆盖到的(真实收录量 / 竞品对比 / 关键词排名 / 外链)一律不臆测,需要时标 unknown。\n"
        "【自我复核】出结论前自问:还有哪些页面 / 模板 / 技术维度 / i18n 泄漏没核到?逐项补全再输出。\n\n"
        "【严重度与优先级判定 · Google 口径】severity:critical=核心模板被 noindex/不可索引 / 大面积死链 / 视频详情页系统性无 VideoObject 致视频不可被 Google 发现;"
        "high=title 大面积重复或缺失、canonical 错误、移动端无 viewport(移动优先)、CWV 实测为差、大面积薄内容、视频站关键页缺结构化数据;"
        "medium=次要优化(alt 覆盖偏低、锚文本通用、desc 过长截断);low/info=轻微提示(含 meta keywords 堆砌仅 info,因 Google 不计权)。"
        "priority 默认 critical→P0、high→P1、medium→P2、low/info→P3。"
        "verdict 与 gate_decision.action 一致映射:通过↔proceed、有条件通过↔proceed_with_warning、不通过↔reject_with_report。\n\n"
        "【★报告主体 = 按页面类型的检查清单 page_type_audits(务必产出,这是用户要的报告格式)】\n"
        "按采集数据里的【模板分类 templates】(如 首页 / 视频详情页 / 分类页 / 标签页 / 列表页 等,**爬到几类就出几节**),"
        "为每个页面类型产出:① overview(综合评定 2-3句该类页SEO总评、核心优势、改进空间按优先级);"
        "② checklist 逐检查项给 状态(通过/警告/不通过)+ 测试详情(引该类页采集真值)+ 优化建议。检查项按维度组织,"
        "以下为标准检查项集(适用该页面类型的必列;**模板原本没有但 Google 标准/视频站需要的 ★ 必须补上**):\n"
        "  · 基础元数据:标题(Title)、关键词(Keywords·注 Google 不计权仅记录)、描述(Description)\n"
        "  · 内容结构:H1标签、H2-H6标签、图片Alt属性、★页面原创性/内容充实度(薄内容 pages_thin)\n"
        "  · ★结构化数据(模板缺·必补):VideoObject(视频页命脉,系统性缺=不通过)、面包屑BreadcrumbList、其它JSON-LD\n"
        "  · 技术性能:HTTPS安全、移动端适配(viewport)、★Core Web Vitals(LCP/CLS实测给值+评级;INP需CrUX现场数据→unknown)、URL规范化\n"
        "  · 爬虫友好与★可索引(模板缺·必补):Robots.txt、站点地图Sitemap、★canonical规范(canonical_issue_pages)、★noindex可索引(pages_noindex)、内链/标签系统\n"
        "  · ★站外:外链/真实收录量/关键词排名 → 不在本工具范围,状态标 unknown 注『需 GSC/Ahrefs』\n"
        "状态判定:该检查项该页类全OK=通过;部分问题/可优化=警告;硬性缺陷(视频页系统性缺VideoObject、核心页noindex、CWV实测差、title/desc大面积重复、Sitemap 404)=不通过。每条测试详情引采集真值(如『该模板120页仅5页含VideoObject』『LCP实测4184ms>4000差』)。\n\n"
        "【输出 · 合法 JSON,以 { 开头,无前后缀、无 markdown 代码块,全部中文】\n"
        "{\n"
        '  "verdict": "通过|有条件通过|不通过",\n'
        '  "verdict_summary": "≤120字一句话核心结论",\n'
        '  "gate_decision": {"action":"proceed|proceed_with_warning|reject_with_report","reasons":["..."]},\n'
        '  "confidence": {"score": 0.0, "rationale": "基于采集覆盖率的保守自评"},\n'
        '  "overview": {"评定":"2-4句全站总体评定", "核心问题":"最关键的1-3个问题", "后续动作":"按优先级的下一步动作"},\n'
        '  "page_type_audits": [{"page_type":"首页|视频详情页|分类页…(来自templates分类)","pages_sampled":0,"overview":{"综合评定":"...","核心优势":"...","改进空间":"..."},"checklist":[{"维度":"基础元数据","检查项":"标题(Title)","测试详情":"引该页类采集真值","状态":"通过|警告|不通过","优化建议":"..."}]}],\n'
        '  "risks": [{"id":"R-001","title":"...","impact":"对收录/排名的影响","why":"基于哪条采集数据","severity":"critical|high|medium|low"}],\n'
        '  "issues": [{"issue_id":"SEO-AREA-0001(AREA∈IDX索引/META元信息/CONTENT内容/SCHEMA结构化/CWV体验/VIDEO视频/SEC安全)","priority":"P0|P1|P2|P3","severity":"critical|high|medium|low|info",'
        '"module":"Google层(①可索引/②T-D元信息/③内容EEAT/④结构化数据/⑤页面体验CWV/⑥视频专项/⑦安全)",'
        '"title":"一句话问题","current_behavior":"实测现状(引采集字段与值)","expected_behavior":"应做到什么",'
        '"impact_scope":"对 SEO/收录/排名的影响","fix_suggestion":"具体怎么修(改哪个文件/标签/响应头/字段)",'
        '"evidence":"采集数据具体字段名:值"}],\n'
        '  "strengths": ["已验证的优点(基于通过项,引采集数据)"]\n'
        "}\n"
        "page_type_audits 必须覆盖采集到的每个页面类型;issues 按 (severity, priority) 双键排序;空数组写 []。"
    )
    state["progress"] = "AI 综合分析 SEO 采集结果…"
    # 超时+重试保护:综合分析 LLM 调用偶发长时间卡住(>20min);每次最多 360s,失败重试一次,
    # 两次都超时则回退最小报告(不挂死)。page_type_audits 输出较大 → max_tokens 提到 32000。
    import asyncio as _aio
    resp = None
    for _att in range(2):
        try:
            resp = await _aio.wait_for(
                ctx.llm.complete(
                    system=system, messages=[{"role": "user", "content": data_text}],
                    max_tokens=20000, allow_degrade=(_att > 0)),
                timeout=600)
            break
        except _aio.TimeoutError:
            state.setdefault("logs", []).append({
                "ts": _time.time(), "event": "seo.synth.timeout", "attempt": _att})
            state["progress"] = f"综合分析超时,重试({_att + 1}/2)…"
        except Exception as _e:
            state.setdefault("logs", []).append({
                "ts": _time.time(), "event": "seo.synth.error", "error": str(_e)[:160]})
            break
    if resp is None:
        # 兜底:两次超时/出错 → 返回基于采集数据的最小报告,保证 run 完成而非挂死
        return {
            "verdict": "有条件通过",
            "verdict_summary": "综合分析 LLM 调用超时,本报告仅含采集事实摘要;请重跑以获取完整 Google 标准分析。",
            "gate_decision": {"action": "proceed_with_warning", "reasons": ["synthesize 超时,未生成完整分析"]},
            "confidence": {"score": 0.2, "rationale": "synthesize 超时,仅采集数据可用"},
            "overview": {"评定": "采集完成但综合分析超时", "核心问题": "需重跑", "后续动作": "重跑 SEO 审计"},
            "page_type_audits": [], "issues": [], "risks": [],
            "cases": [], "strengths": _seo_strengths(sd), "meta": {},
        }
    try:
        ctx.usage.merge(resp.usage)
    except Exception:
        pass
    parsed = {}
    try:
        parsed = resp.json() or {}
    except Exception:
        parsed = {}
    verdict = parsed.get("verdict") or "有条件通过"
    return {
        "verdict": verdict,
        "verdict_summary": parsed.get("verdict_summary") or "(见各 sheet 明细)",
        "gate_decision": parsed.get("gate_decision") or {"action": _verdict_to_gate(verdict), "reasons": []},
        "confidence": parsed.get("confidence") or {},
        "overview": parsed.get("overview") or {},
        "page_type_audits": parsed.get("page_type_audits") or [],
        "issues": parsed.get("issues") or [],
        "risks": parsed.get("risks") or [],
        "cases": [],
        "strengths": parsed.get("strengths") or _seo_strengths(sd),
        "meta": {},
    }


def _seo_strengths(sd: dict[str, Any]) -> list[str]:
    out: list[str] = []
    t = sd.get("tech", {}); s = sd.get("summary", {}); sm = sd.get("sitemap", {})
    if t.get("https"):
        out.append("全站 HTTPS")
    if t.get("hsts"):
        out.append("已配置 HSTS")
    if t.get("http_version") and "2" in str(t.get("http_version")) + str(t.get("alt_svc")):
        out.append(f"HTTP {t.get('http_version')}" + ("+HTTP/3 alt-svc" if t.get("alt_svc") else ""))
    if s.get("dead_links") == 0:
        out.append(f"0 死链、{s.get('orphan_pages', 0)} 孤儿页,内链结构健康")
    if s.get("generic_anchor_pct", 100) < 5:
        out.append(f"内链锚文本通用/空占比仅 {s.get('generic_anchor_pct')}%")
    if sd.get("jsonld_dist"):
        out.append(f"JSON-LD 全站铺开(共 {sum(sd['jsonld_dist'].values())} 项,{len(sd['jsonld_dist'])} 种 @type)")
    if sm.get("index_status") == 200 and sm.get("url_count", 0) > 0:
        out.append(f"sitemap 正常({sm.get('url_count'):,} URL,含 lastmod {sm.get('with_lastmod', 0)})")
    return out


def _build_step5_images(ctx: Any) -> list[dict[str, Any]] | None:
    """把 ctx.screenshots(设计稿 + 实拍)转成带 role caption 的 images,供逐帧比对。
    设计稿排前、实拍排后,便于配对;限量防请求过大。"""
    shots = getattr(ctx, "screenshots", None) or []
    if not shots:
        return None
    sc_dir = Path(settings.evidence_output_dir) / "screenshots"
    images: list[dict[str, Any]] = []
    for s in shots:
        if s.get("error") or not s.get("filename"):
            continue
        p = sc_dir / s["filename"]
        if not p.exists():
            continue
        fn = str(s.get("filename", "")); u = str(s.get("url", "") or "")
        is_design = bool(s.get("is_design") or "figma" in fn.lower() or "figma.com" in u.lower())
        role = ("设计稿(Figma 设计基线 — 目标设计,勿在此图上标问题)" if is_design
                else "实拍(APP/Web 实际界面 — 在此图对照设计稿标偏差)")
        images.append({
            "path": p, "mime": "image/png", "_is_design": is_design,
            "caption": (f"role={role} | viewport_filename={s['filename']} | "
                        f"viewport={s.get('viewport', '?')} | url={u}"),
        })
    # 设计稿排前、实拍排后(用显式字段,不靠 caption 文本);保留全部实拍 + 尽量多设计帧
    images.sort(key=lambda im: 0 if im["_is_design"] else 1)
    return images[:30] or None


def _step5_build_montage(designs: list[dict[str, Any]], out_path: str, cols: int = 8) -> str | None:
    """把设计帧拼成带编号(#1..#N)的网格图,供「视觉配对」——设计帧无有效名字(如导出帧 screen1..N)时用。
    8 列网格;成图任一边超 1900px 则等比缩到 1900(避开多模态 2000px 上限)。"""
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None
    cw, ch, lab, gap = 140, 300, 22, 6
    items = []
    for i, d in enumerate(designs):
        try:
            items.append((i + 1, Image.open(d["path"]).convert("RGB").resize((cw, ch))))
        except Exception:
            continue
    if not items:
        return None
    rows = (len(items) + cols - 1) // cols
    canvas = Image.new("RGB", (cols * (cw + gap) + gap, rows * (ch + lab + gap) + gap), (24, 24, 28))
    draw = ImageDraw.Draw(canvas)
    for idx, (num, im) in enumerate(items):
        r, c = divmod(idx, cols)
        x = gap + c * (cw + gap); y = gap + r * (ch + lab + gap)
        draw.rectangle([x, y, x + cw, y + lab], fill=(245, 200, 0))
        draw.text((x + 4, y + 5), f"#{num}", fill=(0, 0, 0))
        canvas.paste(im, (x, y + lab))
    m = max(canvas.size)
    if m > 1900:
        s = 1900 / m
        canvas = canvas.resize((int(canvas.width * s), int(canvas.height * s)))
    try:
        canvas.save(out_path)
        return out_path
    except Exception:
        return None


async def _step5_synthesize(ctx: Any, state: dict[str, Any]) -> dict[str, Any]:
    """step5:逐帧设计 vs 实拍比对 —— 每个实拍配名字最匹配的设计帧,一对(2图)一次调用,
    多对并行(限并发5),替代 5 子步骤 + 单次塞28图,几分钟出全部差异。"""
    shots = getattr(ctx, "screenshots", None) or []
    sc_dir = Path(settings.evidence_output_dir) / "screenshots"
    designs: list[dict[str, Any]] = []
    actuals: list[dict[str, Any]] = []
    for s in shots:
        if s.get("error") or not s.get("filename"):
            continue
        p = sc_dir / s["filename"]
        if not p.exists():
            continue
        fn = str(s.get("filename", "")); u = str(s.get("url", "") or "")
        is_design = bool(s.get("is_design") or "figma" in fn.lower() or "figma.com" in u.lower())
        nm = (s.get("node_name") or s.get("viewport") or fn) if is_design else (s.get("viewport") or fn)
        item = {"path": p, "filename": s["filename"], "name": str(nm)}
        (designs if is_design else actuals).append(item)
    if not actuals:
        return {"verdict": "不通过", "verdict_summary": "未采集到实拍界面,无法比对(检查模拟器/APK运行)",
                "issues": [], "risks": [], "blockers": [], "cases": [], "meta": {}}

    _KW = ("漫画", "里番", "动漫", "AV", "搜索", "收藏", "追番", "追漫", "点赞", "我的",
           "详情", "列表", "阅读", "播放", "首页", "启动", "登录", "登陆", "个人中心", "设置")

    def _score(a, d):
        return sum(1 for k in _KW if k in a["name"] and k in d["name"])

    pairs: list[tuple] = []
    for a in actuals:
        best = max(designs, key=lambda d: _score(a, d)) if designs else None
        pairs.append((a, best if best and _score(a, best) > 0 else None))

    # 名字配对全军覆没(浏览器导出的设计帧名是 screen1..N,关键词永远 0 分)→ 改用「视觉配对」:
    # 把设计帧拼成带编号拼图,让 AI 逐张实拍判断对应哪个编号(不靠名字)。
    if designs and not any(d for _, d in pairs):
        state["progress"] = f"AI 视觉配对(设计帧无名字,拼图比对 {len(actuals)} 屏)…"
        montage = sc_dir / f"_pair_montage_{ctx.run_id[:8]}.png"
        if _step5_build_montage(designs, str(montage)):
            _msem = _asyncio.Semaphore(5)

            async def _match(a):
                async with _msem:
                    try:
                        resp = await ctx.llm.complete(
                            system=("你在做界面匹配。第一张是若干编号设计稿的拼图(每格左上角 #编号),第二张是一个实拍界面。"
                                    "判断该实拍对应拼图里哪一个编号(同一界面/同一功能页)。只输出 JSON:"
                                    "{\"match\": 编号数字} 或 {\"match\": null}(拿不准/无对应就 null)。"),
                            messages=[{"role": "user", "content": f"实拍界面『{a['name']}』对应拼图里哪个设计稿编号?"}],
                            images=[{"path": str(montage), "mime": "image/png", "caption": "设计稿拼图(每格#编号)"},
                                    {"path": a["path"], "mime": "image/png", "caption": f"实拍:{a['name']}"}],
                            max_tokens=200, allow_degrade=False)
                        _jm = resp.json()
                        m = _jm.get("match") if isinstance(_jm, dict) else None
                        try:
                            ctx.usage.merge(resp.usage)
                        except Exception:
                            pass
                        return designs[int(m) - 1] if isinstance(m, (int, float)) and 1 <= int(m) <= len(designs) else None
                    except Exception:
                        return None

            matched = await _asyncio.gather(*[_match(a) for a, _ in pairs])
            pairs = [(a, matched[i]) for i, (a, _d) in enumerate(pairs)]

    try:
        cmp_md = Path(__file__).resolve().parent.parent.parent / "configs" / "prompts" / "step5_ui" / "_compare_pair.md"
        sysp = cmp_md.read_text(encoding="utf-8")
    except Exception:
        sysp = "比对设计稿与实拍这一对,在实拍图上标出每处差异,输出 {issues:[...]}。"

    state["progress"] = f"AI 逐对比对(实拍 {len(actuals)} 屏 × 设计 {len(designs)} 帧,{len(pairs)} 对并行)…"
    _sem = _asyncio.Semaphore(5)

    async def _cmp(a, d):
        imgs = []
        if d:
            imgs.append({"path": d["path"], "mime": "image/png", "caption": f"role=设计稿(目标设计) | name={d['name']}"})
        imgs.append({"path": a["path"], "mime": "image/png",
                     "caption": f"role=实拍(实际界面) | viewport_filename={a['filename']} | name={a['name']}"})
        async with _sem:
            try:
                resp = await ctx.llm.complete(
                    system=sysp,
                    messages=[{"role": "user", "content": f"比对:设计『{d['name'] if d else '无对应设计帧'}』 vs 实拍『{a['name']}』。输出该对 issues JSON。"}],
                    images=imgs, max_tokens=2500, allow_degrade=False)
                try:
                    ctx.usage.merge(resp.usage)
                except Exception:
                    pass
                _j = resp.json()
                if isinstance(_j, list):
                    return {"issues": _j}        # LLM 直接返回数组 → 包成 {issues:[...]}
                return _j if isinstance(_j, dict) else {}
            except Exception:
                return {}

    results = await _asyncio.gather(*[_cmp(a, d) for a, d in pairs])
    all_issues: list = []
    pairs_checked: list = []
    for (a, d), res in zip(pairs, results):
        iss = (res.get("issues") if isinstance(res, dict)
               else (res if isinstance(res, list) else [])) or []
        iss = [it for it in iss if isinstance(it, dict)]
        for it in iss:
            it.setdefault("module", a["name"])
            it.setdefault("viewport_filename", a["filename"])
            it.setdefault("issue_id", f"UI-{len(all_issues) + len(iss):03d}")
        all_issues.extend(iss)
        pairs_checked.append({"design": d["name"] if d else "(无对应设计帧)", "actual": a["name"], "diff_count": len(iss)})

    sev_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    all_issues.sort(key=lambda i: sev_order.get(str(i.get("severity", "")).lower(), 5))
    has_hi = any(str(i.get("severity", "")).lower() in ("critical", "high") for i in all_issues)
    real = [i for i in all_issues if str(i.get("severity", "")).lower() != "info"]
    verdict = "不通过" if has_hi else ("有条件通过" if real else "通过")
    state.setdefault("logs", []).append({
        "ts": _time.time(), "event": "step5.compare",
        "design_frames": len(designs), "actual_screens": len(actuals),
        "pairs": len(pairs), "issues": len(all_issues)})
    return {
        "verdict": verdict,
        "verdict_summary": f"逐帧比对 {len(pairs)} 对(实拍{len(actuals)}/设计{len(designs)}),共 {len(real)} 处需关注的差异",
        "issues": all_issues, "pairs_checked": pairs_checked,
        "risks": [], "blockers": [], "cases": [], "meta": {},
    }


def _seo_synthesis(report: dict[str, Any], seo_data: dict[str, Any]) -> dict[str, Any]:
    """把 LLM 报告(overview/issues/strengths)+ 采集数据,组装成 Excel 需要的 synthesis。
    优先用 LLM 直接产出的 overview/strengths;缺失则从 verdict/issues/采集数据兜底派生。"""
    issues = report.get("issues") or []
    ov = report.get("overview") or {}
    if not (ov.get("评定") or ov.get("核心问题")):
        core = "；".join(str(i.get("title") or i.get("current_behavior") or "") for i in issues[:3] if (i.get("title") or i.get("current_behavior")))
        actions = "；".join(str(i.get("fix_suggestion") or "") for i in issues[:3] if i.get("fix_suggestion"))
        ov = {
            "评定": report.get("verdict_summary") or report.get("verdict") or "(见各 sheet 明细)",
            "核心问题": core or "(未识别出突出问题)",
            "后续动作": actions or "(无)",
        }
    return {
        "overview": ov,
        "issues": issues,
        "strengths": report.get("strengths") or _seo_strengths(seo_data),
    }


def _build_seo_xlsx(r: dict[str, Any], tool: dict[str, Any], report: dict[str, Any]) -> bytes:
    from packages.reporting.seo_excel import build_seo_xlsx
    meta = report.get("meta") or {}
    seo_data = meta.get("seo_data") or {}
    baseline = meta.get("seo_baseline")
    synthesis = _seo_synthesis(report, seo_data)
    host = seo_data.get("host") or ""
    m = {"site": seo_data.get("base_url", ""), "site_name": host or "站点",
         "run_id": r.get("run_id", ""), "model": meta.get("model", ""),
         "crawled_at": seo_data.get("crawled_at", ""), "baseline": baseline}
    return build_seo_xlsx(seo_data, synthesis, m)


def _build_network_xlsx(r: dict[str, Any], tool: dict[str, Any], report: dict[str, Any]) -> bytes:
    """弱网/断网测试 → Excel(多档位矩阵 + 行为 + 韧性 + 问题清单 + 标准)。"""
    from packages.reporting.network_excel import build_network_xlsx
    meta = report.get("meta") or {}
    nd = meta.get("network_data") or {}
    m = {"site": nd.get("url") or meta.get("target") or "", "run_id": r.get("run_id", ""),
         "model": meta.get("model", ""), "tested_at": nd.get("tested_at") or meta.get("tested_at", "")}
    return build_network_xlsx(report, nd, m)


def _netdata_to_prompt(nd: dict[str, Any]) -> str:
    L = ["## 真实弱网/断网多档位实测结果(结论与问题清单须基于这些真实数据)"]
    for p in nd.get("profiles", []):
        L.append(f"- [{p.get('profile')}] 到达={p.get('reached')} 状态={p.get('status')} "
                 f"加载={p.get('load_ms')}ms FCP={p.get('fcp_ms')}ms 可见{p.get('text_len')}字符 "
                 f"控制台错误{p.get('console_errors')} 加载态={p.get('has_spinner')} 错误提示={p.get('has_error_ui')} "
                 f"超时={p.get('timed_out')} → 判定 {p.get('verdict')}")
    rec = nd.get("recovery", {}) or {}
    L.append(f"- 韧性:在线加载={rec.get('online_ok')},断网重载错误提示={rec.get('offline_has_error_ui')}"
             f"(可见{rec.get('offline_text_len','?')}字符),恢复在线后={'已自动恢复' if rec.get('recovered') else '未自动恢复'}")
    return "\n".join(L)


def _net_matrix_from_nd(nd: dict[str, Any]) -> list[dict[str, Any]]:
    """弱网档位加载矩阵 = 纯脚本(基于 Phase1 CDP 真实实测,不依赖 LLM)。"""
    rows = []
    for p in nd.get("profiles", []) or []:
        reached = p.get("reached")
        rows.append({
            "档位": p.get("profile"),
            "可达": "是" if reached else "否",
            "加载ms": p.get("load_ms") if p.get("load_ms") is not None else "-",
            "FCP": p.get("fcp_ms") if p.get("fcp_ms") is not None else "-",
            "加载态": "有" if p.get("has_spinner") else ("-" if not reached else "无"),
            "错误UI": "有" if p.get("has_error_ui") else ("-" if reached else "无"),
            "超时": "是" if p.get("timed_out") else "否",
            "控制台错误": p.get("console_errors", 0),
        })
    return rows


async def _network_run_collect(ctx: Any, state: dict[str, Any]) -> None:
    """弱网/断网:Playwright CDP 多档位真实限速/断网实测 → ctx/state。"""
    docs = (ctx.inputs or {}).get("documents") or ""
    m = _re.search(r"https?://[^\s\"'<>)]+", docs if isinstance(docs, str) else "")
    if not m:
        state.setdefault("logs", []).append({"ts": _time.time(), "event": "net.skip", "reason": "材料无 URL"})
        return
    url = m.group(0).rstrip(".,;)")
    from packages.core.netprobe import probe_network
    shots_dir = Path(settings.evidence_output_dir) / "screenshots"
    state["progress"] = "弱网/断网多档位实测中…"
    data = await probe_network(
        url, shots_dir, name_prefix=f"net_{ctx.run_id[:8]}",
        on_progress=lambda pn, i, tot: state.update({"progress": f"网络档位实测 {i}/{tot}:{pn}…"}))
    try:
        import datetime as _dt
        data.tested_at = _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    except Exception:
        pass
    nd = data.to_dict()
    state["network_data"] = nd
    # 把各档位截图登记为报告截图(让 HTML/报告能看到弱网/断网实拍)
    shots = [{"filename": p["screenshot"], "url": f"net://{p['profile']}", "viewport": "390x844",
              "width": 390, "height": 844, "label": p["profile"]}
             for p in nd.get("profiles", []) if p.get("screenshot")]
    if shots:
        state["screenshots"] = (state.get("screenshots") or []) + shots
        ctx.screenshots = (getattr(ctx, "screenshots", None) or []) + shots
    ctx.inputs["documents"] = (docs or "") + "\n\n" + _netdata_to_prompt(nd)
    state.setdefault("logs", []).append({
        "ts": _time.time(), "event": "net.collected", "profiles": len(nd.get("profiles", []))})
    # Phase2:限速下驱动前端真实操作 —— 补「用户提示(C深)+ 操作/写/资损(D层)」
    try:
        await _network_drive_agentic(ctx, state, url)
    except Exception as exc:
        state.setdefault("logs", []).append({
            "ts": _time.time(), "event": "net.drive.failed", "error": str(exc)[:200]})


_NET_DRIVE_PROFILES: dict[str, dict[str, Any]] = {
    "online":  {"offline": False, "dl": -1, "ul": -1, "lat": 0},
    "4g":      {"offline": False, "dl": 4 * 1024 * 1024 // 8, "ul": 3 * 1024 * 1024 // 8, "lat": 20},
    "fast_3g": {"offline": False, "dl": 1600 * 1024 // 8, "ul": 750 * 1024 // 8, "lat": 150},
    "slow_3g": {"offline": False, "dl": 400 * 1024 // 8, "ul": 400 * 1024 // 8, "lat": 400},
    "2g":      {"offline": False, "dl": 256 * 1024 // 8, "ul": 256 * 1024 // 8, "lat": 800},
    "offline": {"offline": True, "dl": 0, "ul": 0, "lat": 0},
}


async def _network_drive_agentic(ctx: Any, state: dict[str, Any], url: str) -> None:
    """弱网 Phase2:在各限速档位下真驱动前端操作,观察【用户提示 + 操作结果 + 资损】(C深 + D层)。
    采集器(Phase1)只 load 首页测加载指标;本阶段补「操作/写/提示」层。注入 documents 供 synthesize 分析。
    """
    import asyncio as _aio
    from packages.core.agent import agent_loop
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        return
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        page = await browser.new_page(viewport={"width": 390, "height": 844})
        cdp = await page.context.new_cdp_session(page)
        await cdp.send("Network.enable")
        cur = {"profile": "online"}

        async def navigate(a: dict[str, Any]) -> str:
            tgt = a.get("url") or url
            try:
                r = await page.goto(tgt, timeout=40000, wait_until="domcontentloaded")
                await _aio.sleep(1.5)
                return f"已打开 {page.url} | HTTP {r.status if r else '?'} | 标题: {await page.title()}"
            except Exception as e:
                return f"打开失败(当前【{cur['profile']}】档): {type(e).__name__}: {str(e)[:150]}"

        async def set_network(a: dict[str, Any]) -> str:
            p = (a.get("profile") or "online").lower()
            cfg = _NET_DRIVE_PROFILES.get(p) or _NET_DRIVE_PROFILES["online"]
            cur["profile"] = p
            try:
                await cdp.send("Network.emulateNetworkConditions", {
                    "offline": cfg["offline"], "latency": cfg["lat"],
                    "downloadThroughput": cfg["dl"], "uploadThroughput": cfg["ul"]})
                desc = "断网" if cfg["offline"] else ("下行%dB/s 延迟%dms" % (cfg["dl"], cfg["lat"]))
                return "已切到【%s】档位(%s)。现在做操作,观察这一档下的用户提示与结果。" % (p, desc)
            except Exception as e:
                return "切档失败: " + str(e)[:120]

        async def click(a: dict[str, Any]) -> str:
            try:
                if a.get("text"):
                    await page.get_by_text(a["text"], exact=False).first.click(timeout=10000)
                elif a.get("selector"):
                    await page.click(a["selector"], timeout=10000)
                else:
                    return "需提供 text 或 selector"
                await _aio.sleep(1.6)
                return f"已点击(【{cur['profile']}】档) | 当前 {page.url}"
            except Exception as e:
                return f"点击失败/无响应(【{cur['profile']}】档): {str(e)[:160]}"

        async def form_input(a: dict[str, Any]) -> str:
            try:
                val = str(a.get("value", ""))
                if a.get("placeholder"):
                    await page.get_by_placeholder(a["placeholder"]).first.fill(val, timeout=8000)
                    tgt = a["placeholder"]
                elif a.get("selector"):
                    await page.fill(a["selector"], val, timeout=8000)
                    tgt = a["selector"]
                else:
                    return "需提供 selector 或 placeholder"
                return f"已填 {tgt} = {val[:30]}"
            except Exception as e:
                return f"填写失败: {str(e)[:150]}"

        async def inspect(a: dict[str, Any]) -> str:
            # 重点抽：加载态/弹窗提示/错误文案/按钮 —— 用于判断「用户提示」质量
            try:
                sig = await page.evaluate(
                    "() => ({title:document.title,"
                    "bodyText:(document.body.innerText||'').slice(0,800),"
                    "loadingEls:[...document.querySelectorAll('[class*=load i],[class*=spin i],[class*=skeleton i]')].length,"
                    "toasts:[...document.querySelectorAll('[class*=toast i],[class*=dialog i],[class*=modal i],[class*=alert i],[role=alert]')].map(e=>(e.innerText||'').trim()).filter(Boolean).slice(0,8),"
                    "buttons:[...document.querySelectorAll('button,a')].map(e=>(e.innerText||'').trim()).filter(Boolean).slice(0,20)})")
                return _json.dumps(sig, ensure_ascii=False)[:2000]
            except Exception as e:
                return f"抽取失败: {str(e)[:150]}"

        tools = {"navigate": navigate, "set_network": set_network,
                 "click": click, "form_input": form_input, "inspect": inspect}
        ex_md = Path(__file__).resolve().parent.parent.parent / "configs" / "prompts" / "network_resilience" / "_execute.md"
        try:
            sysp = ex_md.read_text(encoding="utf-8")
        except Exception:
            sysp = ("你在各网络档位真驱动前端操作,观察【用户提示 + 操作结果 + 资损】。动作:set_network(profile)/"
                    "navigate/click/form_input/inspect。每轮输出 JSON {thought, action, args, finding, done}。"
                    "重点:慢网/断网下有无 loading/超时/错误/断网/恢复提示、操作能否完成、弱网超时重试是否重复提交。")
        task = (f"待测前端: {url}\n请在 online/4g/fast_3g/slow_3g/2g/offline 各档位下,"
                f"用 set_network 切档 + navigate/click/form_input 驱动关键操作(如搜索/进入详情/播放/登录/提交),"
                f"每档位 inspect 观察【用户提示】(loading/慢网提示/超时/错误文案/断网提示/恢复提示)与操作结果;"
                f"特别测:断网中操作、弱网超时后重试是否重复提交(资损)。只读为主,不点删除/支付等不可逆按钮。\n\n"
                f"参考材料:\n{(ctx.inputs.get('documents') or '')[:2500]}\n\n现在输出第一步动作的 JSON。")
        state["progress"] = "弱网下驱动操作实测中…"
        res = await agent_loop(
            ctx.llm, sysp, task, tools, max_steps=26,
            on_step=lambda r: state.update({"progress": f"弱网操作实测 第{r.get('step')}步: {r.get('action','')} [{cur['profile']}档]"}),
            files=getattr(ctx, "files", None))
        try:
            await browser.close()
        except Exception:
            pass

    lines = ["", "", "## 操作级弱网实测(Phase2:各档位驱动操作,观察用户提示 + 操作结果 + 资损)"]
    for t in res.get("transcript", []):
        a = t.get("args") or {}
        ident = a.get("profile", "") or a.get("text", "") or a.get("url", "") or a.get("selector", "")
        lines.append(f"- [{t.get('step')}] {t.get('action','')} {ident} → {str(t.get('result',''))[:240]}")
    if res.get("findings"):
        lines.append("\n弱网操作实测已标记的问题:")
        for f in res["findings"]:
            lines.append(f"- {_json.dumps(f, ensure_ascii=False)[:300]}")
    ctx.inputs["documents"] = (ctx.inputs.get("documents") or "") + "\n".join(lines)
    state.setdefault("logs", []).append({
        "ts": _time.time(), "event": "net.drive",
        "steps": res.get("steps"), "findings": len(res.get("findings") or [])})


async def _network_synthesize(ctx: Any, state: dict[str, Any]) -> dict[str, Any]:
    """弱网/断网:基于真实多档位实测数据(Phase1 加载矩阵 + Phase2 操作级实测),产出结论。"""
    # 用户在 UI 勾选测哪类弱网(run_options):页面/接口 与/或 视频播放,单选多选皆可
    mode_page_api = bool((ctx.inputs or {}).get("mode_page_api", True))
    mode_video = bool((ctx.inputs or {}).get("mode_video", False))
    _docs = ctx.inputs.get("documents") or ""
    nd = state.get("network_data") or {}
    data_text = _netdata_to_prompt(nd) if mode_page_api else ""
    if mode_page_api and "## 操作级弱网实测" in _docs:
        data_text += "\n\n" + _docs[_docs.index("## 操作级弱网实测"):][:6500]
    # 视频弱网证据(宿主机 net_video_collect.py 真 Chrome 实测,作为 documents 喂入):仅当用户勾选「视频弱网」才分析
    video_present = "# 视频播放弱网实测证据" in _docs
    has_video = mode_video and video_present
    if has_video:
        data_text += "\n\n" + _docs[_docs.index("# 视频播放弱网实测证据"):][:7500]
    if mode_video and not video_present:
        data_text += "\n\n[提示] 用户勾选了「视频播放弱网」但材料里没有『# 视频播放弱网实测证据』——请在 verdict_summary 提示:需先在宿主机运行 scripts/net_video_collect.py 采集视频证据并粘入材料。"
    if not mode_page_api and not has_video:  # 容错:都没有效 → 退回页面/接口口径,不产空报告
        mode_page_api = True
        data_text = _netdata_to_prompt(nd)
    system = (
        "你是资深弱网 / 容错测试专家。下面是对某前端的【真实实测数据】,含两~三部分:"
        "① Phase1 多档位首页加载矩阵(Playwright CDP 真实限速/断网测得:各档 load_ms/fcp_ms/spinner/error_ui/timed_out + 断网恢复);"
        "② Phase2 各档位下真驱动前端操作的记录(操作动作 + inspect 观察到的用户提示/弹窗文案/结果);"
        "③(若数据里含『# 视频播放弱网实测证据』)宿主机真 Chrome 对视频播放页各档位的弱网实测(标准 <video> 插桩 + CDN 分片网络,播放器无关:起播TTFF/卡顿waiting/自适应码率分辨率/seek/断网/恢复/分片失败)。\n"
        "只基于这些真实数据,**逐档位、逐操作、逐维度穷尽分析**,产出审计结论(合法 JSON)。弱网容错是细节中的细节,必须挖到底。\n\n"
        "【深度铁律 · 逐档位逐场景穷尽,禁止用『等/类似』含糊带过,适用必列尽】\n"
        "1) 逐网络档位(online / 4g / slow_3g / 2g / offline 凡实测到的)分别核:"
        "是否可达 reached、HTTP status、加载耗时 load_ms、首屏 fcp_ms、正文长度 text_len、资源数 resources、控制台报错 console_errors。"
        "对比各档位之间的劣化是否可接受(如 2g 下是否长时间白屏、是否超时无反馈)。\n"
        "2) 加载态反馈:弱网下是否有 loading / 骨架屏 / spinner(has_spinner);慢档位首屏前是否长时间空白无任何提示。\n"
        "3) 超时与错误处理:请求超时 timed_out 是否被捕获;是否有明确错误 UI(has_error_ui)还是静默白屏 / 卡死。\n"
        "4) 断网态(offline):是否给出明确『网络不可用』提示(offline_has_error_ui)还是空白(offline_text_len 过小);"
        "是否有缓存 / Service Worker 兜底内容。\n"
        "5) 断网恢复:恢复网络后页面是否自动重连 / 重新加载 / 数据恢复(recovered / recovered_text_len),还是需手动刷新、卡在错误态。\n"
        "6) 资损与幂等风险:弱网超时重试、断网重连可能导致的重复提交 / 重复扣款 / 数据不一致——**后端幂等页面看不到**,"
        "这类一律写入 risks(需后端确认),**不要凭空判定为已发现 issue**。\n"
        "7) 【用户提示层 · 基于 Phase2 操作记录逐档逐操作核】弱网体验核心,逐条审 11 点:"
        "①加载态提示(慢网有无 loading/骨架,不长时间白屏)②慢网分级提示(超时给『网络较慢正在加载』)"
        "③超时提示(给『加载超时请重试』不静默卡死)④断网提示(明确『网络已断开』不空白)"
        "⑤错误文案质量(说人话/准确区分超时vs服务错vs无网/带重试入口)⑥重试提示(有重试入口+反馈)"
        "⑦降级提示⑧恢复提示(『网络已恢复』+自动重连)⑨写操作弱网提示(进行中/成功/失败/防重复提交)"
        "⑩提示时机一致(不重复弹/不误报)⑪★避免静默失败(操作失败却无任何提示、用户以为成功——尤其支付/提交,弱网红线,最高严重度)。"
        "每条结论引 Phase2 里 inspect 观察到的具体提示文案/动作记录。\n"
        "8) 【操作/写/资损层 · 基于 Phase2】各档关键操作(搜索/详情/播放/登录/提交)能否完成;弱网超时后重试是否重复提交;"
        "断网中操作的处理。能从 Phase2 真实观察到的(如点击无反馈、重复提交)真断言;后端幂等/真实扣款看不到的入 risks。\n"
        + ("9) 【视频播放弱网层 · 数据含『# 视频播放弱网实测证据』必逐项穷尽核——本产品是视频站,这是弱网头号重灾区】"
           "基于宿主机真 Chrome 实测(标准 <video> 插桩 + CDN 分片网络,播放器无关),逐视频页逐档位核 8 项,每条引"
           "『视频页/档位/字段真值(tff/waiting/videoW×videoH/bufferAhead/segments.failed/_resumed/toasts)/截图名』:"
           "①起播:online 起播 TTFF 是否可接受;慢档(slow_3g/2g)起播是否超时黑屏且无封面/loading 提示。"
           "②弱网卡顿:各档 waiting(rebuffer)次数与缓冲余量 bufferAhead,越弱越卡是否在可接受范围、有无『缓冲中/网络不佳』提示(还是画面冻住无提示)。"
           "③自适应码率 ABR:分辨率 videoW×H 是否随网络变差自动降档(如 1080p→480p→320p)、网络恢复是否升回;不降码率硬卡 = 问题;ABR 是否生效要据分辨率轨迹判。"
           "④seek 续播:seek 到未缓冲位后弱网下能否续播(看其后 currentTime/buffered 与 waiting),还是卡死。"
           "⑤★断网断流(视频站头号红线,最高严重度):offline 档是否给『网络断开/重试』明确提示,还是黑屏/卡死/静默失败(看 toasts 为空 + loadingEls=0 + errCode + 截图);静默失败=critical。"
           "⑥恢复续播:recover 档是否自动续播且**断点续**(_resumed=true 且 currentTime 接续断网时位置),还是从头/卡死/需手动。"
           "⑦CDN 分片失败:各档 segments.failed 与 fail_samples——弱网下分片超时(ERR_ABORTED)、token/防盗链 403、断网 ERR_INTERNET_DISCONNECTED,会致花屏/黑屏/起播失败,逐档点名。"
           "⑧各档位综合劣化是否可接受 + 用户提示是否到位(贯穿 ⑤ 的避免静默失败)。"
           "视频弱网 issue 用 issue_id NET-VID-NNNN。证据没覆盖的(iOS 原生 HLS AVPlayer、真机弱网手感、切网流量提醒、付费/DRM 视频)一律标 needs_real_device 入 risks,不写成 issue。\n\n"
           if has_video else "")
        + "【接地气 · 零幻觉】只列实测数据真实体现的问题;每条 issue 的 current_behavior / evidence 必须引用"
        "**具体档位与实测字段值**(如 `profiles[slow_3g].load_ms`、`recovery.offline_has_error_ui=false`);"
        "实测没覆盖的(后端幂等、真实弱网下用户主观感受、真机弱网)一律不臆测,标 unknown 或入 risks。\n"
        "【自我复核】出结论前自问:还有哪些档位 / 场景(断网中操作、恢复时机、重复提交)没核到?逐项补全再输出。\n\n"
        "【严重度与优先级判定】severity:critical=断网或极弱网下核心页卡死 / 白屏无任何提示且不可恢复 / 有资损风险;"
        "high=弱网下长时间无加载反馈或超时无错误提示但可恢复;medium=次要体验(无骨架屏、文案不友好);low/info=轻微提示。"
        "priority 默认 critical→P0、high→P1、medium→P2、low/info→P3。"
        "verdict 与 gate_decision.action 一致映射:通过↔proceed、有条件通过↔proceed_with_warning、不通过↔reject_with_report。\n\n"
        "【输出 · 合法 JSON,以 { 开头,无前后缀、无 markdown 代码块,全部中文】\n"
        "{\n"
        '  "verdict": "通过|有条件通过|不通过",\n'
        '  "verdict_summary": "≤120字核心结论",\n'
        '  "gate_decision": {"action":"proceed|proceed_with_warning|reject_with_report","reasons":["..."]},\n'
        '  "confidence": {"score": 0.0, "rationale": "基于实测档位覆盖的保守自评"},\n'
        '  "overview": {"评定":"2-4句总体评定","核心问题":"最关键1-3个","后续动作":"按优先级下一步"},\n'
        '  "profile_matrix": [{"档位":"online|4g|fast_3g|slow_3g|2g|offline","可达":"是|否","加载ms":0,"FCP":0,"加载态":"有|无(长白屏)","错误UI":"有|无|-","超时":"是|否","控制台错误":0}],\n'
        '  "fault_checklist": [{"维度":"加载态|断网|恢复|超时|重试|提示","检查项":"如 慢网加载提示/断网提示/恢复提示/★避免静默失败","测试详情":"引档位实测真值","状态":"通过|警告|不通过","优化建议":"..."}],\n'
        '  "video_checklist": [{"维度":"起播|卡顿|ABR|seek|断网|恢复|CDN","检查项":"如 起播TTFF/弱网卡顿/ABR降码率/断网断流(静默失败)/断点续播/CDN分片失败","测试详情":"引视频实测真值","状态":"通过|警告|不通过","优化建议":"..."}],\n'
        '  "risks": [{"id":"R-001","title":"...","impact":"对用户/数据的影响","why":"基于哪条实测或为何需后端确认","severity":"critical|high|medium|low"}],\n'
        '  "issues": [{"issue_id":"NET-AREA-0001","priority":"P0|P1|P2|P3","severity":"critical|high|medium|low|info",'
        '"module":"档位/场景(如 断网态/弱网加载/恢复)",'
        '"title":"一句话问题","current_behavior":"实测现状(引档位与字段值)","expected_behavior":"应做到什么",'
        '"impact_scope":"对用户的影响","fix_suggestion":"具体怎么修(如接 Service Worker/骨架屏/超时重试/离线提示)",'
        '"evidence":"档位:字段名:值"}]\n'
        "}\n"
        "profile_matrix 覆盖每个实测到的档位;fault_checklist 覆盖用户提示各点(★避免静默失败必列);"
        "video_checklist 仅当数据含视频弱网证据时填,否则 []。issues 按 (severity, priority) 双键排序;空数组写 []。"
    )
    state["progress"] = "AI 综合分析弱网/断网实测结果…"
    resp = await ctx.llm.complete(
        system=system, messages=[{"role": "user", "content": data_text}],
        max_tokens=16000, allow_degrade=False)
    try:
        ctx.usage.merge(resp.usage)
    except Exception:
        pass
    parsed = {}
    try:
        parsed = resp.json() or {}
    except Exception:
        parsed = {}
    verdict = parsed.get("verdict") or "有条件通过"
    return {
        "verdict": verdict,
        "verdict_summary": parsed.get("verdict_summary") or "(见各 sheet 明细)",
        "gate_decision": parsed.get("gate_decision") or {"action": _verdict_to_gate(verdict), "reasons": []},
        "confidence": parsed.get("confidence") or {},
        "overview": parsed.get("overview") or {},
        "profile_matrix": _net_matrix_from_nd(nd) or parsed.get("profile_matrix") or [],  # 矩阵=脚本(实测),不靠LLM
        "fault_checklist": parsed.get("fault_checklist") or [],
        "video_checklist": parsed.get("video_checklist") or [],
        "issues": parsed.get("issues") or [],
        "risks": parsed.get("risks") or [],
        "cases": [],
        "meta": {},
    }


def _build_markdown_report(r: dict[str, Any], tool: dict[str, Any], report: dict[str, Any]) -> str:
    """中文 Markdown 报告，按 测试结论 / 风险结论 / 问题描述 / 用例执行 4 块组织。"""
    summary = _build_executive_summary(report, tool)
    lines: list[str] = []
    name = tool.get("name", tool.get("id", "?"))
    icon = tool.get("icon", "")
    lines.append(f"# {icon} {name} · 分析报告\n")

    rid = r.get("run_id") or "?"
    lines.append(f"- **Run ID**：`{rid}`")
    meta = report.get("meta", {})
    if meta.get("model_id"):
        lines.append(f"- **模型**：`{meta['model_id']}`")
    if meta.get("produced_at_utc"):
        lines.append(f"- **生成时间**：{meta['produced_at_utc']}")
    lines.append("")

    # 测试结论
    verdict = summary["测试结论"]["判定"]
    icon_map = {"pass": "✅", "warn": "⚠️", "fail": "❌", "skip": "—"}
    vlevel = summary["测试结论"]["level"]
    lines.append(f"## ① 测试结论\n")
    lines.append(f"**{icon_map.get(vlevel, '·')} {verdict}**\n")
    sev = summary["严重度分布"]
    lines.append(f"严重度分布：致命 {sev['critical']} · 高 {sev['high']} · 中 {sev['medium']} · 低 {sev['low']} · 信息 {sev['info']}\n")

    # 风险结论
    lines.append(f"## ② 风险结论\n")
    if summary["风险结论"]:
        for rk in summary["风险结论"]:
            lines.append(f"- {rk}")
    else:
        lines.append("- （无显著风险）")
    lines.append("")

    # 问题描述 — ticket 级
    lines.append(f"## ③ 问题描述\n")
    issues = summary["问题描述"]
    owner_map = {"backend": "后端", "frontend": "前端", "product": "产品",
                 "test": "测试", "devops": "运维", "security": "安全", "data": "数据"}
    if not issues:
        lines.append("（本次未识别到具体问题）\n")
    else:
        for idx, it in enumerate(issues, start=1):
            sev_icon = {"critical": "🔴", "high": "🟠", "medium": "🟡",
                        "low": "🟢", "info": "🔵"}.get(it["severity"], "·")
            head = f"### {idx}. {sev_icon} {it['title']}"
            if it.get("issue_id"):
                head = f"### {idx}. {sev_icon} `{it['issue_id']}` {it['title']}"
            lines.append(head + "\n")
            meta_parts = []
            if it.get("owner") and it["owner"] in owner_map:
                meta_parts.append(f"👤 {owner_map[it['owner']]}")
            if it.get("hours"):
                meta_parts.append(f"⏱ 估时 {it['hours']}h")
            if meta_parts:
                lines.append("> " + " · ".join(meta_parts) + "\n")
            if it.get("module"):
                lines.append(f"**位置**：`{it['module']}`\n")
            if it.get("current"):
                lines.append(f"**现状**：{it['current']}\n")
            if it.get("expected"):
                lines.append(f"**期望**：{it['expected']}\n")
            if it.get("fix"):
                lines.append(f"**修复建议**：{it['fix']}\n")
            if it.get("repro"):
                lines.append("**✅ 复现步骤**：\n")
                for ri, step in enumerate(it["repro"], 1):
                    lines.append(f"  {ri}. {step}")
                lines.append("")
            if it.get("accept"):
                lines.append(f"**🧪 验收标准**：{it['accept']}\n")
            if it.get("cases"):
                lines.append(f"**关联用例**：{', '.join('`'+c+'`' for c in it['cases'])}\n")
            if it.get("impact"):
                lines.append(f"**影响面**：{it['impact']}\n")
            if it.get("evidence"):
                lines.append(f"**证据**：{it['evidence']}\n")
            lines.append("---\n")

    # 用例执行
    cse = summary["用例执行"]
    lines.append(f"## ④ 用例执行情况\n")
    lines.append(f"- 共生成 **{cse['总数']}** 条用例")
    pri = cse["按优先级"]
    if any(pri.values()):
        lines.append(f"- 优先级分布：P0 {pri['P0']} · P1 {pri['P1']} · P2 {pri['P2']} · P3 {pri['P3']}"
                     + (f" · 其他 {pri['其他']}" if pri['其他'] else ""))
    lines.append(f"- 说明：{cse['说明']}\n")

    lines.append("---\n")
    lines.append(f"由 天枢·裁决 生成 · {tool.get('id','?')}\n")
    return "\n".join(lines)


@app.get("/api/screenshots/{filename}")
async def api_screenshot_file(filename: str) -> Any:
    """Serve a captured screenshot PNG. Filename is the stored name from the
    report's `meta.screenshots[*].filename`."""
    # Defense in depth: no path traversal
    if "/" in filename or "\\" in filename or ".." in filename:
        raise HTTPException(400, "invalid filename")
    p = Path(settings.evidence_output_dir) / "screenshots" / filename
    if not p.exists():
        raise HTTPException(404, f"{filename} not found")
    return FileResponse(p, media_type="image/png")


def _annotate_screenshots(report_dump: dict[str, Any], shots: list[dict[str, Any]]) -> None:
    """Walk substep outputs for issue items with bbox + viewport_filename, draw
    red rectangles + numeric labels onto a copy of each screenshot, and add the
    annotated filenames into shots[*].annotated_filename so the report renderer
    can show them.

    The LLM is asked to emit issues in the shape:
        {"viewport_filename": <str>, "bbox": [x, y, w, h], "issue": "...", "severity": ...}
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return
    if not shots:
        return

    # Aggregate issues per filename
    issues_by_file: dict[str, list[dict[str, Any]]] = {}
    def _walk(x: Any) -> None:
        if isinstance(x, dict):
            fn = x.get("viewport_filename") or x.get("filename")
            bbox = x.get("bbox")
            if fn and isinstance(bbox, (list, tuple)) and len(bbox) == 4:
                issues_by_file.setdefault(str(fn), []).append({
                    "bbox": [int(b) for b in bbox],
                    "issue": str(x.get("issue") or x.get("title") or "")[:200],
                    "severity": _normalize_severity(x.get("severity")),
                })
            for v in x.values():
                _walk(v)
        elif isinstance(x, list):
            for v in x:
                _walk(v)
    _walk(report_dump.get("substeps", {}))

    if not issues_by_file:
        return

    out_dir = Path(settings.evidence_output_dir) / "screenshots"
    sev_color = {
        "critical": (220, 38, 38, 255),
        "high":     (220, 38, 38, 255),
        "major":    (220, 38, 38, 255),
        "medium":   (217, 119, 6, 255),
        "low":      (234, 179, 8, 255),
        "minor":    (234, 179, 8, 255),
        "cosmetic": (234, 179, 8, 255),
    }

    for shot in shots:
        fn = shot.get("filename")
        if not fn or fn not in issues_by_file:
            continue
        src = out_dir / fn
        if not src.exists():
            continue
        try:
            img = Image.open(src).convert("RGBA")
            draw = ImageDraw.Draw(img, "RGBA")
            for idx, it in enumerate(issues_by_file[fn], start=1):
                x, y, w, h = it["bbox"]
                color = sev_color.get(it["severity"], (217, 119, 6, 255))
                # Outer box
                for off in (3, 2, 1, 0):
                    draw.rectangle([x - off, y - off, x + w + off, y + h + off],
                                   outline=color, width=1)
                # Number badge top-left
                badge_size = 28
                badge_box = [x, y, x + badge_size, y + badge_size]
                draw.rectangle(badge_box, fill=color)
                try:
                    fnt = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 16)
                except Exception:
                    fnt = ImageFont.load_default()
                draw.text((x + 8, y + 4), str(idx), fill=(255, 255, 255), font=fnt)
            ann_fn = src.stem + "_annotated.png"
            ann_path = out_dir / ann_fn
            img.save(ann_path, "PNG", optimize=True)
            shot["annotated_filename"] = ann_fn
            shot["issue_count"] = len(issues_by_file[fn])
        except Exception:
            continue


async def _run_tool_async(
    *, run_id: str, tool: dict[str, Any], inputs: dict[str, Any], tenant: Tenant
) -> None:
    """Background task — execute orchestrator and stash result in _RUNS."""
    state = _RUNS[run_id]
    state["status"] = "running"
    state["progress"] = "构造执行上下文…"
    # Real-time log stream — frontend polls this list and renders entries
    state.setdefault("logs", [])

    def _emit(event: str, fields: dict[str, Any] | None = None, **kwargs: Any) -> None:
        # Accept both `_emit("ev", {"k":v})` (StepContext callback) and
        # `_emit("ev", k=v)` (direct kwargs from this file).
        if fields is None:
            fields = kwargs
        elif kwargs:
            fields = {**fields, **kwargs}
        entry = {"ts": _time.time(), "event": event, **fields}
        state["logs"].append(entry)
        # Cap log size to avoid memory bloat on long runs
        if len(state["logs"]) > 500:
            del state["logs"][:100]
        # Mirror substep.start / substep.done into the visible "progress" line
        if event == "substep.start":
            sub = fields.get("sub_id", "?")
            name = fields.get("name", "")
            state["progress"] = f"执行子步骤 {sub} · {name}"
        elif event == "substep.done":
            sub = fields.get("sub_id", "?")
            tokens = fields.get("output_tokens", 0)
            state["progress"] = f"完成 {sub} · {tokens:,} tokens"
    try:
        # Strip and apply model/effort/thinking overrides from the request.
        model_override = inputs.pop("__model", None)
        effort_override = inputs.pop("__effort", None)
        thinking_override = inputs.pop("__thinking", None)
        # Per-substep enable list (UI checkboxes). If empty list → run nothing,
        # which usually means user-error so we treat it as "all enabled".
        enabled_substeps = inputs.pop("__enabled_substeps", None)
        # Map the simplified UI form to orchestrator-expected keys.
        inputs = _shape_orchestrator_inputs(tool, inputs)
        if isinstance(enabled_substeps, list) and enabled_substeps:
            inputs["__enabled_substeps"] = enabled_substeps
        # Step 6 needs an ExecutionEnvironment object (dataclass; can't go in JSON body)
        if tool["id"] == "step6":
            inputs.setdefault("execution_env", ExecutionEnvironment(
                evidence_dir=Path(settings.evidence_output_dir) / "step6",
                dry_run=bool(inputs.get("dry_run", True)),
            ))
        ctx = await make_context(tenant, inputs)
        ctx.log_callback = _emit
        # Resolve dropdown key (e.g. "opus-4-7-1m") → actual model + betas.
        # Backwards-compat: bare aliases ("opus"/"sonnet"/"haiku") also accepted.
        model_id: str | None = None
        betas: list[str] = []
        if model_override:
            entry = _model_by_key(model_override)
            if entry:
                model_id = entry["model"]
                betas = list(entry.get("betas") or [])
            else:
                model_id = model_override  # treat as raw alias / model name
        if model_id or effort_override or thinking_override or betas:
            ctx.llm = LlmClient(
                model_override=model_id,
                effort=effort_override or None,
                thinking=thinking_override or None,
                betas=betas,
            )
        # 记下本次用的模型 key,失败时 except 分支用它做运行时黑名单
        state["model_key"] = model_override
        state["model_id"] = model_id
        # 跑之前确保 OAuth token 没过期 — 过期会让所有 LLM 调用拿 401
        try:
            refresh_result = await ensure_fresh_oauth_token()
            if refresh_result.get("status") == "refreshed":
                _emit("oauth.refreshed", detail=refresh_result.get("detail", ""))
            elif refresh_result.get("status") in ("no_refresh_token", "error"):
                _emit("oauth.warn", detail=refresh_result.get("detail", ""))
        except Exception as exc:
            _emit("oauth.warn", detail=f"token 刷新检查异常: {exc}")
        _emit("run.start", tool_id=tool["id"], tool_name=tool["name"])
        # ctx.run_id is independent of our public run_id; we keep ours for tracking
        state["ctx_run_id"] = ctx.run_id
        from packages.workflow.generic import (
            AgentExecutionOrchestrator,
            ApiTestOrchestrator,
            H5AdaptOrchestrator,
            NetworkResilienceOrchestrator,
            RequirementReviewOrchestrator,
            SeoAuditOrchestrator,
            TestCaseDesignOrchestrator,
            UiConsistencyOrchestrator,
        )
        orch_cls = {
            # 全部走 DirectChainOrchestrator —— 每个 substep 独立测试
            "step1": RequirementReviewOrchestrator,
            "step2": TestCaseDesignOrchestrator,
            "step4": ApiTestOrchestrator,
            "step5": UiConsistencyOrchestrator,
            "step6": AgentExecutionOrchestrator,
            "network_resilience": NetworkResilienceOrchestrator,
            "seo_audit": SeoAuditOrchestrator,
            "h5_adapt": H5AdaptOrchestrator,
        }[tool["id"]]
        # ── 各工具的「AI 真执行」前置阶段(全 agentic,把真实执行结果注入材料)──
        # seo/network 保留 Python 采集器(真实全站爬取 / 真实 CDP 限速断网),采集真实
        # 结构化数据供 10-sheet 富 Excel + 多档位矩阵;判断由加深后的综合提示词完成。
        # 注意：h5_adapt 已改为「分析宿主机三端真机/模拟器证据」模式（scripts/h5_device_collect.py
        # 在宿主机跑 iOS 模拟器/Android 模拟器/Chrome 采集真 DOM+真截图，产出 evidence.md 作为 documents
        # 喂入），不再在容器内用桌面 Chromium 模拟视口，故从 _browser_cfg / _TOOL_VIEWPORTS 移除。
        _browser_cfg = {
            "step6":              {"prompt": "step6_agent/_execute.md",        "http": True,  "net": False, "steps": 16},
        }
        # 用户上传的文件(PDF/图片/文本/Office):按 documents 里的 file_ref 标记
        # 加载真实文件 → ctx.files。先于 agentic 执行阶段加载,使 step4/浏览器
        # 规划阶段也能"看到"上传文件;base.run_substep 同样把它们作为
        # document/image 内容块经 stdin 直传 LLM(不读成文本塞进 prompt)。
        try:
            uf = _load_uploaded_files(ctx)
            if uf:
                ctx.files = uf
                state["uploaded_files"] = [{"filename": f["filename"], "kind": f["kind"]} for f in uf]
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "files.loaded", "count": len(uf),
                    "names": [f["filename"] for f in uf][:10]})
        except Exception as exc:
            state.setdefault("logs", []).append({
                "ts": _time.time(), "event": "files.load.failed", "error": str(exc)[:200]})

        # step6:上传的 Excel 用例表 → 智能识别列 → 统一 cases 注入(与 step2 用例设计两路归一)
        if tool["id"] == "step6":
            try:
                _ncs = _inject_excel_cases(ctx)
                if _ncs:
                    state.setdefault("logs", []).append({
                        "ts": _time.time(), "event": "step6.excel_cases.parsed", "count": _ncs})
            except Exception as exc:
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "step6.excel.failed", "error": str(exc)[:200]})

        if tool["id"] == "step4":
            # AI 真发 HTTP 请求看真实响应
            try:
                await _execute_apis_agentic(ctx, state)
            except Exception as exc:
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "api.execute.failed", "error": str(exc)[:200]})
        elif tool["id"] == "seo_audit":
            # SEO:真实全站 BFS 爬取 + 逐页解析 + Core Web Vitals 实测(确定性采集),
            # 采集摘要注入材料供深度综合提示词分析,原始数据存 state 供 10-sheet Excel 渲染。
            try:
                await _seo_run_collect(ctx, state)
            except Exception as exc:
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "seo.collect.failed", "error": str(exc)[:200]})
        elif tool["id"] == "network_resilience":
            # 弱网/断网:仅当用户勾选「页面/接口弱网」时才跑容器 Playwright CDP 多档位实测;
            # 选纯「视频弱网」时跳过(视频证据由宿主机 net_video_collect.py 采集并喂入,容器不重复跑、不卡不可达URL)。
            if bool((ctx.inputs or {}).get("mode_page_api", True)):
                try:
                    await _network_run_collect(ctx, state)
                except Exception as exc:
                    state.setdefault("logs", []).append({
                        "ts": _time.time(), "event": "net.collect.failed", "error": str(exc)[:200]})
            else:
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "net.page_api.skipped", "reason": "用户未勾选页面/接口弱网,仅做视频弱网"})
        elif tool["id"] in _browser_cfg:
            # AI 真驱动浏览器(H5 多视口 / step6 端到端)
            cfg = _browser_cfg[tool["id"]]
            try:
                _bshots: list[dict[str, Any]] = []
                await _run_browser_agent(
                    ctx, state, cfg["prompt"], f"{tool['id']}_{ctx.run_id[:8]}",
                    _bshots, with_http=cfg["http"], with_network=cfg["net"], max_steps=cfg.get("steps", 16))
                if _bshots:
                    state["screenshots"] = (state.get("screenshots") or []) + _bshots
                    ctx.screenshots = (getattr(ctx, "screenshots", None) or []) + _bshots
            except Exception as exc:
                state.setdefault("logs", []).append({
                    "ts": _time.time(), "event": "browser.agent.failed", "error": str(exc)[:200]})

        # For UI-related tools, capture page screenshots before LLM analysis.
        # Attaches PNG paths to ctx so run_substep can pass them as image
        # content blocks (Claude will then actually SEE the page).
        shots = await _capture_screenshots_for_tool(tool["id"], ctx, state)
        if shots:
            state["screenshots"] = shots
            ctx.screenshots = shots  # consumed by base.run_substep
            # Append a summary so the prompt also names the screenshots in text.
            docs = (ctx.inputs or {}).get("documents") or ""
            if docs:
                lines = ["", "", "## 已采集的页面截图（已作为 image content 附在本次请求中）"]
                ok_shots = [s for s in shots if not s.get("error")]
                if ok_shots:
                    for s in ok_shots:
                        lines.append(f"- {s['url']} @ {s['viewport']} ({s['width']}×{s['height']})")
                err_shots = [s for s in shots if s.get("error")]
                if err_shots:
                    lines.append("")
                    lines.append("(以下 URL 截图失败，请在分析时标注无法获取真实视图)")
                    for s in err_shots:
                        lines.append(f"- {s['url']} @ {s['viewport']}: {s['error'][:100]}")
                ctx.inputs["documents"] = docs + "\n".join(lines)

        state["progress"] = "调用 Claude（本地客户端）…"
        # seo/network:采集器已拿到全部真实事实,用加深后的综合提示词(统一契约 + 逐维穷尽)
        # 做深度判断与定稿;采集截图仅留报告/Excel 展示,不喂分析。其余工具走 orchestrator。
        if tool["id"] == "seo_audit":
            report_dump = await _seo_synthesize(ctx, state)
        elif tool["id"] == "network_resilience":
            report_dump = await _network_synthesize(ctx, state)
        elif tool["id"] == "step5":
            # step5:逐帧设计 vs 实拍比对,1 次多模态调用(替代 5 子步骤,~50min→几分钟)
            report_dump = await _step5_synthesize(ctx, state)
        else:
            report = await orch_cls(ctx).execute()
            report_dump = report.model_dump(mode="json")
        # Persist a JSON copy alongside the run-id for download
        out_dir = Path(settings.report_output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        # 把 run 持有人写到 meta — 持久化跨重启可用,/api/reports 据此过滤
        report_meta = report_dump.setdefault("meta", {})
        if state.get("owner_user_id"):
            report_meta["owner_user_id"] = state["owner_user_id"]
            report_meta["owner_username"] = state.get("owner_username")
        # Attach screenshot metadata so HTML/MD renderers can embed them
        if state.get("screenshots"):
            shots_list = state["screenshots"]
            # Phase C: annotate originals with red boxes for any LLM-reported bbox
            try:
                _annotate_screenshots(report_dump, shots_list)
            except Exception:
                pass
            report_meta["screenshots"] = shots_list
        # SEO:把真实采集数据 + 基线一并存进报告 meta,供 Excel 导出渲染 10 sheet
        if state.get("seo_data"):
            report_meta["seo_data"] = state["seo_data"]
            if state.get("seo_baseline"):
                report_meta["seo_baseline"] = state["seo_baseline"]
        # 弱网/断网:多档位实测数据存 meta,供 Excel 渲染
        if state.get("network_data"):
            report_meta["network_data"] = state["network_data"]
        # 多维报告:把本次 token / 成本用量存进 meta(供使用统计聚合;历史报告没有此字段)
        if state.get("usage"):
            report_meta["usage"] = state["usage"]
        report_path = out_dir / f"{tool['id']}_{ctx.run_id}.json"
        report_path.write_text(
            _json.dumps(report_dump, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        state["progress"] = "完成"
        state["report"] = report_dump
        state["report_path"] = str(report_path)
        state["usage"] = {
            "input_tokens": ctx.usage.input_tokens,
            "output_tokens": ctx.usage.output_tokens,
            "cache_read_tokens": ctx.usage.cache_read_tokens,
            "cost_usd": round(ctx.usage.cost_usd, 4),
        }
        state["status"] = "succeeded"
    except Exception as exc:
        state["status"] = "failed"
        # 给前端可读的一句话:常见根因(模型不可用 / 未登录 / 网络)抽出来作业务语言
        raw = f"{type(exc).__name__}: {exc}"
        low = raw.lower()
        hint = None
        if "haiku-4-5" in low or "claude-haiku-4-5" in low:
            hint = "Haiku 4.5 调用失败 — 该模型仍在受限灰度，请确认当前账号已开通；或先切回 Sonnet 4.6 验证工具链。"
        elif "authnotconfigured" in low or "未配置" in raw or "未登录" in raw:
            hint = "认证未就绪 — 请到「设置 → 模型接入」选择 OAuth 或填 API Key 后重试。"
        elif "sdk invocation failed" in low or "claude sdk" in low:
            hint = "Claude SDK 调用失败 — 通常是模型不可用、配额已用尽或本地 CLI 未登录。请先在「设置」做一次模型探测。"
        state["error"] = f"{hint}\n\n{raw}" if hint else raw
        state["traceback"] = _tb.format_exc()
        # 失败模型加入会话黑名单 → 下次 /api/claude/info 返回时 available=false,
        # 下拉自动 disable, 用户不会继续踩同一个模型。
        try:
            failing_model = inputs.get("__model") if isinstance(inputs, dict) else None
            if not failing_model:
                # 走过 inputs.pop 之后 __model 已经被消费,从 state 里拿
                failing_model = state.get("model_key")
            if failing_model:
                _record_model_failure(failing_model, raw)
        except Exception:
            pass
        _emit("run.failed", error=str(exc)[:200])
    finally:
        state["finished_at"] = _time.time()
        if state["status"] == "succeeded":
            _emit("run.done", cost_usd=state.get("usage", {}).get("cost_usd"))


@app.post("/api/tools/{tool_id}/run")
async def api_tool_run(
    tool_id: str,
    body: dict[str, Any],
    request: Request,
    x_tenant_id: str | None = Header(default="default"),
    x_project_id: str | None = Header(default=None),
) -> dict[str, Any]:
    """异步启动一个工具，立即返回 run_id 用于轮询。

    入口闸：用户没在「设置 → 认证模式」选过连接方式（mode == 'unset'）就拒绝
    跑工具，给清晰的 412 错误。前端 toast 出来直接引导去设置页。
    """
    tool = next((t for t in TOOL_CATALOG if t["id"] == tool_id), None)
    if not tool:
        raise HTTPException(404, f"unknown tool: {tool_id}")

    # 全局并发闸:同一工具同一时间只能跑一个 run。
    # 多人同时按运行 → 第二位拿到 409 + 触发者用户名。前端 toast。
    for _r in _RUNS.values():
        if _r.get("tool_id") == tool_id and _r.get("status") in ("queued", "running"):
            owner = _r.get("owner_username") or "其他用户"
            raise HTTPException(
                409,
                f"该工具正被 {owner} 运行,请等它完成再试 (进度:{_r.get('progress') or '初始化'})",
            )

    # 认证闸：未配置/未登录直接 412 拒绝
    # 三种情况都要拦：
    #   - mode=unset            → 用户没选过任何模式
    #   - mode=oauth + 未登录    → 选了 OAuth 但还没浏览器走完授权
    #   - mode=api_key + 无 key  → 选了 API Key 但没填
    try:
        from packages.core.auth_config import (
            get_api_key, get_auth_mode, get_oauth_access_token,
        )
        mode = get_auth_mode()
        if mode == "unset":
            raise HTTPException(
                412,
                "未选择连接方式 — 请到「设置 → 模型接入」选择 OAuth 或 API Key 后再试",
            )
        if mode == "oauth" and not get_oauth_access_token():
            raise HTTPException(
                412,
                "OAuth 模式但还未完成授权 — 请到「设置 → 模型接入」点「OAuth 授权」完成浏览器授权",
            )
        if mode == "api_key" and not get_api_key():
            raise HTTPException(
                412,
                "API Key 模式未填 Key — 请到「设置 → 模型接入」粘贴 sk-ant-... 或换其他模式",
            )
    except HTTPException:
        raise
    except Exception:
        pass  # 配置模块异常不阻塞跑工具，让 LLM 客户端自己抛

    # 强制必填:项目编号 + 项目名称
    project_code = str(body.get("project_code", "") or "").strip()
    project_name = str(body.get("project_name", "") or "").strip()
    if not project_code:
        raise HTTPException(400, "项目编号 (project_code) 必填")
    if not project_name:
        raise HTTPException(400, "项目名称 (project_name) 必填")
    # 简单长度校验
    if len(project_code) > 64:
        raise HTTPException(400, "项目编号过长 (最多 64 字符)")
    if len(project_name) > 128:
        raise HTTPException(400, "项目名称过长 (最多 128 字符)")

    project_id = x_project_id or body.get("project_id") or project_code or "tools-portal"
    tenant = Tenant(tenant_id=x_tenant_id or "default", project_id=project_id)

    # Strip UI-only keys before passing as orchestrator inputs;
    # project_code/name 透传给 orchestrator,后续会被写入 report.meta
    helper_keys = {"project_id"}  # project_code/name 保留给 orchestrator 用
    inputs = {k: v for k, v in body.items() if k not in helper_keys}

    # 把当前登录用户钉到 run 上,用于报告隔离
    current_user = getattr(request.state, "current_user", None)
    owner_user_id = current_user.id if current_user else None
    owner_username = current_user.username if current_user else None

    run_id = str(uuid4())
    _RUNS[run_id] = {
        "run_id": run_id,
        "tool_id": tool_id,
        "tool_name": tool["name"],
        "status": "queued",
        "progress": "等待执行…",
        "started_at": _time.time(),
        "finished_at": None,
        "tenant_id": tenant.tenant_id,
        "project_id": tenant.project_id,
        "project_code": project_code,
        "project_name": project_name,
        "owner_user_id": owner_user_id,
        "owner_username": owner_username,
    }
    # Fire-and-forget; client polls /api/tools/runs/{run_id}
    # 包一层:支持「停止执行」取消任务,取消时标记 cancelled(而非 failed)并清理句柄。
    async def _runner():
        try:
            await _run_tool_async(run_id=run_id, tool=tool, inputs=inputs, tenant=tenant)
        except _asyncio.CancelledError:
            st = _RUNS.get(run_id)
            if st is not None and st.get("status") not in ("succeeded", "failed"):
                st["status"] = "cancelled"
                st["progress"] = st.get("progress") or "已停止"
                st["finished_at"] = _time.time()
        finally:
            _RUN_TASKS.pop(run_id, None)
    _RUN_TASKS[run_id] = _asyncio.create_task(_runner())
    return {"run_id": run_id, "status": "queued"}


def _promote_contract_fields(report: dict[str, Any]) -> dict[str, Any]:
    """读时把 finalize substep 的统一报告契约字段提升到 report 顶层。

    历史 run 是旧 Python 代码跑的，contract 字段只存在于 substeps[finalize_id] 里。
    这里做无副作用的"读时 patch"，让 UI 不用重跑就能看到完整 5 段报告。
    新版 GenericReport 已经在 execute() 里直接写到顶层，此函数对新报告也保持幂等。
    """
    if not isinstance(report, dict):
        return report
    subs = report.get("substeps") or {}
    if not isinstance(subs, dict) or not subs:
        return report
    # 找最后一个非空 substep 当 final
    final: dict[str, Any] = {}
    for sid in list(subs.keys())[::-1]:
        v = subs.get(sid)
        if isinstance(v, dict) and v:
            final = v
            break
    if not final:
        return report

    def merge_list(field: str) -> list[Any]:
        # 永远合并所有 substep 的该数组(去重靠 ID)。
        # 不能"顶层非空就不动"—— finalize 只贡献自己那部分用例时,
        # 会漏掉前 4 个子步骤的用例(实测 step2:8 vs 应有 57)。
        merged: list[Any] = []
        seen_ids: set[str] = set()
        no_id_items: list[Any] = []
        for sid in subs:
            out = subs.get(sid)
            if not isinstance(out, dict):
                continue
            for item in (out.get(field) or []):
                key = ""
                if isinstance(item, dict):
                    key = str(item.get("id") or item.get("issue_id") or item.get("case_id") or "")
                if key:
                    if key in seen_ids:
                        continue
                    seen_ids.add(key)
                    merged.append(item)
                else:
                    no_id_items.append(item)  # 无 ID 的不能去重,但也不丢
        merged.extend(no_id_items)
        # 如果合并结果为空但顶层本来有(旧报告无 substeps 的情况),保留顶层
        cur = report.get(field)
        if not merged and isinstance(cur, list) and cur:
            return cur
        return merged

    if not report.get("verdict") and isinstance(final.get("verdict"), str):
        report["verdict"] = final["verdict"]
    if not report.get("verdict_summary") and isinstance(final.get("verdict_summary"), str):
        report["verdict_summary"] = final["verdict_summary"]
    report["risks"] = merge_list("risks")
    report["blockers"] = merge_list("blockers")
    report["issues"] = merge_list("issues")
    report["cases"] = merge_list("cases")
    return report


@app.get("/api/tools/runs/{run_id}")
async def api_tool_run_status(run_id: str, request: Request) -> dict[str, Any]:
    state = _RUNS.get(run_id)
    if not state:
        _hydrate_runs_from_disk(only_run_id=run_id)   # 重启后内存丢失 → 从磁盘报告补加载
        state = _RUNS.get(run_id)
    if not state:
        raise HTTPException(404, f"run {run_id} not found (server may have restarted)")
    user = require_user(request)
    if not _user_can_see(user, state.get("owner_user_id")):
        raise HTTPException(403, "无权访问此运行")
    # 读时把契约字段提升到 report 顶层（无副作用 / 幂等）
    if isinstance(state.get("report"), dict):
        state["report"] = _promote_contract_fields(state["report"])
    # can_stop:能看到本 run(已过 _user_can_see=持有者或管理员)且仍在跑 → 可停止。
    can_stop = state.get("status") in ("queued", "running")
    return {**state, "can_stop": can_stop}


@app.post("/api/tools/runs/{run_id}/stop")
async def api_tool_run_stop(run_id: str, request: Request) -> dict[str, Any]:
    """停止执行 — 仅「执行人(run 持有者)」或「管理员」可停。"""
    state = _RUNS.get(run_id)
    if not state:
        raise HTTPException(404, f"run {run_id} not found")
    user = require_user(request)
    if not _user_can_see(user, state.get("owner_user_id")):
        raise HTTPException(403, "无权停止此运行（仅执行人或管理员可停）")
    if state.get("status") in ("succeeded", "failed", "cancelled"):
        return {"ok": True, "status": state["status"], "note": "运行已结束，无需停止"}
    # 协作取消标志(供长循环检查)+ 立即标记状态(释放工具锁)+ 取消异步任务。
    state["cancelled"] = True
    state["status"] = "cancelled"
    state["progress"] = f"已被 {user.username} 手动停止"
    state["finished_at"] = _time.time()
    state["error"] = f"运行被 {user.username} 手动停止"
    task = _RUN_TASKS.get(run_id)
    if task is not None and not task.done():
        task.cancel()
    return {"ok": True, "status": "cancelled"}


def _user_can_see(user: UserRecord, owner_user_id: int | None) -> bool:
    """报告可见性:仅超级管理员可看所有;管理员与普通用户只能看自己的。
    legacy (owner_user_id 缺失) 仅对超级管理员可见。"""
    if user.is_superadmin():
        return True
    if owner_user_id is None:
        return False  # 旧数据,非超管看不到
    return int(owner_user_id) == int(user.id)


@app.get("/api/tools/{tool_id}/active")
async def api_tool_active_status(tool_id: str, request: Request) -> dict[str, Any]:
    """查询某个工具是否正在被任意用户运行(全局锁状态)。

    设计:同一工具同时只能跑一个 — 因为底层共用一份 Claude 凭据 + 同一个进程,
    并发跑会触发 Anthropic 速率限制 + 资源争用。所以 UI 侧锁掉「运行」按钮,
    任何人打开页面都看到「运行中 · 由 xxx 触发」。
    完成后此端点 2-3 秒内会返回 active=False,前端轮询自动解锁。

    所有登录用户都能查 — 只暴露 tool/run_id/owner_username/started_at/progress,
    不暴露报告内容,所以不需要 _user_can_see 过滤。
    """
    require_user(request)
    for r in _RUNS.values():
        if r.get("tool_id") != tool_id:
            continue
        if r.get("status") not in ("queued", "running"):
            continue
        return {
            "active": True,
            "run_id": r.get("run_id"),
            "started_at": r.get("started_at"),
            "owner_username": r.get("owner_username"),
            "progress": r.get("progress"),
        }
    return {"active": False}


@app.get("/api/tools/runs")
async def api_tool_run_list(request: Request) -> dict[str, Any]:
    """List recent in-process runs (memory only — clears on server restart).
    普通用户只看自己的;admin 看全部。"""
    user = require_user(request)
    visible = [r for r in _RUNS.values() if _user_can_see(user, r.get("owner_user_id"))]
    runs = sorted(visible, key=lambda r: r["started_at"], reverse=True)
    # 列表已按权限过滤(持有者/管理员可见),故可见即可停 → can_stop 仅看是否在跑。
    summarized = [
        {**{k: v for k, v in r.items() if k not in ("report", "traceback")},
         "can_stop": r.get("status") in ("queued", "running")}
        for r in runs[:50]
    ]
    return {"total": len(visible), "recent": summarized}


# ----- HTML pages -----

TOOLS_INDEX_HTML = """<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>天枢·裁决 · AI 裁决手册</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
/* ======= MUJI 出版物 · 设计 Token ======= */
:root {
  --paper:    #ffffff;
  --paper-2:  #f0f0f0;
  --paper-3:  #e3e3e3;
  --ink:      #0a0a0a;
  --ink-2:    #262626;
  --ink-3:    #4a4a4a;
  --ink-4:    #6e6e6e;
  --line:     #c4c4c4;
  --line-2:   #9e9e9e;
  --accent:   #a8401f;
  --accent-h: #82301a;
  --ok:       #4f6b35;
  --warn:     #8a5300;
  --bad:      #9a3315;
  --serif:    'Noto Serif SC', 'Songti SC', 'STSong', 'SimSun', Georgia, serif;
  --sans:     'Noto Sans SC', 'PingFang SC', -apple-system, 'Microsoft YaHei', sans-serif;
  --mono:     'Inter', 'SF Mono', ui-monospace, Menlo, monospace;
}
*{box-sizing:border-box}
html,body{margin:0;background:var(--paper);color:var(--ink);
  font-family:var(--sans);font-weight:300;font-size:15px;line-height:1.85;
  -webkit-font-smoothing:antialiased;font-feature-settings:"palt" on}
body{height:100vh;display:flex;flex-direction:column;overflow:hidden}
a{color:inherit;text-decoration:none}
button{font-family:inherit;cursor:pointer}
::selection{background:rgba(196,90,58,.22);color:var(--accent)}
::-webkit-scrollbar{width:8px;height:8px}
::-webkit-scrollbar-track{background:transparent}
::-webkit-scrollbar-thumb{background:var(--line-2);border-radius:4px}

/* ======= 顶栏:极简 ======= */
.topbar{position:sticky;top:0;background:rgba(255,255,255,.94);
  backdrop-filter:saturate(140%) blur(10px);-webkit-backdrop-filter:saturate(140%) blur(10px);
  border-bottom:1px solid var(--line);z-index:50;
  display:flex;align-items:center;gap:12px;padding:0 24px;height:56px}
.topbar .brand{font-family:var(--serif);font-size:16px;font-weight:500;
  letter-spacing:.1em;color:var(--ink)}
.topbar .brand .sep{margin:0 6px;color:var(--ink-3)}
.topbar nav{display:flex;gap:2px;margin-left:24px;margin-right:auto}
.topbar nav a{font-size:13px;color:var(--ink-2);letter-spacing:.04em;
  font-family:var(--sans);font-weight:400;padding:6px 12px;border-radius:6px;
  text-decoration:none;transition:all .15s}
.topbar nav a:hover{background:var(--paper-2);color:var(--ink)}
.topbar nav a.active{background:var(--accent-soft,rgba(168,64,31,.12));color:var(--accent)}
.topbar .kbd-hint{margin-left:12px;font-family:var(--mono);font-size:11px;color:var(--ink-3);
  cursor:pointer}
.topbar .kbd-hint:hover{color:var(--accent)}
.topbar .kbd-hint kbd{background:var(--paper-2);border:1px solid var(--line);
  padding:1px 5px;border-radius:3px;font-family:var(--mono);font-size:10.5px;color:var(--ink-2);
  margin:0 1px}

/* ======= Env banner (alert) — 仅必要时显示 ======= */
.env-banner{display:none;margin:24px 48px 0;padding:14px 18px;
  background:rgba(196,90,58,.10);border-left:2px solid var(--accent);
  font-size:13px;color:var(--ink-2);font-weight:400}
.env-banner.show{display:block}
.env-banner h3{margin:0 0 6px;font-family:var(--sans);font-size:13px;
  font-weight:500;color:var(--ink);letter-spacing:.04em}
.env-banner .icon{margin-right:8px;color:var(--accent)}
.env-banner .actions{margin-top:10px;display:flex;gap:14px;font-size:12px}
.env-banner .actions button{background:transparent;border:none;border-bottom:1px solid var(--ink-3);
  color:var(--ink);font-size:12px;padding:1px 0;letter-spacing:.04em}
.env-banner .actions button:hover{border-color:var(--accent);color:var(--accent)}
.env-banner .env-progress{font-family:var(--mono);font-size:11px;
  color:var(--ink-3);margin-top:8px;white-space:pre-wrap;max-height:120px;overflow:auto}

/* ======= Hero ======= */
.hero{flex:1;min-height:0;width:100%;box-sizing:border-box;display:flex;flex-direction:column;justify-content:center;
  padding:0 48px;max-width:1080px;margin:0 auto;align-items:flex-start}
.hero .eyebrow{font-family:var(--mono);font-size:11px;color:var(--ink-3);
  letter-spacing:.24em;text-transform:uppercase;margin-bottom:22px}
.hero h1{font-family:var(--serif);font-size:52px;font-weight:400;
  line-height:1.25;color:var(--ink);margin:0 0 32px;letter-spacing:.02em}
.hero h1 em{font-style:normal;color:var(--accent)}
.hero .lede{font-size:17px;line-height:1.95;color:var(--ink-2);font-weight:300;
  max-width:560px;margin:0 0 56px}
.hero .actions{display:flex;gap:14px;align-items:center}
.hero .btn{display:inline-flex;align-items:center;gap:10px;
  padding:14px 28px;border:1px solid var(--ink);background:transparent;
  font-family:var(--sans);font-size:14px;letter-spacing:.06em;color:var(--ink);
  font-weight:400;transition:all .18s}
.hero .btn:hover{background:var(--ink);color:var(--paper)}
.hero .btn.primary{background:var(--ink);color:var(--paper)}
.hero .btn.primary:hover{background:var(--accent);border-color:var(--accent)}
.hero .btn .arrow{font-family:var(--serif);font-size:14px}
.hero-foot{margin-top:80px;font-family:var(--mono);font-size:11px;color:var(--ink-3);
  letter-spacing:.16em;text-transform:uppercase;display:flex;align-items:center;gap:12px}
.hero-foot .dot{display:inline-block;width:5px;height:5px;border-radius:50%;background:var(--accent)}
.hero-foot .claude-status{margin-left:auto;font-family:var(--mono);font-size:11px}

/* ======= 章节分隔 ======= */
.chapter-divider{max-width:1080px;margin:120px auto 80px;padding:0 48px;
  display:flex;align-items:center;gap:24px}
.chapter-divider .label{font-family:var(--mono);font-size:11px;
  letter-spacing:.32em;color:var(--ink-3);text-transform:uppercase;white-space:nowrap}
.chapter-divider .line{flex:1;height:1px;background:var(--line)}

/* ======= 核心能力 三列 ======= */
.capabilities{max-width:1080px;margin:0 auto;padding:0 48px;
  display:grid;grid-template-columns:repeat(3,1fr);gap:80px}
.cap{}
.cap .num-zh{font-family:var(--serif);font-size:64px;font-weight:300;
  color:var(--ink);line-height:1;margin-bottom:10px;letter-spacing:.04em}
.cap .name{font-family:var(--serif);font-size:18px;font-weight:500;
  color:var(--ink);margin-bottom:18px;letter-spacing:.04em}
.cap .desc{font-size:14px;line-height:1.85;color:var(--ink-2);font-weight:300}

/* ======= 工具目录 ======= */
.directory{max-width:840px;margin:0 auto;padding:0 48px}
.directory-group{margin-bottom:56px}
.directory-group:last-child{margin-bottom:0}
.directory-group .group-label{font-family:var(--mono);font-size:10.5px;
  letter-spacing:.32em;color:var(--ink-3);text-transform:uppercase;
  margin-bottom:28px;padding-bottom:14px;border-bottom:1px solid var(--line)}
.directory-row{display:grid;grid-template-columns:1fr auto;gap:24px;
  align-items:baseline;padding:18px 0;border-bottom:1px dashed var(--line);
  transition:padding .18s ease,background .18s ease;cursor:pointer}
.directory-row:hover{background:linear-gradient(90deg,rgba(196,90,58,.08),transparent);
  padding-left:12px}
.directory-row:hover .name{color:var(--accent)}
.directory-row .chapter{font-family:var(--serif);font-size:14.5px;color:var(--ink-3);
  font-weight:400;letter-spacing:.16em;min-width:80px;white-space:nowrap}
.directory-row .name{font-family:var(--serif);font-size:21px;font-weight:500;
  color:var(--ink);letter-spacing:.04em;transition:color .15s}
.directory-row .desc{font-size:13px;color:var(--ink-2);line-height:1.7;
  display:block;margin-top:6px;font-weight:300}
.directory-row .open{font-family:var(--serif);font-size:18px;color:var(--ink-3);
  transition:color .15s,transform .15s}
.directory-row:hover .open{color:var(--accent);transform:translateX(4px)}

/* ======= 最近报告(索引式) ======= */
.recent{max-width:1080px;margin:0 auto;padding:0 48px}
.recent-month{font-family:var(--serif);font-size:24px;font-weight:400;
  color:var(--ink);margin:0 0 32px;letter-spacing:.08em}
.recent-table{width:100%;border-collapse:collapse}
.recent-table tr{cursor:pointer;transition:background .18s}
.recent-table tr:hover{background:rgba(196,90,58,.08)}
.recent-table td{padding:18px 0;border-bottom:1px solid var(--line);
  font-size:13.5px;color:var(--ink);vertical-align:baseline}
.recent-table .date{font-family:var(--mono);color:var(--ink-3);font-size:12.5px;
  width:120px;letter-spacing:.04em}
.recent-table .chapter-cell{font-family:var(--serif);color:var(--ink-2);
  width:180px;letter-spacing:.06em;font-size:14px}
.recent-table .code-cell{font-family:var(--mono);color:var(--ink-2);font-size:12px;
  width:200px;letter-spacing:.02em}
.recent-table .name-cell{color:var(--ink);font-weight:400}
.recent-table .status-cell{font-family:var(--mono);font-size:11px;
  letter-spacing:.16em;text-transform:uppercase;width:80px;text-align:right;
  color:var(--ink-3)}
.recent-table .status-cell.ok{color:var(--ok)}
.recent-table .status-cell.bad{color:var(--bad)}
.recent-table .status-cell.warn{color:var(--warn)}
.recent-table tr:focus-visible{outline:2px solid var(--accent);outline-offset:-2px}
.directory-row:focus-visible{outline:2px solid var(--accent);outline-offset:4px}
.recent-foot{margin-top:32px;text-align:center}
.recent-foot a{font-family:var(--sans);font-size:13px;color:var(--ink-2);
  letter-spacing:.08em;border-bottom:1px solid var(--ink-3);padding-bottom:2px}
.recent-foot a:hover{color:var(--accent);border-color:var(--accent)}
.recent-empty{text-align:center;color:var(--ink-3);font-size:13px;
  padding:48px 0;font-family:var(--serif);letter-spacing:.04em}

/* ======= Footer ======= */
.footer{width:100%;max-width:1080px;margin:0 auto;padding:20px 48px;
  border-top:1px solid var(--line);
  display:flex;align-items:baseline;justify-content:space-between;
  font-family:var(--mono);font-size:11px;color:var(--ink-3);
  letter-spacing:.16em;text-transform:uppercase}
.footer .left .brand{font-family:var(--serif);font-size:13px;letter-spacing:.16em;
  color:var(--ink);text-transform:none;margin-right:14px}
.footer .right{display:flex;gap:20px}

/* ======= 浮动「运行中」标记 ======= */
.task-fab{position:fixed;bottom:24px;right:24px;z-index:90;display:none;
  background:var(--paper);border:1px solid var(--line);
  padding:14px 18px;min-width:240px;max-width:320px;
  font-family:var(--sans);font-size:12.5px;color:var(--ink-2);
  box-shadow:0 4px 24px rgba(44,44,44,.06)}
.task-fab.show{display:block}
.task-fab .head{display:flex;align-items:center;gap:8px;
  font-family:var(--mono);font-size:11px;color:var(--ink-3);
  letter-spacing:.16em;text-transform:uppercase;margin-bottom:10px;
  padding-bottom:10px;border-bottom:1px solid var(--line)}
.task-fab .head .dot{width:6px;height:6px;border-radius:50%;background:var(--accent);
  animation:pulse 1.8s ease-in-out infinite}
.task-fab .head .count{margin-left:auto;font-family:var(--mono);font-size:11px;
  color:var(--ink-2)}
.task-fab .body{max-height:432px;overflow-y:auto}
.task-fab .row{padding:8px 0;border-bottom:1px dashed var(--line);font-size:12.5px;
  display:flex;align-items:center;gap:10px}
.task-fab .row:last-child{border-bottom:none}
.task-fab .row .row-link{flex:1;min-width:0;display:block;text-decoration:none}
.task-fab .row .title{color:var(--ink);font-family:var(--serif);font-size:13.5px;
  letter-spacing:.04em}
.task-fab .row .progress{color:var(--ink-3);font-family:var(--mono);font-size:11px;
  margin-top:3px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.task-fab .task-stop{flex-shrink:0;cursor:pointer;font-family:var(--mono);font-size:11px;
  padding:4px 10px;border-radius:5px;border:1px solid rgba(220,38,38,.35);
  background:rgba(220,38,38,.06);color:#dc2626;transition:background .15s}
.task-fab .task-stop:hover{background:rgba(220,38,38,.14)}
.task-fab .task-stop:disabled{opacity:.5;cursor:default}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.4}}

/* ======= 命令面板 ======= */
.cmd-overlay{position:fixed;inset:0;background:rgba(44,44,44,.30);
  backdrop-filter:blur(6px);-webkit-backdrop-filter:blur(6px);z-index:100;
  display:none;align-items:flex-start;justify-content:center;padding-top:120px}
.cmd-overlay.open{display:flex}
.cmd-panel{background:var(--paper);border:1px solid var(--line);
  width:min(640px,92vw);padding:6px;
  box-shadow:0 16px 60px rgba(44,44,44,.20)}
.cmd-input{width:100%;padding:14px 18px;background:transparent;border:none;
  font-family:var(--serif);font-size:16px;color:var(--ink);
  letter-spacing:.04em;outline:none;border-bottom:1px solid var(--line)}
.cmd-input::placeholder{color:var(--ink-3)}
.cmd-results{max-height:380px;overflow-y:auto;margin-top:4px}
.cmd-result{display:flex;align-items:baseline;gap:14px;padding:12px 18px;cursor:pointer;
  border-bottom:1px dashed var(--line);font-size:13.5px;color:var(--ink-2)}
.cmd-result:last-child{border-bottom:none}
.cmd-result.active,.cmd-result:hover{background:var(--paper-2);color:var(--ink)}
.cmd-result .label{font-family:var(--serif);color:var(--ink)}
.cmd-result .hint{margin-left:auto;font-family:var(--mono);font-size:11px;color:var(--ink-3)}

/* ======= Toast ======= */
.toast-stack{position:fixed;bottom:24px;left:50%;transform:translateX(-50%);
  z-index:200;display:flex;flex-direction:column;gap:8px;align-items:center}
.toast{background:var(--ink);color:var(--paper);padding:12px 22px;
  font-family:var(--sans);font-size:13px;letter-spacing:.04em;
  box-shadow:0 6px 20px rgba(44,44,44,.18);
  animation:toast-in .22s ease}
@keyframes toast-in{from{opacity:0;transform:translateY(8px)}to{opacity:1;transform:none}}
.toast.leaving{animation:toast-out .2s ease forwards}
@keyframes toast-out{to{opacity:0;transform:translateY(8px)}}
.toast a{color:var(--accent);text-decoration:underline}

/* ======= 响应式 ======= */
@media (max-width:880px){
  .topbar{padding:14px 24px}.topbar .kbd-hint{display:none}
  .hero{padding:0 24px}.hero h1{font-size:36px}
  .chapter-divider{padding:0 24px;margin:80px auto 48px}
  .capabilities{grid-template-columns:1fr;gap:48px;padding:0 24px}
  .directory{padding:0 24px}
  .directory-row{grid-template-columns:1fr;gap:6px}
  .directory-row .chapter{font-size:11px;letter-spacing:.24em}
  .directory-row .open{display:none}
  .recent{padding:0 24px}
  .recent-table .date{width:auto}
  .recent-table .chapter-cell,.recent-table .code-cell{display:none}
  .footer{padding:32px 24px;flex-direction:column;gap:14px;margin-top:80px}
}
</style></head>
<body>

<header class="topbar">
  <a class="brand-link" href="/tools" style="display:inline-flex;align-items:center;gap:10px;text-decoration:none;color:inherit;margin-right:24px;padding:6px 0"><svg class="brand-mark" viewBox="0 0 24 24" fill="currentColor" width="24" height="24" aria-hidden="true" style="color:var(--accent);opacity:1;filter:drop-shadow(0 1px 3px rgba(212,103,74,.3))"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span class="brand" style="font-family:'Noto Serif SC',Georgia,serif;font-size:19px;font-weight:600;letter-spacing:.18em;color:var(--ink)">天枢<span class="sep" style="color:var(--accent);margin:0 6px;font-weight:400">·</span>裁决</span></a>
  <nav>
    <a href="/tools" class="active">开始</a>
    <a href="/catalog">工具</a>
    <a href="/reports">报告</a>
    <a href="/settings" class="superadmin-only" style="display:none">设置</a>
    <a href="/usage" class="admin-only" style="display:none">使用统计</a>
    <a href="/admin/users" class="admin-only" style="display:none">用户管理</a>
  </nav>
  <span class="kbd-hint" id="cmd-trigger" title="打开命令面板"><kbd>⌘</kbd><kbd>K</kbd></span>
</header>
<!-- 顶部用户徽标 + 登出 + admin-only 链接显示由 _SHARED_OVERLAY_SNIPPET 注入,无需在此页脚本里再做 -->

<div class="env-banner" id="env-banner">
  <h3><span class="icon">!</span><span id="env-title">环境检查</span></h3>
  <div id="env-body"></div>
  <div class="actions" id="env-actions"></div>
  <pre class="env-progress" id="env-progress"></pre>
</div>

<!-- HERO -->
<section class="hero">
  <div class="eyebrow">天枢 · 裁决 ── AI Verdict Manual · Edition 0.1</div>
  <h1>AI 驱动的<br>软件<em>质量裁决</em>手册</h1>
  <p class="lede">把软件质量把关交给 AI 智能体 —— 覆盖研发各环节,逐项查验,出具可分派、可追溯的裁决报告。</p>
  <div class="actions">
    <a class="btn primary" href="/catalog">进入工具 <span class="arrow">→</span></a>
    <a class="btn" href="/reports">查看报告</a>
    <a class="btn" href="/guide">使用说明</a>
  </div>
  <div class="hero-foot">
    <span class="dot"></span>
    <span>贯穿研发质量全流程</span>
    <span class="claude-status" id="claude-status">检测中…</span>
  </div>
</section>

<!-- Footer -->
<footer class="footer">
  <div class="left"><span class="brand">天枢 · 裁决</span> v0.1.0</div>
  <div class="right">
    <span>Claude 驱动</span>
    <span>· 二〇二六</span>
  </div>
</footer>


<aside class="task-fab" id="task-fab" aria-label="运行中的任务">
  <div class="head">
    <span class="dot"></span><span>运行中</span><span class="count" id="task-count">0</span>
  </div>
  <div class="body" id="task-list"></div>
</aside>

<div class="cmd-overlay" id="cmd-overlay">
  <div class="cmd-panel">
    <input class="cmd-input" id="cmd-input" placeholder="搜索工具 / 报告 / 设置…" autocomplete="off">
    <div class="cmd-results" id="cmd-results"></div>
  </div>
</div>

<div class="toast-stack" id="toast-stack"></div>

<script>
const PHASE_MAP = {step1:'pre', step2:'pre', step4:'post', step5:'post', step6:'post',
  network_resilience:'post', h5_adapt:'post', seo_audit:'post'};
const CN_CH = ['一','二','三','四','五','六','七','八','九','十'];
let tools = [];
let chapterMap = {};

async function load() {
  try {
    const cat = await fetch('/api/tools').then(r => r.json());
    tools = cat.tools || [];
    checkClaude();
    pollRuns();
    setInterval(pollRuns, 4000);
    checkEnvironment();
  } catch(e) {
    console.error('load failed', e);
  }
}

function renderDirectory() {
  // 按 catalog 顺序分配章节号一至八
  tools.forEach((t, i) => { chapterMap[t.id] = CN_CH[i] || String(i+1); });

  const pre = tools.filter(t => PHASE_MAP[t.id] === 'pre');
  const post = tools.filter(t => PHASE_MAP[t.id] !== 'pre');
  document.getElementById('dir-pre-rows').innerHTML = pre.map(t => rowHtml(t)).join('');
  document.getElementById('dir-post-rows').innerHTML = post.map(t => rowHtml(t)).join('');
  document.querySelectorAll('.directory-row').forEach(r => {
    const go = () => { location.href = '/tools/' + r.dataset.tid; };
    r.onclick = go;
    r.addEventListener('keydown', ev => {
      if (ev.key === 'Enter' || ev.key === ' ') { ev.preventDefault(); go(); }
    });
  });
}

function rowHtml(t){
  const aria = `进入 ${t.name}`;
  return `<div class="directory-row" data-tid="${t.id}" role="link" tabindex="0" aria-label="${escapeHtml(aria)}">
    <div>
      <div class="name">${escapeHtml(t.name)}</div>
      <div class="desc">${escapeHtml(t.tagline || t.description || '')}</div>
    </div>
    <div class="open" aria-hidden="true">→</div>
  </div>`;
}

async function checkClaude() {
  try {
    const ci = await fetch('/api/claude/info').then(r => r.json());
    const el = document.getElementById('claude-status');
    if (ci.bin_found) {
      const auth = ci.auth_state || '';
      const acc = ci.account || {};
      if (auth === 'toolkit_api_key') {
        el.textContent = '· API Key · ' + (acc.display_name || '');
      } else if (auth === 'toolkit_oauth') {
        el.textContent = '· OAuth' + (acc.email ? ' · ' + acc.email : '');
      } else {
        el.textContent = '· 未接入 · 请到设置选择连接方式';
        el.style.color = 'var(--accent)';
      }
    } else {
      el.textContent = '· 未找到 Claude CLI';
      el.style.color = 'var(--accent)';
    }
  } catch(e) {}
}

async function loadRecent() {
  try {
    const d = await fetch('/api/reports').then(r => r.json());
    // 同一个 run_id 可能 in_memory + saved 双重出现 — 先按 run_id 去重,
    // 留下「有真实 status 的内存条目」优先;避免最近报告出现两行同 run_id。
    const seen = new Set();
    const all = [];
    (d.in_memory||[]).forEach(r => { if (r.run_id && !seen.has(r.run_id)){ seen.add(r.run_id); all.push(r); } });
    (d.saved||[]).forEach(r => { if (r.run_id && !seen.has(r.run_id)){ seen.add(r.run_id); all.push(r); } });
    all.sort((a,b) => (b.mtime || b.started_at || 0) - (a.mtime || a.started_at || 0));
    const top = all.slice(0, 6);
    const tbody = document.getElementById('recent-rows');
    if (!top.length) {
      tbody.innerHTML = '<tr><td colspan="5"><div class="recent-empty">尚无报告 · 选一章开始</div></td></tr>';
      document.getElementById('recent-month').textContent = '';
      return;
    }
    // 头部用最新一条的月份
    const newest = top[0];
    const t = new Date((newest.mtime || newest.started_at) * 1000);
    const yMap = {0:'〇',1:'一',2:'二',3:'三',4:'四',5:'五',6:'六',7:'七',8:'八',9:'九'};
    const yStr = String(t.getFullYear()).split('').map(c => yMap[c]||c).join('');
    const monthZh = ['一','二','三','四','五','六','七','八','九','十','十一','十二'][t.getMonth()];
    document.getElementById('recent-month').textContent = `${yStr} ── ${monthZh}月`;

    // 先合并：同一个 run_id 在 in_memory + saved 都出现时，以 in_memory 的真实状态为准
    const memIndex = {};
    (d.in_memory||[]).forEach(m => { if (m.run_id) memIndex[m.run_id] = m; });
    tbody.innerHTML = top.map(r => {
      const tt = new Date((r.mtime || r.started_at) * 1000);
      const dd = String(tt.getDate()).padStart(2,'0');
      const hh = String(tt.getHours()).padStart(2,'0');
      const mm = String(tt.getMinutes()).padStart(2,'0');
      const tm = tools.find(x => x.id === r.tool_id);
      const ch = chapterMap[r.tool_id] || '?';
      const chName = tm ? tm.name : r.tool_id;
      const pc = r.project_code || '—';
      const pn = r.project_name || '—';
      // 状态来源：内存里有真实 status 优先；磁盘 only 才说"已存"
      // 之前用 r.kind === 'memory' 取，但 /api/reports 返回的字段是 r.source — 永远是 false
      const mem = memIndex[r.run_id];
      let status;
      if (mem) {
        status = mem.status || '';
      } else if (r.source === 'memory') {
        status = r.status || '';
      } else {
        status = 'saved';
      }
      const sCls = (status === 'succeeded' || status === 'saved') ? 'ok'
                 : (status === 'failed') ? 'bad'
                 : (status === 'running' || status === 'queued') ? 'warn'
                 : '';
      const sLabel = {succeeded:'成 功',failed:'失 败',running:'运行中',queued:'排 队',saved:'已 存'}[status] || (status || '—');
      return `<tr onclick="location.href='/tools/${r.tool_id}?run=${r.run_id}'" tabindex="0" role="link" aria-label="${escapeHtml(chName)} · ${escapeHtml(pn)} · ${sLabel}">
        <td class="date">${dd}日 ${hh}:${mm}</td>
        <td class="chapter-cell">${escapeHtml(chName)}</td>
        <td class="code-cell">${escapeHtml(pc)}</td>
        <td class="name-cell">${escapeHtml(pn)}</td>
        <td class="status-cell ${sCls}">${sLabel}</td>
      </tr>`;
    }).join('');
    // 行键盘可达：Enter / Space 触发同一个跳转
    tbody.querySelectorAll('tr[onclick]').forEach(tr => {
      tr.addEventListener('keydown', ev => {
        if (ev.key === 'Enter' || ev.key === ' ') { ev.preventDefault(); tr.click(); }
      });
    });
  } catch(e) {}
}

async function pollRuns() {
  try {
    const data = await fetch('/api/tools/runs').then(r => r.json());
    const runs = (data.recent || []).filter(r => r.status === 'running' || r.status === 'queued');
    const fab = document.getElementById('task-fab');
    if (!runs.length) { fab.classList.remove('show'); return; }
    document.getElementById('task-count').textContent = runs.length;
    document.getElementById('task-list').innerHTML = runs.map(t => {
      const stop = t.can_stop
        ? `<button class="task-stop" data-run="${t.run_id}" title="停止执行（仅执行人/管理员）">■ 停止</button>`
        : '';
      return `<div class="row">
        <a class="row-link" href="/tools/${t.tool_id}?run=${t.run_id}">
          <div class="title">${escapeHtml(t.tool_name)}</div>
          <div class="progress">${escapeHtml(t.progress || t.status)}</div>
        </a>
        ${stop}
      </div>`;
    }).join('');
    document.querySelectorAll('#task-list .task-stop').forEach(b => {
      b.onclick = async (e) => {
        e.preventDefault(); e.stopPropagation();
        if (!confirm('确定停止这个运行？已产生的进度会丢失。')) return;
        b.disabled = true; b.textContent = '停止中…';
        try {
          const r = await fetch(`/api/tools/runs/${b.dataset.run}/stop`, {method:'POST'}).then(r => r.json());
          if (r.ok) { pollRuns(); }
          else { alert(r.detail || '停止失败'); b.disabled = false; b.textContent = '■ 停止'; }
        } catch(err) { alert('停止失败：' + err); b.disabled = false; b.textContent = '■ 停止'; }
      };
    });
    fab.classList.add('show');
  } catch(e) {}
}

async function checkEnvironment() {
  try {
    const r = await fetch('/api/settings/ready').then(r => r.json());
    const allMissing = [...(r.missing_required||[]), ...(r.missing_optional||[])];
    if (!allMissing.length) return;
    const hasAuthIssue = (r.missing_required||[]).some(m => m.pkg === 'auth_mode' || m.pkg === 'claude_api_key');
    if (!hasAuthIssue && sessionStorage.getItem('env-skip-once')) return;

    const banner = document.getElementById('env-banner');
    banner.classList.add('show');
    document.getElementById('env-title').textContent =
      hasAuthIssue ? '需要先选择连接方式' : '环境检查';
    document.getElementById('env-body').innerHTML = allMissing.map(m =>
      `<div style="font-size:12.5px;padding:4px 0">— ${escapeHtml(m.label || m.pkg)}:<span style="color:var(--ink-3)">${escapeHtml(m.purpose || '')}</span></div>`
    ).join('');
    document.getElementById('env-actions').innerHTML = hasAuthIssue
      ? `<button onclick="location.href='/settings'">前往设置 →</button>`
      : `<button onclick="sessionStorage.setItem('env-skip-once','1'); document.getElementById('env-banner').classList.remove('show')">本次跳过</button>`;
  } catch(e) {}
}

// 命令面板
const cmdOverlay = document.getElementById('cmd-overlay');
const cmdInput = document.getElementById('cmd-input');
const cmdResults = document.getElementById('cmd-results');
let cmdActive = 0;
function openCmd(){ cmdOverlay.classList.add('open'); cmdInput.value=''; cmdInput.focus(); renderCmd(''); }
function closeCmd(){ cmdOverlay.classList.remove('open'); }
function cmdItems(){
  return [
    ...tools.map(t => ({name:`${t.name}`, hint:t.id, href:`/tools/${t.id}`})),
    {name:'报告索引', hint:'reports', href:'/reports'},
    {name:'设置', hint:'settings', href:'/settings'},
  ];
}
function renderCmd(q){
  const items = cmdItems().filter(x => !q || (x.name+x.hint).toLowerCase().includes(q.toLowerCase()));
  cmdActive = 0;
  cmdResults.innerHTML = items.length
    ? items.map((x,i) => `<div class="cmd-result ${i===0?'active':''}" data-href="${x.href}" data-idx="${i}">
        <span class="label">${escapeHtml(x.name)}</span>
        <span class="hint">${escapeHtml(x.hint)}</span>
      </div>`).join('')
    : `<div style="padding:20px;text-align:center;color:var(--ink-3);font-size:12.5px;font-family:var(--serif)">无匹配</div>`;
  cmdResults.querySelectorAll('.cmd-result').forEach(el => {
    el.onclick = () => { location.href = el.dataset.href; };
  });
}
document.addEventListener('keydown', e => {
  if ((e.metaKey||e.ctrlKey) && e.key === 'k'){ e.preventDefault(); openCmd(); }
  else if (e.key === 'Escape' && cmdOverlay.classList.contains('open')){ closeCmd(); }
  else if (cmdOverlay.classList.contains('open')){
    const rs = cmdResults.querySelectorAll('.cmd-result');
    if (e.key === 'ArrowDown'){ e.preventDefault(); cmdActive = Math.min(rs.length-1, cmdActive+1); rs.forEach((el,i)=>el.classList.toggle('active', i===cmdActive)); }
    else if (e.key === 'ArrowUp'){ e.preventDefault(); cmdActive = Math.max(0, cmdActive-1); rs.forEach((el,i)=>el.classList.toggle('active', i===cmdActive)); }
    else if (e.key === 'Enter'){ const el = rs[cmdActive]; if (el) location.href = el.dataset.href; }
  }
});
cmdInput.addEventListener('input', () => renderCmd(cmdInput.value));
cmdOverlay.addEventListener('click', e => { if (e.target === cmdOverlay) closeCmd(); });
document.getElementById('cmd-trigger').addEventListener('click', openCmd);

function toast(msg){
  const t = document.createElement('div');
  t.className = 'toast'; t.textContent = msg;
  document.getElementById('toast-stack').appendChild(t);
  setTimeout(() => { t.classList.add('leaving'); setTimeout(() => t.remove(), 220); }, 3000);
}

function escapeHtml(s){ return String(s == null ? '' : s).replace(/[&<>]/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;'}[c])); }


load();
</script>
</body></html>
"""


TOOL_DETAIL_HTML_OLD_DEPRECATED = """<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>工具 — 天枢·裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin><link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
  :root{--bg:#07080a;--surface:#0f1216;--surface-2:#161a1f;--surface-3:#1d2228;
    --line:#1f2530;--line-2:#2a3140;
    --fg:#f5f7fa;--fg-2:#a8aeb8;--fg-3:#6c7380;--fg-4:#4a5060;
    --ac:#10b981;--ac-2:#6ee7b7;
    --warn:#fbbf24;--ok:#4ade80;--bad:#f87171;
    --mono:ui-monospace,SFMono-Regular,"JetBrains Mono",Menlo,monospace;
    --sans:-apple-system,BlinkMacSystemFont,"SF Pro Display","SF Pro Text","PingFang SC","Helvetica Neue",Arial,sans-serif;}
  *{box-sizing:border-box}
  html,body{margin:0;background:var(--bg);color:var(--fg);font-family:var(--sans);height:100%;overflow:hidden}

  /* === Top bar (always visible) === */
  .topbar{display:flex;align-items:center;gap:0;height:48px;
    padding:0 16px;border-bottom:1px solid var(--line);
    background:rgba(7,8,10,.92);position:relative;z-index:30}
  .topbar .logo{width:24px;height:24px;border:1.5px solid var(--ac);border-radius:6px;
    display:grid;place-items:center;color:var(--ac);font-family:var(--mono);font-weight:700;
    font-size:11px;margin-right:14px;flex-shrink:0}
  .topbar .crumbs{display:flex;align-items:center;gap:8px;font-size:13px}
  .topbar .crumbs a{color:var(--fg-3);text-decoration:none;transition:color .15s}
  .topbar .crumbs a:hover{color:var(--ac)}
  .topbar .crumbs .sep{color:var(--fg-4)}
  .topbar .crumbs .current{color:var(--fg);font-weight:500;display:flex;align-items:center;gap:6px}
  .topbar .crumbs .icon{font-size:15px}
  .topbar .stats{margin-left:auto;display:flex;gap:14px;align-items:center;
    font-family:var(--mono);font-size:11px;color:var(--fg-3)}
  .topbar .stats{display:flex;gap:14px;align-items:center;
    font-family:"SF Mono",ui-monospace,monospace;font-size:11px;color:var(--fg-3)}
  .topbar .stats .stat{display:flex;align-items:center;gap:5px}
  .topbar .stats .stat .lbl{color:var(--fg-4);font-size:10px;text-transform:uppercase;letter-spacing:.05em}
  .topbar .stats .stat .v{color:var(--fg-2)}
  .topbar .stats .stat.cost .v{color:var(--ac)}
  .topbar .stats .stat.gate.reject .v{color:var(--bad)}
  .topbar .stats .stat.gate.proceed .v{color:var(--ok)}
  .topbar .stats .stat.gate.warn .v{color:var(--warn)}
  .topbar .actions{display:flex;gap:8px;align-items:center;margin-left:14px}
  .topbar .icon-btn{background:transparent;border:1px solid var(--line);color:var(--fg-2);
    width:32px;height:30px;border-radius:6px;cursor:pointer;font-size:13px;
    display:grid;place-items:center;transition:all .15s}
  .topbar .icon-btn:hover{border-color:var(--ac);color:var(--ac)}
  .topbar .icon-btn.active{background:rgba(16,185,129,.10);border-color:var(--ac);color:var(--ac)}
  .topbar button.run{background:var(--ac);color:#001a14;border:none;
    padding:0 18px;height:30px;border-radius:6px;font-family:var(--sans);font-size:13px;
    font-weight:600;cursor:pointer;display:flex;align-items:center;gap:6px;transition:all .15s}
  .topbar button.run:disabled{background:var(--line-2);color:var(--fg-3);cursor:not-allowed}
  .topbar button.run:hover:not(:disabled){background:var(--ac-2)}
  .topbar button.run kbd{background:rgba(0,26,20,.18);padding:1px 5px;border-radius:3px;
    font-family:var(--mono);font-size:10px;font-weight:600;margin-left:4px}

  /* === Workspace === */
  .workspace{display:grid;grid-template-columns:1fr 1px 1fr;height:calc(100vh - 48px);overflow:hidden}
  .pane{display:flex;flex-direction:column;height:100%;overflow:hidden}
  .pane.input{background:var(--bg)}
  .pane.output{background:var(--surface)}
  .divider{background:var(--line);width:1px}

  .pane-head{display:flex;align-items:center;height:38px;padding:0 18px;
    border-bottom:1px solid var(--line);font-size:11.5px;color:var(--fg-3);
    text-transform:uppercase;letter-spacing:.12em;font-weight:600;flex-shrink:0;
    gap:14px}
  .pane-head .label{display:flex;align-items:center;gap:6px}
  .pane-head .dot{width:6px;height:6px;border-radius:50%;background:var(--fg-4)}
  .pane-head .dot.live{background:var(--ac);animation:pulse 1.6s ease-in-out infinite}
  @keyframes pulse{0%,100%{opacity:1}50%{opacity:.35}}
  .pane-head .right{margin-left:auto;display:flex;gap:6px;font-family:var(--mono);
    font-size:11px;color:var(--fg-3);text-transform:none;letter-spacing:0;font-weight:400}

  /* Form */
  .pane-body{flex:1;overflow-y:auto;padding:20px 24px}
  form{display:flex;flex-direction:column;gap:14px}
  .field{display:flex;flex-direction:column;gap:5px}
  .field-head{display:flex;align-items:center;gap:8px;font-size:12.5px}
  .field label{color:var(--fg);font-weight:500}
  .field .req{color:var(--bad)}
  .field .opt{color:var(--fg-4);font-family:var(--mono);font-size:11px;font-weight:400}
  .field .upload-btn{margin-left:auto;background:transparent;border:1px solid var(--line);
    color:var(--fg-3);padding:3px 10px;border-radius:5px;font-family:var(--mono);
    font-size:11px;cursor:pointer;transition:all .15s}
  .field .upload-btn:hover{border-color:var(--ac);color:var(--ac)}
  .field textarea, .field input[type=text]{
    background:var(--surface);border:1px solid var(--line);border-radius:6px;
    color:var(--fg);font-family:var(--mono);font-size:12.5px;padding:10px 12px;
    resize:vertical;min-height:64px;line-height:1.55;width:100%;transition:all .15s}
  .field textarea:focus, .field input[type=text]:focus{
    outline:none;border-color:var(--ac);background:var(--surface-2)}
  .field textarea::placeholder{color:var(--fg-4)}
  .field input[type=file]{display:none}
  .field .checkbox-row{display:flex;align-items:center;gap:10px;
    background:var(--surface);border:1px solid var(--line);border-radius:6px;
    padding:9px 12px;font-size:12.5px}
  .field .checkbox-row input{width:15px;height:15px;accent-color:var(--ac);cursor:pointer}
  .field .checkbox-row label{cursor:pointer}

  /* Output pane */
  .out-empty{display:flex;flex-direction:column;align-items:center;justify-content:center;
    text-align:center;padding:60px 24px;color:var(--fg-3);font-size:13px;height:100%;
    font-family:var(--mono)}
  .out-empty .ascii{font-size:11px;color:var(--fg-4);line-height:1.4;margin-bottom:14px;
    white-space:pre}
  .step-list{display:flex;flex-direction:column;gap:0;padding:0;border-bottom:1px solid var(--line)}
  .step-row{display:grid;grid-template-columns:24px 1fr auto;gap:10px;padding:10px 18px;
    border-bottom:1px solid var(--line);align-items:center;font-size:12.5px;
    transition:background .15s}
  .step-row:last-child{border-bottom:none}
  .step-row .marker{font-family:var(--mono);font-size:11px;color:var(--fg-4);text-align:center}
  .step-row.done .marker{color:var(--ok)}
  .step-row.running .marker{color:var(--ac);animation:pulse 1.4s infinite}
  .step-row.failed .marker{color:var(--bad)}
  .step-row .name{color:var(--fg-2)}
  .step-row.done .name{color:var(--fg)}
  .step-row .info{font-family:var(--mono);font-size:11px;color:var(--fg-3)}

  .gate-banner{padding:14px 20px;border-bottom:1px solid var(--line);display:flex;
    align-items:flex-start;gap:12px;font-size:13px;line-height:1.55}
  .gate-banner.proceed{background:rgba(74,222,128,.05);border-bottom-color:rgba(74,222,128,.2)}
  .gate-banner.reject{background:rgba(248,113,113,.05);border-bottom-color:rgba(248,113,113,.2)}
  .gate-banner.warn{background:rgba(251,191,36,.05);border-bottom-color:rgba(251,191,36,.2)}
  .gate-banner .badge{padding:3px 9px;border-radius:4px;font-family:var(--mono);font-size:11px;
    font-weight:600;text-transform:uppercase;letter-spacing:.05em;flex-shrink:0;margin-top:1px}
  .gate-banner.proceed .badge{background:rgba(74,222,128,.12);color:var(--ok)}
  .gate-banner.reject .badge{background:rgba(248,113,113,.12);color:var(--bad)}
  .gate-banner.warn .badge{background:rgba(251,191,36,.12);color:var(--warn)}
  .gate-banner .reasons{color:var(--fg-2);font-family:var(--mono);font-size:12px}
  .gate-banner .reasons div{margin-top:3px}

  .json-view{flex:1;overflow:auto;padding:14px 18px;background:var(--bg)}
  .json-view pre{margin:0;font-family:var(--mono);font-size:12px;line-height:1.65;
    color:var(--fg);white-space:pre-wrap;word-break:break-word}
  .err-view{padding:14px 18px;font-family:var(--mono);font-size:12px;line-height:1.6;
    color:var(--bad);white-space:pre-wrap;word-break:break-word;flex:1;overflow:auto}

  /* === Drawer (right side, slides in) === */
  .drawer{position:fixed;top:48px;right:-460px;width:460px;height:calc(100vh - 48px);
    background:var(--surface-2);border-left:1px solid var(--line);
    transition:right .25s ease;z-index:25;display:flex;flex-direction:column;
    box-shadow:-8px 0 24px rgba(0,0,0,.4)}
  .drawer.open{right:0}
  .drawer-tabs{display:flex;border-bottom:1px solid var(--line);flex-shrink:0}
  .drawer-tabs button{flex:1;background:transparent;border:none;color:var(--fg-3);
    padding:12px 16px;font-size:12px;cursor:pointer;border-bottom:2px solid transparent;
    transition:all .15s;font-family:var(--sans);font-weight:500}
  .drawer-tabs button:hover{color:var(--fg-2)}
  .drawer-tabs button.active{color:var(--ac);border-bottom-color:var(--ac)}
  .drawer-close{position:absolute;top:8px;right:10px;background:transparent;
    border:none;color:var(--fg-3);font-size:18px;cursor:pointer;width:24px;height:24px;
    border-radius:4px;display:grid;place-items:center}
  .drawer-close:hover{background:var(--surface);color:var(--fg)}
  .drawer-panel{flex:1;overflow-y:auto;display:none;padding:18px 20px}
  .drawer-panel.active{display:block}
  .drawer-panel h4{margin:0 0 4px;font-size:13px;font-weight:600;color:var(--fg);letter-spacing:-.005em}
  .drawer-panel .sub-id{font-family:var(--mono);font-size:11px;color:var(--ac);margin-bottom:8px}
  .drawer-panel .meta-row{display:flex;gap:6px;flex-wrap:wrap;margin-bottom:8px}
  .drawer-panel .chip{font-family:var(--mono);font-size:10px;padding:2px 7px;
    border-radius:3px;background:var(--surface);border:1px solid var(--line);color:var(--fg-3)}
  .drawer-panel .chip.ac{color:var(--ac);border-color:rgba(16,185,129,.3)}
  .drawer-panel pre.body{margin:0 0 18px;background:var(--bg);border:1px solid var(--line);
    border-radius:6px;padding:10px 12px;font-family:var(--mono);font-size:11.5px;
    line-height:1.55;color:var(--fg-2);white-space:pre-wrap;word-break:break-word;
    max-height:240px;overflow-y:auto}
  .drawer-panel .placeholder{color:var(--ac)}
  .drawer-panel hr{border:none;border-top:1px solid var(--line);margin:14px 0}
  .info-grid{display:grid;grid-template-columns:max-content 1fr;gap:7px 14px;
    font-size:12px;margin-bottom:18px}
  .info-grid dt{color:var(--fg-3)}
  .info-grid dd{margin:0;color:var(--fg);font-family:var(--mono);font-size:11.5px}

  /* Status pill in topbar */
  .status-pill{display:inline-flex;align-items:center;gap:6px;font-family:var(--mono);
    font-size:11px;padding:2px 9px;border-radius:3px;font-weight:600;text-transform:uppercase;
    letter-spacing:.05em}
  .status-pill.queued{background:rgba(168,174,184,.10);color:var(--fg-2)}
  .status-pill.running{background:rgba(251,191,36,.10);color:var(--warn)}
  .status-pill.succeeded{background:rgba(74,222,128,.10);color:var(--ok)}
  .status-pill.failed{background:rgba(248,113,113,.10);color:var(--bad)}

  @media(max-width:880px){
    .workspace{grid-template-columns:1fr;grid-template-rows:1fr 1px 1fr}
    .divider{height:1px;width:auto}
    .drawer{width:100%;right:-100%}
  }
</style></head>
<body>
<div class="topbar">
  <a class="brand-link" href="/tools" title="天枢 · 裁决 · 返回主页" style="text-decoration:none;display:inline-flex;align-items:center;gap:10px;font-family:'Noto Serif SC',Georgia,serif;font-size:19px;font-weight:600;letter-spacing:.18em;color:var(--fg);margin-right:24px;padding:4px 0"><svg viewBox="0 0 24 24" fill="currentColor" width="24" height="24" aria-hidden="true" style="color:var(--ac);opacity:1;flex-shrink:0;filter:drop-shadow(0 1px 2px rgba(196,90,58,.28))"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span>天枢</span><span style="color:var(--ac);margin:0 6px;font-weight:400">·</span><span>裁决</span></a>
  <div class="crumbs">
    <a href="/tools">工具</a>
    <span class="sep">/</span>
    <span class="current"><span class="icon" id="tb-icon">·</span><span id="tb-name">…</span></span>
    <span class="sep" id="tb-sep" style="display:none">·</span>
    <span id="tb-status"></span>
  </div>
  <div class="stats">
    <span class="stat"><span>tokens in</span><span class="v" id="s-in">—</span></span>
    <span class="stat"><span>out</span><span class="v" id="s-out">—</span></span>
    <span class="stat"><span>cache</span><span class="v" id="s-cache">—</span></span>
    <span class="stat cost"><span>$</span><span class="v" id="s-cost">—</span></span>
    <span class="stat" id="s-elapsed-wrap" style="display:none"><span>·</span><span class="v" id="s-elapsed"></span></span>
    <span class="stat gate" id="s-gate-wrap" style="display:none"><span>gate</span><span class="v" id="s-gate"></span></span>
  </div>
  <div class="actions">
    <button class="icon-btn" id="btn-info" title="工具信息">ⓘ</button>
    <button class="icon-btn" id="btn-prompts" title="提示词">P</button>
    <button class="icon-btn" id="btn-history" title="历史">H</button>
    <button class="run" id="btn-run">▶ 运行<kbd>⌘↵</kbd></button>
  </div>
</div>

<div class="workspace">
  <div class="pane input">
    <div class="pane-head">
      <span class="label"><span class="dot"></span>输入</span>
      <span class="right" id="input-meta"></span>
    </div>
    <div class="pane-body">
      <form id="form"></form>
    </div>
  </div>
  <div class="divider"></div>
  <div class="pane output">
    <div class="pane-head">
      <span class="label"><span class="dot" id="out-dot"></span>输出</span>
      <span class="right" id="output-meta"></span>
    </div>
    <div id="output-area" style="flex:1;display:flex;flex-direction:column;overflow:hidden">
      <div class="out-empty">
        <div class="ascii">┌──────────────┐
│   no run     │
│   yet        │
└──────────────┘</div>
        <div>填好左侧表单后按 <strong>⌘↵</strong> 或 <strong>▶ 运行</strong></div>
      </div>
    </div>
  </div>
</div>

<!-- Drawer (info / prompts / history) -->
<div class="drawer" id="drawer">
  <button class="drawer-close" onclick="closeDrawer()">×</button>
  <div class="drawer-tabs">
    <button data-tab="info" class="active">信息</button>
    <button data-tab="prompts">提示词</button>
    <button data-tab="history">历史</button>
  </div>
  <div class="drawer-panel active" data-panel="info" id="panel-info"></div>
  <div class="drawer-panel" data-panel="prompts" id="panel-prompts"></div>
  <div class="drawer-panel" data-panel="history" id="panel-history"></div>
</div>

<script>
const TOOL_ID = location.pathname.split('/').filter(Boolean).pop();
let tool = null;
let pollTimer = null;
let elapsedTimer = null;
let currentRunId = new URLSearchParams(location.search).get('run');
let currentRun = null;

async function load(){
  tool = await fetch('/api/tools/' + TOOL_ID).then(r=>r.json());
  document.title = tool.name + ' — 天枢·裁决';
  document.getElementById('tb-icon').textContent = tool.icon;
  document.getElementById('tb-name').textContent = tool.name;
  document.getElementById('input-meta').textContent = tool.fields.length + ' 字段';
  document.getElementById('output-meta').textContent = tool.prompts.length + ' substeps · ' + tool.output.replace(/[《》]/g,'');

  buildForm();
  buildInfoPanel();
  buildPromptsPanel();
  buildHistoryPanel();
  if (currentRunId) startPolling(currentRunId);
}

// ---------- form ----------
function buildForm(){
  const form = document.getElementById('form');
  tool.fields.forEach((f,i) => {
    const wrap = document.createElement('div');
    wrap.className = 'field';
    if (f.type === 'checkbox'){
      wrap.innerHTML = `
        <div class="checkbox-row">
          <input type="checkbox" id="f-${f.key}" ${f.default ? 'checked' : ''}>
          <label for="f-${f.key}">${f.label}</label>
        </div>`;
    } else {
      const isJson = f.type === 'json-file' || f.type === 'json-text';
      const placeholder = isJson ? '粘贴 JSON 或上传 .json' : '粘贴文本或上传文件';
      const minH = isJson ? 'min-height:90px' : 'min-height:64px';
      wrap.innerHTML = `
        <div class="field-head">
          <label for="f-${f.key}">${f.label}</label>
          ${f.required ? '<span class="req">*</span>' : '<span class="opt">optional</span>'}
          <button type="button" class="upload-btn" data-target="f-${f.key}" data-accept="${f.accept || '*'}">↑ 上传</button>
        </div>
        <textarea id="f-${f.key}" placeholder="${placeholder}" style="${minH}">${f.default || ''}</textarea>
      `;
    }
    form.appendChild(wrap);
  });
  // upload-btn handlers
  form.querySelectorAll('.upload-btn').forEach(btn => {
    btn.onclick = () => {
      const inp = document.createElement('input');
      inp.type = 'file';
      inp.accept = btn.dataset.accept;
      inp.onchange = async e => {
        const file = inp.files[0];
        if (!file) return;
        document.getElementById(btn.dataset.target).value = await file.text();
      };
      inp.click();
    };
  });
}

async function runTool(){
  if (!tool) return;
  const body = {};
  for (const f of tool.fields){
    const el = document.getElementById('f-' + f.key);
    if (!el) continue;
    if (f.type === 'checkbox'){ body[f.key] = el.checked; continue; }
    const v = el.value.trim();
    if (!v && !f.required) continue;
    if (f.type === 'json-file' || f.type === 'json-text'){
      try { body[f.key] = JSON.parse(v); }
      catch(err){ alert(`字段「${f.label}」不是合法 JSON：${err.message}`); return; }
    } else {
      if (v) body[f.key] = v;
    }
  }
  const btn = document.getElementById('btn-run');
  btn.disabled = true; btn.innerHTML = '提交中…';
  try {
    const res = await fetch(`/api/tools/${tool.id}/run`, {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(body),
    }).then(r => r.json());
    if (!res.run_id){ alert('启动失败：' + JSON.stringify(res)); return; }
    currentRunId = res.run_id;
    history.replaceState(null, '', `?run=${res.run_id}`);
    startPolling(res.run_id);
  } finally {
    btn.disabled = false; btn.innerHTML = '▶ 运行<kbd>⌘↵</kbd>';
  }
}

document.getElementById('btn-run').onclick = runTool;
document.addEventListener('keydown', e => {
  if ((e.metaKey || e.ctrlKey) && e.key === 'Enter'){ e.preventDefault(); runTool(); }
});
// 右侧 56px 窄条上的 ▶ 圆按钮也能触发运行
document.addEventListener('click', e => {
  const t = e.target;
  if (t && (t.id === 'play-circle' || (t.closest && t.closest('#play-circle')))) {
    runTool();
  }
});

// ---------- polling & rendering ----------
function startPolling(runId){
  if (pollTimer) clearInterval(pollTimer);
  if (elapsedTimer) clearInterval(elapsedTimer);
  document.getElementById('out-dot').classList.add('live');
  poll(runId);
  pollTimer = setInterval(()=>poll(runId), 3000);
  elapsedTimer = setInterval(updateElapsed, 1000);
}

async function poll(runId){
  try {
    const r = await fetch(`/api/tools/runs/${runId}`).then(r => r.json());
    currentRun = r;
    renderRun(r);
    if (r.status === 'succeeded' || r.status === 'failed'){
      clearInterval(pollTimer); pollTimer = null;
      clearInterval(elapsedTimer); elapsedTimer = null;
      document.getElementById('out-dot').classList.remove('live');
      buildHistoryPanel();
    }
  } catch(e){
    clearInterval(pollTimer); pollTimer = null;
    clearInterval(elapsedTimer); elapsedTimer = null;
  }
}

function updateElapsed(){
  if (!currentRun || currentRun.finished_at) return;
  const e = Date.now()/1000 - currentRun.started_at;
  document.getElementById('s-elapsed').textContent = e.toFixed(1) + 's';
  document.getElementById('s-elapsed-wrap').style.display = '';
}

function renderRun(r){
  // top-bar status pill
  const tb = document.getElementById('tb-status');
  document.getElementById('tb-sep').style.display = '';
  tb.innerHTML = `<span class="status-pill ${r.status}">${r.status}</span>`;

  // top-bar usage
  const u = r.usage || {};
  document.getElementById('s-in').textContent = u.input_tokens != null ? u.input_tokens.toLocaleString() : '—';
  document.getElementById('s-out').textContent = u.output_tokens != null ? u.output_tokens.toLocaleString() : '—';
  document.getElementById('s-cache').textContent = u.cache_read_tokens != null ? u.cache_read_tokens.toLocaleString() : '—';
  document.getElementById('s-cost').textContent = u.cost_usd != null ? u.cost_usd.toFixed(4) : '—';
  if (r.finished_at){
    document.getElementById('s-elapsed').textContent = (r.finished_at - r.started_at).toFixed(1) + 's';
    document.getElementById('s-elapsed-wrap').style.display = '';
  }

  // gate
  if (r.report && r.report.gate_decision){
    const g = r.report.gate_decision;
    const cls = gateClass(g.action);
    const wrap = document.getElementById('s-gate-wrap');
    wrap.style.display = '';
    wrap.classList.remove('reject','proceed','warn');
    wrap.classList.add(cls);
    document.getElementById('s-gate').textContent = g.action.replace(/^GateAction\\./,'').toLowerCase();
  }

  // output area
  renderOutput(r);
}

function gateClass(action){
  const a = String(action || '').toLowerCase();
  if (a.includes('reject')) return 'reject';
  if (a.includes('proceed') || a.includes('approve')) return 'proceed';
  return 'warn';
}

function renderOutput(r){
  const area = document.getElementById('output-area');
  if (r.status === 'failed'){
    area.innerHTML = `<div class="err-view">${escapeHtml(r.traceback || r.error || 'failed')}</div>`;
    return;
  }
  // build progress steplist (synthetic — based on prompts)
  const stepHtml = renderSteps(r);
  let gateHtml = '';
  if (r.report && r.report.gate_decision){
    const g = r.report.gate_decision;
    const cls = gateClass(g.action);
    const reasons = (g.reasons || []).map(x => `<div>· ${escapeHtml(x)}</div>`).join('');
    gateHtml = `<div class="gate-banner ${cls}"><span class="badge">${g.action.replace(/^GateAction\\./,'').toLowerCase()}</span><div class="reasons"><strong>${escapeHtml(g.action)}</strong>${reasons}</div></div>`;
  }
  let jsonHtml = '';
  if (r.report){
    jsonHtml = `<div class="json-view"><pre>${escapeHtml(JSON.stringify(r.report, null, 2))}</pre></div>`;
  } else {
    jsonHtml = `<div class="out-empty"><div>${escapeHtml(r.progress || '运行中…')}</div></div>`;
  }
  area.innerHTML = stepHtml + gateHtml + jsonHtml;
}

function renderSteps(r){
  if (!tool || !tool.prompts) return '';
  // Heuristic: if status==succeeded → all done; if running → guess from progress text
  const done = r.status === 'succeeded';
  const failed = r.status === 'failed';
  let html = '<div class="step-list">';
  tool.prompts.forEach((p, i) => {
    let cls = '';
    let marker = '○';
    if (done){ cls = 'done'; marker = '✓'; }
    else if (failed && i === 0){ cls = 'failed'; marker = '✗'; }
    else if (r.status === 'running' && i === 0){ cls = 'running'; marker = '◔'; }
    const shortP = String(p).split('.').pop();
    html += `<div class="step-row ${cls}"><span class="marker">${marker}</span><span class="name" title="${p}">#${shortP}</span><span class="info">${cls === 'done' ? 'OK' : (cls === 'running' ? '执行中' : '')}</span></div>`;
  });
  html += '</div>';
  return html;
}

function escapeHtml(s){ return String(s).replace(/[&<>]/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;'}[c])); }

// ---------- drawer ----------
const drawer = document.getElementById('drawer');
function openDrawer(tab){
  drawer.classList.add('open');
  document.querySelectorAll('.drawer-tabs button').forEach(b => b.classList.toggle('active', b.dataset.tab===tab));
  document.querySelectorAll('.drawer-panel').forEach(p => p.classList.toggle('active', p.dataset.panel===tab));
}
function closeDrawer(){ drawer.classList.remove('open'); }
window.closeDrawer = closeDrawer;

document.getElementById('btn-info').onclick = () => openDrawer('info');
document.getElementById('btn-prompts').onclick = () => openDrawer('prompts');
document.getElementById('btn-history').onclick = () => openDrawer('history');
document.querySelectorAll('.drawer-tabs button').forEach(b => {
  b.onclick = () => openDrawer(b.dataset.tab);
});
document.addEventListener('keydown', e => {
  if (e.key === 'Escape') closeDrawer();
});

function buildInfoPanel(){
  const el = document.getElementById('panel-info');
  el.innerHTML = `
    <div style="display:flex;gap:14px;align-items:flex-start;margin-bottom:18px">
      <div style="font-size:30px">${tool.icon}</div>
      <div>
        <h4>${tool.name}</h4>
        <div class="sub-id">${tool.step}</div>
      </div>
    </div>
    <p style="font-size:13px;color:var(--fg-2);line-height:1.65;margin:0 0 18px">${tool.description}</p>
    <dl class="info-grid">
      <dt>主责</dt><dd>${tool.responsible}</dd>
      <dt>核心输出</dt><dd>${tool.output}</dd>
      <dt>提示词</dt><dd>${tool.prompts.length} × ${tool.prompts[0]}–${tool.prompts.slice(-1)[0]}</dd>
      <dt>API</dt><dd><a href="/api/tools/${tool.id}" target="_blank" style="color:var(--ac)">${tool.endpoint}</a></dd>
    </dl>
    <hr>
    <h4 style="margin-bottom:6px;color:var(--warn)">⚠ 闸门规则</h4>
    <p style="font-size:12px;color:var(--fg-2);line-height:1.6;margin:0">${tool.gate}</p>
  `;
}

async function buildPromptsPanel(){
  const el = document.getElementById('panel-prompts');
  el.innerHTML = '<div style="color:var(--fg-3);font-size:12px;padding:10px 0">加载中…</div>';
  // Use the prompt step's metadata loaded by /api/prompts
  const all = await fetch('/api/prompts').then(r=>r.json());
  const step = all.steps.find(s => s.dir_name === tool.prompt_dir);
  if (!step){ el.innerHTML = '<div>找不到提示词</div>'; return; }
  el.innerHTML = '';
  step.substeps.forEach(s => {
    const div = document.createElement('div');
    const ph = (s.placeholders || []).map(p => `<span class="chip ac">${p}</span>`).join('');
    div.innerHTML = `
      <h4>${s.name}</h4>
      <div class="sub-id">${s.id} · ${s.model_tier} · max_tokens ${s.max_tokens}</div>
      <div class="meta-row">${ph}</div>
      <pre class="body">${highlightPlaceholders(escapeHtml(s.body))}</pre>
    `;
    el.appendChild(div);
  });
}

function highlightPlaceholders(s){
  return s.replace(/\\{\\{[^}]+\\}\\}/g, m => `<span class="placeholder">${m}</span>`);
}

async function buildHistoryPanel(){
  const el = document.getElementById('panel-history');
  el.innerHTML = '<div style="color:var(--fg-3);font-size:12px">加载中…</div>';
  const data = await fetch('/api/tools/runs').then(r=>r.json());
  const mine = data.recent.filter(r => r.tool_id === tool.id);
  if (!mine.length){ el.innerHTML = '<div style="color:var(--fg-3);font-size:12px;font-family:var(--mono)">尚无运行记录</div>'; return; }
  el.innerHTML = '';
  mine.forEach(run => {
    const div = document.createElement('a');
    div.href = `?run=${run.run_id}`;
    div.style.cssText = 'display:block;text-decoration:none;color:inherit;background:var(--surface);border:1px solid var(--line);border-radius:6px;padding:10px 12px;margin-bottom:6px;transition:all .15s';
    div.onmouseover = () => div.style.borderColor = 'var(--ac)';
    div.onmouseout = () => div.style.borderColor = 'var(--line)';
    const t = new Date(run.started_at*1000).toLocaleString();
    div.innerHTML = `
      <div style="display:flex;align-items:center;gap:8px;margin-bottom:4px">
        <span class="status-pill ${run.status}">${run.status}</span>
        <span style="font-family:var(--mono);font-size:11px;color:var(--fg-3)">${t}</span>
      </div>
      <div style="font-family:var(--mono);font-size:11px;color:var(--fg-3);overflow:hidden;text-overflow:ellipsis;white-space:nowrap">${run.run_id}</div>
    `;
    el.appendChild(div);
  });
}

load();
</script>
</body></html>
"""


_SHARED_OVERLAY_SNIPPET = """
<script>
(function sharedOverlays(){
  // Idempotent: skip if the page already has the overlays
  if (document.getElementById('task-fab') && document.getElementById('cmd-overlay')) return;
  const css = `
.cmd-overlay{position:fixed;inset:0;z-index:200;background:rgba(0,0,0,.55);
  -webkit-backdrop-filter:blur(6px);backdrop-filter:blur(6px);
  display:none;align-items:flex-start;justify-content:center;padding-top:14vh}
.cmd-overlay.show{display:flex}
.cmd-panel{width:580px;max-width:90vw;background:var(--surface-2,#171b22);
  border:1px solid var(--line-2,#2a3140);border-radius:14px;
  box-shadow:0 12px 32px rgba(0,0,0,.42);overflow:hidden;
  animation:cmd-in 200ms cubic-bezier(.5,1.4,.3,.95)}
@keyframes cmd-in{from{opacity:0;transform:translateY(-12px) scale(.97)}to{opacity:1;transform:translateY(0) scale(1)}}
.cmd-input{width:100%;padding:16px;border:none;outline:none;background:transparent;
  color:var(--fg,#f5f7fa);font-size:16px;
  font-family:-apple-system,BlinkMacSystemFont,"SF Pro Display","PingFang SC",sans-serif;
  border-bottom:1px solid var(--line-1,#1e242d)}
.cmd-input::placeholder{color:var(--fg-3,#6b7380)}
.cmd-results{max-height:380px;overflow-y:auto;padding:8px}
.cmd-result{padding:10px 12px;display:flex;align-items:center;gap:12px;cursor:pointer;
  border-radius:6px;transition:background .12s}
.cmd-result.active{background:var(--surface-3,#1d222b)}
.cmd-result .icon{font-size:18px;opacity:.85}
.cmd-result .name{font-size:13px;color:var(--fg,#f5f7fa)}
.cmd-result .hint{margin-left:auto;font-family:"SF Mono",ui-monospace,monospace;
  font-size:10.5px;color:var(--fg-3,#6b7380)}
.cmd-empty{padding:24px;text-align:center;color:var(--fg-3,#6b7380);font-size:13px}
.task-fab-shared{position:fixed;bottom:24px;right:24px;z-index:90;display:none;
  background:var(--surface-2,#171b22);border:1px solid var(--line-2,#2a3140);
  border-radius:10px;box-shadow:0 12px 32px rgba(0,0,0,.42);width:300px;overflow:hidden;
  animation:fab-in 200ms cubic-bezier(.5,1.4,.3,.95)}
.task-fab-shared.show{display:block}
@keyframes fab-in{from{opacity:0;transform:translateY(20px) scale(.94)}to{opacity:1;transform:translateY(0) scale(1)}}
.task-fab-shared .head{display:flex;align-items:center;gap:8px;padding:10px 12px;
  border-bottom:1px solid var(--line-1,#1e242d);background:var(--surface-3,#1d222b)}
.task-fab-shared .head .dot{width:8px;height:8px;border-radius:50%;background:#a78bfa;
  animation:fab-pulse 1.4s ease-in-out infinite}
@keyframes fab-pulse{0%,100%{opacity:.4;transform:scale(.85)}50%{opacity:1;transform:scale(1.05)}}
.task-fab-shared .head .label{font-size:11px;font-weight:600;text-transform:uppercase;
  letter-spacing:.04em;color:var(--fg-2,#a8b0bd)}
.task-fab-shared .head .count{margin-left:auto;font-family:"SF Mono",ui-monospace,monospace;
  font-size:10.5px;color:var(--fg-3,#6b7380);padding:2px 7px;border-radius:999px;background:var(--surface,#101216)}
.task-fab-shared .body{padding:8px;max-height:432px;overflow-y:auto;
  scrollbar-width:thin;scrollbar-color:var(--line-2) transparent}
.task-fab-shared .body::-webkit-scrollbar{width:6px}
.task-fab-shared .body::-webkit-scrollbar-thumb{background:var(--line-2);border-radius:3px}
.task-fab-shared .body::-webkit-scrollbar-track{background:transparent}
.task-row-shared{display:block;padding:8px 12px;border-radius:6px;cursor:pointer;
  transition:background .12s;text-decoration:none;color:inherit}
.task-row-shared:hover{background:var(--surface,#101216)}
.task-row-shared .name{font-size:13px;font-weight:500;color:var(--fg,#f5f7fa)}
.task-row-shared .progress{font-size:10.5px;font-family:"SF Mono",ui-monospace,monospace;
  color:var(--fg-3,#6b7380);margin-top:2px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.task-row-shared .bar{height:2px;background:var(--surface,#101216);border-radius:2px;
  overflow:hidden;margin-top:6px;position:relative}
.task-row-shared .bar-fill{height:100%;width:30%;
  background:linear-gradient(90deg,#a78bfa,#10b981);border-radius:2px;position:absolute;
  animation:indet-shared 1.6s ease-in-out infinite}
@keyframes indet-shared{0%{left:-30%}100%{left:100%}}
`;
  const style = document.createElement('style');
  style.textContent = css;
  document.head.appendChild(style);

  // DOM
  const overlay = document.createElement('div');
  overlay.className = 'cmd-overlay';
  overlay.id = 'cmd-overlay';
  overlay.innerHTML = `
    <div class="cmd-panel">
      <input class="cmd-input" id="cmd-input-shared" placeholder="输入工具名 / 报告 / 设置…" autocomplete="off">
      <div class="cmd-results" id="cmd-results-shared"></div>
    </div>`;
  document.body.appendChild(overlay);

  const fab = document.createElement('aside');
  fab.className = 'task-fab-shared';
  fab.id = 'task-fab';
  fab.innerHTML = `
    <div class="head">
      <span class="dot"></span>
      <span class="label">运行中</span>
      <span class="count" id="task-count-shared">0</span>
    </div>
    <div class="body" id="task-list-shared"></div>`;
  document.body.appendChild(fab);

  // Tools data fetched once
  let tools = [];
  fetch('/api/tools').then(r => r.json()).then(c => { tools = c.tools || []; });

  function escapeHtml(s){ return String(s).replace(/[&<>]/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;'}[c])); }

  // ⌘K palette
  const cmdInput = document.getElementById('cmd-input-shared');
  const cmdResults = document.getElementById('cmd-results-shared');
  let cmdActive = 0;
  function openCmd(){
    overlay.classList.add('show');
    setTimeout(() => cmdInput.focus(), 50);
    renderCmd('');
  }
  function closeCmd(){ overlay.classList.remove('show'); cmdInput.value = ''; }
  function renderCmd(q){
    const items = [
      ...tools.map(t => ({icon:t.icon, name:t.name, hint:t.id, href:`/tools/${t.id}`})),
      {icon:'📊', name:'查看报告', hint:'reports', href:'/reports'},
      {icon:'⚙', name:'设置 / 环境', hint:'settings', href:'/settings'},
    ];
    const filtered = q
      ? items.filter(x => x.name.toLowerCase().includes(q.toLowerCase()) || x.hint.toLowerCase().includes(q.toLowerCase()))
      : items;
    cmdActive = 0;
    cmdResults.innerHTML = filtered.length
      ? filtered.map((x, i) => `<div class="cmd-result ${i===0?'active':''}" data-href="${x.href}" data-idx="${i}">
          <span class="icon">${x.icon}</span>
          <span class="name">${escapeHtml(x.name)}</span>
          <span class="hint">${escapeHtml(x.hint)}</span></div>`).join('')
      : '<div class="cmd-empty">无匹配项</div>';
    cmdResults.querySelectorAll('.cmd-result').forEach(el => {
      el.onclick = () => location.href = el.dataset.href;
      el.onmouseenter = () => {
        cmdResults.querySelectorAll('.cmd-result').forEach(x => x.classList.remove('active'));
        el.classList.add('active');
        cmdActive = parseInt(el.dataset.idx, 10);
      };
    });
  }
  document.addEventListener('keydown', e => {
    if ((e.metaKey || e.ctrlKey) && e.key.toLowerCase() === 'k') { e.preventDefault(); openCmd(); }
    else if (overlay.classList.contains('show')) {
      if (e.key === 'Escape') { closeCmd(); return; }
      const items = cmdResults.querySelectorAll('.cmd-result');
      if (e.key === 'ArrowDown') {
        e.preventDefault();
        cmdActive = Math.min(cmdActive + 1, items.length - 1);
        items.forEach((el, i) => el.classList.toggle('active', i === cmdActive));
        items[cmdActive] && items[cmdActive].scrollIntoView({block:'nearest'});
      } else if (e.key === 'ArrowUp') {
        e.preventDefault();
        cmdActive = Math.max(cmdActive - 1, 0);
        items.forEach((el, i) => el.classList.toggle('active', i === cmdActive));
        items[cmdActive] && items[cmdActive].scrollIntoView({block:'nearest'});
      } else if (e.key === 'Enter') {
        e.preventDefault();
        const sel = items[cmdActive];
        if (sel) location.href = sel.dataset.href;
      }
    }
  });
  cmdInput.addEventListener('input', () => renderCmd(cmdInput.value));
  overlay.addEventListener('click', e => { if (e.target === overlay) closeCmd(); });
  // Make any element with id="cmd-trigger" or class "kbd-hint" open the palette
  document.querySelectorAll('#cmd-trigger,.kbd-hint').forEach(el => el.addEventListener('click', openCmd));

  // Floating running tasks widget
  async function pollTasks(){
    try {
      const data = await fetch('/api/tools/runs').then(r => r.json());
      const runs = (data.recent || []).filter(r => r.status === 'running' || r.status === 'queued');
      const list = document.getElementById('task-list-shared');
      if (!runs.length) { fab.classList.remove('show'); return; }
      document.getElementById('task-count-shared').textContent = runs.length;
      // 显示全部任务;最多 6 个在视口内,多了通过滚动条访问
      list.innerHTML = runs.map(t => {
        const elapsed = Math.max(0, Math.floor(Date.now()/1000 - (t.started_at||0)));
        const min = String(Math.floor(elapsed/60)).padStart(2,'0');
        const sec = String(elapsed%60).padStart(2,'0');
        return `<a class="task-row-shared" href="/tools/${t.tool_id}?run=${t.run_id}">
          <div class="name">${escapeHtml(t.tool_name)}</div>
          <div class="progress">${escapeHtml(t.progress || t.status)} · ${min}:${sec}</div>
          <div class="bar"><div class="bar-fill"></div></div></a>`;
      }).join('');
      fab.classList.add('show');
    } catch(e){}
  }
  pollTasks();
  setInterval(pollTasks, 4000);
})();

// ── 共享:登录用户徽标 + 登出按钮 + admin-only 导航链接可见性 ──
(function sharedAuthWidget(){
  if (document.getElementById('shared-auth-widget')) return;
  // 注意:这里改成把 widget 注入到 topbar 内部成为自然 inline 元素,不再 position:fixed,
  // 否则会覆盖 nav 右侧的 admin-only 链接(实测:widget 占据 nav 同一区域,叠在上面挡住"用户管理")。
  const css = `
  .shared-auth-widget{
    display:inline-flex;align-items:center;gap:10px;
    font-family:'Inter','SF Mono',monospace;font-size:12.5px;color:var(--fg-2,#262626);
    background:rgba(255,255,255,.94);
    border:1px solid var(--line-1,#9e9e9e);border-radius:999px;padding:4px 4px 4px 14px;
    margin-left:14px;flex-shrink:0;white-space:nowrap}
  .shared-auth-widget .admin-tag{color:#a8401f;font-size:10.5px;letter-spacing:.18em;font-weight:600}
  .shared-auth-widget .name{font-weight:600;color:var(--fg,#0a0a0a)}
  .shared-auth-widget .logout{background:none;border:1px solid var(--line-1,#9e9e9e);
    color:var(--fg-2,#262626);padding:4px 12px;border-radius:999px;
    font-family:inherit;font-size:11.5px;cursor:pointer;letter-spacing:.04em;
    transition:all .15s;font-weight:500}
  .shared-auth-widget .logout:hover{border-color:#a8401f;color:#a8401f}
  /* fallback: 没找到 topbar 时退回右上浮动 */
  .shared-auth-widget.shared-auth-floating{
    position:fixed;top:14px;right:18px;z-index:80;margin-left:0;
    box-shadow:0 1px 3px rgba(0,0,0,.06)}

  /* === 测试用例工具(step2)报告:用例表 + Excel 主体 === */
  .tc-block{margin:16px 18px;border:1px solid var(--line-1,#c4c4c4);border-radius:8px;overflow:hidden;background:#fff}
  .tc-banner{display:flex;align-items:center;gap:18px;padding:18px 22px;
    background:linear-gradient(135deg,#f4f1e6,#ece8da);border-bottom:1px solid var(--line-1,#c4c4c4)}
  .tc-banner .tc-count{font-family:'Noto Serif SC',Georgia,serif;font-size:38px;font-weight:600;
    color:#a8401f;line-height:1}
  .tc-banner .tc-count-label{font-size:12px;color:var(--fg-2,#262626);letter-spacing:.1em;margin-top:2px}
  .tc-banner .tc-mid{flex:1;font-size:13px;color:var(--fg-2,#262626);line-height:1.7}
  .tc-banner .tc-mid b{color:#a8401f}
  .tc-excel-btn{background:#1a7a3a;color:#fff;border:none;border-radius:5px;
    padding:13px 22px;font-family:'Noto Sans SC',sans-serif;font-size:14px;font-weight:600;
    letter-spacing:.06em;cursor:pointer;white-space:nowrap;transition:background .15s;
    box-shadow:0 2px 6px rgba(26,122,58,.28)}
  .tc-excel-btn:hover{background:#15692f}
  .tc-excel-btn:active{transform:translateY(1px)}
  .tc-table-wrap{overflow-x:auto;max-height:560px;overflow-y:auto}
  .tc-table{width:100%;border-collapse:collapse;font-size:12.5px}
  .tc-table th{position:sticky;top:0;background:#2e2e2e;color:#fff;font-weight:600;
    padding:9px 10px;text-align:left;white-space:nowrap;font-size:11.5px;letter-spacing:.04em;z-index:1}
  .tc-table td{padding:9px 10px;border-bottom:1px solid #e3e3e3;vertical-align:top;
    color:var(--fg,#0a0a0a);line-height:1.6}
  .tc-table tr:hover td{background:#faf8f2}
  .tc-table .tc-id{font-family:'Inter','SF Mono',monospace;font-size:11.5px;color:#a8401f;
    white-space:nowrap;font-weight:600}
  .tc-table .tc-title{min-width:170px;font-weight:500}
  .tc-table .tc-steps{min-width:240px;white-space:pre-line;color:var(--fg-2,#262626)}
  .tc-table .tc-exp{min-width:200px;color:var(--fg-2,#262626)}
  .tc-table .tc-pre{min-width:130px;color:var(--fg-3,#4a4a4a);font-size:12px}
  .tc-pri{display:inline-block;padding:1px 8px;border-radius:3px;font-size:11px;font-weight:700;
    font-family:'Inter',monospace}
  .tc-pri-P0{background:#f4cccc;color:#7a1c1c}
  .tc-pri-P1{background:#fce5cd;color:#7a3d0c}
  .tc-pri-P2{background:#fff2cc;color:#6b5500}
  .tc-pri-P3{background:#efefef;color:#555}
  `;
  const st = document.createElement('style'); st.textContent = css;
  document.head.appendChild(st);

  // 测试用例 Excel 下载 — 委托点击,所有页面通用
  document.addEventListener('click', async (ev) => {
    const btn = ev.target.closest && ev.target.closest('.tc-excel-btn[data-tc-runid]');
    if (!btn) return;
    ev.preventDefault();
    const runId = btn.dataset.tcRunid;
    const toolId = btn.dataset.tcToolid || 'step2';
    const orig = btn.textContent;
    btn.disabled = true; btn.textContent = '导出中…';
    try {
      const resp = await fetch(`/api/reports/${runId}/export.xlsx`, {credentials:'same-origin'});
      if (!resp.ok){
        btn.textContent = '导出失败'; setTimeout(()=>{btn.textContent=orig;btn.disabled=false;}, 1800);
        return;
      }
      const blob = await resp.blob();
      const a = document.createElement('a');
      a.href = URL.createObjectURL(blob);
      a.download = `${toolId}_${runId.slice(0,8)}_测试用例.xlsx`;
      document.body.appendChild(a); a.click(); a.remove();
      setTimeout(()=>URL.revokeObjectURL(a.href), 1000);
      btn.textContent = '✓ 已下载'; setTimeout(()=>{btn.textContent=orig;btn.disabled=false;}, 1800);
    } catch(e){
      btn.textContent = '导出出错'; setTimeout(()=>{btn.textContent=orig;btn.disabled=false;}, 1800);
    }
  });
  fetch('/api/auth/me').then(r => r.json()).then(d => {
    if (!d || !d.authenticated || !d.user) return;
    const u = d.user;
    // 角色门禁:admin/superadmin 显示 .admin-only;仅 superadmin 显示 .superadmin-only
    if (u.role === 'admin' || u.role === 'superadmin'){
      document.querySelectorAll('.admin-only').forEach(el => { el.style.display = ''; });
    }
    if (u.role === 'superadmin'){
      document.querySelectorAll('.superadmin-only').forEach(el => { el.style.display = ''; });
    }
    // 如果页面已经有自己的用户信息区(如 /admin/users 的 user-chip),跳过 widget 注入
    if (document.getElementById('user-chip') || document.getElementById('user-badge')){
      return;
    }
    const wrap = document.createElement('div');
    wrap.className = 'shared-auth-widget';
    wrap.id = 'shared-auth-widget';
    const tag = u.role === 'superadmin'
      ? '<span class="admin-tag">[超管]</span>'
      : (u.role === 'admin' ? '<span class="admin-tag">[admin]</span>' : '');
    wrap.innerHTML = `${tag}<span class="name">${u.display_name || u.username}</span>` +
      `<button class="logout" title="退出登录">登出</button>`;
    // 找 topbar 把 widget 塞进去成自然元素 — 优先级:header.topbar > .topbar > header
    const host = document.querySelector('header.topbar, .topbar, header');
    if (host){
      host.appendChild(wrap);
    } else {
      // 没找到 topbar — 退回 fixed 右上,加 floating class
      wrap.classList.add('shared-auth-floating');
      document.body.appendChild(wrap);
    }
    wrap.querySelector('.logout').onclick = async () => {
      await fetch('/api/auth/logout', {method:'POST'});
      location.href = '/login';
    };
  }).catch(()=>{});
})();
</script>
"""


def _inject_shared_overlays(html: str) -> str:
    """Inject ⌘K palette + floating tasks widget into pages that don't have them.

    Use rsplit to only target the FINAL `</body></html>` (the real page close).
    Naive .replace breaks pages whose JS template literals contain the same
    string (e.g. report download functions that return a full HTML document).
    """
    parts = html.rsplit("</body></html>", 1)
    if len(parts) != 2:
        return html  # malformed input, do nothing
    return parts[0] + _SHARED_OVERLAY_SNIPPET + "</body></html>" + parts[1]


_AUTH_PAGE_CSS = """
:root {
  --paper:#ffffff; --paper-2:#ebebeb; --paper-3:#dcdcdc;
  --ink:#0a0a0a; --ink-2:#262626; --ink-3:#4a4a4a; --ink-4:#6e6e6e;
  --line:#bdbdbd; --line-2:#9e9e9e;
  --accent:#a8401f; --accent-h:#82301a; --accent-soft:rgba(168,64,31,.14);
  --bad:#8a2d12; --ok:#4f6b35; --gold:#876b1f;
  --serif:'Noto Serif SC','Songti SC','STSong',Georgia,serif;
  --sans:'Noto Sans SC','PingFang SC',-apple-system,'Microsoft YaHei',sans-serif;
  --mono:'Inter','SF Mono',ui-monospace,Menlo,monospace;
}
*{box-sizing:border-box}
html,body{margin:0;height:100%;background:var(--paper);color:var(--ink);
  font-family:var(--sans);font-weight:300;font-size:15px;line-height:1.7;
  -webkit-font-smoothing:antialiased;font-feature-settings:"palt" on,"ss01" on}
a{color:inherit;text-decoration:none}
button{font-family:inherit;cursor:pointer}

/* === 两栏布局 === */
.auth-shell{display:grid;grid-template-columns:minmax(0,1.1fr) minmax(0,.9fr);
  min-height:100vh}
@media (max-width:900px){.auth-shell{grid-template-columns:1fr}}

/* === 左侧:品牌叙事区 === */
.brand-side{padding:64px 72px;display:flex;flex-direction:column;justify-content:space-between;
  background:
    radial-gradient(circle at 22% 18%,rgba(184,85,58,.06),transparent 50%),
    radial-gradient(circle at 78% 78%,rgba(169,139,58,.05),transparent 45%),
    linear-gradient(135deg,#f4f1e6 0%,#ebe7d8 100%);
  position:relative;overflow:hidden;border-right:1px solid var(--line)}
.brand-side::before{
  content:"";position:absolute;inset:48px 48px auto auto;width:96px;height:96px;
  background:
    repeating-linear-gradient(45deg,transparent 0 4px,rgba(184,85,58,.05) 4px 5px);
  border-radius:50%;pointer-events:none}
.brand-side::after{
  content:"";position:absolute;left:-160px;bottom:-160px;width:400px;height:400px;
  background:radial-gradient(circle,rgba(184,85,58,.06) 0%,transparent 65%);
  pointer-events:none}
.brand-mark{display:flex;align-items:center;gap:14px;position:relative;z-index:2}
.brand-mark svg{color:var(--accent);filter:drop-shadow(0 2px 6px rgba(184,85,58,.25))}
.brand-name{font-family:var(--serif);font-size:24px;font-weight:600;letter-spacing:.2em;
  color:var(--ink)}
.brand-sep{color:var(--accent);margin:0 6px;font-weight:400}
.brand-body{flex:1;display:flex;flex-direction:column;justify-content:center;
  padding:48px 0;position:relative;z-index:2;max-width:520px}
.eyebrow{font-family:var(--mono);font-size:11.5px;letter-spacing:.4em;color:var(--ink-3);
  text-transform:uppercase;margin-bottom:32px;border-left:2px solid var(--accent);padding-left:14px}
.brand-headline{font-family:var(--serif);font-size:40px;font-weight:500;
  letter-spacing:.04em;line-height:1.35;margin:0 0 24px;color:var(--ink)}
.brand-headline em{font-style:normal;color:var(--accent);font-weight:600}
.brand-lede{font-size:15px;color:var(--ink-2);line-height:1.85;margin:0 0 36px;
  font-weight:300;max-width:460px}
.brand-pillars{display:grid;grid-template-columns:repeat(2,1fr);gap:20px;
  margin-top:24px;max-width:480px}
.pillar{padding:14px 0;border-top:1px solid var(--line);font-size:13px;color:var(--ink-2)}
.pillar-no{font-family:var(--mono);font-size:11px;color:var(--accent);letter-spacing:.18em;
  margin-bottom:4px}
.brand-foot{position:relative;z-index:2;display:flex;align-items:center;justify-content:space-between;
  font-family:var(--mono);font-size:11px;color:var(--ink-3);letter-spacing:.12em}

/* === 右侧:表单区 === */
.form-side{padding:64px 56px;display:flex;flex-direction:column;justify-content:center;
  background:var(--paper);position:relative}
@media (max-width:900px){.form-side{padding:48px 32px}}
.form-card{width:100%;max-width:380px;margin:0 auto}
.notice{padding:14px 16px;background:linear-gradient(180deg,var(--accent-soft),rgba(184,85,58,.02));
  border-left:3px solid var(--accent);border-radius:0 4px 4px 0;
  color:var(--ink);font-size:13px;line-height:1.7;margin-bottom:32px}
.notice strong{color:var(--accent);font-weight:600}
.form-title{font-family:var(--serif);font-size:30px;font-weight:500;letter-spacing:.05em;
  margin:0 0 8px;color:var(--ink)}
.form-sub{font-size:13.5px;color:var(--ink-3);margin-bottom:36px;letter-spacing:.04em;font-weight:300}
form{display:flex;flex-direction:column;gap:20px}
.field{display:flex;flex-direction:column;gap:8px}
.field label{font-family:var(--mono);font-size:10.5px;color:var(--ink-3);
  letter-spacing:.22em;text-transform:uppercase;font-weight:500}
.field input{font-family:var(--sans);font-size:15px;padding:13px 16px;
  border:1px solid var(--line);border-radius:4px;background:#fff;color:var(--ink);
  transition:border .18s,box-shadow .18s;outline:none;font-weight:400}
.field input:hover{border-color:var(--line-2)}
.field input:focus{border-color:var(--accent);box-shadow:0 0 0 3px var(--accent-soft)}
.field .hint{font-size:11.5px;color:var(--ink-4);letter-spacing:.02em;margin-top:2px}
.submit{margin-top:14px;padding:14px 20px;background:var(--ink);color:#fff;border:none;
  border-radius:4px;font-family:var(--serif);font-size:15px;font-weight:500;letter-spacing:.28em;
  transition:background .18s,transform .08s}
.submit:hover{background:var(--accent)}
.submit:active{transform:translateY(1px)}
.submit:disabled{opacity:.5;cursor:not-allowed;transform:none}
.error{margin-bottom:18px;padding:12px 16px;border-left:3px solid var(--bad);
  background:rgba(154,63,41,.06);color:#5a2418;font-size:13px;
  border-radius:0 4px 4px 0;display:none;line-height:1.6}
.error.show{display:block}
.alt{margin-top:28px;text-align:center;font-size:13px;color:var(--ink-3);font-weight:300}
.alt a{color:var(--accent);font-weight:500}
.alt a:hover{color:var(--accent-h);text-decoration:underline}
.form-foot{margin-top:48px;padding-top:24px;border-top:1px solid var(--line);
  text-align:center;font-family:var(--mono);font-size:10.5px;color:var(--ink-4);
  letter-spacing:.18em;text-transform:uppercase}
"""


def _auth_page_shell(*, kind: str) -> str:
    """渲染登录 / 注册页 — 出版物风两栏布局."""
    first_setup = user_store.count_users() == 0
    is_login = (kind == "login")

    if not is_login and not first_setup:
        # 已经有用户了 — 公开注册关闭,直接 302 不会走到这里
        return ""  # caller 应该重定向

    # 左侧根据是否首次安装变文案
    if first_setup and not is_login:
        eyebrow = "首次安装 · INITIAL SETUP"
        headline = '建立你的<em>第一位管理员</em>'
        lede = ("天枢·裁决 是一套面向软件质量的 AI 裁决工具。"
                "现在为这个实例创建第一位管理员账号 — 之后由你掌控用户、凭据与报告。")
    else:
        eyebrow = "AI VERDICT MANUAL · 0.1"
        headline = '面向软件质量的<br><em>AI 裁决平台</em>'
        lede = ("把研发各环节的质量把关交给 AI 智能体 —— "
                "逐项查验,出具可分派、可追溯的裁决报告。")

    pillars = (
        '<div class="pillar"><div class="pillar-no">评审</div>需求质量</div>'
        '<div class="pillar"><div class="pillar-no">设计</div>用例覆盖</div>'
        '<div class="pillar"><div class="pillar-no">测试</div>功能 · 接口</div>'
        '<div class="pillar"><div class="pillar-no">审计</div>体验 · SEO</div>'
    )

    if is_login:
        # 首次安装时,登录页显示提示 + 注册入口
        notice = (
            '<div class="notice">系统里还没有任何用户。请先 '
            '<strong><a href="/register" style="color:var(--accent)">建立第一位管理员</a></strong> '
            '再回来登录。</div>'
            if first_setup else ""
        )
        form_title = "登 录"
        form_sub = "凭用户名与密码进入工作台"
        submit_label = "进 入"
        extra_field = ""
        pwd_hint = ""
        pwd_autocomplete = "current-password"
        alt_html = ""  # 没有"去注册"链接 — 公开注册已关
        form_kind = "login"
    else:
        notice = (
            '<div class="notice"><strong>初始化</strong> · 这位管理员将拥有最高权限:'
            '配置 Claude 凭据、查看所有用户的报告、新建/重置其他账号。</div>'
            if first_setup else ""
        )
        form_title = "建立管理员"
        form_sub = "为本实例的第一位用户"
        submit_label = "创 建"
        extra_field = (
            '<div class="field"><label for="display_name">显示名 (可选)</label>'
            '<input id="display_name" name="display_name" type="text" '
            'autocomplete="nickname" placeholder="留空则同用户名"></div>'
        )
        pwd_hint = '<div class="hint">至少 6 位。bcrypt 哈希存储,不上传外部服务。</div>'
        pwd_autocomplete = "new-password"
        alt_html = ""  # 不引导回登录,先建管理员
        form_kind = "register"

    return f"""<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{'登录' if is_login else '建立管理员'} — 天枢·裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>{_AUTH_PAGE_CSS}</style>
</head><body>
<div class="auth-shell">
  <!-- 左侧:品牌叙事区 -->
  <aside class="brand-side">
    <div class="brand-mark">
      <svg viewBox="0 0 24 24" fill="currentColor" width="32" height="32" aria-hidden="true">
        <path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/>
      </svg>
      <span class="brand-name">天枢 <span class="brand-sep">·</span> 裁决</span>
    </div>
    <div class="brand-body">
      <div class="eyebrow">{eyebrow}</div>
      <h1 class="brand-headline">{headline}</h1>
      <p class="brand-lede">{lede}</p>
      <div class="brand-pillars">{pillars}</div>
    </div>
    <div class="brand-foot">
      <span>EDITION 0.1</span>
      <span>全流程查验 · 一份裁决</span>
    </div>
  </aside>

  <!-- 右侧:表单区 -->
  <main class="form-side">
    <div class="form-card">
      {notice}
      <h2 class="form-title">{form_title}</h2>
      <div class="form-sub">{form_sub}</div>
      <div class="error" id="err"></div>
      <form id="auth-form" autocomplete="on">
        <div class="field">
          <label for="username">用户名</label>
          <input id="username" name="username" type="text" required autocomplete="username"
                 placeholder="如 zhangsan" autofocus>
        </div>
        {extra_field}
        <div class="field">
          <label for="password">密码</label>
          <input id="password" name="password" type="password" required
                 autocomplete="{pwd_autocomplete}" placeholder="至少 6 位">
          {pwd_hint}
        </div>
        <button class="submit" id="submit-btn" type="submit">{submit_label}</button>
      </form>
      {f'<div class="alt">{alt_html}</div>' if alt_html else ''}
      <div class="form-foot">SECURE · BCRYPT · LOCAL ONLY</div>
    </div>
  </main>
</div>
<script>
const FORM_KIND = "{form_kind}";
const errBox = document.getElementById('err');
function showErr(msg){{ errBox.textContent = msg; errBox.classList.add('show'); }}
function clearErr(){{ errBox.classList.remove('show'); errBox.textContent = ''; }}
async function doSubmit(e){{
  e.preventDefault(); clearErr();
  const btn = document.getElementById('submit-btn');
  btn.disabled = true; btn.textContent = '处 理 中...';
  const body = {{
    username: document.getElementById('username').value.trim(),
    password: document.getElementById('password').value,
  }};
  const nm = document.getElementById('display_name');
  if (nm) body.display_name = nm.value.trim();
  try {{
    const path = FORM_KIND === 'login' ? '/api/auth/login' : '/api/auth/register';
    const r = await fetch(path, {{
      method:'POST', headers:{{'Content-Type':'application/json'}},
      body: JSON.stringify(body),
    }});
    if (!r.ok){{
      const d = await r.json().catch(()=>({{}}));
      showErr(d.detail || (FORM_KIND === 'login' ? '登录失败' : '创建失败'));
      btn.disabled = false; btn.textContent = '{submit_label}';
      return;
    }}
    const params = new URLSearchParams(location.search);
    const next = params.get('next') || '/tools';
    location.href = next;
  }} catch(err){{
    showErr('网络错误:' + err.message);
    btn.disabled = false; btn.textContent = '{submit_label}';
  }}
}}
document.getElementById('auth-form').addEventListener('submit', doSubmit);
</script>
</body></html>"""


@app.get("/login", response_class=HTMLResponse)
async def login_page() -> str:
    return _auth_page_shell(kind="login")


@app.get("/register")
async def register_page():
    """注册页 — 仅在系统还没用户时存在。否则 302 到 /login。"""
    if user_store.count_users() > 0:
        return RedirectResponse("/login?msg=registration_disabled", status_code=302)
    return HTMLResponse(_auth_page_shell(kind="register"))


@app.get("/tools", response_class=HTMLResponse)
async def tools_index_page() -> str:
    # 注入共享 overlay:右上角 admin 徽标 + 登出 + admin-only 元素显示逻辑
    return _inject_shared_overlays(TOOLS_INDEX_HTML)


# 二级页:工具列表(平铺网格)。首页只放「开始」,工具/报告各自独立二级页。
CATALOG_HTML = """<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>工具 · 天枢 · 裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;500;600&family=Noto+Sans+SC:wght@300;400;500&display=swap" rel="stylesheet">
<style>
:root{--paper:#fff;--paper-2:#f4f3f1;--ink:#0a0a0a;--ink-2:#3a3a3a;--ink-3:#6e6e6e;--line:#dcdad6;--accent:#a8401f}
*{box-sizing:border-box}
body{margin:0;background:var(--paper);color:var(--ink);font-family:'Noto Sans SC',-apple-system,BlinkMacSystemFont,sans-serif}
header.topbar{display:flex;align-items:center;gap:6px;padding:15px 40px;border-bottom:1px solid var(--line);position:sticky;top:0;background:rgba(255,255,255,.92);backdrop-filter:blur(8px);z-index:20}
.brand-link{display:inline-flex;align-items:center;gap:10px;text-decoration:none;color:inherit;margin-right:20px}
.brand-link svg{color:var(--accent);width:22px;height:22px;flex-shrink:0}
.brand{font-family:'Noto Serif SC',serif;font-size:18px;font-weight:600;letter-spacing:.16em;color:var(--ink)}
.brand .sep{color:var(--accent);margin:0 6px;font-weight:400}
header.topbar nav{display:flex;gap:2px}
header.topbar nav a{padding:6px 14px;border-radius:7px;color:var(--ink-3);text-decoration:none;font-size:14px;transition:.14s}
header.topbar nav a:hover{color:var(--ink);background:var(--paper-2)}
header.topbar nav a.active{color:var(--ink);background:var(--paper-2);font-weight:500}
.wrap{max-width:1140px;margin:0 auto;padding:40px}
.wrap h1{font-family:'Noto Serif SC',serif;font-size:27px;font-weight:500;margin:0 0 6px;letter-spacing:-.01em}
.wrap .sub{color:var(--ink-3);font-size:13.5px;margin:0 0 30px}
.grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(320px,1fr));gap:16px}
.tile{position:relative;display:flex;flex-direction:column;padding:24px 24px 22px;border:1px solid var(--line);border-radius:14px;background:var(--paper);text-decoration:none;color:inherit;transition:.18s;overflow:hidden}
.tile:hover{border-color:var(--accent);box-shadow:0 8px 26px rgba(168,64,31,.10);transform:translateY(-3px)}
.tile .num{position:absolute;top:14px;right:20px;font-family:'Noto Serif SC',serif;font-size:46px;font-weight:500;line-height:1;color:#efece8}
.tile:hover .num{color:rgba(168,64,31,.16)}
.tile .icon{font-size:32px;line-height:1}
.tile h3{font-family:'Noto Serif SC',serif;font-size:20px;font-weight:500;margin:16px 0 8px}
.tile p{font-size:13px;color:var(--ink-2);line-height:1.7;margin:0;flex:1}
.tile .go{margin-top:16px;font-size:13px;color:var(--accent);opacity:.55;transition:.18s}
.tile:hover .go{opacity:1;transform:translateX(3px)}
@media(max-width:640px){.wrap{padding:24px 18px}header.topbar{padding:12px 18px;overflow-x:auto}}
</style></head>
<body>
<header class="topbar">
  <a class="brand-link" href="/tools"><svg viewBox="0 0 24 24" fill="currentColor" aria-hidden="true"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span class="brand">天枢<span class="sep">·</span>裁决</span></a>
  <nav>
    <a href="/tools">开始</a>
    <a href="/catalog" class="active">工具</a>
    <a href="/reports">报告</a>
    <a href="/settings" class="superadmin-only" style="display:none">设置</a>
    <a href="/usage" class="admin-only" style="display:none">使用统计</a>
    <a href="/admin/users" class="admin-only" style="display:none">用户管理</a>
  </nav>
</header>
<div class="wrap">
  <h1>工具</h1>
  <p class="sub">AI 测试智能体 · 点击进入，可单独运行，也可链式接力</p>
  <div class="grid" id="grid"><div style="color:var(--ink-3);font-size:13px">加载中…</div></div>
</div>
<script>
const ORDER=['step1','step2','step4','step6','network_resilience','h5_adapt','seo_audit'];
const esc=s=>String(s==null?'':s).replace(/[&<>"]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}[c]));
fetch('/api/tools',{credentials:'same-origin'}).then(r=>r.json()).then(d=>{
  const tools=d.tools||[];const byId={};tools.forEach(t=>byId[t.id]=t);
  const ordered=ORDER.map(id=>byId[id]).filter(Boolean);
  tools.forEach(t=>{if(ORDER.indexOf(t.id)<0)ordered.push(t);});
  document.getElementById('grid').innerHTML=ordered.map((t,i)=>
    '<a class="tile" href="/tools/'+t.id+'">'+
      '<span class="num">'+(i+1)+'</span>'+
      '<span class="icon">'+(t.icon||'🔧')+'</span>'+
      '<h3>'+esc(t.name)+'</h3>'+
      '<p>'+esc(t.tagline||t.description||'')+'</p>'+
      '<div class="go">进入 →</div>'+
    '</a>').join('');
}).catch(function(e){document.getElementById('grid').innerHTML='<div style="color:#a40e26;font-size:13px">工具加载失败</div>';});
</script>
</body></html>"""


@app.get("/catalog", response_class=HTMLResponse)
async def catalog_page() -> str:
    return _inject_shared_overlays(CATALOG_HTML)


# 二级页:使用说明（每个工具的详细用法）
GUIDE_HTML = """<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>使用说明 · 天枢 · 裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;500;600&family=Noto+Sans+SC:wght@300;400;500&display=swap" rel="stylesheet">
<style>
:root{--paper:#fff;--paper-2:#f4f3f1;--ink:#0a0a0a;--ink-2:#3a3a3a;--ink-3:#6e6e6e;--line:#dcdad6;--accent:#a8401f}
*{box-sizing:border-box}
body{margin:0;background:var(--paper);color:var(--ink);font-family:'Noto Sans SC',-apple-system,BlinkMacSystemFont,sans-serif}
header.topbar{display:flex;align-items:center;gap:6px;padding:15px 40px;border-bottom:1px solid var(--line);position:sticky;top:0;background:rgba(255,255,255,.94);backdrop-filter:blur(8px);z-index:20}
.brand-link{display:inline-flex;align-items:center;gap:10px;text-decoration:none;color:inherit;margin-right:20px}
.brand-link svg{color:var(--accent);width:22px;height:22px;flex-shrink:0}
.brand{font-family:'Noto Serif SC',serif;font-size:18px;font-weight:600;letter-spacing:.16em;color:var(--ink)}
.brand .sep{color:var(--accent);margin:0 6px;font-weight:400}
header.topbar nav{display:flex;gap:2px}
header.topbar nav a{padding:6px 14px;border-radius:7px;color:var(--ink-3);text-decoration:none;font-size:14px;transition:.14s}
header.topbar nav a:hover{color:var(--ink);background:var(--paper-2)}
header.topbar nav a.active{color:var(--ink);background:var(--paper-2);font-weight:500}
.wrap{max-width:880px;margin:0 auto;padding:44px 40px 80px}
.wrap h1{font-family:'Noto Serif SC',serif;font-size:27px;font-weight:500;margin:0 0 6px}
.wrap .sub{color:var(--ink-3);font-size:13.5px;margin:0 0 36px}
.g-tool{padding:26px 0;border-top:1px solid var(--line)}
.g-tool:first-of-type{border-top:none;padding-top:0}
.g-head{display:flex;align-items:center;gap:13px;margin-bottom:16px}
.g-num{font-family:'Noto Serif SC',serif;font-size:22px;font-weight:500;color:var(--accent);min-width:24px}
.g-icon{font-size:24px;line-height:1}
.g-head h2{font-family:'Noto Serif SC',serif;font-size:20px;font-weight:500;margin:0}
.g-tool dl{margin:0;display:grid;grid-template-columns:72px 1fr;gap:9px 18px}
.g-tool dt{font-size:13px;font-weight:600;color:var(--ink-3);padding-top:1px}
.g-tool dd{margin:0;font-size:14px;line-height:1.78;color:var(--ink-2)}
.g-tool dd b{color:var(--ink);font-weight:500}
.g-tip{color:var(--accent)!important}
@media(max-width:600px){.wrap{padding:28px 18px 60px}.g-tool dl{grid-template-columns:1fr;gap:4px}.g-tool dt{padding-top:8px}}
</style></head>
<body>
<header class="topbar">
  <a class="brand-link" href="/tools"><svg viewBox="0 0 24 24" fill="currentColor" aria-hidden="true"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span class="brand">天枢<span class="sep">·</span>裁决</span></a>
  <nav>
    <a href="/tools">开始</a>
    <a href="/catalog">工具</a>
    <a href="/reports">报告</a>
    <a href="/guide" class="active">使用说明</a>
    <a href="/settings" class="superadmin-only" style="display:none">设置</a>
    <a href="/usage" class="admin-only" style="display:none">使用统计</a>
    <a href="/admin/users" class="admin-only" style="display:none">用户管理</a>
  </nav>
</header>
<div class="wrap">
  <h1>使用说明</h1>
  <p class="sub">每个工具的用法 · 都支持单独运行,上一步产出可一键接力给下一步</p>

  <div class="g-tool">
    <div class="g-head"><span class="g-num">1</span><span class="g-icon">📋</span><h2>需求评审</h2></div>
    <dl>
      <dt>用途</dt><dd>把 <b>PRD / 原型 / UI 稿</b>拆成可测结构,提前揪出需求遗漏、歧义与提测门禁。</dd>
      <dt>怎么用</dt><dd>进入工具 → 把需求材料(PRD 文档、原型或 UI 链接/说明)<b>粘贴或上传</b>到输入框 → 选模型 → 点运行。</dd>
      <dt>产出</dt><dd>《需求拆解报告》:模块拆解、主流程/异常流程、风险与遗漏清单、提测门禁。</dd>
      <dt class="g-tip">提示</dt><dd class="g-tip">材料越完整(含验收标准、边界规则)拆得越准;产出可一键接力给「测试用例设计」。</dd>
    </dl>
  </div>

  <div class="g-tool">
    <div class="g-head"><span class="g-num">2</span><span class="g-icon">✏️</span><h2>测试用例设计</h2></div>
    <dl>
      <dt>用途</dt><dd>基于需求,自动生成分优先级(<b>P0/P1/P2</b>)的测试用例集。</dd>
      <dt>怎么用</dt><dd>粘贴<b>需求拆解报告 / 业务场景 / 接口清单</b> → 运行;也可直接接力上一步「需求评审」的产出。</dd>
      <dt>产出</dt><dd>P0/P1/P2 用例表(步骤、预期、数据)+ <b>Excel 下载</b>。</dd>
      <dt class="g-tip">提示</dt><dd class="g-tip">想覆盖异常/边界,材料里写清规则与约束。</dd>
    </dl>
  </div>

  <div class="g-tool">
    <div class="g-head"><span class="g-num">3</span><span class="g-icon">🔌</span><h2>接口测试</h2></div>
    <dl>
      <dt>用途</dt><dd>对接口做<b>功能 / 安全 / 边界 / 契约</b>校验,给放量结论。AI 会<b>真发 HTTP 请求</b>看真实响应。</dd>
      <dt>怎么用</dt><dd>粘贴 <b>API 文档 / OpenAPI / Postman 集合 / 接口清单</b> + 环境信息(base URL、鉴权)→ 运行。</dd>
      <dt>产出</dt><dd>接口测试报告(功能/安全/边界/契约)+ 缺陷清单 + Excel。</dd>
      <dt class="g-tip">提示</dt><dd class="g-tip">请求由真实客户端发起,需保证环境可达、鉴权有效。</dd>
    </dl>
  </div>

  <div class="g-tool">
    <div class="g-head"><span class="g-num">4</span><span class="g-icon">🤖</span><h2>Agent 自动化执行</h2></div>
    <dl>
      <dt>用途</dt><dd>Agent 在<b>真机/模拟器</b>上跑 P0 + 关键 P1/P2,自动归因失败、列阻塞。</dd>
      <dt>怎么用</dt><dd><b>上传 APK</b> + 提供用例/场景,确保 MuMu 模拟器已开机;默认 dry-run,要真跑手动开 <b>--live</b>。</dd>
      <dt>产出</dt><dd>《Agent 自动化执行报告》:执行结果、失败归因(前端/后端/接口/数据/环境)、阻塞清单。</dd>
      <dt class="g-tip">提示</dt><dd class="g-tip">真机执行需模拟器在线、APK 可安装。</dd>
    </dl>
  </div>

  <div class="g-tool">
    <div class="g-head"><span class="g-num">5</span><span class="g-icon">📶</span><h2>弱网/断网测试</h2></div>
    <dl>
      <dt>用途</dt><dd>在 <b>WiFi / 4G / 3G / 弱网 / 2G / 断网</b> 各档下测页面表现 + 断网恢复能力。</dd>
      <dt>怎么用</dt><dd>填<b>目标页面 URL</b> → 运行,工具自动逐档切换网络采集。</dd>
      <dt>产出</dt><dd>弱网深度测试报告(各档表现 + 容错 + 退回判定)+ Excel。</dd>
    </dl>
  </div>

  <div class="g-tool">
    <div class="g-head"><span class="g-num">6</span><span class="g-icon">📱</span><h2>H5 适配初审</h2></div>
    <dl>
      <dt>用途</dt><dd>检查 H5 在不同视口/浏览器的适配:<b>安全区、浏览器矩阵、交互、性能</b>。</dd>
      <dt>怎么用</dt><dd>填 <b>H5 页面 URL</b> +(可选)figma 设计稿链接 → 运行。</dd>
      <dt>产出</dt><dd>H5 适配报告(页面盘点 + 视口/安全区/浏览器矩阵 + 性能)。</dd>
    </dl>
  </div>

  <div class="g-tool">
    <div class="g-head"><span class="g-num">7</span><span class="g-icon">🔍</span><h2>SEO 深度审计</h2></div>
    <dl>
      <dt>用途</dt><dd><b>全站抓取 + 元数据 + 内容 + Core Web Vitals</b> 四维深度审计。</dd>
      <dt>怎么用</dt><dd>填<b>站点 URL</b> → 运行,工具自动 BFS 爬取 + 逐页解析 + 实测 CWV。</dd>
      <dt>产出</dt><dd>SEO 深度审计报告(10 个 sheet 的 Excel)。</dd>
    </dl>
  </div>

</div>
</body></html>"""


@app.get("/guide", response_class=HTMLResponse)
async def guide_page() -> str:
    return _inject_shared_overlays(GUIDE_HTML)


USAGE_HTML = """<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>使用统计 · 天枢 · 裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;500;600&family=Noto+Sans+SC:wght@300;400;500&display=swap" rel="stylesheet">
<style>
:root{--paper:#fff;--paper-2:#f4f3f1;--ink:#0a0a0a;--ink-2:#3a3a3a;--ink-3:#6e6e6e;--ink-4:#8a8a8a;--line:#dcdad6;--accent:#a8401f}
*{box-sizing:border-box}
body{margin:0;background:var(--paper);color:var(--ink);font-family:'Noto Sans SC',-apple-system,BlinkMacSystemFont,sans-serif}
header.topbar{display:flex;align-items:center;gap:6px;padding:15px 40px;border-bottom:1px solid var(--line);position:sticky;top:0;background:rgba(255,255,255,.92);backdrop-filter:blur(8px);z-index:20}
.brand-link{display:inline-flex;align-items:center;gap:10px;text-decoration:none;color:inherit;margin-right:20px}
.brand-link svg{color:var(--accent);width:22px;height:22px;flex-shrink:0}
.brand{font-family:'Noto Serif SC',serif;font-size:18px;font-weight:600;letter-spacing:.16em;color:var(--ink)}
.brand .sep{color:var(--accent);margin:0 6px;font-weight:400}
header.topbar nav{display:flex;gap:2px;flex-wrap:wrap}
header.topbar nav a{padding:6px 14px;border-radius:7px;color:var(--ink-3);text-decoration:none;font-size:14px;transition:.14s}
header.topbar nav a:hover{color:var(--ink);background:var(--paper-2)}
header.topbar nav a.active{color:var(--ink);background:var(--paper-2);font-weight:500}
.wrap{max-width:1140px;margin:0 auto;padding:40px}
.wrap h1{font-family:'Noto Serif SC',serif;font-size:27px;font-weight:500;margin:0 0 6px;letter-spacing:-.01em}
.wrap .sub{color:var(--ink-3);font-size:13.5px;margin:0 0 22px;line-height:1.7}
.mr-head{display:flex;align-items:center;margin-bottom:14px}
.mr-dl{margin-left:auto;font:inherit;font-size:12.5px;padding:6px 14px;border:1px solid var(--line);border-radius:7px;background:transparent;cursor:pointer;color:var(--ink-2)}
.mr-dl:hover{background:var(--paper-2)}
.mr-filters{display:flex;gap:8px;margin:0 0 22px}
.mr-filters button{font:inherit;font-size:12.5px;padding:5px 14px;border:1px solid var(--line);background:transparent;border-radius:7px;cursor:pointer;color:var(--ink-3)}
.mr-filters button.active{background:var(--ink);color:var(--paper);border-color:var(--ink)}
.mr-kpis{display:grid;grid-template-columns:repeat(5,1fr);gap:14px;margin-bottom:26px}
.mr-kpi{border:1px solid var(--line);border-radius:12px;padding:16px 18px;background:var(--paper)}
.mr-kpi .v{font-size:27px;font-weight:700;letter-spacing:-.02em;color:var(--ink);font-variant-numeric:tabular-nums}
.mr-kpi .l{font-size:11.5px;color:var(--ink-4);margin-top:5px}
.mr-tabs{display:flex;gap:2px;border-bottom:1px solid var(--line);margin-bottom:2px;flex-wrap:wrap}
.mr-tabs button{font:inherit;font-size:13px;padding:9px 16px;border:none;background:none;cursor:pointer;color:var(--ink-3);border-bottom:2px solid transparent;margin-bottom:-1px}
.mr-tabs button.active{color:var(--ink);border-bottom-color:var(--accent);font-weight:600}
.mr-table{width:100%;border-collapse:collapse;font-size:13px}
.mr-table th{text-align:left;font-weight:600;color:var(--ink-4);font-size:11px;text-transform:uppercase;letter-spacing:.04em;padding:11px 12px;border-bottom:1px solid var(--line)}
.mr-table td{padding:10px 12px;border-bottom:1px solid var(--paper-2);color:var(--ink-2)}
.mr-table tr:hover td{background:var(--paper-2)}
.mr-table .num{font-variant-numeric:tabular-nums;font-family:ui-monospace,monospace}
.mr-table .pc{font-family:ui-monospace,monospace;font-size:12px;color:var(--accent)}
.mr-empty{padding:30px;text-align:center;color:var(--ink-4);font-size:13px}
.mr-names{max-width:240px;color:var(--ink-2);font-size:12.5px;line-height:1.6;white-space:normal}
.mr-pager{display:flex;align-items:center;justify-content:flex-end;gap:14px;padding:14px 4px;font-size:12.5px;color:var(--ink-3)}
.mr-pager button{font:inherit;font-size:12.5px;padding:5px 12px;border:1px solid var(--line);border-radius:7px;background:transparent;cursor:pointer;color:var(--ink-2)}
.mr-pager button:hover:not(:disabled){background:var(--paper-2)}
.mr-pager button:disabled{opacity:.4;cursor:not-allowed}
@media(max-width:760px){.mr-kpis{grid-template-columns:repeat(2,1fr)}.wrap{padding:24px 18px}header.topbar{padding:12px 18px;overflow-x:auto}}
</style></head>
<body>
<header class="topbar">
  <a class="brand-link" href="/tools"><svg viewBox="0 0 24 24" fill="currentColor" aria-hidden="true"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span class="brand">天枢<span class="sep">·</span>裁决</span></a>
  <nav>
    <a href="/tools">开始</a>
    <a href="/catalog">工具</a>
    <a href="/reports">报告</a>
    <a href="/settings" class="superadmin-only" style="display:none">设置</a>
    <a href="/usage" class="active">使用统计</a>
    <a href="/admin/users" class="admin-only" style="display:none">用户管理</a>
  </nav>
</header>
<div class="wrap">
  <h1>使用统计</h1>
  <p class="sub">跨全员的运行 · 项目（编号）· 工具 · 模型 · 成本 —— 实时统计,删除报告即从中移除。成本按「模型档位单价 × token」实算。</p>
  <section id="multi-report">
    <div class="mr-head"><button class="mr-dl" id="mr-export">↓ 导出 Excel</button></div>
    <div class="mr-filters" id="mr-filters">
      <button data-days="7">近 7 天</button>
      <button data-days="30">近 30 天</button>
      <button data-days="0" class="active">全部</button>
    </div>
    <div class="mr-kpis" id="mr-kpis"></div>
    <div class="mr-tabs" id="mr-tabs">
      <button data-tab="person" class="active">按个人</button>
      <button data-tab="project">按项目(编号)</button>
      <button data-tab="tool">按工具</button>
      <button data-tab="model">按模型</button>
      <button data-tab="detail">明细</button>
    </div>
    <div id="mr-body"></div>
  </section>
</div>
<script>
(function(){
  const sec = document.getElementById('multi-report');
  if (!sec) return;
  let _days = 0, _tab = 'person', _data = null, _page = 1;
  const PS = 15;
  const names = a => (!a||!a.length) ? '—' : a.join('、');
  const fmtT = ts => ts ? new Date(ts*1000).toLocaleDateString('zh-CN') : '—';
  const esc = s => String(s==null?'':s).replace(/[&<>]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;'}[c]));
  const money = v => v ? ('$'+v) : '—';
  const fmtTok = n => !n ? '—' : (n>=1e6 ? (n/1e6).toFixed(1)+'M' : (n>=1e3 ? Math.round(n/1e3)+'K' : ''+n));
  const kpi = (v,l) => `<div class="mr-kpi"><div class="v">${v}</div><div class="l">${l}</div></div>`;
  const tbl = (head, rows) => rows.length
    ? `<table class="mr-table"><thead><tr>${head.map(h=>`<th>${h}</th>`).join('')}</tr></thead><tbody>${rows.join('')}</tbody></table>`
    : '<div class="mr-empty">该时间范围内暂无记录</div>';
  function pager(total, pages){
    if (total <= PS) return '<div class="mr-pager"><span>共 '+total+' 条</span></div>';
    return '<div class="mr-pager">'
      + '<button id="pg-prev"'+(_page<=1?' disabled':'')+'>← 上一页</button>'
      + '<span>第 '+_page+' / '+pages+' 页 · 共 '+total+' 条</span>'
      + '<button id="pg-next"'+(_page>=pages?' disabled':'')+'>下一页 →</button></div>';
  }
  function renderTab(){
    const d = _data; if (!d) return;
    let head, rows;
    if (_tab==='person'){ head=['执行人','运行次数','涉及项目','最常用工具','Tokens','成本','最近活跃'];
      rows = d.by_person.map(p=>`<tr><td><strong>${esc(p.owner)}</strong></td><td class="num">${p.runs}</td><td class="num">${p.projects}</td><td>${esc(p.top_tool)}</td><td class="num">${fmtTok(p.tokens)}</td><td class="num">${money(p.cost)}</td><td>${fmtT(p.last)}</td></tr>`); }
    else if (_tab==='project'){ head=['产品编号','产品名称','运行次数','参与人','涉及工具','Tokens','成本','时间跨度'];
      rows = d.by_project.map(p=>`<tr><td class="pc">${esc(p.code)}</td><td>${esc(p.name||'—')}</td><td class="num">${p.runs}</td><td class="mr-names">${esc(names(p.user_names))}</td><td class="num">${p.tools}</td><td class="num">${fmtTok(p.tokens)}</td><td class="num">${money(p.cost)}</td><td>${fmtT(p.first)} ~ ${fmtT(p.last)}</td></tr>`); }
    else if (_tab==='tool'){ head=['测试类型','运行次数','使用人','涉及项目','Tokens','成本'];
      rows = d.by_tool.map(p=>`<tr><td><strong>${esc(p.name)}</strong></td><td class="num">${p.runs}</td><td class="mr-names">${esc(names(p.user_names))}</td><td class="num">${p.projects}</td><td class="num">${fmtTok(p.tokens)}</td><td class="num">${money(p.cost)}</td></tr>`); }
    else if (_tab==='model'){ head=['模型','运行次数','使用人','Tokens','成本'];
      rows = (d.by_model||[]).map(p=>`<tr><td><strong>${esc(p.model)}</strong></td><td class="num">${p.runs}</td><td class="mr-names">${esc(names(p.user_names))}</td><td class="num">${fmtTok(p.tokens)}</td><td class="num">${money(p.cost)}</td></tr>`); }
    else { head=['时间','执行人','产品编号','测试类型','模型','Tokens','成本','结果'];
      rows = d.detail.map(r=>`<tr><td>${fmtT(r.ts)}</td><td>${esc(r.owner)}</td><td class="pc">${esc(r.project_code||'—')}</td><td>${esc(r.tool_name)}</td><td>${esc(r.model_label||'—')}</td><td class="num">${fmtTok(r.tokens)}</td><td class="num">${money(r.cost_usd)}</td><td>${r.status==='succeeded'?'成功':esc(r.status)}</td></tr>`); }
    const total = rows.length, pages = Math.max(1, Math.ceil(total/PS));
    if (_page > pages) _page = pages;
    const b = document.getElementById('mr-body');
    b.innerHTML = tbl(head, rows.slice((_page-1)*PS, _page*PS)) + pager(total, pages);
    const pv = document.getElementById('pg-prev'), nx = document.getElementById('pg-next');
    if (pv) pv.onclick = ()=>{ if(_page>1){ _page--; renderTab(); } };
    if (nx) nx.onclick = ()=>{ if(_page<pages){ _page++; renderTab(); } };
  }
  async function loadStats(){
    try {
      const r = await fetch('/api/usage/stats?days='+_days, {credentials:'same-origin'});
      if (!r.ok) return;
      _data = await r.json();
      const t = _data.totals;
      document.getElementById('mr-kpis').innerHTML =
        kpi(t.runs,'总运行次数')+kpi(t.users,'活跃执行人')+kpi(t.projects,'覆盖项目')+kpi(t.tools,'用到工具')+kpi(money(t.cost),'总成本 USD');
      renderTab();
    } catch(e){}
  }
  sec.querySelectorAll('#mr-filters button').forEach(btn=>btn.onclick=()=>{
    _days = parseInt(btn.dataset.days); _page = 1;
    sec.querySelectorAll('#mr-filters button').forEach(b=>b.classList.toggle('active', b===btn));
    loadStats();
  });
  sec.querySelectorAll('#mr-tabs button').forEach(btn=>btn.onclick=()=>{
    _tab = btn.dataset.tab; _page = 1;
    sec.querySelectorAll('#mr-tabs button').forEach(b=>b.classList.toggle('active', b===btn));
    renderTab();
  });
  const exp = document.getElementById('mr-export');
  if (exp) exp.onclick = ()=>{ window.location.href = '/api/usage/export.xlsx?days='+_days; };
  loadStats();
})();
</script>
</body></html>"""


@app.get("/usage", response_class=HTMLResponse)
async def usage_page(request: Request):
    # 「使用统计」仅管理员及超管可见;其余访问回工具页
    user = getattr(request.state, "current_user", None)
    if not user or not user.is_admin():
        return RedirectResponse("/tools", status_code=302)
    return _inject_shared_overlays(USAGE_HTML)


@app.get("/tools/{tool_id}", response_class=HTMLResponse)
async def tool_detail_page(tool_id: str) -> str:
    if not any(t["id"] == tool_id for t in TOOL_CATALOG):
        raise HTTPException(404, f"unknown tool: {tool_id}")
    return _inject_shared_overlays(TOOL_DETAIL_HTML)


# =====================================================================
# Settings — 本地 Claude 接入状态、模型与 effort、工具环境需求
# =====================================================================

import shutil as _shutil
import subprocess as _sp
from datetime import datetime, timezone


def _read_json_safe(p: Path) -> dict[str, Any]:
    try:
        return _json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


# Single source of truth for Claude model registry. 档位化设计:每项的 model
# 字段用档位别名(opus/sonnet/haiku),由 Claude CLI 自动解析到账号最新版本
# (4.8/4.9…),无需随 Anthropic 发版手动改版本号。真实版本在探测后读回展示。
# Each entry:
#   key:                 stable identifier used in API requests / localStorage
#   model:               档位别名 (opus/sonnet/haiku) — 自动解析最新版本
#   label:               UI display label (不含版本号,版本由探测动态补)
#   version_badge:       small badge text (1M / Legacy / null)
#   tag:                 short Chinese capability descriptor
#   default:             True for the model with the ✓ in desktop
#   legacy:              True → shown grey at bottom of list
#   betas:               beta flags to pass via ClaudeAgentOptions
#   supports_effort / supports_thinking:  capability flags
#   supported_efforts / supported_thinking:  per-model allowed lists
CLAUDE_MODELS: list[dict[str, Any]] = [
    {
        "key": "opus",
        "model": "opus",  # 档位别名 → Claude CLI 自动解析账号最新 Opus(4.8/4.9…)
        "label": "Opus",
        "version_badge": None,
        "tag": "最强推理 · 慢",
        "default": True,
        "legacy": False,
        "betas": [],
        "supports_effort": True,
        "supports_thinking": True,
        "supported_efforts": ["low", "medium", "high", "xhigh", "max"],
        "supported_thinking": ["disabled", "adaptive", "enabled"],
    },
    {
        "key": "opus-1m",
        "model": "opus",
        "label": "Opus",
        "version_badge": "1M",
        "tag": "100 万上下文 · 适合超长输入",
        "default": False,
        "legacy": False,
        "betas": ["context-1m-2025-08-07"],
        "supports_effort": True,
        "supports_thinking": True,
        "supported_efforts": ["low", "medium", "high", "xhigh", "max"],
        "supported_thinking": ["disabled", "adaptive", "enabled"],
    },
    {
        "key": "sonnet",
        "model": "sonnet",
        "label": "Sonnet",
        "version_badge": None,
        "tag": "平衡 · 默认推荐",
        "default": False,
        "legacy": False,
        "betas": [],
        "supports_effort": True,
        "supports_thinking": True,
        "supported_efforts": ["low", "medium", "high", "xhigh", "max"],
        "supported_thinking": ["disabled", "adaptive", "enabled"],
    },
    {
        "key": "haiku",
        "model": "haiku",
        "label": "Haiku",
        "version_badge": None,
        "tag": "最快 · 简单任务",
        "default": False,
        "legacy": False,
        "betas": [],
        "supports_effort": False,
        "supports_thinking": False,
        "supported_efforts": [],
        "supported_thinking": [],
    },
]


def _model_by_key(key: str) -> dict[str, Any] | None:
    hit = next((m for m in CLAUDE_MODELS if m["key"] == key), None)
    if hit:
        return hit
    # 兼容旧的本地保存键(档位化前写死版本的 key,如 opus-4-7 / sonnet-4-6 / haiku-4-5):
    # 按"是否带 1M"+ 档位词收敛到新键,避免老用户下拉选项失效。
    k = (key or "").lower()
    if "opus" in k:
        return _model_by_key("opus-1m" if "1m" in k else "opus")
    if "haiku" in k:
        return next((m for m in CLAUDE_MODELS if m["key"] == "haiku"), None)
    if "sonnet" in k:
        return next((m for m in CLAUDE_MODELS if m["key"] == "sonnet"), None)
    return None


# ---- 实时模型列表:直接读 Anthropic /v1/models(版本/精度/上下文,全部实时,不写死)----
_LIVE_MODELS_CACHE: dict[str, Any] = {"at": 0.0, "models": None}
_ANTHROPIC_MODELS_URL = "https://api.anthropic.com/v1/models"
_LIVE_MODELS_TTL = 60.0  # 秒;/v1/models 是轻量元数据接口,不耗推理 token


def _tier_word(model_id: str) -> str | None:
    ml = (model_id or "").lower()
    if "opus" in ml:
        return "opus"
    if "haiku" in ml:
        return "haiku"
    if "sonnet" in ml:
        return "sonnet"
    return None


async def _fetch_live_models() -> list[dict[str, Any]] | None:
    """实时从 Anthropic /v1/models 拉账号可用模型(版本 / 精度能力 / 上下文)。

    返回与 CLAUDE_MODELS 同构的列表(前端无需改);失败返回 None,调用方回退静态表。
    结果缓存 _LIVE_MODELS_TTL 秒。
    """
    now = _time.time()
    cache = _LIVE_MODELS_CACHE
    if cache.get("models") is not None and (now - cache.get("at", 0)) < _LIVE_MODELS_TTL:
        return cache["models"]
    try:
        await ensure_fresh_oauth_token()
    except Exception:
        pass
    headers = {"anthropic-version": "2023-06-01"}
    try:
        from packages.core.auth_config import get_api_key, get_oauth_access_token
    except Exception:
        return None
    token = None
    try:
        token = get_oauth_access_token()
    except Exception:
        token = None
    if token:
        headers["authorization"] = f"Bearer {token}"
        headers["anthropic-beta"] = "oauth-2025-04-20"
    else:
        try:
            key = get_api_key()
        except Exception:
            key = None
        if not key:
            return None
        headers["x-api-key"] = key
    import httpx
    try:
        async with httpx.AsyncClient(timeout=15.0) as cli:
            resp = await cli.get(_ANTHROPIC_MODELS_URL, headers=headers, params={"limit": 100})
        if resp.status_code != 200:
            return None
        data = resp.json().get("data", []) or []
    except Exception:
        return None
    import re as _re
    effort_order = ["low", "medium", "high", "xhigh", "max"]
    out: list[dict[str, Any]] = []
    for m in data:
        mid = m.get("id")
        if not mid:
            continue
        cap = m.get("capabilities", {}) or {}
        eff = cap.get("effort", {}) or {}
        efforts = [e for e in effort_order if isinstance(eff.get(e), dict) and eff[e].get("supported")]
        think = cap.get("extended_thinking") or cap.get("thinking") or {}
        supports_think = bool(think.get("supported")) if isinstance(think, dict) else False
        ctx = m.get("max_input_tokens") or 0
        dname = m.get("display_name") or mid
        # 从 display_name 解析版本号用于正确排序("Claude Opus 4.8" → (4,8);"...4" → (4,0))
        vm = _re.search(r"(\d+)(?:\.(\d+))?", dname)
        ver = (int(vm.group(1)), int(vm.group(2) or 0)) if vm else (0, 0)
        out.append({
            "key": mid,
            "model": mid,  # 全 ID → 选哪个版本就跑哪个版本
            "label": dname,
            "version_badge": "1M" if ctx and ctx >= 1_000_000 else None,
            "tag": "",  # 排序后按 tier + 是否最新填
            "default": False,
            "legacy": False,
            "betas": [],
            "supports_effort": bool(efforts),
            "supports_thinking": supports_think,
            "supported_efforts": efforts,
            "supported_thinking": (["disabled", "adaptive", "enabled"] if supports_think else []),
            "context": ctx,
            "_tier": _tier_word(mid) or "",
            "_ver": ver,
        })
    if not out:
        return None
    # 排序:tier 优先级 opus>sonnet>haiku;同 tier 内按版本号新→旧(数值,不是字符串)
    rank = {"opus": 0, "sonnet": 1, "haiku": 2}
    out.sort(key=lambda x: (rank.get(x["_tier"], 9), -x["_ver"][0], -x["_ver"][1]))
    # 标签:每个 tier 第一个(最新)给能力描述 + 「最新」;其余标「上一代」
    tier_desc = {"opus": "最强推理 · 慢", "sonnet": "平衡 · 默认推荐", "haiku": "最快 · 简单任务"}
    seen: set[str] = set()
    for x in out:
        t = x["_tier"]
        if t not in seen:
            seen.add(t)
            base = tier_desc.get(t, "")
            x["tag"] = (base + " · 最新") if base else "最新"
            if t == "opus":
                x["default"] = True  # 最新 Opus = 默认选中
        else:
            x["tag"] = "上一代"
        x.pop("_tier", None)
        x.pop("_ver", None)
    if not any(x.get("default") for x in out):
        out[0]["default"] = True
    cache["models"] = out
    cache["at"] = now
    return out


# ----- 模型可用性探测 -----
# 用最小 token 调一次,把不可用的模型(账号未开通 / 区域限制 / 模型已下线)
# 在前端 disable 掉。降级到运行报告里的"失败"已经在 _run_tool_async 里有兜底,
# 这是个预防——让用户看一眼就知道哪些模型能用。
_MODEL_PROBE_CACHE: dict[str, dict[str, Any]] = {}
# 单次会话内的"运行时黑名单":只要某个 model 在真实运行里报 SDK 失败,
# 立即标 unavailable + reason,避免用户继续踩同一个坑。下次 server 重启清空。
_MODEL_RUNTIME_BLACKLIST: dict[str, str] = {}


def _record_model_failure(model_key: str, reason: str) -> None:
    """运行失败时调用 — 把模型加入会话黑名单。

    只记录"看起来是模型本身不行"的情况(claude-haiku 命名错误、未开通、找不到等),
    不记录通用网络错误,否则会把好模型也错误标记。
    """
    low = (reason or "").lower()
    looks_like_model_issue = any(
        kw in low for kw in (
            "not found", "model_not_found", "404", "permission",
            "unsupported", "not available", "无权限", "未开通",
            "not enabled", "model not enabled",
        )
    )
    if looks_like_model_issue and model_key:
        _MODEL_RUNTIME_BLACKLIST[model_key] = reason[:240]


def _looks_like_model_specific_failure(msg: str) -> bool:
    """判断 SDK 错误信息是不是"这个模型本身不可用"。

    - 是: "model_not_found" / "not enabled" / "permission" / "404"
    - 否: 网络抖动 / 认证失败 / 通用 "returned an error result" (CLI 本身就坏了)
    通用失败时只记 probe ok=False,但 NOT 影响 dropdown 的 available 标记 —
    避免把"CLI 整个挂了"误判成"这一个模型不可用"。
    """
    low = (msg or "").lower()
    return any(kw in low for kw in (
        "model_not_found", "model not found", "not enabled",
        "404", "permission denied", "no access", "not available",
        "未开通", "无权限", "model is not supported",
    ))


async def _probe_model(model_key: str) -> dict[str, Any]:
    """对一个模型做最小成本探测:发一个小 JSON 请求,看 SDK 是否能完成一轮。

    结果缓存 10 分钟。返回 {ok, model, error, model_specific, checked_at}.
    model_specific=True 才会被 /api/claude/info 标 available=False;
    通用失败(网络 / 认证 / CLI 损坏)不归咎于模型本身。
    """
    cached = _MODEL_PROBE_CACHE.get(model_key)
    now = _time.time()
    if cached and (now - cached.get("checked_at", 0)) < 600:
        return cached
    entry = _model_by_key(model_key)
    if not entry and _tier_word(model_key):
        # 实时列表里的全 ID 键(如 claude-opus-4-7)直接当模型 ID 探测
        entry = {"model": model_key}
    if not entry:
        result = {"ok": False, "model": model_key, "error": "unknown model key",
                  "model_specific": True, "checked_at": now}
        _MODEL_PROBE_CACHE[model_key] = result
        return result
    # 探测前先确保 OAuth token 没过期 — 否则探测会因 401 失败,
    # 把"接入没刷新"误判成"模型不可用"。
    try:
        await ensure_fresh_oauth_token()
    except Exception:
        pass
    try:
        from packages.core.llm.client import LlmClient
        client = LlmClient(model_override=entry["model"])
        # 探测请求要"够丰满":LlmClient 默认会注入"必须返回 JSON / 严格运行规则"等长系统约束,
        # 太短的请求(单字 "ok")会让 CLI 子进程返回 is_error=True 而误判。
        # 给一个明确要求 JSON 输出的小问题,任何可用模型都能秒回。
        resp = await client.complete(
            system="你是一个连接测试小助手。请严格按 JSON 输出。",
            messages=[{
                "role": "user",
                "content": '只返回 JSON 对象 {"ok": true}，不要任何其他文字。',
            }],
            allow_degrade=False,
            max_tokens=64,
        )
        sample = (resp.text or "").strip()
        result = {
            "ok": True,
            "model": entry["model"],
            # 探测时别名解析出的真实版本号(如 claude-opus-4-8-…)— 用于 UI 动态显示版本
            "resolved_model": getattr(resp, "model_id", None) or entry["model"],
            "model_key": model_key,
            "checked_at": now,
            "sample_text": sample[:120],
            "model_specific": False,
        }
    except Exception as exc:
        msg = f"{type(exc).__name__}: {str(exc)[:240]}"
        model_specific = _looks_like_model_specific_failure(msg)
        result = {
            "ok": False,
            "model": entry["model"],
            "model_key": model_key,
            "checked_at": now,
            "error": msg,
            "model_specific": model_specific,
        }
        if model_specific:
            _record_model_failure(model_key, msg)
    _MODEL_PROBE_CACHE[model_key] = result
    return result


@app.post("/api/claude/probe_model")
async def api_claude_probe_model(body: dict[str, Any]) -> dict[str, Any]:
    """手动触发对一个模型的可用性探测。前端在用户选实验模型时调用。"""
    key = (body or {}).get("model_key")
    if not key:
        raise HTTPException(400, "missing model_key")
    return await _probe_model(str(key))


@app.get("/api/device/status")
async def api_device_status(request: Request) -> dict[str, Any]:
    """设备自检 — 连接宿主机 Android 模拟器并列出在线设备。

    用于「APP 上传 → 模拟器运行」功能的连通性检查。需登录。
    """
    require_user(request)
    from packages.core.device import adb_status
    return await adb_status()


@app.post("/api/device/screencap")
async def api_device_screencap(request: Request) -> dict[str, Any]:
    """截宿主模拟器当前屏 — 验证容器→模拟器→PNG 通路。需登录。"""
    require_user(request)
    from packages.core.device import _first_online_serial, screencap, AdbError
    serial = await _first_online_serial()
    if not serial:
        raise HTTPException(503, "无在线模拟器设备")
    out_dir = Path(settings.evidence_output_dir) / "screenshots"
    out_dir.mkdir(parents=True, exist_ok=True)
    import uuid as _uuid
    fname = f"screencap_{_uuid.uuid4().hex[:8]}.png"
    fpath = out_dir / fname
    try:
        await screencap(serial, str(fpath))
    except AdbError as exc:
        raise HTTPException(502, f"截图失败: {exc}")
    size = fpath.stat().st_size if fpath.exists() else 0
    return {"ok": size > 0, "serial": serial, "filename": fname, "size": size, "path": str(fpath)}


@app.get("/api/claude/info")
async def api_claude_info() -> dict[str, Any]:
    """探测本地 Claude Code 的接入状态、版本、登录账户、默认 effort/model。

    完全只读：不写任何 ~/.claude 文件。
    用于工具详情页"模型 & 精度"区域和 /settings 页。
    """
    home = Path.home()
    candidates = [
        os.environ.get("CLAUDE_BIN"),
        _shutil.which("claude"),
        str(home / ".local" / "bin" / "claude"),
        str(home / "bin" / "claude"),
        str(home / ".npm-global" / "bin" / "claude"),
        "/opt/homebrew/bin/claude",
        "/usr/local/bin/claude",
    ]
    def _safe_exists(p: str) -> bool:
        try:
            return Path(p).exists()
        except (PermissionError, OSError):
            return False
    bin_path = next(
        (c for c in candidates if c and _safe_exists(c)),
        None,
    )
    bin_exists = bool(bin_path)

    version = None
    if bin_exists:
        try:
            out = _sp.run(
                [bin_path, "--version"], capture_output=True, text=True, timeout=5
            )
            version = (out.stdout or out.stderr).strip()
        except Exception as exc:
            version = f"(error: {exc})"

    # 默认精度由 toolkit 自己管理 — 不再读 ~/.claude/settings.json
    # 这个值只是个"默认 dropdown 选项"，跟登录态、订阅、账号都无关。
    # 用户每次跑工具时仍可单独选其他精度。
    effort = "medium"

    state_path = home / ".claude.json"
    state = _read_json_safe(state_path) if state_path.exists() else {}
    oauth = state.get("oauthAccount") or {}

    # 优先用 toolkit 自己的接入信息(API Key 或 web OAuth);
    # ~/.claude.json 的 oauthAccount 是本机 claude CLI 残留,只是兜底
    toolkit_mode = "unset"
    toolkit_api_masked = None
    toolkit_oauth_acc: dict[str, Any] = {}
    try:
        from packages.core.auth_config import (
            get_api_key, get_auth_mode, get_oauth_account, get_oauth_access_token,
            mask_api_key,
        )
        toolkit_mode = get_auth_mode()
        if toolkit_mode == "api_key":
            k = get_api_key()
            if k:
                toolkit_api_masked = mask_api_key(k)
        elif toolkit_mode == "oauth" and get_oauth_access_token():
            toolkit_oauth_acc = get_oauth_account() or {}
    except Exception:
        pass

    if toolkit_mode == "api_key" and toolkit_api_masked:
        account = {
            "mode_label": "API Key 模式",
            "display_name": toolkit_api_masked,
            "billing_type": "api",
        }
        auth_state = "toolkit_api_key"
    elif toolkit_mode == "oauth" and toolkit_oauth_acc:
        account = {
            "mode_label": "OAuth 模式",
            "email": toolkit_oauth_acc.get("email"),
            "display_name": toolkit_oauth_acc.get("display_name") or toolkit_oauth_acc.get("email"),
            "billing_type": toolkit_oauth_acc.get("billing_type") or "subscription",
            "organization_name": toolkit_oauth_acc.get("organization_name"),
        }
        auth_state = "toolkit_oauth"
    elif oauth:
        # 兜底:本机有残留的 claude login,显示但标"本机 CLI 凭据"
        account = {
            "mode_label": "本机 CLI 凭据",
            "email": oauth.get("emailAddress"),
            "display_name": oauth.get("displayName"),
            "billing_type": oauth.get("billingType"),
        }
        auth_state = "external_claude_cli"
    else:
        account = None
        auth_state = "未接入" if toolkit_mode == "unset" else "未授权"

    install_method = state.get("installMethod")
    first_start = state.get("firstStartTime")
    first_token = state.get("claudeCodeFirstTokenDate")

    # available_models 是 UI 下拉的源头。
    # 优先用 /v1/models 实时列表(账号当前全部可用模型 + 版本 + 精度能力);
    # 拉取失败(无 token / 网络)回退到内置档位静态表。
    # 只在"模型本身明确不可用"时才 disable(model_specific 失败,或运行时黑名单);
    # 通用 SDK 失败(CLI 整个挂、网络抖动)只在 dropdown 显示警告,不禁用。
    source_models = await _fetch_live_models() or CLAUDE_MODELS
    models_live = source_models is not CLAUDE_MODELS
    models_with_avail: list[dict[str, Any]] = []
    for m in source_models:
        key = m["key"]
        probe = _MODEL_PROBE_CACHE.get(key)
        blacklisted = _MODEL_RUNTIME_BLACKLIST.get(key)
        avail = True
        unavail_reason: str | None = None
        last_probe_status: str | None = None
        if blacklisted:
            avail = False
            unavail_reason = blacklisted
        elif probe and probe.get("ok") is False and probe.get("model_specific"):
            avail = False
            unavail_reason = probe.get("error")
        # 通用探测失败也回传给前端做提示,但不 disable
        if probe and probe.get("ok") is False and not probe.get("model_specific"):
            last_probe_status = (probe.get("error") or "")[:240]
        elif probe and probe.get("ok"):
            last_probe_status = "ok"
        models_with_avail.append({
            **m,
            "available": avail,
            "unavailable_reason": unavail_reason,
            "last_probe_status": last_probe_status,
            "last_probed_at": (probe or {}).get("checked_at"),
            # 探测读回的真实版本号(无探测时为 None,UI 退回只显示档位名)
            "resolved_model": (probe or {}).get("resolved_model"),
        })

    return {
        "bin_path": bin_path if bin_exists else None,
        "bin_found": bin_exists,
        "version": version,
        "install_method": install_method,
        "first_start_at": first_start,
        "first_token_at": first_token,
        "auth_state": auth_state,
        "account": account,
        "settings_effort_level": effort,
        "available_models": models_with_avail,
        "models_live": models_live,  # True=来自 /v1/models 实时列表;False=内置静态回退
        "available_efforts": [
            {"key": "low", "label": "Low", "tag": "最快"},
            {"key": "medium", "label": "Medium", "tag": "默认"},
            {"key": "high", "label": "High", "tag": "细致"},
            {"key": "xhigh", "label": "X-High", "tag": "深度推理"},
            {"key": "max", "label": "Max", "tag": "最强 · 最慢"},
        ],
        "available_thinking": [
            {"key": "disabled", "label": "关闭"},
            {"key": "adaptive", "label": "自适应"},
            {"key": "enabled", "label": "开启（自定义预算）"},
        ],
    }


# Per-tool environment requirements: which optional packages each tool benefits from.
# Plus a global pool of file-parsing deps shared by all tools (file upload).
_FILE_PARSING_REQS = [
    {"pkg": "pypdf", "required": False, "purpose": "上传 PDF 自动提取文本"},
    {"pkg": "python-docx", "required": False, "purpose": "上传 DOCX 自动提取文本"},
    {"pkg": "openpyxl", "required": False, "purpose": "上传 XLSX 自动提取文本"},
    {"pkg": "chardet", "required": False, "purpose": "非 UTF-8 文本编码自动识别"},
]
TOOL_ENV_REQUIREMENTS: dict[str, list[dict[str, Any]]] = {
    "step1": [
        {"pkg": "claude-agent-sdk", "required": True, "purpose": "本地 Claude 接入"},
        *_FILE_PARSING_REQS,
    ],
    "step2": [
        {"pkg": "claude-agent-sdk", "required": True, "purpose": "本地 Claude 接入"},
        *_FILE_PARSING_REQS,
    ],
    "step4": [
        {"pkg": "claude-agent-sdk", "required": True, "purpose": "本地 Claude 接入"},
        {"pkg": "httpx", "required": False, "purpose": "live 模式实际请求接口"},
        *_FILE_PARSING_REQS,
    ],
    "step5": [
        {"pkg": "claude-agent-sdk", "required": True, "purpose": "本地 Claude 接入"},
        {"pkg": "Pillow", "required": False, "purpose": "截图差异比对"},
        *_FILE_PARSING_REQS,
    ],
    "step6": [
        {"pkg": "claude-agent-sdk", "required": True, "purpose": "本地 Claude 接入"},
        {"pkg": "playwright", "required": False, "purpose": "live 模式浏览器自动化"},
        *_FILE_PARSING_REQS,
    ],
}


# PyPI distribution name → Python import name (for packages where they differ)
_PKG_IMPORT_NAME: dict[str, str] = {
    "Pillow": "PIL",
    "pillow": "PIL",
    "claude-agent-sdk": "claude_agent_sdk",
    "python-docx": "docx",
    "PyYAML": "yaml",
    "pyyaml": "yaml",
    "beautifulsoup4": "bs4",
    "opencv-python": "cv2",
    "scikit-learn": "sklearn",
}


def _check_pkg(name: str) -> dict[str, Any]:
    """Check installed via importlib.metadata, with import fallback.

    Strategy:
      1) Try importlib.metadata.version() — pip's authoritative record
      2) If that fails, try import. If the module loads we accept it
         (version='unknown')
      3) Otherwise it really is missing
    """
    import importlib
    import importlib.metadata as md

    import_name = _PKG_IMPORT_NAME.get(name, name.replace("-", "_").lower())

    # 1) Pip-registered distribution
    metadata_version: str | None = None
    try:
        metadata_version = md.version(name)
    except md.PackageNotFoundError:
        metadata_version = None

    # 2) Verify the module actually loads
    importable = True
    try:
        importlib.import_module(import_name)
    except ImportError:
        importable = False

    if metadata_version and importable:
        result = {"installed": True, "version": metadata_version}
    elif importable and not metadata_version:
        # Module loadable but no dist-info — accept with placeholder version
        result = {"installed": True, "version": "unknown"}
    elif metadata_version and not importable:
        # Metadata exists but module won't load — partial / broken install
        result = {"installed": False, "version": metadata_version, "broken": True}
    else:
        result = {"installed": False, "version": None}

    if name == "playwright" and result["installed"]:
        result["browsers_installed"] = _playwright_browsers_present()
    return result


def _playwright_browsers_present() -> bool:
    """Check whether playwright's chromium browser is downloaded.

    Mirrors playwright's own resolution:
      - If PLAYWRIGHT_BROWSERS_PATH is set (and not "0"), look ONLY there.
        That's the runtime contract — playwright does not fall back to the
        default cache when the env override is set.
      - Otherwise check the platform defaults.

    Looks for an actual `chromium*` directory so a leftover empty parent
    doesn't make us claim browsers are present.
    """
    env_path = os.environ.get("PLAYWRIGHT_BROWSERS_PATH", "").strip()
    if env_path and env_path != "0":
        candidates: list[Path] = [Path(env_path)]
    else:
        candidates = [
            Path.home() / "Library/Caches/ms-playwright",  # macOS default
            Path.home() / ".cache/ms-playwright",          # Linux
            Path(os.environ.get("LOCALAPPDATA", "")) / "ms-playwright",  # Windows
        ]
    for d in candidates:
        if not d.exists():
            continue
        try:
            for child in d.iterdir():
                if child.is_dir() and child.name.startswith("chromium"):
                    return True
        except OSError:
            continue
    return False


def _check_claude_login() -> dict[str, Any]:
    """Detect whether Claude Code is "ready to call" given the active auth mode.

    认证状态完全由 toolkit 自己的 auth.json 决定:
    - oauth   : 看 toolkit 是否存有 OAuth access_token (web OAuth flow 拿到的)
    - api_key : 看 toolkit 是否存了 sk-ant-... API Key
    - unset   : 用户从未选过模式
    """
    try:
        from packages.core.auth_config import (
            get_api_key, get_auth_mode, get_oauth_access_token,
        )
        mode = get_auth_mode()
    except Exception:
        return {"logged_in": False, "needs_login": True, "mode": "unset"}

    if mode == "unset":
        return {"logged_in": False, "needs_login": True, "mode": "unset"}

    if mode == "api_key":
        has_key = bool(get_api_key())
        return {"logged_in": has_key, "needs_login": not has_key, "mode": "api_key"}

    # mode == "oauth" — toolkit 自己存的 access_token
    has_token = bool(get_oauth_access_token())
    return {"logged_in": has_token, "needs_login": not has_token, "mode": "oauth"}


def _check_claude_cli() -> dict[str, Any]:
    """Locate the user's Claude Code CLI binary.

    Search order (handles macOS Finder-launch PATH sanitization):
      1) $CLAUDE_BIN env override
      2) shutil.which("claude") — covers anything on inherited $PATH
      3) Common user-bin locations (~/.local/bin, ~/bin, brew, etc.)

    Required at runtime — every tool dispatches to it. Without it, the
    env-banner shows a manual install hint.
    """
    home = Path.home()
    candidates = [
        os.environ.get("CLAUDE_BIN"),
        _shutil.which("claude"),
        str(home / ".local" / "bin" / "claude"),
        str(home / "bin" / "claude"),
        str(home / ".npm-global" / "bin" / "claude"),
        "/opt/homebrew/bin/claude",
        "/usr/local/bin/claude",
    ]
    # 用 try/except 包住 .exists() — 容器里 symlink 可能指向当前用户无权访问的路径
    # (cp -a 保留了 symlink target,但目标 /root/... 普通用户读不到 → PermissionError)
    def _safe_exists(p: str) -> bool:
        try:
            return Path(p).exists()
        except (PermissionError, OSError):
            return False
    bin_path = next(
        (c for c in candidates if c and _safe_exists(c)),
        None,
    )
    if not bin_path:
        return {"installed": False, "version": None}
    try:
        out = _sp.run([bin_path, "--version"], capture_output=True, text=True, timeout=5)
        ver = (out.stdout or out.stderr or "").strip()
    except Exception:
        ver = "?"
    return {"installed": True, "version": ver, "path": bin_path}


@app.get("/api/settings/ready")
async def api_settings_ready() -> dict[str, Any]:
    """启动检查 — 检查 Python 依赖 + 系统级二进制（claude）+ 浏览器。"""
    missing_required: list[dict[str, Any]] = []
    optional_missing: list[dict[str, Any]] = []

    # System: Claude Code CLI — 每个工具都依赖，未打进 .app
    claude_stat = _check_claude_cli()
    if not claude_stat["installed"]:
        missing_required.append({
            "pkg": "claude_cli",
            "label": "Claude Code CLI",
            "required": True,
            "purpose": "本地 Claude Code CLI 二进制 — 没有它所有工具都跑不起来",
            "installed": False,
            "auto_installable": True,  # JS uses to show 一键安装 button
            "install_hint": "curl -fsSL https://claude.ai/install.sh | bash",
        })
    else:
        # Claude exists — check 用户是否在设置页选过连接方式 + 凭据是否就绪
        try:
            login_state = _check_claude_login()
            login_mode = login_state.get("mode")
            if login_state.get("needs_login"):
                if login_mode == "unset":
                    # 用户从没选过 — 这是默认初始状态
                    missing_required.append({
                        "pkg": "auth_mode",
                        "label": "认证模式（必选）",
                        "required": True,
                        "purpose": "首次使用必须主动选择连接方式 — 进设置页点选 OAuth（CLI 订阅）或 API Key，本机已有的凭据不会被自动复用",
                        "installed": False,
                        "auto_installable": False,  # 必须人工到设置页选
                        "install_hint": "设置 → 认证模式 → 任选一种连接方式",
                    })
                elif login_mode == "api_key":
                    missing_required.append({
                        "pkg": "claude_api_key",
                        "label": "Anthropic API Key",
                        "required": True,
                        "purpose": "已选择 API Key 模式但还没填 key — 请到设置页输入 sk-ant-... 或切到「本地」「OAuth」模式",
                        "installed": False,
                        "auto_installable": False,
                        "install_hint": "设置 → 连接 Claude → 粘贴 API Key",
                    })
                else:
                    missing_required.append({
                        "pkg": "claude_login",
                        "label": "Claude OAuth 授权",
                        "required": True,
                        "purpose": "已选 OAuth 模式但还没在本工具完成授权 — 到「设置 → 模型接入」点「OAuth 授权」浏览器走授权流程",
                        "installed": False,
                        "auto_installable": False,
                        "install_hint": "设置 → 模型接入 → OAuth 授权（浏览器内完成）",
                    })
        except Exception:
            pass  # don't break env-check on detection failure

    # Python packages — check what each tool needs
    seen = set()
    for tool in TOOL_CATALOG:
        for r in TOOL_ENV_REQUIREMENTS.get(tool["id"], []):
            key = r["pkg"]
            if key in seen:
                continue
            seen.add(key)
            stat = _check_pkg(key)
            if not stat["installed"]:
                entry = {**r, **stat}
                if r["required"]:
                    missing_required.append(entry)
                else:
                    optional_missing.append(entry)
            elif key == "playwright" and not stat.get("browsers_installed"):
                # Special case: playwright python is installed but browsers not
                optional_missing.append({
                    "pkg": "playwright_browsers", "required": False,
                    "purpose": "playwright 已装但 chromium 浏览器未下载",
                    "installed": False,
                })
    # System info — useful for the dialog
    return {
        "ready": len(missing_required) == 0,
        "missing_required": missing_required,
        "missing_optional": optional_missing,
    }


@app.get("/api/settings/tools")
async def api_settings_tools() -> dict[str, Any]:
    """每个工具的环境需求 + 当前安装状态。"""
    out = []
    for tool in TOOL_CATALOG:
        reqs = TOOL_ENV_REQUIREMENTS.get(tool["id"], [])
        checked = []
        ok = True
        for r in reqs:
            stat = _check_pkg(r["pkg"])
            checked.append({**r, **stat})
            if r["required"] and not stat["installed"]:
                ok = False
        out.append({
            "id": tool["id"],
            "name": tool["name"],
            "icon": tool["icon"],
            "ready": ok,
            "requirements": checked,
        })
    return {"tools": out}


@app.get("/api/settings/overrides")
async def api_settings_overrides() -> dict[str, Any]:
    """已激活的提示词覆盖列表。"""
    from packages.core.prompts import list_prompt_overrides
    return {"overrides": list_prompt_overrides()}


# ----- One-click install -----

import platform as _platform
import sys as _sys

# Whitelist of package names we'll allow installing — never accept arbitrary input.
_INSTALLABLE_PKGS = {
    "claude-agent-sdk", "httpx", "Pillow", "playwright",
    "pypdf", "python-docx", "openpyxl", "chardet",
}
_PLAYWRIGHT_BROWSER_INSTALL = "playwright_browsers"  # special token


@app.get("/api/settings/system")
async def api_settings_system() -> dict[str, Any]:
    """检测当前系统 + venv，给一键安装路径。"""
    venv = _sys.prefix if hasattr(_sys, "real_prefix") or _sys.base_prefix != _sys.prefix else None
    return {
        "os": _platform.system(),
        "os_release": _platform.release(),
        "machine": _platform.machine(),
        "python": _sys.version.split()[0],
        "python_executable": _sys.executable,
        "venv": venv,
        "pip_executable": str(Path(_sys.executable).parent / "pip"),
    }


# ─── 认证模式（OAuth / API Key 二选一） ──────────────────────────────────────
# OAuth 走 Claude 订阅；API Key 直接按量计费走 Anthropic API。

class AuthModeReq(BaseModel):
    mode: str   # "oauth" | "api_key"
    api_key: str | None = None   # 仅 mode=api_key 时用


@app.get("/api/settings/auth")
async def api_settings_auth() -> dict[str, Any]:
    """返回当前认证模式 + 两种模式的就绪状态。

    UI 渲染 2 张模式卡时使用：
    - oauth: toolkit 自身是否有 OAuth access_token（web OAuth flow 拿到的）+ 账号信息
    - api_key: 是否已存 API key（带 mask 显示）

    注意：本端点完全不再读 ~/.claude/ 的任何文件 — OAuth 凭据由 toolkit
    自己通过 web OAuth flow 拿到并存进 auth.json。
    """
    try:
        from packages.core.auth_config import (
            get_api_key, get_auth_mode, get_oauth_access_token,
            get_oauth_account, mask_api_key,
        )
    except Exception:
        return {"current_mode": "oauth", "modes": {}, "error": "auth_config 模块加载失败"}

    cli_stat = _check_claude_cli()
    current_mode = get_auth_mode()  # "unset" / "oauth" / "api_key"

    # toolkit 自己存的 OAuth token
    oauth_token_present = bool(get_oauth_access_token())
    oauth_account = get_oauth_account()

    api_key = get_api_key()
    # OAuth ready: mode=oauth + toolkit 有 token + CLI 装好
    oauth_ready  = (current_mode == "oauth") and cli_stat["installed"] and oauth_token_present
    apikey_ready = (current_mode == "api_key") and cli_stat["installed"] and bool(api_key)

    return {
        "current_mode": current_mode,
        "is_unset": current_mode == "unset",
        "modes": {
            "oauth": {
                "label": "OAuth 授权",
                "tag": "推荐 · 含订阅免费",
                "selectable": True,
                "ready": oauth_ready,
                "token_present": oauth_token_present,
                "account": oauth_account,  # ← 授权后展示用
                "needs": (
                    [] if cli_stat["installed"] else ["claude_cli"]
                ) + (
                    [] if oauth_token_present else ["oauth_authorize"]
                ) + (
                    [] if current_mode == "oauth" else ["主动选择此模式"]
                ),
                "summary": "在浏览器内完成 Claude OAuth 授权 — 凭据存在本工具，不依赖本机 claude login。",
            },
            "api_key": {
                "label": "API Key",
                "tag": "按量计费 · 跳过登录",
                "selectable": True,
                "ready": apikey_ready,
                "has_api_key": bool(api_key),
                "api_key_masked": mask_api_key(api_key),
                "needs": (
                    [] if cli_stat["installed"] else ["claude_cli"]
                ) + (
                    [] if api_key else ["api_key"]
                ) + (
                    [] if current_mode == "api_key" else ["主动选择此模式"]
                ),
                "summary": "粘贴 sk-ant-... API Key，跳过 OAuth 登录。按 Anthropic 价目按量计费。",
            },
        },
    }


@app.put("/api/settings/auth")
async def api_settings_auth_put(req: AuthModeReq, request: Request) -> dict[str, Any]:
    require_superadmin(request)
    """切换认证模式 + 可选保存 API key。

    Body 示例：
        {"mode": "api_key", "api_key": "sk-ant-api03-..."}
        {"mode": "oauth"}
    """
    mode = req.mode
    if mode not in ("oauth", "api_key"):
        raise HTTPException(400, f"unsupported mode: {mode}")
    if mode == "api_key" and (not req.api_key or not req.api_key.strip()):
        try:
            from packages.core.auth_config import get_api_key
            if not get_api_key():
                raise HTTPException(400, "切换到 API Key 模式需要至少提供一次 api_key")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(400, "切换到 API Key 模式需要 api_key")
    try:
        from packages.core.auth_config import set_auth_mode
        set_auth_mode(mode, req.api_key)
    except Exception as e:
        raise HTTPException(500, f"保存失败：{e}")
    # 立即回查一遍状态返回 — 前端可直接 refresh
    return await api_settings_auth()


@app.get("/api/settings/figma-token")
async def api_get_figma_token(request: Request) -> dict[str, Any]:
    """查询是否已配置 Figma token(不回传明文)。需登录。"""
    require_user(request)
    from packages.core.auth_config import get_figma_token
    tok = get_figma_token()
    return {"configured": bool(tok), "masked": (tok[:6] + "…" + tok[-4:]) if tok else None}


@app.post("/api/settings/figma-token")
async def api_set_figma_token(request: Request, body: dict[str, Any]) -> dict[str, Any]:
    """保存/清空 Figma PAT。需登录。UI 比对拉取设计图用。"""
    require_user(request)
    from packages.core.auth_config import set_figma_token, get_figma_token
    set_figma_token((body or {}).get("token"))
    tok = get_figma_token()
    return {"ok": True, "configured": bool(tok)}


@app.get("/api/figma/status")
async def api_figma_status(request: Request) -> dict[str, Any]:
    """查询当前用户在宿主 Figma 助手的登录态 + 是否已配只读 Token。step5 轮询用。"""
    user = require_user(request)
    from packages.core.device.figma import host_runner_call
    from packages.core.auth_config import get_figma_token
    # 乙案:读图走宿主机「登录态浏览器导出」(零 REST 额度),会话注入在 default profile。
    res = await host_runner_call("GET", "/status", user="default")
    if not isinstance(res, dict):
        res = {"ok": False}
    res["session_logged_in"] = bool(res.get("logged_in"))
    res["token_configured"] = bool(get_figma_token())   # 兜底:只读 PAT(有账号级额度限制)
    return res


@app.post("/api/figma/login")
async def api_figma_login(request: Request) -> dict[str, Any]:
    """触发宿主助手弹真实 Chrome 让【当前用户】登录 Figma。step5 登录按钮调用。"""
    user = require_user(request)
    from packages.core.device.figma import host_runner_call
    r = await host_runner_call("POST", "/login", user=user.id)
    if not r.get("runner_up"):
        raise HTTPException(503, "宿主读图助手未运行(需在 Mac 上启动 figma_runner)")
    return r


@app.post("/api/figma/logout")
async def api_figma_logout(request: Request) -> dict[str, Any]:
    """退出当前用户的 Figma 登录(清该用户 profile)。"""
    user = require_user(request)
    from packages.core.device.figma import host_runner_call
    return await host_runner_call("POST", "/logout", user=user.id)


@app.post("/api/figma/preview")
async def api_figma_preview(request: Request, body: dict[str, Any]) -> dict[str, Any]:
    """读取一次 Figma 链接预览(走服务端只读 Token + REST API,无需浏览器登录)。step5 用。"""
    require_user(request)
    url = (body or {}).get("url", "").strip()
    if not url:
        raise HTTPException(400, "missing url")
    from packages.core.device.figma import (parse_figma_links, fetch_figma_frames_via_export,
                                             fetch_figma_frames_via_api)
    from packages.core.auth_config import get_figma_token
    links = parse_figma_links(url)
    if not links:
        return {"ok": False, "error": "无法从链接解析出 Figma 文件,请确认是 figma.com/design/... 链接"}
    lk = links[0]
    import uuid as _uuid
    sc_dir = Path(settings.evidence_output_dir) / "screenshots"
    sc_dir.mkdir(parents=True, exist_ok=True)
    fname = f"figma_preview_{_uuid.uuid4().hex[:8]}.png"
    # 首选:登录态浏览器导出 PDF 拆帧(零额度,只读权限也行);失败再退只读 API token
    res = await fetch_figma_frames_via_export(url, str(sc_dir / fname), max_frames=50, user="default")
    if not res.get("ok"):
        token = get_figma_token()
        if token:
            res = await fetch_figma_frames_via_api(
                lk["file_key"], token, str(sc_dir / fname), prefer_node=lk.get("node_id", ""), max_frames=8)
    if not res.get("ok"):
        return {"ok": False, "error": res.get("error")}
    from pathlib import Path as _PP
    urls = [f"/api/screenshots/{_PP(fr['path']).name}" for fr in (res.get("frames") or []) if fr.get("path")]
    return {"ok": True, "image_url": (urls[0] if urls else f"/api/screenshots/{fname}"),
            "image_urls": urls, "frames": len(res.get("frames") or [])}


@app.get("/api/settings/figma-login")
async def api_get_figma_login(request: Request) -> dict[str, Any]:
    """查询是否已配置 Figma 账号(浏览器读图用,不回传密码)。需登录。"""
    require_user(request)
    from packages.core.auth_config import get_figma_login
    fl = get_figma_login()
    return {"configured": bool(fl.get("email")), "email": fl.get("email", "")}


@app.post("/api/settings/figma-login")
async def api_set_figma_login(request: Request, body: dict[str, Any]) -> dict[str, Any]:
    """保存/清空 Figma 账号密码(走前端浏览器持久登录读设计图)。需登录。"""
    require_user(request)
    from packages.core.auth_config import set_figma_login, get_figma_login
    b = body or {}
    set_figma_login(b.get("email"), b.get("password"))
    fl = get_figma_login()
    return {"ok": True, "configured": bool(fl.get("email"))}


@app.get("/api/claude/account")
async def api_claude_account() -> dict[str, Any]:
    """OAuth 登录成功后展示用 — 从 ~/.claude.json#oauthAccount 读详细账号信息。

    UI 在 OAuth 卡 ready 时展示：邮箱 / 订阅类型 / 生效日期 / 组织（如有）。
    没登录返回 logged_in=False，前端不渲染账号块。
    """
    home = Path.home()
    has_creds = False
    account: dict[str, Any] = {}

    state_file = home / ".claude.json"
    if state_file.exists():
        try:
            state = _read_json_safe(state_file) or {}
            oa = state.get("oauthAccount") or {}
            if oa:
                has_creds = True
                account = {
                    "email": oa.get("emailAddress"),
                    "display_name": oa.get("displayName"),
                    "organization_name": oa.get("organizationName"),
                    "billing_type": oa.get("billingType"),  # subscription_pro / subscription_max / api / 等
                    "has_extra_usage": bool(oa.get("hasExtraUsageEnabled")),
                    "account_created_at": oa.get("accountCreatedAt"),
                    "subscription_created_at": oa.get("subscriptionCreatedAt"),
                    "account_uuid": (oa.get("accountUuid") or "")[:8] + "…",
                }
        except Exception:
            pass
    # 也看 ~/.claude/account.json（独立文件版本）
    acct_file = home / ".claude" / "account.json"
    if not has_creds and acct_file.exists():
        try:
            acct = _read_json_safe(acct_file) or {}
            if acct:
                has_creds = True
                account = {
                    "email": acct.get("email") or acct.get("emailAddress"),
                    "display_name": acct.get("displayName") or acct.get("name"),
                    "billing_type": acct.get("billingType") or acct.get("plan"),
                }
        except Exception:
            pass

    return {
        "logged_in": has_creds,
        "account": account,
        # 计费类型友好名（UI 显示用）
        "billing_label": _billing_label(account.get("billing_type") if account else None),
    }


def _billing_label(billing_type: str | None) -> str:
    if not billing_type:
        return ""
    mapping = {
        "subscription_max":          "Claude Max 订阅",
        "subscription_pro":          "Claude Pro 订阅",
        "subscription_team":         "Claude Team",
        "subscription_enterprise":   "Claude Enterprise",
        "google_play_subscription":  "Claude Pro 订阅 (Google Play)",
        "apple_app_store_subscription": "Claude Pro 订阅 (App Store)",
        "api":                       "API 按量计费",
        "free":                      "免费版",
    }
    return mapping.get(billing_type, billing_type)


@app.delete("/api/settings/auth/api-key")
async def api_settings_auth_clear_key(request: Request) -> dict[str, Any]:
    """清掉已存的 API Key（不切模式）。"""
    require_superadmin(request)
    try:
        from packages.core.auth_config import clear_api_key
        clear_api_key()
    except Exception as e:
        raise HTTPException(500, f"清除失败：{e}")
    return await api_settings_auth()


class DisconnectReq(BaseModel):
    """指定要断开的认证目标。

    - mode='oauth'  : 清掉 toolkit 自己存的 OAuth token（不动 ~/.claude/）
    - mode='api_key': 清空已保存的 API key + 切回 unset
    purge 现在没意义（toolkit 凭据天然独立于本机 claude），保留以兼容旧调用。
    """
    mode: str
    purge: bool = False  # 已废弃；忽略其值


@app.post("/api/settings/auth/disconnect")
async def api_settings_auth_disconnect(req: DisconnectReq, request: Request) -> dict[str, Any]:
    require_superadmin(request)
    """断开当前认证 — 只清 toolkit 自身存的凭据，不动本机 claude。

    - api_key 模式: 清掉本工具保存的 sk-ant-... + 切回 unset
    - oauth 模式: 清掉本工具存的 access_token/refresh_token/账号 + 切回 unset
    """
    mode = (req.mode or "").strip()
    if mode not in ("oauth", "api_key"):
        raise HTTPException(400, f"unsupported disconnect target: {mode}")

    deleted: list[str] = []
    errors: list[str] = []

    if mode == "api_key":
        try:
            from packages.core.auth_config import clear_api_key, set_auth_mode
            clear_api_key()
            set_auth_mode("unset")
            deleted.append("toolkit api_key")
        except Exception as e:
            errors.append(f"clear api_key: {e}")
            raise HTTPException(500, f"清除 API Key 失败：{e}")
    else:  # oauth
        try:
            from packages.core.auth_config import clear_oauth_tokens, set_auth_mode
            clear_oauth_tokens()
            set_auth_mode("unset")
            deleted.append("toolkit oauth tokens")
        except Exception as e:
            errors.append(f"clear oauth: {e}")
            raise HTTPException(500, f"清除 OAuth 失败：{e}")

    out = await api_settings_auth()
    out["disconnect_result"] = {"deleted": deleted, "errors": errors}
    return out


# ── Web OAuth flow ─────────────────────────────────────────────────────────
# 借用 Claude Code 的 OAuth client_id 在浏览器跑授权流程。
# token 存到 toolkit auth.json，LLM client 在 mode=oauth 时把 access_token
# 作为 ANTHROPIC_AUTH_TOKEN 注给 CLI 子进程。
#
# 注意：client_id 是 Claude Code 公开值；redirect_uri 是本机回调。如果 Anthropic
# 拒绝本机 redirect 或 client_id 失效，需要换 OOB 流程（用户复制粘贴 code）。

CLAUDE_OAUTH_CLIENT_ID = "9d1c250a-e61b-44d9-88ed-5944d1962f5e"
CLAUDE_OAUTH_AUTHORIZE_URL = "https://claude.ai/oauth/authorize"
CLAUDE_OAUTH_TOKEN_URL = "https://console.anthropic.com/v1/oauth/token"
CLAUDE_OAUTH_SCOPES = "org:create_api_key user:profile user:inference"
# OAuth client_id 的 redirect_uri 被 Anthropic 锁死成这个 OOB callback,
# 用户授权后会看到 code 显示在 console.anthropic.com 内部页面,需要手动复制
# 粘贴回 toolkit。本机 redirect 试过(127.0.0.1:8084)— Anthropic 拒绝。
CLAUDE_OAUTH_REDIRECT_URI = "https://console.anthropic.com/oauth/code/callback"

# 进行中的 OAuth 会话：state → (code_verifier, created_at)
_OAUTH_PENDING: dict[str, dict[str, Any]] = {}
_OAUTH_TTL_SEC = 600  # 10 分钟未完成视为失效

# OAuth token 刷新锁 — 防止并发跑工具时多个协程同时刷新
import asyncio as _asyncio_oauth
_OAUTH_REFRESH_LOCK = _asyncio_oauth.Lock()


async def ensure_fresh_oauth_token() -> dict[str, Any]:
    """确保 OAuth access_token 没过期 — 过期(或 5 分钟内将过期)就用 refresh_token 换新的。

    Anthropic OAuth access_token 只活 8 小时,过期后所有 LLM 调用拿 401。
    refresh_token 活得久(数十天),用 grant_type=refresh_token 可无人值守续期。

    返回 {"status": "fresh|refreshed|expired|not_oauth|no_refresh_token|error", "detail": ...}
    - fresh           : token 还有效,没动
    - refreshed       : 成功刷新
    - not_oauth       : 当前不是 oauth 模式,跳过
    - no_refresh_token: 没存 refresh_token,只能让 admin 重新授权
    - error           : 刷新请求失败(refresh_token 也可能过期了)
    """
    try:
        from packages.core.auth_config import (
            get_auth_mode, get_oauth_access_token, get_oauth_refresh_token,
            get_oauth_expires_at, get_oauth_account, set_oauth_tokens,
        )
    except Exception as exc:
        return {"status": "error", "detail": f"auth_config 加载失败: {exc}"}

    if get_auth_mode() != "oauth":
        return {"status": "not_oauth"}

    # 还有 >30 分钟才到期 → 不用动。阈值留大是因为单个 run 可能跑 8-10 分钟,
    # token 必须在 run 全程有效,不能跑一半过期。
    _REFRESH_THRESHOLD = 1800  # 30 分钟
    expires_at = get_oauth_expires_at() or 0
    now = _time.time()
    if expires_at and (expires_at - now) > _REFRESH_THRESHOLD:
        return {"status": "fresh", "detail": f"还有 {(expires_at-now)/60:.0f} 分钟到期"}

    refresh_token = get_oauth_refresh_token()
    if not refresh_token:
        return {"status": "no_refresh_token",
                "detail": "没有 refresh_token,需 admin 到设置页重新 OAuth 授权"}

    async with _OAUTH_REFRESH_LOCK:
        # 拿到锁后再查一次 — 可能别的协程已经刷过了
        expires_at = get_oauth_expires_at() or 0
        now = _time.time()
        if expires_at and (expires_at - now) > _REFRESH_THRESHOLD:
            return {"status": "fresh", "detail": "并发协程已刷新"}

        import httpx
        payload = {
            "grant_type": "refresh_token",
            "refresh_token": refresh_token,
            "client_id": CLAUDE_OAUTH_CLIENT_ID,
        }
        try:
            async with httpx.AsyncClient(timeout=30.0) as cli:
                resp = await cli.post(
                    CLAUDE_OAUTH_TOKEN_URL, json=payload,
                    headers={"Content-Type": "application/json",
                             "Accept": "application/json"},
                )
        except Exception as exc:
            return {"status": "error", "detail": f"刷新请求失败: {exc}"}

        if resp.status_code != 200:
            return {"status": "error",
                    "detail": f"Anthropic 拒绝刷新 (HTTP {resp.status_code}): {resp.text[:200]}"}
        try:
            data = resp.json()
        except Exception as exc:
            return {"status": "error", "detail": f"刷新返回非 JSON: {exc}"}

        new_access = data.get("access_token")
        if not new_access:
            return {"status": "error", "detail": f"刷新未返回 access_token: {_json.dumps(data)[:200]}"}
        new_refresh = data.get("refresh_token") or refresh_token  # 有些实现不轮换 refresh
        expires_in = data.get("expires_in")
        new_expires_at = int(_time.time()) + int(expires_in) if expires_in else None
        set_oauth_tokens(
            access_token=new_access,
            refresh_token=new_refresh,
            expires_at=new_expires_at,
            account=get_oauth_account() or None,
        )
        print(f"[oauth] token refreshed, expires_in={expires_in}s")
        return {"status": "refreshed",
                "detail": f"已续期,新 token {expires_in}s 后到期"}


def _gc_oauth_pending() -> None:
    now = _time.time()
    stale = [s for s, v in _OAUTH_PENDING.items() if now - v.get("created_at", 0) > _OAUTH_TTL_SEC]
    for s in stale:
        _OAUTH_PENDING.pop(s, None)


def _make_pkce() -> tuple[str, str]:
    """返回 (code_verifier, code_challenge) — PKCE S256。"""
    import base64
    import hashlib
    import secrets
    verifier = secrets.token_urlsafe(48)[:64]
    digest = hashlib.sha256(verifier.encode()).digest()
    challenge = base64.urlsafe_b64encode(digest).rstrip(b"=").decode()
    return verifier, challenge


@app.post("/api/auth/oauth/start")
async def api_auth_oauth_start(request: Request) -> dict[str, Any]:
    """启动 OAuth flow (OOB 模式) — 生成 PKCE + state,返回 authorize URL。
    只有 admin 可以改全局 Claude 凭据。

    Anthropic OAuth client_id 锁死 redirect_uri 在 console.anthropic.com 的
    OOB callback,授权完成后用户会看到 code 显示在网页上,需要手动复制粘贴。
    """
    require_superadmin(request)
    import secrets as _secrets
    _gc_oauth_pending()
    verifier, challenge = _make_pkce()
    state = _secrets.token_urlsafe(24)

    _OAUTH_PENDING[state] = {
        "code_verifier": verifier,
        "redirect_uri": CLAUDE_OAUTH_REDIRECT_URI,
        "created_at": _time.time(),
    }

    from urllib.parse import urlencode
    params = {
        "client_id": CLAUDE_OAUTH_CLIENT_ID,
        "redirect_uri": CLAUDE_OAUTH_REDIRECT_URI,
        "response_type": "code",
        "scope": CLAUDE_OAUTH_SCOPES,
        "state": state,
        "code_challenge": challenge,
        "code_challenge_method": "S256",
    }
    authorize_url = f"{CLAUDE_OAUTH_AUTHORIZE_URL}?{urlencode(params)}"
    return {
        "authorize_url": authorize_url,
        "state": state,
        "expires_in": _OAUTH_TTL_SEC,
        "mode": "oob",  # 提示前端走"粘贴 code"流程
    }


class OAuthExchangeReq(BaseModel):
    code: str
    state: str


@app.post("/api/auth/oauth/exchange")
async def api_auth_oauth_exchange(req: OAuthExchangeReq, request: Request) -> dict[str, Any]:
    """OOB 模式:用户从授权页复制 code 后,POST 到这里换 token。
    只有 admin 可以改全局 Claude 凭据。

    code 格式可能是 "abc-def#state=xxx" 或纯 "abc-def" — 都接受。
    """
    require_superadmin(request)
    _gc_oauth_pending()
    code_raw = (req.code or "").strip()
    state = (req.state or "").strip()
    if not code_raw or not state:
        raise HTTPException(400, "code 和 state 都不能为空")

    # 用户可能直接粘贴完整 URL 或 code#state= 这种格式 — 清理一下
    # Anthropic OAuth code 形如 "uuid#state=xxx",支持把 fragment 部分摘掉
    if "#" in code_raw:
        code = code_raw.split("#", 1)[0].strip()
    else:
        code = code_raw

    pending = _OAUTH_PENDING.pop(state, None)
    if not pending:
        raise HTTPException(
            410,
            "OAuth 会话失效或 state 不匹配(10 分钟有效)。请回到弹窗重新点「OAuth 授权」。"
        )

    import httpx
    payload = {
        "grant_type": "authorization_code",
        "code": code,
        "redirect_uri": pending["redirect_uri"],
        "client_id": CLAUDE_OAUTH_CLIENT_ID,
        "code_verifier": pending["code_verifier"],
        "state": state,
    }
    try:
        async with httpx.AsyncClient(timeout=30.0) as cli:
            resp = await cli.post(
                CLAUDE_OAUTH_TOKEN_URL,
                json=payload,
                headers={
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                },
            )
    except Exception as e:
        raise HTTPException(502, f"调用 Anthropic token 端点失败:{e}")

    if resp.status_code != 200:
        raise HTTPException(
            resp.status_code,
            f"Anthropic 拒绝 token 换取(HTTP {resp.status_code}):{resp.text[:400]}"
        )
    try:
        data = resp.json()
    except Exception as e:
        raise HTTPException(502, f"Anthropic 返回非 JSON:{resp.text[:400]} ({e})")

    access_token = data.get("access_token")
    if not access_token:
        raise HTTPException(502, f"Anthropic 未返回 access_token:{_json.dumps(data)[:400]}")

    refresh_token = data.get("refresh_token")
    expires_in = data.get("expires_in")
    expires_at = int(_time.time()) + int(expires_in) if expires_in else None
    acc = data.get("account") or {}
    account_info = {
        "email": acc.get("email") or data.get("email"),
        "display_name": acc.get("name") or acc.get("display_name"),
        "organization_name": (
            (data.get("organization") or {}).get("name")
            if isinstance(data.get("organization"), dict) else None
        ),
        "billing_type": acc.get("billing_type"),
        "scope": data.get("scope"),
    }
    account_info = {k: v for k, v in account_info.items() if v}

    try:
        from packages.core.auth_config import set_auth_mode, set_oauth_tokens
        set_oauth_tokens(
            access_token=access_token,
            refresh_token=refresh_token,
            expires_at=expires_at,
            account=account_info,
        )
        set_auth_mode("oauth")
    except Exception as e:
        raise HTTPException(500, f"保存 token 失败:{e}")

    return {
        "status": "ok",
        "account": account_info,
        "expires_at": expires_at,
        "has_refresh": bool(refresh_token),
    }


def _html_callback_response(ok: bool, title: str, detail: str) -> HTMLResponse:
    """OAuth callback 完成后给用户看的最小页面。"""
    color = "#10b981" if ok else "#f87171"
    icon = "✓" if ok else "✕"
    return HTMLResponse(
        f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"/>
<title>{title}</title>
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin><link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
  body{{font-family:-apple-system,BlinkMacSystemFont,'PingFang SC',sans-serif;
       background:#0a0a0a;color:#e4e4e7;margin:0;padding:0;
       display:flex;align-items:center;justify-content:center;min-height:100vh}}
  .card{{max-width:480px;padding:36px 40px;border-radius:16px;
        background:#18181b;border:1px solid #27272a;text-align:center}}
  .icon{{font-size:48px;color:{color};margin-bottom:18px}}
  h1{{margin:0 0 12px;font-size:20px;font-weight:600;letter-spacing:-.01em}}
  p{{margin:0;color:#a1a1aa;font-size:13.5px;line-height:1.6}}
  .hint{{margin-top:18px;font-size:12px;color:#71717a}}
</style></head><body>
<div class="card">
  <div class="icon">{icon}</div>
  <h1>{title}</h1>
  <p>{detail}</p>
  <p class="hint">可关闭此 tab，返回原页面查看接入状态。</p>
</div>
</body></html>""",
        status_code=200 if ok else 400,
    )


@app.get("/api/auth/oauth/callback")
async def api_auth_oauth_callback(
    code: str | None = None,
    state: str | None = None,
    error: str | None = None,
    error_description: str | None = None,
) -> HTMLResponse:
    """OAuth 回调 — 拿 code 换 token，存到 toolkit auth.json。"""
    _gc_oauth_pending()
    if error:
        return _html_callback_response(
            False, "OAuth 授权被拒绝",
            f"Anthropic 返回错误：{error}" + (f"<br/>{error_description}" if error_description else "")
        )
    if not code or not state:
        return _html_callback_response(
            False, "OAuth 回调参数缺失",
            "URL 里没有 code 或 state — 可能不是从授权页正常跳回。"
        )
    pending = _OAUTH_PENDING.pop(state, None)
    if not pending:
        return _html_callback_response(
            False, "OAuth 会话失效",
            "state 不匹配或已过期(10 分钟有效)。请回到 toolkit 重新点击「OAuth 授权」。"
        )

    # 用 code + verifier 换 token
    import httpx
    payload = {
        "grant_type": "authorization_code",
        "code": code,
        "redirect_uri": pending["redirect_uri"],
        "client_id": CLAUDE_OAUTH_CLIENT_ID,
        "code_verifier": pending["code_verifier"],
    }
    try:
        async with httpx.AsyncClient(timeout=30.0) as cli:
            resp = await cli.post(
                CLAUDE_OAUTH_TOKEN_URL,
                json=payload,
                headers={"Content-Type": "application/json", "Accept": "application/json"},
            )
        if resp.status_code != 200:
            return _html_callback_response(
                False, "换取 access_token 失败",
                f"Anthropic token endpoint 返回 {resp.status_code}：<br/>"
                f"<code style='font-family:ui-monospace,monospace;font-size:11px'>"
                f"{escape(resp.text[:400])}</code>"
            )
        data = resp.json()
    except Exception as e:
        return _html_callback_response(
            False, "OAuth 网络错误",
            f"调用 Anthropic token endpoint 时出错：<br/>{escape(str(e)[:300])}"
        )

    access_token = data.get("access_token")
    if not access_token:
        return _html_callback_response(
            False, "Anthropic 未返回 access_token",
            f"<code style='font-family:ui-monospace,monospace;font-size:11px'>"
            f"{escape(_json.dumps(data, ensure_ascii=False)[:400])}</code>"
        )
    refresh_token = data.get("refresh_token")
    expires_in = data.get("expires_in")
    expires_at = int(_time.time()) + int(expires_in) if expires_in else None
    # 部分 OAuth 响应直接含账号信息
    acc = data.get("account") or {}
    account_info = {
        "email": acc.get("email") or data.get("email"),
        "display_name": acc.get("name") or acc.get("display_name"),
        "organization_name": (data.get("organization") or {}).get("name") if isinstance(data.get("organization"), dict) else None,
        "billing_type": acc.get("billing_type"),
        "scope": data.get("scope"),
    }
    account_info = {k: v for k, v in account_info.items() if v}

    try:
        from packages.core.auth_config import set_auth_mode, set_oauth_tokens
        set_oauth_tokens(
            access_token=access_token,
            refresh_token=refresh_token,
            expires_at=expires_at,
            account=account_info,
        )
        set_auth_mode("oauth")
    except Exception as e:
        return _html_callback_response(
            False, "保存 token 失败",
            f"换到 token 但写入 auth.json 出错：{escape(str(e)[:300])}"
        )

    return _html_callback_response(
        True, "OAuth 授权成功",
        "已把凭据保存到 toolkit。回到「设置 → 模型接入」页面应已显示 OAuth 模式已接入。"
    )


# Track install jobs so the UI can poll their progress.
_INSTALL_JOBS: dict[str, dict[str, Any]] = {}


class InstallReq(BaseModel):
    target: str  # package name or "playwright_browsers"


async def _run_install_job(job_id: str, target: str) -> None:
    job = _INSTALL_JOBS[job_id]
    job["status"] = "running"
    if target == _PLAYWRIGHT_BROWSER_INSTALL:
        # Use the current playwright module to download chromium. Calls the
        # python entry point of the playwright pkg.
        return await _run_playwright_install_in_process(job)
    if target == "claude_cli":
        # Auto-install Claude Code via official curl pipe. Drops binary at
        # ~/.local/bin/claude. After install, user still needs to run
        # `claude login` (interactive OAuth) — we surface that next.
        return await _run_claude_cli_install(job)

    pip = str(Path(_sys.executable).parent / "pip")
    argv = [pip, "install", "-U", target]
    job["command"] = " ".join(argv)
    try:
        proc = await _asyncio.create_subprocess_exec(
            *argv,
            stdout=_asyncio.subprocess.PIPE,
            stderr=_asyncio.subprocess.STDOUT,
        )
        out_b, _ = await _asyncio.wait_for(proc.communicate(), timeout=600)
        out = out_b.decode("utf-8", errors="replace")
        job["log"] = out[-4000:]
        job["return_code"] = proc.returncode
        job["status"] = "succeeded" if proc.returncode == 0 else "failed"
    except _asyncio.TimeoutError:
        proc.kill()
        await proc.wait()
        job["status"] = "failed"
        job["log"] = (job.get("log") or "") + "\n[timeout after 600s]"
    except Exception as exc:
        job["status"] = "failed"
        job["log"] = f"{type(exc).__name__}: {exc}"
    finally:
        job["finished_at"] = _time.time()


async def _run_playwright_install_in_process(job: dict[str, Any]) -> None:
    """下载 chromium 浏览器到 PLAYWRIGHT_BROWSERS_PATH。

    用 playwright 自带的 CLI（同进程导入）触发，不依赖外部 pip。
    """
    job["command"] = "python -m playwright install chromium"
    job["log"] = ""
    try:
        # The playwright package ships an `install` driver we can invoke.
        # Capture its stdout/stderr so the user sees download progress.
        import io
        import subprocess
        # playwright bundles a Node driver — we need to call its CLI script.
        # The cleanest cross-mode approach: locate the playwright CLI inside
        # the python package and invoke via subprocess so output streams.
        import playwright
        from pathlib import Path as _P
        pkg_dir = _P(playwright.__file__).parent
        node_exe = pkg_dir / "driver" / "node"
        cli_js = pkg_dir / "driver" / "package" / "cli.js"
        if not (node_exe.exists() and cli_js.exists()):
            job["status"] = "failed"
            job["log"] = (
                f"找不到 playwright driver。预期路径:\n"
                f"  {node_exe}\n  {cli_js}\n"
                "可能 playwright 包安装不完整 — 试 `pip install -U playwright` 后重启。"
            )
            return
        argv = [str(node_exe), str(cli_js), "install", "chromium"]
        proc = await _asyncio.create_subprocess_exec(
            *argv,
            stdout=_asyncio.subprocess.PIPE,
            stderr=_asyncio.subprocess.STDOUT,
            env={**os.environ},  # respects PLAYWRIGHT_BROWSERS_PATH if set
        )
        out_b, _ = await _asyncio.wait_for(proc.communicate(), timeout=900)
        out = out_b.decode("utf-8", errors="replace")
        job["log"] = out[-4000:]
        job["return_code"] = proc.returncode
        job["status"] = "succeeded" if proc.returncode == 0 else "failed"
    except _asyncio.TimeoutError:
        try: proc.kill()
        except Exception: pass
        job["status"] = "failed"
        job["log"] = (job.get("log") or "") + "\n[timeout after 900s]"
    except Exception as exc:
        import traceback as _tbm
        job["status"] = "failed"
        job["log"] = f"{type(exc).__name__}: {exc}\n{_tbm.format_exc()}"
    finally:
        job["finished_at"] = _time.time()


async def _run_claude_cli_install(job: dict[str, Any]) -> None:
    """Auto-install Claude Code CLI via the official one-line installer.

    Equivalent to user running:
        curl -fsSL https://claude.ai/install.sh | bash

    Drops the binary at ~/.local/bin/claude. After this the user MUST run
    `claude login` interactively (we can't automate the OAuth) — surfaced
    in the post-install message.
    """
    job["command"] = "curl -fsSL https://claude.ai/install.sh | bash"
    job["log"] = ""
    try:
        proc = await _asyncio.create_subprocess_shell(
            "curl -fsSL https://claude.ai/install.sh | bash",
            stdout=_asyncio.subprocess.PIPE,
            stderr=_asyncio.subprocess.STDOUT,
            env={**os.environ},
        )
        out_b, _ = await _asyncio.wait_for(proc.communicate(), timeout=300)
        out = out_b.decode("utf-8", errors="replace")
        job["log"] = out[-4000:]
        job["return_code"] = proc.returncode

        # Verify by running --version on the installed binary
        installed = _check_claude_cli()
        if installed.get("installed"):
            job["status"] = "succeeded"
            job["log"] += (
                "\n\n[✓ Claude Code 二进制已安装]\n"
                f"路径：{installed.get('path')}\n"
                f"版本：{installed.get('version')}\n\n"
                "下一步：请在终端执行 `claude login` 完成账户授权 — "
                "这一步必须用户手动做（OAuth 浏览器流程）。\n"
                "完成后回到本应用「设置」页点「重新检测」。"
            )
            job["next_action"] = "manual_login"
            job["next_command"] = "claude login"
        else:
            job["status"] = "failed"
            job["log"] += (
                "\n\n[✗ 安装脚本退出但未找到 claude 二进制]\n"
                "请手动执行：\n"
                "  curl -fsSL https://claude.ai/install.sh | bash\n"
                "  claude login"
            )
    except _asyncio.TimeoutError:
        try:
            proc.kill()
        except Exception:
            pass
        job["status"] = "failed"
        job["log"] = (job.get("log") or "") + "\n[timeout after 300s — 检查网络]"
    except Exception as exc:
        import traceback as _tbm
        job["status"] = "failed"
        job["log"] = f"{type(exc).__name__}: {exc}\n{_tbm.format_exc()}"
    finally:
        job["finished_at"] = _time.time()


@app.post("/api/settings/install")
async def api_settings_install(req: InstallReq, request: Request) -> dict[str, Any]:
    """一键安装某个工具依赖。返回 job_id 后异步轮询。仅超级管理员。"""
    require_superadmin(request)
    target = req.target.strip()
    # System-level targets we can self-install on macOS:
    #   claude_cli   — curl pipe → ~/.local/bin/claude
    #   playwright_browsers — playwright + chromium download
    if target == "claude_cli":
        pass
    elif target not in _INSTALLABLE_PKGS and target != _PLAYWRIGHT_BROWSER_INSTALL:
        raise HTTPException(400, f"package not in install allowlist: {target}")
    job_id = str(uuid4())
    _INSTALL_JOBS[job_id] = {
        "job_id": job_id,
        "target": target,
        "status": "queued",
        "started_at": _time.time(),
        "finished_at": None,
        "log": "",
    }
    _asyncio.create_task(_run_install_job(job_id, target))
    return {"job_id": job_id}


@app.get("/api/settings/install/{job_id}")
async def api_settings_install_status(job_id: str) -> dict[str, Any]:
    job = _INSTALL_JOBS.get(job_id)
    if not job:
        raise HTTPException(404, f"install job {job_id} not found")
    # 过滤内部字段（asyncio.subprocess / Task 等不可 JSON 序列化）
    # 任何 _ 开头的 key 都视作内部字段
    return {k: v for k, v in job.items() if not k.startswith("_")}


# ----- Reports listing -----

@app.get("/api/reports")
async def api_reports(request: Request) -> dict[str, Any]:
    """合并：本会话内存 runs + 持久化 JSON 报告文件。
    普通用户只看自己的;admin 看全部。"""
    user = require_user(request)
    out_dir = Path(settings.report_output_dir)
    saved: list[dict[str, Any]] = []
    if out_dir.exists():
        # 匹配 step*, h5_adapt, network_resilience, seo_audit, tdr_*
        # 标准格式 `<tool_id>_<run_id>.json`
        valid_tool_ids = {t["id"] for t in TOOL_CATALOG}
        seen_paths: set[str] = set()
        # 最长前缀匹配优先，避免把 network_resilience_xxx.json 解成 network。
        for p in sorted(out_dir.glob("*.json"), key=lambda x: -x.stat().st_mtime):
            stem = p.stem
            if "_" not in stem:
                continue
            tool_id: str | None = None
            run_id: str | None = None
            for tid in sorted(valid_tool_ids | {"tdr"}, key=lambda s: -len(s)):
                if stem.startswith(tid + "_"):
                    tool_id = tid
                    run_id = stem[len(tid) + 1 :]
                    break
            if tool_id is None:
                tool_id, run_id = stem.split("_", 1)
            seen_paths.add(str(p))
            # 轻量预读 project info (只读 meta 字段,避免完整 JSON 反序列化大文件)
            pc = pn = None
            owner_uid: int | None = None
            owner_un: str | None = None
            try:
                # 文件可能很大;读前 4 KB 用 json.loads 兜底失败再扫一次
                data = _json.loads(p.read_text(encoding="utf-8"))
                m = (data or {}).get("meta") or {}
                pc = m.get("project_code")
                pn = m.get("project_name")
                owner_uid = m.get("owner_user_id")
                owner_un = m.get("owner_username") or m.get("owner_email")
            except Exception:
                pass
            # 权限过滤:普通用户跳过别人的报告 (legacy 无 owner 的报告也只对 admin 可见)
            if not _user_can_see(user, owner_uid):
                continue
            saved.append({
                "source": "file",
                "tool_id": tool_id,
                "run_id": run_id,
                "filename": p.name,
                "path": str(p),
                "size": p.stat().st_size,
                "mtime": p.stat().st_mtime,
                "project_code": pc,
                "project_name": pn,
                "owner_user_id": owner_uid,
                "owner_username": owner_un,
            })

        # 也扫嵌套 TDR 目录：<report_dir>/tdr/<run_id>/review.json
        # （TdrWorkstation.finalize 默认写到这里，跟 /api/tdr/submit mirror 文件不同）
        tdr_dir = out_dir / "tdr"
        if tdr_dir.exists():
            for review_file in sorted(tdr_dir.glob("*/review.json"), key=lambda x: -x.stat().st_mtime):
                if str(review_file) in seen_paths:
                    continue
                run_id = review_file.parent.name
                tdr_owner = None
                try:
                    tdr_data = _json.loads(review_file.read_text(encoding="utf-8"))
                    tdr_owner = (tdr_data.get("meta") or {}).get("owner_user_id")
                except Exception:
                    pass
                if not _user_can_see(user, tdr_owner):
                    continue
                saved.append({
                    "source": "file",
                    "tool_id": "tdr",
                    "run_id": run_id,
                    "filename": f"tdr/{run_id}/review.json",
                    "path": str(review_file),
                    "size": review_file.stat().st_size,
                    "mtime": review_file.stat().st_mtime,
                    "owner_user_id": tdr_owner,
                })
    in_memory = sorted(
        (r for r in _RUNS.values() if _user_can_see(user, r.get("owner_user_id"))),
        key=lambda r: -r["started_at"],
    )
    in_memory_summarized = [
        {
            "source": "memory",
            "tool_id": r["tool_id"],
            "tool_name": r.get("tool_name"),
            "run_id": r["run_id"],
            "status": r["status"],
            "started_at": r["started_at"],
            "finished_at": r.get("finished_at"),
            "report_path": r.get("report_path"),
            "usage": r.get("usage"),
            "project_code": r.get("project_code"),
            "project_name": r.get("project_name"),
            "owner_user_id": r.get("owner_user_id"),
            "owner_username": r.get("owner_username"),
        }
        for r in in_memory
    ]
    return {
        "in_memory": in_memory_summarized,
        "saved": saved,
        "total_saved": len(saved),
    }


# ══════════════════ 多维报告:使用统计(跨用户聚合,实时从当前报告集算)══════════════════
def _usage_parse_ts(meta: dict[str, Any], mtime: float) -> float:
    """报告时间戳 → epoch 秒。优先 produced_at_utc(ISO),否则文件 mtime。"""
    s = meta.get("produced_at_utc") or meta.get("produced_at")
    if s:
        try:
            import datetime as _dt
            return _dt.datetime.fromisoformat(str(s).replace("Z", "+00:00")).timestamp()
        except Exception:
            pass
    return mtime


def _collect_usage_records() -> list[dict[str, Any]]:
    """扫描当前全部报告(磁盘 + 内存)→ 扁平用量记录。跨用户(供 admin+ 聚合);
    删除的报告因文件已不在,自然不计入。
    成本 = 模型档位单价 × token 数(走 packages.core.llm.pricing.cost_usd,
    opus $15/$75、sonnet $3/$15、haiku $1/$5 每百万 token,缓存另计)。"""
    from packages.core.llm.pricing import cost_usd as _price
    recs: list[dict[str, Any]] = []
    seen: set[str] = set()
    valid = {t["id"] for t in TOOL_CATALOG}

    def _mk(model: str, u: dict[str, Any]):
        it = int(u.get("input_tokens") or 0)
        ot = int(u.get("output_tokens") or 0)
        cw = int(u.get("cache_write_tokens") or u.get("cache_creation_tokens") or 0)
        cr = int(u.get("cache_read_tokens") or 0)
        # 有模型 + 有 token → 按模型单价实算;否则退回已存 cost_usd(兜底)
        if model and (it or ot or cr):
            cost = _price(model, it, ot, cw, cr)
        else:
            cost = float(u.get("cost_usd") or 0)
        return it, ot, cr, round(cost, 4)

    out_dir = Path(settings.report_output_dir)
    if out_dir.exists():
        for p in out_dir.glob("*.json"):
            stem = p.stem
            if "_" not in stem:
                continue
            tool_id = run_id = None
            for tid in sorted(valid | {"tdr"}, key=lambda s: -len(s)):
                if stem.startswith(tid + "_"):
                    tool_id, run_id = tid, stem[len(tid) + 1:]
                    break
            if tool_id is None:
                tool_id, run_id = stem.split("_", 1)
            try:
                d = _json.loads(p.read_text(encoding="utf-8"))
            except Exception:
                continue
            m = (d or {}).get("meta") or {}
            model = m.get("model_id") or m.get("model_version") or ""
            it, ot, cr, cost = _mk(model, m.get("usage") or {})
            seen.add(run_id)
            recs.append({
                "run_id": run_id, "tool_id": tool_id,
                "owner": m.get("owner_username") or m.get("owner_email") or "(未知)",
                "project_code": m.get("project_code") or "", "project_name": m.get("project_name") or "",
                "ts": _usage_parse_ts(m, p.stat().st_mtime), "status": "succeeded",
                "model": model, "in_tok": it, "out_tok": ot, "tokens": it + ot, "cost_usd": cost,
            })
    for r in _RUNS.values():
        if r.get("run_id") in seen:
            continue
        model = r.get("model_id") or r.get("model") or ""
        it, ot, cr, cost = _mk(model, r.get("usage") or {})
        recs.append({
            "run_id": r.get("run_id"), "tool_id": r.get("tool_id"),
            "owner": r.get("owner_username") or "(未知)",
            "project_code": r.get("project_code") or "", "project_name": r.get("project_name") or "",
            "ts": r.get("started_at") or 0, "status": r.get("status") or "running",
            "model": model, "in_tok": it, "out_tok": ot, "tokens": it + ot, "cost_usd": cost,
        })
    return recs


def _usage_model_label(mid: str) -> str:
    """claude-opus-4-8 → Opus 4.8;claude-sonnet-4-6 → Sonnet 4.6。"""
    if not mid:
        return "(未知模型)"
    s = str(mid); low = s.lower()
    tier = "Opus" if "opus" in low else ("Sonnet" if "sonnet" in low else ("Haiku" if "haiku" in low else ""))
    mm = _re.search(r"(\d+)[-.](\d+)", s)
    if tier and mm:
        return f"{tier} {mm.group(1)}.{mm.group(2)}"
    return tier or s


def _usage_aggregate(days: int = 0) -> dict[str, Any]:
    recs = _collect_usage_records()
    if days and days > 0:
        cutoff = _time.time() - days * 86400
        recs = [r for r in recs if (r["ts"] or 0) >= cutoff]
    tname = {t["id"]: t.get("name", t["id"]) for t in TOOL_CATALOG}

    def agg(keyf):
        g: dict[Any, dict[str, Any]] = {}
        for r in recs:
            e = g.setdefault(keyf(r), {"runs": 0, "cost": 0.0, "tokens": 0, "users": set(), "projects": set(), "tools": {}, "first": None, "last": None})
            e["runs"] += 1; e["cost"] += r["cost_usd"]; e["tokens"] += r["tokens"]; e["users"].add(r["owner"])
            if r["project_code"]:
                e["projects"].add(r["project_code"])
            e["tools"][r["tool_id"]] = e["tools"].get(r["tool_id"], 0) + 1
            ts = r["ts"] or 0
            e["first"] = ts if e["first"] is None else min(e["first"], ts)
            e["last"] = ts if e["last"] is None else max(e["last"], ts)
        return g

    by_person = []
    for owner, e in agg(lambda r: r["owner"]).items():
        tt = max(e["tools"].items(), key=lambda x: x[1])[0] if e["tools"] else ""
        by_person.append({"owner": owner, "runs": e["runs"], "projects": len(e["projects"]),
                          "top_tool": tname.get(tt, tt), "tokens": e["tokens"], "cost": round(e["cost"], 2), "last": e["last"]})
    by_person.sort(key=lambda x: -x["cost"] if x["cost"] else -x["runs"])

    by_project = []
    for code, e in agg(lambda r: r["project_code"] or "(无编号)").items():
        pn = next((r["project_name"] for r in recs if (r["project_code"] or "(无编号)") == code and r["project_name"]), "")
        by_project.append({"code": code, "name": pn, "runs": e["runs"], "users": len(e["users"]),
                          "user_names": sorted(e["users"]),
                          "tools": len(e["tools"]), "tokens": e["tokens"], "cost": round(e["cost"], 2), "first": e["first"], "last": e["last"]})
    by_project.sort(key=lambda x: -x["runs"])

    by_tool = []
    for tid, e in agg(lambda r: r["tool_id"]).items():
        by_tool.append({"tool_id": tid, "name": tname.get(tid, tid), "runs": e["runs"],
                       "users": len(e["users"]), "user_names": sorted(e["users"]), "projects": len(e["projects"]), "tokens": e["tokens"], "cost": round(e["cost"], 2)})
    by_tool.sort(key=lambda x: -x["runs"])

    by_model = []
    for mid, e in agg(lambda r: r["model"] or "").items():
        by_model.append({"model": _usage_model_label(mid), "runs": e["runs"],
                        "users": len(e["users"]), "user_names": sorted(e["users"]), "tokens": e["tokens"], "cost": round(e["cost"], 2)})
    by_model.sort(key=lambda x: -x["cost"])

    totals = {"runs": len(recs), "users": len({r["owner"] for r in recs}),
              "projects": len({r["project_code"] for r in recs if r["project_code"]}),
              "tools": len({r["tool_id"] for r in recs}),
              "cost": round(sum(r["cost_usd"] for r in recs), 2),
              "tokens": sum(r["tokens"] for r in recs)}
    detail = sorted(recs, key=lambda r: -(r["ts"] or 0))[:1000]
    for r in detail:
        r["tool_name"] = tname.get(r["tool_id"], r["tool_id"])
        r["model_label"] = _usage_model_label(r["model"])
    return {"totals": totals, "by_person": by_person, "by_project": by_project,
            "by_tool": by_tool, "by_model": by_model, "detail": detail, "days": days}


@app.get("/api/usage/stats")
async def api_usage_stats(request: Request, days: int = 0) -> dict[str, Any]:
    """多维报告聚合:个人 / 项目 / 工具 / 总计。管理员及超管可见(跨用户)。"""
    require_admin(request)
    return _usage_aggregate(days)


@app.get("/api/usage/export.xlsx")
async def api_usage_export(request: Request, days: int = 0) -> Any:
    """多维报告导出 Excel(总览 / 按个人 / 按项目 / 按工具 / 明细)。管理员及超管。"""
    require_admin(request)
    data = _usage_aggregate(days)
    import io as _io, datetime as _dt
    from urllib.parse import quote as _q
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from packages.reporting.seo_excel import _F, _header, _row, _title_row, _widths

    def _d(ts):
        try:
            return _dt.datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M") if ts else "—"
        except Exception:
            return "—"

    rng = "全部" if not days else f"近{days}天"
    wb = Workbook(); wb.remove(wb.active)
    ws = wb.create_sheet("01 总览"); _widths(ws, [20, 26])
    _title_row(ws, "多维报告 · 使用统计总览", 2, f"统计范围:{rng}")
    t = data["totals"]; r = 4
    for lbl, val in [("总运行次数", t["runs"]), ("活跃执行人", t["users"]), ("覆盖项目", t["projects"]),
                     ("用到工具", t["tools"]), ("总成本(USD)", t["cost"]), ("总 tokens", t["tokens"])]:
        c = ws.cell(r, 1, lbl); c.font = Font(name=_F, bold=True, color="1F4E5F")
        c.fill = PatternFill("solid", fgColor="E8F1F2")
        ws.cell(r, 2, val); r += 1
    ws = wb.create_sheet("02 按个人"); _widths(ws, [20, 12, 12, 22, 12, 18])
    r = _title_row(ws, "按个人", 6); _header(ws, r, ["执行人", "运行次数", "涉及项目", "最常用工具", "成本USD", "最近活跃"]); r += 1
    for p in data["by_person"]:
        _row(ws, r, [p["owner"], p["runs"], p["projects"], p["top_tool"], p["cost"], _d(p["last"])]); r += 1
    ws = wb.create_sheet("03 按项目"); _widths(ws, [22, 24, 12, 12, 12, 12, 32])
    r = _title_row(ws, "按项目(编号)", 7); _header(ws, r, ["产品编号", "产品名称", "运行次数", "参与人数", "涉及工具", "成本USD", "时间跨度"]); r += 1
    for p in data["by_project"]:
        _row(ws, r, [p["code"], p["name"], p["runs"], p["users"], p["tools"], p["cost"], f'{_d(p["first"])} ~ {_d(p["last"])}']); r += 1
    ws = wb.create_sheet("04 按工具"); _widths(ws, [22, 12, 12, 12, 14, 12])
    r = _title_row(ws, "按工具(测试类型)", 6); _header(ws, r, ["测试类型", "运行次数", "使用人数", "涉及项目", "Tokens", "成本USD"]); r += 1
    for p in data["by_tool"]:
        _row(ws, r, [p["name"], p["runs"], p["users"], p["projects"], p["tokens"], p["cost"]]); r += 1
    ws = wb.create_sheet("05 按模型"); _widths(ws, [20, 12, 12, 14, 12])
    r = _title_row(ws, "按模型(成本 = 模型单价 × token 数)", 5); _header(ws, r, ["模型", "运行次数", "使用人数", "Tokens", "成本USD"]); r += 1
    for p in data.get("by_model", []):
        _row(ws, r, [p["model"], p["runs"], p["users"], p["tokens"], p["cost"]]); r += 1
    ws = wb.create_sheet("06 明细"); _widths(ws, [17, 16, 18, 20, 16, 12, 10, 8])
    r = _title_row(ws, "逐条明细(近 1000 条)", 8); _header(ws, r, ["时间", "执行人", "产品编号", "测试类型", "模型", "Tokens", "成本USD", "结果"]); r += 1
    for d in data["detail"]:
        _row(ws, r, [_d(d["ts"]), d["owner"], d["project_code"] or "—", d.get("tool_name") or d["tool_id"],
                     d.get("model_label") or "—", d.get("tokens") or 0, d.get("cost_usd") or 0,
                     "成功" if d["status"] == "succeeded" else d["status"]]); r += 1
    buf = _io.BytesIO(); wb.save(buf)
    fn = f"多维报告_使用统计_{rng}.xlsx"
    disp = f"attachment; filename=\"usage_report.xlsx\"; filename*=UTF-8''{_q(fn)}"
    return Response(content=buf.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": disp})


@app.get("/api/reports/export")
async def api_reports_export() -> StreamingResponse:
    """打包全部已保存报告为一个 zip 下载。

    通过 _iter_saved_report_files() 同时打包顶层 *.json 和嵌套 tdr/<rid>/review.json，
    保证 raw /tdr 提交的报告也在导出范围内。

    NOTE: 必须注册在 /api/reports/{run_id} 之前，否则 run_id 会匹配 "export"。
    """
    import io as _io
    import zipfile as _zip
    out_dir = Path(settings.report_output_dir)
    buf = _io.BytesIO()
    with _zip.ZipFile(buf, mode="w", compression=_zip.ZIP_DEFLATED) as zf:
        for p in sorted(_iter_saved_report_files(), key=lambda x: -x.stat().st_mtime):
            try:
                arc = p.relative_to(out_dir).as_posix()
            except ValueError:
                arc = p.name
            zf.write(p, arcname=arc)
    buf.seek(0)
    fname = f"tianshu-reports-{int(_time.time())}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@app.get("/api/reports/{run_id}")
async def api_report_detail(run_id: str, request: Request) -> dict[str, Any]:
    """按 run_id 拿报告：先查内存，再用 _find_report_files_by_run_id 扫盘。

    读时把 finalize substep 的统一报告契约字段提升到顶层，让 UI 直接读到
    verdict / risks / blockers / issues / cases — 兼容旧 run。
    """
    user = require_user(request)
    if run_id in _RUNS:
        state = _RUNS[run_id]
        if not _user_can_see(user, state.get("owner_user_id")):
            raise HTTPException(403, "无权访问此报告")
        if isinstance(state.get("report"), dict):
            state["report"] = _promote_contract_fields(state["report"])
        return state
    files = _find_report_files_by_run_id(run_id)
    for p in files:
        try:
            stem = p.stem
            if p.parent.parent.name == "tdr" and p.name == "review.json":
                tid = "tdr"
            else:
                tid = _parse_tool_id_from_stem(stem, run_id)
            data = _json.loads(p.read_text(encoding="utf-8"))
            owner_uid = (data.get("meta") or {}).get("owner_user_id")
            if not _user_can_see(user, owner_uid):
                raise HTTPException(403, "无权访问此报告")
            return {
                "source": "file",
                "run_id": run_id,
                "tool_id": tid,
                "report": _promote_contract_fields(data),
                "report_path": str(p),
            }
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(500, f"failed to read {p}: {exc}")
    raise HTTPException(404, f"report {run_id} not found")


@app.delete("/api/reports/{run_id}")
async def api_report_delete(run_id: str, request: Request) -> dict[str, Any]:
    """删除某条报告（内存 + 磁盘 + 嵌套 review.json）。owner 或 admin 才能删。"""
    user = require_user(request)
    # 先检查权限
    if run_id in _RUNS:
        state = _RUNS[run_id]
        if not _user_can_see(user, state.get("owner_user_id")):
            raise HTTPException(403, "无权删除此报告")
    for p in _find_report_files_by_run_id(run_id):
        try:
            data = _json.loads(p.read_text(encoding="utf-8"))
            owner_uid = (data.get("meta") or {}).get("owner_user_id")
            if not _user_can_see(user, owner_uid):
                raise HTTPException(403, "无权删除此报告")
        except HTTPException:
            raise
        except Exception:
            pass  # 读失败不阻塞,继续走删
    deleted_files: list[str] = []
    deleted_memory = False
    if run_id in _RUNS:
        del _RUNS[run_id]
        deleted_memory = True
    for p in _find_report_files_by_run_id(run_id):
        try:
            p.unlink()
            deleted_files.append(p.name)
            # 嵌套目录也清掉空 dir（tdr/<run_id>/）
            try:
                if p.parent.name == run_id and p.parent.parent.name == "tdr":
                    if not any(p.parent.iterdir()):
                        p.parent.rmdir()
            except Exception:
                pass
        except Exception as exc:
            raise HTTPException(500, f"failed to delete {p}: {exc}")
    if not deleted_memory and not deleted_files:
        raise HTTPException(404, f"report {run_id} not found")
    return {"ok": True, "memory": deleted_memory, "files": deleted_files}


# ---------- 批量删除 ----------

class BatchDeleteReq(BaseModel):
    run_ids: list[str] = []
    all_visible: bool = False  # True = 删当前用户能看到的所有报告(admin=全部, user=本人)


def _delete_one_report(run_id: str, user: UserRecord) -> tuple[bool, str | None]:
    """删除单条报告 — 返回 (成功?, 错误原因)。"""
    # 权限检查 + 收集要删的文件
    has_memory = run_id in _RUNS
    if has_memory:
        state = _RUNS[run_id]
        if not _user_can_see(user, state.get("owner_user_id")):
            return False, "无权删除"
    files = _find_report_files_by_run_id(run_id)
    for p in files:
        try:
            data = _json.loads(p.read_text(encoding="utf-8"))
            owner_uid = (data.get("meta") or {}).get("owner_user_id")
            if not _user_can_see(user, owner_uid):
                return False, "无权删除"
        except Exception:
            pass
    if not has_memory and not files:
        return False, "不存在"
    # 真删
    if has_memory:
        del _RUNS[run_id]
    for p in files:
        try:
            p.unlink()
            try:
                if p.parent.name == run_id and p.parent.parent.name == "tdr":
                    if not any(p.parent.iterdir()):
                        p.parent.rmdir()
            except Exception:
                pass
        except Exception as exc:
            return False, f"删文件失败: {exc}"
    return True, None


@app.post("/api/reports/batch_delete")
async def api_reports_batch_delete(req: BatchDeleteReq, request: Request) -> dict[str, Any]:
    """批量删除。两种模式:
      - run_ids = [...] : 删指定的这些 run
      - all_visible=True : 删当前用户能看到的所有报告(admin=系统所有,user=自己所有)
    返回 {deleted: N, failed: [{run_id, reason}], total: M}
    """
    user = require_user(request)
    targets: list[str] = []
    if req.all_visible:
        # 收集所有可见 run_id
        for r in list(_RUNS.values()):
            if _user_can_see(user, r.get("owner_user_id")):
                targets.append(r["run_id"])
        out_dir = Path(settings.report_output_dir)
        if out_dir.exists():
            valid_tool_ids = {t["id"] for t in TOOL_CATALOG} | {"tdr"}
            for p in out_dir.glob("*.json"):
                stem = p.stem
                if "_" not in stem:
                    continue
                rid = None
                for tid in sorted(valid_tool_ids, key=lambda s: -len(s)):
                    if stem.startswith(tid + "_"):
                        rid = stem[len(tid)+1:]
                        break
                if not rid or rid in targets:
                    continue
                try:
                    data = _json.loads(p.read_text(encoding="utf-8"))
                    owner_uid = (data.get("meta") or {}).get("owner_user_id")
                except Exception:
                    owner_uid = None
                if _user_can_see(user, owner_uid):
                    targets.append(rid)
    else:
        # 去重保留顺序
        seen = set()
        for rid in req.run_ids or []:
            if rid and rid not in seen:
                seen.add(rid); targets.append(rid)
    if not targets:
        return {"deleted": 0, "failed": [], "total": 0}
    deleted = 0
    failed: list[dict[str, Any]] = []
    for rid in targets:
        ok, reason = _delete_one_report(rid, user)
        if ok:
            deleted += 1
        else:
            failed.append({"run_id": rid, "reason": reason or "未知错误"})
    return {"deleted": deleted, "failed": failed, "total": len(targets)}


# ----- URL fetcher (input convenience) -----

class FetchUrlReq(BaseModel):
    url: str


_BINARY_PARSERS_AVAILABLE: dict[str, bool] = {}


def _parser_available(name: str) -> bool:
    if name in _BINARY_PARSERS_AVAILABLE:
        return _BINARY_PARSERS_AVAILABLE[name]
    try:
        if name == "pypdf":
            import pypdf  # noqa: F401
        elif name == "python-docx":
            import docx  # noqa: F401
        elif name == "openpyxl":
            import openpyxl  # noqa: F401
        elif name == "chardet":
            import chardet  # noqa: F401
        _BINARY_PARSERS_AVAILABLE[name] = True
    except ImportError:
        _BINARY_PARSERS_AVAILABLE[name] = False
    return _BINARY_PARSERS_AVAILABLE[name]


def _extract_pdf(blob: bytes) -> str:
    if not _parser_available("pypdf"):
        return "(pypdf 未安装；无法解析 PDF — 请到 /settings 安装)"
    import io
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(blob))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        try:
            txt = page.extract_text() or ""
        except Exception as exc:
            txt = f"[页 {i} 解析失败：{exc}]"
        parts.append(f"--- page {i} ---\n{txt.strip()}")
    return "\n\n".join(parts) if parts else "(空 PDF)"


def _extract_docx(blob: bytes) -> str:
    if not _parser_available("python-docx"):
        return "(python-docx 未安装；无法解析 DOCX)"
    import io
    import docx
    doc = docx.Document(io.BytesIO(blob))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables
    for ti, t in enumerate(doc.tables, 1):
        rows = []
        for row in t.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paras.append(f"\n--- table {ti} ---\n" + "\n".join(rows))
    return "\n".join(paras) if paras else "(空 DOCX)"


def _extract_xlsx(blob: bytes) -> str:
    if not _parser_available("openpyxl"):
        return "(openpyxl 未安装；无法解析 XLSX)"
    import io
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(blob), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"--- sheet: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts) if parts else "(空 XLSX)"


def _extract_text_with_encoding(blob: bytes) -> str:
    """Decode arbitrary text bytes with chardet fallback."""
    # Try UTF-8 first (BOM-tolerant)
    for enc in ("utf-8-sig", "utf-8"):
        try:
            return blob.decode(enc)
        except UnicodeDecodeError:
            continue
    # Then chardet detect
    if _parser_available("chardet"):
        import chardet
        det = chardet.detect(blob)
        enc = det.get("encoding") or "gbk"
        try:
            return blob.decode(enc, errors="replace")
        except (UnicodeDecodeError, LookupError):
            pass
    # Last resort
    return blob.decode("utf-8", errors="replace")


def _extract_image_meta(blob: bytes, name: str) -> str:
    """Return metadata stub for images — actual OCR is out of scope."""
    if not _parser_available_pillow():
        return f"(图片 {name} · {len(blob)}B — Pillow 未安装无法读取尺寸)"
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(blob))
        return f"[图片 {name}: {img.format} · {img.width}×{img.height}px · 模式 {img.mode} · {len(blob)}B]\n(图片内容暂不解析；请用文字描述图中信息)"
    except Exception as exc:
        return f"[图片 {name} · {len(blob)}B · 读取失败：{exc}]"


def _parser_available_pillow() -> bool:
    try:
        from PIL import Image  # noqa: F401
        return True
    except ImportError:
        return False


try:
    from fastapi import File, UploadFile
except ImportError:
    pass


@app.post("/api/extract-file")
async def api_extract_file(file: UploadFile = File(...)) -> dict[str, Any]:
    """从上传文件中提取文本内容。

    支持：
      文本类：.md .txt .json .yaml .yml .csv .tsv .html .xml .py .js 等
      二进制文档：.pdf / .docx / .xlsx
      图片：返回元信息摘要（暂不 OCR）

    大小限制（防止单大文件 OOM）：
      - 文本类：10 MB
      - PDF/DOCX/XLSX：50 MB
      - 图片：20 MB
    超限直接 413，不会读完整个文件到内存。
    """
    name = file.filename or "uploaded"
    ext = Path(name).suffix.lower().lstrip(".")
    ct = (file.content_type or "").lower()

    # 按类型决定上限
    is_pdf_or_doc = ext in ("pdf", "docx", "xlsx", "xlsm") or "officedocument" in ct or ct == "application/pdf"
    is_image = ct.startswith("image/") or ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp", "svg")
    # APP 安装包(APK/IPA):不读内容,识别为"待运行物料",给大上限
    is_app = ext in ("apk", "ipa") or ct in ("application/vnd.android.package-archive",)
    if is_app:
        max_bytes, kind = 500 * 1024 * 1024, "APP 安装包"
    elif is_pdf_or_doc:
        max_bytes, kind = 50 * 1024 * 1024, "PDF/DOCX/XLSX"
    elif is_image:
        max_bytes, kind = 20 * 1024 * 1024, "图片"
    else:
        max_bytes, kind = 10 * 1024 * 1024, "文本"

    # 优先走 Content-Length 提前拒绝
    cl = file.headers.get("content-length") if hasattr(file, "headers") else None
    if cl:
        try:
            if int(cl) > max_bytes:
                raise HTTPException(413, f"{kind}文件超过 {max_bytes // 1024 // 1024}MB 上限（Content-Length={cl}）")
        except (ValueError, TypeError):
            pass

    # 流式累计读取，超限立即 413
    chunks: list[bytes] = []
    received = 0
    while True:
        chunk = await file.read(64 * 1024)  # 64K 块
        if not chunk:
            break
        received += len(chunk)
        if received > max_bytes:
            raise HTTPException(413, f"{kind}文件超过 {max_bytes // 1024 // 1024}MB 上限")
        chunks.append(chunk)
    blob = b"".join(chunks)

    # Decide handler
    # APP 安装包:不读内容 — APP 上传 = "要把它跑起来",不是读文本。
    if is_app:
        if ext == "ipa":
            # iOS 包跑不了安卓模拟器
            return {
                "filename": name, "content_type": ct or "application/octet-stream",
                "size": len(blob), "kind": "app", "platform": "ios", "runnable": False,
                "text": f"[iOS 安装包 {name}（{len(blob)//1024//1024}MB）— 安卓模拟器无法运行 .ipa，"
                        f"暂不支持运行 iOS 包]",
            }
        # APK:存盘供后续在模拟器运行(P2 用),返回 app 标记
        import uuid as _uuid
        apps_dir = settings.report_output_dir.parent.parent / "uploads" / "apps"
        apps_dir.mkdir(parents=True, exist_ok=True)
        app_id = _uuid.uuid4().hex[:12]
        app_path = apps_dir / f"{app_id}.apk"
        app_path.write_bytes(blob)
        # 尝试解析包名/启动 Activity(androguard 可选,缺失则留空,装包后再从设备拿)
        pkg = launch_activity = None
        try:
            from androguard.core.apk import APK as _APK  # type: ignore
            _ap = _APK(str(app_path))
            pkg = _ap.get_package()
            launch_activity = _ap.get_main_activity()
        except Exception:
            pass
        return {
            "filename": name, "content_type": "application/vnd.android.package-archive",
            "size": len(blob), "kind": "app", "platform": "android", "runnable": True,
            "app_id": app_id, "app_path": str(app_path),
            "package": pkg, "launch_activity": launch_activity,
            "text": f"[待运行 Android APP：{name}（{len(blob)//1024//1024}MB）"
                    + (f"，包名 {pkg}" if pkg else "")
                    + f"。运行工具时将在模拟器中安装并启动以做 UI 比对。app_ref={app_id}]",
        }
    # ── 统一:所有文档/图片都"存盘 + 返回 file_ref 短标记" ──
    # 运行工具时,后端按 file_ref 取出真实文件,作为 document/image 内容块经
    # stdin 直传 LLM(PDF/图片让模型原生读,文本/Office 抽文本作命名文本文档),
    # 不再把全文回吐到前端输入框(避免污染 + 撑爆命令行参数)。
    import uuid as _uuid2
    files_dir = settings.report_output_dir.parent.parent / "uploads" / "files"
    files_dir.mkdir(parents=True, exist_ok=True)
    file_id = _uuid2.uuid4().hex[:12]

    def _store(ext_: str, kind_: str, mime_: str, data_: bytes) -> str:
        p = files_dir / f"{file_id}.{ext_}"
        p.write_bytes(data_)
        (files_dir / f"{file_id}.json").write_text(
            _json.dumps({
                "file_id": file_id, "filename": name, "kind": kind_,
                "mime": mime_, "path": str(p), "size": len(data_),
            }, ensure_ascii=False),
            encoding="utf-8",
        )
        return str(p)

    if ext == "pdf" or ct == "application/pdf":
        _store("pdf", "document", "application/pdf", blob)
        kind, mime_out = "document", "application/pdf"
    elif ct.startswith("image/") or ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp"):
        out_ext = ext if ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp") else "png"
        mime_out = ct if ct.startswith("image/") else f"image/{out_ext}"
        _store(out_ext, "image", mime_out, blob)
        kind = "image"
    elif ext == "docx" or "wordprocessingml" in ct:
        _store("txt", "document_text", "text/plain", _extract_docx(blob).encode("utf-8"))
        kind, mime_out = "document_text", "text/plain"
    elif ext in ("xlsx", "xlsm") or "spreadsheetml" in ct:
        _store("txt", "document_text", "text/plain", _extract_xlsx(blob).encode("utf-8"))
        kind, mime_out = "document_text", "text/plain"
    elif ext in ("doc", "ppt", "pptx"):
        note = (f"[.{ext} 暂不支持自动解析，建议另存为 .pdf / .docx 后重传] "
                f"文件名：{name}，大小 {len(blob)}B")
        _store("txt", "document_text", "text/plain", note.encode("utf-8"))
        kind, mime_out = "document_text", "text/plain"
    else:
        # 文本类(md/txt/json/csv/yaml/html/源码/log…):原样存盘,运行时按文本读取
        _store(ext or "txt", "document_text", "text/plain", blob)
        kind, mime_out = "document_text", "text/plain"

    _sz = len(blob)
    _hsz = f"{_sz/1024/1024:.1f}MB" if _sz >= 1024 * 1024 else f"{max(1, _sz//1024)}KB"
    return {
        "filename": name,
        "content_type": ct or "application/octet-stream",
        "size": _sz,
        "kind": kind,
        "mime": mime_out,
        "file_id": file_id,
        "file_ref": file_id,
        # 短标记:进输入框/材料文本,运行时据此加载真实文件直传 LLM
        "text": f"[📎 {name}（{_hsz}） | file_ref={file_id}]",
    }


def _is_safe_public_host(host: str) -> tuple[bool, str]:
    """判断 host 解析后是否所有 IP 都是公网 — SSRF 防御。

    返回 (safe, reason)。拒绝：loopback / private / link-local / multicast / unspecified /
    reserved。元数据地址 169.254.169.254 也被 link-local 覆盖。
    """
    import ipaddress, socket
    if not host:
        return False, "host 为空"
    # 同时拒绝直接给的字面 IP（即使域名也会被解析后再判一次）
    try:
        # getaddrinfo 拿全部解析结果（一个域名可能解析到多个 IP）
        infos = socket.getaddrinfo(host, None)
        for info in infos:
            ip_str = info[4][0]
            try:
                ip = ipaddress.ip_address(ip_str)
            except ValueError:
                continue
            if (ip.is_loopback or ip.is_private or ip.is_link_local
                or ip.is_multicast or ip.is_unspecified or ip.is_reserved):
                return False, f"{host} 解析到非公网地址 {ip_str}"
    except socket.gaierror as e:
        return False, f"DNS 解析失败：{e}"
    return True, ""


@app.post("/api/fetch-url")
async def api_fetch_url(req: FetchUrlReq) -> dict[str, Any]:
    """把任意可读 URL 抓回来作为输入文本。

    SSRF 防御：
      - 仅允许 http/https
      - 域名解析后所有 IP 必须是公网（拒绝 127.x / 10.x / 169.254.x 等）
      - 跟随重定向后**重新校验**最终落脚 URL
      - 响应大小硬上限 5MB（不只是字符串截断）
    """
    import httpx
    from urllib.parse import urlparse

    if not req.url.startswith(("http://", "https://")):
        raise HTTPException(400, "URL 必须是 http:// 或 https:// 开头")

    # 第一道：发送前校验 host
    parsed = urlparse(req.url)
    safe, reason = _is_safe_public_host(parsed.hostname or "")
    if not safe:
        raise HTTPException(403, f"拒绝抓取内网/loopback 地址：{reason}")

    MAX_BYTES = 5 * 1024 * 1024  # 5MB 硬上限
    try:
        # 改流式 — 实时限流，不等响应完整下载
        async with httpx.AsyncClient(timeout=20.0, follow_redirects=True) as cli:
            async with cli.stream("GET", req.url, headers={"User-Agent": "ai-test-toolkit/0.1"}) as r:
                r.raise_for_status()
                # 第二道：重定向后重新校验最终 URL
                final_parsed = urlparse(str(r.url))
                safe2, reason2 = _is_safe_public_host(final_parsed.hostname or "")
                if not safe2:
                    raise HTTPException(403, f"重定向到内网/loopback 地址：{reason2}")
                # Content-Length 提前拒绝（如果 server 给了准确值）
                cl = r.headers.get("content-length")
                if cl:
                    try:
                        if int(cl) > MAX_BYTES:
                            raise HTTPException(413, f"响应过大（Content-Length={cl} 超过 {MAX_BYTES // 1024 // 1024}MB 上限）")
                    except ValueError:
                        pass
                # 流式累计 — 超限立即中断（即使 server 不给 Content-Length 也能护住）
                chunks: list[bytes] = []
                received = 0
                async for chunk in r.aiter_bytes():
                    received += len(chunk)
                    if received > MAX_BYTES:
                        raise HTTPException(413, f"响应过大（流式累计超过 {MAX_BYTES // 1024 // 1024}MB 上限）")
                    chunks.append(chunk)
                raw = b"".join(chunks)
                # 解码（按 charset；fallback utf-8）
                charset = "utf-8"
                ct_header = r.headers.get("content-type", "").lower()
                if "charset=" in ct_header:
                    charset = ct_header.split("charset=", 1)[1].split(";", 1)[0].strip() or "utf-8"
                try:
                    text = raw.decode(charset, errors="replace")
                except LookupError:
                    text = raw.decode("utf-8", errors="replace")
                return {
                    "url": str(r.url),
                    "status": r.status_code,
                    "content_type": ct_header,
                    "size": received,
                    "text": text[:200_000],  # frontend cap
                }
    except HTTPException:
        raise
    except httpx.HTTPError as exc:
        raise HTTPException(502, f"fetch failed: {exc}")


# =====================================================================
# Redesigned tool detail page + Settings page
# =====================================================================

TOOL_DETAIL_HTML = r"""<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>工具 — 天枢·裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin><link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
  :root{--bg:#ffffff;--surface:#f0f0f0;--surface-2:#ebebeb;--surface-3:#dcdcdc;
    --line:#c4c4c4;--line-2:#9e9e9e;
    --fg:#0a0a0a;--fg-2:#262626;--fg-3:#4a4a4a;--fg-4:#6e6e6e;
    --ac:#a8401f;--ac-2:#c45a3a;--ac-bg:rgba(168,64,31,.14);--ac-line:rgba(168,64,31,.58);
    --warn:#8a5300;--ok:#4f6b35;--bad:#8a2d12;--info:#3f5560;
    --running:#7a4f00;
    --mono:'Inter','SF Mono',ui-monospace,Menlo,monospace;
    --sans:'Noto Sans SC','PingFang SC',-apple-system,'Microsoft YaHei',sans-serif;
    --serif:'Noto Serif SC','Songti SC','STSong',Georgia,serif;}
  @media (prefers-reduced-motion: reduce){*,*::before,*::after{animation-duration:.01ms!important;transition-duration:.01ms!important}}
  *{box-sizing:border-box}
  html,body{margin:0;background:var(--bg);color:var(--fg);font-family:var(--sans);min-height:100%;
    -webkit-font-smoothing:antialiased}
  body{background:
    radial-gradient(ellipse 90% 50% at 50% -10%, rgba(196,90,58,.07), transparent 65%) fixed,
    radial-gradient(ellipse 80% 40% at 50% 110%, rgba(255,255,255,.02), transparent 60%) fixed,
    var(--bg);}
  ::selection{background:var(--ac-bg);color:var(--ac-2)}
  ::-webkit-scrollbar{width:10px;height:10px}
  ::-webkit-scrollbar-track{background:transparent}
  ::-webkit-scrollbar-thumb{background:var(--surface-3);border-radius:5px;border:2px solid var(--bg)}
  ::-webkit-scrollbar-thumb:hover{background:var(--line-2)}

  .topbar{display:flex;align-items:center;gap:0;height:56px;
    padding:0 24px;border-bottom:1px solid var(--line);
    background:rgba(255,255,255,.94);position:sticky;top:0;z-index:30;
    backdrop-filter:saturate(180%) blur(20px);-webkit-backdrop-filter:saturate(180%) blur(20px)}
  .topbar .logo{width:28px;height:28px;flex-shrink:0;
    background:linear-gradient(135deg,#262626,#1a1a1a);border-radius:6px;
    display:grid;place-items:center;color:#001f1a;font-weight:700;font-size:14px;letter-spacing:0;
    font-family:"PingFang SC",-apple-system,sans-serif;
    margin-right:14px;
    box-shadow:0 1px 2px rgba(0,0,0,.4),inset 0 1px 0 rgba(255,255,255,.18)}
  .topbar .logo .logo-mark{width:18px;height:18px;display:block;
    filter:drop-shadow(0 .5px 0 rgba(255,255,255,.18))}
  .topbar .logo-old-replaced{width:24px;height:24px;border:1.5px solid var(--ac);border-radius:6px;
    display:grid;place-items:center;color:var(--ac);font-family:var(--mono);font-weight:700;
    font-size:11px;margin-right:14px}
  .topbar .crumbs{display:flex;align-items:center;gap:8px;font-size:13px}
  .topbar .crumbs a{color:var(--fg-3);text-decoration:none}
  .topbar .crumbs a:hover{color:var(--ac)}
  .topbar .crumbs .sep{color:var(--fg-4)}
  .topbar .crumbs .current{color:var(--fg);font-weight:500;display:flex;align-items:center;gap:6px}
  .topbar .stats{margin-left:auto;display:flex;gap:14px;align-items:center;
    font-family:var(--mono);font-size:11px;color:var(--fg-3)}
  .topbar .stats .stat .v{color:var(--fg-2);margin-left:5px}
  .topbar .stats .stat.cost .v{color:var(--ac)}
  .topbar .stats .stat.gate.reject .v{color:var(--bad)}
  .topbar .stats .stat.gate.proceed .v{color:var(--ok)}
  .topbar button.run{background:linear-gradient(180deg,var(--ac-2),var(--ac));color:#001a14;border:none;
    padding:0 18px;height:34px;border-radius:8px;font-family:var(--sans);font-size:13px;
    font-weight:600;cursor:pointer;display:flex;align-items:center;gap:8px;margin-left:14px;
    box-shadow:0 1px 2px rgba(0,0,0,.4),inset 0 1px 0 rgba(255,255,255,.24);
    transition:transform .12s,filter .12s}
  .topbar button.run:hover:not(:disabled){filter:brightness(1.06);transform:translateY(-.5px)}
  .topbar button.run:active:not(:disabled){transform:translateY(0);filter:brightness(.97)}
  .topbar button.run:disabled{background:var(--surface-3);color:var(--fg-4);cursor:not-allowed;box-shadow:none}
  .topbar button.run kbd{background:rgba(0,26,20,.20);padding:2px 6px;border-radius:4px;
    font-family:var(--mono);font-size:10.5px;font-weight:600;letter-spacing:.02em}
  .topbar a.set{color:var(--fg-3);font-size:13px;text-decoration:none;padding:6px 10px;
    border-radius:6px;margin-left:6px;transition:background .12s,color .12s}
  .topbar a.set:hover{background:var(--surface-2);color:var(--ac-2)}
  .topbar .back-btn{
    display:inline-flex;align-items:center;gap:6px;
    margin-right:14px;padding:6px 12px;border-radius:6px;
    background:var(--surface-2);border:1px solid var(--line-2);
    color:var(--fg-2);text-decoration:none;font-size:12.5px;font-weight:500;
    transition:background .12s,color .12s,border-color .12s}
  .topbar .back-btn:hover{border-color:var(--ac-line);color:var(--ac-2);background:var(--ac-bg)}

  /* === Premium hero card === */
  .hero{position:relative;padding:32px 36px 28px;
    border-bottom:1px solid var(--line);overflow:hidden}
  .hero::before{
    content:"";position:absolute;inset:0;
    background:radial-gradient(circle 380px at 12% 30%, var(--ac-bg) 0%, transparent 70%);
    opacity:.85;pointer-events:none}
  .hero-inner{position:relative;display:flex;align-items:flex-start;gap:22px;max-width:1200px;margin:0 auto}
  .hero-icon-frame{
    flex-shrink:0;width:88px;height:88px;border-radius:20px;
    background:linear-gradient(135deg,var(--surface-3) 0%,var(--surface-2) 100%);
    border:1px solid var(--line-2);
    display:grid;place-items:center;
    box-shadow:var(--shadow-2),inset 0 1px 0 rgba(255,255,255,.04);
    position:relative;overflow:hidden}
  .hero-icon-frame::before{
    content:"";position:absolute;inset:0;
    background:radial-gradient(circle at 30% 20%, rgba(110,231,183,.18), transparent 60%);
    pointer-events:none}
  .hero-icon-frame .ic{font-size:42px;line-height:1;position:relative;z-index:1;
    filter:drop-shadow(0 2px 8px rgba(0,0,0,.4))}
  .hero-body{flex:1;min-width:0}
  .hero-meta{display:flex;align-items:center;gap:8px;
    font-family:"SF Mono",ui-monospace,monospace;
    font-size:11px;color:var(--fg-3);margin-top:10px;
    text-transform:uppercase;letter-spacing:.06em}
  /* hero 右侧:接入状态紧凑显示 */
  .hero-right{margin-left:auto;display:flex;flex-direction:column;align-items:flex-end;gap:6px;
    font-family:var(--mono);font-size:11.5px;color:var(--fg-3);min-width:160px}
  .hero-right .row{display:flex;align-items:center;gap:6px}
  .hero-right .row.ok{color:var(--ok)}
  .hero-right .row.warn{color:var(--warn)}
  .hero-right .row.bad{color:var(--bad)}
  .hero-right .dot{width:6px;height:6px;border-radius:50%;background:currentColor}
  .hero-right a{color:var(--fg-3);text-decoration:none;font-size:11px}
  .hero-right a:hover{color:var(--ac)}
  .hero-meta .step{
    background:var(--ac-bg);color:var(--ac-2);
    padding:3px 10px;border-radius:4px;font-weight:600;
    border:1px solid var(--ac-line)}
  .hero-meta .resp{color:var(--fg-2)}
  .hero h2{margin:0 0 4px;font-size:24px;letter-spacing:-.02em;font-weight:600;line-height:1.2;
    color:var(--fg)}
  .hero .tag{color:var(--fg-3);font-size:13px;line-height:1.6;margin:0;max-width:780px}
  .hero-pills{display:flex;gap:8px;flex-wrap:wrap;margin-top:4px}
  .hero-pill{
    display:inline-flex;align-items:center;gap:6px;
    padding:5px 11px;border-radius:999px;
    background:var(--surface-2);border:1px solid var(--line-2);
    font-size:11.5px;color:var(--fg-2);font-family:"SF Mono",ui-monospace,monospace;
    transition:border-color .12s,color .12s}
  .hero-pill:hover{border-color:var(--ac-line);color:var(--fg)}
  .hero-pill .lbl{color:var(--fg-3);font-size:10.5px;text-transform:uppercase;letter-spacing:.05em}
  .hero-pill.success{border-color:rgba(52,211,153,.28);color:var(--ok);background:rgba(52,211,153,.06)}
  .hero-pill.fail{border-color:rgba(248,113,113,.28);color:var(--bad);background:rgba(248,113,113,.06)}
  .hero-pill.warn{border-color:rgba(251,191,36,.28);color:var(--warn);background:rgba(251,191,36,.06)}

  main{max-width:1280px;margin:0 auto;padding:0 28px 80px}
  .runner-grid{
    display:grid;
    grid-template-columns:minmax(0,1fr) 420px;
    gap:20px;
    padding-top:20px;
    align-items:start;
    transition:grid-template-columns .25s ease}
  /* 初态(无运行任务):右侧面板折叠成 56px 窄条,只显示状态 + 快捷运行提示 */
  .runner-grid.idle{grid-template-columns:minmax(0,1fr) 56px}
  .runner-grid.idle .run-panel{padding:0;cursor:pointer}
  .runner-grid.idle .run-panel:hover{border-color:var(--ac-line)}
  .runner-grid.idle .run-panel-body{display:flex;flex-direction:column;align-items:center;
    justify-content:center;gap:14px;padding:18px 4px;
    writing-mode:vertical-rl;text-orientation:upright;
    font-family:var(--mono);font-size:11px;color:var(--fg-3);letter-spacing:.1em}
  .runner-grid.idle #run-area{display:flex;flex-direction:column;align-items:center;
    justify-content:center;gap:14px}
  .runner-grid.idle #run-area > *{writing-mode:horizontal-tb;text-orientation:mixed}
  /* 隐藏空态长文本,只露 ▶ 圆按钮 + 「⌘↵」提示 */
  .runner-grid.idle .run-empty-pretty .title,
  .runner-grid.idle .run-empty-pretty .desc{display:none}
  .runner-grid.idle .run-empty-pretty{padding:14px 4px;gap:10px}
  .runner-grid.idle .run-empty-pretty .play-circle{width:36px;height:36px;font-size:16px}
  .runner-grid.idle .run-empty-pretty .hint-row{font-size:10px;flex-direction:column;
    gap:3px;writing-mode:horizontal-tb}
  .runner-grid.idle .run-empty-pretty .hint-row kbd{font-size:9.5px;padding:1px 4px}
  .runner-grid.idle .run-panel-head{
    flex-direction:column;padding:14px 6px;gap:10px;writing-mode:vertical-rl;
    text-orientation:upright;height:auto;background:transparent;border-bottom:none}
  .runner-grid.idle .run-panel-head .meta-tail{display:none}
  .workspace{min-width:0;display:flex;flex-direction:column;gap:18px}
  .run-panel{
    position:sticky;top:76px;
    max-height:calc(100vh - 96px);
    background:var(--surface);border:1px solid var(--line);
    border-radius:14px;overflow:hidden;
    display:flex;flex-direction:column;transition:none}
  .run-panel-head{
    display:flex;align-items:center;gap:10px;
    padding:16px 20px;border-bottom:1px solid var(--line);
    background:var(--surface-2);
    font-size:11px;color:var(--fg-3);
    text-transform:uppercase;letter-spacing:.08em;font-weight:600}
  .run-panel-head .live-dot{
    width:8px;height:8px;border-radius:50%;background:var(--fg-4);flex-shrink:0}
  .run-panel-head .live-dot.queued{background:var(--fg-3)}
  .run-panel-head .live-dot.running{background:var(--running);
    box-shadow:0 0 0 3px rgba(167,139,250,.18);
    animation:dot-pulse 1.4s ease-in-out infinite}
  .run-panel-head .live-dot.succeeded{background:var(--ok)}
  .run-panel-head .live-dot.failed{background:var(--bad)}
  @keyframes dot-pulse{
    0%,100%{transform:scale(1);box-shadow:0 0 0 3px rgba(167,139,250,.18)}
    50%{transform:scale(1.15);box-shadow:0 0 0 6px rgba(167,139,250,.10)}
  }
  .run-panel-head .label-text{color:var(--fg-2)}
  .run-panel-head .meta-tail{margin-left:auto;font-family:"SF Mono",ui-monospace,monospace;
    font-size:10.5px;color:var(--fg-3);text-transform:none;letter-spacing:0}
  .run-panel-body{flex:1;overflow-y:auto;padding:0}
  /* Idle/empty state inside panel */
  #run-area .run-empty-pretty{
    padding:48px 28px;text-align:center;
    display:flex;flex-direction:column;align-items:center;gap:14px}
  #run-area .run-empty-pretty .play-circle{
    width:72px;height:72px;border-radius:50%;
    background:linear-gradient(135deg,var(--ac-bg),transparent);
    border:1px solid var(--ac-line);
    display:grid;place-items:center;
    color:var(--ac-2);font-size:24px;
    box-shadow:inset 0 1px 0 rgba(255,255,255,.04)}
  #run-area .run-empty-pretty .title{
    font-size:16px;font-weight:600;color:var(--fg);margin:0}
  #run-area .run-empty-pretty .desc{
    font-size:13px;color:var(--fg-2);line-height:1.55;margin:0;max-width:280px}
  #run-area .run-empty-pretty .hint-row{
    display:flex;align-items:center;gap:8px;
    font-family:"SF Mono",ui-monospace,monospace;
    font-size:11px;color:var(--fg-3);margin-top:8px}
  #run-area .run-empty-pretty kbd{
    background:var(--surface-3);border:1px solid var(--line-2);
    border-bottom-width:2px;
    padding:2px 7px;border-radius:5px;font-size:10.5px;color:var(--fg-2);
    font-family:"SF Mono",ui-monospace,monospace;font-weight:600}
  /* Responsive: stack on narrow viewport */
  @media (max-width:1080px){
    .runner-grid{grid-template-columns:1fr}
    .run-panel{position:static;max-height:none}
  }

  /* Section card */
  .sec{background:var(--surface);border:1px solid var(--line);
    border-radius:14px;overflow:hidden;
    transition:border-color .15s,box-shadow .25s}
  .sec:hover{border-color:var(--line-2)}
  .sec-head{display:flex;align-items:center;gap:14px;padding:18px 22px;
    border-bottom:1px solid var(--line);background:var(--surface-2);
    position:relative}
  .sec-head .num{
    width:22px;height:22px;border-radius:50%;flex-shrink:0;
    display:grid;place-items:center;
    font-family:var(--mono);font-size:11px;font-weight:600;
    color:var(--bg);background:var(--fg-2);border:none}
  .sec-head h3{margin:0;font-size:14px;font-weight:600;letter-spacing:.02em;color:var(--fg-3);
    text-transform:uppercase}
  .sec-head .sub{margin-left:auto;font-size:12.5px;color:var(--fg-3);max-width:50%;
    text-align:right;line-height:1.45}
  /* 折叠式填写指引 — 标题右侧按钮,点开展示 hint */
  .hint-toggle{margin-left:auto;background:transparent;border:1px solid var(--line);
    color:var(--fg-3);padding:4px 10px;border-radius:5px;
    font-size:11.5px;cursor:pointer;font-family:var(--mono);
    transition:border-color .12s,color .12s}
  .hint-toggle:hover{border-color:var(--ac-line);color:var(--ac)}
  .hint-toggle.open{color:var(--ac);border-color:var(--ac-line);background:var(--ac-bg)}
  .hint-panel{display:none;padding:0 22px;margin:0 0 -4px}
  .hint-panel.open{display:block}
  .hint-body{padding:12px 14px;background:var(--surface-2);border-radius:6px;
    border-left:3px solid var(--ac-line);color:var(--fg-2);font-size:12.5px;
    line-height:1.7;white-space:pre-wrap;font-family:var(--sans)}
  .sec-body{padding:20px 22px}

  /* Section 1: Inputs (single textarea + upload + drag-drop) */
  @keyframes aitk-spin{to{transform:rotate(360deg)}}
  .input-zone{position:relative;border-radius:8px;transition:all .15s}
  .input-zone.dragging{outline:2px dashed var(--ac);outline-offset:4px}
  .input-zone .drop-overlay{
    position:absolute;inset:0;
    background:rgba(7,8,10,.85);backdrop-filter:blur(4px);
    border:2px dashed var(--ac);border-radius:8px;
    display:none;flex-direction:column;align-items:center;justify-content:center;
    z-index:10;pointer-events:none;
  }
  .input-zone.dragging .drop-overlay{display:flex}
  .drop-overlay-inner{display:flex;flex-direction:column;align-items:center;gap:8px}
  .drop-icon{font-size:42px;color:var(--ac);line-height:1;animation:bounce 1.4s ease-in-out infinite}
  @keyframes bounce{0%,100%{transform:translateY(0)}50%{transform:translateY(-6px)}}
  .drop-text{color:var(--fg);font-size:14px;font-weight:500}
  .drop-sub{color:var(--fg-3);font-size:11.5px;font-family:var(--mono)}
  .input-toolbar{display:flex;gap:8px;align-items:center;margin-bottom:10px}
  .input-toolbar button{background:transparent;border:1px solid var(--line-2);
    color:var(--fg-2);padding:0 14px;height:32px;border-radius:6px;
    font-family:var(--mono);font-size:12px;cursor:pointer;white-space:nowrap;
    transition:all .15s}
  .input-toolbar button:hover{border-color:var(--ac);color:var(--ac)}
  .input-toolbar button:disabled{opacity:.5;cursor:wait}
  .input-toolbar .drag-hint{font-family:var(--mono);font-size:11px;color:var(--fg-3);
    padding-left:4px;display:flex;align-items:center;gap:4px}
  .input-toolbar .drag-hint::before{content:"⇲";color:var(--fg-4)}
  .input-toolbar .size-hint{font-family:var(--mono);font-size:11px;color:var(--fg-3);
    padding:0 4px}
  #doc-input{
    background:var(--surface-2);border:1px solid var(--line);border-radius:6px;
    color:var(--fg);font-family:var(--mono);font-size:12.5px;padding:12px 14px;
    resize:vertical;min-height:240px;line-height:1.55;width:100%;
    transition:border-color .15s}
  #doc-input:focus{outline:none;border-color:var(--ac);background:var(--bg)}
  #doc-input::placeholder{color:var(--fg-4)}
  .run-options{display:flex;gap:14px;margin-top:10px;flex-wrap:wrap}
  .run-options label{display:inline-flex;align-items:center;gap:8px;
    background:var(--surface-2);border:1px solid var(--line);border-radius:6px;
    padding:8px 12px;font-size:12.5px;color:var(--fg-2);cursor:pointer}
  .run-options label input{width:14px;height:14px;accent-color:var(--ac);cursor:pointer}

  /* Section 2: Prompts (collapsible + editable + optional check) */
  .prompt-row{border-bottom:1px solid var(--line);transition:opacity .15s}
  .prompt-row:last-child{border-bottom:none}
  .prompt-row.disabled{opacity:.45}
  .prompt-row.disabled .prompt-head{cursor:default}
  .prompt-head{display:flex;align-items:center;gap:12px;padding:10px 18px;
    cursor:pointer;font-size:13px;transition:background .12s}
  .prompt-head:hover{background:var(--surface-2)}
  .prompt-head .check{accent-color:var(--ac);width:14px;height:14px;cursor:pointer;
    flex-shrink:0;margin:0}
  .prompt-head .check:disabled{cursor:not-allowed;accent-color:var(--fg-4)}
  .prompt-head .twirl{color:var(--fg-4);font-family:var(--mono);font-size:9px;
    transition:transform .15s;flex-shrink:0;width:10px}
  .prompt-row.open .prompt-head .twirl{transform:rotate(90deg);color:var(--fg-3)}
  .prompt-head .id{font-family:var(--mono);color:var(--fg-3);min-width:64px;font-size:11.5px}
  .prompt-head .name{color:var(--fg)}
  .prompt-row.disabled .id{color:var(--fg-4)}
  .prompt-row.disabled .name{color:var(--fg-3)}
  /* chip 极简:无底色,只灰字 */
  .prompt-head .chip{font-family:var(--mono);font-size:11px;padding:0;
    border-radius:0;background:transparent;border:none;
    color:var(--fg-4);margin-left:auto;letter-spacing:.02em;text-transform:uppercase}
  .prompt-head .chip.dirty{color:var(--warn);background:transparent;border:none}
  .prompt-head .chip.override{color:var(--ac);background:transparent;border:none}
  .prompt-body{display:none;padding:10px 18px 14px 18px;background:transparent;
    border-top:1px dashed var(--line)}
  .prompt-row.open .prompt-body{display:block}
  .prompt-body textarea{
    background:var(--bg);border:1px solid var(--line);border-radius:6px;
    color:var(--fg);font-family:var(--mono);font-size:12px;padding:11px 13px;
    resize:vertical;min-height:200px;line-height:1.62;width:100%}
  .prompt-body textarea:focus{outline:none;border-color:var(--ac)}
  .prompt-body .btn-row{display:flex;gap:8px;margin-top:8px;align-items:center}
  .prompt-body .btn-row button{
    background:transparent;border:1px solid var(--line-2);color:var(--fg-2);
    padding:5px 12px;border-radius:5px;font-family:var(--mono);font-size:11.5px;cursor:pointer}
  .prompt-body .btn-row button.save{background:var(--ac);color:#001a14;
    border-color:var(--ac);font-weight:600}
  .prompt-body .btn-row button.save:hover{background:var(--ac-2)}
  .prompt-body .btn-row button:hover:not(.save){border-color:var(--ac);color:var(--ac)}
  .prompt-body .btn-row .info{margin-left:auto;font-family:var(--mono);font-size:11px;color:var(--fg-3)}

  /* Section 3: Model & precision */
  .model-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:14px}
  .model-cell{background:var(--surface-2);border:1px solid var(--line);border-radius:8px;
    padding:14px}
  .model-cell .label{font-size:11.5px;color:var(--fg-3);text-transform:uppercase;
    letter-spacing:.1em;margin-bottom:8px;font-weight:600}
  .model-cell .label .from{font-family:var(--mono);font-size:10px;color:var(--fg-4);
    text-transform:none;letter-spacing:0;font-weight:400;margin-left:6px}
  .model-cell select{
    width:100%;background:var(--bg);border:1px solid var(--line-2);border-radius:6px;
    color:var(--fg);font-family:var(--mono);font-size:13px;padding:8px 10px;cursor:pointer}
  .model-cell select:focus{outline:none;border-color:var(--ac)}
  .model-cell select option.legacy{color:var(--fg-3);font-style:italic}
  .model-cell .hint{margin-top:6px;font-size:11px;color:var(--fg-3);font-family:var(--mono)}
  .model-cell.unsupported{background:rgba(196,90,58,.06);
    border:1px dashed rgba(16,185,129,.25)}
  .model-cell.unsupported .label{color:var(--fg-2)}
  .model-cell.unsupported .hint{color:var(--fg-2);font-family:var(--sans);font-size:12px;line-height:1.55}
  .claude-status{display:flex;align-items:center;gap:8px;font-family:var(--mono);font-size:11.5px;
    margin-top:14px;padding:10px 12px;background:var(--surface-2);border-radius:6px;
    border:1px solid var(--line)}
  .claude-status.ok{border-color:rgba(74,222,128,.3);background:rgba(74,222,128,.04)}
  .claude-status.bad{border-color:rgba(248,113,113,.3);background:rgba(248,113,113,.04)}
  .claude-status .dot{width:8px;height:8px;border-radius:50%;background:var(--fg-4);flex-shrink:0}
  .claude-status.ok .dot{background:var(--ok)}
  .claude-status.bad .dot{background:var(--bad)}
  .claude-status .v{color:var(--fg-2)}
  .claude-status a{color:var(--ac);text-decoration:none;margin-left:auto}

  /* Section 4: Run + report */
  .run-empty{padding:40px;text-align:center;color:var(--fg-3);font-size:13px;
    font-family:var(--mono);border:1px dashed var(--line);border-radius:8px}
  .step-list{display:flex;flex-direction:column}
  .step-row{display:grid;grid-template-columns:24px 1fr auto;gap:10px;padding:10px 16px;
    border-bottom:1px solid var(--line);align-items:center;font-size:12.5px}
  .step-row:last-child{border-bottom:none}
  .step-row .marker{font-family:var(--mono);font-size:11px;color:var(--fg-4);text-align:center}
  .step-row.done .marker{color:var(--ok)}
  .step-row.running .marker{color:var(--ac);animation:pulse 1.4s infinite}
  .step-row.failed .marker{color:var(--bad)}
  .step-row .name{color:var(--fg-2);font-family:var(--mono);font-size:12px}
  .step-row.done .name{color:var(--fg)}
  .step-row .info{font-family:var(--mono);font-size:11px;color:var(--fg-3)}
  @keyframes pulse{0%,100%{opacity:1}50%{opacity:.4}}

  .run-stats{display:grid;grid-template-columns:repeat(5,1fr);gap:10px;padding:14px 16px;
    border-bottom:1px solid var(--line);background:var(--surface-2)}
  .run-stats .cell{text-align:center}
  .run-stats .cell .v{font-family:var(--mono);font-size:14px;color:var(--ac);font-weight:600}
  .run-stats .cell .l{font-size:10px;color:var(--fg-3);text-transform:uppercase;letter-spacing:.08em;margin-top:2px}

  .gate-banner{padding:14px 18px;border-bottom:1px solid var(--line);display:flex;
    align-items:flex-start;gap:12px;font-size:13px;line-height:1.55}
  .gate-banner.proceed{background:rgba(74,222,128,.05)}
  .gate-banner.reject{background:rgba(248,113,113,.05)}
  .gate-banner.warn{background:rgba(251,191,36,.05)}
  .gate-banner .badge{padding:3px 9px;border-radius:4px;font-family:var(--mono);font-size:11px;
    font-weight:600;text-transform:uppercase;letter-spacing:.05em;flex-shrink:0;margin-top:1px}
  .gate-banner.proceed .badge{background:rgba(74,222,128,.12);color:var(--ok)}
  .gate-banner.reject .badge{background:rgba(248,113,113,.12);color:var(--bad)}
  .gate-banner.warn .badge{background:rgba(251,191,36,.12);color:var(--warn)}
  .gate-banner .reasons{color:var(--fg-2);font-family:var(--mono);font-size:12px}
  .gate-banner .reasons div{margin-top:3px}

  .json-view{padding:14px 16px;background:var(--bg);max-height:560px;overflow:auto}
  .json-view pre{margin:0;font-family:var(--mono);font-size:11.5px;line-height:1.65;
    color:var(--fg);white-space:pre-wrap;word-break:break-word}
  /* Real-time log panel */
  .log-panel{
    background:var(--bg);border-bottom:1px solid var(--line);
    max-height:280px;overflow-y:auto;padding:10px 14px;font-family:var(--mono);
    font-size:11.5px;line-height:1.5;color:var(--fg-2)}
  .log-panel-head{display:flex;align-items:center;gap:8px;padding:8px 16px 6px;
    font-family:var(--mono);font-size:10.5px;color:var(--fg-3);
    text-transform:uppercase;letter-spacing:.1em;font-weight:600;
    border-bottom:1px solid var(--line);background:var(--surface-2)}
  .log-panel-head .live{display:inline-flex;align-items:center;gap:5px}
  .log-panel-head .live .dot{width:6px;height:6px;border-radius:50%;background:var(--ac);
    animation:pulse 1.4s ease-in-out infinite}
  .log-panel-head .count{margin-left:auto;color:var(--fg-2)}
  .log-line{display:grid;grid-template-columns:64px 110px 1fr;gap:10px;
    padding:3px 0;align-items:start}
  .log-line .ts{color:var(--fg-4);font-size:10.5px}
  .log-line .ev{font-weight:600;font-size:10.5px;text-transform:uppercase;letter-spacing:.05em}
  .log-line .ev.run-start, .log-line .ev.substep-start{color:var(--ac-2)}
  .log-line .ev.substep-done, .log-line .ev.run-done{color:var(--ok)}
  .log-line .ev.llm-call{color:var(--fg-3)}
  .log-line .ev.substep-parse-failed{color:var(--warn)}
  .log-line .ev.substep-skipped{color:var(--fg-4)}
  .log-line .ev.run-failed{color:var(--bad)}
  .log-line .ev.llm-thinking{color:#c084fc}     /* purple — extended thinking */
  .log-line .ev.llm-thinking-final{color:#c084fc;font-weight:700}
  .log-line .ev.llm-text{color:var(--fg-2)}     /* normal output streaming */
  .log-line .ev.llm-text-final{color:var(--ac);font-weight:700}
  /* Stream lines: dim italic for the tail snippet */
  .log-line.stream .msg .v{color:var(--fg-2);font-style:normal}
  .log-line.stream .msg .tail{display:block;margin-top:2px;padding:4px 8px;
    background:rgba(192,132,252,.06);border-left:2px solid #c084fc;
    color:var(--fg-2);font-size:11px;line-height:1.5;
    max-height:60px;overflow:hidden;text-overflow:ellipsis;
    white-space:pre-wrap;word-break:break-word}
  .log-line.stream.text .msg .tail{background:rgba(16,185,129,.05);border-left-color:var(--ac)}
  .log-line .msg{color:var(--fg-2);white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
  .log-line .msg .k{color:var(--fg-3)}
  .log-line .msg .v{color:var(--fg)}
  .log-line .msg .v.ac{color:var(--ac)}

  /* Report view (post-run, structured) */
  .report-tabs{display:flex;gap:0;background:var(--surface-2);
    border-bottom:1px solid var(--line);padding:0 16px}
  .report-tabs button{background:transparent;border:none;color:var(--fg-3);
    padding:10px 16px;cursor:pointer;font-size:12px;font-family:var(--sans);
    border-bottom:2px solid transparent;transition:all .15s}
  .report-tabs button:hover{color:var(--fg-2)}
  .report-tabs button.active{color:var(--ac);border-bottom-color:var(--ac)}
  .report-tabs .spacer{flex:1}
  .report-tabs .export{padding:6px 12px;margin:auto 4px;background:transparent;
    border:1px solid var(--line-2);color:var(--fg-2);font-size:11px;
    border-radius:5px;cursor:pointer;font-family:var(--mono)}
  .report-tabs .export:hover{border-color:var(--ac);color:var(--ac)}

  .report-hero{display:flex;align-items:flex-start;gap:14px;padding:16px 20px;
    border-bottom:1px solid var(--line);background:var(--surface)}
  .report-hero .report-icon{font-size:22px;width:40px;height:40px;border-radius:8px;
    background:var(--surface-2);border:1px solid var(--line-2);
    display:grid;place-items:center;flex-shrink:0}
  /* === 执行摘要 — 卡片化、有层次 === */
  .exec-block{background:var(--surface);border:1px solid var(--line);border-radius:10px;
    padding:18px 22px;margin:0 22px 16px}
  .exec-block:last-of-type{margin-bottom:22px}
  .exec-head{display:flex;align-items:center;gap:12px;margin-bottom:14px;padding-bottom:12px;
    border-bottom:1px solid var(--line)}
  .exec-head h3{margin:0;font-size:13px;font-weight:700;color:var(--fg);
    text-transform:uppercase;letter-spacing:.08em}
  .exec-num{display:inline-grid;place-items:center;width:24px;height:24px;border-radius:50%;
    background:var(--fg);color:var(--bg);font-size:11.5px;font-weight:700;
    border:none;font-family:var(--mono);flex-shrink:0}
  /* verdict — 醒目左色条 + 浅底色 */
  .verdict{padding:14px 18px;background:var(--surface-2);border:1px solid var(--line);
    border-left:4px solid var(--line-2);border-radius:0 8px 8px 0;
    display:flex;align-items:center;gap:12px;font-size:20px;font-weight:700;letter-spacing:-.01em}
  .verdict.ok{color:var(--ok);border-left-color:var(--ok);background:rgba(22,163,74,.06)}
  .verdict.warn{color:var(--warn);border-left-color:var(--warn);background:rgba(202,138,4,.06)}
  .verdict.bad{color:var(--bad);border-left-color:var(--bad);background:rgba(220,38,38,.06)}
  .verdict.skip{color:var(--fg-3)}
  .verdict-icon{font-size:18px;display:inline-block}
  .verdict-summary{margin-top:12px;padding:12px 16px;background:var(--bg);border:1px solid var(--line);
    border-radius:6px;font-size:14px;line-height:1.7;color:var(--fg)}
  .sev-strip{margin-top:12px;padding:10px 14px;background:var(--bg);border:1px solid var(--line);
    border-radius:6px;font-family:var(--mono);font-size:12.5px;color:var(--fg-2);letter-spacing:.02em}
  .sev-strip strong{color:var(--fg);font-weight:700}
  .exec-muted{color:var(--fg-3);font-size:13.5px;margin:0;padding:10px 14px;
    background:var(--bg);border:1px dashed var(--line);border-radius:6px;font-style:normal}
  .exec-head .exec-count{margin-left:auto;font-family:var(--mono);font-size:12px;font-weight:600;
    color:var(--fg-2);background:var(--surface-2);border:1px solid var(--line);
    padding:3px 10px;border-radius:999px;text-transform:none;letter-spacing:0}
  .exec-head .exec-count.danger{color:#fff;background:var(--bad);border-color:var(--bad)}
  .exec-head .exec-count-note{font-size:11.5px;color:var(--fg-3);font-style:normal;
    font-weight:500;text-transform:none;letter-spacing:0;margin-left:6px}
  /* KPI 卡片 — 大字 + 卡边 */
  .exec-kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-top:14px}
  .exec-kpi{display:flex;flex-direction:column;align-items:flex-start;justify-content:center;
    padding:16px 18px;background:var(--bg);border:1px solid var(--line);border-radius:8px}
  .exec-kpi-num{font-size:30px;font-weight:700;color:var(--fg);font-family:var(--sans);
    font-variant-numeric:tabular-nums;line-height:1;letter-spacing:-.025em}
  .exec-kpi-lbl{font-size:11.5px;color:var(--fg-2);margin-top:8px;letter-spacing:.08em;
    text-transform:uppercase;font-weight:600}
  /* 严重度分布条 — 加厚到 10px 可见 */
  .sev-bar{display:flex;height:10px;border-radius:5px;overflow:hidden;margin-top:14px;
    background:var(--surface-2);border:1px solid var(--line)}
  .sev-bar-seg{display:flex;align-items:center;justify-content:center;color:#fff;font-size:0;
    line-height:0;min-width:0;transition:flex .2s}
  .sev-bar-critical{background:#dc2626}
  .sev-bar-high{background:#ea580c}
  .sev-bar-medium{background:#ca8a04}
  .sev-bar-low{background:#16a34a}
  .sev-bar-info{background:#0891b2}
  /* 优先级分布条 — 24px 大字带文字 */
  .pri-bar{display:flex;height:24px;border-radius:6px;overflow:hidden;margin-top:8px;
    margin-bottom:14px;background:var(--surface-2);border:1px solid var(--line)}
  .pri-bar-seg{display:flex;align-items:center;justify-content:center;color:#fff;
    font-size:11.5px;font-weight:700;padding:0 6px;min-width:0;
    font-family:var(--sans);letter-spacing:.04em}
  .pri-bar-P0{background:#dc2626}
  .pri-bar-P1{background:#ea580c}
  .pri-bar-P2{background:#ca8a04}
  .pri-bar-P3{background:#737373}
  /* sev-tag — 带底色的小药丸 */
  .sev-tag{font-size:10.5px;font-weight:700;padding:2px 8px;background:var(--surface-2);
    color:var(--fg-2);font-family:var(--sans);letter-spacing:.06em;
    white-space:nowrap;flex-shrink:0;text-transform:uppercase;border-radius:4px;border:1px solid var(--line)}
  .sev-tag.sev-critical{color:#dc2626;background:rgba(220,38,38,.10);border-color:rgba(220,38,38,.30)}
  .sev-tag.sev-high{color:#ea580c;background:rgba(234,88,12,.10);border-color:rgba(234,88,12,.30)}
  .sev-tag.sev-medium{color:#ca8a04;background:rgba(202,138,4,.10);border-color:rgba(202,138,4,.30)}
  .sev-tag.sev-low{color:#16a34a;background:rgba(22,163,74,.10);border-color:rgba(22,163,74,.30)}
  .sev-tag.sev-info{color:#0891b2;background:rgba(8,145,178,.10);border-color:rgba(8,145,178,.30)}
  /* 优先级标签 — 带底色 */
  .pri-tag,.meta-chip.pri-P0,.meta-chip.pri-P1,.meta-chip.pri-P2,.meta-chip.pri-P3{
    font-size:10.5px;font-weight:700;padding:2px 8px;border-radius:4px;
    background:var(--surface-2);color:var(--fg-2);font-family:var(--sans);
    letter-spacing:.04em;white-space:nowrap;border:1px solid var(--line)}
  .pri-tag.pri-P0,.meta-chip.pri-P0{color:#dc2626;background:rgba(220,38,38,.10);border-color:rgba(220,38,38,.30)}
  .pri-tag.pri-P1,.meta-chip.pri-P1{color:#ea580c;background:rgba(234,88,12,.10);border-color:rgba(234,88,12,.30)}
  .pri-tag.pri-P2,.meta-chip.pri-P2{color:#ca8a04;background:rgba(202,138,4,.10);border-color:rgba(202,138,4,.30)}
  .pri-tag.pri-P3,.meta-chip.pri-P3{color:#525252;background:var(--surface-2);border-color:var(--line)}
  /* 风险卡 — 左侧黄条 */
  .exec-risk-list{display:flex;flex-direction:column;gap:10px;margin-top:4px}
  .exec-risk-item{padding:12px 16px;background:var(--bg);border:1px solid var(--line);
    border-left:4px solid #ca8a04;border-radius:0 6px 6px 0;
    transition:transform .12s ease,box-shadow .12s ease}
  .exec-risk-item:hover{transform:translateX(2px);box-shadow:0 2px 8px rgba(0,0,0,.06)}
  .exec-risk-head{display:flex;align-items:baseline;gap:10px;margin-bottom:6px;flex-wrap:wrap}
  .exec-risk-title{font-weight:700;color:var(--fg);font-size:14.5px;line-height:1.5;flex:1;min-width:0}
  .exec-risk-line{font-size:13px;color:var(--fg-2);line-height:1.7;margin-top:4px;
    display:flex;align-items:baseline;gap:8px}
  .exec-risk-line .lbl{display:inline-block;min-width:48px;font-size:10.5px;color:var(--fg-3);
    font-weight:700;letter-spacing:.06em;text-transform:uppercase;flex-shrink:0;
    background:var(--surface-2);padding:2px 6px;border-radius:3px}
  /* 阻碍卡 — 红色强调 */
  .exec-blocker-list{display:flex;flex-direction:column;gap:10px;margin-top:4px}
  .exec-blocker-item{padding:14px 16px;background:rgba(220,38,38,.04);
    border:1px solid rgba(220,38,38,.20);border-left:4px solid #dc2626;border-radius:0 6px 6px 0;
    transition:transform .12s ease,box-shadow .12s ease}
  .exec-blocker-item:hover{transform:translateX(2px);box-shadow:0 2px 8px rgba(220,38,38,.10)}
  .exec-blocker-head{display:flex;align-items:center;gap:10px;flex-wrap:wrap;margin-bottom:6px}
  .blocker-tag{font-size:10px;font-weight:700;padding:2px 8px;background:#dc2626;color:#fff;
    letter-spacing:.10em;text-transform:uppercase;border-radius:3px}
  .exec-blocker-title{font-weight:700;color:var(--fg);font-size:14.5px;line-height:1.5;flex:1;min-width:0}
  .exec-blocker-line{font-size:13px;color:var(--fg-2);line-height:1.7;margin-top:4px;
    display:flex;align-items:baseline;gap:8px}
  .exec-blocker-line.fix{color:var(--fg)}
  .exec-blocker-line .lbl{display:inline-block;min-width:66px;font-size:10.5px;color:var(--fg-3);
    font-weight:700;letter-spacing:.06em;text-transform:uppercase;flex-shrink:0;
    background:var(--surface-2);padding:2px 6px;border-radius:3px}
  .exec-blocker-line.fix .lbl{color:#16a34a;background:rgba(22,163,74,.10)}
  /* Bug 卡 — 卡片化 + 左色条按严重度 */
  .exec-issue{border:1px solid var(--line);border-left:4px solid var(--line-2);
    border-radius:0 6px 6px 0;padding:14px 16px;margin-bottom:10px;
    background:var(--bg);transition:transform .12s ease,box-shadow .12s ease}
  .exec-issue:hover{transform:translateX(2px);box-shadow:0 2px 8px rgba(0,0,0,.06)}
  .exec-issue:last-of-type{margin-bottom:0}
  .exec-issue.sev-critical{border-left-color:#dc2626;background:rgba(220,38,38,.03)}
  .exec-issue.sev-high{border-left-color:#ea580c;background:rgba(234,88,12,.03)}
  .exec-issue.sev-medium{border-left-color:#ca8a04}
  .exec-issue.sev-low{border-left-color:#16a34a}
  .exec-issue.sev-info{border-left-color:#0891b2}
  .exec-issue-head{display:flex;align-items:baseline;gap:10px;margin-bottom:6px;flex-wrap:wrap}
  .exec-issue-title{font-weight:700;color:var(--fg);font-size:14.5px;flex:1;min-width:0;line-height:1.5}
  .exec-issue-loc{color:var(--fg-3);font-size:12px;margin:4px 0;font-family:var(--mono);
    padding:3px 8px;background:var(--surface-2);border-radius:4px;display:inline-block}
  .exec-issue-meta{display:flex;flex-wrap:wrap;gap:8px;margin:6px 0 10px;
    font-size:11.5px;color:var(--fg-3);align-items:center}
  .exec-issue-meta .meta-chip{font-size:11px;font-family:var(--mono);font-weight:600;
    padding:2px 8px;background:var(--surface-2);border:1px solid var(--line);
    color:var(--fg-2);border-radius:4px}
  .exec-issue-meta .meta-chip.role{color:var(--fg);background:var(--surface-2);border-color:var(--line-2)}
  .exec-issue-section{margin:8px 0 0;padding:10px 12px;background:var(--surface-2);
    border:1px solid var(--line);border-radius:6px}
  .exec-issue-section.fix{background:rgba(22,163,74,.04);border-color:rgba(22,163,74,.20)}
  .exec-issue-section.verify{background:rgba(8,145,178,.04);border-color:rgba(8,145,178,.20)}
  .exec-issue-section .sec-body{padding-left:0;border-left:none}
  .sec-lbl{font-size:10.5px;font-weight:700;color:var(--fg-3);margin-bottom:6px;
    letter-spacing:.08em;text-transform:uppercase}
  .exec-issue-section.fix .sec-lbl{color:#16a34a}
  .exec-issue-section.verify .sec-lbl{color:#0891b2}
  .sec-body{font-size:13.5px;color:var(--fg);line-height:1.7}
  .repro-list{margin:0;padding-left:20px;font-size:13.5px;line-height:1.7;color:var(--fg)}
  .repro-list li{margin-bottom:3px}
  .accept-line{margin-top:6px;padding:10px 12px;background:rgba(8,145,178,.04);
    border:1px solid rgba(8,145,178,.20);border-radius:6px;font-size:13.5px;color:var(--fg);line-height:1.7}
  .accept-line::before{content:"验收: ";font-weight:700;color:#0891b2;font-size:11.5px;
    letter-spacing:.06em;text-transform:uppercase;margin-right:6px}
  .related-cases{margin-top:8px;font-size:12px;color:var(--fg-3);font-family:var(--mono);
    display:flex;flex-wrap:wrap;gap:6px;align-items:center}
  .related-cases::before{content:"关联用例";font-size:10.5px;font-weight:700;color:var(--fg-3);
    letter-spacing:.06em;text-transform:uppercase;font-family:var(--sans);margin-right:4px}
  .related-cases code{background:var(--surface-2);padding:2px 8px;color:var(--fg-2);
    font-size:11.5px;border:1px solid var(--line);border-radius:3px}
  .exec-issue-impact,.exec-issue-evidence{margin-top:6px;font-size:12.5px;color:var(--fg-2);
    padding:6px 10px;background:var(--surface-2);border:1px solid var(--line);border-radius:4px;line-height:1.6}
  .exec-issue-impact::before{content:"影响范围: ";font-weight:700;color:var(--fg-3);font-size:10.5px;
    letter-spacing:.06em;text-transform:uppercase}
  .exec-issue-evidence::before{content:"证据: ";font-weight:700;color:var(--fg-3);font-size:10.5px;
    letter-spacing:.06em;text-transform:uppercase}
  /* 用例表 — 加边框、斑马纹 */
  .case-table-wrap{margin-top:8px;border:1px solid var(--line);border-radius:8px;overflow-x:auto;
    background:var(--bg)}
  .case-table{width:100%;border-collapse:collapse;font-size:13.5px}
  .case-table th{padding:10px 12px;text-align:left;background:var(--surface-2);
    color:var(--fg);font-weight:700;font-size:11px;text-transform:uppercase;
    letter-spacing:.08em;border-bottom:1px solid var(--line);white-space:nowrap;position:sticky;top:0}
  .case-table td{padding:10px 12px;border-bottom:1px solid var(--line);
    vertical-align:top;color:var(--fg);font-size:13.5px;line-height:1.6}
  .case-table tr:last-child td{border-bottom:none}
  .case-table tr:nth-child(even) td{background:var(--surface-2)}
  .case-table tr.pri-P0 td{background:rgba(220,38,38,.04)}
  .case-table tr.pri-P1 td{background:rgba(234,88,12,.03)}
  .case-table tr:hover td{background:rgba(0,0,0,.02)}
  .case-table .case-idx{font-family:var(--mono);color:var(--fg-3);font-size:12px;width:36px;font-weight:600}
  .case-table .case-id{background:var(--surface-2);padding:2px 8px;color:var(--fg);
    font-size:12px;font-family:var(--mono);border:1px solid var(--line);border-radius:4px;display:inline-block}
  .case-table .case-title{max-width:420px;line-height:1.55;color:var(--fg);font-size:14px;font-weight:500}
  .case-type,.case-auto{font-size:11px;padding:2px 8px;border-radius:4px;font-weight:600;
    background:var(--surface-2);color:var(--fg-2);font-family:var(--sans);white-space:nowrap;
    border:1px solid var(--line);display:inline-block}
  .case-status{font-size:12px;font-family:var(--sans);font-weight:600;white-space:nowrap;
    color:var(--fg-3);padding:2px 8px;background:var(--surface-2);border:1px solid var(--line);
    border-radius:4px;display:inline-block}
  .case-status-ok{color:#16a34a;background:rgba(22,163,74,.10);border-color:rgba(22,163,74,.30)}
  .case-status-bad{color:#dc2626;background:rgba(220,38,38,.10);border-color:rgba(220,38,38,.30)}
  .case-status-muted{color:var(--fg-3)}
  .exec-case-note{margin-top:12px;padding:10px 14px;background:var(--surface-2);
    border:1px solid var(--line);border-left:3px solid var(--ac-line);
    border-radius:0 6px 6px 0;color:var(--fg-2);font-size:13px;line-height:1.65}
  /* === 子步骤原始输出区(默认折叠) === */
  .substep-section-head{display:flex;align-items:baseline;gap:10px;
    margin:0 24px 8px;padding:16px 0 6px 0;border-top:1px solid var(--line)}
  .substep-section-title{font-size:12px;font-weight:600;color:var(--fg-3);letter-spacing:.08em;
    text-transform:uppercase}
  .substep-section-hint{font-size:11px;color:var(--fg-4);font-style:normal;margin-left:auto}
  .report-sub{margin:0 24px}
  .report-sub-head{cursor:pointer;user-select:none}
  .report-sub-head:hover{background:transparent}
  .report-sub-twirl{display:inline-block;width:12px;color:var(--fg-3);font-size:10px;
    transition:transform .15s ease}
  .report-sub.open .report-sub-twirl{transform:rotate(90deg)}
  /* KPI 网格在窄屏自适应 */
  @media (max-width: 720px){
    .exec-kpis{grid-template-columns:repeat(2,1fr);gap:14px}
    .exec-block{padding:0 16px 16px}
    .substep-section-head{margin:0 16px 8px}
    .report-sub{margin:0 16px}
  }
  /* === Screenshot evidence === */
  .report-screenshots{margin:14px 18px 0;padding:14px;background:var(--surface-2);
    border:1px solid var(--line);border-radius:10px}
  .screenshots-head{font-size:13px;font-weight:600;color:var(--fg);margin-bottom:10px}
  .screenshots-hint{font-size:10.5px;font-weight:400;color:var(--fg-3);margin-left:6px;font-style:italic}
  .shot-group{margin-top:10px;padding-top:10px;border-top:1px dashed var(--line)}
  .shot-group:first-of-type{margin-top:0;padding-top:0;border-top:none}
  .shot-url{font-family:"SF Mono",ui-monospace,monospace;font-size:11.5px;color:var(--fg-3);margin-bottom:8px}
  .shot-url code{background:var(--bg);padding:2px 8px;border-radius:4px;color:var(--ac-2)}
  .shot-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(180px,1fr));gap:10px}
  .shot-cell{display:block;background:var(--bg);border:1px solid var(--line);border-radius:8px;
    overflow:hidden;text-decoration:none;color:var(--fg-2);
    transition:transform .15s,border-color .15s,box-shadow .15s}
  .shot-cell:hover{transform:translateY(-2px);border-color:var(--ac-line);
    box-shadow:0 8px 22px rgba(0,0,0,.35)}
  .shot-cell img{display:block;width:100%;height:auto;max-height:240px;object-fit:cover;background:#000}
  .shot-cap{padding:6px 10px;font-size:11px;font-family:"SF Mono",ui-monospace,monospace;
    color:var(--fg-3);border-top:1px solid var(--line);
    display:flex;align-items:center;justify-content:space-between;gap:6px}
  .shot-cap .issue-badge{background:rgba(220,38,38,.14);color:#dc2626;
    padding:1px 7px;border-radius:999px;font-weight:600;font-size:10px}
  .report-hero h4{margin:0;font-size:15px;font-weight:600;letter-spacing:-.005em;color:var(--fg)}
  .report-hero .meta{font-family:var(--mono);font-size:11px;color:var(--fg-3);margin-top:4px}
  .report-hero .meta code{background:var(--bg);padding:1px 6px;border-radius:3px;
    color:var(--ac-2);border:1px solid var(--line)}
  .report-project{margin-top:8px;display:flex;flex-wrap:wrap;gap:14px;
    padding:6px 0 0;border-top:1px dashed var(--line);font-size:12.5px;
    align-items:center}
  .report-project .lbl{color:var(--fg-3);font-size:10.5px;letter-spacing:.04em;
    text-transform:uppercase;font-weight:500;margin-right:2px}
  .report-project code{background:var(--surface-2);padding:2px 7px;border-radius:4px;
    color:var(--fg);font-family:var(--mono);font-size:12px}
  .report-project .val{color:var(--fg);font-weight:500}
  .report-hero .stat-pills{display:flex;gap:6px;flex-wrap:wrap;margin-top:8px}
  .report-hero .stat-pills .pill{font-family:var(--mono);font-size:10.5px;
    padding:2px 8px;border-radius:3px;background:var(--surface-2);border:1px solid var(--line);
    color:var(--fg-2)}
  .report-hero .stat-pills .pill .v{color:var(--ac);font-weight:600}

  .report-sub{border-bottom:1px solid var(--line)}
  .report-sub:last-child{border-bottom:none}
  .report-sub-head{display:flex;align-items:center;gap:10px;padding:10px 18px;
    cursor:pointer;font-size:12.5px;background:var(--surface);
    transition:background .15s}
  .report-sub-head:hover{background:var(--surface-2)}
  .report-sub.open .report-sub-head{background:var(--surface-2)}
  .report-sub-twirl{color:var(--fg-3);font-family:var(--mono);font-size:10px;
    width:10px;transition:transform .15s}
  .report-sub.open .report-sub-twirl{transform:rotate(90deg);color:var(--ac)}
  .report-sub-num{
    font-family:var(--mono);font-size:10.5px;font-weight:600;color:var(--ac);
    background:rgba(16,185,129,.10);border:1px solid rgba(16,185,129,.25);
    width:20px;height:20px;border-radius:50%;display:inline-grid;place-items:center;flex-shrink:0}
  .report-sub-name{color:var(--fg);font-weight:500}
  .report-sub-stats{margin-left:auto;font-family:var(--mono);font-size:11px;color:var(--fg-3);
    display:flex;gap:8px}
  .report-sub-stats .chip{padding:2px 7px;border-radius:3px;background:var(--surface-2);
    border:1px solid var(--line-2)}
  .report-sub-body{display:none;padding:14px 20px;background:var(--bg)}
  .report-sub.open .report-sub-body{display:block}

  /* Smart renderer outputs */
  .report-kv{display:grid;grid-template-columns:max-content 1fr;gap:6px 16px;margin:0;font-size:12.5px}
  .report-kv dt{color:var(--fg-3);font-family:var(--mono);font-size:11.5px;align-self:start;
    padding-top:1px}
  .report-kv dd{margin:0;color:var(--fg);min-width:0;line-height:1.55}
  .report-kv dd.muted{color:var(--fg-3)}

  .report-table{width:100%;border-collapse:collapse;font-size:11.5px;margin:8px 0;
    border:1px solid var(--line);border-radius:6px;overflow:hidden}
  .report-table th{font-family:var(--mono);font-size:10px;text-align:left;
    padding:8px 10px;background:var(--surface-2);color:var(--fg-3);
    text-transform:uppercase;letter-spacing:.07em;font-weight:600;
    border-bottom:1px solid var(--line);white-space:nowrap}
  .report-table td{padding:8px 10px;border-bottom:1px solid var(--line);
    font-family:var(--mono);font-size:11px;color:var(--fg);vertical-align:top;
    word-break:break-word;max-width:320px}
  .report-table tr:last-child td{border-bottom:none}
  .report-table tr:hover td{background:var(--surface-2)}
  .report-table tfoot td{font-family:var(--mono);color:var(--fg-3);font-style:italic;
    text-align:center;font-size:10.5px;background:var(--surface-2)}

  .sev{font-family:var(--mono);font-size:9.5px;font-weight:700;padding:1px 6px;
    border-radius:3px;text-transform:uppercase;letter-spacing:.05em;display:inline-block;
    line-height:1.5}
  .sev-critical{background:rgba(248,113,113,.18);color:var(--bad);
    border:1px solid rgba(248,113,113,.4)}
  .sev-high{background:rgba(248,113,113,.10);color:var(--bad)}
  .sev-major{background:rgba(248,113,113,.10);color:var(--bad)}
  .sev-medium{background:rgba(251,191,36,.10);color:var(--warn)}
  .sev-warn{background:rgba(251,191,36,.10);color:var(--warn)}
  .sev-low{background:rgba(168,174,184,.10);color:var(--fg-2)}
  .sev-minor{background:rgba(168,174,184,.10);color:var(--fg-2)}
  .sev-cosmetic{background:rgba(168,174,184,.10);color:var(--fg-3)}
  .sev-info{background:rgba(16,185,129,.08);color:var(--ac)}

  .confbar{display:inline-flex;align-items:center;gap:6px;font-family:var(--mono);font-size:11px}
  .confbar .track{width:80px;height:5px;border-radius:3px;background:var(--line);overflow:hidden}
  .confbar .fill{height:100%;background:var(--ac)}
  .confbar .pct{color:var(--ac);font-weight:600}

  .empty-array{color:var(--fg-4);font-family:var(--mono);font-size:11px;font-style:italic}
  details.report-detail{margin:4px 0}
  details.report-detail summary{cursor:pointer;font-family:var(--mono);font-size:11px;
    color:var(--ac);user-select:none}
  details.report-detail summary:hover{color:var(--ac-2)}
  details.report-detail[open] > div{margin:6px 0 6px 16px;padding:8px 10px;
    background:var(--surface-2);border-radius:5px;border-left:2px solid var(--line-2)}

  ul.report-list{margin:0;padding-left:18px;font-size:12px;line-height:1.7}
  ul.report-list li{color:var(--fg-2)}
  .err-view{padding:14px 16px;font-family:var(--mono);font-size:11.5px;line-height:1.6;
    color:var(--bad);white-space:pre-wrap;word-break:break-word;background:var(--bg)}

  .status-pill{display:inline-flex;align-items:center;gap:5px;font-family:var(--mono);
    font-size:10px;padding:2px 7px;border-radius:3px;font-weight:600;text-transform:uppercase;
    letter-spacing:.05em}
  .status-pill.queued{background:rgba(168,174,184,.10);color:var(--fg-2)}
  .status-pill.running{background:rgba(251,191,36,.10);color:var(--warn)}
  .status-pill.succeeded{background:rgba(74,222,128,.10);color:var(--ok)}
  .status-pill.failed{background:rgba(248,113,113,.10);color:var(--bad)}
  .status-pill.cancelled{background:rgba(168,174,184,.14);color:var(--fg-2)}
  .status-stop{cursor:pointer;font-family:var(--mono);font-size:10px;font-weight:600;
    padding:2px 8px;border-radius:3px;border:1px solid rgba(220,38,38,.4);
    background:rgba(220,38,38,.07);color:#dc2626;text-transform:uppercase;
    letter-spacing:.04em;vertical-align:middle;margin-left:6px;transition:background .15s}
  .status-stop:hover{background:rgba(220,38,38,.16)}
  .status-stop:disabled{opacity:.5;cursor:default}

  .toast{position:fixed;bottom:24px;left:50%;transform:translateX(-50%);
    background:var(--surface-2);border:1px solid var(--ac);color:var(--fg);
    padding:9px 16px;border-radius:8px;font-size:12.5px;font-family:var(--mono);
    z-index:50;box-shadow:0 12px 30px rgba(0,0,0,.4);opacity:0;transition:opacity .2s}
  .toast.show{opacity:1}

  /* Floating section anchor nav (right rail) */
  .anchor-nav{
    position:fixed;top:50%;right:24px;transform:translateY(-50%);
    display:flex;flex-direction:column;gap:2px;
    background:rgba(20,24,30,.78);
    backdrop-filter:saturate(180%) blur(14px);
    -webkit-backdrop-filter:saturate(180%) blur(14px);
    border:1px solid var(--line-2);border-radius:12px;padding:6px;
    z-index:40;box-shadow:var(--shadow-2);
    opacity:0;transition:opacity .2s}
  .anchor-nav.show{opacity:1}
  .anchor-nav a{
    display:flex;align-items:center;gap:8px;
    padding:8px 12px;border-radius:8px;
    color:var(--fg-3);text-decoration:none;
    font-size:11.5px;font-weight:500;letter-spacing:.02em;
    transition:background .12s,color .12s;
    white-space:nowrap}
  .anchor-nav a:hover{background:var(--surface-2);color:var(--fg)}
  .anchor-nav a.current{background:var(--ac-bg);color:var(--ac-2)}
  .anchor-nav a .num{
    width:20px;height:20px;border-radius:6px;flex-shrink:0;
    display:grid;place-items:center;font-family:var(--mono);font-size:10.5px;font-weight:600;
    background:var(--surface-3);color:var(--fg-3);border:1px solid var(--line-1)}
  .anchor-nav a.current .num{background:var(--ac);color:#001a14;border-color:transparent}
  @media (max-width:1080px){.anchor-nav{display:none}}

  /* ===== 执行摘要组件 — 统一为定稿样式(覆盖旧规则,与下载/导出一致) ===== */
  .exec-block{border:1px solid var(--line);border-radius:14px;padding:20px 22px;margin:0 0 16px;background:#fff;box-shadow:0 1px 3px rgba(16,24,40,.04)}
  .exec-verdict-block{background:linear-gradient(180deg,#fafdfd,#fff)}
  .exec-head{display:flex;align-items:center;gap:11px;margin-bottom:16px}
  .exec-num{font-family:var(--mono);font-size:12px;font-weight:700;color:#fff;background:var(--ac);width:25px;height:25px;border-radius:8px;display:grid;place-items:center;flex-shrink:0;box-shadow:0 2px 6px rgba(13,148,136,.3)}
  .exec-head h3{margin:0;font-size:15px;font-weight:650;letter-spacing:-.01em}
  .exec-count{margin-left:auto;font-family:var(--mono);font-size:12px;font-weight:600;color:var(--fg-3);background:var(--surface);padding:3px 11px;border-radius:20px;border:1px solid var(--line)}
  .exec-count.danger{color:var(--bad);background:rgba(220,38,38,.07);border-color:rgba(220,38,38,.22)}
  .verdict{padding:9px 18px;display:inline-flex;align-items:center;gap:9px;font-size:18px;font-weight:680;border-radius:11px;margin-bottom:14px;border:none}
  .verdict.ok,.verdict.proceed{background:rgba(5,150,105,.09);color:var(--ok)}
  .verdict.warn,.verdict.proceed_with_warning{background:rgba(180,83,9,.1);color:var(--warn)}
  .verdict.bad,.verdict.reject{background:rgba(220,38,38,.09);color:var(--bad)}
  .verdict-icon{font-size:19px}
  .verdict-summary{font-size:14px;color:var(--fg-2);line-height:1.78;margin-bottom:18px}
  .exec-kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:11px;margin-bottom:16px}
  .exec-kpi{flex-direction:column;align-items:center;justify-content:center;background:var(--surface);border:1px solid var(--line);border-radius:11px;padding:15px 12px;text-align:center}
  .exec-kpi-num{display:block;font-size:28px;font-weight:720;color:var(--fg);line-height:1;font-family:var(--mono);letter-spacing:-.02em}
  .exec-kpi-lbl{display:block;font-size:11.5px;color:var(--fg-3);margin-top:7px;font-weight:500}
  .sev-bar{display:flex;height:30px;border-radius:8px;overflow:hidden;border:1px solid var(--line);margin-top:14px;box-shadow:inset 0 1px 2px rgba(0,0,0,.04)}
  .sev-bar-seg{display:grid;place-items:center;color:#fff;font-family:var(--mono);font-size:12px;font-weight:700;min-width:26px}
  .sev-bar-critical{background:#dc2626}.sev-bar-high{background:#f97316}
  .sev-bar-medium{background:#eab308;color:#422006}.sev-bar-low{background:#94a3b8}.sev-bar-info{background:#06b6d4}
  .sev-tag{font-family:var(--mono);font-size:10px;font-weight:700;padding:2px 8px;border-radius:5px;text-transform:uppercase;letter-spacing:.05em;flex-shrink:0}
  .sev-tag.sev-critical{background:#fee2e2;color:#991b1b}
  .sev-tag.sev-high{background:#ffedd5;color:#c2410c}
  .sev-tag.sev-medium{background:#fef9c3;color:#854d0e}
  .sev-tag.sev-low{background:#f1f5f9;color:#475569}
  .sev-tag.sev-info{background:#cffafe;color:#0e7490}
  .exec-risk-list,.exec-blocker-list{display:flex;flex-direction:column;gap:10px}
  .exec-risk-item,.exec-blocker-item{border:1px solid var(--line);border-left:3px solid var(--line-2);border-radius:9px;padding:13px 15px;background:var(--surface)}
  .exec-blocker-item{border-left-color:var(--bad)}
  .exec-risk-head,.exec-blocker-head{display:flex;align-items:center;gap:8px;flex-wrap:wrap;margin-bottom:7px}
  .exec-risk-title,.exec-blocker-title{font-weight:620;font-size:13.5px;color:var(--fg)}
  .exec-risk-line,.exec-blocker-line{font-size:12.5px;color:var(--fg-2);line-height:1.65;margin-top:5px;display:flex;gap:9px}
  .exec-risk-line .lbl,.exec-blocker-line .lbl{color:var(--fg-3);font-size:10.5px;font-weight:600;text-transform:uppercase;letter-spacing:.04em;flex-shrink:0;width:52px;padding-top:1px}
  .exec-blocker-line.fix .lbl{color:var(--ac)}
  .blocker-tag{font-family:var(--mono);font-size:10px;font-weight:700;padding:2px 8px;border-radius:5px;background:rgba(220,38,38,.1);color:var(--bad);text-transform:uppercase;letter-spacing:.05em;flex-shrink:0}
  .meta-chip{font-family:var(--mono);font-size:11px;color:var(--fg-2);background:#fff;border:1px solid var(--line);padding:2px 8px;border-radius:5px;font-weight:500}
  .meta-chip.role{background:var(--surface-2);border-color:transparent}
  .meta-chip.pri-P0{background:#fee2e2;color:#991b1b;border-color:transparent}
  .meta-chip.pri-P1{background:#ffedd5;color:#c2410c;border-color:transparent}
  .meta-chip.pri-P2{background:#fef9c3;color:#854d0e;border-color:transparent}
  .meta-chip.pri-P3{background:#f1f5f9;color:#475569;border-color:transparent}
  .exec-issue{border:1px solid var(--line);border-radius:11px;padding:16px 18px;margin-bottom:12px;background:#fff;border-left:4px solid var(--line-2)}
  .exec-issue.sev-critical{border-left-color:#dc2626}
  .exec-issue.sev-high{border-left-color:#f97316}
  .exec-issue.sev-medium{border-left-color:#eab308}
  .exec-issue.sev-low{border-left-color:#94a3b8}
  .exec-issue-head{display:flex;align-items:center;gap:9px;margin-bottom:10px}
  .exec-issue-title{font-weight:650;font-size:14.5px;color:var(--fg);letter-spacing:-.01em;line-height:1.4}
  .exec-issue-meta{display:flex;flex-wrap:wrap;gap:6px;margin-bottom:11px}
  .exec-issue-loc{font-size:12px;color:var(--fg-3);margin-bottom:6px}
  .exec-issue-loc code{font-size:11px}
  .exec-issue-impact,.exec-issue-evidence{font-size:12.5px;color:var(--fg-2);line-height:1.65;margin-top:5px}
  .exec-issue-section{display:grid;grid-template-columns:68px 1fr;gap:13px;padding:9px 0;border-top:1px dashed var(--line)}
  .exec-issue-section .sec-lbl{font-size:11px;font-weight:600;color:var(--fg-3);text-transform:uppercase;letter-spacing:.04em;padding-top:1px}
  .exec-issue-section.fix .sec-lbl{color:var(--ac)}
  .exec-issue-section.verify .sec-lbl{color:var(--ok)}
  .exec-issue-section .sec-body{font-size:13px;color:var(--fg-2);line-height:1.72;min-width:0}
  .repro-list{margin:0 0 8px;padding-left:18px;font-size:12.5px;line-height:1.7;color:var(--fg-2)}
  .repro-list li{margin:2px 0}
  .accept-line{font-size:12.5px;color:#047857;background:rgba(5,150,105,.06);border-radius:6px;padding:7px 11px;line-height:1.6;margin-top:4px}
  .related-cases{font-size:12px;color:var(--fg-3);margin-top:12px;padding-top:11px;border-top:1px dashed var(--line)}
  .related-cases code{font-size:11px}
  .pri-tag{font-family:var(--mono);font-size:10px;font-weight:700;padding:1px 7px;border-radius:4px}
  .pri-tag.pri-P0{background:#fee2e2;color:#991b1b}.pri-tag.pri-P1{background:#ffedd5;color:#c2410c}
  .pri-tag.pri-P2{background:#fef9c3;color:#854d0e}.pri-tag.pri-P3{background:#f1f5f9;color:#475569}
  .case-type{font-family:var(--mono);font-size:10.5px;color:var(--fg-2);background:var(--surface-2);padding:1px 7px;border-radius:4px}
  .case-auto{font-family:var(--mono);font-size:10.5px;color:var(--ac-2)}
  .case-title{font-family:var(--sans);font-size:12.5px;color:var(--fg);font-weight:500}
</style></head>
<body>
<div class="topbar">
  <a class="brand-link" href="/tools" title="天枢 · 裁决 · 返回主页" style="text-decoration:none;display:inline-flex;align-items:center;gap:10px;font-family:'Noto Serif SC',Georgia,serif;font-size:19px;font-weight:600;letter-spacing:.18em;color:var(--fg);margin-right:24px;padding:4px 0"><svg viewBox="0 0 24 24" fill="currentColor" width="24" height="24" aria-hidden="true" style="color:var(--ac);opacity:1;flex-shrink:0;filter:drop-shadow(0 1px 2px rgba(196,90,58,.28))"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span>天枢</span><span style="color:var(--ac);margin:0 6px;font-weight:400">·</span><span>裁决</span></a>
  <a class="back-btn" href="/tools" title="返回工具列表">
    <span style="font-size:14px;line-height:1">←</span>
    <span>返回</span>
  </a>
  <div class="crumbs">
    <a href="/tools">工具</a><span class="sep">/</span>
    <span class="current"><span id="tb-icon">·</span><span id="tb-name">…</span></span>
    <span class="sep" id="tb-sep" style="display:none">·</span>
    <span id="tb-status"></span>
  </div>
  <div class="stats" id="tb-stats" style="display:none">
    <span class="stat"><span class="lbl">入</span><span class="v" id="s-in">—</span></span>
    <span class="stat"><span class="lbl">出</span><span class="v" id="s-out">—</span></span>
    <span class="stat"><span class="lbl">缓存</span><span class="v" id="s-cache">—</span></span>
    <span class="stat cost"><span class="lbl">$</span><span class="v" id="s-cost">—</span></span>
    <span class="stat" id="s-elapsed-wrap" style="display:none"><span class="v" id="s-elapsed"></span></span>
  </div>
  <a href="/settings" class="set superadmin-only" style="display:none">设置</a>
  <button class="run" id="btn-run">▶ 运行<kbd>⌘↵</kbd></button>
</div>

<div class="hero">
  <div class="hero-inner">
    <div class="hero-icon-frame"><span class="ic" id="hero-icon">·</span></div>
    <div class="hero-body">
      <h2 id="hero-name">…</h2>
      <p class="tag" id="hero-tag">·</p>
      <div class="hero-meta">
        <span class="step" id="hero-step">·</span>
        <span class="resp" id="hero-resp">·</span>
        <span id="hero-prompts">·</span>
      </div>
    </div>
    <div class="hero-right" id="hero-claude-status"></div>
  </div>
</div>

<nav class="anchor-nav" id="anchor-nav">
  <a href="#sec-0" data-target="sec-0"><span class="num">0</span>项目</a>
  <a href="#sec-1" data-target="sec-1"><span class="num">1</span>材料</a>
  <a href="#sec-2" data-target="sec-2"><span class="num">2</span>步骤</a>
  <a href="#sec-3" data-target="sec-3"><span class="num">3</span>配置</a>
</nav>

<main>
<div class="runner-grid idle">
<div class="workspace">
  <!-- Section 0: 项目信息 (必填) -->
  <div class="sec" id="sec-0">
    <div class="sec-head">
      <span class="num">0</span><h3>项目信息</h3>
      <span style="margin-left:auto;font-size:11px;color:var(--fg-4)">必填 · 会写入报告</span>
    </div>
    <div class="sec-body" style="display:grid;grid-template-columns:240px 1fr;gap:14px">
      <div>
        <label for="proj-code" style="display:block;font-size:11px;color:var(--fg-3);margin-bottom:5px;letter-spacing:.04em;text-transform:uppercase">项目编号 *</label>
        <input id="proj-code" type="text" maxlength="64" placeholder="例:PROJ-2026-Q1-001"
          style="width:100%;background:var(--surface-2);border:1px solid var(--line);border-radius:6px;
                 color:var(--fg);font-family:var(--mono);font-size:12.5px;padding:9px 11px;outline:none">
      </div>
      <div>
        <label for="proj-name" style="display:block;font-size:11px;color:var(--fg-3);margin-bottom:5px;letter-spacing:.04em;text-transform:uppercase">项目名称 *</label>
        <input id="proj-name" type="text" maxlength="128" placeholder="例:Miku 18 红包活动 H5"
          style="width:100%;background:var(--surface-2);border:1px solid var(--line);border-radius:6px;
                 color:var(--fg);font-family:var(--sans);font-size:13px;padding:9px 11px;outline:none">
      </div>
    </div>
  </div>

  <!-- Section 1: 输入 -->
  <div class="sec" id="sec-1">
    <div class="sec-head">
      <span class="num">1</span><h3 id="input-label">需求材料</h3>
      <button type="button" id="hint-toggle" class="hint-toggle">填写指引 ▾</button>
    </div>
    <details id="hint-panel" class="hint-panel">
      <summary style="display:none"></summary>
      <div id="input-hint" class="hint-body">粘贴文本、上传文档或拖入 — 这是整个分析的源头</div>
    </details>
    <div class="sec-body">
      <div class="input-zone" id="input-zone">
        <div class="input-toolbar">
          <button type="button" id="upload-btn">↑ 上传文件</button>
          <input type="file" id="file-input" multiple style="display:none"
            accept=".md,.txt,.json,.yaml,.yml,.csv,.tsv,.html,.xml,.pdf,.png,.jpg,.jpeg,.gif,.webp,.docx,.xlsx,.pptx,.py,.js,.ts,.tsx,.go,.rs,.java,.cpp,.c,.swift,.kt,.rb,.sh,.sql,.log">
          <button type="button" id="import-btn" title="把上游工具(step1/2/4 等)的报告作为输入">⇡ 从上游报告导入</button>
          <span class="drag-hint">支持拖拽文件到此处</span>
          <span class="size-hint" id="size-hint" style="margin-left:auto">0 B</span>
        </div>
        <!-- step5 专属:Figma 设计稿 链接 + 读取 + 登录切换 -->
        <div id="figma-bar" style="display:none;margin:8px 0 0;padding:10px 12px;
          background:var(--surface-2);border:1px solid var(--line);border-radius:6px">
          <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap">
            <span style="font-size:12px;color:var(--fg-2);font-weight:600;white-space:nowrap">Figma 设计稿</span>
            <input type="text" id="figma-url" placeholder="粘贴 Figma 链接(.../design/...?node-id=...)"
              style="flex:1;min-width:240px;padding:6px 10px;border:1px solid var(--line-2);border-radius:6px;
                font-family:var(--mono);font-size:12px;background:var(--surface);color:var(--fg)">
            <button type="button" id="figma-read-btn" style="background:var(--ac);color:#fff;border:none;
              padding:6px 14px;border-radius:6px;font-size:12px;cursor:pointer;white-space:nowrap">读取预览</button>
          </div>
          <div style="display:flex;align-items:center;gap:10px;margin-top:8px">
            <span id="figma-status" style="flex:1;font-size:11.5px;color:var(--fg-3);font-family:var(--mono)">设计稿读取走服务端只读 Token</span>
          </div>
          <div id="figma-preview-gallery" style="display:none;margin-top:8px;gap:8px;flex-wrap:wrap;max-height:360px;overflow-y:auto;padding:4px;background:var(--bg-2);border:1px solid var(--line);border-radius:6px"></div>
        </div>
        <div id="import-panel" style="display:none;margin:8px 0 0;padding:10px 12px;
          background:var(--surface-2);border:1px solid var(--line);border-radius:6px">
          <div style="display:flex;align-items:center;gap:10px;margin-bottom:8px">
            <span style="font-size:12px;color:var(--fg-2);font-weight:500">选择要作为输入的上游报告</span>
            <button type="button" id="import-close" style="margin-left:auto;background:transparent;border:none;color:var(--fg-3);cursor:pointer;font-size:12px">关闭</button>
          </div>
          <div id="import-list" style="max-height:240px;overflow-y:auto;
            display:flex;flex-direction:column;gap:4px"></div>
        </div>
        <div id="upload-progress" style="display:none;margin:8px 0;padding:10px 14px;
          background:var(--surface-2);border:1px solid var(--line);border-radius:6px">
          <div style="display:flex;align-items:center;gap:8px">
            <span style="display:inline-block;width:13px;height:13px;border:2px solid var(--ac);
              border-top-color:transparent;border-radius:50%;animation:aitk-spin .7s linear infinite"></span>
            <span id="upload-progress-label" style="font-size:12.5px;color:var(--fg-2)">上传中…</span>
            <span id="upload-progress-pct" style="margin-left:auto;font-family:var(--mono);font-size:12px;color:var(--ac)"></span>
          </div>
          <div style="height:6px;background:var(--surface-3,#e5e5e5);border-radius:3px;margin-top:8px;overflow:hidden">
            <div id="upload-progress-bar" style="height:100%;width:0%;background:var(--ac);transition:width .15s"></div>
          </div>
        </div>
        <textarea id="doc-input" placeholder="粘贴 PRD / 接口定义 / 页面信息&#10;&#10;或拖一个 .md / .pdf / .docx / 截图 / .apk 进来"></textarea>
        <div class="drop-overlay">
          <div class="drop-overlay-inner">
            <div class="drop-icon">⤓</div>
            <div class="drop-text">释放鼠标 — 文件将自动加载</div>
            <div class="drop-sub" id="drop-sub">.md .txt .json .yaml .csv .pdf .png .docx …</div>
          </div>
        </div>
      </div>
      <div class="run-options" id="run-options"></div>
    </div>
  </div>

  <!-- Section 2: 提示词 -->
  <div class="sec" id="sec-2">
    <div class="sec-head">
      <span class="num">2</span><h3>执行步骤</h3>
      <span class="sub" id="prompts-sub">展开查看 prompt · 修改自动保存为覆盖</span>
    </div>
    <div id="prompts-toolbar" style="display:none;padding:9px 18px;border-bottom:1px solid var(--line);background:var(--surface-2);font-size:11.5px;color:var(--fg-3);font-family:var(--mono);display:flex;align-items:center;gap:10px">
      <span>勾选要执行的子步骤：</span>
      <button id="prompt-all" style="background:transparent;border:1px solid var(--line-2);color:var(--fg-2);padding:3px 9px;border-radius:5px;font-family:var(--mono);font-size:11px;cursor:pointer">全选</button>
      <button id="prompt-none" style="background:transparent;border:1px solid var(--line-2);color:var(--fg-2);padding:3px 9px;border-radius:5px;font-family:var(--mono);font-size:11px;cursor:pointer">全不选</button>
      <span style="margin-left:auto" id="prompt-count">0/0</span>
    </div>
    <div id="prompts-list"></div>
  </div>

  <!-- Section 3: 模型 & 精度 -->
  <div class="sec" id="sec-3">
    <div class="sec-head">
      <span class="num">3</span><h3>运行配置</h3>
      <span class="sub">所有子步骤共用一份模型与精度 · 自动记忆</span>
    </div>
    <div class="sec-body">
      <div class="model-grid">
        <div class="model-cell" id="cell-model">
          <div class="label">模型 <span class="from" id="model-from">来源：本地</span></div>
          <select id="sel-model"></select>
          <div class="hint" id="model-hint">所有子步骤统一使用</div>
          <div class="probe-status" id="model-probe-status" style="margin-top:6px;font-family:var(--mono);font-size:11px;letter-spacing:.02em;display:none"></div>
        </div>
        <div class="model-cell" id="cell-effort">
          <div class="label">精度 <span class="from" id="effort-from">默认</span></div>
          <select id="sel-effort"></select>
          <div class="hint">越高越细 · max 最强最慢</div>
        </div>
        <div class="model-cell" id="cell-thinking">
          <div class="label">扩展思考</div>
          <select id="sel-thinking"></select>
          <div class="hint">复杂推理选 adaptive</div>
        </div>
        <div class="model-cell unsupported" id="cell-unsupported" style="display:none;grid-column:span 2">
          <div class="label">当前模型不支持精度 / 思考调节</div>
          <div class="hint" id="unsupported-hint">Haiku 系列定位高速低成本。如需调精度请切到 Sonnet 或 Opus。</div>
        </div>
      </div>
      <div class="claude-status" id="claude-status">检测中…</div>
    </div>
  </div>
</div><!-- /workspace -->

<aside class="run-panel" aria-label="运行状态">
  <div class="run-panel-head">
    <span class="live-dot" id="run-status-dot"></span>
    <span class="label-text" id="run-status-label">就绪</span>
    <span class="meta-tail" id="run-status-meta"></span>
  </div>
  <div class="run-panel-body">
    <div id="run-area">
      <div class="run-empty-pretty">
        <div class="play-circle" id="play-circle" title="运行" style="cursor:pointer">▶</div>
        <h4 class="title">准备开始运行</h4>
        <p class="desc">填写左侧三块（材料 · 步骤 · 配置）后启动 — 这里会实时显示进度与报告。</p>
        <div class="hint-row"><kbd>⌘</kbd><kbd>↵</kbd>&nbsp;快速运行</div>
      </div>
    </div>
  </div>
</aside>
</div><!-- /runner-grid -->
</main>

<div class="toast" id="toast"></div>

<script>
const TOOL_ID = location.pathname.split('/').filter(Boolean).pop();
let tool = null;
let claudeInfo = null;
let pollTimer = null, elapsedTimer = null;
let currentRunId = new URLSearchParams(location.search).get('run');
let currentRun = null;

async function load(){
  const [t, ci, cat] = await Promise.all([
    fetch('/api/tools/' + TOOL_ID).then(r=>r.json()),
    fetch('/api/claude/info').then(r=>r.json()),
    fetch('/api/tools').then(r=>r.json()).catch(()=>({tools:[]})),
  ]);
  tool = t; claudeInfo = ci;
  document.title = tool.name + ' — 天枢·裁决';
  document.getElementById('tb-icon').textContent = tool.icon;
  document.getElementById('tb-name').textContent = tool.name;
  document.getElementById('hero-icon').textContent = tool.icon;
  // 顶部 meta:负责方 · 子步骤数(不再显示"第 X 章")
  document.getElementById('hero-step').textContent = (tool.responsible || '');
  document.getElementById('hero-step').title = (tool.step || tool.id).toUpperCase() + ' · 内部步骤标识';
  document.getElementById('hero-resp').textContent = '· ' + tool.prompts.length + ' 子步骤';
  document.getElementById('hero-prompts').textContent = '';
  document.getElementById('hero-name').textContent = tool.name;
  document.getElementById('hero-tag').textContent = tool.description;
  // step5 / h5_adapt:显示 Figma 设计稿登录条
  if (tool.id === 'step5' || tool.id === 'h5_adapt') { try { initFigmaBar(); } catch(e){} }
  // 取消 hero-pills(已合并到 meta 一行)
  // hero 右侧:接入状态 + 输出物
  const hr = document.getElementById('hero-claude-status');
  if (hr) {
    const auth = (claudeInfo && claudeInfo.auth_state) || '';
    const acc = (claudeInfo && claudeInfo.account) || {};
    let row1 = '';
    if (auth === 'toolkit_api_key') {
      row1 = `<span class="row ok"><span class="dot"></span>API Key · ${acc.display_name || ''}</span>`;
    } else if (auth === 'toolkit_oauth') {
      row1 = `<span class="row ok"><span class="dot"></span>OAuth${acc.email ? ' · ' + acc.email : ''}</span>`;
    } else if (auth === 'external_claude_cli' && acc.email) {
      row1 = `<span class="row warn"><span class="dot"></span>本机 CLI · ${acc.email}</span>`;
    } else {
      row1 = `<span class="row warn"><span class="dot"></span>未接入 · <a href="/settings">去设置</a></span>`;
    }
    const out = (tool.output || '').replace(/[《》]/g, '');
    hr.innerHTML = row1 + `<span class="row" style="font-size:10.5px">输出 · ${out || '—'}</span>`;
  }

  // Section 1 — single field label + hint + run options
  const inp = tool.input || {};
  // 项目编号 + 名称:从 localStorage 预填,聚焦时清除红边
  try {
    const last = JSON.parse(localStorage.getItem('toolkit.last_project') || '{}');
    if (last.code) document.getElementById('proj-code').value = last.code;
    if (last.name) document.getElementById('proj-name').value = last.name;
  } catch(_){}
  ['proj-code','proj-name'].forEach(id => {
    const el = document.getElementById(id);
    if (el) el.addEventListener('focus', () => { el.style.borderColor = ''; });
  });

  document.getElementById('input-label').textContent = inp.label || '需求材料';
  const hintEl = document.getElementById('input-hint');
  if (hintEl) hintEl.textContent = inp.hint || '粘贴文本、上传文档或拖入 — 这是整个分析的源头';
  // 折叠式 hint 切换
  const hintBtn = document.getElementById('hint-toggle');
  const hintPanel = document.getElementById('hint-panel');
  if (hintBtn && hintPanel) {
    hintBtn.onclick = () => {
      const open = hintPanel.classList.toggle('open');
      hintBtn.classList.toggle('open', open);
      hintBtn.textContent = open ? '填写指引 ▴' : '填写指引 ▾';
    };
  }
  document.getElementById('input-hint').textContent = inp.hint || '粘贴文本、上传文档或拖入 — 这是整个分析的源头';
  buildRunOptions();
  wireInputToolbar();

  await buildPrompts();
  buildModelControls();
  if (currentRunId) startPolling(currentRunId);
  // 全局运行锁:不管谁打开、是否刷新,都按后端真实状态显示"运行中" / "▶ 运行"
  startActiveStatusPolling();
}

function buildRunOptions(){
  const root = document.getElementById('run-options');
  root.innerHTML = '';
  const opts = tool.run_options || [];
  opts.forEach(o => {
    const lbl = document.createElement('label');
    lbl.innerHTML = `<input type="checkbox" id="ro-${o.key}" ${o.default ? 'checked' : ''}> ${o.label}`;
    root.appendChild(lbl);
  });
}

function wireInputToolbar(){
  const ta = document.getElementById('doc-input');
  const sizeEl = document.getElementById('size-hint');
  const updateSize = () => {
    const len = new Blob([ta.value]).size;
    sizeEl.textContent = len < 1024 ? `${len} B` :
                         len < 1024*1024 ? `${(len/1024).toFixed(1)} KB` :
                         `${(len/1024/1024).toFixed(2)} MB`;
  };
  ta.addEventListener('input', updateSize);
  updateSize();

  // Files which are safe to read as UTF-8 in the browser
  const TEXT_EXTS = new Set([
    'md','txt','json','yaml','yml','csv','tsv','html','htm','xml','svg',
    'py','js','ts','tsx','jsx','vue','go','rs','java','cpp','c','h','swift','kt','rb','sh','bash',
    'sql','log','env','toml','ini','conf','properties','dockerfile','makefile',
    'patch','diff'
  ]);
  // Files which need server-side parsing (binary or encoding-tricky)
  const BINARY_EXTS = new Set(['pdf','docx','xlsx','xlsm','png','jpg','jpeg','gif','webp','bmp','doc','ppt','pptx']);

  function fileExt(name){
    const i = name.lastIndexOf('.');
    return i >= 0 ? name.slice(i+1).toLowerCase() : '';
  }

  // 用 XHR 上传(fetch 拿不到上传进度)— 大文件(APK 等)显示真实上传 %
  function uploadViaServer(f, onProgress){
    return new Promise((resolve, reject) => {
      const fd = new FormData(); fd.append('file', f, f.name);
      const xhr = new XMLHttpRequest();
      xhr.open('POST', '/api/extract-file');
      xhr.upload.onprogress = (e) => { if (e.lengthComputable && onProgress) onProgress(e.loaded, e.total); };
      xhr.onload = () => {
        if (xhr.status >= 200 && xhr.status < 300){
          try { resolve(JSON.parse(xhr.responseText)); }
          catch(err){ reject(new Error('返回解析失败')); }
        } else {
          let msg = 'HTTP ' + xhr.status;
          try { const d = JSON.parse(xhr.responseText); if (d.detail) msg = d.detail; } catch(_){}
          reject(new Error(msg));
        }
      };
      xhr.onerror = () => reject(new Error('网络错误'));
      xhr.send(fd);
    });
  }

  async function uploadFileGetMarker(f, onProgress){
    // 所有文件都上传到服务端存盘,拿回短标记(含 file_ref / app_ref)。
    // 不再在前端读全文塞输入框 —— 真实文件运行时作为内容块直传 LLM。
    const data = await uploadViaServer(f, onProgress);
    return data.text;
  }

  // Helper: append a list of files into the textarea with separators
  async function loadFiles(files){
    if (!files || !files.length) return;
    const parts = [];
    let skipped = 0;
    let failed = 0;
    const prog = document.getElementById('upload-progress');
    const bar = document.getElementById('upload-progress-bar');
    const lbl = document.getElementById('upload-progress-label');
    const pctEl = document.getElementById('upload-progress-pct');
    const mb = (n) => (n/1024/1024).toFixed(1);
    for (const f of files){
      // Skip directories (dragged folders show up as 0-byte type-empty entries)
      if (f.size === 0 && !f.type){
        skipped++;
        continue;
      }
      // 大文件(>1MB,典型如 APK)显示上传进度;远程上传尤其需要
      const showProg = prog && f.size > 1024 * 1024;
      if (showProg){
        prog.style.display = 'block';
        bar.style.width = '0%'; pctEl.textContent = '0%';
        lbl.textContent = `上传 ${f.name}（${mb(f.size)}MB）…`;
      }
      let text;
      try {
        text = await uploadFileGetMarker(f, (loaded, total) => {
          if (!showProg) return;
          const pct = Math.round(loaded / total * 100);
          bar.style.width = pct + '%'; pctEl.textContent = pct + '%';
          if (pct >= 100){ lbl.textContent = `服务端处理中…（${f.name}）`; pctEl.textContent = '处理中'; }
        });
      } catch(err){
        text = `(上传失败：${f.name} — ${err.message || err})`;
        failed++;
      }
      if (showProg){ prog.style.display = 'none'; }
      // 只插短标记(📎 文件名 | file_ref=…),真实文件运行时直传 LLM
      parts.push(text);
    }
    if (parts.length){
      const sep = ta.value.trim() ? '\n' : '';
      ta.value += sep + parts.join('\n');
      updateSize();
      const bits = [`已附加 ${parts.length} 个文件（运行时直传给 AI）`];
      if (skipped) bits.push(`跳过 ${skipped} 个空/目录`);
      if (failed) bits.push(`${failed} 个上传失败`);
      toast(bits.join(' · '));
    } else if (skipped > 0){
      toast(`跳过 ${skipped} 个目录或空文件，未加载任何内容`);
    }
  }

  // Click to upload
  document.getElementById('upload-btn').onclick = () => {
    document.getElementById('file-input').click();
  };
  document.getElementById('file-input').onchange = async (e) => {
    await loadFiles(Array.from(e.target.files || []));
    e.target.value = '';
  };

  // 从上游报告导入 — 列出最近 8 个其他工具的报告,选一个把核心内容塞 textarea
  const importBtn = document.getElementById('import-btn');
  const importPanel = document.getElementById('import-panel');
  const importList = document.getElementById('import-list');
  document.getElementById('import-close').onclick = () => { importPanel.style.display='none'; };
  importBtn.onclick = async () => {
    importPanel.style.display = 'block';
    importList.innerHTML = '<div style="color:var(--fg-3);font-size:11.5px;padding:8px">加载中…</div>';
    try {
      const [reps, catRaw] = await Promise.all([
        fetch('/api/reports').then(r => r.json()),
        fetch('/api/tools').then(r => r.json()),
      ]);
      const toolNameMap = {};
      (catRaw.tools || []).forEach(t => { toolNameMap[t.id] = (t.icon || '') + ' ' + t.name; });
      // 合并 in-memory + saved,排除当前工具的 run(导入上游是另一个工具)
      const all = [];
      (reps.in_memory || []).forEach(r => { if (r.tool_id !== tool.id) all.push(r); });
      (reps.saved || []).forEach(r => {
        if (r.tool_id !== tool.id && !all.find(x => x.run_id === r.run_id)) all.push(r);
      });
      all.sort((a, b) => (b.mtime || b.started_at || 0) - (a.mtime || a.started_at || 0));
      const top = all.slice(0, 12);
      if (!top.length) {
        importList.innerHTML = '<div style="color:var(--fg-3);font-size:11.5px;padding:8px">还没有其他工具的报告，先跑一个上游工具(比如 step1 需求评审)</div>';
        return;
      }
      importList.innerHTML = top.map(r => {
        const label = toolNameMap[r.tool_id] || r.tool_id;
        const t = r.mtime ? new Date(r.mtime * 1000) : new Date((r.started_at || 0) * 1000);
        const ts = t.toLocaleString('zh-CN', {month:'2-digit',day:'2-digit',hour:'2-digit',minute:'2-digit'});
        return `<button type="button" class="import-row" data-rid="${r.run_id}" data-tid="${r.tool_id}"
          style="display:flex;align-items:center;gap:8px;padding:6px 10px;border:1px solid var(--line);
                 background:transparent;border-radius:4px;cursor:pointer;text-align:left;
                 font-size:12px;color:var(--fg-2)">
          <span style="font-weight:500;color:var(--fg)">${label}</span>
          <span style="font-family:var(--mono);font-size:10.5px;color:var(--fg-3)">${r.run_id.slice(0,8)}</span>
          <span style="margin-left:auto;font-size:10.5px;color:var(--fg-3)">${ts}</span>
        </button>`;
      }).join('');
      importList.querySelectorAll('.import-row').forEach(b => {
        b.onclick = async () => {
          const rid = b.dataset.rid;
          const tid = b.dataset.tid;
          b.disabled = true;
          b.textContent = '加载中…';
          try {
            const rep = await fetch('/api/reports/' + rid).then(r => r.json());
            const r = rep.report || {};
            const ta = document.getElementById('doc-input');
            // 把上游报告的核心契约字段格式化成 markdown 注入,LLM 比单纯 JSON 友好
            const lines = [];
            lines.push('# 上游报告 — ' + (toolNameMap[tid] || tid));
            lines.push('Run ID: `' + rid + '`');
            lines.push('');
            if (r.verdict) lines.push('## 测试结论\n' + r.verdict + (r.verdict_summary ? '\n\n' + r.verdict_summary : ''));
            if (Array.isArray(r.risks) && r.risks.length) {
              lines.push('\n## 风险');
              r.risks.forEach(x => {
                const t = typeof x === 'string' ? x : (x.title || '');
                const imp = typeof x === 'object' ? (x.impact || '') : '';
                lines.push('- ' + t + (imp ? ' — ' + imp : ''));
              });
            }
            if (Array.isArray(r.blockers) && r.blockers.length) {
              lines.push('\n## 阻碍');
              r.blockers.forEach(x => {
                lines.push('- **' + (x.title || '') + '** (' + (x.owner_role || '') + ')');
                if (x.why_blocking) lines.push('  - 为何阻碍:' + x.why_blocking);
                if (x.what_to_unblock) lines.push('  - 如何解开:' + x.what_to_unblock);
              });
            }
            if (Array.isArray(r.issues) && r.issues.length) {
              lines.push('\n## 已知问题 (top ' + Math.min(15, r.issues.length) + ')');
              r.issues.slice(0, 15).forEach(it => {
                lines.push('- [' + (it.severity || '?') + '/' + (it.priority || '?') + '] ' + (it.title || ''));
                if (it.module) lines.push('  - 位置:' + it.module);
              });
            }
            if (Array.isArray(r.cases) && r.cases.length) {
              lines.push('\n## 用例集 (top ' + Math.min(20, r.cases.length) + ')');
              r.cases.slice(0, 20).forEach(c => {
                lines.push('- [' + (c.priority || '?') + '] `' + (c.id || '') + '` ' + (c.title || ''));
              });
            }
            const text = lines.join('\n');
            // 追加到 textarea 末尾,不覆盖现有输入
            if (ta.value && ta.value.trim()) {
              ta.value += '\n\n---\n\n' + text;
            } else {
              ta.value = text;
            }
            ta.dispatchEvent(new Event('input', {bubbles:true}));
            importPanel.style.display = 'none';
            // 提示
            const hint = document.getElementById('size-hint');
            if (hint) hint.textContent = '已导入上游报告';
          } catch (e) {
            b.textContent = '加载失败';
          }
        };
      });
    } catch (e) {
      importList.innerHTML = '<div style="color:var(--bad);font-size:11.5px;padding:8px">加载失败:' + (e.message || e) + '</div>';
    }
  };

  // Drag & drop — entire input-zone receives files
  const zone = document.getElementById('input-zone');
  let dragDepth = 0;
  // Stop the browser from "navigating away" if a file is dropped outside the zone
  ['dragenter','dragover','drop'].forEach(evt => {
    window.addEventListener(evt, e => {
      // Only intercept if files are being dragged
      if (e.dataTransfer && [...(e.dataTransfer.types||[])].includes('Files')){
        e.preventDefault();
      }
    });
  });
  zone.addEventListener('dragenter', e => {
    if (e.dataTransfer && [...(e.dataTransfer.types||[])].includes('Files')){
      e.preventDefault();
      dragDepth++;
      zone.classList.add('dragging');
    }
  });
  zone.addEventListener('dragover', e => {
    if (e.dataTransfer && [...(e.dataTransfer.types||[])].includes('Files')){
      e.preventDefault();
      e.dataTransfer.dropEffect = 'copy';
    }
  });
  zone.addEventListener('dragleave', e => {
    e.preventDefault();
    dragDepth = Math.max(0, dragDepth - 1);
    if (dragDepth === 0) zone.classList.remove('dragging');
  });
  zone.addEventListener('drop', async e => {
    e.preventDefault();
    dragDepth = 0;
    zone.classList.remove('dragging');
    const files = Array.from(e.dataTransfer?.files || []);
    await loadFiles(files);
  });
}

// === Section 2: prompts (collapsible + editable + optional checkbox) ===
async function buildPrompts(){
  const root = document.getElementById('prompts-list');
  const toolbar = document.getElementById('prompts-toolbar');
  root.innerHTML = '';
  const optional = !!tool.substeps_optional;
  toolbar.style.display = optional ? 'flex' : 'none';

  // Side effect: populate a substep id → name map used by the HTML report renderer.
  tool._substepNames = tool._substepNames || {};

  for (const sub_id of tool.prompts){
    const data = await fetch(`/api/prompts/${tool.prompt_dir}/${sub_id}`).then(r=>r.json());
    tool._substepNames[sub_id] = data.name;
    const row = document.createElement('div');
    row.className = 'prompt-row';
    row.dataset.subId = sub_id;
    const checkboxHtml = optional
      ? `<input type="checkbox" class="check" data-sub="${sub_id}" checked title="勾选 = 执行该子步骤；取消 = 跳过">`
      : '';
    const shortId = String(data.id).split('.').pop();
    row.innerHTML = `
      <div class="prompt-head">
        ${checkboxHtml}
        <span class="twirl">▶</span>
        <span class="id" title="${data.id}">#${shortId}</span>
        <span class="name">${data.name}</span>
        <span class="chip" style="margin-left:auto" title="模板字段数 · 输出格式">${data.placeholders.length} 字段 · ${data.output_format}</span>
        <span class="chip override" style="margin-left:6px;display:${data.is_override ? '' : 'none'}">已覆盖</span>
        <span class="chip dirty" style="margin-left:6px;display:none">未保存</span>
      </div>
      <div class="prompt-body">
        <textarea>${escapeHtml(data.body)}</textarea>
        <div class="btn-row">
          <button class="save">保存覆盖</button>
          <button class="reset" style="display:${data.is_override ? '' : 'none'}">重置回原版</button>
          <span class="info">占位符 ${data.placeholders.length} · 输出 ${data.output_format}</span>
        </div>
      </div>`;
    const head = row.querySelector('.prompt-head');
    const cb = row.querySelector('.check');
    head.onclick = (e) => {
      // Don't toggle expand when clicking the checkbox
      if (e.target === cb) return;
      row.classList.toggle('open');
    };
    if (cb){
      cb.onclick = (e) => e.stopPropagation();
      cb.onchange = () => {
        row.classList.toggle('disabled', !cb.checked);
        updatePromptCount();
      };
    }
    const ta = row.querySelector('textarea');
    const saveBtn = row.querySelector('.save');
    const resetBtn = row.querySelector('.reset');
    const overrideChip = row.querySelectorAll('.chip.override')[0];
    const dirtyChip = row.querySelectorAll('.chip.dirty')[0];
    let original = data.body;
    ta.oninput = () => {
      dirtyChip.style.display = (ta.value !== original) ? '' : 'none';
    };
    saveBtn.onclick = async (ev) => {
      ev.stopPropagation();
      saveBtn.disabled = true; saveBtn.textContent = '保存中…';
      try {
        await fetch(`/api/prompts/${tool.prompt_dir}/${sub_id}`, {
          method:'PUT', headers:{'Content-Type':'application/json'},
          body: JSON.stringify({body: ta.value}),
        });
        original = ta.value;
        dirtyChip.style.display = 'none';
        overrideChip.style.display = '';
        resetBtn.style.display = '';
        toast(`${sub_id} 已保存覆盖`);
      } finally {
        saveBtn.disabled = false; saveBtn.textContent = '保存覆盖';
      }
    };
    resetBtn.onclick = async (ev) => {
      ev.stopPropagation();
      if (!confirm(`确认重置 ${sub_id} 到原版？当前编辑会丢失。`)) return;
      await fetch(`/api/prompts/${tool.prompt_dir}/${sub_id}`, {method:'DELETE'});
      // Reload original body
      const r2 = await fetch(`/api/prompts/${tool.prompt_dir}/${sub_id}`).then(r=>r.json());
      ta.value = r2.body;
      original = r2.body;
      dirtyChip.style.display = 'none';
      overrideChip.style.display = 'none';
      resetBtn.style.display = 'none';
      toast(`${sub_id} 已重置`);
    };
    root.appendChild(row);
  }
  // Wire toolbar select-all / select-none
  if (tool.substeps_optional){
    document.getElementById('prompt-all').onclick = () => {
      root.querySelectorAll('.check').forEach(cb => {
        if (!cb.checked){ cb.checked = true; cb.dispatchEvent(new Event('change')); }
      });
    };
    document.getElementById('prompt-none').onclick = () => {
      root.querySelectorAll('.check').forEach(cb => {
        if (cb.checked){ cb.checked = false; cb.dispatchEvent(new Event('change')); }
      });
    };
    updatePromptCount();
  }
}

function updatePromptCount(){
  const all = document.querySelectorAll('#prompts-list .check');
  const checked = document.querySelectorAll('#prompts-list .check:checked');
  const el = document.getElementById('prompt-count');
  if (el) el.textContent = `${checked.length}/${all.length}`;
}

function getEnabledSubsteps(){
  if (!tool.substeps_optional) return null;
  const ids = [];
  document.querySelectorAll('#prompts-list .check:checked').forEach(cb => ids.push(cb.dataset.sub));
  return ids;
}

// === Section 3: model controls ===
function buildModelControls(){
  const stored = JSON.parse(localStorage.getItem('toolkit.model_prefs') || '{}');

  // Model dropdown — match Claude desktop list (specific versions + Legacy badge + 1M)
  const ms = document.getElementById('sel-model');
  ms.innerHTML = '';
  // Default to whichever entry the registry marks as default (currently Opus 4.7)
  // 跳过已知不可用的模型,避免默认选到一个会失败的 model
  const availableModels = claudeInfo.available_models.filter(m => m.available !== false);
  const defaultPick = availableModels.find(m => m.default) || availableModels[0]
                      || claudeInfo.available_models[0];
  const desktopDefault = defaultPick.key;
  // 用户上次选的如果已经不可用,也别再默认到它
  const storedAvailable = stored.model && claudeInfo.available_models.find(
    m => m.key === stored.model && m.available !== false);
  const defaultModel = storedAvailable ? stored.model : desktopDefault;
  claudeInfo.available_models.forEach(m => {
    const sel = m.key === defaultModel ? 'selected' : '';
    const badge = m.version_badge ? ` ${m.version_badge}` : '';
    // 探测读回的真实版本号 → 动态附在档位名后(如 "Opus · 4.8");未探测则只显示档位名
    const _vm = (m.resolved_model || '').match(/([0-9]+)[.-]([0-9]+)/);
    const ver = _vm ? ` · ${_vm[1]}.${_vm[2]}` : '';
    const classes = [];
    if (m.legacy) classes.push('legacy');
    if (m.available === false) classes.push('unavailable');
    if (m.experimental) classes.push('experimental');
    const disabled = m.available === false ? 'disabled' : '';
    // 不可用的模型在 label 前缀里挂个标记,disabled 浏览器自带样式会变灰
    const prefix = m.available === false ? '✕ ' : (m.experimental ? '⚠ ' : '');
    const tail = m.available === false ? '（账号未开通 / 已下线）' : '';
    ms.insertAdjacentHTML('beforeend',
      `<option value="${m.key}" ${sel} ${disabled} class="${classes.join(' ')}">${prefix}${m.label}${ver}${badge} — ${m.tag}${tail}</option>`);
  });
  const defaultLabel = (claudeInfo.available_models.find(m => m.key === desktopDefault) || {}).label || 'Opus';
  document.getElementById('model-hint').textContent = '· 所有子步骤统一使用此模型';
  document.getElementById('model-from').textContent = stored.model
    ? '来源：上次选择'
    : `来源：默认 ${defaultLabel}`;

  // Effort dropdown — only filled if model supports it (handled by syncCapabilities)
  const ef = document.getElementById('sel-effort');
  const defaultEffort = stored.effort || claudeInfo.settings_effort_level || 'medium';

  // Thinking dropdown
  const th = document.getElementById('sel-thinking');
  const defaultThinking = stored.thinking || 'disabled';

  // Build effort/thinking options ONCE; visibility & population per model
  function fillEffort(model){
    const allowed = new Set(model.supported_efforts || []);
    ef.innerHTML = '';
    claudeInfo.available_efforts
      .filter(e => allowed.has(e.key))
      .forEach(e => {
        const sel = e.key === defaultEffort ? 'selected' : '';
        ef.insertAdjacentHTML('beforeend', `<option value="${e.key}" ${sel}>${e.label} — ${e.tag}</option>`);
      });
    // If stored default isn't in allowed set, just pick first
    if (!allowed.has(defaultEffort) && ef.options.length){
      ef.selectedIndex = 0;
    }
  }
  function fillThinking(model){
    const allowed = new Set(model.supported_thinking || []);
    th.innerHTML = '';
    claudeInfo.available_thinking
      .filter(t => allowed.has(t.key))
      .forEach(t => {
        const sel = t.key === defaultThinking ? 'selected' : '';
        th.insertAdjacentHTML('beforeend', `<option value="${t.key}" ${sel}>${t.label}</option>`);
      });
    if (!allowed.has(defaultThinking) && th.options.length){
      th.selectedIndex = 0;
    }
  }

  function syncCapabilities(){
    const sel = document.getElementById('sel-model').value;
    const m = claudeInfo.available_models.find(x => x.key === sel) || claudeInfo.available_models[0];
    const supportsEffort = !!m.supports_effort;
    const supportsThinking = !!m.supports_thinking;

    document.getElementById('cell-effort').style.display = supportsEffort ? '' : 'none';
    document.getElementById('cell-thinking').style.display = supportsThinking ? '' : 'none';
    // experimental / unavailable 也强制显示 — 不光是"两个都不支持时才提示"
    const forceShow = !!m.experimental || m.available === false;
    document.getElementById('cell-unsupported').style.display =
      (forceShow || (!supportsEffort && !supportsThinking)) ? 'block' : 'none';

    if (supportsEffort) fillEffort(m);
    if (supportsThinking) fillThinking(m);

    // Tailor unsupported hint to current model
    const isHaikuFamily = String(m.model || '').includes('haiku');
    let hint = isHaikuFamily
      ? `Haiku 系列定位是高速 / 低成本任务，桌面端也不暴露 effort 与 thinking 选项。如需调精度请切换到 Sonnet 或 Opus。`
      : `${m.label} 不支持精度 / 思考调节。`;
    if (m.experimental) {
      hint = `⚠️ ${m.label} 当前为实验级。若你的账号未开通该模型，运行会因 SDK 调用失败而无报告 — 建议先用 Sonnet / Opus 验证流程，或点 [测试该模型] 现场探测。\n` + hint;
    }
    if (m.available === false) {
      hint = `✕ ${m.label} 在当前账号上不可用。原因：${m.unavailable_reason || '上次调用失败已被会话禁用'}。请改选 Sonnet / Opus，或到设置页修复后重启。\n` + hint;
    }
    const hintEl = document.getElementById('unsupported-hint');
    hintEl.textContent = hint;
    hintEl.style.whiteSpace = 'pre-line';
    // 给实验模型加一个"现场探测"按钮 — 现场打一个 1-token 请求看模型能不能跑
    const cell = document.getElementById('cell-unsupported');
    let btn = cell.querySelector('button.probe-btn');
    if (m.experimental && m.available !== false){
      if (!btn){
        btn = document.createElement('button');
        btn.className = 'probe-btn';
        btn.type = 'button';
        btn.style.cssText = 'margin-top:8px;padding:6px 12px;border:1px solid var(--accent);background:transparent;color:var(--accent);border-radius:3px;font-family:var(--mono);font-size:11.5px;letter-spacing:.04em;cursor:pointer';
        cell.appendChild(btn);
      }
      btn.textContent = '测试该模型 (1-token 探测)';
      btn.onclick = async () => {
        btn.disabled = true; btn.textContent = '探测中…';
        try {
          const r = await fetch('/api/claude/probe_model', {
            method:'POST', headers:{'Content-Type':'application/json'},
            body: JSON.stringify({model_key: m.key})
          }).then(r=>r.json());
          if (r.ok){
            btn.textContent = `✓ ${m.label} 可用`;
            btn.style.color = 'var(--ok)'; btn.style.borderColor = 'var(--ok)';
          } else if (r.model_specific){
            btn.textContent = `✕ ${m.label} 当前账号未开通 — 已禁用`;
            btn.style.color = 'var(--bad)'; btn.style.borderColor = 'var(--bad)';
            // 刷新 claudeInfo,让下拉立即 disable
            try {
              claudeInfo = await fetch('/api/claude/info').then(r => r.json());
              buildModelControls();
            } catch(_){}
          } else {
            // 非模型问题(认证/网络/CLI 损坏)— 不禁用,只警告,让用户去设置页排查
            btn.textContent = '⚠ 探测失败 — 可能是认证/网络,请到设置排查';
            btn.title = String(r.error || '').slice(0, 400);
            btn.style.color = 'var(--warn)'; btn.style.borderColor = 'var(--warn)';
          }
        } catch(e){
          btn.textContent = '探测请求失败';
        } finally {
          setTimeout(() => { btn.disabled = false; }, 1500);
        }
      };
    } else if (btn){
      btn.remove();
    }

    document.getElementById('effort-from').textContent = stored.effort && supportsEffort
      ? `沿用上次：${ef.value}`
      : `默认：${claudeInfo.settings_effort_level || 'medium'}`;
  }
  syncCapabilities();

  // ── 选模型 → 自动探测该模型 + 刷新接入,即时显示 ✓/✗ ──
  // 用户选什么模型就跑什么,选完立刻知道这个模型通不通,不用跑完整工具才发现。
  let _probeTimer = null;
  async function autoProbeModel(){
    const sel = document.getElementById('sel-model');
    const statusEl = document.getElementById('model-probe-status');
    if (!sel || !statusEl) return;
    const m = claudeInfo.available_models.find(x => x.key === sel.value);
    const label = m ? m.label : sel.value;
    statusEl.style.display = '';
    statusEl.style.color = 'var(--fg-3)';
    statusEl.textContent = `· 正在探测 ${label} 接入…`;
    try {
      const r = await fetch('/api/claude/probe_model', {
        method:'POST', headers:{'Content-Type':'application/json'},
        body: JSON.stringify({model_key: sel.value}),
      }).then(x => x.json());
      if (r.ok){
        statusEl.style.color = 'var(--ok)';
        statusEl.textContent = `✓ ${label} 接入正常,可运行`;
      } else if (r.model_specific){
        statusEl.style.color = 'var(--bad)';
        statusEl.textContent = `✗ ${label} 当前账号不可用 — 请换其他模型`;
      } else {
        statusEl.style.color = 'var(--warn)';
        statusEl.textContent = `⚠ ${label} 探测失败(认证/网络),运行可能受影响`;
        statusEl.title = String(r.error || '').slice(0, 300);
      }
    } catch(e){
      statusEl.style.color = 'var(--warn)';
      statusEl.textContent = `⚠ 探测请求失败:${e.message}`;
    }
  }
  function scheduleProbe(){
    // debounce 600ms — 用户连续切换时只探最后一次
    if (_probeTimer) clearTimeout(_probeTimer);
    _probeTimer = setTimeout(autoProbeModel, 600);
  }

  // Persist + re-sync on any change
  document.getElementById('sel-model').addEventListener('change', () => {
    syncCapabilities();
    persistPrefs();
    scheduleProbe();   // 换模型 → 自动探测
  });
  ['sel-effort','sel-thinking'].forEach(id => {
    document.getElementById(id).addEventListener('change', () => {
      persistPrefs();
      scheduleProbe();  // 换精度/思考 → 也重新确认一次接入
    });
  });
  // 页面加载后先探一次当前选中的模型
  scheduleProbe();

  function persistPrefs(){
    const cur = JSON.parse(localStorage.getItem('toolkit.model_prefs') || '{}');
    cur.model = document.getElementById('sel-model').value;
    // Only persist effort/thinking if their cells are visible
    if (document.getElementById('cell-effort').style.display !== 'none'){
      cur.effort = document.getElementById('sel-effort').value;
    }
    if (document.getElementById('cell-thinking').style.display !== 'none'){
      cur.thinking = document.getElementById('sel-thinking').value;
    }
    localStorage.setItem('toolkit.model_prefs', JSON.stringify(cur));
  }

  // Claude status line —— 优先显示 toolkit 接入模式(API Key / OAuth),
  // 没接入则提示去设置;本机 CLI 版本作为辅助信息
  const cs = document.getElementById('claude-status');
  if (claudeInfo.bin_found){
    const acc = claudeInfo.account || {};
    const auth = claudeInfo.auth_state || '';
    let modeChip = '';
    if (auth === 'toolkit_api_key') {
      modeChip = `<span>API Key</span><span>·</span><span style="font-family:var(--mono);font-size:11.5px">${acc.display_name || ''}</span>`;
    } else if (auth === 'toolkit_oauth') {
      const email = acc.email ? ` · ${acc.email}` : '';
      modeChip = `<span>OAuth</span>${email ? '<span>·</span><span>'+email+'</span>' : ''}`;
    } else if (auth === 'external_claude_cli' && acc.email) {
      modeChip = `<span>本机 CLI 凭据</span><span>·</span><span>${acc.email}</span>`;
    } else {
      modeChip = `<span style="color:var(--warn)">未接入 — 请先到 <a href="/settings">设置</a> 选 OAuth 或 API Key</span>`;
    }
    const okOrWarn = (auth === 'toolkit_api_key' || auth === 'toolkit_oauth') ? 'ok' : 'warn';
    cs.className = 'claude-status ' + okOrWarn;
    cs.innerHTML = `<span class="dot"></span>
      <span>Claude CLI <span class="v">${claudeInfo.version || '?'}</span></span>
      <span>·</span>${modeChip}
      <a href="/settings" style="margin-left:auto">详情</a>`;
  } else {
    cs.className = 'claude-status bad';
    cs.innerHTML = `<span class="dot"></span><span>未找到 Claude Code CLI</span><a href="/settings">查看设置</a>`;
  }
}

// === Figma 设计稿条(step5 / h5_adapt)=== 读图走宿主机「登录态浏览器导出」(零 REST 额度)
async function refreshFigmaStatus(){
  const st=document.getElementById('figma-status');
  if(!st) return;
  try{
    const d=await fetch('/api/figma/status').then(r=>r.json());
    if(d.session_logged_in){
      st.textContent='✅ 已登录 Figma 会话 — 粘链接直接运行 / 读取预览(登录态浏览器导出,零额度,只读权限也支持)';
      st.style.color='var(--ok)';
    } else if(d.token_configured){
      st.textContent='⚠ 未导入会话,暂用只读 Token(REST,有账号级额度限制)— 建议导入会话走零额度';
      st.style.color='var(--warn)';
    } else {
      st.textContent='⚠ 未登录 Figma 会话 — 需把 figma.com 会话导入宿主机(零额度读图)';
      st.style.color='var(--warn)';
    }
  }catch(e){ st.textContent='Figma 状态查询失败'; }
}
function initFigmaBar(){
  const bar=document.getElementById('figma-bar'); if(!bar) return;
  bar.style.display='block';
  const readBtn=document.getElementById('figma-read-btn');
  const urlInp=document.getElementById('figma-url');
  const gal=document.getElementById('figma-preview-gallery');
  const st=document.getElementById('figma-status');
  // 刷新清空(输入框永不禁用)
  urlInp.value=''; urlInp.disabled=false;
  // 换链接 → 「读取完成」还原为「读取预览」
  urlInp.oninput=()=>{ if(readBtn.textContent==='读取完成'){ readBtn.textContent='读取预览'; gal.style.display='none'; } };
  // 读取 figma 内容(登录态浏览器导出):成功→展示全部帧;失败→说明原因
  readBtn.onclick=async()=>{
    const u=(urlInp.value||'').trim();
    if(!u){ st.textContent='先粘贴 Figma 链接'; st.style.color='var(--warn)'; return; }
    readBtn.disabled=true; readBtn.textContent='读取中(首次约1分钟,之后秒开)…';
    try{
      const r=await fetch('/api/figma/preview',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({url:u})}).then(r=>r.json());
      if(r.ok){
        const urls=(r.image_urls&&r.image_urls.length)?r.image_urls:(r.image_url?[r.image_url]:[]);
        gal.innerHTML=urls.map((u,i)=>`<img src="${u}" title="设计帧 ${i+1}/${urls.length}" style="height:220px;border:1px solid var(--line);border-radius:6px;background:#fff"/>`).join('');
        gal.style.display='flex';
        readBtn.textContent='读取完成';
        st.textContent='✓ 读取成功'+(r.frames?`(共 ${r.frames} 帧设计图 — 下方可滚动查看全部)`:''); st.style.color='var(--ok)';
      }
      else { readBtn.textContent='读取预览'; st.textContent='✕ 读取失败:'+(r.error||'未知原因'); st.style.color='var(--bad)'; }
    }catch(e){ readBtn.textContent='读取预览'; st.textContent='✕ 读取失败:网络错误'; st.style.color='var(--bad)'; }
    readBtn.disabled=false;
  };
  refreshFigmaStatus();
}

// === Run ===
async function runTool(){
  if (!tool) return;
  // 强校验:项目编号 + 项目名称必填
  const projCodeEl = document.getElementById('proj-code');
  const projNameEl = document.getElementById('proj-name');
  const projCode = (projCodeEl ? projCodeEl.value : '').trim();
  const projName = (projNameEl ? projNameEl.value : '').trim();
  if (!projCode){
    toast('请填写项目编号');
    if (projCodeEl) { projCodeEl.focus(); projCodeEl.style.borderColor = 'var(--bad)'; }
    return;
  }
  if (!projName){
    toast('请填写项目名称');
    if (projNameEl) { projNameEl.focus(); projNameEl.style.borderColor = 'var(--bad)'; }
    return;
  }
  // 记住上次填的项目,下次预填
  try {
    localStorage.setItem('toolkit.last_project', JSON.stringify({code:projCode, name:projName}));
  } catch(_){}

  const body = {project_code: projCode, project_name: projName};
  let text = document.getElementById('doc-input').value.trim();
  // step5:把 Figma 链接框的值并入文档(后端据此读设计图)
  const _figEl = document.getElementById('figma-url');
  if (_figEl && _figEl.value.trim() && !text.includes(_figEl.value.trim())){
    text = (text ? text + '\n\n## 设计稿\n' : '') + _figEl.value.trim();
  }
  if (!text){ toast('请先填入文档内容 / Figma 链接'); return; }
  const inp = tool.input || {};
  const fmt = inp.format || 'text';
  const key = inp.primary_key || 'documents';

  if (fmt === 'json'){
    try { body[key] = JSON.parse(text); }
    catch(err){ toast(`输入不是合法 JSON：${err.message}`); return; }
  } else if (fmt === 'auto'){
    // Try JSON first, fall back to raw text under '_documents'
    try { body[key] = JSON.parse(text); }
    catch(_){ body[key] = text; }
  } else {
    body[key] = text;
  }

  // Run options (e.g. dry_run for step6)
  (tool.run_options || []).forEach(o => {
    const el = document.getElementById('ro-' + o.key);
    if (el) body[o.key] = el.checked;
  });

  // Always use the user's selected model/effort/thinking — applied to all substeps.
  // The per-prompt model_tier in frontmatter is now reference-only.
  const m = document.getElementById('sel-model').value;
  const ef = document.getElementById('sel-effort').value;
  const th = document.getElementById('sel-thinking').value;
  if (m) body.__model = m;
  if (ef) body.__effort = ef;
  if (th) body.__thinking = th;
  // Optional per-substep filter — only sent when the tool supports it
  const enabled = getEnabledSubsteps();
  if (enabled !== null){
    if (enabled.length === 0){
      toast('至少勾选一个子步骤');
      return;
    }
    if (enabled.length < tool.prompts.length){
      body.__enabled_substeps = enabled;
    }
  }

  const btn = document.getElementById('btn-run');
  btn.disabled = true; btn.innerHTML = '提交中…';
  let started = false;
  try {
    const r = await fetch(`/api/tools/${tool.id}/run`, {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify(body),
    });
    // 412 = 认证未配置 — 给清晰指引而不是丢一坨 JSON
    if (r.status === 412){
      const errBody = await r.json().catch(()=>({}));
      const msg = errBody.detail || '需要先到设置页选择连接方式';
      if (confirm(`无法启动:${msg}\n\n是否现在跳到设置页?`)){
        location.href = '/settings#auth';
      }
      return;
    }
    if (r.status === 409){
      // 409 = 已有正在运行的同 tool 任务 — 提示用户等待
      const errBody = await r.json().catch(()=>({}));
      toast(errBody.detail || '该工具已被其他人触发,请等它完成');
      return;
    }
    const res = await r.json();
    if (!res.run_id){ toast('启动失败:' + JSON.stringify(res)); return; }
    started = true;
    currentRunId = res.run_id;
    history.replaceState(null, '', `?run=${res.run_id}`);
    // reset tab state for new run
    _activeReportTab = null;
    window._userSwitchedTab = false;
    _lastLogCount = 0;
    startPolling(res.run_id);
    document.getElementById('run-area').scrollIntoView({behavior:'smooth', block:'start'});
    // 启动成功 — 按钮维持锁定状态(显示"运行中"),由 pollActiveStatus 接管最终解锁
    btn.disabled = true; btn.innerHTML = '<span>运行中</span>';
    // 立刻拉一次 active 状态确保 UI 同步(也避免 3s 间隔的空窗期能再次点击)
    if (typeof pollActiveStatus === 'function') pollActiveStatus();
  } catch(err){
    toast('启动请求出错: ' + err.message);
  } finally {
    if (!started){
      // 没成功启动 — 解锁让用户能修正后重试
      btn.disabled = false; btn.innerHTML = '▶ 运行<kbd>⌘↵</kbd>';
    }
  }
}

document.getElementById('btn-run').onclick = runTool;
document.addEventListener('keydown', e => {
  if ((e.metaKey || e.ctrlKey) && e.key === 'Enter'){ e.preventDefault(); runTool(); }
});
// 右侧 56px 窄条上的 ▶ 圆按钮也能触发运行
document.addEventListener('click', e => {
  const t = e.target;
  if (t && (t.id === 'play-circle' || (t.closest && t.closest('#play-circle')))) {
    runTool();
  }
});

// ============== 全局运行状态:不管谁触发,不管刷不刷新,都能看到 "运行中" ==============
// 后端 /api/tools/{tool_id}/active 返回当前是否有任意用户在跑这个工具。
// 前端每 3s 轮询一次:active=true → 锁按钮 + "运行中";active=false → 解锁。
let _activeStatusTimer = null;
let _lastActiveSig = null;

async function pollActiveStatus(){
  if (!tool || !tool.id) return;
  let r;
  try {
    r = await fetch(`/api/tools/${tool.id}/active`).then(x => x.json());
  } catch(_) { return; }
  const btn = document.getElementById('btn-run');
  if (!btn) return;
  const sig = r.active ? ('active:' + (r.run_id || '?')) : 'idle';
  if (sig === _lastActiveSig) return;  // 没变化就不更新
  _lastActiveSig = sig;
  if (r.active){
    btn.disabled = true;
    const owner = r.owner_username ? ` · 由 ${escapeHtml(r.owner_username)} 触发` : '';
    btn.innerHTML = `<span>运行中</span><span style="opacity:.7;font-size:11px;margin-left:6px">${owner}</span>`;
    btn.title = (r.progress || '') + (owner ? '\n' + r.owner_username : '');
  } else {
    btn.disabled = false;
    btn.innerHTML = '▶ 运行<kbd>⌘↵</kbd>';
    btn.title = '';
  }
}

function startActiveStatusPolling(){
  pollActiveStatus();  // 立即拉一次,页面加载完就反映当前状态
  if (_activeStatusTimer) clearInterval(_activeStatusTimer);
  _activeStatusTimer = setInterval(pollActiveStatus, 3000);
}

function startPolling(runId){
  if (pollTimer) clearInterval(pollTimer);
  if (elapsedTimer) clearInterval(elapsedTimer);
  poll(runId);
  pollTimer = setInterval(()=>poll(runId), 1500);
  elapsedTimer = setInterval(updateElapsed, 1000);
}
async function poll(runId){
  // 内存里没找着先回退看磁盘:URL 带 ?run=xxx 的"分享 / 回看历史"场景
  // 之前直接报 404 → 顶栏 UNDEFINED + 暂无日志,看上去像报告丢了。
  try {
    const resp = await fetch(`/api/tools/runs/${runId}`);
    if (resp.status === 404){
      clearInterval(pollTimer); pollTimer = null;
      clearInterval(elapsedTimer); elapsedTimer = null;
      await loadHistoryFromDisk(runId);
      return;
    }
    if (!resp.ok) throw new Error('http ' + resp.status);
    const r = await resp.json();
    currentRun = r;
    renderRun(r);
    if (r.status === 'succeeded' || r.status === 'failed' || r.status === 'cancelled'){
      clearInterval(pollTimer); pollTimer = null;
      clearInterval(elapsedTimer); elapsedTimer = null;
    }
  } catch(e){
    clearInterval(pollTimer); pollTimer = null;
    clearInterval(elapsedTimer); elapsedTimer = null;
    // 网络/JSON 错误也尝试从磁盘读 — 如果是 history 链接至少能渲染。
    try { await loadHistoryFromDisk(runId); } catch(_){}
  }
}

async function stopCurrentRun(runId){
  if (!confirm('确定停止这个运行？已产生的进度会丢失。')) return;
  const btns = document.querySelectorAll('.status-stop');
  btns.forEach(b => { b.disabled = true; b.textContent = '停止中…'; });
  try {
    const r = await fetch(`/api/tools/runs/${runId}/stop`, {method:'POST'}).then(r => r.json());
    if (r.ok){ poll(runId); }
    else { alert(r.detail || '停止失败'); btns.forEach(b => { b.disabled = false; b.textContent = '■ 停止'; }); }
  } catch(e){ alert('停止失败：' + e); btns.forEach(b => { b.disabled = false; b.textContent = '■ 停止'; }); }
}

async function loadHistoryFromDisk(runId){
  // 把 /api/reports/{run_id} 的返回包成 poll 已经认识的 run 形状再喂给 renderRun。
  let payload;
  try {
    const resp = await fetch('/api/reports/' + encodeURIComponent(runId));
    if (!resp.ok) throw new Error('http ' + resp.status);
    payload = await resp.json();
  } catch(e){
    const area = document.getElementById('run-area');
    if (area) {
      area.innerHTML = `<div class="run-empty">报告未找到 — 可能服务重启后内存清掉了，且磁盘也没保存。<br><a href="/reports">去 /reports 查找</a></div>`;
    }
    const rsl = document.getElementById('run-status-label');
    const rsd = document.getElementById('run-status-dot');
    if (rsl) rsl.textContent = '未找到';
    if (rsd) rsd.className = 'live-dot failed';
    return;
  }
  // 磁盘 payload: {source, run_id, tool_id, report, report_path}
  const report = payload.report || {};
  const meta = report.meta || {};
  const finalizedAt = meta.produced_at_epoch || ((Date.now()/1000) - 0.1);
  const fakeRun = {
    run_id: payload.run_id || runId,
    tool_id: payload.tool_id || (tool && tool.id) || '',
    status: 'succeeded',
    progress: '历史报告',
    started_at: finalizedAt,
    finished_at: finalizedAt,
    report,
    logs: [],
    usage: meta.usage || {},
    project_code: meta.project_code,
    project_name: meta.project_name,
    __historical: true,
  };
  currentRun = fakeRun;
  renderRun(fakeRun);
}
function updateElapsed(){
  if (!currentRun || currentRun.finished_at) return;
  const e = Date.now()/1000 - currentRun.started_at;
  document.getElementById('s-elapsed').textContent = e.toFixed(1) + 's';
  document.getElementById('s-elapsed-wrap').style.display = '';
}
function gateClass(action){
  const a = String(action || '').toLowerCase();
  if (a.includes('reject')) return 'reject';
  if (a.includes('proceed') || a.includes('approve')) return 'proceed';
  return 'warn';
}
// Track active tab so re-renders preserve user choice. Default rules:
//   running → 'logs'   succeeded → 'report'   failed → 'logs'
let _activeReportTab = null;

function renderRun(r){
  // 有运行任务 → 展开右侧运行面板
  const grid = document.querySelector('.runner-grid');
  if (grid) grid.classList.remove('idle');
  // Top-bar status + token/cost stats (+ 停止按钮:仅执行人/管理员、运行中)
  document.getElementById('tb-sep').style.display = '';
  const _stop = r.can_stop
    ? ` <button class="status-stop" onclick="stopCurrentRun('${r.run_id}')" title="停止执行（仅执行人/管理员）">■ 停止</button>`
    : '';
  document.getElementById('tb-status').innerHTML = `<span class="status-pill ${r.status}">${r.status}</span>${_stop}`;
  const u = r.usage || {};
  document.getElementById('s-in').textContent = u.input_tokens != null ? u.input_tokens.toLocaleString() : '—';
  document.getElementById('s-out').textContent = u.output_tokens != null ? u.output_tokens.toLocaleString() : '—';
  document.getElementById('s-cache').textContent = u.cache_read_tokens != null ? u.cache_read_tokens.toLocaleString() : '—';
  document.getElementById('s-cost').textContent = u.cost_usd != null ? u.cost_usd.toFixed(4) : '—';
  if (r.finished_at){
    document.getElementById('s-elapsed').textContent = (r.finished_at - r.started_at).toFixed(1) + 's';
    document.getElementById('s-elapsed-wrap').style.display = '';
  }
  // Reveal stats once we have a run (idle: hidden)
  const stats = document.getElementById('tb-stats');
  if (stats) stats.style.display = 'flex';
  // Right-panel header: live status dot + label + meta
  const rsd = document.getElementById('run-status-dot');
  const rsl = document.getElementById('run-status-label');
  const rsm = document.getElementById('run-status-meta');
  if (rsd && rsl) {
    rsd.className = 'live-dot ' + r.status;
    const labels = {queued:'排队中', running:'执行中', succeeded:'已完成', failed:'失败'};
    rsl.textContent = labels[r.status] || r.status;
    if (rsm) {
      const elapsed = r.finished_at ? (r.finished_at - r.started_at) : (Date.now()/1000 - r.started_at);
      const min = String(Math.floor(elapsed/60)).padStart(2,'0');
      const sec = String(Math.floor(elapsed%60)).padStart(2,'0');
      rsm.textContent = (r.status === 'queued') ? '' : `${min}:${sec}${u.cost_usd != null ? ' · $'+u.cost_usd.toFixed(4) : ''}`;
    }
  }

  // Default tab choice based on status
  if (_activeReportTab === null){
    _activeReportTab = (r.status === 'succeeded') ? 'report' : 'logs';
  }
  // If just completed and still on logs, auto-switch to report
  if (r.status === 'succeeded' && _activeReportTab === 'logs' && !window._userSwitchedTab){
    _activeReportTab = 'report';
  }

  const area = document.getElementById('run-area');

  // Failed: show error + logs tab
  if (r.status === 'failed'){
    const tabs = renderReportTabs(r, 'logs');
    const errBlock = `<div class="err-view">${escapeHtml(r.traceback || r.error || 'failed')}</div>`;
    area.innerHTML = tabs.head + errBlock + tabs.bodyOf('logs', r);
    bindTabClicks(r);
    autoScrollLogs(r);
    return;
  }

  // Still running: tabs (logs default + step list at top)
  if (r.status !== 'succeeded'){
    const tabs = renderReportTabs(r, _activeReportTab);
    const stepList = renderStepList(r);
    area.innerHTML = tabs.head + stepList + tabs.bodyOf(_activeReportTab, r);
    bindTabClicks(r);
    autoScrollLogs(r);
    return;
  }

  // Succeeded — show structured report by default, logs+JSON in tabs
  const tabs = renderReportTabs(r, _activeReportTab);
  area.innerHTML = tabs.head + tabs.bodyOf(_activeReportTab, r);
  bindTabClicks(r);
  bindReportInteractions();
  autoScrollLogs(r);
  // 报告渲染后异步把所有 <img data-screenshot-filename=...> 的 src 替换成 data: URI
  inlineScreenshotsInArea(area);
}

// 把当前 area 下所有标记了 data-screenshot-filename 的 img,异步 fetch 内容
// 然后把 src 替换为 base64 data URI — 报告就完全自包含了。
// 已经是 data: 开头的跳过。同一文件名缓存避免重复请求。
const _screenshotInlineCache = {};
async function inlineScreenshotsInArea(rootEl){
  if (!rootEl) return;
  const imgs = rootEl.querySelectorAll('img[data-screenshot-filename]');
  if (!imgs.length) return;
  await Promise.all(Array.from(imgs).map(async img => {
    if (img.src && img.src.startsWith('data:')) return;
    const fn = img.dataset.screenshotFilename;
    if (!fn) return;
    if (_screenshotInlineCache[fn]) {
      img.src = _screenshotInlineCache[fn];
      return;
    }
    try {
      const resp = await fetch('/api/screenshots/' + encodeURIComponent(fn));
      if (!resp.ok) return;
      const blob = await resp.blob();
      const reader = new FileReader();
      const dataUri = await new Promise((res, rej) => {
        reader.onload = () => res(reader.result);
        reader.onerror = () => rej(reader.error);
        reader.readAsDataURL(blob);
      });
      _screenshotInlineCache[fn] = dataUri;
      img.src = dataUri;
    } catch(_){}
  }));
}

function renderStepList(r){
  let html = '<div class="step-list">';
  tool.prompts.forEach(p => {
    let cls='', m='○';
    if (r.status === 'succeeded'){ cls='done'; m='✓'; }
    else if (r.status === 'running'){ cls='running'; m='◔'; }
    const shortP = String(p).split('.').pop();
    html += `<div class="step-row ${cls}"><span class="marker">${m}</span><span class="name" title="${p}">#${shortP}</span><span class="info">${cls === 'done' ? 'OK' : (cls === 'running' ? '执行中' : '')}</span></div>`;
  });
  html += '</div>';
  return html;
}

function renderReportTabs(r, active){
  // step2 测试用例工具:产出就是用例,不是"报告" —— tab 叫"用例",只给 Excel 下载。
  const isStep2 = (typeof tool !== 'undefined' && tool && tool.id === 'step2');
  const reportTabLabel = isStep2 ? '用例' : '报告';
  const dlButtons = `${r.report ? `<button class="export" data-action="download-xlsx">↓ 下载 Excel</button>` : ''}
       ${r.report ? `<button class="export" data-action="download-md">↓ Markdown</button>` : ''}
       ${r.report ? `<button class="export" data-action="download-json">↓ JSON</button>` : ''}
       ${r.report ? `<button class="export" data-action="copy">⧉ 复制</button>` : ''}`;
  const head = `<div class="report-tabs">
    <button class="${active==='report'?'active':''}" data-tab="report">${reportTabLabel}</button>
    <button class="${active==='logs'?'active':''}" data-tab="logs">日志 ${(r.logs||[]).length ? '('+r.logs.length+')' : ''}</button>
    <button class="${active==='json'?'active':''}" data-tab="json">{ } 原始 JSON</button>
    <span class="spacer"></span>
    ${dlButtons}
  </div>`;
  const bodyOf = (tab, run) => {
    if (tab === 'logs') return renderLogs(run) || '<div class="run-empty">暂无日志</div>';
    if (tab === 'json'){
      if (!run.report) return '<div class="run-empty">报告尚未生成</div>';
      return `<div class="json-view"><pre>${escapeHtml(JSON.stringify(run.report, null, 2))}</pre></div>`;
    }
    // report
    if (!run.report) return '<div class="run-empty">报告尚未生成</div>';
    return renderHtmlReport(run);
  };
  return { head, bodyOf };
}

function bindTabClicks(r){
  document.querySelectorAll('#run-area .report-tabs button[data-tab]').forEach(btn => {
    btn.onclick = () => {
      _activeReportTab = btn.dataset.tab;
      window._userSwitchedTab = true;
      renderRun(r);
    };
  });
  document.querySelectorAll('#run-area .report-tabs button[data-action]').forEach(btn => {
    const action = btn.dataset.action;
    if (action === 'download-json'){
      btn.onclick = () => downloadBlob(
        JSON.stringify(r.report, null, 2),
        `${tool.id}_${r.run_id.slice(0,8)}.json`,
        'application/json'
      );
    } else if (action === 'download-md'){
      btn.onclick = () => downloadBlob(
        buildMarkdownReport(r),
        `${tool.id}_${r.run_id.slice(0,8)}.md`,
        'text/markdown;charset=utf-8'
      );
    } else if (action === 'download-xlsx'){
      // Excel 用例表 — 服务端 openpyxl 生成,走 export.xlsx 端点
      btn.onclick = async () => {
        try {
          const resp = await fetch(`/api/reports/${r.run_id}/export.xlsx`, {credentials:'same-origin'});
          if (!resp.ok){
            const d = await resp.json().catch(()=>({}));
            toast('Excel 导出失败:' + (d.detail || resp.status));
            return;
          }
          const blob = await resp.blob();
          const a = document.createElement('a');
          a.href = URL.createObjectURL(blob);
          a.download = `${tool.id}_${r.run_id.slice(0,8)}.xlsx`;
          document.body.appendChild(a); a.click(); a.remove();
          setTimeout(()=>URL.revokeObjectURL(a.href), 1000);
        } catch(e){ toast('Excel 导出出错:' + e.message); }
      };
    } else if (action === 'copy'){
      btn.onclick = async () => {
        try {
          await navigator.clipboard.writeText(JSON.stringify(r.report, null, 2));
          toast('已复制完整 JSON');
        } catch(e){ toast('复制失败'); }
      };
    }
  });
}

function downloadBlob(content, filename, mime){
  const blob = new Blob([content], {type: mime});
  const url = URL.createObjectURL(blob);
  const a = document.createElement('a');
  a.href = url; a.download = filename;
  document.body.appendChild(a);
  a.click();
  document.body.removeChild(a);
  setTimeout(() => URL.revokeObjectURL(url), 1000);
  toast(`已下载 ${filename}`);
}

function bindReportInteractions(){
  document.querySelectorAll('#run-area .report-sub-head').forEach(h => {
    h.onclick = () => h.parentElement.classList.toggle('open');
  });
}

// === Exporters: Markdown ===

function buildMarkdownReport(r){
  // Produce a clean markdown version for sharing/version control.
  const rep = r.report || {};
  const meta = rep.meta || {};
  const u = r.usage || {};
  const lines = [];
  lines.push(`# ${tool.name} · 分析报告`);
  lines.push('');
  lines.push(`- **生成时间**：${new Date().toLocaleString('zh-CN', {hour12:false})}`);
  lines.push(`- **Run ID**：\`${r.run_id}\``);
  lines.push(`- **模型**：\`${meta.model_id || '?'}\``);
  if (r.finished_at && r.started_at) lines.push(`- **耗时**：${(r.finished_at - r.started_at).toFixed(1)}s`);
  if (u.cost_usd != null) lines.push(`- **成本**：$${u.cost_usd}`);
  if (u.output_tokens != null) lines.push(`- **输出 tokens**：${u.output_tokens.toLocaleString()}`);
  lines.push('');

  // Gate
  if (rep.gate_decision){
    const g = rep.gate_decision;
    lines.push(`## ⚠️ 闸门决策：\`${g.action}\``);
    (g.reasons || []).forEach(rs => lines.push(`- ${rs}`));
    lines.push('');
  }

  // Substeps
  let idx = 0;
  for (const sid of tool.prompts){
    idx++;
    const data = (rep.substeps || {})[sid];
    const title = (tool._substepNames && tool._substepNames[sid]) || extractTitle(data) || `子分析 ${idx}`;
    lines.push(`## ${idx}. ${title}`);
    if (data == null){
      lines.push('_已跳过_');
      lines.push('');
      continue;
    }
    lines.push('');
    lines.push('```json');
    lines.push(JSON.stringify(data, null, 2));
    lines.push('```');
    lines.push('');
  }
  return lines.join('\n');
}

// === Structured report renderer ===

function buildExecutiveSummary(rep, tool){
  // 收集所有 substep 输出里的 issue/case 节点(向后兼容旧报告)
  const walkedIssues = [], walkedCases = [];
  function walk(x){
    if (x && typeof x === 'object'){
      if (Array.isArray(x)){ x.forEach(walk); return; }
      const hasCase = ('expected' in x || 'scenario' in x) && ('id' in x);
      const hasIssue = 'severity' in x && (('issue' in x) || ('title' in x) || ('description' in x) || ('name' in x));
      if (hasCase) walkedCases.push(x);
      else if (hasIssue) walkedIssues.push(x);
      Object.values(x).forEach(walk);
    }
  }
  walk(rep.substeps || {});

  // 顶层契约字段(新报告)优先,缺失则 fallback 到 walk 结果
  const topIssues = Array.isArray(rep.issues) ? rep.issues : walkedIssues;
  const topCases = Array.isArray(rep.cases) ? rep.cases : walkedCases;
  const topRisks = Array.isArray(rep.risks) ? rep.risks : [];
  const topBlockers = Array.isArray(rep.blockers) ? rep.blockers : [];

  // severity & priority 排序权重
  const sevRank = {critical:0, high:1, medium:2, low:3, info:4};
  const priRank = {P0:0, P1:1, P2:2, P3:3};
  function getSev(x){ return String(x.severity||'medium').toLowerCase(); }
  function getPri(x){ return String(x.priority||'P2').toUpperCase(); }

  // 整形 issues
  const naturalIssues = topIssues.map(it => ({
    issue_id: String(it.issue_id || it.id || ''),
    title: String(it.title || it.name || it.issue || it.description || '未命名问题').slice(0, 200),
    severity: getSev(it),
    priority: getPri(it),
    module: String(it.module || it.endpoint || it.file_path || it.location || it.viewport || it.page || it.viewport_filename || '').slice(0, 200),
    current: String(it.current_behavior || it.current || it.observed || it.description || it.issue || '').slice(0, 800),
    expected: String(it.expected_behavior || it.expected || it.requirement || '').slice(0, 800),
    fix: String(it.fix_suggestion || it.fix || it.recommendation || it.suggestion || it.remediation || '').slice(0, 1000),
    repro: Array.isArray(it.reproduce_steps) ? it.reproduce_steps : (it.reproduce_steps ? [String(it.reproduce_steps)] : []),
    accept: String(it.acceptance_criteria || it.acceptance || it.verify || '').slice(0, 600),
    cases: Array.isArray(it.related_test_cases) ? it.related_test_cases : (it.related_test_cases ? [String(it.related_test_cases)] : []),
    owner: String(it.owner_role || it.owner || it.assignee || '').toLowerCase(),
    hours: it.estimated_hours || it.effort || null,
    impact: String(it.impact_scope || it.impact || '').slice(0, 400),
    evidence: String(it.evidence || it.source || '').slice(0, 400),
  })).sort((a,b) => {
    const sa = sevRank[a.severity] ?? 9, sb = sevRank[b.severity] ?? 9;
    if (sa !== sb) return sa - sb;
    const pa = priRank[a.priority] ?? 9, pb = priRank[b.priority] ?? 9;
    return pa - pb;
  }).slice(0, 60);

  // 整形 risks (新报告是对象; 旧报告 gate.reasons 是字符串数组)
  let naturalRisks = topRisks.map(r => {
    if (typeof r === 'string') return {title: r, impact:'', why:'', severity:'medium'};
    return {
      id: String(r.id || ''),
      title: String(r.title || r.name || r.risk || '未命名风险').slice(0, 200),
      impact: String(r.impact || r.affects || '').slice(0, 400),
      why: String(r.why || r.reason || r.detail || '').slice(0, 400),
      severity: getSev(r),
    };
  });
  // fallback: 旧报告从 gate_decision.reasons 抽 risks
  if (!naturalRisks.length){
    const gate = rep.gate_decision || {};
    if (Array.isArray(gate.reasons)){
      naturalRisks = gate.reasons.filter(Boolean).map(s => ({title:String(s), impact:'', why:'', severity:'medium'}));
    }
  }

  // 整形 blockers
  const naturalBlockers = topBlockers.map(b => ({
    id: String(b.id || ''),
    title: String(b.title || b.name || '未命名阻碍').slice(0, 200),
    why_blocking: String(b.why_blocking || b.reason || b.why || '').slice(0, 500),
    what_to_unblock: String(b.what_to_unblock || b.action || b.fix || '').slice(0, 500),
    owner_role: String(b.owner_role || b.owner || '').toLowerCase(),
    hours: b.estimated_hours || b.effort || null,
  }));

  // 汇总 severity & priority 计数(基于排序后的 issues + walkedCases)
  const sev = {critical:0, high:0, medium:0, low:0, info:0};
  naturalIssues.forEach(it => { if (sev[it.severity] !== undefined) sev[it.severity]++; });
  const pri = {P0:0, P1:0, P2:0, P3:0, 其他:0};
  topCases.forEach(c => {
    const p = String(c.priority||'').toUpperCase();
    if (pri[p] !== undefined) pri[p]++;
    else if (p) pri.其他++;
  });

  // 整形 cases (按 priority 排序)
  const naturalCases = topCases.map(c => ({
    id: String(c.id || c.case_id || c.tc_id || ''),
    title: String(c.title || c.name || c.scenario || '').slice(0, 200),
    priority: String(c.priority || 'P2').toUpperCase(),
    type: String(c.type || c.kind || '').toLowerCase(),
    status: String(c.status || 'designed').toLowerCase(),
    automation: String(c.automation_tag || c.automation || '').toLowerCase(),
    preconditions: String(c.preconditions || c.precondition || '').slice(0, 300),
    steps: Array.isArray(c.steps) ? c.steps : (c.steps ? [String(c.steps)] : []),
    expected: String(c.expected || c.expected_result || '').slice(0, 400),
  })).sort((a,b) => {
    const pa = priRank[a.priority] ?? 9, pb = priRank[b.priority] ?? 9;
    return pa - pb;
  }).slice(0, 200);

  // verdict / verdict_summary
  const gate = rep.gate_decision || {};
  const action = String(gate.action || '').toLowerCase();
  let verdict, vlevel;
  if (rep.verdict){
    const v = String(rep.verdict);
    verdict = v;
    if (v.includes('不通过')) vlevel = 'fail';
    else if (v.includes('有条件') || v.includes('警告') || v.includes('部分')) vlevel = 'warn';
    else vlevel = 'pass';
  } else {
    if (action.includes('reject') || sev.critical > 0){ verdict = '不通过'; vlevel = 'fail'; }
    else if (action.includes('warn') || sev.high > 2 || naturalBlockers.length){ verdict = '有条件通过'; vlevel = 'warn'; }
    else if (!action && !naturalIssues.length && !naturalCases.length){ verdict = '未产出'; vlevel = 'skip'; }
    else { verdict = '通过'; vlevel = 'pass'; }
  }
  const verdictSummary = String(rep.verdict_summary || '').slice(0, 200);

  return {
    verdict, vlevel, verdictSummary,
    risks: naturalRisks,
    blockers: naturalBlockers,
    issues: naturalIssues,
    cases: naturalCases,
    casesCount: naturalCases.length,
    pri, sev,
    agent: tool ? tool.name : '',
  };
}

function renderExecutiveSummary(s){
  const vmap = {pass:['✅','通过','ok'], warn:['⚠️','有条件通过','warn'], fail:['❌','不通过','bad'], skip:['—','未产出','skip']};
  const [vicon, vtext, vcls] = vmap[s.vlevel] || ['·','—','skip'];
  const sevMap = {critical:'CRITICAL', high:'HIGH', medium:'MEDIUM', low:'LOW', info:'INFO'};
  const priMap = {P0:'P0', P1:'P1', P2:'P2', P3:'P3'};
  const ownerMap = {backend:'后端', frontend:'前端', product:'产品', test:'测试',
                    devops:'运维', security:'安全', data:'数据'};
  const sevTotal = s.sev.critical + s.sev.high + s.sev.medium + s.sev.low + s.sev.info;
  const priTotal = s.pri.P0 + s.pri.P1 + s.pri.P2 + s.pri.P3 + s.pri.其他;

  // ── 严重度可视化条 ──
  function sevBar(){
    if (!sevTotal) return '<span class="exec-muted">（暂无问题）</span>';
    const cells = ['critical','high','medium','low','info'].map(k => {
      const n = s.sev[k]; if (!n) return '';
      const pct = (n / sevTotal * 100).toFixed(0);
      return `<span class="sev-bar-seg sev-bar-${k}" style="flex:${n}" title="${sevMap[k]} · ${n} 个 (${pct}%)">${n}</span>`;
    }).filter(Boolean).join('');
    return `<div class="sev-bar">${cells}</div>`;
  }
  function priBar(){
    if (!priTotal) return '';
    const cells = ['P0','P1','P2','P3'].map(k => {
      const n = s.pri[k]; if (!n) return '';
      return `<span class="pri-bar-seg pri-bar-${k}" style="flex:${n}" title="${k} · ${n} 条">${k}: ${n}</span>`;
    }).filter(Boolean).join('');
    return `<div class="pri-bar">${cells}</div>`;
  }

  let html = '';

  // ── ① 测试结论(verdict + verdict_summary 大字)──
  html += `<div class="exec-block exec-verdict-block">
    <div class="exec-head"><span class="exec-num">①</span><h3>测试结论</h3></div>
    <div class="verdict ${vcls}"><span class="verdict-icon">${vicon}</span><span>${escapeHtml(vtext)}</span></div>
    ${s.verdictSummary ? `<div class="verdict-summary">${escapeHtml(s.verdictSummary)}</div>` : ''}
    <div class="exec-kpis">
      <div class="exec-kpi"><span class="exec-kpi-num">${sevTotal}</span><span class="exec-kpi-lbl">问题总数</span></div>
      <div class="exec-kpi"><span class="exec-kpi-num">${s.sev.critical + s.sev.high}</span><span class="exec-kpi-lbl">需立即处理</span></div>
      <div class="exec-kpi"><span class="exec-kpi-num">${s.blockers.length}</span><span class="exec-kpi-lbl">阻碍</span></div>
      <div class="exec-kpi"><span class="exec-kpi-num">${s.casesCount}</span><span class="exec-kpi-lbl">用例</span></div>
    </div>
    ${sevBar()}
  </div>`;

  // ── ② 风险结论 ──
  let risksHtml;
  if (s.risks.length){
    risksHtml = '<div class="exec-risk-list">' + s.risks.map(r => {
      const sevBadge = r.severity ? `<span class="sev-tag sev-${r.severity}">${escapeHtml(r.severity)}</span>` : '';
      return `<div class="exec-risk-item">
        <div class="exec-risk-head">${sevBadge}<span class="exec-risk-title">${escapeHtml(r.title)}</span></div>
        ${r.impact ? `<div class="exec-risk-line"><span class="lbl">影响</span>${escapeHtml(r.impact)}</div>` : ''}
        ${r.why ? `<div class="exec-risk-line"><span class="lbl">原因</span>${escapeHtml(r.why)}</div>` : ''}
      </div>`;
    }).join('') + '</div>';
  } else {
    risksHtml = '<p class="exec-muted">（无显著风险）</p>';
  }
  html += `<div class="exec-block">
    <div class="exec-head"><span class="exec-num">②</span><h3>风险结论</h3><span class="exec-count">${s.risks.length}</span></div>
    ${risksHtml}
  </div>`;

  // ── ③ 阻碍 ──
  let blockersHtml;
  if (s.blockers.length){
    blockersHtml = '<div class="exec-blocker-list">' + s.blockers.map((b, idx) => `
      <div class="exec-blocker-item">
        <div class="exec-blocker-head">
          <span class="blocker-tag">BLOCKER</span>
          <span class="exec-blocker-title">${idx+1}. ${escapeHtml(b.title)}</span>
          ${b.id ? `<span class="meta-chip">${escapeHtml(b.id)}</span>` : ''}
          ${b.owner_role && ownerMap[b.owner_role] ? `<span class="meta-chip role">👤 ${ownerMap[b.owner_role]}</span>` : ''}
          ${b.hours ? `<span class="meta-chip">⏱ ${b.hours}h</span>` : ''}
        </div>
        ${b.why_blocking ? `<div class="exec-blocker-line"><span class="lbl">为何阻碍</span>${escapeHtml(b.why_blocking)}</div>` : ''}
        ${b.what_to_unblock ? `<div class="exec-blocker-line fix"><span class="lbl">如何解开</span>${escapeHtml(b.what_to_unblock)}</div>` : ''}
      </div>`).join('') + '</div>';
  } else {
    blockersHtml = '<p class="exec-muted">（无阻碍）</p>';
  }
  html += `<div class="exec-block">
    <div class="exec-head"><span class="exec-num">③</span><h3>阻碍</h3><span class="exec-count danger">${s.blockers.length}</span></div>
    ${blockersHtml}
  </div>`;

  // ── ④ Bug 表(按严重度+优先级排序)──
  function renderIssueCard(it, idx){
    const metaChips = [];
    if (it.issue_id) metaChips.push(`<span class="meta-chip">${escapeHtml(it.issue_id)}</span>`);
    if (it.priority) metaChips.push(`<span class="meta-chip pri-${it.priority}">${priMap[it.priority] || it.priority}</span>`);
    if (it.owner && ownerMap[it.owner]) metaChips.push(`<span class="meta-chip role">👤 ${ownerMap[it.owner]}</span>`);
    if (it.hours) metaChips.push(`<span class="meta-chip">⏱ ${it.hours}h</span>`);
    const reproHtml = it.repro.length
      ? '<ol class="repro-list">' + it.repro.map(s => `<li>${escapeHtml(s)}</li>`).join('') + '</ol>' : '';
    const casesHtml = it.cases.length
      ? `<div class="related-cases">关联用例:${it.cases.map(c => `<code>${escapeHtml(c)}</code>`).join(' ')}</div>` : '';
    return `
      <div class="exec-issue sev-${it.severity}">
        <div class="exec-issue-head">
          <span class="sev-tag sev-${it.severity}">${sevMap[it.severity] || '·'}</span>
          <span class="exec-issue-title">${idx+1}. ${escapeHtml(it.title)}</span>
        </div>
        <div class="exec-issue-meta">${metaChips.join('')}</div>
        ${it.module ? `<div class="exec-issue-loc">位置:<code>${escapeHtml(it.module)}</code></div>` : ''}
        ${it.current ? `<div class="exec-issue-section"><div class="sec-lbl">现状</div><div class="sec-body">${escapeHtml(it.current)}</div></div>` : ''}
        ${it.expected ? `<div class="exec-issue-section"><div class="sec-lbl">期望</div><div class="sec-body">${escapeHtml(it.expected)}</div></div>` : ''}
        ${it.fix ? `<div class="exec-issue-section fix"><div class="sec-lbl">修复建议</div><div class="sec-body">${escapeHtml(it.fix)}</div></div>` : ''}
        ${(reproHtml || it.accept) ? `<div class="exec-issue-section verify"><div class="sec-lbl">验收</div><div class="sec-body">${reproHtml}${it.accept ? `<div class="accept-line">验收标准：${escapeHtml(it.accept)}</div>` : ''}</div></div>` : ''}
        ${casesHtml}
        ${it.impact ? `<div class="exec-issue-impact">影响面:${escapeHtml(it.impact)}</div>` : ''}
        ${it.evidence ? `<div class="exec-issue-evidence">证据:${escapeHtml(it.evidence)}</div>` : ''}
      </div>`;
  }
  const issuesHtml = s.issues.length
    ? s.issues.map((it, i) => renderIssueCard(it, i)).join('')
    : '<p class="exec-muted">（本次未识别到具体问题）</p>';
  html += `<div class="exec-block">
    <div class="exec-head">
      <span class="exec-num">④</span>
      <h3>Bug 表</h3>
      <span class="exec-count">${s.issues.length}</span>
      <span class="exec-count-note">(按严重度 × 优先级排序)</span>
    </div>
    ${issuesHtml}
  </div>`;

  // ── ⑤ 执行用例记录(按 P0→P3)──
  let casesListHtml;
  if (s.cases.length){
    const rows = s.cases.map((c, i) => {
      const statusMap = {
        designed: ['','已设计','muted'],
        executed_pass: ['','已执行通过','ok'],
        executed_fail: ['','执行失败','bad'],
        skipped: ['','已跳过','muted'],
        blocked: ['','阻塞','bad'],
      };
      const [sIcon, sLabel, sCls] = statusMap[c.status] || ['','未定义','muted'];
      return `<tr class="case-row pri-${c.priority}">
        <td class="case-idx">${i+1}</td>
        <td><span class="pri-tag pri-${c.priority}">${priMap[c.priority]||c.priority}</span></td>
        <td><code class="case-id">${escapeHtml(c.id)}</code></td>
        <td class="case-title">${escapeHtml(c.title)}</td>
        <td>${c.type ? `<span class="case-type">${escapeHtml(c.type)}</span>` : ''}</td>
        <td>${c.automation ? `<span class="case-auto">${escapeHtml(c.automation)}</span>` : ''}</td>
        <td><span class="case-status case-status-${sCls}">${sIcon} ${sLabel}</span></td>
      </tr>`;
    }).join('');
    casesListHtml = `<div class="case-table-wrap">
      <table class="case-table">
        <thead><tr><th>#</th><th>优先级</th><th>用例 ID</th><th>用例标题</th><th>类型</th><th>自动化</th><th>状态</th></tr></thead>
        <tbody>${rows}</tbody>
      </table>
    </div>`;
  } else if (s.casesCount){
    casesListHtml = `<div class="exec-case-total"><span class="num">${s.casesCount}</span><span class="lbl">条已生成用例(未提供详细列表)</span></div>`;
  } else {
    casesListHtml = '<p class="exec-muted">（本次未生成用例）</p>';
  }
  html += `<div class="exec-block">
    <div class="exec-head">
      <span class="exec-num">⑤</span>
      <h3>执行用例记录</h3>
      <span class="exec-count">${s.casesCount}</span>
    </div>
    ${priBar()}
    ${casesListHtml}
  </div>`;

  return html;
}

// 测试用例工具:把用例表 + Excel 下载做成报告主体
function buildUiComparisonBlock(rep, runId){
  const meta = (rep && rep.meta) || {};
  const shots = (meta.screenshots || []).filter(s => !s.error);
  const isDesign = (s) => s.is_design || /figma\.com/i.test(s.url||'');
  const designShots = shots.filter(isDesign);
  const actualShots = shots.filter(s => !isDesign(s));
  const issues = (rep.issues || []);
  const scOf = (s) => ({critical:'crit',high:'hi',major:'hi',medium:'med',low:'lo',minor:'lo',info:'lo',cosmetic:'lo'})[(s||'').toLowerCase()]||'med';
  const zhOf = (s) => ({critical:'严重',high:'高',major:'高',medium:'中',low:'低',minor:'低',info:'提示',cosmetic:'细节'})[(s||'').toLowerCase()]||'中';
  const verdict = rep.verdict || '';
  const vsum = rep.verdict_summary || '';
  const vCls = /不通过|reject|fail/i.test(verdict)?'bad':(/有条件|warn|conditional/i.test(verdict)?'warn':'ok');
  const imgTag = (fn) => `<img data-screenshot-filename="${escapeHtml(fn)}" src="/api/screenshots/${encodeURIComponent(fn)}" loading="lazy">`;
  const css = `<style>
  .ui-cmp{margin:0 0 18px}
  .ui-cmp-verdict{padding:14px 18px;border-radius:10px;border:1px solid var(--line,#ddd);margin-bottom:16px;font-size:14px;line-height:1.7}
  .ui-cmp-verdict.ok{background:rgba(22,163,74,.06);border-color:rgba(22,163,74,.3)}
  .ui-cmp-verdict.warn{background:rgba(202,138,4,.06);border-color:rgba(202,138,4,.3)}
  .ui-cmp-verdict.bad{background:rgba(220,38,38,.06);border-color:rgba(220,38,38,.3)}
  .ui-cmp-pair{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-bottom:18px;align-items:start}
  .ui-cmp-side{border:1px solid var(--line,#ddd);border-radius:10px;overflow:hidden;background:var(--surface,#f8fafc)}
  .ui-cmp-label{font-size:12.5px;font-weight:600;padding:8px 12px;border-bottom:1px solid var(--line,#ddd);color:var(--fg-2,#555)}
  .ui-cmp-label.actual{color:#dc2626}
  .ui-cmp-side img{display:block;width:100%;height:auto}
  .ui-cmp-nodes{padding:40px 20px;text-align:center;color:var(--fg-3,#888);font-size:13px}
  .ui-cmp-list-head{font-size:14px;font-weight:700;margin:8px 0 12px;color:var(--fg,#222)}
  .ui-cmp-item{border:1px solid var(--line,#ddd);border-left:4px solid var(--line-2,#ccc);border-radius:0 8px 8px 0;padding:12px 16px;margin-bottom:10px;background:var(--surface,#f8fafc)}
  .ui-cmp-item.crit{border-left-color:#dc2626}.ui-cmp-item.hi{border-left-color:#ea580c}.ui-cmp-item.med{border-left-color:#ca8a04}.ui-cmp-item.lo{border-left-color:#16a34a}
  .ui-cmp-ih{display:flex;align-items:center;gap:10px;margin-bottom:8px}
  .ui-cmp-num{width:22px;height:22px;border-radius:50%;background:var(--fg,#222);color:#fff;display:grid;place-items:center;font-size:12px;font-weight:700;flex-shrink:0}
  .ui-cmp-t{font-weight:700;font-size:14.5px;flex:1;color:var(--fg,#222)}
  .ui-cmp-sev{font-size:11px;font-weight:700;padding:2px 8px;border-radius:4px}
  .ui-cmp-sev.crit{background:rgba(220,38,38,.12);color:#dc2626}.ui-cmp-sev.hi{background:rgba(234,88,12,.12);color:#ea580c}.ui-cmp-sev.med{background:rgba(202,138,4,.12);color:#a16207}.ui-cmp-sev.lo{background:rgba(22,163,74,.12);color:#16a34a}
  .ui-cmp-r{display:flex;gap:8px;font-size:13.5px;line-height:1.7;margin-top:3px;color:var(--fg,#222)}
  .ui-cmp-r .l{flex-shrink:0;width:56px;font-size:10.5px;font-weight:700;color:var(--fg-3,#888);background:var(--surface-2,#eef);padding:2px 6px;border-radius:3px;height:fit-content;text-align:center}
  .ui-cmp-r.design .l{color:#0891b2}.ui-cmp-r.actual .l{color:#dc2626}
  @media(max-width:720px){.ui-cmp-pair{grid-template-columns:1fr}}
  </style>`;
  let h = css + '<div class="ui-cmp">';
  if (verdict || vsum){
    h += `<div class="ui-cmp-verdict ${vCls}"><b>设计符合度：${escapeHtml(verdict||'—')}</b>${vsum?' — '+escapeHtml(vsum):''}</div>`;
  }
  // 按真实配对一对一展示:每个实拍配它自己对应的设计帧(不再全部贴第一张 designShots[0])
  const pcMap = {}; (rep.pairs_checked||[]).forEach(p => { pcMap[p.actual] = p.design; });
  const designByName = {};
  designShots.forEach(d => { const nm = String(d.node_name||d.viewport||''); if(nm && !(nm in designByName)) designByName[nm]=d; });
  if (actualShots.length){
    actualShots.forEach(act => {
      const actFn = act.annotated_filename || act.filename;
      const an = act.viewport || act.filename;
      const dn = pcMap[an];
      const paired = dn && dn !== '(无对应设计帧)';
      const des = paired ? designByName[String(dn)] : null;
      h += `<div class="ui-cmp-pair">
        <div class="ui-cmp-side"><div class="ui-cmp-label">🎨 设计稿${paired?'『'+escapeHtml(dn)+'』':''}</div>${des?imgTag(des.filename):'<div class="ui-cmp-nodes">无对应设计帧</div>'}</div>
        <div class="ui-cmp-side"><div class="ui-cmp-label actual">📱 实际产品 · ${escapeHtml(an)} · 红框=与设计不一致</div>${imgTag(actFn)}</div>
      </div>`;
    });
  } else if (designShots.length){
    h += `<div class="ui-cmp-pair"><div class="ui-cmp-side"><div class="ui-cmp-label">🎨 设计稿（Figma）</div>${imgTag(designShots[0].filename)}</div><div class="ui-cmp-side"><div class="ui-cmp-nodes">未捕获到实际产品截图</div></div></div>`;
  }
  h += `<div class="ui-cmp-list-head">不一致清单（${issues.length} 处）</div>`;
  if (!issues.length){
    h += '<div class="ui-cmp-nodes" style="border:1px dashed var(--line,#ddd);border-radius:8px">未发现与设计稿的明显不一致 🎉</div>';
  } else {
    issues.forEach((it, i) => {
      const sc = scOf(it.severity);
      h += `<div class="ui-cmp-item ${sc}">
        <div class="ui-cmp-ih"><span class="ui-cmp-num">${i+1}</span><span class="ui-cmp-t">${escapeHtml(it.title||it.issue_id||'差异')}</span><span class="ui-cmp-sev ${sc}">${zhOf(it.severity)}</span></div>
        <div class="ui-cmp-r design"><span class="l">设计要求</span><span>${escapeHtml(it.expected_behavior||it.expected||'—')}</span></div>
        <div class="ui-cmp-r actual"><span class="l">实际</span><span>${escapeHtml(it.current_behavior||it.current||'—')}</span></div>
        ${it.fix_suggestion?`<div class="ui-cmp-r"><span class="l">建议</span><span>${escapeHtml(it.fix_suggestion)}</span></div>`:''}
      </div>`;
    });
  }
  h += '</div>';
  return h;
}

function buildTestCaseBlock(rep, runId, toolId){
  const cases = (rep && rep.cases) || [];
  if (!cases.length) return '';
  const rows = cases.map(c => {
    const stepsArr = Array.isArray(c.steps) ? c.steps : (c.steps ? [c.steps] : []);
    const steps = stepsArr.filter(Boolean).map(s => escapeHtml(String(s))).join('\n');
    let exp = c.expected;
    if (Array.isArray(exp)) exp = exp.map(e => typeof e==='string'?e:JSON.stringify(e)).join('\n');
    exp = escapeHtml(String(exp || ''));
    const pri = String(c.priority||'').toUpperCase();
    return '<tr>' +
      '<td class="tc-id">' + escapeHtml(c.id||'') + '</td>' +
      '<td>' + escapeHtml(c.module||'') + '</td>' +
      '<td class="tc-title">' + escapeHtml(c.title||c.name||'') + '</td>' +
      '<td><span class="tc-pri tc-pri-' + pri + '">' + escapeHtml(pri||'-') + '</span></td>' +
      '<td>' + escapeHtml(c.type||'') + '</td>' +
      '<td class="tc-pre">' + escapeHtml(c.preconditions||'') + '</td>' +
      '<td class="tc-steps">' + steps + '</td>' +
      '<td class="tc-exp">' + exp + '</td>' +
    '</tr>';
  }).join('');
  return '<div class="tc-block">' +
    '<div class="tc-banner">' +
      '<div><div class="tc-count">' + cases.length + '</div><div class="tc-count-label">条测试用例</div></div>' +
      '<div class="tc-mid">本工具产出 <b>纯人工执行</b> 测试用例,步骤是自然语言操作清单。' +
        '点右侧按钮导出标准 Excel(含「执行结果 / 实际结果」空列),可直接发给测试团队执行。</div>' +
      '<button class="tc-excel-btn" data-tc-runid="' + runId + '" data-tc-toolid="' + (toolId||'step2') + '">↓ 下载 Excel 用例表</button>' +
    '</div>' +
    '<div class="tc-table-wrap"><table class="tc-table"><thead><tr>' +
      '<th>用例编号</th><th>模块</th><th>用例标题</th><th>优先级</th><th>类型</th>' +
      '<th>前置条件</th><th>测试步骤</th><th>预期结果</th>' +
    '</tr></thead><tbody>' + rows + '</tbody></table></div>' +
  '</div>';
}

function renderHtmlReport(r, opts){
  opts = opts || {};
  const expandAll = !!opts.expandAll;
  const rep = r.report || {};
  const meta = rep.meta || {};
  const u = r.usage || {};
  let html = '';

  // Hero card
  const elapsed = r.finished_at ? (r.finished_at - r.started_at).toFixed(1) : '—';
  const pc = meta.project_code || r.project_code || '';
  const pn = meta.project_name || r.project_name || '';
  const projectRow = (pc || pn)
    ? `<div class="report-project">
        <span class="lbl">项目编号</span><code>${escapeHtml(pc || '—')}</code>
        <span class="lbl">项目名称</span><span class="val">${escapeHtml(pn || '—')}</span>
      </div>` : '';
  html += `<div class="report-hero">
    <div class="report-icon">${tool.icon}</div>
    <div style="flex:1;min-width:0">
      <h4>${escapeHtml(tool.name)} · 报告</h4>
      <div class="meta">
        ${meta.produced_at_utc ? new Date(meta.produced_at_utc).toLocaleString('zh-CN', {hour12:false}) : ''}
        · 模型 <code>${escapeHtml(meta.model_id || u.model || '?')}</code>
        · run <code>${r.run_id.slice(0, 12)}</code>
      </div>
      ${projectRow}
      <div class="stat-pills">
        <span class="pill">用时 <span class="v">${elapsed}s</span></span>
        <span class="pill">成本 <span class="v">$${u.cost_usd ?? '—'}</span></span>
        <span class="pill">输出 <span class="v">${(u.output_tokens||0).toLocaleString()}</span> tokens</span>
        <span class="pill">缓存 <span class="v">${(u.cache_read_tokens||0).toLocaleString()}</span></span>
      </div>
    </div>
  </div>`;

  // 测试用例工具(step2):只出用例表 + Excel,不要 verdict / 风险 / Bug 那套"报告"。
  // 用例就是用例 —— 报告主体只有用例表,到此为止直接返回。
  if (tool && tool.id === 'step2'){
    const tcBlock = buildTestCaseBlock(rep, r.run_id, tool.id);
    html += tcBlock || '<div class="run-empty" style="padding:48px;text-align:center">本次未生成测试用例</div>';
    return html;
  }

  // UI 一致性比对(step5):并排「设计稿 vs 实际产品」对比图 + 一条条不一致清单
  if (tool && tool.id === 'step5'){
    html += buildUiComparisonBlock(rep, r.run_id);
    return html;
  }

  // === 4 块结构（测试结论 / 风险结论 / 问题描述 / 用例执行情况）===
  // 适用于其余 7 个 Agent 的报告
  const summary = buildExecutiveSummary(rep, tool);
  html += renderExecutiveSummary(summary);

  // Inline page screenshots (step5 / h5_adapt only)
  // 用 data-screenshot-filename 占位,渲染后 inlineScreenshotsInArea() 把 src
  // 替换成 data: URI,报告自包含、不暴露任何本地文件夹路径。
  const shots = (meta.screenshots || []).filter(s => !s.error);
  if (shots.length){
    const groupedByUrl = {};
    shots.forEach(s => {(groupedByUrl[s.url] = groupedByUrl[s.url] || []).push(s);});
    const imgMap = (opts && opts.imgMap) || {};
    const _shotGroupLabel = (u) => {
      if(!u) return '截图';
      if(u.indexOf('app://')===0) return '📱 APP 实拍(模拟器)· ' + u.replace('app://','');
      if(/figma\.com/i.test(u)) return '🎨 设计基线(Figma)';
      return '🌐 ' + u;
    };
    let shotsHtml = '<div class="report-screenshots">';
    shotsHtml += '<div class="screenshots-head">截图证据 · 设计基线 vs 实拍对照 <span class="screenshots-hint">（已嵌入报告，无本地文件依赖）</span></div>';
    Object.entries(groupedByUrl).forEach(([url, arr]) => {
      shotsHtml += `<div class="shot-group"><div class="shot-url">${escapeHtml(_shotGroupLabel(url))}</div>`;
      shotsHtml += '<div class="shot-grid">';
      arr.forEach(s => {
        const annotated = s.annotated_filename;
        const fnPrimary = annotated || s.filename;
        // 已 inline 直接用;否则放占位 src,后续 inliner 异步替换为 data: URI
        const initialSrc = imgMap[fnPrimary] || `/api/screenshots/${encodeURIComponent(fnPrimary)}`;
        const issueBadge = s.issue_count
          ? `<span class="issue-badge">${s.issue_count} 个问题</span>` : '';
        const dim = (s.width && s.height) ? ` · ${s.width}×${s.height}` : '';
        shotsHtml += `<div class="shot-cell" title="${escapeHtml(s.viewport)}${dim}${annotated?' · 已标注':''}">
          <img src="${initialSrc}" data-screenshot-filename="${escapeHtml(fnPrimary)}" alt="${escapeHtml(s.viewport)}" loading="lazy">
          <div class="shot-cap">${escapeHtml(s.viewport)}${dim}${issueBadge}</div>
        </div>`;
      });
      shotsHtml += '</div></div>';
    });
    shotsHtml += '</div>';
    html += shotsHtml;
  }

  // Substep cards — labeled by Chinese name + sequence number, NOT internal step id
  // 当顶层契约字段已含 issues/cases 时,substep 数据已被合并入卡片,默认折叠
  // 避免内容重复展示导致报告"很乱"。点击标题可手动展开看原始数据。
  const hasContractData = (Array.isArray(rep.issues) && rep.issues.length) ||
                          (Array.isArray(rep.cases) && rep.cases.length);
  const subs = rep.substeps || {};
  let firstWithContent = null;
  for (const sid of tool.prompts){
    if (subs[sid] && Object.keys(subs[sid]).length){ firstWithContent = sid; break; }
  }
  // 折叠区标题
  html += `<div class="substep-section-head">
    <span class="substep-section-title">各子步骤原始输出</span>
    <span class="substep-section-hint">${hasContractData ? '已聚合到上方 5 段;点击展开看原始数据' : '点击展开'}</span>
  </div>`;
  let idx = 0;
  for (const sid of tool.prompts){
    idx++;
    const data = subs[sid];
    const stats = data ? quickStats(data) : null;
    // 有契约字段时一律默认折叠;否则保留旧逻辑(首个有内容的展开)
    const open = expandAll || (!hasContractData && sid === firstWithContent);
    const title = (tool._substepNames && tool._substepNames[sid]) || extractTitle(data) || `子分析 ${idx}`;
    if (data == null){
      html += `<div class="report-sub">
        <div class="report-sub-head">
          <span class="report-sub-twirl">▶</span>
          <span class="report-sub-num">${idx}</span>
          <span class="report-sub-name" style="color:var(--fg-3)">${escapeHtml(title)} · 已跳过</span>
        </div>
      </div>`;
      continue;
    }
    html += `<div class="report-sub${open?' open':''}">
      <div class="report-sub-head">
        <span class="report-sub-twirl">▶</span>
        <span class="report-sub-num">${idx}</span>
        <span class="report-sub-name">${escapeHtml(title)}</span>
        <span class="report-sub-stats">${stats || ''}</span>
      </div>
      <div class="report-sub-body">
        ${renderSmart(data)}
      </div>
    </div>`;
  }
  return html;
}

// Pull a friendly title from common substep schemas
function extractTitle(data){
  if (!data || typeof data !== 'object') return '';
  // Look for common labels
  const keys = ['name','title'];
  for (const k of keys){
    if (typeof data[k] === 'string') return data[k];
  }
  return '';
}

// Quick stats chip on the substep header
function quickStats(data){
  if (!data || typeof data !== 'object') return '';
  const chips = [];
  // Common high-signal keys
  if (Array.isArray(data.cases)) chips.push(`<span class="chip">${data.cases.length} 用例</span>`);
  if (Array.isArray(data.issues)) chips.push(`<span class="chip">${data.issues.length} 问题</span>`);
  if (Array.isArray(data.scenarios)) chips.push(`<span class="chip">${data.scenarios.length} 场景</span>`);
  if (Array.isArray(data.endpoints)) chips.push(`<span class="chip">${data.endpoints.length} 接口</span>`);
  if (Array.isArray(data.pages)) chips.push(`<span class="chip">${data.pages.length} 页面</span>`);
  if (Array.isArray(data.matrix)) chips.push(`<span class="chip">${data.matrix.length} 矩阵项</span>`);
  if (Array.isArray(data.fix_list)) chips.push(`<span class="chip">${data.fix_list.length} 待修</span>`);
  if (data.confidence && typeof data.confidence.score === 'number'){
    chips.push(`<span class="chip">conf ${(data.confidence.score*100).toFixed(0)}%</span>`);
  }
  return chips.join('');
}

// === Smart auto-renderer for arbitrary JSON ===

function renderSmart(data, depth){
  depth = depth || 0;
  if (data === null || data === undefined) return '<span class="empty-array">—</span>';
  if (typeof data === 'string'){
    if (!data) return '<span class="empty-array">""</span>';
    return escapeHtml(data);
  }
  if (typeof data === 'number' || typeof data === 'boolean'){
    return `<span style="color:var(--ac)">${escapeHtml(String(data))}</span>`;
  }
  if (Array.isArray(data)){
    if (!data.length) return '<span class="empty-array">[]</span>';
    // Array of objects with shared keys → table
    const allObjects = data.every(x => x && typeof x === 'object' && !Array.isArray(x));
    if (allObjects && data.length >= 2){
      return renderObjectArrayAsTable(data);
    }
    if (allObjects){
      // Single object — just unwrap
      return renderSmart(data[0], depth+1);
    }
    // Array of primitives
    return '<ul class="report-list">' +
      data.slice(0, 50).map(x => `<li>${renderSmart(x, depth+1)}</li>`).join('') +
      (data.length > 50 ? `<li class="empty-array">…剩余 ${data.length - 50} 条已折叠</li>` : '') +
      '</ul>';
  }
  if (typeof data === 'object'){
    return renderObjectAsKv(data, depth);
  }
  return escapeHtml(String(data));
}

function renderObjectAsKv(obj, depth){
  const entries = Object.entries(obj);
  if (!entries.length) return '<span class="empty-array">{}</span>';
  return '<dl class="report-kv">' + entries.map(([k, v]) => {
    return `<dt>${escapeHtml(k)}</dt><dd>${renderFieldValue(k, v, depth)}</dd>`;
  }).join('') + '</dl>';
}

function renderFieldValue(key, v, depth){
  // Special schemas
  if ((key === 'severity' || key.endsWith('_severity')) && typeof v === 'string'){
    return `<span class="sev sev-${escapeHtml(v.toLowerCase())}">${escapeHtml(v)}</span>`;
  }
  if (key === 'confidence' && v && typeof v === 'object' && typeof v.score === 'number'){
    const pct = (v.score * 100).toFixed(0);
    return `<span class="confbar"><span class="track"><span class="fill" style="width:${pct}%"></span></span><span class="pct">${pct}%</span></span>${
      v.rationale ? ` <span style="color:var(--fg-3)">— ${escapeHtml(String(v.rationale).slice(0,140))}</span>` : ''
    }`;
  }
  if (key === 'gate_decision' && v && typeof v === 'object'){
    const cls = gateClass(v.action);
    const rs = (v.reasons || []).map(x => `<div>· ${escapeHtml(x)}</div>`).join('');
    return `<div class="gate-banner ${cls}" style="margin:0">
      <span class="badge">${escapeHtml(String(v.action).toLowerCase())}</span>
      <div class="reasons">${rs}</div>
    </div>`;
  }
  if (Array.isArray(v) && v.length > 5){
    // Big array → wrap in details
    return `<details class="report-detail" ${depth === 0 ? 'open' : ''}>
      <summary>${v.length} 项 · 点击展开</summary>
      <div>${renderSmart(v, depth+1)}</div>
    </details>`;
  }
  if (typeof v === 'object' && v !== null && !Array.isArray(v) && Object.keys(v).length > 6){
    return `<details class="report-detail">
      <summary>对象（${Object.keys(v).length} 字段）· 点击展开</summary>
      <div>${renderSmart(v, depth+1)}</div>
    </details>`;
  }
  return renderSmart(v, depth+1);
}

function renderObjectArrayAsTable(arr){
  // Collect all keys, prioritize common ones
  const keySet = new Set();
  arr.forEach(o => Object.keys(o).forEach(k => keySet.add(k)));
  const priority = ['id','title','name','severity','status','category','kind','endpoint','page','area','module','expected','actual','fix','impact','effort_hours','severity_if_fails'];
  const keys = [...keySet].sort((a,b)=>{
    const ai = priority.indexOf(a), bi = priority.indexOf(b);
    if (ai >= 0 && bi >= 0) return ai - bi;
    if (ai >= 0) return -1;
    if (bi >= 0) return 1;
    return 0;
  }).slice(0, 6);  // cap to 6 columns

  const head = `<thead><tr>${keys.map(k=>`<th>${escapeHtml(k)}</th>`).join('')}</tr></thead>`;
  const body = arr.slice(0, 30).map(o => {
    const cells = keys.map(k => `<td>${renderTableCell(k, o[k])}</td>`).join('');
    return `<tr>${cells}</tr>`;
  }).join('');
  const foot = arr.length > 30
    ? `<tfoot><tr><td colspan="${keys.length}">… 共 ${arr.length} 条，已显示前 30 条；查看全部请切到 JSON tab</td></tr></tfoot>`
    : '';
  return `<table class="report-table">${head}<tbody>${body}</tbody>${foot}</table>`;
}

function renderTableCell(key, v){
  if (v === null || v === undefined) return '<span class="empty-array">—</span>';
  if ((key === 'severity' || key.endsWith('_severity')) && typeof v === 'string'){
    return `<span class="sev sev-${escapeHtml(v.toLowerCase())}">${escapeHtml(v)}</span>`;
  }
  if (typeof v === 'string'){
    const s = v.length > 100 ? v.slice(0, 97) + '…' : v;
    return escapeHtml(s);
  }
  if (typeof v === 'number' || typeof v === 'boolean'){
    return `<span style="color:var(--ac)">${escapeHtml(String(v))}</span>`;
  }
  if (Array.isArray(v)){
    if (!v.length) return '<span class="empty-array">[]</span>';
    if (v.length <= 3 && v.every(x => typeof x === 'string' || typeof x === 'number')){
      return v.map(x => escapeHtml(String(x))).join(', ');
    }
    return `<details class="report-detail"><summary>${v.length} 项</summary><div>${renderSmart(v)}</div></details>`;
  }
  if (typeof v === 'object'){
    return `<details class="report-detail"><summary>{${Object.keys(v).length} 字段}</summary><div>${renderSmart(v)}</div></details>`;
  }
  return escapeHtml(String(v));
}

function escapeHtml(s){ return String(s).replace(/[&<>]/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;'}[c])); }

function fmtLogTs(ts, baseTs){
  const dt = ts - baseTs;
  if (dt < 0) return '+0s';
  if (dt < 60) return '+' + dt.toFixed(1) + 's';
  return '+' + Math.floor(dt/60) + 'm' + Math.round(dt%60) + 's';
}
function fmtLogMsg(entry){
  const skip = new Set(['ts','event']);
  const parts = [];
  for (const [k, v] of Object.entries(entry)){
    if (skip.has(k) || v == null) continue;
    let vs = String(v);
    if (vs.length > 60) vs = vs.slice(0, 57) + '…';
    const cls = (k === 'cost_usd' || k === 'output_tokens' || k === 'model') ? 'v ac' : 'v';
    parts.push(`<span class="k">${escapeHtml(k)}</span>=<span class="${cls}">${escapeHtml(vs)}</span>`);
  }
  return parts.join(' · ');
}
function renderLogs(r){
  const logs = r.logs || [];
  if (!logs.length) return '';
  const base = r.started_at || logs[0].ts;
  const recent = logs.slice(-120);  // streaming events spike count, allow more
  const lines = recent.map(e => {
    const evClass = e.event.replace(/\./g, '-');
    const isStream = e.event === 'llm.thinking' || e.event === 'llm.text';
    const isStreamFinal = e.event === 'llm.thinking.final' || e.event === 'llm.text.final';
    let msgHtml;
    if (isStream && e.tail){
      // Render streaming with prominent tail snippet on its own line
      const head = `<span class="k">sub</span>=<span class="v">${escapeHtml(e.sub_id||'?')}</span> · <span class="k">chars</span>=<span class="v ac">${(e.chars||0).toLocaleString()}</span>`;
      const tail = `<span class="tail">${escapeHtml(e.tail)}</span>`;
      msgHtml = head + tail;
    } else {
      msgHtml = fmtLogMsg(e);
    }
    const cls = isStream ? `log-line stream ${e.event === 'llm.text' ? 'text' : 'thinking'}` : 'log-line';
    return `<div class="${cls}">
      <span class="ts">${fmtLogTs(e.ts, base)}</span>
      <span class="ev ${evClass}">${escapeHtml(e.event)}</span>
      <span class="msg">${msgHtml}</span>
    </div>`;
  }).join('');
  const liveBadge = (r.status === 'running')
    ? '<span class="live"><span class="dot"></span>实时</span>'
    : '<span style="color:var(--fg-3)">已停止</span>';
  // Show streaming progress count too (helpful for long thinking sequences)
  const streamingCount = logs.filter(x => x.event === 'llm.thinking' || x.event === 'llm.text').length;
  return `<div class="log-panel-head">
    ${liveBadge}
    <span style="color:var(--fg-3)">日志</span>
    <span class="count">${logs.length} 条${streamingCount ? ` · ${streamingCount} 流式` : ''}</span>
  </div>
  <div class="log-panel" id="log-panel">${lines}</div>`;
}
let _lastLogCount = 0;
function autoScrollLogs(r){
  const lp = document.getElementById('log-panel');
  if (!lp) return;
  // Only auto-scroll if there are new entries since last render and user hasn't scrolled up
  const total = (r.logs || []).length;
  const atBottom = (lp.scrollHeight - lp.scrollTop - lp.clientHeight) < 40;
  if (total !== _lastLogCount && atBottom){
    lp.scrollTop = lp.scrollHeight;
  }
  _lastLogCount = total;
}

function toast(msg){
  const t = document.getElementById('toast');
  t.textContent = msg;
  t.classList.add('show');
  clearTimeout(window._tt);
  window._tt = setTimeout(() => t.classList.remove('show'), 2400);
}

// === Anchor nav (right rail) ===
(function(){
  const nav = document.getElementById('anchor-nav');
  if (!nav) return;
  const links = nav.querySelectorAll('a');
  links.forEach(a => {
    a.addEventListener('click', e => {
      e.preventDefault();
      const target = document.getElementById(a.dataset.target);
      if (target) target.scrollIntoView({behavior:'smooth', block:'start'});
    });
  });
  function update(){
    const heroBottom = document.querySelector('.hero')?.getBoundingClientRect().bottom ?? 200;
    nav.classList.toggle('show', heroBottom < 80);
    let activeIdx = 0;
    ['sec-1','sec-2','sec-3'].forEach((id, i) => {
      const el = document.getElementById(id);
      if (!el) return;
      const rect = el.getBoundingClientRect();
      if (rect.top < window.innerHeight * 0.45) activeIdx = i;
    });
    links.forEach((a, i) => a.classList.toggle('current', i === activeIdx));
  }
  window.addEventListener('scroll', update, {passive:true});
  window.addEventListener('resize', update);
  update();
})();

load();
</script>
</body></html>
"""


SETTINGS_HTML = r"""<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>设置 — 天枢·裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin><link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
  :root{--bg:#ffffff;--surface:#f0f0f0;--surface-2:#ebebeb;--surface-3:#dcdcdc;
    --line:#c4c4c4;--line-2:#9e9e9e;
    --fg:#0a0a0a;--fg-2:#262626;--fg-3:#4a4a4a;--fg-4:#6e6e6e;
    --ac:#a8401f;--ac-2:#c45a3a;--ac-bg:rgba(168,64,31,.14);--ac-line:rgba(168,64,31,.58);
    --warn:#8a5300;--ok:#4f6b35;--bad:#8a2d12;--info:#3f5560;
    --running:#7a4f00;
    --mono:'Inter','SF Mono',ui-monospace,Menlo,monospace;
    --sans:'Noto Sans SC','PingFang SC',-apple-system,'Microsoft YaHei',sans-serif;
    --serif:'Noto Serif SC','Songti SC','STSong',Georgia,serif;}
  @media (prefers-reduced-motion: reduce){*,*::before,*::after{animation-duration:.01ms!important;transition-duration:.01ms!important}}
  *{box-sizing:border-box}
  html,body{margin:0;background:var(--bg);color:var(--fg);font-family:var(--sans);min-height:100%;
    -webkit-font-smoothing:antialiased}
  body{background:
    radial-gradient(ellipse 90% 50% at 50% -10%, rgba(196,90,58,.07), transparent 65%) fixed,
    radial-gradient(ellipse 80% 40% at 50% 110%, rgba(255,255,255,.02), transparent 60%) fixed,
    var(--bg);}
  ::selection{background:var(--ac-bg);color:var(--ac-2)}
  ::-webkit-scrollbar{width:10px;height:10px}
  ::-webkit-scrollbar-track{background:transparent}
  ::-webkit-scrollbar-thumb{background:var(--surface-3);border-radius:5px;border:2px solid var(--bg)}
  ::-webkit-scrollbar-thumb:hover{background:var(--line-2)}
  .topbar{display:flex;align-items:center;gap:12px;height:56px;padding:0 24px;
    border-bottom:1px solid var(--line);background:rgba(255,255,255,.94);
    position:sticky;top:0;z-index:100;backdrop-filter:saturate(180%) blur(20px);
    -webkit-backdrop-filter:saturate(180%) blur(20px)}
  .topbar .logo{width:28px;height:28px;flex-shrink:0;
    background:linear-gradient(135deg,#262626,#1a1a1a);border-radius:6px;
    display:grid;place-items:center;color:#001f1a;font-weight:700;font-size:11px;
    letter-spacing:-.02em;
    box-shadow:0 1px 2px rgba(0,0,0,.4),inset 0 1px 0 rgba(255,255,255,.18)}
  .topbar .logo .logo-mark{width:18px;height:18px;display:block;
    filter:drop-shadow(0 .5px 0 rgba(255,255,255,.18))}
  .topbar h1{margin:0;font-size:15px;font-weight:600;letter-spacing:-.015em}
  .topbar h1.brand{display:inline-flex;align-items:baseline;gap:0}
  .topbar h1.brand strong{font-weight:600;color:var(--fg);letter-spacing:-.02em}
  .topbar h1.brand .brand-sub{font-weight:400;color:var(--fg-3);margin-left:8px;font-size:13px;letter-spacing:.005em}
  .topbar nav{display:flex;gap:2px;margin-left:24px}
  .topbar nav a{padding:6px 12px;border-radius:6px;color:var(--fg-2);text-decoration:none;
    font-size:13px;transition:background .12s,color .12s}
  .topbar nav a:hover{background:var(--surface-2);color:var(--fg)}
  .topbar nav a.active{background:var(--ac-bg);color:var(--ac-2)}
  main{max-width:1080px;margin:0 auto;padding:40px 24px 80px}
  h2{margin:0 0 8px;font-size:24px;letter-spacing:-.025em;font-weight:600}
  .sub{color:var(--fg-2);font-size:13px;margin-bottom:32px;line-height:1.55}
  .sec{background:var(--surface);border:1px solid var(--line);border-radius:14px;
    margin-top:14px;overflow:hidden}
  .sec-head{padding:16px 22px;border-bottom:1px solid var(--line);display:flex;
    align-items:center;gap:12px;background:var(--surface-2)}
  .sec-head h3{margin:0;font-size:15px;font-weight:600;letter-spacing:-.005em}
  .sec-head .dot{width:8px;height:8px;border-radius:50%;background:var(--fg-4)}
  .sec-head .dot.ok{background:var(--ok)}
  .sec-head .dot.bad{background:var(--bad)}
  .sec-head .dot.warn{background:var(--warn)}
  .sec-body{padding:14px 18px}
  dl.kv{display:grid;grid-template-columns:160px 1fr;gap:8px 16px;margin:0;font-size:13px}
  dl.kv dt{color:var(--fg-3)}
  dl.kv dd{margin:0;color:var(--fg);font-family:var(--mono);font-size:12.5px}
  dl.kv dd code{background:var(--surface-2);padding:2px 8px;border-radius:4px;border:1px solid var(--line)}
  dl.kv dd .ok{color:var(--ok)}
  dl.kv dd .bad{color:var(--bad)}
  dl.kv dd .warn{color:var(--warn)}
  table{width:100%;border-collapse:collapse;font-size:12.5px}
  table th{font-size:11px;text-transform:uppercase;letter-spacing:.1em;color:var(--fg-3);
    text-align:left;padding:10px 12px;border-bottom:1px solid var(--line);font-weight:600}
  table td{padding:12px;border-bottom:1px solid var(--line);font-family:var(--mono);font-size:12px;
    color:var(--fg-2)}
  table tr:last-child td{border-bottom:none}
  table .name{color:var(--fg)}
  table .pkg{color:var(--ac-2)}
  table .ok{color:var(--ok)}
  table .bad{color:var(--bad)}
  table .req{font-size:10px;padding:1px 7px;border-radius:3px;
    background:rgba(248,113,113,.08);color:var(--bad);border:1px solid rgba(248,113,113,.3)}
  table .opt{font-size:10px;padding:1px 7px;border-radius:3px;
    background:var(--surface-2);color:var(--fg-3);border:1px solid var(--line-2)}
  table .install-btn{background:rgba(16,185,129,.10);border:1px solid var(--ac);color:var(--ac);
    padding:3px 11px;border-radius:5px;font-family:var(--mono);font-size:11px;cursor:pointer;
    transition:all .15s}
  table .install-btn:hover{background:var(--ac);color:#001a14}
  table .install-btn:disabled{background:transparent;border-color:var(--line-2);color:var(--fg-3);
    cursor:wait}
  /* === Hero & KPI tiles for Reports === */
  .reports-hero-wrap{margin-bottom:32px}
  .hero-eyebrow{display:inline-flex;align-items:center;gap:6px;
    font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:.08em;
    color:var(--ac-2);margin-bottom:12px;padding:5px 12px;border-radius:999px;
    background:var(--ac-bg);border:1px solid var(--ac-line)}
  .hero-eyebrow::before{content:"";width:6px;height:6px;border-radius:50%;
    background:var(--ac-2);box-shadow:0 0 0 2px var(--ac-bg)}
  .reports-hero-wrap h2{font-size:32px;letter-spacing:-.03em;margin:0 0 8px;line-height:1.1;
    background:linear-gradient(135deg,var(--fg) 0%,var(--fg-2) 75%);
    -webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
  .reports-hero-wrap .sub{color:var(--fg-2);font-size:14.5px;line-height:1.55;margin:0 0 24px;max-width:680px}
  .kpi-strip{display:grid;grid-template-columns:repeat(4,1fr);gap:24px;margin-bottom:14px;
    padding:14px 0;border-top:1px solid var(--line);border-bottom:1px solid var(--line)}
  .kpi{background:transparent;border:none;border-radius:0;
    padding:0;position:relative;overflow:visible;
    transition:none}
  .kpi:hover{transform:none;border-color:transparent}
  .kpi::before{display:none}
  .kpi .num{font-size:20px;font-weight:600;letter-spacing:-.02em;color:var(--fg);
    font-variant-numeric:tabular-nums;display:flex;align-items:baseline;gap:6px}
  .kpi .num.ok{color:var(--ok)}
  .kpi .num.bad{color:var(--bad)}
  .kpi .num.warn{color:var(--warn)}
  .kpi .lbl{font-size:11px;color:var(--fg-3);text-transform:uppercase;letter-spacing:.06em;
    margin-top:4px;font-weight:500}
  .kpi .icon-bg{display:none;position:absolute;right:-8px;bottom:-8px;font-size:60px;opacity:.04;line-height:1;
    pointer-events:none}
  /* === Hero & KPI tiles === */
  .settings-hero{margin-bottom:32px}
  .hero-eyebrow{display:inline-flex;align-items:center;gap:6px;
    font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:.08em;
    color:var(--ac-2);margin-bottom:12px;padding:5px 12px;border-radius:999px;
    background:var(--ac-bg);border:1px solid var(--ac-line)}
  .hero-eyebrow::before{content:"";width:6px;height:6px;border-radius:50%;
    background:var(--ac-2);box-shadow:0 0 0 2px var(--ac-bg)}
  .settings-hero h2{font-size:32px;letter-spacing:-.03em;margin:0 0 8px;line-height:1.1;
    background:linear-gradient(135deg,var(--fg) 0%,var(--fg-2) 75%);
    -webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
  .settings-hero .sub{color:var(--fg-2);font-size:14.5px;line-height:1.55;margin:0 0 24px;max-width:680px}
  .kpi-strip{display:grid;grid-template-columns:repeat(4,1fr);gap:24px;margin-bottom:14px;
    padding:14px 0;border-top:1px solid var(--line);border-bottom:1px solid var(--line)}
  .kpi{background:transparent;border:none;border-radius:0;
    padding:0;position:relative;overflow:visible;
    transition:none}
  .kpi:hover{transform:none;border-color:transparent}
  .kpi::before{display:none}
  .kpi .num{font-size:20px;font-weight:600;letter-spacing:-.02em;color:var(--fg);
    font-variant-numeric:tabular-nums;display:flex;align-items:baseline;gap:6px}
  .kpi .num.ok{color:var(--ok)}
  .kpi .num.bad{color:var(--bad)}
  .kpi .num.warn{color:var(--warn)}
  .kpi .lbl{font-size:11px;color:var(--fg-3);text-transform:uppercase;letter-spacing:.06em;
    margin-top:4px;font-weight:500}
  .kpi .icon-bg{display:none;position:absolute;right:-8px;bottom:-8px;font-size:60px;opacity:.04;line-height:1;
    pointer-events:none}
  .install-toast{position:fixed;bottom:24px;right:24px;width:480px;background:var(--surface-2);
    border:1px solid var(--ac);border-radius:10px;padding:14px 18px;z-index:50;
    box-shadow:0 12px 40px rgba(0,0,0,.4);display:none;font-size:12.5px}
  .install-toast.show{display:block}
  .install-toast h4{margin:0 0 6px;font-size:13px;font-weight:600;display:flex;align-items:center;gap:8px}
  .install-toast .cmd{font-family:var(--mono);color:var(--fg-3);font-size:11px;
    padding:6px 8px;background:var(--bg);border-radius:4px;margin:6px 0}
  .install-toast pre{margin:0;font-family:var(--mono);font-size:11px;line-height:1.5;
    color:var(--fg-2);max-height:240px;overflow-y:auto;white-space:pre-wrap;word-break:break-word;
    background:var(--bg);padding:8px 10px;border-radius:4px;border:1px solid var(--line)}
  .install-toast .row{display:flex;align-items:center;gap:8px;margin-top:8px}
  .install-toast button{background:transparent;border:1px solid var(--line-2);color:var(--fg-2);
    padding:4px 12px;border-radius:5px;font-family:var(--mono);font-size:11.5px;cursor:pointer}
  .install-toast button:hover{border-color:var(--ac);color:var(--ac)}
  .install-toast .status-tag{font-family:var(--mono);font-size:10.5px;padding:2px 8px;
    border-radius:3px;font-weight:600;text-transform:uppercase}
  .install-toast .status-tag.queued{background:rgba(168,174,184,.15);color:var(--fg-2)}
  .install-toast .status-tag.running{background:rgba(251,191,36,.15);color:var(--warn)}
  .install-toast .status-tag.succeeded{background:rgba(74,222,128,.15);color:var(--ok)}
  .install-toast .status-tag.failed{background:rgba(248,113,113,.15);color:var(--bad)}
  .pill{display:inline-block;padding:1px 8px;border-radius:3px;font-family:var(--mono);
    font-size:10.5px;font-weight:600;text-transform:uppercase;letter-spacing:.05em;
    background:rgba(74,222,128,.12);color:var(--ok)}
  .pill.bad{background:rgba(248,113,113,.12);color:var(--bad)}
  .pill.warn{background:rgba(251,191,36,.12);color:var(--warn)}
  .empty{color:var(--fg-3);font-size:12.5px;font-family:var(--mono);padding:14px}
  .ovr-row{display:grid;grid-template-columns:1fr auto;gap:14px;padding:11px 14px;
    border-bottom:1px solid var(--line);align-items:center;font-size:12.5px}
  .ovr-row:last-child{border-bottom:none}
  .ovr-row code{font-family:var(--mono);color:var(--fg-2);font-size:12px}
  .ovr-row .reset{background:transparent;border:1px solid var(--line-2);color:var(--fg-2);
    padding:4px 10px;border-radius:5px;font-family:var(--mono);font-size:11px;cursor:pointer}
  .ovr-row .reset:hover{border-color:var(--bad);color:var(--bad)}
  /* === 连接 Claude 三模式纵向卡 === */
  .modes{display:flex;flex-direction:column;gap:10px}
  .mode-card{position:relative;padding:16px 18px 16px 50px;border-radius:10px;
    border:1px solid var(--line);background:var(--surface-2);cursor:pointer;
    transition:border-color .15s,background .15s,transform .12s}
  .mode-card:hover{border-color:var(--line-2);background:var(--surface-3)}
  .mode-card.active{border-color:var(--ac);background:linear-gradient(135deg,rgba(16,185,129,.08),rgba(16,185,129,.02));
    cursor:default}
  .mode-card.active:hover{background:linear-gradient(135deg,rgba(16,185,129,.10),rgba(16,185,129,.03))}
  .mode-card .radio{position:absolute;top:18px;left:18px;width:18px;height:18px;
    border-radius:50%;border:2px solid var(--line-2);background:var(--bg);transition:all .15s}
  .mode-card.active .radio{border-color:var(--ac);background:var(--ac);
    box-shadow:inset 0 0 0 4px var(--bg)}
  .mode-card .head{display:flex;align-items:center;gap:10px;flex-wrap:wrap}
  .mode-card .name{font-size:14.5px;font-weight:600;color:var(--fg);letter-spacing:-.005em}
  .mode-card .badge{font-size:10px;font-weight:600;letter-spacing:.04em;text-transform:uppercase;
    padding:2px 8px;border-radius:999px;background:var(--surface-3);color:var(--fg-3);
    border:1px solid var(--line-2)}
  .mode-card .badge.recommend{background:linear-gradient(135deg,var(--ac),#0ea671);
    color:#001a14;border-color:transparent}
  .mode-card .badge.simple{background:rgba(96,165,250,.12);color:#7eb6ff;
    border-color:rgba(96,165,250,.3)}
  .mode-card .badge.metered{background:rgba(251,191,36,.12);color:var(--warn);
    border-color:rgba(251,191,36,.3)}
  .mode-card .status{margin-left:auto;font-size:11.5px;font-family:var(--mono);
    padding:3px 10px;border-radius:999px;border:1px solid var(--line-2);
    background:var(--surface);color:var(--fg-3);white-space:nowrap}
  .mode-card .status.ok{color:var(--ok);border-color:rgba(52,211,153,.36);
    background:rgba(52,211,153,.10)}
  .mode-card .status.warn{color:var(--warn);border-color:rgba(251,191,36,.36);
    background:rgba(251,191,36,.10)}
  .mode-card .status.bad{color:var(--bad);border-color:rgba(248,113,113,.36);
    background:rgba(248,113,113,.10)}
  .mode-card .desc{font-size:12.5px;color:var(--fg-2);line-height:1.55;margin-top:6px}
  /* 激活后展开的操作区 */
  .mode-card .action-zone{margin-top:12px;padding:12px 14px;border-radius:8px;
    background:var(--surface);border:1px solid var(--line);
    display:none;flex-direction:column;gap:10px}
  .mode-card.active .action-zone{display:flex}
  .mode-card .info-line{font-size:12px;color:var(--fg-2);line-height:1.55}
  .mode-card .info-line code{background:var(--surface-2);padding:1px 7px;border-radius:4px;
    font-size:11px;color:var(--fg);font-family:var(--mono)}
  .mode-card .btn-row{display:flex;gap:8px;align-items:center;flex-wrap:wrap}
  .mode-btn{padding:8px 14px;border-radius:7px;font-size:12.5px;font-weight:500;
    cursor:pointer;border:1px solid var(--line-2);background:transparent;
    color:var(--fg-2);transition:all .12s;white-space:nowrap}
  .mode-btn:hover{background:var(--surface-2);color:var(--fg)}
  .mode-btn.primary{background:var(--ac);color:#001a14;border-color:var(--ac);font-weight:600}
  .mode-btn.primary:hover{background:#0ea671;border-color:#0ea671}
  .mode-btn.danger{color:var(--bad);border-color:rgba(248,113,113,.3)}
  .mode-btn.danger:hover{background:rgba(248,113,113,.12);border-color:var(--bad)}
  .mode-btn:disabled{opacity:.55;cursor:wait}
  .mode-card .api-input{flex:1;padding:8px 12px;background:var(--surface-2);
    border:1px solid var(--line-2);border-radius:6px;color:var(--fg);
    font-family:var(--mono);font-size:12.5px;outline:none;transition:border-color .12s;
    min-width:200px}
  .mode-card .api-input:focus{border-color:var(--ac)}
  .mode-card .checkbox-line{display:flex;align-items:center;gap:5px;
    font-size:11.5px;color:var(--fg-3)}
  /* OAuth 已登录后的账号信息块 */
  .mode-card .account-block{background:linear-gradient(135deg,rgba(16,185,129,.08),rgba(16,185,129,.02));
    border:1px solid var(--ac-line);border-radius:8px;padding:12px 14px;
    display:flex;flex-direction:column;gap:6px;margin-bottom:4px}
  .mode-card .account-row{display:grid;grid-template-columns:80px 1fr;gap:10px;
    font-size:12.5px;line-height:1.5}
  .mode-card .account-row .lbl{color:var(--fg-3);font-size:11.5px;font-family:var(--mono)}
  .mode-card .account-row .val{color:var(--fg);word-break:break-word}
  .mode-card .account-row code{font-size:11px;background:var(--surface-3);padding:1px 6px;
    border-radius:4px}

  /* === 模型接入主区（单按钮 + 状态条） === */
  .access-pane{display:flex;flex-direction:column;gap:14px}
  .access-status{padding:16px 18px;border-radius:10px;border:1px solid var(--line);
    background:var(--surface-2);display:flex;align-items:center;gap:14px}
  .access-status.ready{border-color:var(--ac-line);
    background:linear-gradient(135deg,rgba(16,185,129,.08),rgba(16,185,129,.02))}
  .access-status.unset{border-color:rgba(251,191,36,.32);
    background:linear-gradient(135deg,rgba(251,191,36,.06),rgba(251,191,36,.02))}
  .access-status .icon{font-size:24px;line-height:1;flex-shrink:0}
  .access-status .info{flex:1;min-width:0}
  .access-status .title{font-size:14px;font-weight:600;color:var(--fg);margin-bottom:3px}
  .access-status .sub{font-size:12px;color:var(--fg-3);line-height:1.5;word-break:break-word}
  .access-status code{font-family:var(--mono);font-size:11.5px;background:var(--surface-3);
    padding:1px 7px;border-radius:4px;color:var(--fg-2)}
  .access-btn-row{display:flex;gap:10px;align-items:center}
  .access-btn{padding:11px 22px;border-radius:8px;font-size:13.5px;font-weight:600;
    cursor:pointer;border:1px solid var(--ac);background:var(--ac);color:#001a14;
    transition:all .15s;letter-spacing:-.005em}
  .access-btn:hover{background:#0ea671;border-color:#0ea671;transform:translateY(-1px);
    box-shadow:0 4px 14px rgba(16,185,129,.25)}
  .access-btn.danger{background:transparent;color:var(--bad);border-color:rgba(248,113,113,.35)}
  .access-btn.danger:hover{background:rgba(248,113,113,.10);border-color:var(--bad);
    transform:translateY(-1px);box-shadow:0 4px 14px rgba(248,113,113,.18)}
  .access-btn:disabled{opacity:.55;cursor:wait;transform:none;box-shadow:none}

  /* === 模型接入弹窗 === */
  .auth-modal-backdrop{position:fixed;inset:0;background:rgba(0,0,0,.55);
    backdrop-filter:blur(4px);-webkit-backdrop-filter:blur(4px);
    display:flex;align-items:center;justify-content:center;z-index:9000;
    animation:authModalFade .15s ease-out}
  .auth-modal-backdrop[hidden]{display:none}
  @keyframes authModalFade{from{opacity:0}to{opacity:1}}
  .auth-modal{width:min(540px,92vw);max-height:88vh;background:var(--surface);
    border:1px solid var(--line-2);border-radius:14px;overflow:hidden;
    display:flex;flex-direction:column;
    box-shadow:0 24px 80px rgba(0,0,0,.5);animation:authModalRise .18s ease-out}
  @keyframes authModalRise{from{transform:translateY(12px);opacity:.6}to{transform:none;opacity:1}}
  .auth-modal-head{display:flex;align-items:center;padding:16px 20px;
    border-bottom:1px solid var(--line)}
  .auth-modal-head h4{margin:0;font-size:15px;font-weight:600;color:var(--fg);
    letter-spacing:-.005em}
  .auth-modal-close{margin-left:auto;background:transparent;border:none;color:var(--fg-3);
    font-size:22px;line-height:1;cursor:pointer;padding:2px 8px;border-radius:6px;
    transition:all .12s}
  .auth-modal-close:hover{background:var(--surface-2);color:var(--fg)}
  .auth-modal-body{padding:18px 20px;overflow-y:auto;display:flex;flex-direction:column;gap:12px}

  /* 兼容旧选择器（防 grep 漏） */
  .auth-primary{padding:20px 22px;border-radius:12px;
    background:linear-gradient(135deg,rgba(16,185,129,.08),rgba(16,185,129,.02));
    border:1px solid var(--ac-line);position:relative}
  .auth-primary.ready{border-color:rgba(16,185,129,.45)}
  .auth-primary.unset{background:linear-gradient(135deg,rgba(251,191,36,.06),rgba(251,191,36,.02));
    border-color:rgba(251,191,36,.32)}
  .ap-header{display:flex;align-items:flex-start;gap:14px;margin-bottom:10px}
  .ap-emoji{font-size:26px;line-height:1.1;flex-shrink:0}
  .ap-title-block{flex:1;min-width:0}
  .ap-title{font-size:16px;font-weight:600;color:var(--fg);letter-spacing:-.01em;
    display:flex;align-items:center;gap:10px}
  .ap-badge{font-size:10px;font-weight:700;letter-spacing:.04em;text-transform:uppercase;
    padding:3px 9px;border-radius:999px;
    background:linear-gradient(135deg,var(--ac),#0ea671);color:#001a14;
    box-shadow:0 1px 2px rgba(16,185,129,.4)}
  .ap-subtitle{font-size:12.5px;color:var(--fg-2);line-height:1.55;margin-top:5px}
  .ap-status{flex-shrink:0;font-family:var(--mono);font-size:11.5px;padding:3px 10px;
    border-radius:999px;border:1px solid var(--line-2);background:var(--surface-2);color:var(--fg-3)}
  .ap-status.ok{color:var(--ok);border-color:rgba(52,211,153,.36);background:rgba(52,211,153,.10)}
  .ap-status.warn{color:var(--warn);border-color:rgba(251,191,36,.36);background:rgba(251,191,36,.10)}
  .ap-status.bad{color:var(--bad);border-color:rgba(248,113,113,.36);background:rgba(248,113,113,.10)}
  .ap-body{padding:12px 14px;border-radius:8px;background:var(--surface-2);border:1px solid var(--line);
    font-size:12.5px;color:var(--fg-2);line-height:1.6;margin-bottom:12px}
  .ap-body.empty{display:none}
  .ap-actions{display:flex;gap:10px;align-items:center;flex-wrap:wrap}
  .ap-btn{padding:9px 20px;border-radius:8px;font-size:13px;font-weight:600;
    cursor:pointer;border:none;transition:all .15s;letter-spacing:-.005em}
  .ap-btn.primary{background:var(--ac);color:#001a14;
    box-shadow:0 1px 3px rgba(16,185,129,.3),inset 0 1px 0 rgba(255,255,255,.18)}
  .ap-btn.primary:hover{background:#0ea671;transform:translateY(-1px);
    box-shadow:0 3px 8px rgba(16,185,129,.4),inset 0 1px 0 rgba(255,255,255,.2)}
  .ap-btn.danger{background:transparent;color:var(--bad);
    border:1px solid rgba(248,113,113,.3)}
  .ap-btn.danger:hover{background:rgba(248,113,113,.12);border-color:var(--bad)}
  .ap-btn.ghost{background:transparent;color:var(--fg-2);border:1px solid var(--line-2)}
  .ap-btn.ghost:hover{background:var(--surface-2);color:var(--fg);border-color:var(--line-2)}
  .ap-btn:disabled{opacity:.55;cursor:wait;transform:none!important}

  /* === 高级选项（折叠） === */
  .auth-advanced{margin-top:14px;border:1px solid var(--line);border-radius:10px;
    background:var(--surface);overflow:hidden}
  .auth-advanced summary{padding:12px 16px;cursor:pointer;display:flex;align-items:center;gap:10px;
    font-size:12.5px;color:var(--fg-2);user-select:none;list-style:none;outline:none;
    transition:background .12s}
  .auth-advanced summary::-webkit-details-marker{display:none}
  .auth-advanced summary:hover{background:var(--surface-2)}
  .auth-advanced[open] summary{border-bottom:1px solid var(--line);background:var(--surface-2)}
  .auth-advanced[open] .adv-icon{transform:rotate(90deg)}
  .adv-icon{display:inline-block;color:var(--fg-3);transition:transform .15s}
  .adv-text{font-weight:500;color:var(--fg)}
  .adv-hint{margin-left:auto;font-size:11.5px;color:var(--fg-3);font-family:var(--mono)}
  .adv-body{padding:14px 16px;display:flex;flex-direction:column;gap:12px}
  .adv-card{padding:14px 16px;border-radius:8px;background:var(--surface-2);border:1px solid var(--line)}
  .adv-card-head{display:flex;align-items:center;gap:10px;margin-bottom:8px;flex-wrap:wrap}
  .adv-card-head strong{font-size:13.5px;color:var(--fg)}
  .adv-tag{font-size:10.5px;padding:2px 8px;border-radius:999px;background:var(--surface-3);
    color:var(--fg-3);border:1px solid var(--line-2)}
  .adv-card-state{font-size:11.5px;font-family:var(--mono);color:var(--fg-3)}
  .adv-card-state.ok{color:var(--ok)}
  .adv-desc{margin:0 0 10px;font-size:12px;color:var(--fg-2);line-height:1.55}
  .adv-input-row{display:flex;gap:8px;align-items:center}
  .adv-input-row input[type=password],.adv-input-row input[type=text]{flex:1;padding:8px 12px;
    background:var(--surface-3);border:1px solid var(--line-2);border-radius:6px;
    color:var(--fg);font-family:var(--mono);font-size:12.5px;outline:none;transition:border-color .12s}
  .adv-input-row input:focus{border-color:var(--ac)}
  .adv-btn{padding:8px 14px;border-radius:6px;font-size:12px;font-weight:500;cursor:pointer;
    border:1px solid var(--line-2);background:transparent;color:var(--fg-2);transition:all .12s;white-space:nowrap}
  .adv-btn:hover{background:var(--surface-3);color:var(--fg)}
  .adv-btn.primary{background:var(--ac);color:#001a14;border-color:var(--ac)}
  .adv-btn.primary:hover{background:#0ea671}
  .adv-btn.danger{color:var(--bad);border-color:rgba(248,113,113,.3)}
  .adv-btn.danger:hover{background:rgba(248,113,113,.12)}
  .adv-btn:disabled{opacity:.55;cursor:wait}
  .adv-checkbox{display:flex;align-items:center;gap:5px;font-size:11.5px;color:var(--fg-3)}
  .adv-checkbox input{margin:0}
  .adv-info{padding:11px 14px;border-radius:7px;background:var(--surface-2);
    border:1px solid var(--line);display:flex;gap:10px;align-items:flex-start}
  .adv-info-icon{font-size:16px;line-height:1.2;flex-shrink:0;opacity:.65}
  .adv-info-text{font-size:12px;color:var(--fg-2);line-height:1.6}
  .adv-info-text strong{color:var(--fg)}
  .adv-info-text u{text-decoration:underline;text-decoration-color:rgba(248,113,113,.5);
    text-underline-offset:2px}
</style></head>
<body>
<div class="topbar">
  <a class="brand-link" href="/tools" title="天枢 · 裁决 · 返回主页" style="text-decoration:none;display:inline-flex;align-items:center;gap:10px;font-family:'Noto Serif SC',Georgia,serif;font-size:19px;font-weight:600;letter-spacing:.18em;color:var(--fg);margin-right:24px;padding:4px 0"><svg viewBox="0 0 24 24" fill="currentColor" width="24" height="24" aria-hidden="true" style="color:var(--ac);opacity:1;flex-shrink:0;filter:drop-shadow(0 1px 2px rgba(196,90,58,.28))"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span>天枢</span><span style="color:var(--ac);margin:0 6px;font-weight:400">·</span><span>裁决</span></a>
  
  <nav>
    <a href="/tools">开始</a>
    <a href="/catalog">工具</a>
    <a href="/reports">报告</a>
    <a href="/settings" class="active">设置</a>
    <a href="/usage" class="admin-only" style="display:none">使用统计</a>
    <a href="/admin/users" class="admin-only" style="display:none">用户管理</a>
  </nav>
  <div class="right" style="margin-left:auto">
    <span class="kbd-hint" id="cmd-trigger" title="打开命令面板"
      style="display:inline-flex;align-items:center;gap:4px;font-size:11px;color:var(--fg-3);
        padding:4px 8px;border-radius:6px;cursor:pointer;transition:all .12s">
      <kbd style="background:var(--surface-3);border:1px solid var(--line-1);border-bottom-width:2px;
        padding:1px 6px;border-radius:4px;font-size:10.5px;color:var(--fg-2);
        font-family:'SF Mono',ui-monospace,monospace;min-width:18px;text-align:center">⌘</kbd>
      <kbd style="background:var(--surface-3);border:1px solid var(--line-1);border-bottom-width:2px;
        padding:1px 6px;border-radius:4px;font-size:10.5px;color:var(--fg-2);
        font-family:'SF Mono',ui-monospace,monospace;min-width:18px;text-align:center">K</kbd>
      跳转
    </span>
  </div>
</div>
<main>
  <section class="settings-hero">
    <span class="hero-eyebrow">偏好与环境</span>
    <h2>设置</h2>
    <p class="sub">本地 Claude 接入、默认模型、依赖健康度、提示词覆盖 — 一处管理。</p>
    <div class="kpi-strip">
      <div class="kpi"><div class="num ok" id="kpi-claude">—</div><div class="lbl" id="kpi-claude-lbl">Claude</div><div class="icon-bg">🤖</div></div>
      <div class="kpi"><div class="num" id="kpi-models">—</div><div class="lbl">可用模型</div><div class="icon-bg">⚡</div></div>
      <div class="kpi"><div class="num" id="kpi-env">—</div><div class="lbl">工具就绪</div><div class="icon-bg">🛠</div></div>
      <div class="kpi"><div class="num" id="kpi-overrides">—</div><div class="lbl">提示词覆盖</div><div class="icon-bg">📝</div></div>
    </div>
  </section>

  <div class="sec" id="auth">
    <div class="sec-head">
      <span class="dot" id="auth-dot"></span>
      <h3>模型接入</h3>
      <span style="margin-left:auto;font-size:11.5px;color:var(--fg-3);font-family:var(--mono)" id="auth-mode-current">—</span>
    </div>
    <div class="sec-body" id="auth-body">
      <div style="color:var(--fg-3);font-size:12.5px;padding:6px 0">加载中…</div>
    </div>
  </div>

  <div class="sec" id="sec-defaults" hidden>
    <div class="sec-head">
      <span class="dot ok"></span>
      <h3>默认模型与精度（Claude Code 一致）</h3>
    </div>
    <div class="sec-body" id="defaults">加载中…</div>
  </div>

  <div class="sec">
    <div class="sec-head" id="env-head" style="cursor:pointer;user-select:none">
      <span class="dot ok"></span>
      <h3>每个工具的环境需求</h3>
      <span id="env-chevron" style="margin-left:8px;color:var(--fg-3);font-size:12px;display:inline-block;transition:transform .15s">▸</span>
      <button id="env-refresh" style="margin-left:auto;background:transparent;border:1px solid var(--line-2);color:var(--fg-2);padding:4px 12px;border-radius:5px;font-family:var(--mono);font-size:11px;cursor:pointer">↻ 重新检测</button>
    </div>
    <div id="env-body" hidden>
      <div id="system-info" style="padding:10px 18px;border-bottom:1px solid var(--line);
        font-family:var(--mono);font-size:11.5px;color:var(--fg-3);background:var(--surface-2)"></div>
      <table id="tool-env-table"><thead>
        <tr><th>工具</th><th>包</th><th>必需</th><th>已安装</th><th>版本</th><th>用途</th><th></th></tr>
      </thead><tbody id="tool-env-tbody"></tbody></table>
    </div>
  </div>

  <div class="sec">
    <div class="sec-head">
      <span class="dot warn"></span>
      <h3>提示词覆盖（已激活）</h3>
    </div>
    <div id="overrides"></div>
  </div>
</main>

<div class="install-toast" id="install-toast">
  <h4 id="it-title">安装中…</h4>
  <div class="cmd" id="it-cmd">·</div>
  <div class="row">
    <span class="status-tag" id="it-status">queued</span>
    <span style="font-family:var(--mono);font-size:11px;color:var(--fg-3)" id="it-elapsed"></span>
    <button style="margin-left:auto" onclick="document.getElementById('install-toast').classList.remove('show')">×</button>
  </div>
  <pre id="it-log" style="margin-top:8px"></pre>
</div>

<!-- 模型接入选择弹窗 — 一次渲染，按需 show/hide -->
<div class="auth-modal-backdrop" id="auth-modal" hidden>
  <div class="auth-modal" role="dialog" aria-modal="true">
    <div class="auth-modal-head">
      <h4>选择模型接入方式</h4>
      <button class="auth-modal-close" data-act="close" aria-label="关闭">×</button>
    </div>
    <div class="auth-modal-body" id="auth-modal-body">
      <!-- renderAuthModal() 填充 -->
    </div>
  </div>
</div>

<script>
// 局部 escapeHtml — SETTINGS 页内的 renderAuthSection 等多处依赖。
// 不能依赖 _inject_shared_overlays 注入的版本，因为那个 script 块在本块之后执行，
// 而本块末尾的 load() 调用在它之前；hoisting 跨 script 不共享。
function escapeHtml(s){ return String(s == null ? '' : s).replace(/[&<>"']/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c])); }

function fmtDate(s){
  if (!s) return '—';
  const d = new Date(s);
  return d.toLocaleString('zh-CN', {year:'numeric',month:'2-digit',day:'2-digit',hour:'2-digit',minute:'2-digit'});
}

// ── Figma 设计图接入区块已移除（UI 比对工具已下线）──

async function load(){
  const ci = await fetch('/api/claude/info').then(r=>r.json());
  // 顶部 KPI：Claude 状态 — 状态语义已经迁移到「连接 Claude」section；这里仅作信息汇总
  const k1 = document.getElementById('kpi-claude');
  const k1l = document.getElementById('kpi-claude-lbl');
  if (ci.bin_found && ci.account){
    k1.textContent = '✓'; k1.className = 'num ok'; k1l.textContent = 'CLI 就绪';
  } else if (ci.bin_found){
    k1.textContent = '◔'; k1.className = 'num warn'; k1l.textContent = 'CLI 已找到 · 未登录';
  } else {
    k1.textContent = '✗'; k1.className = 'num bad'; k1l.textContent = '未找到 CLI';
  }
  document.getElementById('kpi-models').textContent = (ci.available_models || []).length;

  document.getElementById('defaults').innerHTML = `
    <dl class="kv">
      <dt>默认精度</dt><dd><code>${ci.settings_effort_level}</code> <span style="color:var(--fg-3)">每次跑工具可单独调整</span></dd>
      <dt>可用模型</dt><dd>${ci.available_models.map(m=>{
        const isDefault = m.default ? ' <span style="color:var(--ac-2);font-size:10px;font-weight:600;text-transform:uppercase;letter-spacing:.05em;margin-left:6px">默认</span>' : '';
        const legacy = m.legacy ? ' <span style="color:var(--fg-3);font-size:10px;text-transform:uppercase;margin-left:6px">旧</span>' : '';
        const key = m.key ? ` <span style="color:var(--fg-4);font-family:'SF Mono',ui-monospace,monospace;font-size:10px;margin-left:6px">${m.key}</span>` : '';
        return `<div style="margin:5px 0;display:flex;align-items:center;gap:4px;flex-wrap:wrap"><code style="font-weight:600">${m.label || m.key}</code>${key}<span style="color:var(--fg-3);font-size:11.5px;margin-left:8px">${m.tag || ''}</span>${isDefault}${legacy}</div>`;
      }).join('')}</dd>
      <dt>可选精度</dt><dd>${ci.available_efforts.map(e=>`<code style="margin-right:6px">${e.key}</code>`).join('')}</dd>
      <dt>扩展思考</dt><dd>${ci.available_thinking.map(t=>`<code style="margin-right:6px">${t.key}</code>`).join('')}</dd>
    </dl>`;

  await renderToolEnv();
  await renderAuthSection();

  const sys = await fetch('/api/settings/system').then(r=>r.json());
  document.getElementById('system-info').innerHTML =
    `OS <code>${sys.os} ${sys.os_release}</code> · arch <code>${sys.machine}</code> · python <code>${sys.python}</code> · venv <code>${sys.venv || '系统'}</code>`;

  // KPI: tools env health (pass count)
  try {
    const tEnv = await fetch('/api/settings/tools').then(r => r.json());
    const total = tEnv.tools.length;
    const okCount = tEnv.tools.filter(t => t.ready).length;
    const kEnv = document.getElementById('kpi-env');
    kEnv.innerHTML = `${okCount}<span class="unit" style="font-size:14px;font-weight:500;color:var(--fg-3)">/${total}</span>`;
    kEnv.className = okCount === total ? 'num ok' : (okCount > 0 ? 'num warn' : 'num bad');
  } catch(e){}

  const ovr = await fetch('/api/settings/overrides').then(r=>r.json());
  const ovrEl = document.getElementById('overrides');
  document.getElementById('kpi-overrides').textContent = (ovr.overrides || []).length;
  if (!ovr.overrides.length){
    ovrEl.innerHTML = '<div class="empty">没有覆盖 · 所有工具使用 configs/prompts/ 中的原版提示词</div>';
  } else {
    ovrEl.innerHTML = '';
    ovr.overrides.forEach(o => {
      const id = o.filename.replace('.md','').replace(/^(\d)_(\d)_.*/, 'step$1.$2');
      const div = document.createElement('div');
      div.className = 'ovr-row';
      div.innerHTML = `
        <div>
          <code>${o.step_dir}/${o.filename}</code>
          <span style="color:var(--fg-3);font-family:var(--mono);font-size:11px;margin-left:10px">${(o.size/1024).toFixed(1)}KB · ${fmtDate(new Date(o.mtime*1000).toISOString())}</span>
        </div>
        <button class="reset" data-step="${o.step_dir}" data-id="${id}">重置回原版</button>`;
      ovrEl.appendChild(div);
    });
    ovrEl.querySelectorAll('button.reset').forEach(b => {
      b.onclick = async () => {
        if (!confirm('确认重置 ' + b.dataset.id + ' 到原版？')) return;
        await fetch(`/api/prompts/${b.dataset.step}/${b.dataset.id}`, {method:'DELETE'});
        load();
      };
    });
  }
}

// === 模型接入 主区 ===
// 状态机：
//   unset      → 显示「+ 模型接入」按钮（点开弹窗 picker）
//   oauth/ok   → 显示账号信息 + 「✕ 模型移除」按钮
//   api_key/ok → 显示 key masked + 「✕ 模型移除」按钮
// OAuth 和 API Key 互斥；同一时间只能选一种。
async function renderAuthSection(){
  let data;
  try { data = await fetch('/api/settings/auth').then(r=>r.json()); }
  catch(e){
    document.getElementById('auth-body').innerHTML =
      `<div style="color:var(--bad);font-size:12.5px">加载失败：${escapeHtml(String(e))}</div>`;
    return;
  }
  // Stash latest snapshot for the modal renderer (避免再发一次请求)
  window.__authSnap = data;

  const cur = data.current_mode;
  const o = data.modes.oauth;
  const ak = data.modes.api_key;
  const oauthReady = !!o.ready;
  const apiReady = !!ak.ready;
  const ready = oauthReady || apiReady;

  // 顶部 dot + tag
  document.getElementById('auth-dot').className = 'dot ' + (ready ? 'ok' : 'warn');
  document.getElementById('auth-mode-current').textContent =
    oauthReady ? 'OAuth · 已接入' :
    apiReady   ? 'API Key · 已接入' :
                 '未接入';

  // 「默认模型与精度」section — 只在接入后显示
  const secDefaults = document.getElementById('sec-defaults');
  if (secDefaults) secDefaults.hidden = !ready;

  // 主区 — 状态条 + 单个按钮
  const body = document.getElementById('auth-body');
  if (oauthReady) {
    const acc = o.account || {};
    const billingMap = {
      'subscription_max'              : 'Claude Max 订阅',
      'subscription_pro'              : 'Claude Pro 订阅',
      'subscription_team'             : 'Claude Team',
      'subscription_enterprise'       : 'Claude Enterprise',
      'google_play_subscription'      : 'Claude Pro 订阅 (Google Play)',
      'apple_app_store_subscription'  : 'Claude Pro 订阅 (App Store)',
      'api'                           : 'API 按量计费',
      'free'                          : '免费版',
    };
    const billingLabel = billingMap[acc.billing_type] || acc.billing_type || '订阅类型未知';
    const sub = acc.email
      ? `账号 <code>${escapeHtml(acc.email)}</code>${acc.billing_type ? ` · 订阅 <strong style="color:var(--ac-2)">${escapeHtml(billingLabel)}</strong>` : ''}${acc.organization_name ? ` · 组织 ${escapeHtml(acc.organization_name)}` : ''}`
      : `已通过 OAuth 接入 · 凭据存于本工具`;
    body.innerHTML = `
      <div class="access-pane">
        <div class="access-status ready">
          <span class="icon">🔗</span>
          <div class="info">
            <div class="title">OAuth 模式已接入</div>
            <div class="sub">${sub}</div>
          </div>
          <div class="access-btn-row">
            <button class="access-btn danger" id="btn-access-remove">✕ 模型移除</button>
          </div>
        </div>
        <div style="font-size:11.5px;color:var(--fg-3);line-height:1.55">
          移除会清掉本工具内的 OAuth 凭据（access_token / refresh_token），不影响其他客户端。想换 API Key 模式，先移除再点接入。
        </div>
      </div>`;
    document.getElementById('btn-access-remove').onclick = () => disconnectAccess('oauth');
    return;
  }
  if (apiReady) {
    body.innerHTML = `
      <div class="access-pane">
        <div class="access-status ready">
          <span class="icon">🔑</span>
          <div class="info">
            <div class="title">API Key 模式已接入</div>
            <div class="sub">当前 Key <code>${escapeHtml(ak.api_key_masked || '')}</code> · 调用按 Anthropic 价目计费</div>
          </div>
          <div class="access-btn-row">
            <button class="access-btn danger" id="btn-access-remove">✕ 模型移除</button>
          </div>
        </div>
        <div style="font-size:11.5px;color:var(--fg-3);line-height:1.55">
          移除会清除本机存储的 API Key；想换 OAuth 订阅模式，先移除再点接入。
        </div>
      </div>`;
    document.getElementById('btn-access-remove').onclick = () => disconnectAccess('api_key');
    return;
  }
  // 未接入
  body.innerHTML = `
    <div class="access-pane">
      <div class="access-status unset">
        <span class="icon">⚠️</span>
        <div class="info">
          <div class="title">尚未接入模型</div>
          <div class="sub">所有工具都依赖 Claude 模型；点右侧按钮选择接入方式（OAuth 订阅或 API Key）。</div>
        </div>
        <div class="access-btn-row">
          <button class="access-btn" id="btn-access-add">+ 模型接入</button>
        </div>
      </div>
    </div>`;
  document.getElementById('btn-access-add').onclick = () => openAccessModal();
}


// === 弹窗 ===
function openAccessModal(){
  renderAccessModal();
  const m = document.getElementById('auth-modal');
  m.hidden = false;
  // 绑定一次性 listener（避免重复绑定）
  if (!m.dataset.bound) {
    m.dataset.bound = '1';
    m.addEventListener('click', e => {
      // 点空白处或关闭按钮关闭
      if (e.target === m || e.target.dataset.act === 'close') closeAccessModal();
    });
    document.addEventListener('keydown', e => {
      if (e.key === 'Escape' && !m.hidden) closeAccessModal();
    });
  }
}

function closeAccessModal(){
  document.getElementById('auth-modal').hidden = true;
}

function renderAccessModal(){
  const data = window.__authSnap;
  if (!data) { document.getElementById('auth-modal-body').innerHTML = '<div style="color:var(--fg-3)">加载中…</div>'; return; }
  const o = data.modes.oauth;
  const ak = data.modes.api_key;
  // OAuth 子卡内容 — 走 web OAuth 流程,不读本机 claude login
  let oauthCard;
  if (o.token_present) {
    oauthCard = `
      <div class="info-line">
        ✓ 本工具已存有 OAuth 凭据。点下方按钮即可激活;或换号请先点「OAuth 授权」重新走流程。
      </div>
      <div class="btn-row">
        <button class="mode-btn primary" data-act="oauth-authorize">↻ 重新授权(换号)</button>
      </div>`;
  } else {
    oauthCard = `
      <div class="info-line">
        在浏览器内完成 Claude 账号授权 — 凭据由本工具保存,不依赖本机 <code>claude login</code>。
        <div style="margin-top:6px;color:var(--fg-3);font-size:11.5px">点击下方按钮 → 新 tab 跳 claude.ai 授权 → 完成后自动回写本工具。</div>
      </div>
      <div class="btn-row">
        <button class="mode-btn primary" data-act="oauth-authorize">↗ OAuth 授权</button>
      </div>`;
  }
  // API Key 子卡内容
  const keyHasStored = ak.has_api_key === true;
  const apiCard = `
    <div class="info-line">
      从 <a href="https://console.anthropic.com/settings/keys" target="_blank" style="color:var(--ac-2)">console.anthropic.com</a> 获取 Key（<code>sk-ant-</code> 开头）。本机明文保存于 <code style="font-size:11px">~/Library/Application Support/AITestToolkit/configs/auth.json</code>（权限 0600）。
    </div>
    <div class="btn-row">
      <input class="api-input" id="modal-api-input" type="password" placeholder="${keyHasStored ? '已存 Key（' + escapeHtml(ak.api_key_masked || '') + '）— 粘贴新 Key 可替换' : 'sk-ant-api03-...'}"
             autocomplete="off" spellcheck="false">
      <label class="checkbox-line"><input id="modal-api-show" type="checkbox">明文</label>
    </div>
    <div class="btn-row">
      ${keyHasStored ? '<button class="mode-btn primary" data-act="use-stored-key">✓ 用已存 Key 接入</button>' : ''}
      <button class="mode-btn ${keyHasStored ? '' : 'primary'}" data-act="save-key">${keyHasStored ? '替换并接入' : '保存并接入'}</button>
    </div>`;

  document.getElementById('auth-modal-body').innerHTML = `
    <div style="font-size:12px;color:var(--fg-3);line-height:1.55;margin-bottom:4px">
      二选一 — 接入后可随时点「模型移除」再换。
    </div>
    <div class="mode-card active" data-mode="oauth">
      <div class="radio"></div>
      <div class="head">
        <span class="name">OAuth 登录</span>
        <span class="badge recommend">推荐 · 含订阅</span>
      </div>
      <div class="desc">${escapeHtml(o.summary)}</div>
      <div class="action-zone" style="display:flex">${oauthCard}</div>
    </div>
    <div class="mode-card active" data-mode="api_key">
      <div class="radio"></div>
      <div class="head">
        <span class="name">API Key</span>
        <span class="badge metered">${escapeHtml(ak.tag || '按量计费')}</span>
      </div>
      <div class="desc">${escapeHtml(ak.summary)}</div>
      <div class="action-zone" style="display:flex">${apiCard}</div>
    </div>`;

  // 绑事件 — 弹窗内按钮
  document.getElementById('auth-modal-body').querySelectorAll('button[data-act]').forEach(btn => {
    btn.onclick = async (e) => {
      e.stopPropagation();
      const act = btn.dataset.act;
      btn.disabled = true;
      try {
        if (act === 'oauth-authorize') {
          await doOAuthAuthorize();
        } else if (act === 'use-stored-key') {
          await switchAuthMode('api_key', null);
          closeAccessModal();
        } else if (act === 'save-key') {
          const inp = document.getElementById('modal-api-input');
          const v = (inp ? inp.value : '').trim();
          if (!v) { alert('请粘贴 API Key（以 sk-ant- 开头）'); btn.disabled = false; return; }
          if (!v.startsWith('sk-') && !confirm('Key 不是以 sk- 开头，确认保存？')) { btn.disabled = false; return; }
          await switchAuthMode('api_key', v);
          closeAccessModal();
        }
      } catch(err) {
        alert('操作失败：' + (err.detail || err.message || JSON.stringify(err)));
      } finally {
        btn.disabled = false;
      }
    };
  });
  const showCk = document.getElementById('modal-api-show');
  if (showCk) showCk.onchange = (e) => {
    const inp = document.getElementById('modal-api-input');
    if (inp) inp.type = e.target.checked ? 'text' : 'password';
  };
}


// Web OAuth flow (OOB 模式):
//   1. POST /api/auth/oauth/start → 拿 authorize_url + state
//   2. window.open() 打开新 tab → 用户在 claude.ai 授权
//   3. Anthropic 把 code 显示在 console.anthropic.com 内部页面，让用户复制
//   4. 用户回 toolkit 弹窗，粘到输入框 → POST /api/auth/oauth/exchange 换 token
async function doOAuthAuthorize(){
  let startResp;
  try {
    const r = await fetch('/api/auth/oauth/start', {method:'POST'});
    if (!r.ok) {
      const err = await r.json().catch(()=>({detail:r.statusText}));
      alert('启动 OAuth 失败：' + (err.detail || JSON.stringify(err)));
      return;
    }
    startResp = await r.json();
  } catch(e) {
    alert('启动 OAuth 失败：' + (e.message || e));
    return;
  }
  const authWin = window.open(startResp.authorize_url, '_blank');
  if (!authWin) {
    alert('浏览器拦截了新 tab。请允许弹窗后重试，或复制下面 URL 手动打开：\n\n' + startResp.authorize_url);
  }
  // 弹窗内换成"粘贴 code"状态
  const body = document.getElementById('auth-modal-body');
  if (!body) return;
  body.innerHTML = `
    <div class="info-line" style="line-height:1.6">
      <div style="font-size:14px;color:var(--fg);margin-bottom:10px"><strong>1.</strong> 在新打开的 tab 完成 Anthropic 账号登录 + 授权</div>
      <div style="font-size:14px;color:var(--fg);margin-bottom:10px"><strong>2.</strong> 授权完成后，Anthropic 会显示一段授权码 — 复制整段</div>
      <div style="font-size:14px;color:var(--fg);margin-bottom:6px"><strong>3.</strong> 把授权码粘贴到下方输入框，点「完成接入」</div>
      <div style="font-size:11.5px;color:var(--fg-3);margin-bottom:4px">
        授权 tab 没自动打开？
        <a href="${startResp.authorize_url}" target="_blank" style="color:var(--ac-2);text-decoration:underline">点这里手动打开</a>
      </div>
    </div>
    <div class="btn-row" style="margin-top:6px">
      <input class="api-input" id="oauth-code-input" type="text"
             placeholder="粘贴授权码（形如 abc-def#state=...）"
             autocomplete="off" spellcheck="false" style="flex:1">
    </div>
    <div class="btn-row" style="margin-top:2px">
      <button class="mode-btn primary" data-act="submit-code">✓ 完成接入</button>
      <button class="mode-btn" data-act="cancel-oauth">取消</button>
    </div>
    <div id="oauth-exchange-msg" style="font-size:12px;color:var(--fg-3);margin-top:6px;min-height:16px"></div>`;

  const inp = document.getElementById('oauth-code-input');
  const msg = document.getElementById('oauth-exchange-msg');
  if (inp) inp.focus();

  const submit = async () => {
    const v = (inp.value || '').trim();
    if (!v) { msg.style.color='var(--warn)'; msg.textContent='请粘贴授权码'; return; }
    msg.style.color='var(--fg-3)'; msg.textContent='换取 access_token…';
    try {
      const r = await fetch('/api/auth/oauth/exchange', {
        method:'POST', headers:{'Content-Type':'application/json'},
        body: JSON.stringify({code:v, state: startResp.state})
      });
      if (r.ok) {
        const d = await r.json();
        msg.style.color='var(--ok)'; msg.textContent='✓ 接入成功';
        toastSettings('✓ OAuth 接入成功' + (d.account && d.account.email ? ` · ${d.account.email}` : ''));
        await new Promise(r=>setTimeout(r,400));
        closeAccessModal();
        await renderAuthSection();
        await load();
      } else {
        const err = await r.json().catch(()=>({detail:r.statusText}));
        msg.style.color='var(--bad)'; msg.textContent='✗ ' + (err.detail || JSON.stringify(err));
      }
    } catch(e) {
      msg.style.color='var(--bad)'; msg.textContent='✗ ' + (e.message || e);
    }
  };

  body.querySelector('button[data-act="submit-code"]').onclick = submit;
  body.querySelector('button[data-act="cancel-oauth"]').onclick = () => renderAccessModal();
  if (inp) inp.addEventListener('keydown', e => { if (e.key==='Enter') submit(); });
}


// 移除（断开）当前模式
async function disconnectAccess(mode){
  const modeLabel = mode === 'oauth' ? 'OAuth' : 'API Key';
  if (!confirm(`确认移除「${modeLabel}」模式？\n\n` +
               (mode === 'oauth'
                 ? '会清除本工具内存的 OAuth 凭据（access_token / refresh_token），不影响其他客户端。'
                 : '会清除本工具存储的 API Key — 想保留可先复制。'))) return;
  try {
    await fetch('/api/settings/auth/disconnect', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({mode, purge: mode === 'api_key'})
    });
    if (mode === 'api_key') {
      // 同时清除存储的 key
      await fetch('/api/settings/auth/api-key', {method:'DELETE'}).catch(()=>{});
    }
    toastSettings('已移除「' + modeLabel + '」模式');
  } catch(e) {
    alert('移除失败：' + (e.message || JSON.stringify(e)));
  }
  await renderAuthSection();
  await load();
}

async function switchAuthMode(mode, apiKey){
  try {
    const r = await fetch('/api/settings/auth', {
      method:'PUT', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({ mode, api_key: apiKey })
    });
    if (!r.ok) {
      const e = await r.json().catch(()=>({detail:r.statusText}));
      alert('切换失败：' + (e.detail || JSON.stringify(e)));
      return;
    }
    await renderAuthSection();
    // 刷新主 KPI 和 Claude 信息卡（loading state 改了）
    await load();
    const labels = {oauth:'OAuth 订阅', api_key:'API Key'};
    toastSettings(`已接入「${labels[mode] || mode}」模式 · 立即可用`);
  } catch(e) {
    alert('切换失败：' + e);
  }
}

function toastSettings(msg){
  // 复用全局 toast 若存在；否则简单 alert
  if (typeof toast === 'function') { toast(msg); return; }
  const t = document.createElement('div');
  t.textContent = msg;
  t.style.cssText = 'position:fixed;bottom:24px;left:50%;transform:translateX(-50%);background:var(--surface-2);border:1px solid var(--ac-line);color:var(--fg);padding:10px 18px;border-radius:8px;font-size:13px;z-index:9999;box-shadow:0 8px 24px rgba(0,0,0,.4)';
  document.body.appendChild(t);
  setTimeout(()=>t.remove(), 2400);
}

async function renderToolEnv(){
  const tEnv = await fetch('/api/settings/tools').then(r=>r.json());
  const tbody = document.getElementById('tool-env-tbody');
  tbody.innerHTML = '';
  tEnv.tools.forEach(t => {
    if (!t.requirements.length){
      tbody.insertAdjacentHTML('beforeend',
        `<tr><td><span class="name">${t.icon} ${t.name}</span></td><td colspan="6" style="color:var(--fg-3)">无额外环境需求</td></tr>`);
      return;
    }
    t.requirements.forEach((r, i) => {
      const ok = r.installed;
      const reqClass = r.required ? 'req' : 'opt';
      const reqLabel = r.required ? '必需' : '可选';
      let action = '';
      if (!ok){
        action = `<button class="install-btn" data-target="${r.pkg}">↓ 安装</button>`;
      } else if (r.pkg === 'playwright' && r.browsers_installed === false){
        action = `<button class="install-btn" data-target="playwright_browsers" title="playwright 已装但缺浏览器">↓ 装浏览器</button>`;
      } else {
        action = '';
      }
      tbody.insertAdjacentHTML('beforeend', `
        <tr>
          <td>${i === 0 ? `<span class="name">${t.icon} ${t.name}</span>` : ''}</td>
          <td><span class="pkg">${r.pkg}</span></td>
          <td><span class="${reqClass}">${reqLabel}</span></td>
          <td><span class="${ok ? 'ok' : 'bad'}">${ok ? '✓' : '✗'}</span>${(r.pkg==='playwright' && ok && r.browsers_installed===false) ? ' <span class="bad" style="font-size:10px">浏览器缺</span>' : ''}</td>
          <td>${r.version || '—'}</td>
          <td style="color:var(--fg-3)">${r.purpose}</td>
          <td>${action}</td>
        </tr>`);
    });
  });
  // wire install buttons
  tbody.querySelectorAll('.install-btn').forEach(btn => {
    btn.onclick = () => installPackage(btn.dataset.target, btn);
  });
}

async function installPackage(target, btn){
  if (!confirm(`确认安装 ${target}？将在当前 venv 内执行 pip 安装，可能需要 1-3 分钟。`)) return;
  btn.disabled = true; btn.textContent = '安装中…';
  const toast = document.getElementById('install-toast');
  document.getElementById('it-title').textContent = `安装 ${target}`;
  document.getElementById('it-cmd').textContent = '提交中…';
  document.getElementById('it-status').className = 'status-tag queued';
  document.getElementById('it-status').textContent = 'queued';
  document.getElementById('it-log').textContent = '';
  document.getElementById('it-elapsed').textContent = '';
  toast.classList.add('show');

  const startedAt = Date.now();
  let timer = null;
  try {
    const j = await fetch('/api/settings/install', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({target}),
    }).then(r=>r.ok ? r.json() : r.json().then(e=>Promise.reject(e)));

    const poll = async () => {
      try {
        const s = await fetch('/api/settings/install/' + j.job_id).then(r=>r.json());
        document.getElementById('it-cmd').textContent = '$ ' + (s.command || '·');
        document.getElementById('it-status').className = 'status-tag ' + s.status;
        document.getElementById('it-status').textContent = s.status;
        document.getElementById('it-log').textContent = s.log || '';
        document.getElementById('it-elapsed').textContent = ((Date.now() - startedAt)/1000).toFixed(1) + 's';
        if (s.status === 'succeeded' || s.status === 'failed'){
          clearInterval(timer);
          btn.disabled = false; btn.textContent = '↓ 安装';
          if (s.status === 'succeeded'){
            await renderToolEnv();
          }
        }
      } catch(e){ clearInterval(timer); }
    };
    poll();
    timer = setInterval(poll, 2000);
  } catch(e){
    document.getElementById('it-status').className = 'status-tag failed';
    document.getElementById('it-status').textContent = 'failed';
    document.getElementById('it-log').textContent = JSON.stringify(e);
    btn.disabled = false; btn.textContent = '↓ 安装';
  }
}

document.getElementById('env-refresh').onclick = (e)=>{ e.stopPropagation(); renderToolEnv(); };
(function initEnvCollapse(){
  const head=document.getElementById('env-head'), body=document.getElementById('env-body'), chev=document.getElementById('env-chevron');
  if(!head||!body) return;
  head.addEventListener('click',()=>{
    const willOpen=body.hasAttribute('hidden');
    if(willOpen){ body.removeAttribute('hidden'); if(chev) chev.style.transform='rotate(90deg)'; }
    else { body.setAttribute('hidden',''); if(chev) chev.style.transform=''; }
  });
})();

load();
</script>
</body></html>
"""


@app.get("/settings", response_class=HTMLResponse)
async def settings_page(request: Request):
    # 「设置」整页仅超级管理员可见;管理员/普通用户访问直接回工具页
    user = getattr(request.state, "current_user", None)
    if not user or not user.is_superadmin():
        return RedirectResponse("/tools", status_code=302)
    return _inject_shared_overlays(SETTINGS_HTML)


# =====================================================================
# Admin · 用户管理页 (仅 admin 可访问)
# =====================================================================

ADMIN_USERS_HTML = r"""<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>用户管理 — 天枢·裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
:root {
  --paper:#ffffff;--paper-2:#ebebeb;--ink:#0a0a0a;--ink-2:#262626;--ink-3:#4a4a4a;--ink-4:#6e6e6e;
  --line:#bdbdbd;--line-2:#9e9e9e;--accent:#a8401f;--accent-h:#82301a;--accent-soft:rgba(168,64,31,.14);
  --bad:#8a2d12;--ok:#4f6b35;--warn:#876b1f;
  --serif:'Noto Serif SC',Georgia,serif;
  --sans:'Noto Sans SC','PingFang SC',-apple-system,sans-serif;
  --mono:'Inter','SF Mono',ui-monospace,Menlo,monospace;
}
*{box-sizing:border-box}
html,body{margin:0;background:var(--paper);color:var(--ink);font-family:var(--sans);
  font-weight:300;font-size:15px;line-height:1.7;-webkit-font-smoothing:antialiased}
a{color:inherit;text-decoration:none}
button{font-family:inherit;cursor:pointer}

.topbar{display:flex;align-items:center;gap:12px;padding:0 24px;height:56px;
  border-bottom:1px solid var(--line);background:#fff;position:sticky;top:0;z-index:10}
.brand-link{display:inline-flex;align-items:center;gap:10px;margin-right:24px;text-decoration:none;color:inherit}
.brand-link svg{color:var(--accent);width:24px;height:24px}
.brand{font-family:var(--serif);font-size:19px;font-weight:600;letter-spacing:.18em;color:var(--ink)}
.brand .sep{color:var(--accent);margin:0 6px;font-weight:400}
.topbar nav{display:flex;gap:2px;margin-left:24px;margin-right:auto;font-family:var(--sans)}
.topbar nav a{font-size:13px;color:var(--ink-2);padding:6px 12px;border-radius:6px;
  text-decoration:none;letter-spacing:.04em;transition:all .15s}
.topbar nav a:hover{background:var(--paper-2,#ebebeb);color:var(--ink)}
.topbar nav a.active{background:var(--accent-soft,rgba(168,64,31,.12));color:var(--accent);font-weight:500}
.spacer{display:none}
.user-chip{font-family:var(--sans);font-size:12.5px;color:var(--ink-2);display:inline-flex;
  align-items:center;gap:8px;font-weight:600}
.user-chip .admin-tag{color:var(--accent);font-size:10.5px;letter-spacing:.18em;font-weight:600}
.logout-btn{background:none;border:1px solid var(--line);color:var(--ink-3);
  padding:4px 14px;border-radius:999px;font-family:var(--sans);font-size:11.5px;letter-spacing:.04em;
  margin-left:4px;transition:all .15s}
.logout-btn:hover{border-color:var(--accent);color:var(--accent)}

.shell{max-width:1080px;margin:0 auto;padding:48px 48px}
.page-head{display:flex;align-items:baseline;justify-content:space-between;
  border-bottom:1px solid var(--line);padding-bottom:18px;margin-bottom:36px}
h1{font-family:var(--serif);font-size:30px;font-weight:500;letter-spacing:.05em;margin:0}
.page-sub{font-size:13px;color:var(--ink-3);font-family:var(--mono);letter-spacing:.06em}

.section{margin-bottom:48px}
.section-title{font-family:var(--mono);font-size:11px;letter-spacing:.22em;color:var(--ink-3);
  text-transform:uppercase;margin-bottom:18px;padding-bottom:8px;border-bottom:1px solid var(--line)}

/* 新建用户表单 */
.create-form{background:#fff;border:1px solid var(--line);border-radius:6px;
  padding:28px 32px;display:grid;grid-template-columns:repeat(4,1fr);gap:18px}
.create-form .field{display:flex;flex-direction:column;gap:6px}
.create-form .field.span2{grid-column:span 2}
.create-form label{font-family:var(--mono);font-size:10.5px;color:var(--ink-3);
  letter-spacing:.18em;text-transform:uppercase}
.create-form input,.create-form select{font-family:var(--sans);font-size:14px;
  padding:9px 12px;border:1px solid var(--line);border-radius:3px;background:#fff;
  color:var(--ink);outline:none;transition:border .15s}
.create-form input:focus,.create-form select:focus{border-color:var(--accent);
  box-shadow:0 0 0 3px var(--accent-soft)}
.create-form .submit-cell{grid-column:span 4;display:flex;justify-content:flex-end;align-items:end;gap:14px}
.create-form .submit-cell .form-err{flex:1;color:var(--bad);font-size:12.5px}
.create-form button{padding:11px 24px;background:var(--ink);color:#fff;border:none;
  border-radius:3px;font-family:var(--serif);font-size:14px;letter-spacing:.22em;
  transition:background .18s}
.create-form button:hover{background:var(--accent)}
.create-form button:disabled{opacity:.5;cursor:not-allowed}

/* 用户表格 */
table.users{width:100%;border-collapse:collapse;background:#fff;
  border:1px solid var(--line);border-radius:6px;overflow:hidden}
table.users th{background:var(--paper-2);font-family:var(--mono);font-size:10.5px;
  color:var(--ink-3);letter-spacing:.18em;text-transform:uppercase;
  text-align:left;padding:14px 18px;font-weight:500;border-bottom:1px solid var(--line)}
table.users td{padding:14px 18px;border-bottom:1px solid var(--line);font-size:14px}
table.users tr:last-child td{border-bottom:none}
table.users tr:hover{background:var(--paper-2)}
.role-tag{font-family:var(--mono);font-size:10.5px;letter-spacing:.16em;padding:2px 8px;
  border-radius:3px;text-transform:uppercase}
.role-tag.admin{color:#fff;background:var(--accent)}
.role-tag.superadmin{color:#fff;background:#b91c1c}
.role-tag.user{color:var(--ink-3);border:1px solid var(--line)}
.role-sel{font-family:var(--mono);font-size:11.5px;padding:3px 6px;border:1px solid var(--line);border-radius:5px;background:#fff;color:var(--ink,#111);cursor:pointer}
.row-actions{display:flex;gap:8px;justify-content:flex-end}
.row-actions button{padding:5px 10px;font-family:var(--mono);font-size:11px;
  border-radius:3px;border:1px solid var(--line);background:#fff;color:var(--ink-2);
  letter-spacing:.04em}
.row-actions button:hover{border-color:var(--accent);color:var(--accent)}
.row-actions button.danger:hover{border-color:var(--bad);color:var(--bad)}
.mono{font-family:var(--mono);font-size:12.5px;color:var(--ink-2)}
.muted{color:var(--ink-3)}
.empty{padding:48px;text-align:center;color:var(--ink-3);font-family:var(--serif);letter-spacing:.04em}

/* Modal */
.modal-overlay{position:fixed;inset:0;background:rgba(26,26,26,.42);display:none;
  align-items:center;justify-content:center;z-index:100;backdrop-filter:blur(2px)}
.modal-overlay.open{display:flex}
.modal{background:#fff;border-radius:6px;padding:32px 36px;width:100%;max-width:420px;
  box-shadow:0 12px 48px rgba(0,0,0,.2)}
.modal h3{font-family:var(--serif);font-size:20px;font-weight:500;margin:0 0 6px}
.modal .modal-sub{color:var(--ink-3);font-size:13px;margin-bottom:20px}
.modal .field{display:flex;flex-direction:column;gap:6px;margin-bottom:14px}
.modal label{font-family:var(--mono);font-size:10.5px;color:var(--ink-3);letter-spacing:.18em;text-transform:uppercase}
.modal input{font-family:var(--sans);font-size:14px;padding:10px 12px;border:1px solid var(--line);
  border-radius:3px;background:#fff;color:var(--ink);outline:none}
.modal input:focus{border-color:var(--accent);box-shadow:0 0 0 3px var(--accent-soft)}
.modal-err{color:var(--bad);font-size:12.5px;margin-top:8px;display:none}
.modal-err.show{display:block}
.modal-actions{display:flex;justify-content:flex-end;gap:12px;margin-top:24px}
.modal-actions button{padding:9px 20px;font-family:var(--mono);font-size:12px;letter-spacing:.1em;
  border-radius:3px;border:1px solid var(--line);background:#fff;color:var(--ink-2)}
.modal-actions button.primary{background:var(--ink);color:#fff;border-color:var(--ink)}
.modal-actions button.primary:hover{background:var(--accent);border-color:var(--accent)}
.toast{position:fixed;top:24px;right:24px;background:var(--ink);color:#fff;padding:12px 20px;
  border-radius:4px;font-size:13px;z-index:200;opacity:0;transform:translateY(-8px);
  transition:opacity .2s,transform .2s;font-family:var(--mono);letter-spacing:.04em}
.toast.show{opacity:1;transform:translateY(0)}
.toast.bad{background:var(--bad)}
.toast.ok{background:var(--ok)}
</style>
</head><body>
<header class="topbar">
  <a class="brand-link" href="/tools">
    <svg viewBox="0 0 24 24" fill="currentColor" width="26" height="26"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg>
    <span class="brand">天枢<span class="sep">·</span>裁决</span>
  </a>
  <nav>
    <a href="/tools">开始</a>
    <a href="/catalog">工具</a>
    <a href="/reports">报告</a>
    <a href="/settings" class="superadmin-only" style="display:none">设置</a>
    <a href="/usage" class="admin-only" style="display:none">使用统计</a>
    <a href="/admin/users" class="active">用户管理</a>
  </nav>
  <div class="spacer"></div>
  <div class="user-chip" id="user-chip"></div>
  <button class="logout-btn" id="logout-btn">登出</button>
</header>

<main class="shell">
  <div class="page-head">
    <h1>用户管理</h1>
    <div class="page-sub" id="user-count">— 位用户</div>
  </div>

  <!-- 新建用户 -->
  <section class="section">
    <div class="section-title">
      <span>新建用户</span>
      <button type="button" id="open-bulk-btn"
        style="float:right;font-family:var(--mono);font-size:11px;letter-spacing:.12em;
               padding:5px 14px;border:1px solid var(--line-2);background:#fff;
               color:var(--ink-2);border-radius:3px;cursor:pointer;transition:all .15s">
        批量创建 →
      </button>
    </div>
    <form class="create-form" id="create-form" autocomplete="off">
      <div class="field">
        <label>用户名</label>
        <input id="c-username" required placeholder="如 lisi" autocomplete="off">
      </div>
      <div class="field">
        <label>初始密码</label>
        <input id="c-password" type="text" required placeholder="至少 6 位" autocomplete="off">
      </div>
      <div class="field">
        <label>显示名 (可选)</label>
        <input id="c-display" placeholder="留空则同用户名" autocomplete="off">
      </div>
      <div class="field">
        <label>角色</label>
        <select id="c-role">
          <option value="user" selected>普通用户</option>
          <option value="admin">管理员</option>
          <option value="superadmin">超级管理员</option>
        </select>
      </div>
      <div class="submit-cell">
        <span class="form-err" id="form-err"></span>
        <button type="submit" id="create-btn">创 建 用 户</button>
      </div>
    </form>
  </section>

  <!-- 用户列表 -->
  <section class="section">
    <div class="section-title">用户列表</div>
    <div id="list-wrap">加载中…</div>
  </section>
</main>

<!-- 重置密码 modal -->
<div class="modal-overlay" id="reset-modal">
  <div class="modal">
    <h3 id="reset-title">重置密码</h3>
    <div class="modal-sub" id="reset-sub"></div>
    <div class="field">
      <label>新密码</label>
      <input id="reset-pwd" type="text" placeholder="至少 6 位">
    </div>
    <div class="modal-err" id="reset-err"></div>
    <div class="modal-actions">
      <button id="reset-cancel">取消</button>
      <button class="primary" id="reset-ok">确认重置</button>
    </div>
  </div>
</div>

<!-- 批量创建 modal — 多行可叠加表单 -->
<style>
.bulk-modal{max-width:900px;max-height:90vh;overflow-y:auto}
.bulk-modal .bulk-tip{background:var(--paper-2);border-left:3px solid var(--accent);
  padding:10px 14px;font-size:12.5px;color:var(--ink-2);line-height:1.7;margin:6px 0 14px;
  border-radius:0 4px 4px 0}
.bulk-modal .bulk-defaults{display:flex;align-items:center;gap:14px;margin:0 0 14px;
  padding:10px 14px;background:var(--paper-2);border-radius:4px;border:1px solid var(--line)}
.bulk-modal .bulk-defaults label{font-family:var(--mono);font-size:11px;color:var(--ink-3);
  letter-spacing:.14em;text-transform:uppercase}
.bulk-modal .bulk-defaults select{font-family:var(--sans);font-size:13px;padding:5px 10px;
  border:1px solid var(--line);border-radius:3px;background:#fff;color:var(--ink)}

/* 行容器 */
.bulk-rows{display:flex;flex-direction:column;gap:8px;max-height:380px;overflow-y:auto;
  padding:2px;margin-bottom:10px}
.bulk-row-item{display:grid;
  grid-template-columns:30px minmax(0,1.4fr) minmax(0,1.4fr) minmax(0,1fr) 110px 28px;
  gap:10px;align-items:center;padding:8px;border-radius:4px;
  transition:background .12s;border:1px solid transparent}
.bulk-row-item:hover{background:var(--paper-2);border-color:var(--line)}
.bulk-row-item .row-num{font-family:var(--mono);font-size:11px;color:var(--ink-3);
  text-align:center;letter-spacing:.1em}
.bulk-row-item input,.bulk-row-item select{font-family:var(--sans);font-size:13.5px;
  padding:8px 10px;border:1px solid var(--line);border-radius:3px;background:#fff;
  color:var(--ink);outline:none;transition:border .15s;width:100%}
.bulk-row-item input:focus,.bulk-row-item select:focus{
  border-color:var(--accent);box-shadow:0 0 0 2px var(--accent-soft)}
.bulk-row-item input.has-error{border-color:var(--bad);background:rgba(138,45,18,.04)}
.bulk-row-item .del-row{width:28px;height:28px;border:1px solid var(--line);
  background:#fff;color:var(--ink-3);border-radius:3px;cursor:pointer;
  font-size:14px;line-height:1;padding:0;display:flex;align-items:center;justify-content:center;
  transition:all .15s}
.bulk-row-item .del-row:hover{border-color:var(--bad);color:var(--bad)}
.bulk-row-item .del-row:disabled{opacity:.3;cursor:not-allowed}
.bulk-row-headers{display:grid;
  grid-template-columns:30px minmax(0,1.4fr) minmax(0,1.4fr) minmax(0,1fr) 110px 28px;
  gap:10px;padding:0 8px 6px;border-bottom:1px solid var(--line);margin-bottom:6px}
.bulk-row-headers div{font-family:var(--mono);font-size:10.5px;color:var(--ink-3);
  letter-spacing:.16em;text-transform:uppercase}

.add-row-btn{display:flex;align-items:center;justify-content:center;gap:8px;
  width:100%;padding:10px;border:1px dashed var(--line-2);background:transparent;
  color:var(--ink-2);border-radius:4px;cursor:pointer;font-family:var(--mono);
  font-size:12px;letter-spacing:.1em;transition:all .15s;margin-bottom:8px}
.add-row-btn:hover{border-color:var(--accent);color:var(--accent);background:var(--accent-soft)}
.add-row-btn .plus{font-size:14px}

/* 结果 */
.bulk-results{margin-top:20px;border-top:1px solid var(--line);padding-top:18px}
.bulk-results .bulk-summary{display:flex;gap:18px;font-family:var(--mono);font-size:12px;
  margin-bottom:14px;color:var(--ink-2)}
.bulk-results .bulk-summary .ok{color:var(--ok);font-weight:600}
.bulk-results .bulk-summary .fail{color:var(--bad);font-weight:600}
.bulk-results table{width:100%;border-collapse:collapse;font-size:12.5px;margin-top:8px}
.bulk-results th{background:var(--paper-2);font-family:var(--mono);font-size:10.5px;
  letter-spacing:.16em;color:var(--ink-3);text-transform:uppercase;
  text-align:left;padding:8px 10px;border-bottom:1px solid var(--line);font-weight:500}
.bulk-results td{padding:8px 10px;border-bottom:1px solid var(--line)}
.bulk-results .pwd-cell{font-family:var(--mono);font-size:12px;color:var(--accent);
  font-weight:600;letter-spacing:.04em}
.bulk-results .fail-cell{color:var(--bad);font-size:12px}
.bulk-export-row{display:flex;gap:10px;margin-top:14px;flex-wrap:wrap;align-items:center}
.bulk-export-row button{padding:7px 14px;font-family:var(--mono);font-size:11.5px;
  letter-spacing:.08em;border:1px solid var(--line);background:#fff;color:var(--ink-2);
  border-radius:3px;cursor:pointer;transition:all .15s}
.bulk-export-row button:hover{border-color:var(--accent);color:var(--accent)}
</style>
<div class="modal-overlay" id="bulk-modal">
  <div class="modal bulk-modal">
    <h3>批量创建用户</h3>
    <div class="modal-sub">每行一个用户。失败的不会阻塞其他行,完成后可下载 CSV 把账号密码发给员工。</div>

    <div class="bulk-tip">
      ① 密码留空 → 按下方「默认密码长度」自动生成 &nbsp;·&nbsp;
      ② 角色留空 → 用「默认角色」&nbsp;·&nbsp;
      ③ 用户名重复会失败但不阻塞其他行
    </div>

    <div class="bulk-defaults">
      <label>默认角色</label>
      <select id="bulk-default-role">
        <option value="user" selected>普通用户</option>
        <option value="admin">管理员</option>
      </select>
      <label style="margin-left:auto">自动密码长度</label>
      <select id="bulk-pwd-length">
        <option value="8">8 位</option>
        <option value="10" selected>10 位</option>
        <option value="12">12 位</option>
        <option value="16">16 位</option>
      </select>
    </div>

    <!-- 表头 -->
    <div class="bulk-row-headers">
      <div>#</div>
      <div>用户名 *</div>
      <div>初始密码</div>
      <div>显示名</div>
      <div>角色</div>
      <div></div>
    </div>

    <!-- 动态行容器 -->
    <div class="bulk-rows" id="bulk-rows"></div>

    <!-- 加行按钮 -->
    <button type="button" class="add-row-btn" id="add-row-btn">
      <span class="plus">+</span><span>添加一行</span>
    </button>

    <div class="modal-err" id="bulk-err"></div>

    <div class="modal-actions">
      <button id="bulk-cancel">取消</button>
      <button class="primary" id="bulk-submit">开 始 创 建</button>
    </div>

    <div class="bulk-results" id="bulk-results" style="display:none"></div>
  </div>
</div>

<div class="toast" id="toast"></div>

<script>
let currentUser = null;
let users = [];

function fmtTs(t){
  if (!t) return '—';
  const d = new Date(t * 1000);
  return d.toLocaleString('zh-CN',{hour12:false}).replace(/\//g,'-');
}

function toast(msg, kind){
  const t = document.getElementById('toast');
  t.textContent = msg;
  t.className = 'toast show' + (kind ? ' ' + kind : '');
  setTimeout(() => { t.className = 'toast'; }, 2500);
}

async function fetchMe(){
  const r = await fetch('/api/auth/me').then(r=>r.json());
  if (!r.authenticated){ location.href = '/login'; return; }
  if (r.user.role !== 'admin' && r.user.role !== 'superadmin'){ alert('需要管理员权限'); location.href = '/tools'; return; }
  currentUser = r.user;
  const isSuper = r.user.role === 'superadmin';
  const tag = isSuper ? '<span class="admin-tag">[超管]</span>' : '<span class="admin-tag">[admin]</span>';
  document.getElementById('user-chip').innerHTML = tag + (r.user.display_name || r.user.username);
  // 本页不经 _inject_shared_overlays,手动按角色显隐顶部导航链接
  document.querySelectorAll('.admin-only').forEach(el => { el.style.display = ''; });
  if (isSuper){ document.querySelectorAll('.superadmin-only').forEach(el => { el.style.display = ''; }); }
  // 非超管(管理员):创建用户只能选「普通用户」
  if (!isSuper){
    const sel = document.getElementById('c-role');
    if (sel) sel.innerHTML = '<option value="user" selected>普通用户</option>';
  }
}

async function loadUsers(){
  try {
    const r = await fetch('/api/auth/users').then(r=>r.json());
    users = r.users || [];
    document.getElementById('user-count').textContent = users.length + ' 位用户';
    renderUsers();
  } catch(e){ toast('加载用户失败:' + e.message, 'bad'); }
}

function renderUsers(){
  const w = document.getElementById('list-wrap');
  if (!users.length){
    w.innerHTML = '<div class="empty">暂无用户</div>';
    return;
  }
  const amSuper = currentUser && currentUser.role === 'superadmin';
  const roleLabel = {user:'USER', admin:'ADMIN', superadmin:'SUPER'};
  const rows = users.map(u => {
    const isMe = currentUser && u.id === currentUser.id;
    const isLastSuper = u.role === 'superadmin' && users.filter(x => x.role === 'superadmin').length === 1;
    // 非超管只能操作普通用户
    const canManage = amSuper || u.role === 'user';
    const delDisabled = isMe || isLastSuper || !canManage;
    const delTitle = isMe ? '不能删除自己'
      : (isLastSuper ? '不能删除最后一个超级管理员'
      : (!canManage ? '只有超级管理员能操作管理员/超管' : '删除用户'));
    const manageTitle = !canManage ? '只有超级管理员能操作管理员/超管' : '';
    // 角色:超管用下拉直接改;其余只读标签
    const roleCell = amSuper
      ? `<select class="role-sel" onchange="changeRole(${u.id}, this.value, '${escapeHtml(u.username)}')"${isLastSuper ? ' title="最后一个超管不能降级"' : ''}>
           <option value="user"${u.role==='user'?' selected':''}>普通用户</option>
           <option value="admin"${u.role==='admin'?' selected':''}>管理员</option>
           <option value="superadmin"${u.role==='superadmin'?' selected':''}>超级管理员</option>
         </select>`
      : `<span class="role-tag ${u.role}">${roleLabel[u.role] || 'USER'}</span>`;
    return `<tr>
      <td class="mono">${u.id}</td>
      <td><strong>${escapeHtml(u.username)}</strong></td>
      <td class="muted">${escapeHtml(u.display_name || '—')}</td>
      <td>${roleCell}</td>
      <td class="mono muted">${fmtTs(u.created_at)}</td>
      <td class="mono muted">${fmtTs(u.last_login_at)}</td>
      <td>
        <div class="row-actions">
          <button onclick="openReset(${u.id}, '${escapeHtml(u.username)}')"
                  ${!canManage ? `disabled style="opacity:.4;cursor:not-allowed" title="${manageTitle}"` : ''}>重置密码</button>
          <button class="danger" onclick="deleteUser(${u.id}, '${escapeHtml(u.username)}')"
                  ${delDisabled ? 'disabled style="opacity:.4;cursor:not-allowed"' : ''}
                  title="${delTitle}">删除</button>
        </div>
      </td>
    </tr>`;
  }).join('');
  w.innerHTML = `<table class="users">
    <thead><tr>
      <th>ID</th><th>用户名</th><th>显示名</th><th>角色</th>
      <th>创建</th><th>最近登录</th><th style="text-align:right">操作</th>
    </tr></thead>
    <tbody>${rows}</tbody>
  </table>`;
}

function escapeHtml(s){
  return String(s == null ? '' : s).replace(/[&<>"']/g, c =>
    ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c]));
}

document.getElementById('create-form').onsubmit = async (e) => {
  e.preventDefault();
  const err = document.getElementById('form-err');
  err.textContent = '';
  const btn = document.getElementById('create-btn');
  btn.disabled = true;
  const body = {
    username: document.getElementById('c-username').value.trim(),
    password: document.getElementById('c-password').value,
    display_name: document.getElementById('c-display').value.trim(),
    role: document.getElementById('c-role').value,
  };
  try {
    const r = await fetch('/api/auth/users', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify(body),
    });
    if (!r.ok){
      const d = await r.json().catch(()=>({}));
      err.textContent = d.detail || '创建失败';
      btn.disabled = false;
      return;
    }
    const d = await r.json();
    toast('✓ 已创建 ' + d.user.username, 'ok');
    document.getElementById('create-form').reset();
    btn.disabled = false;
    await loadUsers();
  } catch(e){
    err.textContent = e.message;
    btn.disabled = false;
  }
};

let resetTargetId = null;
function openReset(id, username){
  resetTargetId = id;
  document.getElementById('reset-title').textContent = '重置密码 · ' + username;
  document.getElementById('reset-sub').textContent =
    '不需要原密码,生效后该用户必须用新密码重新登录。';
  document.getElementById('reset-pwd').value = '';
  document.getElementById('reset-err').classList.remove('show');
  document.getElementById('reset-modal').classList.add('open');
  setTimeout(() => document.getElementById('reset-pwd').focus(), 50);
}
document.getElementById('reset-cancel').onclick = () => {
  document.getElementById('reset-modal').classList.remove('open');
};
document.getElementById('reset-ok').onclick = async () => {
  const pwd = document.getElementById('reset-pwd').value;
  const err = document.getElementById('reset-err');
  if (pwd.length < 6){
    err.textContent = '密码至少 6 位';
    err.classList.add('show');
    return;
  }
  try {
    const r = await fetch('/api/auth/users/' + resetTargetId + '/reset_password', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({new_password: pwd}),
    });
    if (!r.ok){
      const d = await r.json().catch(()=>({}));
      err.textContent = d.detail || '重置失败';
      err.classList.add('show');
      return;
    }
    toast('✓ 密码已更新', 'ok');
    document.getElementById('reset-modal').classList.remove('open');
    await loadUsers();
  } catch(e){
    err.textContent = e.message;
    err.classList.add('show');
  }
};

async function changeRole(id, role, username){
  const label = {user:'普通用户', admin:'管理员', superadmin:'超级管理员'}[role] || role;
  if (!confirm('确认把「' + username + '」改为「' + label + '」?\n改后该用户需重新登录才生效。')){ await loadUsers(); return; }
  try {
    const r = await fetch('/api/auth/users/' + id + '/role', {
      method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({role})
    });
    const d = await r.json().catch(()=>({}));
    if (!r.ok){ toast('改角色失败:' + (d.detail || r.status), 'bad'); await loadUsers(); return; }
    toast('✓ ' + (d.msg || ('已将 ' + username + ' 设为 ' + label)), 'ok');
    await loadUsers();
  } catch(e){ toast('改角色请求失败:' + e.message, 'bad'); await loadUsers(); }
}

async function deleteUser(id, username){
  if (!confirm('确认删除 ' + username + '?\n该用户的所有报告归属也会留在系统里,但他无法再登录。')) return;
  try {
    const r = await fetch('/api/auth/users/' + id, {method:'DELETE'});
    if (!r.ok){
      const d = await r.json().catch(()=>({}));
      toast('删除失败:' + (d.detail || r.status), 'bad');
      return;
    }
    toast('✓ 已删除 ' + username, 'ok');
    await loadUsers();
  } catch(e){ toast('删除请求失败:' + e.message, 'bad'); }
}

document.getElementById('logout-btn').onclick = async () => {
  await fetch('/api/auth/logout', {method:'POST'});
  location.href = '/login';
};

// ============== 批量创建(行式表单) ==============
let bulkRowCounter = 0;

function makeBulkRow(){
  bulkRowCounter += 1;
  const id = 'br-' + bulkRowCounter;
  const div = document.createElement('div');
  div.className = 'bulk-row-item';
  div.dataset.rid = String(bulkRowCounter);
  div.innerHTML = `
    <div class="row-num"></div>
    <input class="bulk-username" type="text" placeholder="必填,如 lisi" autocomplete="off">
    <input class="bulk-password" type="text" placeholder="留空自动生成" autocomplete="off">
    <input class="bulk-display"  type="text" placeholder="留空同用户名" autocomplete="off">
    <select class="bulk-role">
      <option value="">默认</option>
      <option value="user">普通用户</option>
      <option value="admin">管理员</option>
    </select>
    <button class="del-row" type="button" title="删除此行">×</button>
  `;
  div.querySelector('.del-row').onclick = () => {
    const rows = document.querySelectorAll('#bulk-rows .bulk-row-item');
    if (rows.length <= 1){
      // 最后一行只清空,不删
      div.querySelectorAll('input').forEach(i => i.value = '');
      div.querySelector('select').selectedIndex = 0;
      div.querySelectorAll('input').forEach(i => i.classList.remove('has-error'));
    } else {
      div.remove();
      renumberRows();
    }
  };
  // 按 Tab 在最后一格按下时自动加新行
  div.querySelector('.bulk-role').addEventListener('keydown', e => {
    if (e.key === 'Tab' && !e.shiftKey){
      const all = [...document.querySelectorAll('#bulk-rows .bulk-row-item')];
      if (all[all.length-1] === div){
        // 是最后一行 → 加新行
        e.preventDefault();
        addBulkRow();
      }
    }
  });
  // 按下 Enter 也加一行
  div.querySelectorAll('input').forEach(inp => {
    inp.addEventListener('keydown', e => {
      if (e.key === 'Enter'){
        e.preventDefault();
        const all = [...document.querySelectorAll('#bulk-rows .bulk-row-item')];
        if (all[all.length-1] === div) addBulkRow();
        else {
          const next = all[all.indexOf(div)+1];
          next.querySelector('.bulk-username').focus();
        }
      }
    });
  });
  return div;
}

function addBulkRow(){
  const row = makeBulkRow();
  document.getElementById('bulk-rows').appendChild(row);
  renumberRows();
  row.querySelector('.bulk-username').focus();
  return row;
}

function renumberRows(){
  document.querySelectorAll('#bulk-rows .bulk-row-item').forEach((el, idx) => {
    el.querySelector('.row-num').textContent = String(idx + 1);
  });
}

function resetBulkModal(){
  document.getElementById('bulk-rows').innerHTML = '';
  // 默认开 3 行 — 大多数批量场景至少几个用户
  for (let i = 0; i < 3; i++) addBulkRow();
  document.getElementById('bulk-err').classList.remove('show');
  document.getElementById('bulk-results').style.display = 'none';
  document.getElementById('bulk-results').innerHTML = '';
  const sb = document.getElementById('bulk-submit');
  sb.disabled = false;
  sb.textContent = '开 始 创 建';
  sb.onclick = submitBulk;  // 上一次完成后被改成 closeBulkAndReset,重新打开要恢复
  document.getElementById('bulk-cancel').textContent = '取 消';
}

function openBulk(){
  resetBulkModal();
  document.getElementById('bulk-modal').classList.add('open');
  setTimeout(() => {
    const first = document.querySelector('#bulk-rows .bulk-username');
    if (first) first.focus();
  }, 50);
}

document.getElementById('open-bulk-btn').onclick = openBulk;
document.getElementById('add-row-btn').onclick = addBulkRow;
document.getElementById('bulk-cancel').onclick = () => {
  document.getElementById('bulk-modal').classList.remove('open');
  loadUsers();
};

function collectBulkRows(){
  // 把表单里的行收集成提交 payload + 本地校验
  const items = [];
  const localErrors = [];
  const rows = [...document.querySelectorAll('#bulk-rows .bulk-row-item')];
  rows.forEach((row, idx) => {
    const unInput = row.querySelector('.bulk-username');
    const pwInput = row.querySelector('.bulk-password');
    const dnInput = row.querySelector('.bulk-display');
    const roleSel = row.querySelector('.bulk-role');
    const un = unInput.value.trim();
    const pw = pwInput.value;
    const dn = dnInput.value.trim();
    const role = roleSel.value || null;
    // 清掉旧的错误样式
    [unInput, pwInput].forEach(i => i.classList.remove('has-error'));
    // 整行全空 — 跳过(允许用户留空白行)
    if (!un && !pw && !dn && !role) return;
    if (!un){
      unInput.classList.add('has-error');
      localErrors.push({line: idx+1, username: '(空)', error: '用户名为空'});
      return;
    }
    if (pw && pw.length < 6){
      pwInput.classList.add('has-error');
      localErrors.push({line: idx+1, username: un, error: '密码至少 6 位 (留空可自动生成)'});
      return;
    }
    const item = {username: un};
    if (pw) item.password = pw;
    if (dn) item.display_name = dn;
    if (role) item.role = role;
    items.push(item);
  });
  return {items, localErrors};
}

function closeBulkAndReset(){
  // 关掉模态 + 复原按钮状态,让下次打开还是干净的"开始创建"
  document.getElementById('bulk-modal').classList.remove('open');
  const btn = document.getElementById('bulk-submit');
  btn.onclick = submitBulk;
  btn.disabled = false;
  btn.textContent = '开 始 创 建';
  document.getElementById('bulk-cancel').textContent = '取 消';
  loadUsers();
}

async function submitBulk(){
  const err = document.getElementById('bulk-err');
  err.classList.remove('show');
  const {items, localErrors} = collectBulkRows();
  if (!items.length){
    err.textContent = localErrors.length
      ? '请修正红色标出的字段后再提交'
      : '至少填一行用户名';
    err.classList.add('show'); return;
  }
  if (localErrors.length){
    err.textContent = `有 ${localErrors.length} 行格式不合法 — 已标红,请修正后重试`;
    err.classList.add('show'); return;
  }
  const btn = document.getElementById('bulk-submit');
  btn.disabled = true; btn.textContent = '创建中…';
  try {
    const r = await fetch('/api/auth/users/bulk', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({
        users: items,
        default_role: document.getElementById('bulk-default-role').value,
        default_password_length: parseInt(document.getElementById('bulk-pwd-length').value, 10) || 10,
      }),
    });
    if (!r.ok){
      const d = await r.json().catch(()=>({}));
      err.textContent = d.detail || '批量创建失败';
      err.classList.add('show');
      btn.disabled = false; btn.textContent = '开 始 创 建';
      return;
    }
    const d = await r.json();
    renderBulkResults(d.created || [], d.failed || []);
    // 完成态:换文案 + 改 onclick 为"关闭模态"
    btn.textContent = '完 成';
    btn.disabled = false;
    document.getElementById('bulk-cancel').textContent = '关 闭';
    btn.onclick = closeBulkAndReset;
    await loadUsers();
  } catch(e){
    err.textContent = '请求出错:' + e.message;
    err.classList.add('show');
    btn.disabled = false; btn.textContent = '开 始 创 建';
  }
}

document.getElementById('bulk-submit').onclick = submitBulk;

function renderBulkResults(created, failed){
  const wrap = document.getElementById('bulk-results');
  let html = `<div class="bulk-summary">
    <span>合计 ${created.length + failed.length}</span>
    <span class="ok">✓ 成功 ${created.length}</span>
    <span class="fail">✕ 失败 ${failed.length}</span>
  </div>`;
  if (created.length){
    html += '<table><thead><tr><th>用户名</th><th>密码 (一次性显示)</th><th>角色</th></tr></thead><tbody>';
    created.forEach(u => {
      const role = u.role === 'superadmin' ? '<span class="role-tag superadmin">SUPER</span>'
        : (u.role === 'admin' ? '<span class="role-tag admin">ADMIN</span>' : '<span class="role-tag user">USER</span>');
      html += `<tr>
        <td><strong>${escapeHtml(u.username)}</strong></td>
        <td class="pwd-cell">${escapeHtml(u.password)}${u.password_auto_generated ? ' <span style="color:var(--ink-3);font-weight:400">(自动)</span>' : ''}</td>
        <td>${role}</td>
      </tr>`;
    });
    html += '</tbody></table>';
    html += '<div class="bulk-export-row">';
    html += '<button id="bulk-copy-csv">复制 CSV</button>';
    html += '<button id="bulk-download-csv">下载 CSV</button>';
    html += '<span style="margin-left:auto;font-size:11.5px;color:var(--ink-3);align-self:center">' +
            '⚠ 密码只在此次响应中显示,关掉窗口就再也拿不到原文 — 请立刻保存' +
            '</span>';
    html += '</div>';
  }
  if (failed.length){
    html += '<div style="margin-top:18px"><div class="section-title" style="margin-bottom:8px">失败明细</div>';
    html += '<table><thead><tr><th>用户名</th><th>失败原因</th></tr></thead><tbody>';
    failed.forEach(f => {
      html += `<tr>
        <td>${escapeHtml(f.username)}</td>
        <td class="fail-cell">${escapeHtml(f.error)}</td>
      </tr>`;
    });
    html += '</tbody></table></div>';
  }
  wrap.innerHTML = html;
  wrap.style.display = 'block';

  // 复制 / 下载 CSV
  if (created.length){
    const csv = ['username,password,display_name,role']
      .concat(created.map(u =>
        [u.username, u.password, u.display_name || '', u.role]
          .map(s => /[,"\n]/.test(String(s)) ? '"' + String(s).replace(/"/g,'""') + '"' : s)
          .join(',')
      )).join('\n');
    document.getElementById('bulk-copy-csv').onclick = async () => {
      try {
        await navigator.clipboard.writeText(csv);
        toast('✓ 已复制到剪贴板', 'ok');
      } catch(e){ toast('复制失败 — 浏览器拦截了剪贴板', 'bad'); }
    };
    document.getElementById('bulk-download-csv').onclick = () => {
      const blob = new Blob(['﻿' + csv], {type:'text/csv;charset=utf-8'});
      const a = document.createElement('a');
      a.href = URL.createObjectURL(blob);
      a.download = 'tianshu-users-' + new Date().toISOString().slice(0,10) + '.csv';
      document.body.appendChild(a); a.click(); a.remove();
      URL.revokeObjectURL(a.href);
    };
  }
}

(async () => {
  await fetchMe();
  await loadUsers();
})();
</script>
</body></html>
"""


@app.get("/admin/users", response_class=HTMLResponse)
async def admin_users_page(request: Request) -> str:
    # middleware 已经验证登录;这里再确认 admin
    user = require_user(request)
    if not user.is_admin():
        raise HTTPException(403, "需要管理员权限")
    return ADMIN_USERS_HTML


REPORTS_HTML = r"""<!doctype html>
<html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>报告 — 天枢·裁决</title>
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin><link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
  :root{--bg:#ffffff;--surface:#f0f0f0;--surface-2:#ebebeb;--surface-3:#dcdcdc;
    --line:#c4c4c4;--line-2:#9e9e9e;
    --fg:#0a0a0a;--fg-2:#262626;--fg-3:#4a4a4a;--fg-4:#6e6e6e;
    --ac:#a8401f;--ac-2:#c45a3a;--ac-bg:rgba(168,64,31,.14);--ac-line:rgba(168,64,31,.58);
    --warn:#8a5300;--ok:#4f6b35;--bad:#8a2d12;--info:#3f5560;
    --running:#7a4f00;
    --mono:'Inter','SF Mono',ui-monospace,Menlo,monospace;
    --sans:'Noto Sans SC','PingFang SC',-apple-system,'Microsoft YaHei',sans-serif;
    --serif:'Noto Serif SC','Songti SC','STSong',Georgia,serif;}
  @media (prefers-reduced-motion: reduce){*,*::before,*::after{animation-duration:.01ms!important;transition-duration:.01ms!important}}
  *{box-sizing:border-box}
  html,body{margin:0;background:var(--bg);color:var(--fg);font-family:var(--sans);min-height:100%;
    -webkit-font-smoothing:antialiased}
  body{background:
    radial-gradient(ellipse 90% 50% at 50% -10%, rgba(196,90,58,.07), transparent 65%) fixed,
    radial-gradient(ellipse 80% 40% at 50% 110%, rgba(255,255,255,.02), transparent 60%) fixed,
    var(--bg);}
  ::selection{background:var(--ac-bg);color:var(--ac-2)}
  ::-webkit-scrollbar{width:10px;height:10px}
  ::-webkit-scrollbar-track{background:transparent}
  ::-webkit-scrollbar-thumb{background:var(--surface-3);border-radius:5px;border:2px solid var(--bg)}
  ::-webkit-scrollbar-thumb:hover{background:var(--line-2)}
  .topbar{display:flex;align-items:center;gap:12px;height:56px;padding:0 24px;
    border-bottom:1px solid var(--line);background:rgba(255,255,255,.94);
    position:sticky;top:0;z-index:100;backdrop-filter:saturate(180%) blur(20px);
    -webkit-backdrop-filter:saturate(180%) blur(20px)}
  .topbar .logo{width:28px;height:28px;flex-shrink:0;
    background:linear-gradient(135deg,#262626,#1a1a1a);border-radius:6px;
    display:grid;place-items:center;color:#001f1a;font-weight:700;font-size:11px;
    letter-spacing:-.02em;
    box-shadow:0 1px 2px rgba(0,0,0,.4),inset 0 1px 0 rgba(255,255,255,.18)}
  .topbar .logo .logo-mark{width:18px;height:18px;display:block;
    filter:drop-shadow(0 .5px 0 rgba(255,255,255,.18))}
  .topbar h1{margin:0;font-size:15px;font-weight:600;letter-spacing:-.015em}
  .topbar h1.brand{display:inline-flex;align-items:baseline;gap:0}
  .topbar h1.brand strong{font-weight:600;color:var(--fg);letter-spacing:-.02em}
  .topbar h1.brand .brand-sub{font-weight:400;color:var(--fg-3);margin-left:8px;font-size:13px;letter-spacing:.005em}
  .topbar nav{display:flex;gap:2px;margin-left:24px}
  .topbar nav a{padding:6px 12px;border-radius:6px;color:var(--fg-2);text-decoration:none;
    font-size:13px;transition:background .12s,color .12s}
  .topbar nav a:hover{background:var(--surface-2);color:var(--fg)}
  .topbar nav a.active{background:var(--ac-bg);color:var(--ac-2)}
  main{max-width:1200px;margin:0 auto;padding:40px 24px 80px}
  h2{margin:0 0 8px;font-size:24px;letter-spacing:-.025em;font-weight:600}
  .sub{color:var(--fg-2);font-size:13px;margin-bottom:32px;line-height:1.55}
  /* 批量删除工具栏 */
  .bulk-bar{display:flex;align-items:center;gap:12px;padding:10px 14px;
    background:var(--surface);border:1px solid var(--line);border-radius:6px;
    margin-bottom:12px;font-size:13px;color:var(--fg-2)}
  .bulk-bar .bulk-all{display:inline-flex;align-items:center;gap:8px;cursor:pointer;user-select:none}
  .bulk-bar .bulk-all input{margin:0;cursor:pointer;width:14px;height:14px}
  .bulk-bar .bulk-sep{color:var(--line-2)}
  .bulk-bar .bulk-count{font-family:var(--mono);font-size:12px;color:var(--fg-3)}
  .bulk-bar .bulk-count.has-sel{color:var(--ac);font-weight:600}
  .bulk-bar .spacer{flex:1}
  .bulk-bar .bulk-btn{padding:6px 14px;border-radius:4px;font-family:var(--sans);
    font-size:12.5px;cursor:pointer;border:1px solid;transition:all .15s;font-weight:500}
  .bulk-bar .bulk-del{background:#fff;border-color:#c44a2e;color:#c44a2e}
  .bulk-bar .bulk-del:hover:not(:disabled){background:#c44a2e;color:#fff}
  .bulk-bar .bulk-del:disabled{opacity:.4;cursor:not-allowed;border-color:var(--line)}
  .bulk-bar .bulk-clear{background:#fff;border-color:var(--line-2);color:var(--fg-2)}
  .bulk-bar .bulk-clear:hover{border-color:#8a2d12;color:#8a2d12;background:rgba(138,45,18,.04)}
  /* 表里的复选框列 */
  td.row-pick, th.row-pick{width:36px;text-align:center;padding-left:14px}
  td.row-pick input, th.row-pick input{cursor:pointer;width:14px;height:14px;margin:0}
  tr.row.row-selected{background:rgba(196,90,58,.06)}

  .filter-row{
    display:grid;
    grid-template-columns:1fr auto auto;
    grid-template-rows:auto auto;
    gap:10px;
    margin-bottom:18px;
    align-items:center}
  .filter-row .search{grid-column:1;grid-row:1;
    background:var(--surface);border:1px solid var(--line);
    border-radius:8px;color:var(--fg);font-family:var(--sans);font-size:13px;
    padding:0 12px;height:36px;width:100%}
  .filter-row #export-all{grid-column:2;grid-row:1;text-decoration:none}
  .filter-row #refresh{grid-column:3;grid-row:1}
  .filter-row .search::placeholder{color:var(--fg-3)}
  .filter-row .search:focus{outline:none;border-color:var(--ac-line);
    box-shadow:0 0 0 3px var(--ac-bg)}
  /* Filter chips — 紧凑横向滚动,无 wrap;activated 用下划线而非整体填色 */
  .filter-row .filters{
    grid-column:1 / -1;grid-row:2;
    display:flex;gap:2px;align-items:center;
    background:transparent;border:none;border-bottom:1px solid var(--line);
    border-radius:0;padding:0;
    overflow-x:auto;scrollbar-width:none;max-width:100%}
  .filter-row .filters::-webkit-scrollbar{display:none}
  .filter-row .filters button{
    background:transparent;border:none;color:var(--fg-3);
    padding:9px 10px;font-size:12.5px;border-radius:0;cursor:pointer;
    font-family:var(--sans);font-weight:500;white-space:nowrap;flex-shrink:0;
    display:inline-flex;align-items:center;gap:6px;line-height:1;
    border-bottom:2px solid transparent;margin-bottom:-1px;
    transition:color .12s,border-color .12s}
  .filter-row .filters button:hover{color:var(--fg)}
  .filter-row .filters button.active{color:var(--fg);border-bottom-color:var(--ac);font-weight:600}
  .filter-row .filters button .ic{display:none}
  .filter-row .refresh{grid-column:2;grid-row:1;
    background:transparent;border:1px solid var(--line-2);color:var(--fg-2);
    padding:0 14px;height:36px;border-radius:8px;
    font-family:var(--sans);font-size:12.5px;font-weight:500;cursor:pointer;
    white-space:nowrap;flex-shrink:0;
    display:inline-flex;align-items:center;gap:6px;
    transition:background .12s,color .12s,border-color .12s}
  .filter-row .refresh:hover{border-color:var(--ac-line);color:var(--ac-2);background:var(--surface-2)}
  table{width:100%;border-collapse:collapse;font-size:12.5px;background:var(--surface);
    border:1px solid var(--line);border-radius:10px;overflow:hidden}
  table th{font-size:10.5px;text-transform:uppercase;letter-spacing:.1em;color:var(--fg-3);
    text-align:left;padding:11px 14px;border-bottom:1px solid var(--line);font-weight:600;
    background:var(--surface-2)}
  table td{padding:12px 14px;border-bottom:1px solid var(--line);font-family:var(--mono);
    font-size:12px;color:var(--fg-2)}
  table tr:last-child td{border-bottom:none}
  table tr.row{cursor:pointer;transition:background .15s}
  table tr.row:hover{background:var(--surface-2)}
  /* 项目编号 + 项目名称 单元格样式 */
  .proj-code{background:var(--surface-2);padding:2px 8px;border-radius:4px;
    color:var(--fg);font-family:var(--mono);font-size:11.5px;border:1px solid var(--line);
    display:inline-block;max-width:160px;overflow:hidden;text-overflow:ellipsis;
    white-space:nowrap;vertical-align:middle}
  .proj-name{font-size:12.5px;color:var(--fg);max-width:220px;
    display:inline-block;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;
    vertical-align:middle}
  .muted{color:var(--fg-4);font-size:11.5px}
  table .name{color:var(--fg)}
  table .id{color:var(--fg-3);overflow:hidden;text-overflow:ellipsis;white-space:nowrap;max-width:240px;display:block}
  .status-pill{display:inline-flex;align-items:center;gap:4px;font-family:var(--mono);
    font-size:10px;padding:3px 9px;border-radius:999px;font-weight:600;text-transform:uppercase;
    letter-spacing:.05em;line-height:1}
  .status-pill.queued{background:rgba(168,174,184,.10);color:var(--fg-2)}
  .status-pill.running{background:rgba(167,139,250,.14);color:#c4b5fd}
  .status-pill.succeeded{background:rgba(52,211,153,.10);color:var(--ok)}
  .status-pill.failed{background:rgba(248,113,113,.10);color:var(--bad)}
  .status-pill.saved{background:rgba(168,176,189,.08);color:var(--fg-2);font-weight:500}
  .empty{padding:40px;text-align:center;color:var(--fg-3);font-size:13px;
    border:1px dashed var(--line);border-radius:10px;font-family:var(--mono)}
  .reports-pre{background:var(--surface);border:1px solid var(--line);border-radius:8px;
    padding:14px 18px;font-family:var(--mono);font-size:11.5px;line-height:1.6;color:var(--fg);
    max-height:520px;overflow-y:auto;white-space:pre-wrap;word-break:break-word;margin-top:14px}
  .modal-overlay{position:fixed;inset:0;background:rgba(0,0,0,.7);display:none;
    align-items:flex-start;justify-content:center;padding:40px 20px;z-index:50}
  .modal-overlay.open{display:flex}
  .modal{background:var(--surface);border:1px solid var(--line);border-radius:12px;
    max-width:1080px;width:100%;max-height:calc(100vh - 80px);display:flex;flex-direction:column;
    overflow:hidden;box-shadow:0 24px 60px rgba(0,0,0,.6)}
  .modal-head{display:flex;align-items:center;gap:10px;padding:14px 20px;
    border-bottom:1px solid var(--line);background:var(--surface-2)}
  .modal-head .close{background:transparent;border:none;color:var(--fg-3);
    font-size:18px;cursor:pointer;width:28px;height:28px;border-radius:5px}
  .modal-head .close:hover{background:var(--surface);color:var(--fg)}
  .modal-head .back{
    background:transparent;border:1px solid var(--line-2);color:var(--fg-2);
    padding:5px 12px;border-radius:6px;cursor:pointer;
    font-family:var(--sans);font-size:12.5px;font-weight:500;
    display:inline-flex;align-items:center;gap:6px;
    transition:background .12s,color .12s,border-color .12s}
  .modal-head .back:hover{border-color:var(--ac-line);color:var(--ac-2);background:var(--ac-bg)}
  .modal-divider{width:1px;height:18px;background:var(--line-2);margin:0 4px}
  .modal-head .download{background:transparent;border:1px solid var(--line-2);color:var(--fg-2);
    padding:5px 12px;border-radius:5px;font-family:var(--mono);font-size:11.5px;cursor:pointer}
  .modal-head .download:hover{border-color:var(--ac);color:var(--ac)}
  .modal-tabs{display:flex;background:var(--surface-2);border-bottom:1px solid var(--line);padding:0 18px}
  .modal-tabs button{background:transparent;border:none;color:var(--fg-3);
    padding:9px 14px;cursor:pointer;font-size:12px;font-family:var(--sans);
    border-bottom:2px solid transparent;transition:all .15s}
  .modal-tabs button:hover{color:var(--fg-2)}
  .modal-tabs button.active{color:var(--ac);border-bottom-color:var(--ac)}
  .modal-body{flex:1;overflow-y:auto;background:var(--bg)}
  .modal-body .pane-json{padding:14px 20px}
  .modal-body .pane-json pre{margin:0;font-family:var(--mono);font-size:11.5px;line-height:1.65;
    color:var(--fg);white-space:pre-wrap;word-break:break-word}
  /* === 失败状态卡片 — 不要把 failed 包成"5 段都是 0"的报告 === */
  .fail-card{padding:24px 28px;font-family:var(--sans),inherit;color:var(--fg)}
  .fail-banner{display:flex;align-items:center;gap:12px;margin-bottom:18px}
  .fail-banner .fail-tag{background:#b9482e;color:#fff;font-family:var(--mono);font-size:11px;
    letter-spacing:.12em;padding:4px 10px;border-radius:3px;font-weight:600}
  .fail-banner .fail-title{font-family:var(--serif,inherit);font-size:18px;font-weight:500;color:var(--fg)}
  .fail-meta{display:flex;gap:18px;font-family:var(--mono);font-size:12px;color:var(--fg-3);
    border-bottom:1px solid var(--line);padding-bottom:14px;margin-bottom:18px;flex-wrap:wrap}
  .fail-meta code{background:var(--surface-2);padding:1px 6px;border-radius:3px;color:var(--fg-2)}
  .fail-section{margin-bottom:14px}
  .fail-section .fail-label{font-size:11.5px;color:var(--fg-3);letter-spacing:.16em;
    text-transform:uppercase;margin-bottom:6px}
  .fail-section .fail-err{margin:0;padding:12px 14px;background:rgba(185,72,46,.06);
    border-left:3px solid #b9482e;color:#7a2d1c;font-family:var(--mono);font-size:12.5px;
    line-height:1.65;white-space:pre-wrap;word-break:break-word;border-radius:0 4px 4px 0}
  .fail-trace{margin:14px 0;border:1px solid var(--line);border-radius:4px;
    background:var(--surface)}
  .fail-trace summary{cursor:pointer;padding:10px 14px;font-size:12px;color:var(--fg-2);
    font-family:var(--mono);user-select:none}
  .fail-trace summary:hover{color:var(--fg)}
  .fail-trace pre{margin:0;padding:14px;border-top:1px solid var(--line);font-family:var(--mono);
    font-size:11.5px;color:var(--fg-2);line-height:1.6;white-space:pre-wrap;word-break:break-word}
  .fail-logs{margin-top:14px}
  .fail-logs .fail-logs-head{font-size:11.5px;color:var(--fg-3);letter-spacing:.16em;
    text-transform:uppercase;margin-bottom:6px}
  .fail-logs pre{margin:0;padding:12px 14px;background:var(--surface);border:1px solid var(--line);
    border-radius:4px;font-family:var(--mono);font-size:11.5px;color:var(--fg-2);
    line-height:1.6;white-space:pre-wrap;word-break:break-word}
  .fail-hint{margin-top:18px;padding:12px 14px;background:rgba(161,98,7,.08);border-left:3px solid var(--warn);
    color:var(--fg-2);font-size:12.5px;line-height:1.7;border-radius:0 4px 4px 0}

  /* Embedded report renderer styles (mirror tool detail page) */
  .report-hero{display:flex;align-items:flex-start;gap:14px;padding:16px 20px;
    border-bottom:1px solid var(--line);background:var(--surface)}
  .report-hero .report-icon{font-size:22px;width:40px;height:40px;border-radius:8px;
    background:var(--surface-2);border:1px solid var(--line-2);display:grid;place-items:center}
  .report-hero h4{margin:0;font-size:15px;font-weight:600;color:var(--fg)}
  .report-hero .meta{font-family:var(--mono);font-size:11px;color:var(--fg-3);margin-top:4px}
  .report-hero .meta code{background:var(--bg);padding:1px 6px;border-radius:3px;color:var(--ac-2);border:1px solid var(--line)}
  .stat-pills{display:flex;gap:6px;flex-wrap:wrap;margin-top:8px}
  .stat-pills .pill{font-family:var(--mono);font-size:10.5px;padding:2px 8px;border-radius:3px;
    background:var(--surface-2);border:1px solid var(--line);color:var(--fg-2)}
  .stat-pills .pill .v{color:var(--ac);font-weight:600}
  .gate-banner{padding:14px 18px;border-bottom:1px solid var(--line);display:flex;
    align-items:flex-start;gap:12px;font-size:13px;line-height:1.55}
  .gate-banner.proceed{background:rgba(74,222,128,.05)}
  .gate-banner.reject{background:rgba(248,113,113,.05)}
  .gate-banner.warn{background:rgba(251,191,36,.05)}
  .gate-banner .badge{padding:3px 9px;border-radius:4px;font-family:var(--mono);font-size:11px;
    font-weight:600;text-transform:uppercase;letter-spacing:.05em;flex-shrink:0;margin-top:1px}
  .gate-banner.proceed .badge{background:rgba(74,222,128,.12);color:var(--ok)}
  .gate-banner.reject .badge{background:rgba(248,113,113,.12);color:var(--bad)}
  .gate-banner.warn .badge{background:rgba(251,191,36,.12);color:var(--warn)}
  .gate-banner .reasons{color:var(--fg-2);font-family:var(--mono);font-size:12px}
  .gate-banner .reasons div{margin-top:3px}
  .report-sub{border-bottom:1px solid var(--line)}
  .report-sub:last-child{border-bottom:none}
  .report-sub-head{display:flex;align-items:center;gap:10px;padding:10px 18px;
    cursor:pointer;font-size:12.5px;background:var(--surface);transition:background .15s}
  .report-sub-head:hover{background:var(--surface-2)}
  .report-sub.open .report-sub-head{background:var(--surface-2)}
  .report-sub-twirl{color:var(--fg-3);font-family:var(--mono);font-size:10px;width:10px;transition:transform .15s}
  .report-sub.open .report-sub-twirl{transform:rotate(90deg);color:var(--ac)}
  .report-sub-num{font-family:var(--mono);font-size:10.5px;font-weight:600;color:var(--ac);
    background:rgba(16,185,129,.10);border:1px solid rgba(16,185,129,.25);
    width:20px;height:20px;border-radius:50%;display:inline-grid;place-items:center;flex-shrink:0}
  .report-sub-name{color:var(--fg);font-weight:500}
  .report-sub-stats{margin-left:auto;font-family:var(--mono);font-size:11px;color:var(--fg-3);display:flex;gap:8px}
  .report-sub-stats .chip{padding:2px 7px;border-radius:3px;background:var(--surface-2);border:1px solid var(--line-2)}
  .report-sub-body{display:none;padding:14px 20px;background:var(--bg)}
  .report-sub.open .report-sub-body{display:block}
  .report-kv{display:grid;grid-template-columns:max-content 1fr;gap:6px 16px;margin:0;font-size:12.5px}
  .report-kv dt{color:var(--fg-3);font-family:var(--mono);font-size:11.5px;align-self:start}
  .report-kv dd{margin:0;color:var(--fg);min-width:0;line-height:1.55}
  .report-table{width:100%;border-collapse:collapse;font-size:11.5px;margin:8px 0;
    border:1px solid var(--line);border-radius:6px;overflow:hidden}
  .report-table th{font-family:var(--mono);font-size:10px;text-align:left;padding:8px 10px;
    background:var(--surface-2);color:var(--fg-3);text-transform:uppercase;letter-spacing:.07em;
    font-weight:600;border-bottom:1px solid var(--line);white-space:nowrap}
  .report-table td{padding:8px 10px;border-bottom:1px solid var(--line);font-family:var(--mono);
    font-size:11px;color:var(--fg);vertical-align:top;word-break:break-word;max-width:320px}
  .report-table tr:last-child td{border-bottom:none}
  .sev{font-family:var(--mono);font-size:9.5px;font-weight:700;padding:1px 6px;border-radius:3px;
    text-transform:uppercase;letter-spacing:.05em;display:inline-block}
  .sev-critical{background:rgba(248,113,113,.18);color:var(--bad);border:1px solid rgba(248,113,113,.4)}
  .sev-high,.sev-major{background:rgba(248,113,113,.10);color:var(--bad)}
  .sev-medium,.sev-warn{background:rgba(251,191,36,.10);color:var(--warn)}
  .sev-low,.sev-minor,.sev-cosmetic{background:rgba(168,174,184,.10);color:var(--fg-2)}
  .sev-info{background:rgba(16,185,129,.08);color:var(--ac)}
  .confbar{display:inline-flex;align-items:center;gap:6px;font-family:var(--mono);font-size:11px}
  .confbar .track{width:80px;height:5px;border-radius:3px;background:var(--line);overflow:hidden}
  .confbar .fill{height:100%;background:var(--ac)}
  .confbar .pct{color:var(--ac);font-weight:600}
  .empty-array{color:var(--fg-4);font-family:var(--mono);font-size:11px;font-style:italic}
  details.report-detail summary{cursor:pointer;font-family:var(--mono);font-size:11px;color:var(--ac)}
  details.report-detail[open] > div{margin:6px 0 6px 16px;padding:8px 10px;background:var(--surface-2);
    border-radius:5px;border-left:2px solid var(--line-2)}
  ul.report-list{margin:0;padding-left:18px;font-size:12px;line-height:1.7}
  /* === 执行摘要 4 块结构 === */
  .exec-block{background:var(--surface);border:1px solid var(--line);border-radius:10px;
    padding:18px 22px;margin:14px 0;page-break-inside:avoid}
  .exec-head{display:flex;align-items:center;gap:10px;margin-bottom:14px}
  .exec-head h3{margin:0;font-size:15px;font-weight:600;color:var(--fg)}
  .exec-num{display:inline-grid;place-items:center;width:28px;height:28px;border-radius:8px;
    background:rgba(16,185,129,.10);color:var(--ac);font-size:14px;font-weight:700;
    border:1px solid rgba(16,185,129,.32);font-family:var(--mono)}
  .verdict{padding:14px 18px;border-radius:8px;display:flex;align-items:center;gap:14px;
    font-size:16px;font-weight:600}
  .verdict.ok{background:rgba(74,222,128,.10);color:var(--ok);border:1px solid rgba(74,222,128,.4)}
  .verdict.warn{background:rgba(251,191,36,.10);color:var(--warn);border:1px solid rgba(251,191,36,.4)}
  .verdict.bad{background:rgba(248,113,113,.10);color:var(--bad);border:1px solid rgba(248,113,113,.4)}
  .verdict.skip{background:var(--surface-2);color:var(--fg-2);border:1px solid var(--line-2)}
  .verdict-icon{font-size:22px}
  .verdict-summary{margin-top:12px;padding:10px 14px;background:var(--surface-2);
    border-left:3px solid var(--line-2);border-radius:0 6px 6px 0;
    font-size:13.5px;line-height:1.7;color:var(--fg)}
  .sev-strip{margin-top:10px;padding:8px 14px;background:var(--surface-2);border-radius:6px;
    font-family:var(--mono);font-size:12px;color:var(--fg-2)}
  .sev-strip strong{color:var(--fg);font-weight:600}
  .exec-muted{color:var(--fg-3);font-size:13px;margin:4px 0;font-style:normal}

  /* === KPI 卡片(测试结论里的 4 个数字)=== */
  .exec-kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-top:14px}
  .exec-kpi{display:flex;flex-direction:column;align-items:flex-start;justify-content:center;
    padding:14px 16px;background:var(--bg);border:1px solid var(--line);border-radius:8px}
  .exec-kpi-num{font-size:28px;font-weight:700;color:var(--fg);font-family:var(--sans);
    font-variant-numeric:tabular-nums;line-height:1;letter-spacing:-.02em}
  .exec-kpi-lbl{font-size:11px;color:var(--fg-3);margin-top:6px;letter-spacing:.08em;
    text-transform:uppercase;font-weight:600}

  /* === 严重度分布条 === */
  .sev-bar{display:flex;height:8px;border-radius:4px;overflow:hidden;margin-top:14px;
    background:var(--surface-2);border:1px solid var(--line)}
  .sev-bar-seg{display:flex;align-items:center;justify-content:center;color:#fff;
    font-size:10px;font-weight:600;line-height:8px;min-width:8px;transition:flex .2s}
  .sev-bar-critical{background:#dc2626}
  .sev-bar-high{background:#ea580c}
  .sev-bar-medium{background:#ca8a04}
  .sev-bar-low{background:#16a34a}
  .sev-bar-info{background:#0891b2}

  /* === 优先级分布条(用例段)=== */
  .pri-bar{display:flex;height:24px;border-radius:6px;overflow:hidden;margin-top:8px;margin-bottom:12px;
    background:var(--surface-2);border:1px solid var(--line)}
  .pri-bar-seg{display:flex;align-items:center;justify-content:center;color:#fff;
    font-size:11px;font-weight:600;min-width:36px;letter-spacing:.04em;padding:0 8px}
  .pri-bar-P0{background:#dc2626}
  .pri-bar-P1{background:#ea580c}
  .pri-bar-P2{background:#ca8a04}
  .pri-bar-P3{background:#737373}

  /* === 风险卡片 === */
  .exec-risk-list{display:flex;flex-direction:column;gap:10px;margin-top:4px}
  .exec-risk-item{padding:12px 14px 12px 14px;background:var(--bg);border:1px solid var(--line);
    border-left:3px solid #ca8a04;border-radius:6px}
  .exec-risk-head{display:flex;align-items:baseline;gap:10px;margin-bottom:6px;flex-wrap:wrap}
  .exec-risk-title{font-weight:600;color:var(--fg);font-size:14px;line-height:1.5;flex:1;min-width:0}
  .exec-risk-line{font-size:12.5px;color:var(--fg-2);line-height:1.65;margin-top:3px;
    display:flex;align-items:baseline;gap:8px}
  .exec-risk-line .lbl{display:inline-block;min-width:42px;font-size:11px;color:var(--fg-3);
    font-weight:600;letter-spacing:.04em;text-transform:uppercase;flex-shrink:0}

  /* === 阻碍卡片 === */
  .exec-blocker-list{display:flex;flex-direction:column;gap:10px;margin-top:4px}
  .exec-blocker-item{padding:12px 14px;background:rgba(220,38,38,.04);border:1px solid rgba(220,38,38,.20);
    border-left:3px solid #dc2626;border-radius:6px}
  .exec-blocker-head{display:flex;align-items:center;gap:10px;flex-wrap:wrap;margin-bottom:6px}
  .blocker-tag{font-size:10.5px;font-weight:700;padding:2px 8px;background:rgba(220,38,38,.12);
    color:#dc2626;letter-spacing:.1em;text-transform:uppercase;border-radius:4px}
  .exec-blocker-title{font-weight:600;color:var(--fg);font-size:14px;flex:1;min-width:0}
  .exec-blocker-line{font-size:12.5px;color:var(--fg-2);line-height:1.65;margin-top:3px;
    display:flex;align-items:baseline;gap:8px}
  .exec-blocker-line .lbl{display:inline-block;min-width:62px;font-size:11px;color:var(--fg-3);
    font-weight:600;letter-spacing:.04em;text-transform:uppercase;flex-shrink:0}
  .exec-blocker-line.fix .lbl{color:var(--ok)}
  .exec-blocker-line.fix{color:var(--fg)}

  /* === 优先级 pill(用例表 + Bug 卡)=== */
  .pri-tag,.exec-issue-meta .meta-chip.pri-P0,.exec-issue-meta .meta-chip.pri-P1,
  .exec-issue-meta .meta-chip.pri-P2,.exec-issue-meta .meta-chip.pri-P3,
  .exec-blocker-head .meta-chip.pri-P0,.exec-blocker-head .meta-chip.pri-P1,
  .exec-blocker-head .meta-chip.pri-P2,.exec-blocker-head .meta-chip.pri-P3{
    font-size:10.5px;font-weight:700;padding:2px 8px;border-radius:4px;letter-spacing:.04em;
    background:var(--surface-2);color:var(--fg-2);font-family:var(--sans)}
  .pri-tag.pri-P0,.exec-issue-meta .meta-chip.pri-P0{background:rgba(220,38,38,.12);color:#dc2626}
  .pri-tag.pri-P1,.exec-issue-meta .meta-chip.pri-P1{background:rgba(234,88,12,.12);color:#ea580c}
  .pri-tag.pri-P2,.exec-issue-meta .meta-chip.pri-P2{background:rgba(202,138,4,.12);color:#a16207}
  .pri-tag.pri-P3{background:var(--surface-2);color:var(--fg-3)}
  .exec-issue{border:1px solid var(--line);border-radius:8px;padding:12px 16px;margin-bottom:10px;
    background:var(--bg);border-left:3px solid var(--line-2)}
  .exec-issue.sev-critical{border-left-color:var(--bad)}
  .exec-issue.sev-high{border-left-color:var(--warn)}
  .exec-issue.sev-medium{border-left-color:#eab308}
  .exec-issue.sev-low{border-left-color:var(--ok)}
  .exec-issue.sev-info{border-left-color:var(--ac)}
  .exec-issue-head{display:flex;align-items:center;gap:10px;margin-bottom:6px}
  .sev-tag{font-size:11px;font-weight:600;padding:2px 8px;border-radius:999px;
    background:rgba(16,185,129,.10);color:var(--ac);border:1px solid rgba(16,185,129,.32)}
  .exec-issue-title{font-weight:600;color:var(--fg);font-size:13.5px}
  .exec-issue-loc{color:var(--fg-3);font-size:12px;margin:6px 0;font-family:var(--mono)}
  .exec-issue-desc{margin:6px 0;color:var(--fg);line-height:1.6;font-size:13px}
  .exec-issue-fix{margin-top:8px;padding:8px 12px;background:rgba(16,185,129,.06);
    border-radius:5px;font-size:12.5px;color:var(--fg-2);
    border:1px solid rgba(16,185,129,.18)}
  .exec-issue-meta{display:flex;flex-wrap:wrap;gap:6px;margin:4px 0 8px}
  .exec-issue-meta .meta-chip{font-size:11px;font-family:var(--mono);
    padding:2px 9px;border-radius:999px;background:var(--surface-2);
    border:1px solid var(--line-2);color:var(--fg-2)}
  .exec-issue-meta .meta-chip.role{color:var(--ac);border-color:var(--ac-line);background:rgba(16,185,129,.08)}
  .exec-issue-section{margin:10px 0;padding:10px 14px;border-radius:6px;background:var(--surface-2)}
  .exec-issue-section.fix{background:rgba(16,185,129,.06);border:1px solid rgba(16,185,129,.20)}
  .exec-issue-section.verify{background:rgba(96,165,250,.05);border:1px solid rgba(96,165,250,.20)}
  .sec-lbl{font-size:12px;font-weight:600;color:var(--fg-2);margin-bottom:6px;letter-spacing:.02em}
  .sec-body{font-size:13px;color:var(--fg);line-height:1.65}
  .repro-list{margin:0;padding-left:22px;font-size:12.5px;line-height:1.75;color:var(--fg-2)}
  .accept-line{margin-top:8px;padding:8px 12px;background:rgba(96,165,250,.08);
    border-left:3px solid #60a5fa;border-radius:0 5px 5px 0;font-size:12.5px;color:var(--fg-2)}
  .related-cases{margin-top:10px;font-size:12px;color:var(--fg-3);font-family:var(--mono)}
  .related-cases code{background:var(--surface-2);padding:2px 7px;border-radius:3px;
    color:var(--ac);font-size:11.5px;margin-right:4px}
  .exec-issue-impact,.exec-issue-evidence{margin-top:8px;font-size:12px;color:var(--fg-3)}
  .exec-case-total{display:flex;align-items:baseline;gap:10px}
  .exec-case-total .num{font-size:32px;font-weight:700;color:var(--fg);font-family:var(--mono);
    font-variant-numeric:tabular-nums}
  .exec-case-total .lbl{font-size:13px;color:var(--fg-2)}
  .exec-pri{display:flex;flex-wrap:wrap;gap:10px;margin-top:12px}
  .exec-pri-cell{display:flex;flex-direction:column;align-items:center;
    background:var(--surface-2);border:1px solid var(--line);border-radius:8px;
    padding:10px 14px;min-width:80px}
  .exec-pri-cell .num{font-size:18px;font-weight:700;color:var(--fg);font-family:var(--mono)}
  .exec-pri-cell .lbl{font-size:10.5px;color:var(--fg-3);text-transform:uppercase;letter-spacing:.04em;margin-top:2px}
  .exec-case-note{margin-top:12px;padding:10px 14px;background:rgba(251,191,36,.06);
    border-left:3px solid var(--warn);border-radius:0 6px 6px 0;
    color:var(--fg-2);font-size:12.5px}
  /* === Screenshot evidence (step5 / h5_adapt) === */
  .report-screenshots{margin:14px 0 0;padding:14px;background:var(--surface-2);
    border:1px solid var(--line);border-radius:10px}
  .screenshots-head{font-size:13px;font-weight:600;color:var(--fg);margin-bottom:10px}
  .screenshots-hint{font-size:10.5px;font-weight:400;color:var(--fg-3);margin-left:6px;font-style:italic}
  .shot-group{margin-top:10px;padding-top:10px;border-top:1px dashed var(--line)}
  .shot-group:first-of-type{margin-top:0;padding-top:0;border-top:none}
  .shot-url{font-family:"SF Mono",ui-monospace,monospace;font-size:11.5px;color:var(--fg-3);margin-bottom:8px}
  .shot-url code{background:var(--bg);padding:2px 8px;border-radius:4px;color:var(--ac-2)}
  .shot-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(180px,1fr));gap:10px}
  .shot-cell{display:block;background:var(--bg);border:1px solid var(--line);border-radius:8px;
    overflow:hidden;text-decoration:none;color:var(--fg-2);
    transition:transform .15s,border-color .15s,box-shadow .15s}
  .shot-cell:hover{transform:translateY(-2px);border-color:var(--ac-line);
    box-shadow:0 8px 22px rgba(0,0,0,.35)}
  .shot-cell img{display:block;width:100%;height:auto;max-height:240px;object-fit:cover;
    background:#000}
  .shot-cap{padding:6px 10px;font-size:11px;font-family:"SF Mono",ui-monospace,monospace;
    color:var(--fg-3);border-top:1px solid var(--line);
    display:flex;align-items:center;justify-content:space-between;gap:6px}
  .shot-cap .issue-badge{background:rgba(220,38,38,.14);color:#dc2626;
    padding:1px 7px;border-radius:999px;font-weight:600;font-size:10px}
  /* === Row action buttons (download / delete) === */
  td.actions{white-space:nowrap;display:flex;gap:6px;align-items:center}
  td.actions a{
    display:inline-flex;align-items:center;justify-content:center;
    padding:4px 10px;border-radius:6px;
    font-family:"SF Mono",ui-monospace,monospace;font-size:11px;
    text-decoration:none;cursor:pointer;
    transition:background .12s,color .12s,border-color .12s;
    border:1px solid var(--line-2);color:var(--fg-2);background:transparent}
  td.actions a:hover{border-color:var(--ac-line);color:var(--ac-2);background:var(--ac-bg)}
  td.actions a.del{color:var(--fg-3)}
  td.actions a.del:hover{border-color:rgba(248,113,113,.4);color:var(--bad);background:rgba(248,113,113,.08)}
  /* === Hero & KPI tiles === */
  .reports-hero{margin-bottom:32px}
  .hero-eyebrow{display:inline-flex;align-items:center;gap:6px;
    font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:.08em;
    color:var(--ac-2);margin-bottom:12px;padding:5px 12px;border-radius:999px;
    background:var(--ac-bg);border:1px solid var(--ac-line)}
  .hero-eyebrow::before{content:"";width:6px;height:6px;border-radius:50%;
    background:var(--ac-2);box-shadow:0 0 0 2px var(--ac-bg)}
  .reports-hero h2{font-size:32px;letter-spacing:-.03em;margin:0 0 8px;line-height:1.1;
    background:linear-gradient(135deg,var(--fg) 0%,var(--fg-2) 75%);
    -webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
  .reports-hero .sub{color:var(--fg-2);font-size:14.5px;line-height:1.55;margin:0 0 24px;max-width:680px}
  .kpi-strip{display:grid;grid-template-columns:repeat(4,1fr);gap:24px;margin-bottom:14px;
    padding:14px 0;border-top:1px solid var(--line);border-bottom:1px solid var(--line)}
  .kpi{background:transparent;border:none;border-radius:0;
    padding:0;position:relative;overflow:visible;
    transition:none}
  .kpi:hover{transform:none;border-color:transparent}
  .kpi::before{display:none}
  .kpi .num{font-size:20px;font-weight:600;letter-spacing:-.02em;color:var(--fg);
    font-variant-numeric:tabular-nums;display:flex;align-items:baseline;gap:6px}
  .kpi .num.ok{color:var(--ok)}
  .kpi .num.bad{color:var(--bad)}
  .kpi .num.warn{color:var(--warn)}
  .kpi .lbl{font-size:11px;color:var(--fg-3);text-transform:uppercase;letter-spacing:.06em;
    margin-top:4px;font-weight:500}
  .kpi .icon-bg{display:none;position:absolute;right:-8px;bottom:-8px;font-size:60px;opacity:.04;line-height:1;
    pointer-events:none}
</style></head>
<body>
<div class="topbar">
  <a class="brand-link" href="/tools" title="天枢 · 裁决 · 返回主页" style="text-decoration:none;display:inline-flex;align-items:center;gap:10px;font-family:'Noto Serif SC',Georgia,serif;font-size:19px;font-weight:600;letter-spacing:.18em;color:var(--fg);margin-right:24px;padding:4px 0"><svg viewBox="0 0 24 24" fill="currentColor" width="24" height="24" aria-hidden="true" style="color:var(--ac);opacity:1;flex-shrink:0;filter:drop-shadow(0 1px 2px rgba(196,90,58,.28))"><path d="M12 1.5L13.4 10.6L22.5 12L13.4 13.4L12 22.5L10.6 13.4L1.5 12L10.6 10.6Z"/></svg><span>天枢</span><span style="color:var(--ac);margin:0 6px;font-weight:400">·</span><span>裁决</span></a>
  
  <nav>
    <a href="/tools">工具</a>
    <a href="/reports" class="active">报告</a>
    <a href="/settings" class="superadmin-only" style="display:none">设置</a>
    <a href="/usage" class="admin-only" style="display:none">使用统计</a>
    <a href="/admin/users" class="admin-only" style="display:none">用户管理</a>
  </nav>
  <div class="right" style="margin-left:auto">
    <span class="kbd-hint" id="cmd-trigger" title="打开命令面板"
      style="display:inline-flex;align-items:center;gap:4px;font-size:11px;color:var(--fg-3);
        padding:4px 8px;border-radius:6px;cursor:pointer;transition:all .12s">
      <kbd style="background:var(--surface-3);border:1px solid var(--line-1);border-bottom-width:2px;
        padding:1px 6px;border-radius:4px;font-size:10.5px;color:var(--fg-2);
        font-family:'SF Mono',ui-monospace,monospace;min-width:18px;text-align:center">⌘</kbd>
      <kbd style="background:var(--surface-3);border:1px solid var(--line-1);border-bottom-width:2px;
        padding:1px 6px;border-radius:4px;font-size:10.5px;color:var(--fg-2);
        font-family:'SF Mono',ui-monospace,monospace;min-width:18px;text-align:center">K</kbd>
      跳转
    </span>
  </div>
</div>
<main>
  <section class="reports-hero">
    <span class="hero-eyebrow">运行档案</span>
    <h2>报告</h2>
    <p class="sub">所有跑过的任务、生成的用例与漏测分析。本地保存,点开任意一条进入预览,在预览里下载（Excel / Markdown / JSON）。</p>
    <div class="kpi-strip">
      <div class="kpi"><div class="num" id="kpi-total">—</div><div class="lbl">累计运行</div><div class="icon-bg">📊</div></div>
      <div class="kpi"><div class="num ok" id="kpi-success">—</div><div class="lbl">成功</div><div class="icon-bg">✓</div></div>
      <div class="kpi"><div class="num bad" id="kpi-fail">—</div><div class="lbl">失败</div><div class="icon-bg">✗</div></div>
      <div class="kpi"><div class="num" id="kpi-recent">—</div><div class="lbl">近 24h</div><div class="icon-bg">⏱</div></div>
    </div>
  </section>
  <div class="filter-row">
    <input class="search" id="search" placeholder="搜索 报告名称 / 测试类型 / 产品编号 / 产品名称 / 执行人 / run id">
    <div class="filters" id="filters"></div>
    <button class="refresh" id="refresh">↻ 刷新</button>
  </div>
  <!-- 多选批量操作工具栏 — 只在表格有内容时显示 -->
  <div class="bulk-bar" id="bulk-bar" style="display:none">
    <label class="bulk-all"><input type="checkbox" id="bulk-all-cb"><span id="bulk-all-label">全选当前页</span></label>
    <span class="bulk-sep">·</span>
    <span class="bulk-count" id="bulk-count">已选 0 条</span>
    <span class="spacer"></span>
    <button class="bulk-btn bulk-del" id="bulk-del-sel" disabled>✕ 删除选中</button>
    <button class="bulk-btn bulk-clear" id="bulk-clear-all">⚠ 清空所有</button>
  </div>
  <div id="content"></div>
</main>

<div class="modal-overlay" id="modal-overlay">
  <div class="modal">
    <div class="modal-head">
      <button class="back" onclick="closeModal()" title="返回报告列表（Esc）">
        <span style="font-size:14px;line-height:1">←</span>
        <span>返回</span>
      </button>
      <span class="modal-divider"></span>
      <span id="modal-icon" style="font-size:18px">·</span>
      <span id="modal-tool" style="font-size:14px;font-weight:500">·</span>
      <span style="font-family:var(--mono);font-size:11px;color:var(--fg-3)" id="modal-id">·</span>
      <span style="margin-left:auto"></span>
      <button class="download" data-act="xlsx" id="modal-dl-xlsx">↓ 下载 Excel</button>
      <button class="download" data-act="md" id="modal-dl-md">↓ Markdown</button>
      <button class="download" data-act="json" id="modal-dl-json">↓ JSON</button>
      <button class="close" onclick="closeModal()" title="关闭（Esc）">×</button>
    </div>
    <div class="modal-tabs">
      <button class="active" data-tab="report" id="modal-tab-report">报告</button>
      <button data-tab="json">{ } 原始 JSON</button>
    </div>
    <div class="modal-body" id="modal-body"><div style="padding:40px;text-align:center;color:var(--fg-3)">加载中…</div></div>
  </div>
</div>

<script>
let toolCatalog = [];                    // {id, icon, name, prompts, prompt_dir, ...}
const toolMap = {};                      // tool_id → catalog entry
const substepNamesCache = {};            // tool_id → {sub_id: name}
let allReports = [];
let currentFilter = 'all';
// 多选批量删除状态
const selectedRunIds = new Set();
let currentVisibleRunIds = [];
let currentRun = null;                   // currently-open run (for download buttons)
let currentTool = null;                  // tool meta of current run
let activeModalTab = 'report';

async function load(){
  // Load catalog first so renderer can look up tool icon/prompts
  if (!toolCatalog.length){
    const cat = await fetch('/api/tools').then(r=>r.json());
    toolCatalog = cat.tools;
    toolCatalog.forEach(t => toolMap[t.id] = t);
    buildFilters();
  }
  const data = await fetch('/api/reports').then(r=>r.json());
  const seen = new Set();
  allReports = [];
  data.in_memory.forEach(r => { seen.add(r.run_id); allReports.push({...r, kind:'memory'}); });
  data.saved.forEach(r => { if (!seen.has(r.run_id)) allReports.push({...r, kind:'saved'}); });
  // KPI tiles
  const kt = document.getElementById('kpi-total');
  const ks = document.getElementById('kpi-success');
  const kf = document.getElementById('kpi-fail');
  const kr = document.getElementById('kpi-recent');
  if (kt && ks && kf && kr) {
    const total = allReports.length;
    // saved-on-disk reports always count as success (they only persist if completed)
    const okN = allReports.filter(r => r.status === 'succeeded' || r.kind === 'saved').length;
    const failN = allReports.filter(r => r.status === 'failed').length;
    const dayAgo = Date.now()/1000 - 86400;
    const recentN = allReports.filter(r => (r.started_at || r.mtime || 0) > dayAgo).length;
    kt.textContent = total;
    ks.textContent = okN;
    kf.textContent = failN;
    kr.textContent = recentN;
  }
  render();
}

function buildFilters(){
  const f = document.getElementById('filters');
  f.innerHTML = '<button class="active" data-f="all">全部</button>';
  toolCatalog.forEach(t => {
    f.insertAdjacentHTML('beforeend',
      `<button data-f="${t.id}"><span class="ic">${t.icon}</span>${t.name}</button>`);
  });
}

function render(){
  const q = document.getElementById('search').value.toLowerCase().trim();
  const filtered = allReports.filter(r => {
    if (currentFilter !== 'all' && r.tool_id !== currentFilter) return false;
    if (q) {
      const _tn = (toolMap[r.tool_id]||{}).name || r.tool_name || '';
      const hay = (r.run_id + ' ' + _tn + ' ' + (r.project_code||'') + ' ' + (r.project_name||'') + ' ' + (r.owner_username||'')).toLowerCase();
      if (!hay.includes(q)) return false;
    }
    return true;
  });
  const c = document.getElementById('content');
  const bulkBar = document.getElementById('bulk-bar');
  if (!filtered.length){
    c.innerHTML = '<div class="empty">暂无报告 · 去 <a href="/tools" style="color:var(--ac);text-decoration:none">工具</a> 触发一次运行</div>';
    if (bulkBar) bulkBar.style.display = 'none';
    selectedRunIds.clear();
    return;
  }
  if (bulkBar) bulkBar.style.display = '';
  c.innerHTML = `<table><thead><tr>
    <th class="row-pick"></th>
    <th>状态</th><th>报告名称</th><th>测试类型</th><th>产品编号</th><th>产品名称</th><th>执行人</th><th>操作</th>
  </tr></thead><tbody id="tbody"></tbody></table>`;
  const tbody = document.getElementById('tbody');
  // 当前页 run_ids 给"全选"用
  currentVisibleRunIds = filtered.map(r => r.run_id);
  filtered.forEach(r => {
    const t = r.started_at || r.mtime;
    const _p2 = n => String(n).padStart(2,'0');
    let stamp = '';
    if (t) { const d = new Date(t*1000); stamp = `${d.getFullYear()}-${_p2(d.getMonth()+1)}-${_p2(d.getDate())} ${_p2(d.getHours())}:${_p2(d.getMinutes())}:${_p2(d.getSeconds())}`; }
    const status = r.kind === 'memory' ? r.status : 'saved';
    const tm = toolMap[r.tool_id];
    const typePlain = tm ? tm.name : r.tool_id;
    const name = tm ? `${tm.icon} ${tm.name}` : r.tool_id;
    // 报告名称 = 产品名称 · 测试类型 · 日期时间(精确到秒)
    const reportName = [r.project_name, typePlain, stamp].filter(Boolean).join(' · ') || (r.run_id || '').slice(0,8);
    const executor = r.owner_username ? escapeHtml(r.owner_username) : '<span class="muted">—</span>';
    const pc = r.project_code ? `<code class="proj-code">${escapeHtml(r.project_code)}</code>` : '<span class="muted">—</span>';
    const pn = r.project_name ? `<span class="proj-name">${escapeHtml(r.project_name)}</span>` : '<span class="muted">—</span>';
    const row = document.createElement('tr');
    row.className = 'row' + (selectedRunIds.has(r.run_id) ? ' row-selected' : '');
    row.dataset.runid = r.run_id;
    row.onclick = (e) => {
      // 复选框点击不要触发打开报告
      if (e.target && (e.target.closest('.row-pick') || e.target.matches('a,button,input'))) return;
      openReport(r);
    };
    const checked = selectedRunIds.has(r.run_id) ? 'checked' : '';
    row.innerHTML = `
      <td class="row-pick"><input type="checkbox" class="row-cb" data-runid="${r.run_id}" ${checked}></td>
      <td><span class="status-pill ${status}">${status}</span></td>
      <td class="report-name" style="max-width:300px;white-space:normal;line-height:1.45;font-weight:500">${escapeHtml(reportName)}</td>
      <td class="name">${name}</td>
      <td>${pc}</td>
      <td>${pn}</td>
      <td class="executor">${executor}</td>
      <td class="actions">
        <a href="javascript:void(0)" data-runid="${r.run_id}" class="open">预览</a>
        <a href="javascript:void(0)" data-runid="${r.run_id}" class="del">删除</a>
      </td>
    `;
    tbody.appendChild(row);
  });
  // 复选框 change → 维护 selectedRunIds + 行高亮 + bulk bar 状态
  tbody.querySelectorAll('input.row-cb').forEach(cb => {
    cb.onclick = (e) => e.stopPropagation();
    cb.onchange = () => {
      const rid = cb.dataset.runid;
      if (cb.checked) selectedRunIds.add(rid); else selectedRunIds.delete(rid);
      cb.closest('tr').classList.toggle('row-selected', cb.checked);
      updateBulkBar();
    };
  });
  updateBulkBar();
  // Wire 预览 links — 打开报告预览弹窗(下载在弹窗内)
  tbody.querySelectorAll('a.open').forEach(a => {
    a.onclick = (e) => {
      e.stopPropagation();
      const rec = allReports.find(x => x.run_id === a.dataset.runid);
      if (rec) openReport(rec);
    };
  });
  // Wire delete buttons
  tbody.querySelectorAll('a.del').forEach(a => {
    a.onclick = async (e) => {
      e.stopPropagation();
      const runId = a.dataset.runid;
      if (!confirm('确认删除这条报告？\n\n此操作会同时删除内存与磁盘记录，不可恢复。')) return;
      try {
        const res = await fetch('/api/reports/' + runId, {method: 'DELETE'});
        if (!res.ok) {
          const err = await res.text();
          alert('删除失败：' + err);
          return;
        }
        // Reload list
        await load();
      } catch (e) {
        alert('删除请求出错：' + e.message);
      }
    };
  });
}

async function quickDownload(runId, toolId, fmt){
  // 下载流程:先 preflight 一次 HEAD/GET 看是否 200,再触发真正的下载。
  // 这样 401 (未登录) / 403 (无权限) / 404 (报告丢失) 会显示明确提示,
  // 避免浏览器把错误 JSON 当成 .html 保存导致 "无法从网站上提取文件"。
  if (!['html', 'md', 'json', 'xlsx'].includes(fmt)) { alert('不支持的格式:' + fmt); return; }
  const url = `/api/reports/${runId}/export.${fmt}`;
  try {
    const resp = await fetch(url, {method:'GET', credentials:'same-origin'});
    if (!resp.ok){
      let detail = '';
      try {
        const j = await resp.json();
        detail = j.detail || j.error || '';
      } catch(_) {}
      const msg = ({
        401: '会话已过期 — 请重新登录',
        403: detail || '无权下载此报告',
        404: '报告不存在 — 可能服务重启后内存清掉了,且磁盘也没保存',
      })[resp.status] || ('下载失败: HTTP ' + resp.status + (detail ? ' · ' + detail : ''));
      // 用页面已有的 toast (如果有),否则降级 alert
      if (typeof toast === 'function') toast(msg, 'bad');
      else alert(msg);
      return;
    }
    // 200 — 拿到 blob 触发下载
    const blob = await resp.blob();
    // 文件名以服务端 Content-Disposition 为准:SEO/弱网/用例的 html/md 会被后端重定向成 xlsx,
    // 扩展名必须跟着真实内容走,否则存成 .html 的 xlsx → 双击当 HTML 打开就是乱码。
    let fn = `${toolId}_${runId.slice(0,8)}.${fmt}`;
    const cd = resp.headers.get('Content-Disposition') || '';
    let m = cd.match(/filename\*=UTF-8''([^;]+)/i);
    if (m) { try { fn = decodeURIComponent(m[1]); } catch(_){} }
    else if ((m = cd.match(/filename="?([^";]+)"?/i))) { fn = m[1]; }
    else if ((resp.headers.get('Content-Type')||'').includes('spreadsheet')) { fn = `${toolId}_${runId.slice(0,8)}.xlsx`; }
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = fn;
    document.body.appendChild(a); a.click(); a.remove();
    setTimeout(() => URL.revokeObjectURL(a.href), 1000);
  } catch(err){
    const msg = '下载请求失败: ' + err.message;
    if (typeof toast === 'function') toast(msg, 'bad');
    else alert(msg);
  }
}

async function prefetchScreenshotsAsDataUri(r){
  const shots = ((r.report && r.report.meta && r.report.meta.screenshots) || []).filter(s => !s.error && s.filename);
  if (!shots.length) return {};
  const map = {};
  await Promise.all(shots.map(async s => {
    try {
      const resp = await fetch('/api/screenshots/' + encodeURIComponent(s.filename));
      if (!resp.ok) return;
      const blob = await resp.blob();
      const reader = new FileReader();
      const dataUri = await new Promise((res, rej) => {
        reader.onload = () => res(reader.result);
        reader.onerror = () => rej(reader.error);
        reader.readAsDataURL(blob);
      });
      map[s.filename] = dataUri;
    } catch(_){}
  }));
  return map;
}

async function ensureSubstepNames(tool){
  if (substepNamesCache[tool.id]) { tool._substepNames = substepNamesCache[tool.id]; return; }
  const m = {};
  for (const sid of tool.prompts){
    try {
      const d = await fetch(`/api/prompts/${tool.prompt_dir}/${sid}`).then(r=>r.json());
      m[sid] = d.name;
    } catch(_){ m[sid] = sid; }
  }
  substepNamesCache[tool.id] = m;
  tool._substepNames = m;
}

async function openReport(r){
  const tool = toolMap[r.tool_id];
  if (!tool){ alert('未知工具：' + r.tool_id); return; }
  document.getElementById('modal-icon').textContent = tool.icon;
  document.getElementById('modal-tool').textContent = tool.name;
  document.getElementById('modal-id').textContent = r.run_id;
  document.getElementById('modal-body').innerHTML =
    '<div style="padding:40px;text-align:center;color:var(--fg-3);font-family:var(--mono)">加载中…</div>';
  document.getElementById('modal-overlay').classList.add('open');
  activeModalTab = 'report';
  document.querySelectorAll('.modal-tabs button').forEach(b => b.classList.toggle('active', b.dataset.tab==='report'));

  try {
    const data = await fetch('/api/reports/' + r.run_id).then(r=>r.json());
    const run = (data.tool_id) ? data : {
      tool_id: r.tool_id, run_id: r.run_id,
      report: data.report || data,
      usage: data.usage || {}, started_at: r.started_at || 0, finished_at: r.finished_at || 0,
    };
    await ensureSubstepNames(tool);
    currentRun = run;
    currentTool = tool;
    // step2 测试用例工具:产出就是用例,只给 Excel,藏掉 HTML/MD/JSON 那套"报告"下载。
    const caseCount = ((run.report || {}).cases || []).length;
    const isStep2 = tool.id === 'step2';
    const setBtn = (id, show) => {
      const el = document.getElementById(id);
      if (el) el.style.display = show ? '' : 'none';
    };
    setBtn('modal-dl-xlsx', true);
    setBtn('modal-dl-md',   true);
    setBtn('modal-dl-json', true);
    // step2:tab 标签「报告」→「用例」
    const tabReport = document.getElementById('modal-tab-report');
    if (tabReport) tabReport.textContent = isStep2 ? '用例' : '报告';
    renderModalBody();
  } catch(e){
    document.getElementById('modal-body').innerHTML =
      `<div style="padding:40px;text-align:center;color:var(--bad)">加载失败：${e}</div>`;
  }
}

function renderModalBody(){
  if (!currentRun) return;
  const body = document.getElementById('modal-body');
  if (activeModalTab === 'json'){
    body.innerHTML = `<div class="pane-json"><pre>${escapeHtml(JSON.stringify(currentRun.report, null, 2))}</pre></div>`;
    return;
  }
  // 失败任务短路 — 不要把 failed 包装成正常的"5 段都是 0"报告:
  // 这是 Codex AI-FE-002 报告的:用户点 FAILED 行,弹窗显示"未产出/0 问题",
  // 真实失败原因被掩盖。
  if (currentRun.status === 'failed'){
    body.innerHTML = renderFailureCard(currentRun, currentTool);
    return;
  }
  // expandAll=false:子步骤默认折叠,executive summary 5 段直接展示,避免重复
  body.innerHTML = renderHtmlReport(currentRun, currentTool, {expandAll: false});
  body.querySelectorAll('.report-sub-head').forEach(h => {
    h.onclick = () => h.parentElement.classList.toggle('open');
  });
  inlineScreenshotsInBody(body);
}

function renderFailureCard(r, tool){
  const errRaw = r.traceback || r.error || '(未捕获到错误详情)';
  const lastProgress = r.progress || '';
  const logs = Array.isArray(r.logs) ? r.logs.slice(-6) : [];
  const elapsed = (r.finished_at && r.started_at) ? (r.finished_at - r.started_at).toFixed(1) + 's' : '—';
  const errFirstLine = String(errRaw).split('\n')[0].slice(0, 240);
  const logsHtml = logs.length
    ? '<div class="fail-logs"><div class="fail-logs-head">最近日志</div><pre>' +
      logs.map(l => {
        const t = l.ts ? new Date(l.ts * 1000).toLocaleTimeString('zh-CN', {hour12:false}) : '';
        return escapeHtml(`[${t}] ${l.event || ''} ${JSON.stringify(l.fields || {})}`);
      }).join('\n') + '</pre></div>'
    : '';
  return `
    <div class="fail-card">
      <div class="fail-banner">
        <span class="fail-tag">FAILED</span>
        <span class="fail-title">${escapeHtml(tool && tool.name || r.tool_id || '?')} · 未生成有效报告</span>
      </div>
      <div class="fail-meta">
        <span>run <code>${escapeHtml((r.run_id || '').slice(0,12))}</code></span>
        <span>用时 ${elapsed}</span>
        ${lastProgress ? `<span>最后进度 · ${escapeHtml(lastProgress)}</span>` : ''}
      </div>
      <div class="fail-section">
        <div class="fail-label">错误摘要</div>
        <pre class="fail-err">${escapeHtml(errFirstLine)}</pre>
      </div>
      <details class="fail-trace">
        <summary>完整 traceback / stderr (点击展开)</summary>
        <pre>${escapeHtml(String(errRaw))}</pre>
      </details>
      ${logsHtml}
      <div class="fail-hint">
        说明：本次运行未产出结构化报告。请确认 ① 所选模型是否可用（如 Haiku 4.5 需当前账户已开通），
        ② 输入是否触发了上游限制，再按相同输入重试。
      </div>
    </div>`;
}

// 详情页 modal 内的截图 inliner — 与 list 页同逻辑(两份 script 块各自定义)
const _screenshotInlineCache2 = {};
async function inlineScreenshotsInBody(rootEl){
  if (!rootEl) return;
  const imgs = rootEl.querySelectorAll('img[data-screenshot-filename]');
  if (!imgs.length) return;
  await Promise.all(Array.from(imgs).map(async img => {
    if (img.src && img.src.startsWith('data:')) return;
    const fn = img.dataset.screenshotFilename;
    if (!fn) return;
    if (_screenshotInlineCache2[fn]) { img.src = _screenshotInlineCache2[fn]; return; }
    try {
      const resp = await fetch('/api/screenshots/' + encodeURIComponent(fn));
      if (!resp.ok) return;
      const blob = await resp.blob();
      const reader = new FileReader();
      const dataUri = await new Promise((res, rej) => {
        reader.onload = () => res(reader.result);
        reader.onerror = () => rej(reader.error);
        reader.readAsDataURL(blob);
      });
      _screenshotInlineCache2[fn] = dataUri;
      img.src = dataUri;
    } catch(_){}
  }));
}

function closeModal(){
  document.getElementById('modal-overlay').classList.remove('open');
  currentRun = null; currentTool = null;
}
window.closeModal = closeModal;

document.getElementById('refresh').onclick = load;
document.getElementById('search').oninput = render;
document.getElementById('filters').onclick = e => {
  if (e.target.tagName !== 'BUTTON') return;
  document.querySelectorAll('#filters button').forEach(b => b.classList.remove('active'));
  e.target.classList.add('active');
  currentFilter = e.target.dataset.f;
  render();
};

// ===================== 批量删除工具栏 =====================
function updateBulkBar(){
  const countEl = document.getElementById('bulk-count');
  const delBtn = document.getElementById('bulk-del-sel');
  const allCb = document.getElementById('bulk-all-cb');
  const allLabel = document.getElementById('bulk-all-label');
  if (!countEl || !delBtn || !allCb) return;
  // 当前页选中数
  const visibleSel = currentVisibleRunIds.filter(rid => selectedRunIds.has(rid)).length;
  const n = selectedRunIds.size;
  countEl.textContent = '已选 ' + n + ' 条';
  countEl.classList.toggle('has-sel', n > 0);
  delBtn.disabled = n === 0;
  delBtn.textContent = '✕ 删除选中' + (n > 0 ? '(' + n + ')' : '');
  // 「全选当前页」状态
  const allVisChecked = currentVisibleRunIds.length > 0 && visibleSel === currentVisibleRunIds.length;
  const someVisChecked = visibleSel > 0 && visibleSel < currentVisibleRunIds.length;
  allCb.checked = allVisChecked;
  allCb.indeterminate = someVisChecked;
  allLabel.textContent = '全选当前页 (' + currentVisibleRunIds.length + ' 条)';
}

document.getElementById('bulk-all-cb').onchange = (e) => {
  const checked = e.target.checked;
  currentVisibleRunIds.forEach(rid => {
    if (checked) selectedRunIds.add(rid); else selectedRunIds.delete(rid);
  });
  // 直接 re-render 更新所有行的高亮
  render();
};

document.getElementById('bulk-del-sel').onclick = async () => {
  const ids = Array.from(selectedRunIds);
  if (!ids.length) return;
  if (!confirm('确认删除选中的 ' + ids.length + ' 条报告?\n\n该操作不可恢复:内存运行 + 磁盘 JSON + 截图全部删除。')) return;
  const btn = document.getElementById('bulk-del-sel');
  btn.disabled = true; btn.textContent = '删除中…';
  try {
    const r = await fetch('/api/reports/batch_delete', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({run_ids: ids}),
    });
    const d = await r.json();
    if (!r.ok){ alert('批量删除失败:' + (d.detail || r.status)); return; }
    const failedTxt = d.failed && d.failed.length ? `\n失败 ${d.failed.length} 条:` + d.failed.slice(0,5).map(f=>`\n  ${f.run_id.slice(0,12)} — ${f.reason}`).join('') : '';
    alert(`已删除 ${d.deleted} / ${d.total} 条${failedTxt}`);
    selectedRunIds.clear();
    await load();
  } catch(e){
    alert('请求出错:' + e.message);
  } finally {
    btn.disabled = false;
  }
};

document.getElementById('bulk-clear-all').onclick = async () => {
  // 二次确认 — 强制让用户输入 "DELETE"
  const total = allReports.length;
  if (total === 0){ alert('没有报告可删'); return; }
  const c1 = confirm('⚠ 警告:即将清空你能看到的【全部 ' + total + ' 条】报告。\n\n继续吗?');
  if (!c1) return;
  const c2 = prompt('请输入「DELETE」(大写)确认清空所有报告:');
  if (c2 !== 'DELETE'){ alert('已取消'); return; }
  const btn = document.getElementById('bulk-clear-all');
  btn.disabled = true; btn.textContent = '清空中…';
  try {
    const r = await fetch('/api/reports/batch_delete', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({all_visible: true}),
    });
    const d = await r.json();
    if (!r.ok){ alert('清空失败:' + (d.detail || r.status)); return; }
    alert(`已清空 ${d.deleted} / ${d.total} 条报告${d.failed && d.failed.length ? '(失败 '+d.failed.length+' 条)' : ''}`);
    selectedRunIds.clear();
    await load();
  } catch(e){
    alert('请求出错:' + e.message);
  } finally {
    btn.disabled = false;
    btn.textContent = '⚠ 清空所有';
  }
};
document.querySelectorAll('.modal-tabs button').forEach(b => {
  b.onclick = () => {
    activeModalTab = b.dataset.tab;
    document.querySelectorAll('.modal-tabs button').forEach(x => x.classList.toggle('active', x === b));
    renderModalBody();
  };
});
document.querySelectorAll('.modal-head .download').forEach(b => {
  b.onclick = async () => {
    if (!currentRun || !currentTool) return;
    const r = currentRun;
    const fmt = b.dataset.act;
    if (!['html', 'md', 'json', 'xlsx'].includes(fmt)) return;
    // 委托给 quickDownload 统一处理 preflight + 错误提示
    if (typeof quickDownload === 'function'){
      await quickDownload(r.run_id, r.tool_id, fmt);
      return;
    }
    // 兜底:没有 quickDownload(其他页面)时走老路径。
    // 不设 a.download —— 让服务端 Content-Disposition 决定文件名(扩展名随真实内容,避免 xlsx 存成 .html)。
    const url = `/api/reports/${r.run_id}/export.${fmt}`;
    const a = document.createElement('a');
    a.href = url;
    document.body.appendChild(a); a.click(); a.remove();
  };
});
document.addEventListener('keydown', e => { if (e.key === 'Escape') closeModal(); });
document.getElementById('modal-overlay').onclick = e => {
  if (e.target.id === 'modal-overlay') closeModal();
};

// === Shared renderer functions (mirror tool detail page, parametric on `tool`) ===
function escapeHtml(s){ return String(s).replace(/[&<>]/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;'}[c])); }
function gateClass(action){
  const a = String(action || '').toLowerCase();
  if (a.includes('reject')) return 'reject';
  if (a.includes('proceed') || a.includes('approve') || a === 'pass') return 'proceed';
  return 'warn';
}
function extractTitle(data){
  if (!data || typeof data !== 'object') return '';
  for (const k of ['name','title']) if (typeof data[k] === 'string') return data[k];
  return '';
}
function quickStats(data){
  if (!data || typeof data !== 'object') return '';
  const chips = [];
  if (Array.isArray(data.cases)) chips.push(`<span class="chip">${data.cases.length} 用例</span>`);
  if (Array.isArray(data.issues)) chips.push(`<span class="chip">${data.issues.length} 问题</span>`);
  if (Array.isArray(data.scenarios)) chips.push(`<span class="chip">${data.scenarios.length} 场景</span>`);
  if (Array.isArray(data.endpoints)) chips.push(`<span class="chip">${data.endpoints.length} 接口</span>`);
  if (Array.isArray(data.pages)) chips.push(`<span class="chip">${data.pages.length} 页面</span>`);
  if (Array.isArray(data.matrix)) chips.push(`<span class="chip">${data.matrix.length} 矩阵</span>`);
  if (Array.isArray(data.fix_list)) chips.push(`<span class="chip">${data.fix_list.length} 待修</span>`);
  if (data.confidence && typeof data.confidence.score === 'number'){
    chips.push(`<span class="chip">conf ${(data.confidence.score*100).toFixed(0)}%</span>`);
  }
  return chips.join('');
}
// 详情页的 build/render 与列表页保持一致 — 复用同一实现。
// (两份 script 块只能各自定义；逻辑必须同步;若改一份请同步另一份。)
function buildExecutiveSummary(rep, tool){
  const walkedIssues = [], walkedCases = [];
  function walk(x){
    if (x && typeof x === 'object'){
      if (Array.isArray(x)){ x.forEach(walk); return; }
      const hasCase = ('expected' in x || 'scenario' in x) && ('id' in x);
      const hasIssue = 'severity' in x && (('issue' in x) || ('title' in x) || ('description' in x) || ('name' in x));
      if (hasCase) walkedCases.push(x);
      else if (hasIssue) walkedIssues.push(x);
      Object.values(x).forEach(walk);
    }
  }
  walk(rep.substeps || {});
  const topIssues = Array.isArray(rep.issues) ? rep.issues : walkedIssues;
  const topCases = Array.isArray(rep.cases) ? rep.cases : walkedCases;
  const topRisks = Array.isArray(rep.risks) ? rep.risks : [];
  const topBlockers = Array.isArray(rep.blockers) ? rep.blockers : [];
  const sevRank = {critical:0, high:1, medium:2, low:3, info:4};
  const priRank = {P0:0, P1:1, P2:2, P3:3};
  function getSev(x){ return String(x.severity||'medium').toLowerCase(); }
  function getPri(x){ return String(x.priority||'P2').toUpperCase(); }
  const naturalIssues = topIssues.map(it => ({
    issue_id: String(it.issue_id || it.id || ''),
    title: String(it.title || it.name || it.issue || it.description || '未命名问题').slice(0, 200),
    severity: getSev(it), priority: getPri(it),
    module: String(it.module || it.endpoint || it.file_path || it.location || it.viewport || it.page || it.viewport_filename || '').slice(0, 200),
    current: String(it.current_behavior || it.current || it.observed || it.description || it.issue || '').slice(0, 800),
    expected: String(it.expected_behavior || it.expected || it.requirement || '').slice(0, 800),
    fix: String(it.fix_suggestion || it.fix || it.recommendation || it.suggestion || it.remediation || '').slice(0, 1000),
    repro: Array.isArray(it.reproduce_steps) ? it.reproduce_steps : (it.reproduce_steps ? [String(it.reproduce_steps)] : []),
    accept: String(it.acceptance_criteria || it.acceptance || it.verify || '').slice(0, 600),
    cases: Array.isArray(it.related_test_cases) ? it.related_test_cases : (it.related_test_cases ? [String(it.related_test_cases)] : []),
    owner: String(it.owner_role || it.owner || it.assignee || '').toLowerCase(),
    hours: it.estimated_hours || it.effort || null,
    impact: String(it.impact_scope || it.impact || '').slice(0, 400),
    evidence: String(it.evidence || it.source || '').slice(0, 400),
  })).sort((a,b) => {
    const sa = sevRank[a.severity] ?? 9, sb = sevRank[b.severity] ?? 9;
    if (sa !== sb) return sa - sb;
    const pa = priRank[a.priority] ?? 9, pb = priRank[b.priority] ?? 9;
    return pa - pb;
  }).slice(0, 60);
  let naturalRisks = topRisks.map(r => {
    if (typeof r === 'string') return {title: r, impact:'', why:'', severity:'medium'};
    return {
      id: String(r.id || ''),
      title: String(r.title || r.name || r.risk || '未命名风险').slice(0, 200),
      impact: String(r.impact || r.affects || '').slice(0, 400),
      why: String(r.why || r.reason || r.detail || '').slice(0, 400),
      severity: getSev(r),
    };
  });
  if (!naturalRisks.length){
    const gate = rep.gate_decision || {};
    if (Array.isArray(gate.reasons)){
      naturalRisks = gate.reasons.filter(Boolean).map(s => ({title:String(s), impact:'', why:'', severity:'medium'}));
    }
  }
  const naturalBlockers = topBlockers.map(b => ({
    id: String(b.id || ''),
    title: String(b.title || b.name || '未命名阻碍').slice(0, 200),
    why_blocking: String(b.why_blocking || b.reason || b.why || '').slice(0, 500),
    what_to_unblock: String(b.what_to_unblock || b.action || b.fix || '').slice(0, 500),
    owner_role: String(b.owner_role || b.owner || '').toLowerCase(),
    hours: b.estimated_hours || b.effort || null,
  }));
  const sev = {critical:0, high:0, medium:0, low:0, info:0};
  naturalIssues.forEach(it => { if (sev[it.severity] !== undefined) sev[it.severity]++; });
  const pri = {P0:0, P1:0, P2:0, P3:0, 其他:0};
  topCases.forEach(c => {
    const p = String(c.priority||'').toUpperCase();
    if (pri[p] !== undefined) pri[p]++;
    else if (p) pri.其他++;
  });
  const naturalCases = topCases.map(c => ({
    id: String(c.id || c.case_id || c.tc_id || ''),
    title: String(c.title || c.name || c.scenario || '').slice(0, 200),
    priority: String(c.priority || 'P2').toUpperCase(),
    type: String(c.type || c.kind || '').toLowerCase(),
    status: String(c.status || 'designed').toLowerCase(),
    automation: String(c.automation_tag || c.automation || '').toLowerCase(),
    preconditions: String(c.preconditions || c.precondition || '').slice(0, 300),
    steps: Array.isArray(c.steps) ? c.steps : (c.steps ? [String(c.steps)] : []),
    expected: String(c.expected || c.expected_result || '').slice(0, 400),
  })).sort((a,b) => {
    const pa = priRank[a.priority] ?? 9, pb = priRank[b.priority] ?? 9;
    return pa - pb;
  }).slice(0, 200);
  const gate = rep.gate_decision || {};
  const action = String(gate.action || '').toLowerCase();
  let verdict, vlevel;
  if (rep.verdict){
    const v = String(rep.verdict);
    verdict = v;
    if (v.includes('不通过')) vlevel = 'fail';
    else if (v.includes('有条件') || v.includes('警告') || v.includes('部分')) vlevel = 'warn';
    else vlevel = 'pass';
  } else {
    if (action.includes('reject') || sev.critical > 0){ verdict = '不通过'; vlevel = 'fail'; }
    else if (action.includes('warn') || sev.high > 2 || naturalBlockers.length){ verdict = '有条件通过'; vlevel = 'warn'; }
    else if (!action && !naturalIssues.length && !naturalCases.length){ verdict = '未产出'; vlevel = 'skip'; }
    else { verdict = '通过'; vlevel = 'pass'; }
  }
  const verdictSummary = String(rep.verdict_summary || '').slice(0, 200);
  return {
    verdict, vlevel, verdictSummary,
    risks: naturalRisks, blockers: naturalBlockers,
    issues: naturalIssues, cases: naturalCases, casesCount: naturalCases.length,
    pri, sev, agent: tool.name,
  };
}

function renderExecutiveSummary(s){
  const vmap = {pass:['✅','通过','ok'], warn:['⚠️','有条件通过','warn'], fail:['❌','不通过','bad'], skip:['—','未产出','skip']};
  const [vicon, vtext, vcls] = vmap[s.vlevel] || ['·','—','skip'];
  const sevMap = {critical:'CRITICAL', high:'HIGH', medium:'MEDIUM', low:'LOW', info:'INFO'};
  const priMap = {P0:'P0', P1:'P1', P2:'P2', P3:'P3'};
  const ownerMap = {backend:'后端', frontend:'前端', product:'产品', test:'测试',
                    devops:'运维', security:'安全', data:'数据'};
  const sevTotal = s.sev.critical + s.sev.high + s.sev.medium + s.sev.low + s.sev.info;
  const priTotal = s.pri.P0 + s.pri.P1 + s.pri.P2 + s.pri.P3 + s.pri.其他;
  function sevBar(){
    if (!sevTotal) return '<span class="exec-muted">（暂无问题）</span>';
    const cells = ['critical','high','medium','low','info'].map(k => {
      const n = s.sev[k]; if (!n) return '';
      const pct = (n / sevTotal * 100).toFixed(0);
      return `<span class="sev-bar-seg sev-bar-${k}" style="flex:${n}" title="${sevMap[k]} · ${n} 个 (${pct}%)">${n}</span>`;
    }).filter(Boolean).join('');
    return `<div class="sev-bar">${cells}</div>`;
  }
  function priBar(){
    if (!priTotal) return '';
    const cells = ['P0','P1','P2','P3'].map(k => {
      const n = s.pri[k]; if (!n) return '';
      return `<span class="pri-bar-seg pri-bar-${k}" style="flex:${n}" title="${k} · ${n} 条">${k}: ${n}</span>`;
    }).filter(Boolean).join('');
    return `<div class="pri-bar">${cells}</div>`;
  }
  let html = '';
  html += `<div class="exec-block exec-verdict-block">
    <div class="exec-head"><span class="exec-num">①</span><h3>测试结论</h3></div>
    <div class="verdict ${vcls}"><span class="verdict-icon">${vicon}</span><span>${escapeHtml(vtext)}</span></div>
    ${s.verdictSummary ? `<div class="verdict-summary">${escapeHtml(s.verdictSummary)}</div>` : ''}
    <div class="exec-kpis">
      <div class="exec-kpi"><span class="exec-kpi-num">${sevTotal}</span><span class="exec-kpi-lbl">问题总数</span></div>
      <div class="exec-kpi"><span class="exec-kpi-num">${s.sev.critical + s.sev.high}</span><span class="exec-kpi-lbl">需立即处理</span></div>
      <div class="exec-kpi"><span class="exec-kpi-num">${s.blockers.length}</span><span class="exec-kpi-lbl">阻碍</span></div>
      <div class="exec-kpi"><span class="exec-kpi-num">${s.casesCount}</span><span class="exec-kpi-lbl">用例</span></div>
    </div>
    ${sevBar()}
  </div>`;
  let risksHtml;
  if (s.risks.length){
    risksHtml = '<div class="exec-risk-list">' + s.risks.map(r => {
      const sevBadge = r.severity ? `<span class="sev-tag sev-${r.severity}">${escapeHtml(r.severity)}</span>` : '';
      return `<div class="exec-risk-item">
        <div class="exec-risk-head">${sevBadge}<span class="exec-risk-title">${escapeHtml(r.title)}</span></div>
        ${r.impact ? `<div class="exec-risk-line"><span class="lbl">影响</span>${escapeHtml(r.impact)}</div>` : ''}
        ${r.why ? `<div class="exec-risk-line"><span class="lbl">原因</span>${escapeHtml(r.why)}</div>` : ''}
      </div>`;
    }).join('') + '</div>';
  } else {
    risksHtml = '<p class="exec-muted">（无显著风险）</p>';
  }
  html += `<div class="exec-block">
    <div class="exec-head"><span class="exec-num">②</span><h3>风险结论</h3><span class="exec-count">${s.risks.length}</span></div>
    ${risksHtml}
  </div>`;
  let blockersHtml;
  if (s.blockers.length){
    blockersHtml = '<div class="exec-blocker-list">' + s.blockers.map((b, idx) => `
      <div class="exec-blocker-item">
        <div class="exec-blocker-head">
          <span class="blocker-tag">BLOCKER</span>
          <span class="exec-blocker-title">${idx+1}. ${escapeHtml(b.title)}</span>
          ${b.id ? `<span class="meta-chip">${escapeHtml(b.id)}</span>` : ''}
          ${b.owner_role && ownerMap[b.owner_role] ? `<span class="meta-chip role">👤 ${ownerMap[b.owner_role]}</span>` : ''}
          ${b.hours ? `<span class="meta-chip">⏱ ${b.hours}h</span>` : ''}
        </div>
        ${b.why_blocking ? `<div class="exec-blocker-line"><span class="lbl">为何阻碍</span>${escapeHtml(b.why_blocking)}</div>` : ''}
        ${b.what_to_unblock ? `<div class="exec-blocker-line fix"><span class="lbl">如何解开</span>${escapeHtml(b.what_to_unblock)}</div>` : ''}
      </div>`).join('') + '</div>';
  } else {
    blockersHtml = '<p class="exec-muted">（无阻碍）</p>';
  }
  html += `<div class="exec-block">
    <div class="exec-head"><span class="exec-num">③</span><h3>阻碍</h3><span class="exec-count danger">${s.blockers.length}</span></div>
    ${blockersHtml}
  </div>`;
  function renderIssueCard(it, idx){
    const metaChips = [];
    if (it.issue_id) metaChips.push(`<span class="meta-chip">${escapeHtml(it.issue_id)}</span>`);
    if (it.priority) metaChips.push(`<span class="meta-chip pri-${it.priority}">${priMap[it.priority] || it.priority}</span>`);
    if (it.owner && ownerMap[it.owner]) metaChips.push(`<span class="meta-chip role">👤 ${ownerMap[it.owner]}</span>`);
    if (it.hours) metaChips.push(`<span class="meta-chip">⏱ ${it.hours}h</span>`);
    const reproHtml = it.repro.length
      ? '<ol class="repro-list">' + it.repro.map(s => `<li>${escapeHtml(s)}</li>`).join('') + '</ol>' : '';
    const casesHtml = it.cases.length
      ? `<div class="related-cases">关联用例:${it.cases.map(c => `<code>${escapeHtml(c)}</code>`).join(' ')}</div>` : '';
    return `
      <div class="exec-issue sev-${it.severity}">
        <div class="exec-issue-head">
          <span class="sev-tag sev-${it.severity}">${sevMap[it.severity] || '·'}</span>
          <span class="exec-issue-title">${idx+1}. ${escapeHtml(it.title)}</span>
        </div>
        <div class="exec-issue-meta">${metaChips.join('')}</div>
        ${it.module ? `<div class="exec-issue-loc">位置:<code>${escapeHtml(it.module)}</code></div>` : ''}
        ${it.current ? `<div class="exec-issue-section"><div class="sec-lbl">现状</div><div class="sec-body">${escapeHtml(it.current)}</div></div>` : ''}
        ${it.expected ? `<div class="exec-issue-section"><div class="sec-lbl">期望</div><div class="sec-body">${escapeHtml(it.expected)}</div></div>` : ''}
        ${it.fix ? `<div class="exec-issue-section fix"><div class="sec-lbl">修复建议</div><div class="sec-body">${escapeHtml(it.fix)}</div></div>` : ''}
        ${(reproHtml || it.accept) ? `<div class="exec-issue-section verify"><div class="sec-lbl">验收</div><div class="sec-body">${reproHtml}${it.accept ? `<div class="accept-line">验收标准：${escapeHtml(it.accept)}</div>` : ''}</div></div>` : ''}
        ${casesHtml}
        ${it.impact ? `<div class="exec-issue-impact">影响面:${escapeHtml(it.impact)}</div>` : ''}
        ${it.evidence ? `<div class="exec-issue-evidence">证据:${escapeHtml(it.evidence)}</div>` : ''}
      </div>`;
  }
  const issuesHtml = s.issues.length
    ? s.issues.map((it, i) => renderIssueCard(it, i)).join('')
    : '<p class="exec-muted">（本次未识别到具体问题）</p>';
  html += `<div class="exec-block">
    <div class="exec-head">
      <span class="exec-num">④</span>
      <h3>Bug 表</h3>
      <span class="exec-count">${s.issues.length}</span>
      <span class="exec-count-note">(按严重度 × 优先级排序)</span>
    </div>
    ${issuesHtml}
  </div>`;
  let casesListHtml;
  if (s.cases.length){
    const rows = s.cases.map((c, i) => {
      const statusMap = {
        designed: ['','已设计','muted'],
        executed_pass: ['','已执行通过','ok'],
        executed_fail: ['','执行失败','bad'],
        skipped: ['','已跳过','muted'],
        blocked: ['','阻塞','bad'],
      };
      const [sIcon, sLabel, sCls] = statusMap[c.status] || ['','未定义','muted'];
      return `<tr class="case-row pri-${c.priority}">
        <td class="case-idx">${i+1}</td>
        <td><span class="pri-tag pri-${c.priority}">${priMap[c.priority]||c.priority}</span></td>
        <td><code class="case-id">${escapeHtml(c.id)}</code></td>
        <td class="case-title">${escapeHtml(c.title)}</td>
        <td>${c.type ? `<span class="case-type">${escapeHtml(c.type)}</span>` : ''}</td>
        <td>${c.automation ? `<span class="case-auto">${escapeHtml(c.automation)}</span>` : ''}</td>
        <td><span class="case-status case-status-${sCls}">${sIcon} ${sLabel}</span></td>
      </tr>`;
    }).join('');
    casesListHtml = `<div class="case-table-wrap">
      <table class="case-table">
        <thead><tr><th>#</th><th>优先级</th><th>用例 ID</th><th>用例标题</th><th>类型</th><th>自动化</th><th>状态</th></tr></thead>
        <tbody>${rows}</tbody>
      </table>
    </div>`;
  } else if (s.casesCount){
    casesListHtml = `<div class="exec-case-total"><span class="num">${s.casesCount}</span><span class="lbl">条已生成用例(未提供详细列表)</span></div>`;
  } else {
    casesListHtml = '<p class="exec-muted">（本次未生成用例）</p>';
  }
  html += `<div class="exec-block">
    <div class="exec-head">
      <span class="exec-num">⑤</span>
      <h3>执行用例记录</h3>
      <span class="exec-count">${s.casesCount}</span>
    </div>
    ${priBar()}
    ${casesListHtml}
  </div>`;
  return html;
}

// 测试用例工具:用例表 + Excel 下载做成报告主体(reports 弹窗版)
function buildUiComparisonBlock(rep, runId){
  const meta = (rep && rep.meta) || {};
  const shots = (meta.screenshots || []).filter(s => !s.error);
  const isDesign = (s) => s.is_design || /figma\.com/i.test(s.url||'');
  const designShots = shots.filter(isDesign);
  const actualShots = shots.filter(s => !isDesign(s));
  const issues = (rep.issues || []);
  const scOf = (s) => ({critical:'crit',high:'hi',major:'hi',medium:'med',low:'lo',minor:'lo',info:'lo',cosmetic:'lo'})[(s||'').toLowerCase()]||'med';
  const zhOf = (s) => ({critical:'严重',high:'高',major:'高',medium:'中',low:'低',minor:'低',info:'提示',cosmetic:'细节'})[(s||'').toLowerCase()]||'中';
  const verdict = rep.verdict || '';
  const vsum = rep.verdict_summary || '';
  const vCls = /不通过|reject|fail/i.test(verdict)?'bad':(/有条件|warn|conditional/i.test(verdict)?'warn':'ok');
  const imgTag = (fn) => `<img data-screenshot-filename="${escapeHtml(fn)}" src="/api/screenshots/${encodeURIComponent(fn)}" loading="lazy">`;
  const css = `<style>
  .ui-cmp{margin:0 0 18px}
  .ui-cmp-verdict{padding:14px 18px;border-radius:10px;border:1px solid var(--line,#ddd);margin-bottom:16px;font-size:14px;line-height:1.7}
  .ui-cmp-verdict.ok{background:rgba(22,163,74,.06);border-color:rgba(22,163,74,.3)}
  .ui-cmp-verdict.warn{background:rgba(202,138,4,.06);border-color:rgba(202,138,4,.3)}
  .ui-cmp-verdict.bad{background:rgba(220,38,38,.06);border-color:rgba(220,38,38,.3)}
  .ui-cmp-pair{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-bottom:18px;align-items:start}
  .ui-cmp-side{border:1px solid var(--line,#ddd);border-radius:10px;overflow:hidden;background:var(--surface,#f8fafc)}
  .ui-cmp-label{font-size:12.5px;font-weight:600;padding:8px 12px;border-bottom:1px solid var(--line,#ddd);color:var(--fg-2,#555)}
  .ui-cmp-label.actual{color:#dc2626}
  .ui-cmp-side img{display:block;width:100%;height:auto}
  .ui-cmp-nodes{padding:40px 20px;text-align:center;color:var(--fg-3,#888);font-size:13px}
  .ui-cmp-list-head{font-size:14px;font-weight:700;margin:8px 0 12px;color:var(--fg,#222)}
  .ui-cmp-item{border:1px solid var(--line,#ddd);border-left:4px solid var(--line-2,#ccc);border-radius:0 8px 8px 0;padding:12px 16px;margin-bottom:10px;background:var(--surface,#f8fafc)}
  .ui-cmp-item.crit{border-left-color:#dc2626}.ui-cmp-item.hi{border-left-color:#ea580c}.ui-cmp-item.med{border-left-color:#ca8a04}.ui-cmp-item.lo{border-left-color:#16a34a}
  .ui-cmp-ih{display:flex;align-items:center;gap:10px;margin-bottom:8px}
  .ui-cmp-num{width:22px;height:22px;border-radius:50%;background:var(--fg,#222);color:#fff;display:grid;place-items:center;font-size:12px;font-weight:700;flex-shrink:0}
  .ui-cmp-t{font-weight:700;font-size:14.5px;flex:1;color:var(--fg,#222)}
  .ui-cmp-sev{font-size:11px;font-weight:700;padding:2px 8px;border-radius:4px}
  .ui-cmp-sev.crit{background:rgba(220,38,38,.12);color:#dc2626}.ui-cmp-sev.hi{background:rgba(234,88,12,.12);color:#ea580c}.ui-cmp-sev.med{background:rgba(202,138,4,.12);color:#a16207}.ui-cmp-sev.lo{background:rgba(22,163,74,.12);color:#16a34a}
  .ui-cmp-r{display:flex;gap:8px;font-size:13.5px;line-height:1.7;margin-top:3px;color:var(--fg,#222)}
  .ui-cmp-r .l{flex-shrink:0;width:56px;font-size:10.5px;font-weight:700;color:var(--fg-3,#888);background:var(--surface-2,#eef);padding:2px 6px;border-radius:3px;height:fit-content;text-align:center}
  .ui-cmp-r.design .l{color:#0891b2}.ui-cmp-r.actual .l{color:#dc2626}
  @media(max-width:720px){.ui-cmp-pair{grid-template-columns:1fr}}
  </style>`;
  let h = css + '<div class="ui-cmp">';
  if (verdict || vsum){
    h += `<div class="ui-cmp-verdict ${vCls}"><b>设计符合度：${escapeHtml(verdict||'—')}</b>${vsum?' — '+escapeHtml(vsum):''}</div>`;
  }
  // 按真实配对一对一展示:每个实拍配它自己对应的设计帧(不再全部贴第一张 designShots[0])
  const pcMap = {}; (rep.pairs_checked||[]).forEach(p => { pcMap[p.actual] = p.design; });
  const designByName = {};
  designShots.forEach(d => { const nm = String(d.node_name||d.viewport||''); if(nm && !(nm in designByName)) designByName[nm]=d; });
  if (actualShots.length){
    actualShots.forEach(act => {
      const actFn = act.annotated_filename || act.filename;
      const an = act.viewport || act.filename;
      const dn = pcMap[an];
      const paired = dn && dn !== '(无对应设计帧)';
      const des = paired ? designByName[String(dn)] : null;
      h += `<div class="ui-cmp-pair">
        <div class="ui-cmp-side"><div class="ui-cmp-label">🎨 设计稿${paired?'『'+escapeHtml(dn)+'』':''}</div>${des?imgTag(des.filename):'<div class="ui-cmp-nodes">无对应设计帧</div>'}</div>
        <div class="ui-cmp-side"><div class="ui-cmp-label actual">📱 实际产品 · ${escapeHtml(an)} · 红框=与设计不一致</div>${imgTag(actFn)}</div>
      </div>`;
    });
  } else if (designShots.length){
    h += `<div class="ui-cmp-pair"><div class="ui-cmp-side"><div class="ui-cmp-label">🎨 设计稿（Figma）</div>${imgTag(designShots[0].filename)}</div><div class="ui-cmp-side"><div class="ui-cmp-nodes">未捕获到实际产品截图</div></div></div>`;
  }
  h += `<div class="ui-cmp-list-head">不一致清单（${issues.length} 处）</div>`;
  if (!issues.length){
    h += '<div class="ui-cmp-nodes" style="border:1px dashed var(--line,#ddd);border-radius:8px">未发现与设计稿的明显不一致 🎉</div>';
  } else {
    issues.forEach((it, i) => {
      const sc = scOf(it.severity);
      h += `<div class="ui-cmp-item ${sc}">
        <div class="ui-cmp-ih"><span class="ui-cmp-num">${i+1}</span><span class="ui-cmp-t">${escapeHtml(it.title||it.issue_id||'差异')}</span><span class="ui-cmp-sev ${sc}">${zhOf(it.severity)}</span></div>
        <div class="ui-cmp-r design"><span class="l">设计要求</span><span>${escapeHtml(it.expected_behavior||it.expected||'—')}</span></div>
        <div class="ui-cmp-r actual"><span class="l">实际</span><span>${escapeHtml(it.current_behavior||it.current||'—')}</span></div>
        ${it.fix_suggestion?`<div class="ui-cmp-r"><span class="l">建议</span><span>${escapeHtml(it.fix_suggestion)}</span></div>`:''}
      </div>`;
    });
  }
  h += '</div>';
  return h;
}

function buildTestCaseBlock(rep, runId, toolId){
  const cases = (rep && rep.cases) || [];
  if (!cases.length) return '';
  const rows = cases.map(c => {
    const stepsArr = Array.isArray(c.steps) ? c.steps : (c.steps ? [c.steps] : []);
    const steps = stepsArr.filter(Boolean).map(s => escapeHtml(String(s))).join('\n');
    let exp = c.expected;
    if (Array.isArray(exp)) exp = exp.map(e => typeof e==='string'?e:JSON.stringify(e)).join('\n');
    exp = escapeHtml(String(exp || ''));
    const pri = String(c.priority||'').toUpperCase();
    return '<tr>' +
      '<td class="tc-id">' + escapeHtml(c.id||'') + '</td>' +
      '<td>' + escapeHtml(c.module||'') + '</td>' +
      '<td class="tc-title">' + escapeHtml(c.title||c.name||'') + '</td>' +
      '<td><span class="tc-pri tc-pri-' + pri + '">' + escapeHtml(pri||'-') + '</span></td>' +
      '<td>' + escapeHtml(c.type||'') + '</td>' +
      '<td class="tc-pre">' + escapeHtml(c.preconditions||'') + '</td>' +
      '<td class="tc-steps">' + steps + '</td>' +
      '<td class="tc-exp">' + exp + '</td>' +
    '</tr>';
  }).join('');
  return '<div class="tc-block">' +
    '<div class="tc-banner">' +
      '<div><div class="tc-count">' + cases.length + '</div><div class="tc-count-label">条测试用例</div></div>' +
      '<div class="tc-mid">本工具产出 <b>纯人工执行</b> 测试用例,步骤是自然语言操作清单。' +
        '点右侧按钮导出标准 Excel(含「执行结果 / 实际结果」空列),可直接发给测试团队执行。</div>' +
      '<button class="tc-excel-btn" data-tc-runid="' + runId + '" data-tc-toolid="' + (toolId||'step2') + '">↓ 下载 Excel 用例表</button>' +
    '</div>' +
    '<div class="tc-table-wrap"><table class="tc-table"><thead><tr>' +
      '<th>用例编号</th><th>模块</th><th>用例标题</th><th>优先级</th><th>类型</th>' +
      '<th>前置条件</th><th>测试步骤</th><th>预期结果</th>' +
    '</tr></thead><tbody>' + rows + '</tbody></table></div>' +
  '</div>';
}

function renderHtmlReport(r, tool, opts){
  opts = opts || {};
  const expandAll = !!opts.expandAll;
  const rep = r.report || {};
  const meta = rep.meta || {};
  const u = r.usage || {};
  let html = '';
  const elapsed = (r.finished_at && r.started_at) ? (r.finished_at - r.started_at).toFixed(1) : '—';

  // === 执行摘要 4 块结构 (测试结论 / 风险结论 / 问题描述 / 用例执行) ===
  const summary = buildExecutiveSummary(rep, tool);

  const pc2 = meta.project_code || r.project_code || '';
  const pn2 = meta.project_name || r.project_name || '';
  const projectRow2 = (pc2 || pn2)
    ? `<div class="report-project">
        <span class="lbl">项目编号</span><code>${escapeHtml(pc2 || '—')}</code>
        <span class="lbl">项目名称</span><span class="val">${escapeHtml(pn2 || '—')}</span>
      </div>` : '';
  html += `<div class="report-hero">
    <div class="report-icon">${tool.icon}</div>
    <div style="flex:1;min-width:0">
      <h4>${escapeHtml(tool.name)} · 报告</h4>
      <div class="meta">
        ${meta.produced_at_utc ? new Date(meta.produced_at_utc).toLocaleString('zh-CN', {hour12:false}) : ''}
        · 模型 <code>${escapeHtml(meta.model_id || '?')}</code>
        · run <code>${r.run_id.slice(0,12)}</code>
      </div>
      ${projectRow2}
      <div class="stat-pills">
        <span class="pill">用时 <span class="v">${elapsed}s</span></span>
        <span class="pill">成本 <span class="v">$${u.cost_usd ?? '—'}</span></span>
        <span class="pill">输出 <span class="v">${(u.output_tokens||0).toLocaleString()}</span></span>
        <span class="pill">缓存 <span class="v">${(u.cache_read_tokens||0).toLocaleString()}</span></span>
      </div>
    </div>
  </div>`;

  // 测试用例工具(step2):只出用例表 + Excel,不要 verdict / 风险 / Bug 那套"报告"。
  if (tool && tool.id === 'step2'){
    const tcBlock = buildTestCaseBlock(rep, r.run_id, tool.id);
    html += tcBlock || '<div class="run-empty" style="padding:48px;text-align:center">本次未生成测试用例</div>';
    return html;
  }

  // UI 一致性比对(step5):并排「设计稿 vs 实际产品」对比图 + 一条条不一致清单
  if (tool && tool.id === 'step5'){
    html += buildUiComparisonBlock(rep, r.run_id);
    return html;
  }

  // 测试结论 + 风险结论 + 问题描述 + 用例执行
  html += renderExecutiveSummary(summary);
  // Inline page screenshots (step5 / h5_adapt only)
  // 用 data-screenshot-filename 占位,渲染后 inlineScreenshotsInArea() 把 src
  // 替换成 data: URI,报告自包含、不暴露任何本地文件夹路径。
  const shots = (meta.screenshots || []).filter(s => !s.error);
  if (shots.length){
    const groupedByUrl = {};
    shots.forEach(s => {(groupedByUrl[s.url] = groupedByUrl[s.url] || []).push(s);});
    const imgMap = (opts && opts.imgMap) || {};
    const _shotGroupLabel = (u) => {
      if(!u) return '截图';
      if(u.indexOf('app://')===0) return '📱 APP 实拍(模拟器)· ' + u.replace('app://','');
      if(/figma\.com/i.test(u)) return '🎨 设计基线(Figma)';
      return '🌐 ' + u;
    };
    let shotsHtml = '<div class="report-screenshots">';
    shotsHtml += '<div class="screenshots-head">截图证据 · 设计基线 vs 实拍对照 <span class="screenshots-hint">（已嵌入报告，无本地文件依赖）</span></div>';
    Object.entries(groupedByUrl).forEach(([url, arr]) => {
      shotsHtml += `<div class="shot-group"><div class="shot-url">${escapeHtml(_shotGroupLabel(url))}</div>`;
      shotsHtml += '<div class="shot-grid">';
      arr.forEach(s => {
        const annotated = s.annotated_filename;
        const fnPrimary = annotated || s.filename;
        // 已 inline 直接用;否则放占位 src,后续 inliner 异步替换为 data: URI
        const initialSrc = imgMap[fnPrimary] || `/api/screenshots/${encodeURIComponent(fnPrimary)}`;
        const issueBadge = s.issue_count
          ? `<span class="issue-badge">${s.issue_count} 个问题</span>` : '';
        const dim = (s.width && s.height) ? ` · ${s.width}×${s.height}` : '';
        shotsHtml += `<div class="shot-cell" title="${escapeHtml(s.viewport)}${dim}${annotated?' · 已标注':''}">
          <img src="${initialSrc}" data-screenshot-filename="${escapeHtml(fnPrimary)}" alt="${escapeHtml(s.viewport)}" loading="lazy">
          <div class="shot-cap">${escapeHtml(s.viewport)}${dim}${issueBadge}</div>
        </div>`;
      });
      shotsHtml += '</div></div>';
    });
    shotsHtml += '</div>';
    html += shotsHtml;
  }
  const hasContractData = (Array.isArray(rep.issues) && rep.issues.length) ||
                          (Array.isArray(rep.cases) && rep.cases.length);
  const subs = rep.substeps || {};
  let firstWithContent = null;
  for (const sid of tool.prompts){
    if (subs[sid] && Object.keys(subs[sid]).length){ firstWithContent = sid; break; }
  }
  html += `<div class="substep-section-head">
    <span class="substep-section-title">各子步骤原始输出</span>
    <span class="substep-section-hint">${hasContractData ? '已聚合到上方 5 段;点击展开看原始数据' : '点击展开'}</span>
  </div>`;
  let idx = 0;
  for (const sid of tool.prompts){
    idx++;
    const data = subs[sid];
    const title = (tool._substepNames && tool._substepNames[sid]) || extractTitle(data) || `子分析 ${idx}`;
    const open = expandAll || (!hasContractData && sid === firstWithContent);
    if (data == null){
      html += `<div class="report-sub">
        <div class="report-sub-head">
          <span class="report-sub-twirl">▶</span>
          <span class="report-sub-num">${idx}</span>
          <span class="report-sub-name" style="color:var(--fg-3)">${escapeHtml(title)} · 已跳过</span>
        </div>
      </div>`;
      continue;
    }
    html += `<div class="report-sub${open?' open':''}">
      <div class="report-sub-head">
        <span class="report-sub-twirl">▶</span>
        <span class="report-sub-num">${idx}</span>
        <span class="report-sub-name">${escapeHtml(title)}</span>
        <span class="report-sub-stats">${quickStats(data) || ''}</span>
      </div>
      <div class="report-sub-body">${renderSmart(data)}</div>
    </div>`;
  }
  return html;
}
function renderSmart(data, depth){
  depth = depth || 0;
  if (data === null || data === undefined) return '<span class="empty-array">—</span>';
  if (typeof data === 'string'){ return data ? escapeHtml(data) : '<span class="empty-array">""</span>'; }
  if (typeof data === 'number' || typeof data === 'boolean') return `<span style="color:var(--ac)">${escapeHtml(String(data))}</span>`;
  if (Array.isArray(data)){
    if (!data.length) return '<span class="empty-array">[]</span>';
    const allObj = data.every(x => x && typeof x === 'object' && !Array.isArray(x));
    if (allObj && data.length >= 2) return renderObjectArrayAsTable(data);
    if (allObj) return renderSmart(data[0], depth+1);
    return '<ul class="report-list">' + data.slice(0,50).map(x => `<li>${renderSmart(x, depth+1)}</li>`).join('') +
      (data.length > 50 ? `<li class="empty-array">…剩余 ${data.length - 50} 已折叠</li>` : '') + '</ul>';
  }
  if (typeof data === 'object') return renderObjectAsKv(data, depth);
  return escapeHtml(String(data));
}
function renderObjectAsKv(obj, depth){
  const entries = Object.entries(obj);
  if (!entries.length) return '<span class="empty-array">{}</span>';
  return '<dl class="report-kv">' + entries.map(([k,v]) =>
    `<dt>${escapeHtml(k)}</dt><dd>${renderFieldValue(k, v, depth)}</dd>`
  ).join('') + '</dl>';
}
function renderFieldValue(key, v, depth){
  if ((key === 'severity' || key.endsWith('_severity')) && typeof v === 'string'){
    return `<span class="sev sev-${escapeHtml(v.toLowerCase())}">${escapeHtml(v)}</span>`;
  }
  if (key === 'confidence' && v && typeof v === 'object' && typeof v.score === 'number'){
    const pct = (v.score*100).toFixed(0);
    return `<span class="confbar"><span class="track"><span class="fill" style="width:${pct}%"></span></span><span class="pct">${pct}%</span></span>${
      v.rationale ? ` <span style="color:var(--fg-3)">— ${escapeHtml(String(v.rationale).slice(0,140))}</span>` : ''
    }`;
  }
  if (key === 'gate_decision' && v && typeof v === 'object'){
    const cls = gateClass(v.action);
    const rs = (v.reasons || []).map(x => `<div>· ${escapeHtml(x)}</div>`).join('');
    return `<div class="gate-banner ${cls}" style="margin:0;border-bottom:none">
      <span class="badge">${escapeHtml(String(v.action).toLowerCase())}</span>
      <div class="reasons">${rs}</div></div>`;
  }
  if (Array.isArray(v) && v.length > 5){
    return `<details class="report-detail" ${depth === 0 ? 'open' : ''}><summary>${v.length} 项</summary><div>${renderSmart(v, depth+1)}</div></details>`;
  }
  if (typeof v === 'object' && v !== null && !Array.isArray(v) && Object.keys(v).length > 6){
    return `<details class="report-detail"><summary>对象（${Object.keys(v).length} 字段）</summary><div>${renderSmart(v, depth+1)}</div></details>`;
  }
  return renderSmart(v, depth+1);
}
function renderObjectArrayAsTable(arr){
  const keySet = new Set();
  arr.forEach(o => Object.keys(o).forEach(k => keySet.add(k)));
  const priority = ['id','title','name','severity','status','category','kind','endpoint','page','area','module','expected','actual','fix','impact','effort_hours','severity_if_fails'];
  const keys = [...keySet].sort((a,b)=>{
    const ai = priority.indexOf(a), bi = priority.indexOf(b);
    if (ai >= 0 && bi >= 0) return ai - bi;
    if (ai >= 0) return -1; if (bi >= 0) return 1; return 0;
  }).slice(0,6);
  const head = `<thead><tr>${keys.map(k => `<th>${escapeHtml(k)}</th>`).join('')}</tr></thead>`;
  const body = arr.slice(0,30).map(o => `<tr>${keys.map(k => `<td>${renderTableCell(k, o[k])}</td>`).join('')}</tr>`).join('');
  const foot = arr.length > 30
    ? `<tfoot><tr><td colspan="${keys.length}" style="text-align:center;color:var(--fg-3)">… 共 ${arr.length} 条，已显示前 30 条</td></tr></tfoot>` : '';
  return `<table class="report-table">${head}<tbody>${body}</tbody>${foot}</table>`;
}
function renderTableCell(key, v){
  if (v == null) return '<span class="empty-array">—</span>';
  if ((key === 'severity' || key.endsWith('_severity')) && typeof v === 'string'){
    return `<span class="sev sev-${escapeHtml(v.toLowerCase())}">${escapeHtml(v)}</span>`;
  }
  if (typeof v === 'string'){ return escapeHtml(v.length > 100 ? v.slice(0,97)+'…' : v); }
  if (typeof v === 'number' || typeof v === 'boolean') return `<span style="color:var(--ac)">${escapeHtml(String(v))}</span>`;
  if (Array.isArray(v)){
    if (!v.length) return '<span class="empty-array">[]</span>';
    if (v.length <= 3 && v.every(x => typeof x === 'string' || typeof x === 'number')) return v.map(x => escapeHtml(String(x))).join(', ');
    return `<details class="report-detail"><summary>${v.length} 项</summary><div>${renderSmart(v)}</div></details>`;
  }
  if (typeof v === 'object') return `<details class="report-detail"><summary>{${Object.keys(v).length}}</summary><div>${renderSmart(v)}</div></details>`;
  return escapeHtml(String(v));
}

// === Exporters (mirror tool detail page) ===
function downloadBlob(content, filename, mime){
  const blob = new Blob([content], {type: mime});
  const url = URL.createObjectURL(blob);
  const a = document.createElement('a');
  a.href = url; a.download = filename;
  document.body.appendChild(a); a.click(); document.body.removeChild(a);
  setTimeout(() => URL.revokeObjectURL(url), 1000);
}
function buildMarkdownReport(r, tool){
  const rep = r.report || {};
  const meta = rep.meta || {};
  const u = r.usage || {};
  const lines = [];
  lines.push(`# ${tool.name} · 分析报告`); lines.push('');
  lines.push(`- **生成时间**：${new Date().toLocaleString('zh-CN', {hour12:false})}`);
  lines.push(`- **Run ID**：\`${r.run_id}\``);
  lines.push(`- **模型**：\`${meta.model_id || '?'}\``);
  if (r.finished_at && r.started_at) lines.push(`- **耗时**：${(r.finished_at - r.started_at).toFixed(1)}s`);
  if (u.cost_usd != null) lines.push(`- **成本**：$${u.cost_usd}`);
  lines.push('');
  if (rep.gate_decision){
    const g = rep.gate_decision;
    lines.push(`## ⚠️ 闸门决策：\`${g.action}\``);
    (g.reasons || []).forEach(rs => lines.push(`- ${rs}`));
    lines.push('');
  }
  let idx = 0;
  for (const sid of tool.prompts){
    idx++;
    const data = (rep.substeps || {})[sid];
    const title = (tool._substepNames && tool._substepNames[sid]) || extractTitle(data) || `子分析 ${idx}`;
    lines.push(`## ${idx}. ${title}`);
    if (data == null){ lines.push('_已跳过_'); lines.push(''); continue; }
    lines.push(''); lines.push('```json'); lines.push(JSON.stringify(data, null, 2)); lines.push('```'); lines.push('');
  }
  return lines.join('\n');
}

load();
</script>
</body></html>
"""


@app.get("/reports", response_class=HTMLResponse)
async def reports_page() -> str:
    return _inject_shared_overlays(REPORTS_HTML)
